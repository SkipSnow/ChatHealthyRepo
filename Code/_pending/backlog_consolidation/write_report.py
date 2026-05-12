"""
Write the Word report for the EPIC-008-F-002 backlog consolidation pass.
Output: architecture/EngineeringRuleEnforcement/ArchitectureDesignAndAuditDocs/EPIC-008-F-002-backlog-consolidation-V1.docx
"""
import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parents[3]
OUT = ROOT / "architecture" / "EngineeringRuleEnforcement" / "ArchitectureDesignAndAuditDocs" / "EPIC-008-F-002-backlog-consolidation-V1.docx"
MAPPING = ROOT / "Code" / "_pending" / "backlog_consolidation" / "id_mapping.json"
REPOINT_LOG = ROOT / "Code" / "_pending" / "backlog_consolidation" / "repoint_log.json"
BACKLOG = ROOT / "brain" / "machine_artifacts" / "content" / "agile_backlog.json"

with open(MAPPING, encoding="utf-8") as f:
    mapping = json.load(f)
with open(REPOINT_LOG, encoding="utf-8") as f:
    repoint_log = json.load(f)
with open(BACKLOG, encoding="utf-8") as f:
    backlog = json.load(f)

# ---- Compute counts ----
# Before: count flagged + count REQs in F-002 before
with open(ROOT / "_oneshots/test_output" / "scan_validate_relevance.json", encoding="utf-8") as f:
    flagged = json.load(f)
flagged_count = len(flagged)
flagged_reqs = sum(1 for e in flagged if e["entity_type"] == "requirement")

retired_count = len(mapping["retired_to_new"])
retired_reqs = sum(1 for k in mapping["retired_to_new"] if "-REQ-" in k)
retired_stories = retired_count - retired_reqs

new_stories = mapping["new_stories"]
# Count new requirements in agile_backlog.json under those new stories
new_req_count = 0
new_story_obj_map = {}
for epic in backlog["epics"]["epic"]:
    if epic.get("epic_id") == "EPIC-008":
        for feat in epic.get("features", {}).get("feature", []):
            if feat.get("feature_id") == "EPIC-008-F-002":
                for story in feat.get("stories", {}).get("story", []):
                    sid = story.get("story_id")
                    if sid in new_stories:
                        reqs = story.get("requirements", {}).get("requirement", [])
                        new_req_count += len(reqs)
                        new_story_obj_map[sid] = (story, reqs)

# F-002 totals
f002_total_stories = 0
f002_total_reqs = 0
for epic in backlog["epics"]["epic"]:
    if epic.get("epic_id") == "EPIC-008":
        for feat in epic.get("features", {}).get("feature", []):
            if feat.get("feature_id") == "EPIC-008-F-002":
                stories = feat.get("stories", {}).get("story", [])
                f002_total_stories = len(stories)
                for s in stories:
                    f002_total_reqs += len(s.get("requirements", {}).get("requirement", []))

# ---- Build the doc ----
doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

doc.add_heading("EPIC-008-F-002 Backlog Consolidation — V1", level=1)
p = doc.add_paragraph()
p.add_run("Author: ").bold = True
p.add_run("Claude (autonomous backlog consolidation pass)")
p = doc.add_paragraph()
p.add_run("Date: ").bold = True
p.add_run("2026-04-27")
p = doc.add_paragraph()
p.add_run("Source of truth: ").bold = True
p.add_run("CH-EPIC8-Feachure-002-EngineeringRulesEnforcement-designV18.docx (BR-1..14, TR-1..17)")
p = doc.add_paragraph()
p.add_run("Inventory: ").bold = True
p.add_run(f"_oneshots/test_output/scan_validate_relevance.json ({flagged_count} flagged entities)")

# Executive summary
doc.add_heading("1. Executive summary", level=2)
doc.add_paragraph(
    "Per the architect's directive, this pass refactors the agile backlog so that "
    "EPIC-008-F-002 (Engineering Rules Enforcement) reflects the V1 framework that "
    "shipped on 2026-04-27. Every business and technical requirement from V18 is "
    "reified as exactly one backlog requirement under one of four new stories under "
    "F-002. Old story-level records that the framework supersedes have been "
    "physically removed; all foreign-key pointers have been repointed."
)

p = doc.add_paragraph()
p.add_run("Milestone: ").bold = True
p.add_run(
    "Once Skip approves the status field flip, Rule-008-ENF-001 (the first "
    "enforcement entry in engineering_rules.json) becomes the first rule in the "
    "ChatHealthy.ai brain with TWO completed enforcement functions — _scan_http "
    "and _validate_json — both anchored to fresh requirements under "
    "EPIC-008-F-002-S-005 and EPIC-008-F-002-S-006."
)

doc.add_heading("Counts", level=3)
table = doc.add_table(rows=1, cols=2)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
hdr[0].text = "Metric"
hdr[1].text = "Value"
for label, value in [
    ("Classifier-flagged candidates (input)", str(flagged_count)),
    ("  ... of which were requirements", str(flagged_reqs)),
    ("Retired records (physically removed)", f"{retired_count} ({retired_stories} stories + {retired_reqs} requirements)"),
    ("New requirements created under F-002", str(new_req_count)),
    ("New stories created under F-002", f"{len(new_stories)} (S-003 through S-006)"),
    ("Total stories under F-002 after pass", str(f002_total_stories)),
    ("Total requirements under F-002 after pass", str(f002_total_reqs)),
    ("V18 BR + TR slots covered", "31 (BR-1..14 + TR-1..17), each in exactly one bucket"),
]:
    row = table.add_row().cells
    row[0].text = label
    row[1].text = value

# Mapping table
doc.add_heading("2. Mapping table — old_req_id -> new_req_id", level=2)
doc.add_paragraph(
    "Every retired identifier and where its semantic content (or its FK pointers) "
    "now lives. There are no superseded_by markers, no tombstone records, and no "
    "ghost IDs in the live JSON; the entries below are all that remains of the "
    "audit trail."
)
table = doc.add_table(rows=1, cols=3)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
hdr[0].text = "Retired ID"
hdr[1].text = "New ID"
hdr[2].text = "Reason"
for old_id, info in mapping["retired_to_new"].items():
    row = table.add_row().cells
    row[0].text = old_id
    row[1].text = info["new_id"]
    row[2].text = info["reason"]

# 4 new stories
doc.add_heading("3. The four new stories under EPIC-008-F-002", level=2)
for sid in new_stories:
    story, reqs = new_story_obj_map[sid]
    doc.add_heading(f"{sid} — {story['name']}", level=3)
    doc.add_paragraph(story.get("description", ""))
    p = doc.add_paragraph()
    p.add_run("Status: ").bold = True
    p.add_run(story.get("status", "?"))
    p.add_run("  |  Approval: ").bold = True
    p.add_run(story.get("approval", "?"))
    p.add_run("  |  Size: ").bold = True
    p.add_run(story.get("size", "?"))
    p = doc.add_paragraph()
    p.add_run("V18 BR/TR slots assigned: ").bold = True
    p.add_run(", ".join(mapping["br_tr_assignment"][sid]))
    doc.add_paragraph(f"Requirements ({len(reqs)}):")
    rt = doc.add_table(rows=1, cols=3)
    rt.style = "Light Grid Accent 1"
    h = rt.rows[0].cells
    h[0].text = "req_id"
    h[1].text = "name"
    h[2].text = "Anchored to V18"
    for req in reqs:
        r = rt.add_row().cells
        r[0].text = req["req_id"]
        r[1].text = req["name"]
        # Pull the V18 anchor out of the name (we set it as "BR: BR-N" or "TR: TR-N")
        r[2].text = req["name"].split(":", 1)[-1].strip() if ":" in req["name"] else ""

# Pointer-rewrite ledger
doc.add_heading("4. Pointer-rewrite ledger", level=2)
doc.add_paragraph(
    "Every git-tracked file that contained a retired identifier was rewritten in "
    "place. Counts below are taken directly from the repoint_fks.py log."
)
total_files = len(repoint_log)
total_replacements = sum(item["replacements"] for item in repoint_log)
p = doc.add_paragraph()
p.add_run(f"Files rewritten: {total_files}").bold = True
p.add_run(f"   Total identifier replacements: {total_replacements}")

table = doc.add_table(rows=1, cols=3)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
hdr[0].text = "File"
hdr[1].text = "Replacements"
hdr[2].text = "Details"
for item in repoint_log:
    row = table.add_row().cells
    row[0].text = item["path"]
    row[1].text = str(item["replacements"])
    row[2].text = "\n".join(item["details"])

doc.add_paragraph(
    "Plus agile_backlog.json itself, which was edited by migrate.py: 9 retired "
    "records removed (2 stories + 7 requirements), 4 new stories with 39 new "
    "requirements added, 0 dangling depends_on entries (a backlog-wide scan "
    "found no requirements whose depends_on pointed at any of the retired IDs)."
)

# FK integrity
doc.add_heading("5. FK integrity verdict", level=2)
doc.add_paragraph("Verdict: CLEAN. After the pass:")
doc.add_paragraph(
    "  - agile_backlog.json validates against ChatHealthyAgileBacklogSchema.json (Draft 2020-12).",
    style="List Bullet",
)
doc.add_paragraph(
    "  - bugs.json validates against ChatHealthyBugsSchema.json. Every bug.req_id "
    "that follows the EPIC-NNN-F-NNN-S-NNN-REQ-* pattern resolves to an existing "
    "requirement in the backlog, except for three pre-existing orphan FKs "
    "(BUG-GOV-008 -> EPIC-008-F-007-S-007-REQ-B-001; BUG-DEFER-002 -> "
    "EPIC-008-F-007-S-011-REQ-T-002; BUG-DEFER-003 -> EPIC-002-F-003-S-002-REQ-T-001) "
    "which predate this consolidation and are out of scope for this pass.",
    style="List Bullet",
)
doc.add_paragraph(
    "  - engineering_rules.json validates against EngineeringRulesSchema.json. "
    "Every rule_statement.req_id that is not the literal 'Orphan' resolves to an "
    "existing requirement (Rule-008's three rule_statements currently carry "
    "req_id='Orphan' awaiting Skip's anchoring step).",
    style="List Bullet",
)
doc.add_paragraph(
    "  - version.json validates against ChatHealthyVersionSchema.json. Three "
    "historical build records that previously cited retired IDs now cite the new "
    "F-002 IDs.",
    style="List Bullet",
)
doc.add_paragraph(
    "  - Repo-wide grep for any of the nine retired identifiers returns zero hits "
    "outside the migration scripts under Code/_pending/backlog_consolidation/ "
    "(which document the mapping by design).",
    style="List Bullet",
)

# Ambiguity
doc.add_heading("6. Ambiguity calls and judgment", level=2)
doc.add_paragraph("Two judgment calls worth recording:")

doc.add_paragraph(
    "(a) The classifier output flagged 71 entities — but most of them (the EPIC-001, "
    "EPIC-002, EPIC-006, EPIC-009 reqs about Pydantic input models in product code, "
    "the EPIC-008-F-003-S-001 boot/mode requirements, the EPIC-008-F-004 deploy "
    "requirements, and the EPIC-005 schema-design requirements) are out of scope for "
    "the V18 framework. They describe schema validation in product code paths, not "
    "the EnforcementWorker dispatch system. Per Skip's directive that the four "
    "buckets cover \"HTTP vs HTTPS, file scanning, rules infrastructure, or JSON "
    "validation\" as defined by V18, those entities stay where they are. The "
    "classifier was over-broad on \"json_validation\" because it matched any "
    "appearance of schema discipline anywhere in the codebase.",
    style="List Bullet",
)

doc.add_paragraph(
    "(b) EPIC-008-F-008 (Code and Security Scanning) was flagged as a feature "
    "candidate for retirement. The two flagged stories under it (S-002 Security "
    "Code Scan and S-004 JSON Scan) were retired because they semantically "
    "duplicate _scan_http and _validate_json — both of which now live as concrete "
    "requirements under F-002-S-005 and F-002-S-006. The other stories under "
    "F-008 (S-001 Bug Scan, S-003 Pytest Scans, S-005 General Scanning "
    "Requirements) were NOT flagged by the classifier and are NOT semantically "
    "redundant with V18; they remain in F-008 untouched. F-008 the feature stays "
    "alive with its remaining three stories. Killing the parent feature would "
    "have orphaned them.",
    style="List Bullet",
)

doc.add_paragraph(
    "(c) One classifier entry, EPIC-008-F-004-S-001-REQ-T-014, has an empty "
    "name in the source data and the corresponding requirement could not be "
    "located in the backlog by the same id. It was treated as already-orphaned "
    "and ignored; no action taken.",
    style="List Bullet",
)

# Bugs status
doc.add_heading("7. Bug-record status after this pass", level=2)
doc.add_paragraph(
    "Three bugs that pointed at retired requirements were repointed by repoint_fks.py. "
    "Their statuses were not modified by this pass; status updates remain Skip's "
    "to apply once he's reviewed the new ID anchors. The repointed bugs:"
)
doc.add_paragraph("  - BUG-BACKLOG-SCHEMA-UNENFORCED-001  (status=new, was req_id=EPIC-008-F-008-S-004-REQ-T-002, now EPIC-008-F-002-S-006-REQ-T-001)", style="List Bullet")
doc.add_paragraph("  - BUG-RULES-NO-REQ-TRACE-001         (status=deferred, was req_id=EPIC-008-F-008-S-004-REQ-T-003, now EPIC-008-F-002-S-003-REQ-T-001)", style="List Bullet")
doc.add_paragraph("  - BUG-SCANNER-NOT-SINGLE-GATE-001    (status=new, was req_id=EPIC-008-F-008-S-004-REQ-B-002, now EPIC-008-F-002-S-003-REQ-B-001)", style="List Bullet")

# Save
OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(f"Wrote {OUT}")
