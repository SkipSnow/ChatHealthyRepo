# Copyright (c) 2026 ChatHealthy.ai LLC. All rights reserved.
# Licensed under the FindCare Evaluation License (FEL-1.0).
"""
One-shot writer for the compliance/engineering-rules DEFERRAL document.

Repackages the work built in the past 24 hours, formally specifies the two
deferred work products as proposed B+T requirement pairs awaiting human
approval, and reaffirms the alpha focus per v4-022. Output is a .docx
under brain/BusinessArtifacts/governance/.
"""

from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

REPO_ROOT = Path(__file__).resolve().parents[4]
OUT = (REPO_ROOT / "brain" / "BusinessArtifacts" / "governance" /
       "compliance_crosswalk_deferral_and_proposed_requirements.docx")


def _p(doc, text, italic=False, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.size = Pt(11)
    if italic: r.italic = True
    if bold: r.bold = True
    return p


def _bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text).font.size = Pt(11)
    return p


def _kv_table(doc, rows):
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Light Grid Accent 1"
    for i, (k, v) in enumerate(rows):
        t.rows[i].cells[0].text = str(k)
        t.rows[i].cells[1].text = str(v)
    return t


def main():
    doc = Document()

    title = doc.add_heading(
        "Compliance Crosswalk + Engineering-Rules Schema — Deferral & Proposed Requirements",
        level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _p(doc, f"Generated {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}.",
       italic=True)

    # ── 1. Decision ──────────────────────────────────────────────
    doc.add_heading("1. Human decision (2026-04-19)", level=1)
    _p(doc, "Both work products described below are FORMALLY SPECIFIED as "
            "proposed requirements but DEFERRED until after alpha (per "
            "v4-022 alpha-freeze focus). No further build work proceeds on "
            "either until human approves the requirements and lifts the "
            "freeze for them. Focus returns to application code (FindCare "
            "scoring engine, measures, end-to-end wiring).")

    # ── 2. What's already built (the asset to repackage) ─────────
    doc.add_heading("2. What is already built (assets — last 24 hours)", level=1)
    _kv_table(doc, [
        ("Compliance Crosswalk Agent (single class file)",
         "Code/Shared/ops/tools/compliance_crosswalk_agent.py"),
        ("Pytest suite (23 tests, all green; LLM mocked)",
         "Code/Shared/ops/tools/tests/test_compliance_crosswalk_agent.py"),
        ("Master crosswalk JSON (1,204 controls; 48 with implementations)",
         "brain/machine_artifacts/content/SecurityAuditControls.json"),
        ("Compliance crosswalk review (.docx)",
         "brain/BusinessArtifacts/governance/compliance_crosswalk_review.docx"),
        ("Engineering rules review V2 (the schema design source)",
         "brain/BusinessArtifacts/governance/all_engineering_rules_review_V2.docx"),
        ("This deferral document",
         "brain/BusinessArtifacts/governance/compliance_crosswalk_deferral_and_proposed_requirements.docx"),
    ])
    _p(doc, "Public agent surface (per human directive 2026-04-19):")
    _bullet(doc, "ComplianceCrosswalkAgent.__init__(scan_dirs, model, api_key, ...)")
    _bullet(doc, "ComplianceCrosswalkAgent.run(file_types=FileType.ALL) "
                 "→ full pipeline; returns Path to .docx review")
    _bullet(doc, "ComplianceCrosswalkAgent.ComplianceCrossWalk(file_types) "
                 "→ self (chainable)")
    _bullet(doc, "ComplianceCrosswalkAgent.ComplianceCrosswalkReport(out_path=None) "
                 "→ Path")
    _bullet(doc, "FileType enum: ALL (default), PYTHON, YAML, JSON, HTML, "
                 "SHELL, BAT, POWERSHELL, JAVASCRIPT, TYPESCRIPT, TSX, TF, CFG, INI")

    # ── 3. Work product 1 — engineering_rules schema redesign ────
    doc.add_heading("3. Work product 1 — engineering_rules.json schema redesign",
                    level=1)
    _p(doc, "Designed in: brain/BusinessArtifacts/governance/all_engineering_rules_review_V2.docx")
    _p(doc, "Parent backlog requirement (already in agile_backlog.json): "
            "BRAIN-ENGINEERING-RULES-REQ-001 [not_started] — "
            "\"engineering_rules.json MUST be controlled by a JSON schema and "
            "in strict compliance at all times.\" The schema below is the "
            "design that fulfills that requirement.", italic=True)

    doc.add_heading("3.1 Target schema shape", level=2)
    _kv_table(doc, [
        ("Top-level", "rules: rule[] (array, no other root keys beyond metadata)"),
        ("rule.id", "Rule-NNN (sequential 3-digit; rename from v4-NNN / DR-NNN; "
                    "all references in non-loose-schema JSONs and bugs.json must be updated)"),
        ("rule.name", "Short string, ≤80 chars (replaces current 'title')"),
        ("rule.description", "Non-binding string (replaces current 'rule' body, "
                              "reframed as descriptive)"),
        ("rule.rule_statements[]",
         "List of {statement_number: int (sequential from 1), "
         "statement_body: string (Boolean-testable)}. The binding part of the rule."),
        ("rule.date", "Empty values must be filled with the redesign date (today)"),
        ("rule.type", "Keep ('standard' | 'incident-driven')"),
        ("rule.source", "Keep"),
    ])

    doc.add_heading("3.2 Proposed B+T requirement pair (for human approval)", level=2)
    _kv_table(doc, [
        ("req_id", "EPIC-008-F-007-S-011-REQ-B-001"),
        ("type", "business"),
        ("priority", "must-have"),
        ("status", "proposed (awaiting human approval)"),
        ("requirement",
         "engineering_rules.json MUST express each rule as (a) a short "
         "human-readable name, (b) a non-binding description, and (c) one or "
         "more individually Boolean-testable statements. Rule identifiers "
         "MUST be uniform (Rule-NNN). Every rule MUST carry a date. The "
         "schema MUST be enforced — non-compliant content MUST NOT be "
         "committed."),
    ])
    _kv_table(doc, [
        ("req_id", "EPIC-008-F-007-S-011-REQ-T-002"),
        ("type", "technical"),
        ("priority", "must-have"),
        ("status", "proposed (awaiting human approval)"),
        ("parent_req_id", "EPIC-008-F-007-S-011-REQ-B-001 + the existing "
                          "BRAIN-ENGINEERING-RULES-REQ-001"),
        ("requirement",
         "Publish a strict JSON schema (Draft 2020-12) on chathealthy.ai/schemas "
         "that defines rules[].rule with fields {id (^Rule-\\d{3}$), name "
         "(maxLength 80), description (string), rule_statements[] of "
         "{statement_number (integer ≥ 1), statement_body (string)}, date "
         "(date), type (enum: standard|incident-driven), source (string)}. "
         "Reference the schema from engineering_rules.json $schema. Migrate "
         "the file to the new shape in one commit; in the same commit, "
         "rename every v4-NNN / DR-NNN identifier to Rule-NNN and update "
         "every reference in non-loose-schema brain JSONs and in bugs.json. "
         "preCommitScan.py MUST be updated to iterate rule_statements[] "
         "instead of the prose 'rule' body. "
         "Pytests cover schema validation, ID-rename completeness, and "
         "no-broken-references across all consumers."),
    ])

    doc.add_heading("3.3 Acceptance criteria", level=2)
    _bullet(doc, "engineering_rules.json validates against the published "
                 "strict schema (not the relaxed schema).")
    _bullet(doc, "Every rule has Rule-NNN id, name ≤80 chars, "
                 "rule_statements[] with sequential statement_number from 1.")
    _bullet(doc, "grep -r 'v4-' across the repo returns zero hits in active "
                 "JSONs and zero hits in bugs.json (legacy mentions in "
                 "human-language docs allowed but flagged).")
    _bullet(doc, "preCommitScan.py exercises rule_statements[] when "
                 "dispatching enforcement.")
    _bullet(doc, "controls.json (now SecurityAuditControls.json) is "
                 "regenerated against the new rule IDs after the rename.")

    # ── 4. Work product 2 — agent ingestion of engineering_rules ─
    doc.add_heading("4. Work product 2 — Compliance crosswalk agent: "
                    "third framework (engineering rules)", level=1)
    _p(doc, "Closes the loop: external compliance frameworks ↔ internal "
            "engineering rules ↔ source code, all in one master crosswalk.")

    doc.add_heading("4.1 Proposed B+T requirement pair (for human approval)",
                    level=2)
    _kv_table(doc, [
        ("req_id", "EPIC-002-F-003-S-002-REQ-B-001"),
        ("type", "business"),
        ("priority", "should-have"),
        ("status", "proposed (deferred — awaiting human approval and post-alpha)"),
        ("requirement",
         "The system MUST be able to demonstrate, for any source file or any "
         "external compliance control (NIST 800-53, HITRUST CSF), the "
         "internal engineering rules that govern it. The crosswalk MUST be a "
         "single artifact, machine-readable, audit-presentable."),
    ])
    _kv_table(doc, [
        ("req_id", "EPIC-002-F-003-S-002-REQ-T-001"),
        ("type", "technical"),
        ("priority", "should-have"),
        ("status", "proposed (deferred — awaiting human approval and post-alpha)"),
        ("requirement",
         "ComplianceCrosswalkAgent gains a third framework: engineering_rules. "
         "Pass 1 ingests engineering_rules.json (post-redesign, so the unit of "
         "ingestion is rule_statements[]) as its own framework block. Pass 2 "
         "extends the classifier prompt so each implementation entry can "
         "reference both external control IDs and engineering rule IDs. "
         "controls.json gains framework: 'engineering_rules' alongside "
         "nist_800_53 and hitrust_csf. The .docx review surfaces a fourth "
         "section: rules-without-implementation (potential dead rules)."),
    ])

    doc.add_heading("4.2 Dependencies", level=2)
    _bullet(doc, "Hard dependency on Work Product 1 (engineering_rules.json "
                 "must be in the new shape with rule_statements[] before "
                 "ingestion makes sense).")
    _bullet(doc, "Soft dependency on a HITRUST CSF v11 catalog source (the "
                 "current seed-only HITRUST gap remains either way).")

    # ── 5. Alpha focus reaffirmation ─────────────────────────────
    doc.add_heading("5. Alpha focus reaffirmation (v4-022)", level=1)
    _p(doc, "Both work products above are governance / brain-housekeeping "
            "work. Per v4-022 (alpha freeze 2026-04-07 → 2026-05-01), "
            "neither proceeds without human approval AND GPT (Enterprise "
            "Architect) concurrence. The compliance crosswalk agent and "
            "SecurityAuditControls.json that already exist were built "
            "overnight under human's risk-acceptance for the engineering-"
            "rules curation work; they remain as-is (not committed pending "
            "the requirement approval per v4-023).")
    _p(doc, "Effective immediately, focus returns to alpha-critical work — "
            "the application: scoring engine, measures, data loading, "
            "end-to-end wiring. No further work on either deferred product "
            "until human explicitly lifts the deferral.", bold=True)

    # ── 6. Status of unsubmitted artifacts ───────────────────────
    doc.add_heading("6. Pending git status (nothing committed overnight)", level=1)
    _bullet(doc, "compliance_crosswalk_agent.py — present in working tree, "
                 "not committed (v4-023 — needs ARCH-COMPLIANCE-CROSSWALK-001 "
                 "B+T pair approved first).")
    _bullet(doc, "tests/test_compliance_crosswalk_agent.py — same.")
    _bullet(doc, "brain/machine_artifacts/content/SecurityAuditControls.json "
                 "— renamed (was controls.json), populated; not committed.")
    _bullet(doc, "brain/BusinessArtifacts/governance/compliance_crosswalk_review.docx "
                 "— regenerated against renamed file; not committed.")
    _bullet(doc, "This deferral document — not committed.")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT} ({OUT.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
