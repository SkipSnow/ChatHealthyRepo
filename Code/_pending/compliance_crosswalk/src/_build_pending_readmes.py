# Copyright (c) 2026 ChatHealthy.ai LLC. All rights reserved.
# One-shot writer for the README.docx files under Code/_pending/.
# Lives in the compliance_crosswalk src/ because the script that knows
# the pending structure is part of the deferred work itself.

from __future__ import annotations

from pathlib import Path
from datetime import datetime, timezone

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

REPO = Path(r"c:/chatHealthy/findCare")
PENDING = REPO / "Code" / "_pending"
DATE = datetime.now(timezone.utc).strftime("%Y-%m-%d")


def make_readme(out_path, title, status, requirements, tracking_bugs,
                inputs, outputs, src_files, how_to_resume,
                dependencies="", extra_notes=""):
    doc = Document()
    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _p(text, italic=False, bold=False):
        p = doc.add_paragraph()
        r = p.add_run(text); r.font.size = Pt(11)
        if italic: r.italic = True
        if bold: r.bold = True
    def _bullet(text):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(text).font.size = Pt(11)
    def _kv(rows):
        t = doc.add_table(rows=len(rows), cols=2)
        t.style = "Light Grid Accent 1"
        for i, (k, v) in enumerate(rows):
            t.rows[i].cells[0].text = str(k)
            t.rows[i].cells[1].text = str(v)

    _p(f"Generated {DATE}.", italic=True)
    _p("This is a deferred-work directory under Code/_pending/. "
       "Per engineering rule v4-041, every deferred unit of work MUST "
       "live here with input/, output/, src/, and a README. Do NOT "
       "promote to QA or prod until the linked requirement is "
       "approved-and-completed and the linked tracking bug is closed.")

    doc.add_heading("Status", level=1)
    _kv([("Status", status), ("Date deferred", DATE)])

    doc.add_heading("Linked requirements (agile_backlog.json)", level=1)
    if requirements:
        for r in requirements: _bullet(r)
    else:
        _p("(none yet)", italic=True)

    doc.add_heading("Linked tracking bugs (bugs.json)", level=1)
    if tracking_bugs:
        for b in tracking_bugs: _bullet(b)
    else:
        _p("(none yet)", italic=True)

    if dependencies:
        doc.add_heading("Dependencies", level=1)
        _p(dependencies)

    doc.add_heading("Inputs (live brain paths - do not copy)", level=1)
    if inputs:
        for inp in inputs: _bullet(inp)
    else:
        _p("(none required)", italic=True)

    doc.add_heading("Outputs (live brain paths)", level=1)
    if outputs:
        for o in outputs: _bullet(o)
    else:
        _p("(none yet - work not started)", italic=True)

    doc.add_heading("Source files in src/", level=1)
    if src_files:
        for s in src_files: _bullet(s)
    else:
        _p("(empty - work not started)", italic=True)

    doc.add_heading("How to resume", level=1)
    _p(how_to_resume)

    if extra_notes:
        doc.add_heading("Notes", level=1)
        _p(extra_notes)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def main():
    # 1. compliance_crosswalk
    make_readme(
        PENDING / "compliance_crosswalk" / "README.docx",
        "Pending: Compliance Crosswalk Agent (NIST 800-53 + HITRUST CSF v11)",
        "deferred - built but not approved into prod",
        ["EPIC-002-F-002-S-002-REQ-B-001 (business, must-have)",
         "EPIC-002-F-002-S-002-REQ-T-001 (technical, must-have)"],
        ["BUG-DEFER-001 - severity_by_environment={dev: low, qa: critical, prod: show_stopper}"],
        ["NIST OSCAL JSON catalog (downloaded by the agent at runtime)",
         "brain/machine_artifacts/content/security.json - HITRUST seed",
         "brain/machine_artifacts/content/engineering_rules.json - referenced",
         "Code/.env Anthropic_API_KEY - required for Pass 2 LLM calls"],
        ["brain/machine_artifacts/content/SecurityAuditControls.json - master crosswalk",
         "brain/BusinessArtifacts/governance/compliance_crosswalk_review.docx",
         "brain/BusinessArtifacts/governance/compliance_crosswalk_deferral_and_proposed_requirements.docx"],
        ["compliance_crosswalk_agent.py - single class file (FileType enum + ComplianceCrosswalkAgent)",
         "tests/test_compliance_crosswalk_agent.py - 23 tests, LLM mocked",
         "_build_deferral_doc.py - one-shot writer for the deferral .docx",
         "_apply_brain_edits_2026_04_19.py - one-shot writer that applied the four brain edits human approved 2026-04-19",
         "_build_pending_readmes.py - this README generator"],
        "Run from this src/: python compliance_crosswalk_agent.py status - confirms "
        "SecurityAuditControls.json is reachable. Then python compliance_crosswalk_agent.py "
        "run for full pipeline. Tests: python -m pytest tests/ -v. To approve into prod, "
        "complete ARCH-COMPLIANCE-CROSSWALK-001 requirements, close BUG-DEFER-001, "
        "and move src/ files back into Code/Shared/ops/tools/ in one commit."
    )

    # 2. compliance_crosswalk_third_framework
    make_readme(
        PENDING / "compliance_crosswalk_third_framework" / "README.docx",
        "Pending: Compliance Crosswalk Agent - Third Framework (engineering rules)",
        "deferred - design captured, work not started",
        ["EPIC-002-F-003-S-002-REQ-B-001 (business, should-have)",
         "EPIC-002-F-003-S-002-REQ-T-001 (technical, should-have)"],
        ["BUG-DEFER-003"],
        ["brain/machine_artifacts/content/engineering_rules.json (POST-redesign - "
         "depends on engineering_rules_redesign)"],
        ["Extension to brain/machine_artifacts/content/SecurityAuditControls.json "
         "(frameworks.engineering_rules block)"],
        [],
        "Hard-blocked on engineering_rules_redesign completing first (must consume "
        "rule_statements[]). When ready, extend the agent at "
        "../compliance_crosswalk/src/compliance_crosswalk_agent.py to ingest "
        "engineering_rules.json as a third framework. Update the classifier prompt "
        "to include rule_statements as classification targets. Add tests.",
        dependencies="Hard dependency on EPIC-008-F-007-S-011-REQ-T-002 "
                     "(engineering rules schema redesign)."
    )

    # 3. engineering_rules_redesign
    make_readme(
        PENDING / "engineering_rules_redesign" / "README.docx",
        "Pending: engineering_rules.json schema redesign",
        "deferred - design captured in V2 review .docx, migration not started",
        ["EPIC-008-F-007-S-011-REQ-B-001 (business, must-have)",
         "EPIC-008-F-007-S-011-REQ-T-002 (technical, must-have)",
         "(parent) BRAIN-ENGINEERING-RULES-REQ-001 - already in backlog as not_started"],
        ["BUG-DEFER-002"],
        ["brain/BusinessArtifacts/governance/all_engineering_rules_review_V2.docx - the design source",
         "brain/machine_artifacts/content/engineering_rules.json - the file to migrate"],
        ["Migrated engineering_rules.json (Rule-NNN ids, name, description, "
         "rule_statements[], dates)",
         "New strict JSON schema published at chathealthy.ai/schemas/dev/EngineeringRulesSchema.json",
         "Updates to bugs.json + every other governed JSON that references v4-NNN/DR-NNN rule ids"],
        [],
        "Read the V2 design .docx. Author the strict JSON schema. Migrate "
        "engineering_rules.json. In the SAME commit, rename every v4-NNN/DR-NNN "
        "reference across non-loose-schema brain JSONs and bugs.json to Rule-NNN. "
        "Update preCommitScan.py + pre_deploy_rule_check.py to iterate "
        "rule_statements[]. Add pytests for schema validation and ID-rename completeness."
    )

    # 4. deferred_work_pattern
    make_readme(
        PENDING / "deferred_work_pattern" / "README.docx",
        "Pending: Deferred work as a first-class concept (incl. BugManagerAgent)",
        "deferred - pattern codified in v4-041; auto-escalation agent not built",
        ["EPIC-008-F-009-S-001-REQ-B-001 (business, must-have)",
         "EPIC-008-F-009-S-001-REQ-T-001 (technical, must-have)"],
        ["BUG-DEFER-004"],
        ["brain/machine_artifacts/content/engineering_rules.json - v4-041 (this pattern) + v4-042 (assignment scope)",
         "brain/machine_artifacts/content/agile_backlog.json - every requirement with status=deferred",
         "brain/machine_artifacts/content/bugs.json - every BUG-DEFER-* tracking bug",
         "Website/schemas/dev/ChatHealthyBugsSchema.json - severity_by_environment field"],
        ["BugManagerAgent (future) - auto-escalates bug severity on environment promotion",
         "(possibly) a pre-deploy gate that refuses promotion to QA/prod if any "
         "deferred-tracking bug for must-have requirements is still open above "
         "the severity threshold for the target environment"],
        [],
        "Build the BugManagerAgent class. It watches every requirement transition "
        "(status, environment) and updates the paired tracking bug severity per "
        "the severity_by_environment map. It also auto-files new BUG-DEFER-* bugs "
        "when a new requirement is marked status=deferred. Optionally wire into "
        "the pre-push hook so severity escalation is checked at deploy time.",
        extra_notes="human directive 2026-04-19 (refinement to v4-042): each "
                    "engineering rule should carry per-rule enforcement config "
                    "declaring whether (and which) human assignments override "
                    "it. The BugManagerAgent and v4-042 refinement should land "
                    "together so 'assignment-over-rule' is configurable per rule "
                    "rather than a coarse blanket. Capture this iteration as a "
                    "follow-on requirement when this unit is reactivated."
    )

    # Top-level _pending/README.docx
    make_readme(
        PENDING / "README.docx",
        "Code/_pending/ - Deferred work directory (governed by v4-041)",
        "active",
        ["Each subdirectory = one deferred unit of work. Each unit has paired "
         "requirement(s) in agile_backlog.json and a tracking bug in bugs.json. "
         "See the per-unit README.docx for details."],
        [],
        [],
        [],
        ["compliance_crosswalk/ - built, awaiting prod approval",
         "compliance_crosswalk_third_framework/ - design captured, blocked on engineering_rules_redesign",
         "engineering_rules_redesign/ - design captured in V2 .docx, migration not started",
         "deferred_work_pattern/ - pattern codified in v4-041, BugManagerAgent not built"],
        "Per v4-041: a deferred unit closes only when its requirement is "
        "fulfilled AND its tracking bug is closed AND the src/ files are moved "
        "out of Code/_pending/. Per v4-042: a human assignment overrides ordinary "
        "engineering rules during execution of that assignment (with named "
        "exceptions for prod sign-off and governance matrix changes)."
    )

    print("README.docx files written:")
    for p in sorted(PENDING.glob("**/README.docx")):
        print(f"  {p.relative_to(REPO)} ({p.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
