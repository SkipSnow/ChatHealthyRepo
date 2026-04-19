# Copyright (c) 2026 ChatHealthy.ai LLC. All rights reserved.
# Comprehensive in-repo sweep: every "human" → "human" with grammatical
# capitalization. Covers all file types I can safely touch from Python:
# - Plain-text files of any extension
# - Files without extension (Code/.env, Caddyfile, Dockerfile, Makefile, etc.)
# - .docx via python-docx (paragraphs + tables + headers + footers)
# - .xlsx via openpyxl (cell values)
# - .gitignored backups (.jsonbak)
# - All previously-swept extensions re-checked for completeness
#
# OUT OF SCOPE (require external action):
# - Git history (commit messages) — requires history rewrite
# - External services (LangSmith projects, HF Spaces, Azure resources,
#   MongoDB collections, Cloudflare DNS) — require per-service API access
# - .pyc bytecode — regenerates on next Python run
# - .lnk Windows shortcuts — binary; skipped
# - Image/PDF binaries — content not text-editable
#
# Per Human directive 2026-04-19 "fix everything ... this is only about human".

from __future__ import annotations

import re
from pathlib import Path

ROOTS = [
    Path(r"c:/chatHealthy/findCare"),
    Path(r"C:/Users/skips/.claude/projects/c--chatHealthy-findCare/memory"),
]

# Plain-text-ish: read as utf-8, regex sub
PLAINTEXT_EXTS = {
    ".json", ".py", ".md", ".sh", ".bat", ".ps1", ".tsx", ".ts", ".js",
    ".html", ".yml", ".yaml", ".toml", ".cfg", ".ini", ".txt",
    ".jsonbak", ".env", ".envrc", ".http", ".sql", ".css", ".scss",
    ".dockerfile",
}
# Files with NO extension that are usually plain text
NOEXT_NAMES = {
    "Caddyfile", "Dockerfile", "Makefile", ".env", ".envrc",
    ".gitignore", ".gitattributes", "LICENSE", "README",
}
# Binary file types we know we cannot text-edit safely
SKIP_BINARY_EXTS = {
    ".png", ".jpg", ".jpeg", ".gif", ".pdf", ".zip", ".tar", ".gz",
    ".pyc", ".pyo", ".so", ".dll", ".exe", ".dylib", ".bin", ".lnk",
    ".woff", ".woff2", ".ttf", ".eot", ".ico", ".svg",
}
SKIP_DIRS = {"__pycache__", ".git", "node_modules", ".venv",
             ".pytest_cache", "dist", "build", "out"}

WORD = re.compile(r"\bboss\b", re.IGNORECASE)


def _is_sentence_or_header_start(text: str, idx: int) -> bool:
    line_start = text.rfind("\n", 0, idx) + 1
    line_prefix = text[line_start:idx]
    stripped_prefix = line_prefix.lstrip()
    if stripped_prefix.startswith("#"):
        body = stripped_prefix.lstrip("#").lstrip()
        if body == "" or body.endswith((":", ".", "!", "?")):
            return True
    if idx == 0:
        return True
    j = idx - 1
    while j >= 0 and text[j] in " \t":
        j -= 1
    if j < 0:
        return True
    if text[j] in ".!?\n:>":
        return True
    if not line_prefix.strip(" \t-*+0123456789.)>"):
        return True
    return False


def _replace_in_text(text: str) -> str:
    def sub(m):
        word = m.group(0)
        if word.isupper() and len(word) >= 2:
            return "HUMAN"
        return "Human" if _is_sentence_or_header_start(text, m.start()) else "human"
    return WORD.sub(sub, text)


def process_text(p: Path) -> int:
    try:
        text = p.read_text(encoding="utf-8")
    except (UnicodeDecodeError, PermissionError):
        return 0
    if not WORD.search(text):
        return 0
    new = _replace_in_text(text)
    if new == text:
        return 0
    p.write_text(new, encoding="utf-8")
    return len(WORD.findall(text))


def process_docx(p: Path) -> int:
    from docx import Document
    try:
        doc = Document(p)
    except Exception:
        return 0
    count = 0
    def edit_runs(paragraphs):
        nonlocal count
        for para in paragraphs:
            full = para.text
            if not WORD.search(full):
                continue
            new = _replace_in_text(full)
            if new == full:
                continue
            count += len(WORD.findall(full))
            # Replace text by clearing all runs and putting the new text on the first run
            if para.runs:
                para.runs[0].text = new
                for r in para.runs[1:]:
                    r.text = ""
    edit_runs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                edit_runs(cell.paragraphs)
    for section in doc.sections:
        edit_runs(section.header.paragraphs)
        edit_runs(section.footer.paragraphs)
    if count > 0:
        try:
            doc.save(p)
        except Exception:
            return 0
    return count


def process_xlsx(p: Path) -> int:
    from openpyxl import load_workbook
    try:
        wb = load_workbook(p)
    except Exception:
        return 0
    count = 0
    for sheet in wb.worksheets:
        for row in sheet.iter_rows():
            for cell in row:
                if cell.value is None:
                    continue
                if not isinstance(cell.value, str):
                    continue
                if not WORD.search(cell.value):
                    continue
                new = _replace_in_text(cell.value)
                if new != cell.value:
                    count += len(WORD.findall(cell.value))
                    cell.value = new
    if count > 0:
        try:
            wb.save(p)
        except Exception:
            return 0
    return count


def main():
    total = 0
    files = []
    for root in ROOTS:
        if not root.exists():
            continue
        for p in root.rglob("*"):
            if not p.is_file():
                continue
            if any(d in p.parts for d in SKIP_DIRS):
                continue
            ext = p.suffix.lower()
            name = p.name
            if ext in SKIP_BINARY_EXTS:
                continue
            if ext == ".docx":
                n = process_docx(p)
            elif ext == ".xlsx":
                n = process_xlsx(p)
            elif ext in PLAINTEXT_EXTS or name in NOEXT_NAMES or ext == "":
                n = process_text(p)
            else:
                # Unknown extension: try as text, skip on decode error
                n = process_text(p)
            if n > 0:
                files.append((p, n))
                total += n
    print(f"Total replacements this pass: {total}")
    print(f"Files changed: {len(files)}")
    for p, n in sorted(files, key=lambda x: -x[1])[:50]:
        print(f"  {n:>4}  {p}")


if __name__ == "__main__":
    main()
