# Copyright (c) 2026 ChatHealthy.ai LLC. All rights reserved.
# Licensed under the FindCare Evaluation License (FEL-1.0).
"""
ComplianceCrosswalkAgent — single-class compliance crosswalk agent.

Builds and maintains a master compliance control crosswalk that binds
NIST SP 800-53 Rev 5 + HITRUST CSF v11 controls to backlog requirements
and the source code blocks that implement them.

Output: brain/machine_artifacts/content/SecurityAuditControls.json
Review: brain/BusinessArtifacts/governance/compliance_crosswalk_review.docx

Public interface (per human directive 2026-04-19):
    agent = ComplianceCrosswalkAgent()
    agent.run()                                       # full pipeline
    agent.ComplianceCrossWalk(FileType.ALL)           # build catalog + classify
    agent.ComplianceCrosswalkReport()                 # write the .docx
    # Chained:
    agent.ComplianceCrossWalk(FileType.PYTHON).ComplianceCrosswalkReport()

Two-pass internal architecture:
- Pass 1 (deterministic): NIST OSCAL ingestion + HITRUST seed from brain/security.json
- Pass 2 (Anthropic SDK, Sonnet 4.6 with prompt caching): semantic file→control mapping
"""

from __future__ import annotations

import argparse
import json
import sys
import urllib.request
from dataclasses import dataclass
from datetime import datetime, timezone
from enum import Enum
from pathlib import Path
from typing import Iterable, Optional, Union

NIST_OSCAL_URL = (
    "https://raw.githubusercontent.com/usnistgov/oscal-content/main/"
    "nist.gov/SP800-53/rev5/json/NIST_SP-800-53_rev5_catalog.json"
)

REPO_ROOT = Path(__file__).resolve().parents[4]
BRAIN_DIR = REPO_ROOT / "brain" / "machine_artifacts" / "content"
CONTROLS_OUT = BRAIN_DIR / "SecurityAuditControls.json"
SECURITY_IN = BRAIN_DIR / "security.json"
ENV_FILE = REPO_ROOT / "Code" / ".env"
REVIEW_DOC_OUT = (REPO_ROOT / "brain" / "BusinessArtifacts" /
                  "governance" / "compliance_crosswalk_review.docx")

DEFAULT_SCAN_DIRS = [
    "Code/Shared", "Code/ConversationalUX", "Code/DataPipelines",
    "Code/evaluate_care", "Code/shared_services", "Code/deploy",
    "Website", ".github", ".claude",
]
SCAN_SKIP_NAMES = {"node_modules", ".venv", "__pycache__", ".git",
                   ".pytest_cache", "dist", "build", "out"}

DEFAULT_MODEL = "claude-sonnet-4-6"
MAX_FILE_BYTES = 200_000


# ── FileType enum (human-specified public surface) ────────────────


class FileType(str, Enum):
    """Which file extensions Pass 2 walks. ALL is the default."""
    ALL = "all"
    PYTHON = "py"
    YAML = "yml"   # also matches .yaml
    JSON = "json"
    HTML = "html"
    SHELL = "sh"
    BAT = "bat"
    POWERSHELL = "ps1"
    JAVASCRIPT = "js"
    TYPESCRIPT = "ts"
    TSX = "tsx"
    TF = "tf"
    CFG = "cfg"
    INI = "ini"

    def extensions(self) -> set[str]:
        if self is FileType.ALL:
            return {".py", ".yml", ".yaml", ".sh", ".bat", ".ps1", ".tf",
                    ".json", ".html", ".js", ".ts", ".tsx", ".cfg", ".ini"}
        if self is FileType.YAML:
            return {".yml", ".yaml"}
        return {f".{self.value}"}


# ── Internal data classes ────────────────────────────────────────


@dataclass
class _NistControl:
    id: str
    family: str
    title: str
    control_text: str


@dataclass
class _HitrustControl:
    id: str
    domain: str
    control: str
    how_seed: str = ""


# ── ComplianceCrosswalkAgent — single class ──────────────────────


class ComplianceCrosswalkAgent:
    """Single-class compliance crosswalk agent. See module docstring."""

    # Public-source NIST→HITRUST seed crosswalk (many-to-many; many gaps).
    # Real crosswalk needs a HITRUST license; this seed covers the controls
    # most likely to be implemented by code-level scanners and middleware.
    NIST_TO_HITRUST_SEED: dict[str, list[str]] = {
        "SC-8": ["09.s", "09.m"], "SC-8(1)": ["09.s"],
        "SC-13": ["06.d"], "SC-23": ["09.s"],
        "CM-2": ["10.h"], "CM-3": ["10.k"], "CM-3(2)": ["10.k"],
        "CM-7": ["07.a"],
        "SA-11": ["10.l"], "SA-11(1)": ["10.l"], "SA-15": ["10.l"],
        "SI-10": ["10.h"],
        "AC-3": ["05.a"], "AC-17": ["09.s"], "AC-17(2)": ["09.s"],
        "IA-3": ["05.a"], "AU-2": ["09.aa"], "CA-7": ["06.h"],
    }

    def __init__(self,
                 scan_dirs: Optional[list[str]] = None,
                 model: str = DEFAULT_MODEL,
                 api_key: Optional[str] = None,
                 controls_path: Path = CONTROLS_OUT,
                 review_path: Path = REVIEW_DOC_OUT,
                 nist_oscal_url: str = NIST_OSCAL_URL,
                 security_json: Path = SECURITY_IN):
        """
        scan_dirs:    repo-relative directories Pass 2 walks; default DEFAULT_SCAN_DIRS
        model:        Anthropic model id; default claude-sonnet-4-6
        api_key:      Anthropic_API_KEY; if None, loaded from Code/.env
        controls_path: where SecurityAuditControls.json is written
        review_path:  where the .docx review doc is written
        nist_oscal_url: source for the NIST 800-53 catalog (overridable for tests)
        security_json: source for the HITRUST seed
        """
        self.scan_dirs = scan_dirs or DEFAULT_SCAN_DIRS
        self.model = model
        self.api_key = api_key or self._load_api_key()
        self.controls_path = controls_path
        self.review_path = review_path
        self.nist_oscal_url = nist_oscal_url
        self.security_json = security_json
        # Internal state populated by ComplianceCrossWalk
        self._nist: list[_NistControl] = []
        self._hitrust: list[_HitrustControl] = []
        self._doc: Optional[dict] = None
        self._client = None
        # Pass 2 token accounting
        self.input_tokens = 0
        self.output_tokens = 0
        self.cache_read_tokens = 0
        self.cache_creation_tokens = 0

    # ── Public surface ───────────────────────────────────────────

    def run(self,
            file_types: Union[FileType, str] = FileType.ALL,
            sample: Optional[list[Path]] = None,
            max_files: Optional[int] = None,
            with_pass2: bool = True) -> Path:
        """Full pipeline: ComplianceCrossWalk(...) then ComplianceCrosswalkReport().
        Returns the path to the .docx review.
        """
        self.ComplianceCrossWalk(file_types=file_types,
                                 sample=sample,
                                 max_files=max_files,
                                 with_pass2=with_pass2)
        return self.ComplianceCrosswalkReport()

    def ComplianceCrossWalk(self,
                            file_types: Union[FileType, str] = FileType.ALL,
                            sample: Optional[list[Path]] = None,
                            max_files: Optional[int] = None,
                            with_pass2: bool = True) -> "ComplianceCrosswalkAgent":
        """Pass 1 (catalog ingestion) + optional Pass 2 (semantic classification).
        Returns self for chaining."""
        if isinstance(file_types, str):
            file_types = FileType(file_types)
        # Pass 1
        self._pass_1()
        # Pass 2
        if with_pass2:
            self._pass_2(file_types=file_types, sample=sample, max_files=max_files)
        return self

    def ComplianceCrosswalkReport(self,
                                  out_path: Optional[Path] = None) -> Path:
        """Generate the .docx review of the current crosswalk state. Returns the path."""
        target = out_path or self.review_path
        self._write_docx_report(target)
        return target

    # ── Pass 1 — deterministic catalog ingestion ─────────────────

    def _pass_1(self) -> dict:
        self._nist = self._load_nist_oscal()
        self._hitrust = self._load_hitrust_seed()
        doc = self._build_crosswalk(self._nist, self._hitrust)
        self.controls_path.parent.mkdir(parents=True, exist_ok=True)
        with open(self.controls_path, "w", encoding="utf-8") as f:
            json.dump(doc, f, indent=2)
        self._doc = doc
        return doc

    def _load_nist_oscal(self) -> list[_NistControl]:
        req = urllib.request.Request(self.nist_oscal_url,
                                     headers={"User-Agent": "Mozilla/5.0"})
        with urllib.request.urlopen(req, timeout=60) as r:
            oscal = json.loads(r.read().decode("utf-8"))
        catalog = oscal.get("catalog", oscal)
        out: list[_NistControl] = []
        for group in catalog.get("groups", []):
            family_id = (group.get("id") or "").upper()
            family_title = group.get("title", "")
            self._extract_oscal_controls(group.get("controls", []),
                                         family_id, family_title, out)
        return out

    def _extract_oscal_controls(self, controls_list, family_id, family_title, out):
        for c in controls_list:
            cid = (c.get("id") or "").upper()
            title = c.get("title", "")
            text = ""
            for part in c.get("parts", []):
                if part.get("name") == "statement":
                    text = self._part_to_text(part)
                    break
            out.append(_NistControl(id=cid,
                                    family=f"{family_id} {family_title}".strip(),
                                    title=title, control_text=text))
            if c.get("controls"):
                self._extract_oscal_controls(c["controls"],
                                             family_id, family_title, out)

    def _part_to_text(self, part) -> str:
        out: list[str] = []
        if part.get("prose"):
            out.append(part["prose"])
        for sub in part.get("parts", []):
            out.append(self._part_to_text(sub))
        return "\n".join(t for t in out if t)

    def _load_hitrust_seed(self) -> list[_HitrustControl]:
        if not self.security_json.exists():
            return []
        with open(self.security_json, encoding="utf-8") as f:
            sec = json.load(f)
        out: list[_HitrustControl] = []
        for rec in sec.get("records", []):
            hitrust = rec.get("compliance_mapping", {}).get("hitrust_csf", {})
            for domain in hitrust.get("domains", []):
                dname = domain.get("domain", "")
                for c in domain.get("controls_satisfied", []):
                    out.append(_HitrustControl(
                        id=c.get("id", ""), domain=dname,
                        control=c.get("control", ""),
                        how_seed=c.get("how", "")))
        return out

    def _build_crosswalk(self, nist, hitrust) -> dict:
        controls_out: list[dict] = []
        hitrust_by_id = {h.id: h for h in hitrust}
        cross_walked: set[str] = set()
        for c in nist:
            cw_ids = self.NIST_TO_HITRUST_SEED.get(c.id, [])
            frameworks = ["nist_800_53"]
            hitrust_block = None
            if cw_ids:
                frameworks.append("hitrust_csf")
                cross_walked.update(cw_ids)
                hitrust_block = {
                    "ids": cw_ids,
                    "details": [self._hitrust_to_dict(hitrust_by_id[h])
                                for h in cw_ids if h in hitrust_by_id],
                }
            controls_out.append({
                "control_id": c.id, "frameworks": frameworks,
                "nist_800_53": {"id": c.id, "family": c.family,
                                "title": c.title, "control_text": c.control_text},
                "hitrust_csf": hitrust_block,
                "implementations": [],
            })
        for h in hitrust:
            if h.id and h.id not in cross_walked:
                controls_out.append({
                    "control_id": f"HITRUST:{h.id}",
                    "frameworks": ["hitrust_csf"],
                    "nist_800_53": None,
                    "hitrust_csf": {"ids": [h.id],
                                    "details": [self._hitrust_to_dict(h)]},
                    "implementations": [],
                })
        return {
            "$schema": "https://dev.chathealthy.ai/schemas/dev/chathealthy-relaxed.schema.json",
            "collection": "SecurityAuditControls",
            "description": ("Master compliance control catalog and crosswalk: "
                            "NIST SP 800-53 Rev 5 + HITRUST CSF v11. Each "
                            "control entry binds to backlog requirements and "
                            "code blocks that implement it. Maintained by "
                            "ComplianceCrosswalkAgent."),
            "version": "1",
            "date": datetime.now(timezone.utc).strftime("%Y-%m-%d"),
            "frameworks": {
                "nist_800_53": {
                    "name": "NIST SP 800-53 Rev 5",
                    "source_url": self.nist_oscal_url,
                    "loaded_from": "OSCAL JSON catalog (open)",
                    "control_count": len(nist),
                },
                "hitrust_csf": {
                    "name": "HITRUST CSF v11",
                    "source": "Seeded from brain/security.json compliance_mapping.hitrust_csf",
                    "control_count": len(hitrust),
                    "gap": ("Full HITRUST CSF v11 catalog requires a HITRUST "
                            "license. Until then, only the controls already "
                            "mapped in security.json are present plus the "
                            "NIST→HITRUST seed crosswalk in this agent."),
                },
            },
            "crosswalk_seed_size": len(self.NIST_TO_HITRUST_SEED),
            "control_count": len(controls_out),
            "controls": controls_out,
        }

    @staticmethod
    def _hitrust_to_dict(h: _HitrustControl) -> dict:
        return {"id": h.id, "domain": h.domain,
                "control": h.control, "how_seed": h.how_seed}

    # ── Pass 2 — semantic classification (Anthropic SDK) ─────────

    def _ensure_client(self):
        if self._client is None:
            import anthropic
            if not self.api_key:
                raise RuntimeError("Anthropic_API_KEY not found in Code/.env")
            self._client = anthropic.Anthropic(api_key=self.api_key)

    def _build_catalog_context(self, doc: dict, max_controls: int = 250) -> str:
        lines = ["NIST 800-53 + HITRUST CSF Control Catalog "
                 "(id | title | one-line statement):"]
        for c in doc["controls"][:max_controls]:
            cid = c["control_id"]
            nist = c.get("nist_800_53") or {}
            title = nist.get("title") or ""
            text = (nist.get("control_text") or "").replace("\n", " ").strip()
            text = text[:200] + ("…" if len(text) > 200 else "")
            lines.append(f"{cid} | {title} | {text}")
        return "\n".join(lines)

    def _classify_file(self, filepath: Path, catalog_context: str) -> dict:
        try:
            content = filepath.read_text(encoding="utf-8", errors="replace")
        except Exception as e:
            return {"file": str(filepath), "error": f"read failed: {e}",
                    "controls": []}
        if len(content.encode("utf-8")) > MAX_FILE_BYTES:
            return {"file": str(filepath),
                    "error": f"file > {MAX_FILE_BYTES} bytes; flagged for human review",
                    "controls": []}
        self._ensure_client()
        system = [{
            "type": "text",
            "text": ("You are a compliance crosswalk classifier. Given a "
                     "control catalog and a source code file, decide which "
                     "controls (if any) the file substantively implements. "
                     "Be conservative: only emit a control when there is "
                     "concrete evidence in the code (not aspirational). "
                     "Output strict JSON only — no prose.\n\n" + catalog_context),
            "cache_control": {"type": "ephemeral"},
        }]
        user = (
            "Classify the following source file. Output JSON of shape:\n"
            "{\n"
            '  "file": "<filepath>",\n'
            '  "controls": [\n'
            '    {"control_id": "SC-8", "evidence": "<why>", '
            '"code_blocks": [{"start_line": int, "end_line": int, '
            '"summary": "<≤80 char summary>"}], '
            '"confidence": "high|medium|low"}\n'
            "  ]\n}\n\n"
            f"File path: {filepath}\nFile contents:\n```\n{content}\n```"
        )
        resp = self._client.messages.create(
            model=self.model, max_tokens=2000,
            system=system, messages=[{"role": "user", "content": user}],
        )
        usage = resp.usage
        self.input_tokens += usage.input_tokens
        self.output_tokens += usage.output_tokens
        self.cache_read_tokens += getattr(usage, "cache_read_input_tokens", 0) or 0
        self.cache_creation_tokens += getattr(usage, "cache_creation_input_tokens", 0) or 0
        text = (resp.content[0].text if resp.content else "").strip()
        if text.startswith("```"):
            text = text.strip("`")
            if text.startswith("json"):
                text = text[4:]
            text = text.strip()
        try:
            return json.loads(text)
        except json.JSONDecodeError:
            return {"file": str(filepath), "error": "non-JSON response",
                    "raw": text[:500], "controls": []}

    def _pass_2(self,
                file_types: FileType,
                sample: Optional[list[Path]] = None,
                max_files: Optional[int] = None) -> dict:
        if self._doc is None:
            with open(self.controls_path, encoding="utf-8") as f:
                self._doc = json.load(f)
        catalog_ctx = self._build_catalog_context(self._doc)
        files = list(sample) if sample else list(
            self._iter_scannable_files(REPO_ROOT,
                                       scan_dirs=self.scan_dirs,
                                       exts=file_types.extensions()))
        if max_files:
            files = files[:max_files]
        by_id = {c["control_id"]: c for c in self._doc["controls"]}
        results: list[dict] = []
        for fp in files:
            res = self._classify_file(fp, catalog_ctx)
            results.append(res)
            for entry in res.get("controls", []):
                cid = entry.get("control_id")
                if not cid or cid not in by_id:
                    continue
                by_id[cid]["implementations"].append({
                    "source_file": str(fp.relative_to(REPO_ROOT)).replace("\\", "/"),
                    "code_blocks": entry.get("code_blocks", []),
                    "evidence": entry.get("evidence", ""),
                    "confidence": entry.get("confidence", "medium"),
                    "classified_by": "agent",
                    "classifier_model": self.model,
                    "classified_at": datetime.now(timezone.utc).isoformat(),
                })
        with open(self.controls_path, "w", encoding="utf-8") as f:
            json.dump(self._doc, f, indent=2)
        return {
            "files_scanned": len(files),
            "files_with_findings": sum(1 for r in results if r.get("controls")),
            "input_tokens": self.input_tokens,
            "output_tokens": self.output_tokens,
            "cache_read_tokens": self.cache_read_tokens,
            "cache_creation_tokens": self.cache_creation_tokens,
            "model": self.model,
            "results": results,
        }

    @staticmethod
    def _iter_scannable_files(root: Path,
                              scan_dirs: Iterable[str],
                              exts: set[str]) -> Iterable[Path]:
        for sub in scan_dirs:
            d = root / sub
            if not d.exists():
                continue
            for p in d.rglob("*"):
                if not p.is_file():
                    continue
                if any(part in SCAN_SKIP_NAMES for part in p.parts):
                    continue
                if p.suffix.lower() in exts:
                    yield p

    # ── Report (.docx) — folded in from former _build_compliance_morning_report.py ──

    def _write_docx_report(self, out_path: Path) -> None:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        if not self.controls_path.exists():
            raise RuntimeError(
                f"{self.controls_path} not found — run ComplianceCrossWalk first")
        with open(self.controls_path, encoding="utf-8") as f:
            data = json.load(f)
        out_path.parent.mkdir(parents=True, exist_ok=True)
        doc = Document()
        title = doc.add_heading("Compliance Crosswalk Review", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        def _p(text, italic=False):
            p = doc.add_paragraph()
            r = p.add_run(text); r.font.size = Pt(11)
            if italic: r.italic = True
            return p

        def _bullet(text):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(text).font.size = Pt(11)
            return p

        def _kv_table(rows):
            t = doc.add_table(rows=len(rows), cols=2)
            t.style = "Light Grid Accent 1"
            for i, (k, v) in enumerate(rows):
                t.rows[i].cells[0].text = str(k)
                t.rows[i].cells[1].text = str(v)
            return t

        _p(f"Generated {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}.",
           italic=True)

        doc.add_heading("1. Catalog summary", level=1)
        _kv_table([
            ("controls.json path", str(self.controls_path)),
            ("file size", f"{self.controls_path.stat().st_size:,} bytes"),
            ("NIST 800-53 controls",
             data["frameworks"]["nist_800_53"]["control_count"]),
            ("HITRUST seed controls",
             data["frameworks"]["hitrust_csf"]["control_count"]),
            ("Total control entries (after crosswalk)", data["control_count"]),
            ("NIST→HITRUST crosswalk seed pairs", data["crosswalk_seed_size"]),
        ])

        doc.add_heading("2. Membership distribution", level=1)
        from collections import Counter
        mems = Counter(tuple(sorted(c["frameworks"])) for c in data["controls"])
        _kv_table([(", ".join(k), v) for k, v in mems.most_common()])

        doc.add_heading("3. Implementation coverage", level=1)
        with_impl = [c for c in data["controls"] if c["implementations"]]
        impl_count = sum(len(c["implementations"]) for c in with_impl)
        sources = {impl["source_file"]
                   for c in with_impl for impl in c["implementations"]}
        _kv_table([
            ("Controls with at least one implementation", len(with_impl)),
            ("Total implementation entries", impl_count),
            ("Distinct source files referenced", len(sources)),
        ])

        if with_impl:
            doc.add_heading("4. Top-20 controls by implementation count",
                            level=1)
            for c in sorted(with_impl,
                            key=lambda x: -len(x["implementations"]))[:20]:
                title_str = (c.get("nist_800_53") or {}).get("title", "") or "HITRUST"
                _bullet(f"{c['control_id']} ({len(c['implementations'])}): "
                        f"{title_str}")

        doc.add_heading("5. Limits to know", level=1)
        _bullet("HITRUST CSF v11 catalog gap — only seeded controls + the "
                "in-agent NIST→HITRUST seed crosswalk are populated.")
        _bullet("Classifier truncates the catalog to 250 controls per call — "
                "tail-of-catalog controls are not visible to the LLM per call.")
        _bullet("Files larger than 200 KB are flagged for human review and "
                "skipped by the LLM.")

        doc.save(out_path)

    # ── Misc helpers ─────────────────────────────────────────────

    @staticmethod
    def _load_api_key() -> Optional[str]:
        if not ENV_FILE.exists():
            return None
        with open(ENV_FILE, encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if not line or line.startswith("#") or "=" not in line:
                    continue
                k, _, v = line.partition("=")
                if k.strip() in ("Anthropic_API_KEY", "ANTHROPIC_API_KEY"):
                    return v.strip()
        return None


# ── CLI ──────────────────────────────────────────────────────────


def _cmd_run(args):
    agent = ComplianceCrosswalkAgent(model=args.model)
    sample = _resolve_sample_paths(args.sample)
    target = agent.run(file_types=FileType(args.file_types),
                       sample=sample, max_files=args.max_files,
                       with_pass2=not args.no_pass2)
    print(f"Wrote {agent.controls_path}")
    print(f"Wrote {target}")


def _cmd_crosswalk(args):
    agent = ComplianceCrosswalkAgent(model=args.model)
    sample = _resolve_sample_paths(args.sample)
    agent.ComplianceCrossWalk(file_types=FileType(args.file_types),
                              sample=sample, max_files=args.max_files,
                              with_pass2=not args.no_pass2)
    print(f"Wrote {agent.controls_path}")


def _cmd_report(args):
    agent = ComplianceCrosswalkAgent()
    out = agent.ComplianceCrosswalkReport()
    print(f"Wrote {out}")


def _cmd_status(_args):
    p = CONTROLS_OUT
    if not p.exists():
        print(f"{p} not found — run crosswalk first.")
        return
    with open(p, encoding="utf-8") as f:
        data = json.load(f)
    n = data["control_count"]
    with_impl = sum(1 for c in data["controls"] if c["implementations"])
    print(f"{p.name}: {n} controls; {with_impl} have implementations.")


def _resolve_sample_paths(items: Optional[list[str]]) -> Optional[list[Path]]:
    if not items:
        return None
    out: list[Path] = []
    for s in items:
        p = Path(s)
        if not p.is_absolute():
            p = REPO_ROOT / s
        if p.exists():
            out.append(p)
        else:
            print(f"WARN: sample file not found: {p}", file=sys.stderr)
    return out or None


def _add_common_pass2_args(p):
    p.add_argument("--file-types", default="all",
                   choices=[ft.value for ft in FileType])
    p.add_argument("--sample", nargs="*",
                   help="Specific files (relative to repo root)")
    p.add_argument("--max-files", type=int, default=None)
    p.add_argument("--no-pass2", action="store_true",
                   help="Pass 1 only — do not call the LLM")
    p.add_argument("--model", default=DEFAULT_MODEL)


def main():
    ap = argparse.ArgumentParser(description="Compliance Crosswalk Agent")
    sub = ap.add_subparsers(dest="cmd", required=True)

    p_run = sub.add_parser("run", help="Full pipeline: crosswalk + report")
    _add_common_pass2_args(p_run)
    p_run.set_defaults(func=_cmd_run)

    p_cw = sub.add_parser("crosswalk", help="Build catalog + classify (no report)")
    _add_common_pass2_args(p_cw)
    p_cw.set_defaults(func=_cmd_crosswalk)

    p_rep = sub.add_parser("report", help="Generate the .docx review only")
    p_rep.set_defaults(func=_cmd_report)

    p_st = sub.add_parser("status", help="Show SecurityAuditControls.json status")
    p_st.set_defaults(func=_cmd_status)

    args = ap.parse_args()
    args.func(args)


if __name__ == "__main__":
    main()
