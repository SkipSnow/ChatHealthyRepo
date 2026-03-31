# Copyright (c) 2026 Skip Snow. All rights reserved.
# Licensed under the FindCare Evaluation License (FEL-1.0).
import json, os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

with open("brain/machine_artifacts/document_type_json/biz_arch_investor_v2.json", "r", encoding="utf-8") as f:
    data = json.load(f)

DIAGRAM_DIR = "brain/BusinessArtifacts/diagrams"
doc = Document()

# Default style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# -- TITLE PAGE --
for _ in range(4):
    doc.add_paragraph("")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("ChatHealthy.ai")
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = RGBColor(13, 148, 136)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Find Care. Evaluate Care. Discuss Care.")
run.font.size = Pt(16)
run.font.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Business Architecture")
run.font.size = Pt(12)

doc.add_paragraph("")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run = p.add_run(data.get("vision", ""))
run.font.size = Pt(11)
run.font.italic = True

for _ in range(3):
    doc.add_paragraph("")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Copyright 2026 Skip Snow. All rights reserved.")
run.font.size = Pt(9)

doc.add_page_break()

# -- EXECUTIVE SUMMARY --
doc.add_heading("Executive Summary", level=1)
es = data.get("executive_summary", "")
text = "\n\n".join(es) if isinstance(es, list) else str(es)
for para in text.split("\n\n"):
    doc.add_paragraph(para)

# Value chain diagram
vc = os.path.join(DIAGRAM_DIR, "value_chain.png")
if os.path.exists(vc):
    doc.add_paragraph("")
    doc.add_picture(vc, width=Inches(5.5))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Fig. 1: The consumer orchestrates the care journey. Three components serve the decision, not a linear pipeline.")
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_page_break()

# -- PLATFORM ARCHITECTURE --
doc.add_heading("Platform Architecture", level=1)
doc.add_paragraph("ChatHealthy.ai is built on three business components. The consumer is at the center, orchestrating the relationship between all three. This is not a sequential journey -- it is a decision-making environment.")

comp_diag = os.path.join(DIAGRAM_DIR, "component_diagram.png")
if os.path.exists(comp_diag):
    doc.add_picture(comp_diag, width=Inches(5.5))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Fig. 2: Three components share a governance layer and data infrastructure.")
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_page_break()

# -- COMPONENTS --
fig_num = 3
for key in ["find_care", "evaluate_care", "discuss_care"]:
    comp = data.get("components", {}).get(key, {})
    if not comp:
        continue

    doc.add_heading(comp.get("name", key), level=1)
    if comp.get("tagline"):
        p = doc.add_paragraph()
        run = p.add_run(comp["tagline"])
        run.font.italic = True
        run.font.size = Pt(12)

    doc.add_paragraph(comp.get("description", ""))

    for section in ["capabilities", "features"]:
        items = comp.get(section, [])
        if items:
            doc.add_heading(section.title(), level=2)
            for item in items:
                if isinstance(item, dict):
                    doc.add_paragraph(
                        item.get("id", "") + ": " + item.get("name", item.get("feature", "")) + " -- " + item.get("description", ""),
                        style="List Bullet"
                    )
                else:
                    doc.add_paragraph(str(item), style="List Bullet")

    p = doc.add_paragraph()
    run = p.add_run("Pricing: " + comp.get("pricing", ""))
    run.font.bold = True

    # Inline diagrams per component
    diagram_map = {
        "find_care": ("find_care_diagram.png", "Chaos in, clarity out. The AI engine translates natural language into precise provider matches."),
        "evaluate_care": ("evaluate_care_diagram.png", "Public data flows through a graph network. The consumer tunes the algorithm. Patent pending."),
        "discuss_care": ("discuss_care_room.png", "Inside a Discuss Care session: humans and AI models collaborate with shared tools."),
    }
    if key in diagram_map:
        fname, caption = diagram_map[key]
        fpath = os.path.join(DIAGRAM_DIR, fname)
        if os.path.exists(fpath):
            doc.add_paragraph("")
            doc.add_picture(fpath, width=Inches(5.5))
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"Fig. {fig_num}: {caption}")
            run.font.size = Pt(8)
            run.font.italic = True
            run.font.color.rgb = RGBColor(100, 100, 100)
            fig_num += 1

    doc.add_page_break()

# -- REVENUE MODEL --
doc.add_heading("Revenue Model", level=1)
rm = data.get("revenue_model", {})
for s in rm.get("streams", []):
    if isinstance(s, dict):
        p = doc.add_paragraph()
        run = p.add_run(s.get("name", ""))
        run.font.bold = True
        doc.add_paragraph("  " + s.get("description", ""))
    else:
        doc.add_paragraph(str(s), style="List Bullet")

rev_diag = os.path.join(DIAGRAM_DIR, "revenue_flow.png")
if os.path.exists(rev_diag):
    doc.add_paragraph("")
    doc.add_picture(rev_diag, width=Inches(5.0))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Fig. {fig_num}: Two-sided platform. Consumers pay token premiums. AI vendors contribute capital.")
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = RGBColor(100, 100, 100)
    fig_num += 1

if rm.get("gov_001_compliance"):
    p = doc.add_paragraph()
    run = p.add_run(rm["gov_001_compliance"])
    run.font.italic = True
    run.font.size = Pt(9)

doc.add_page_break()

# -- COMPETITIVE MOAT --
doc.add_heading("Competitive Moat", level=1)
moat = data.get("moat", {})
for b in moat.get("barriers", []):
    if isinstance(b, dict):
        p = doc.add_paragraph()
        run = p.add_run(b.get("name", ""))
        run.font.bold = True
        doc.add_paragraph("  " + b.get("description", ""))
    else:
        doc.add_paragraph(str(b), style="List Bullet")

moat_diag = os.path.join(DIAGRAM_DIR, "moat_diagram.png")
if os.path.exists(moat_diag):
    doc.add_paragraph("")
    doc.add_picture(moat_diag, width=Inches(5.0))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Fig. {fig_num}: Five concentric barriers. IP Protection & Trade Secrets at the core.")
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = RGBColor(100, 100, 100)
    fig_num += 1

doc.add_page_break()

# -- COMPETITIVE LANDSCAPE --
doc.add_heading("Competitive Landscape", level=1)
cl = data.get("competitive_landscape", {})
if isinstance(cl, dict):
    if cl.get("positioning"):
        doc.add_paragraph(cl["positioning"])
    if cl.get("the_gap"):
        doc.add_paragraph(cl["the_gap"])
    if cl.get("incumbents"):
        for inc in cl["incumbents"]:
            p = doc.add_paragraph()
            run = p.add_run(inc.get("name", ""))
            run.font.bold = True
            doc.add_paragraph("  " + inc.get("problem", ""))

    logo_grid = os.path.join(DIAGRAM_DIR, "competitor_logos.png")
    if os.path.exists(logo_grid):
        doc.add_paragraph("")
        doc.add_picture(logo_grid, width=Inches(5.5))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Fig. {fig_num}: The healthcare navigation market. ChatHealthy.ai is the only platform that takes zero revenue from the care chain.")
        run.font.size = Pt(8)
        run.font.italic = True
        run.font.color.rgb = RGBColor(100, 100, 100)
        fig_num += 1

    if cl.get("our_position"):
        doc.add_paragraph(cl["our_position"])

doc.add_page_break()

# -- SECURITY --
doc.add_heading("Security Positioning", level=1)
sec = data.get("security_positioning", {})
if isinstance(sec, dict):
    if sec.get("headline"):
        p = doc.add_paragraph()
        run = p.add_run(sec["headline"])
        run.font.bold = True
        run.font.size = Pt(12)
    if sec.get("description"):
        doc.add_paragraph(sec["description"])
    for c in sec.get("certifications", sec.get("certifications_path", [])):
        if isinstance(c, dict):
            p = doc.add_paragraph()
            run = p.add_run(c.get("standard", ""))
            run.font.bold = True
            doc.add_paragraph("  " + c.get("purpose", c.get("why", "")))

comp_flow = os.path.join(DIAGRAM_DIR, "compliance_flow.png")
if os.path.exists(comp_flow):
    doc.add_paragraph("")
    doc.add_picture(comp_flow, width=Inches(5.5))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Fig. {fig_num}: Licensed provider enters the room -- mandatory disclosure to all participants.")
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = RGBColor(100, 100, 100)
    fig_num += 1

doc.add_page_break()

# -- INVESTMENT THESIS --
doc.add_heading("Investment Thesis", level=1)
thesis = data.get("investment_thesis", "")
if isinstance(thesis, dict):
    for v in thesis.values():
        doc.add_paragraph(str(v))
elif isinstance(thesis, list):
    for t in thesis:
        doc.add_paragraph(str(t))
else:
    doc.add_paragraph(str(thesis))

# -- TIMELINE --
doc.add_heading("Timeline", level=2)
for m in data.get("timeline", {}).get("milestones", data.get("timeline", {}).get("events", [])):
    if isinstance(m, dict):
        doc.add_paragraph(
            m.get("date", "") + "  --  " + m.get("milestone", m.get("event", "")),
            style="List Bullet"
        )

# -- FOOTER --
doc.add_paragraph("")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("ChatHealthy.ai Business Architecture  |  2026-03-31  |  Authored by Claude Code")
run.font.size = Pt(8)
run.font.italic = True
run.font.color.rgb = RGBColor(150, 150, 150)

doc.save("brain/BusinessArtifacts/biz_arch_investor_v2.docx")
print("DOCX generated")
