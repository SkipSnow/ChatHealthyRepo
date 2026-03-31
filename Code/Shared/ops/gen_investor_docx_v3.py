# Copyright (c) 2026 Skip Snow. All rights reserved.
# Licensed under the FindCare Evaluation License (FEL-1.0).
#
# gen_investor_docx_v3.py — Responds to Boss SGML markup, produces clean draft.
# Resolves all <Skip> tags. Inserts <Claude> tags for feedback.

import json, os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

with open("brain/machine_artifacts/document_type_json/biz_arch_investor_v2.json", "r", encoding="utf-8") as f:
    data = json.load(f)
with open("brain/machine_artifacts/document_type_json/component_discuss_care.json", "r", encoding="utf-8") as f:
    dc = json.load(f)

DIAGRAM_DIR = "brain/BusinessArtifacts/diagrams"
LOGO_PATH = "Website/images/logo.png"  # if exists
BUILD = "269"  # current git build

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    # Footer
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run(f"Copyright \u00a9 ChatHealthy.ai  |  Page ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(150, 150, 150)

# Default style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
style.paragraph_format.space_after = Pt(6)

# Heading styles
for level in [1, 2, 3]:
    hs = doc.styles[f"Heading {level}"]
    hs.font.color.rgb = RGBColor(13, 148, 136)  # teal
    hs.font.name = "Calibri"

def add_diagram(fname, caption, width=5.5):
    fpath = os.path.join(DIAGRAM_DIR, fname)
    if os.path.exists(fpath):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(fpath, width=Inches(width))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(caption)
        run.font.size = Pt(8)
        run.font.italic = True
        run.font.color.rgb = RGBColor(100, 100, 100)

# ═══════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════

# Logo - try to use the one Boss inserted
if os.path.exists(LOGO_PATH):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(LOGO_PATH, width=Inches(2.0))

for _ in range(3):
    doc.add_paragraph("")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("ChatHealthy.ai")
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = RGBColor(13, 148, 136)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Find Care. Evaluate Care. Discuss Care.")
run.font.size = Pt(18)
run.font.bold = True
run.font.color.rgb = RGBColor(245, 158, 11)  # gold

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Operational Business Architecture")
run.font.size = Pt(14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
today = datetime.now().strftime("%B %d, %Y")
run = p.add_run(f"{today}  |  Build {BUILD}")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph("")

# Vision - justified, italic
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run = p.add_run(data.get("vision", dc.get("vision", "")))
run.font.size = Pt(11)
run.font.italic = True

for _ in range(4):
    doc.add_paragraph("")

# Bottom credits
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Copyright 2026 Skip Snow. All rights reserved.")
run.font.size = Pt(9)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Authored by ")
run.font.size = Pt(9)
run = p.add_run("Claude Code")
run.font.size = Pt(9)
run.font.bold = True
run = p.add_run(" (Claude Opus 4.6, training cutoff May 2025)")
run.font.size = Pt(9)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("With assistance from ")
run.font.size = Pt(9)
run = p.add_run("GPT-5.3")
run.font.size = Pt(9)
run.font.bold = True
run = p.add_run(" (OpenAI, Enterprise Architect)")
run.font.size = Pt(9)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Supervised by ")
run.font.size = Pt(9)
run = p.add_run("Skip Snow")
run.font.size = Pt(9)
run.font.bold = True
run = p.add_run(", Founder & CEO")
run.font.size = Pt(9)

doc.add_page_break()

# ═══════════════════════════════════════════════
# EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════
doc.add_heading("Executive Summary", level=1)

es_paragraphs = [
    "Healthcare is the largest consumer market in the world, yet the individual remains the least empowered participant. Decisions are fragmented across search, opaque quality data, and isolated conversations with no coordination. The result is predictable: overpayment, underperformance, and preventable harm. ChatHealthy.ai exists because this gap is not a feature of the system -- it is the system's failure point.",
    "ChatHealthy.ai unifies the healthcare decision process into a single consumer-controlled environment: find care, evaluate it with real quality signals, and decide collaboratively with people and AI. This is not another point solution. It is the decision layer sitting above the entire healthcare system, where choices are actually made. Once this layer exists, control shifts from institutions to individuals.",
    "The business model follows the control point. We own the decision room. Frontier AI vendors contribute capital to launch and grow in exchange for equity and future token sales. As healthcare decisions become AI-assisted and collaborative by default, ChatHealthy.ai becomes the gateway through which those decisions flow. This is not a feature set. It is inevitable infrastructure.",
]
for para_text in es_paragraphs:
    doc.add_paragraph(para_text)

add_diagram("value_chain.png",
    "Fig. 1: The consumer orchestrates the care journey. Three components serve the decision, not a linear pipeline.", 5.0)

# ═══════════════════════════════════════════════
# PLATFORM ARCHITECTURE
# ═══════════════════════════════════════════════
doc.add_heading("Platform Architecture", level=1)
doc.add_paragraph("ChatHealthy.ai is built on three business components. The consumer is at the center, orchestrating the relationship between all three. This is not a sequential journey -- it is a decision-making environment.")

add_diagram("component_diagram.png",
    "Fig. 2: Three components share a governance layer and data infrastructure.", 5.5)

# ═══════════════════════════════════════════════
# FIND CARE
# ═══════════════════════════════════════════════
doc.add_heading("Find Care", level=1)
p = doc.add_paragraph()
run = p.add_run("The entry point to the healthcare system, redesigned for clarity and control.")
run.font.italic = True

doc.add_paragraph("Users discover providers, specialties, and care pathways without friction or bias.")

doc.add_heading("Capabilities", level=2)
for cap in ["Provider and specialty discovery", "Navigation of care pathways",
            "Geographic and access-based filtering", "Integrated entry into evaluation and discussion"]:
    doc.add_paragraph(cap, style="List Bullet")

p = doc.add_paragraph()
run = p.add_run("Pricing: Free forever")
run.font.bold = True

add_diagram("find_care_diagram.png",
    "Fig. 3: Chaos in, clarity out. The AI engine translates natural language into precise provider matches.", 5.0)

# ═══════════════════════════════════════════════
# EVALUATE CARE
# ═══════════════════════════════════════════════
doc.add_heading("Evaluate Care", level=1)
p = doc.add_paragraph()
run = p.add_run("Transforms fragmented public and clinical data into clear, consumer-comprehensible quality signals using proprietary AI synthesis.")
run.font.italic = True

doc.add_heading("Capabilities", level=2)
for cap in ["Clinical trials discovery and matching", "Provider credential analysis",
            "Quality signal aggregation and interpretation", "Comparative provider evaluation"]:
    doc.add_paragraph(cap, style="List Bullet")

p = doc.add_paragraph()
run = p.add_run("Pricing: Free initially, with premium paid tiers for advanced insights and deeper analysis")
run.font.bold = True

add_diagram("evaluate_care_diagram.png",
    "Fig. 4: Public data flows through a graph network. The consumer tunes the algorithm. Patent pending.", 5.5)

# ═══════════════════════════════════════════════
# DISCUSS CARE
# ═══════════════════════════════════════════════
doc.add_heading("Discuss Care", level=1)
p = doc.add_paragraph()
run = p.add_run("A real-time, multi-participant decision environment where individuals bring family, caregivers, and multiple AI models into a single structured conversation.")
run.font.italic = True

doc.add_heading("Capabilities", level=2)
for cap in ["Multi-human, multi-AI shared sessions",
            "Invite-based collaboration with no account friction",
            "AI model selection and competition within the same session",
            "Shared tools across participants (search, evaluation, trials)",
            "Persistent sessions with exportable transcripts",
            "Shared research bot for autonomous analysis",
            "Consent-based participation and governance controls"]:
    doc.add_paragraph(cap, style="List Bullet")

doc.add_heading("Features", level=2)
# All 15 features from the component doc
for feat in dc.get("features", []):
    doc.add_paragraph(
        f"{feat['id']}: {feat.get('feature', feat.get('name', ''))} -- {feat.get('description', '')}",
        style="List Bullet"
    )

p = doc.add_paragraph()
run = p.add_run("Pricing: Premium via AI token usage markup; users pay only when they choose advanced AI participation")
run.font.bold = True

add_diagram("discuss_care_room.png",
    "Fig. 5: Inside a Discuss Care session: humans and AI models collaborate with shared tools.", 5.0)

# ═══════════════════════════════════════════════
# REVENUE MODEL
# ═══════════════════════════════════════════════
doc.add_heading("Revenue Model", level=1)
for stream in ["AI token markup on premium model usage within Discuss Care",
               "AI vendor access fees to participate in decision sessions",
               "Future premium subscriptions for advanced evaluation insights"]:
    doc.add_paragraph(stream, style="List Bullet")

add_diagram("revenue_flow.png",
    "Fig. 6: Two-sided platform. Consumers pay token premiums. AI vendors contribute capital.", 4.5)

# ═══════════════════════════════════════════════
# COMPETITIVE MOAT
# ═══════════════════════════════════════════════
doc.add_heading("Competitive Moat", level=1)
moat_items = [
    ("Regulatory Willingness", "AI vendors avoid FDA classification and HIPAA compliance. They need someone else to own the healthcare room."),
    ("Governance Framework", "Multi-AI safety, consent, floor control, moderation. Nobody else has this."),
    ("Data + Network", "Provider, trial, and quality signals compound over time. Takes years to build."),
    ("Consumer Trust", "Years to build, seconds to lose. Trust is the ultimate barrier."),
    ("Patent", "LLM-based synthesis of provider quality signals. Method patent pending."),
    ("IP Protection & Trade Secrets", "Proprietary algorithms, data processing methods, and governance frameworks."),
]
for title, desc in moat_items:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.bold = True
    doc.add_paragraph(f"  {desc}")

add_diagram("moat_diagram.png",
    "Fig. 7: Five concentric barriers. IP Protection & Trade Secrets at the core.", 4.5)

# ═══════════════════════════════════════════════
# COMPETITIVE LANDSCAPE
# ═══════════════════════════════════════════════
doc.add_heading("Competitive Landscape", level=1)
doc.add_paragraph("The healthcare navigation market is dominated by platforms that earn revenue from the providers they claim to evaluate. ChatHealthy.ai is the only platform built entirely on the consumer side of the table.")
doc.add_paragraph("Every incumbent earns money from the care chain -- providers, hospitals, or payers. This creates a structural conflict of interest. The consumer cannot trust a platform that is paid by the people it is supposed to evaluate.")

incumbents = [
    ("Yelp Health", "Earns revenue from sponsored listings. Providers who pay appear first. Consumer trust is the product being sold."),
    ("ZocDoc", "Charges providers for patient bookings. The platform is a referral engine disguised as a search tool."),
    ("Healthgrades", "Revenue from provider advertising and hospital partnerships. Ratings influenced by who pays."),
    ("Angi", "Merged with HomeAdvisor. Lead generation model. Providers pay for leads. Consumer is the product."),
]
for name, problem in incumbents:
    p = doc.add_paragraph()
    run = p.add_run(name)
    run.font.bold = True
    doc.add_paragraph(f"  {problem}")

add_diagram("competitor_logos.png",
    "Fig. 8: The healthcare navigation market. ChatHealthy.ai is the only platform that takes zero revenue from the care chain.", 5.0)

doc.add_paragraph("ChatHealthy.ai will never take money from providers, hospitals, payers, pharmaceutical companies, PBMs, or any entity in the care delivery or payment chain. This is not a guideline. It is governance policy GOV-001, permanent and non-negotiable. Our revenue comes exclusively from consumers (token premiums) and technology partners (AI vendor capital). This is our moat. No incumbent can match it without destroying their business model.")

# ═══════════════════════════════════════════════
# SECURITY
# ═══════════════════════════════════════════════
doc.add_heading("Security Positioning", level=1)
p = doc.add_paragraph()
run = p.add_run("We are not required to protect your data. We choose to. And we do it better than the ones who are required to.")
run.font.bold = True
run.font.size = Pt(12)

doc.add_paragraph("ChatHealthy.ai is not a HIPAA covered entity. We voluntarily adopt enterprise-grade security standards because our business depends on consumer trust, not regulatory compulsion.")

certs = [
    ("HITRUST CSF", "Healthcare-specific security framework. Gold standard for health tech."),
    ("SOC 2 Type II", "Independent audit of security controls. Enterprise table stakes."),
    ("NIST SP 800-171", "DOD standard. 110 security controls."),
    ("Annual External Penetration Test", "Third-party security audit. Published results. We publish what hospitals hide."),
]
for standard, purpose in certs:
    p = doc.add_paragraph()
    run = p.add_run(standard)
    run.font.bold = True
    doc.add_paragraph(f"  {purpose}")

add_diagram("compliance_flow.png",
    "Fig. 9: Licensed provider enters the room -- mandatory disclosure to all participants.", 5.0)

# ═══════════════════════════════════════════════
# INVESTMENT THESIS
# ═══════════════════════════════════════════════
doc.add_heading("Investment Thesis", level=1)
doc.add_paragraph("ChatHealthy.ai sits at the intersection of two inevitable shifts: consumer control of healthcare decisions and AI participation in those decisions. The company captures value by owning the environment where both converge. Investors are backing a platform that becomes the default interface for navigating a $4 trillion market. Strategic value is immediate for AI companies seeking distribution, and long-term for any entity that benefits from influencing healthcare decisions without regulatory exposure.")

# ═══════════════════════════════════════════════
# TIMELINE
# ═══════════════════════════════════════════════
doc.add_heading("Timeline", level=2)
milestones = [
    ("April 30, 2026", "Alpha Launch (Find Care + Evaluate Care)"),
    ("May 2026", "Discuss Care Beta Sprint 1"),
    ("July 31, 2026", "Seed Funding Close"),
    ("August 25, 2026", "Production Launch"),
    ("September 20, 2026", "Beta Launch"),
    ("September 30, 2026", "iPhone App Launch"),
]
for date, event in milestones:
    p = doc.add_paragraph()
    run = p.add_run(date)
    run.font.bold = True
    p.add_run(f"  --  {event}")

# Save
output_path = "brain/BusinessArtifacts/ChatHealthyOperationalBusinessArchitecture.docx"
doc.save(output_path)
print(f"DOCX saved: {output_path}")
