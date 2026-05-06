# Copyright (c) 2026 ChatHealthy.ai LLC. All rights reserved.
# Generate ops-manager-design-EPIC3-V1.docx from design content.

import os

from docx import Document
from docx.shared import Pt, RGBColor

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def h1(text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x0B, 0x7A, 0x75)


def h2(text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x0B, 0x7A, 0x75)


def h3(text):
    doc.add_heading(text, level=3)


def para(text):
    doc.add_paragraph(text)


def add_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for r_idx, row in enumerate(rows):
        for c_idx, cell in enumerate(row):
            t.rows[r_idx + 1].cells[c_idx].text = str(cell)


# ── TITLE ──
h1("Pipeline Operations Manager")
para("EPIC-3 \u2014 Infrastructure Lifecycle Management")
para("Design Document | V1 (dev) | April 2, 2026 | Build 389")
para("Copyright 2026 Skip Snow. All rights reserved.")

# ── OVERVIEW ──
h2("Overview")
para(
    "The Pipeline Operations Manager is an AI agent that manages infrastructure "
    "lifecycle for ChatHealthy.ai data pipelines. It manages clusters, not data. "
    "It reads errors, not records. It escalates to the right person, not fixes code."
)
para("The agent uses tools. Tools don\u2019t think. The agent thinks.")

# ── OWNERSHIP BOUNDARY ──
h2("Ownership Boundary \u2014 Ops Manager vs Dev Manager")
para(
    "The Ops Manager and Dev Manager are separate concerns with a strict boundary. "
    "The Ops Manager never reads pipeline code. The Dev Manager never touches "
    "infrastructure. They communicate through one handoff point: an evidence package."
)

h3("Ops Manager")
add_table(
    ["Category", "Details"],
    [
        ["OWNS", "Cluster lifecycle (wake/pause/scale), Reservation array, Timer (overdue/orphan), Error triage (infra vs logic), Self-repair (infra-only)"],
        ["TOOLS", "ClusterTool (Atlas API), AlertTool (Pushover+email), TriageTool (GPT-4.1-mini), CostTool (future)"],
        ["CANNOT", "Read pipeline source code, Import pipeline modules, Fix logic errors"],
        ["TIMER", "Every 5 min, infrastructure checks only, no task execution"],
    ],
)

h3("Dev Manager")
add_table(
    ["Category", "Details"],
    [
        ["OWNS", "Pipeline business logic, Data schemas/transforms, Exception handling, Code fixes, Tests"],
        ["TOOLS", "Code read/edit/write, Test runner, Git, Deploy (dev only)"],
        ["CANNOT", "Wake/pause/scale clusters, Modify reservations, Send operational alerts"],
        ["ACTIVATED BY", "Evidence package from Ops Manager only"],
    ],
)

# ── EVIDENCE PACKAGE ──
h3("Evidence Package \u2014 The Only Handoff")
add_table(
    ["Field", "Source", "Purpose"],
    [
        ["error_code", "Failing tool\u2019s ToolResult", "What broke"],
        ["error_message", "Full exception text", "Exact error"],
        ["classification", "TriageTool (GPT-4.1-mini)", "Why Ops thinks it\u2019s a logic error"],
        ["confidence", "TriageTool", "How sure the classifier is"],
        ["reasoning", "TriageTool", "Evidence chain for the classification"],
        ["cluster_state", "ClusterTool", "Infra was healthy (rules out infra)"],
        ["recent_events", "Brain log (JSON, max 10)", "What happened leading up to the error"],
        ["job_id / requester", "Reservation array", "Which pipeline task failed"],
        ["first_seen_at", "Brain log", "New regression vs known flaky"],
        ["occurrence_count", "Brain log", "How many times this error occurred"],
    ],
)

# ── NORMAL FLOW ──
h2("Normal Flow \u2014 Reservation Lifecycle")
para(
    "Pipeline Task Starts \u2192 WakeCluster (add to reservation array, send Atlas resume) "
    "\u2192 ClusterStatus (poll every 15s) \u2192 cluster IDLE \u2192 Do Work "
    "\u2192 Release (remove from array, pause if empty) \u2192 Done"
)

h3("Reservation Duration Parameters")
add_table(
    ["Field", "Type", "Purpose"],
    [
        ["expected_min_minutes", "int", "Job faster than this = WARNING, Dev Manager verifies output"],
        ["expected_max_minutes", "int", "Job longer than this = OVERDUE, Notify human"],
    ],
)

h3("Duration Check Logic (timer, every 5 min)")
add_table(
    ["Condition", "Classification", "Action"],
    [
        ["actual < min", "WARNING: suspiciously fast", "Notify Dev Manager \u2014 verify output"],
        ["min <= actual <= max", "Normal", "No action"],
        ["actual > max", "OVERDUE", "Notify human (cooldown: 60 min/job)"],
    ],
)

h3("Concurrent Reservations Example")
add_table(
    ["Reservation", "Requester", "Min", "Max", "Status"],
    [
        ["copy_47_states", "CopyToFrontEnd", "30 min", "120 min", "active"],
        ["county_enrich", "CountyEnrichment", "15 min", "60 min", "active"],
        ["embed_va", "EmbeddingWorker", "5 min", "30 min", "active"],
    ],
)

# ── ERROR FLOW ──
h2("Error Flow \u2014 Diagnose Loop with Tool Calls")
para("1. Error Event")
para("2. NOTIFY human (email: error detected, investigating)")
para("3. GPT-4.1-mini: DIAGNOSE (call tools, classify: infra or ETL code?) \u2014 max N iterations (configurable)")
para("4. Diamond: Infrastructure? ETL code? Can\u2019t determine?")

h3("Left Branch \u2014 Infrastructure (Ops Manager)")
para(
    "GPT-4.1-mini picks repair \u2192 Execute repair tool \u2192 Read logs + verify \u2192 "
    "Fixed? Yes: NOTIFY human, END. No: retry < M? Yes: loop back. No: ESCALATE human, END."
)

h3("Right Branch \u2014 ETL Code (Dev Manager)")
para(
    "Evidence Package received \u2192 Investigate ETL code \u2192 Fix + Run Tests \u2192 "
    "Pass? Yes: NOTIFY human, END. No: retry < M? Yes: loop back. No: ESCALATE human, END."
)

h3("Bottom \u2014 Can\u2019t Determine")
para("ESCALATE human (can\u2019t classify after N tries) \u2192 END")
para("N = diagnose iterations (configurable). M = repair/fix attempts (configurable, independent per path).")

# ── RULE-BASED PRECHECK ──
h2("Error Classification \u2014 Rule-Based Precheck")
para("Before calling GPT-4.1-mini, deterministic rules are applied. If a rule matches, classification is instant and free.")
add_table(
    ["Signal", "Auto-Classification", "Rationale"],
    [
        ["ConnectionFailure, ServerSelectionTimeoutError", "INFRASTRUCTURE", "Cluster unreachable"],
        ["AuthenticationFailed, TLS error", "INFRASTRUCTURE", "Credential or network layer"],
        ["Cluster state = PAUSED or REPAIRING", "INFRASTRUCTURE", "Cluster not ready"],
        ["KeyError, TypeError, ValidationError in pipeline code", "PIPELINE", "Schema or data format"],
        ["IndexError, ValueError in Code/DataPipelines/", "PIPELINE", "ETL logic error"],
        ["None of the above", "\u2192 LLM TRIAGE", "Ambiguous \u2014 send to GPT-4.1-mini"],
    ],
)

# ── LLM TRIAGE ──
h2("Error Classification \u2014 LLM Triage (when rules inconclusive)")
para("GPT-4.1-mini classifies errors using the full error context.")
add_table(
    ["Field", "Value"],
    [
        ["Model", "gpt-4.1-mini (cost-optimized, proven on safety classifier)"],
        ["Role", "Error triage agent for pipeline infrastructure"],
        ["Output", "JSON: category, confidence, reasoning, self_repair_possible, repair_action"],
        ["Categories", "INFRASTRUCTURE, PIPELINE, UNKNOWN"],
    ],
)

h3("Allowed Repair Actions (config-driven allowlist)")
add_table(
    ["Action", "What It Does", "Risk"],
    [
        ["restart_cluster", "Resume paused cluster via Atlas API", "Low \u2014 idempotent"],
        ["reconnect", "Drop and re-establish connection pool", "Low \u2014 brief disruption"],
        ["scale_up", "Increase cluster tier (within cap)", "Medium \u2014 cost increase"],
        ["flush_connection_pool", "Clear stale connections", "Low \u2014 brief disruption"],
    ],
)

h3("Repair Loop Guardrail \u2014 Evidence Must Change")
para(
    "After each repair attempt, compare new evidence to previous. "
    "If unchanged, skip remaining retries and escalate immediately. Prevents churn."
)

# ── NOTIFICATION ──
h2("Notification Model")
para("Notify \u2260 Escalate. Notify is information. Escalate is \u201cI need your help.\u201d")
add_table(
    ["Event", "Recipient", "Type", "Content"],
    [
        ["Error detected", "human (configurable)", "Notify", "What happened, what we\u2019re doing"],
        ["Self-repaired", "human", "Notify", "What broke, what fixed it"],
        ["Sent to Dev Manager", "Dev Manager", "Action", "Error details, please resolve"],
        ["Dev resolved", "human", "Notify", "Dev fixed it, here\u2019s what they did"],
        ["Nobody can fix", "human", "Escalate", "We need your help"],
        ["Job overdue", "human", "Notify", "Job X running longer than expected"],
    ],
)

h3("Notification Cooldowns")
add_table(
    ["Event", "Cooldown", "Rule"],
    [
        ["Job overdue", "60 min per job", "One notification per job per hour"],
        ["Cluster stuck", "30 min per cluster", "One per cluster per 30 min"],
        ["Error detected", "None", "Always notify immediately"],
        ["Self-repaired / Dev resolved", "None", "Always immediately (good news)"],
        ["Escalation", "None", "Always escalate immediately"],
    ],
)

h3("Configurable Recipients")
add_table(
    ["Role", "Current", "Future"],
    [
        ["human", "skip.snow@gmail.com", "Same until COO hired"],
        ["Ops escalation", "skip.snow@gmail.com", "COO when hired"],
        ["Dev escalation", "Claude (via Brain)", "Dev lead when hired"],
    ],
)

# ── TIMER ──
h2("Timer \u2014 Infrastructure Only (every 5 min)")
add_table(
    ["Check", "Action"],
    [
        ["Reservation overdue?", "Notify human (cooldown: 60 min per job)"],
        ["Cluster running, zero reservations?", "Pause cluster"],
        ["Cluster stuck (non-IDLE > 30 min)?", "Notify human (cooldown: 30 min per cluster)"],
        ["NO task execution. NO pipeline imports. NO data access.", ""],
    ],
)

# ── AUDIT TRAIL ──
h2("Audit Trail \u2014 Append-Only (V1 Minimal)")
para("Every triage decision, repair action, and escalation is logged as an immutable record. No updates, no deletes.")
add_table(
    ["Event Type", "Fields Logged"],
    [
        ["Triage decision", "timestamp, error_code, classification, confidence, source (rule/LLM), reasoning"],
        ["Repair action", "timestamp, action, attempt_number, outcome, cluster_state_before/after"],
        ["Escalation", "timestamp, recipient, reason, full EvidencePackage"],
        ["Notification", "timestamp, recipient, event_type, content_summary"],
    ],
)

# ── COMPONENTS ──
h2("Components")
add_table(
    ["Component", "Location", "Responsibility"],
    [
        ["BaseAgent", "Code/Shared/agent_framework/", "Abstract agent with tool registry"],
        ["BaseTool", "Code/Shared/agent_framework/", "Tool interface: execute, repair"],
        ["ToolRegistry", "Code/Shared/agent_framework/", "Allowlist enforcement (GOV-004)"],
        ["OpsManagerAgent", "Code/DataPipelines/ops_manager/", "Pipeline infrastructure agent"],
        ["ClusterTool", "Code/DataPipelines/ops_manager/", "Wake, pause, status, reserve, release"],
        ["AlertTool", "Code/DataPipelines/ops_manager/", "Email (SparkPost) + Pushover"],
        ["TriageTool", "Code/DataPipelines/ops_manager/", "LLM error classification (GPT-4.1-mini)"],
        ["Config", "Brain pseudo-collection", "Recipients, thresholds, channels"],
    ],
)

# ── VERSIONING ──
h2("Versioning")
add_table(
    ["Version", "Capability"],
    [
        ["v0.1", "Rule-based. ClusterTool + AlertTool only. Reservation array. Timer."],
        ["v0.2", "AI triage. GPT-4.1-mini classifies errors. Self-repair. Notifications."],
        ["v0.3", "Multi-agent. Ops Manager + Research Agent + DB Agent coordinating."],
    ],
)

# ── AUDIT RESPONSE ──
h2("External Audit Response")
para(
    "Audit: ExternalAuditOfPipelineOpsManagerEPIC3V0.1.docx \u2014 GPT (Enterprise Architect), "
    "April 2, 2026. Overall opinion: \u201cThe design is sound and implementable for an "
    "infra-only agent.\u201d Signed off for V1 dev."
)
add_table(
    ["Finding", "Title", "Accepted/Rejected", "V1 or Deferred", "Feature/Story"],
    [
        ["F1", "Boundary clarity \u2014 enforce in CI", "Accepted", "Deferred (v0.2)", "OPS-BOUNDARY"],
        ["F2", "EvidencePackage schema hardening", "Accepted", "Implemented (V1)", "OPS-EVIDENCE / S1"],
        ["F3", "LLM triage guardrails", "Accepted", "Implemented (V1)", "OPS-TRIAGE / S1"],
        ["F4", "Self-repair loop safety limits", "Accepted", "Implemented (V1)", "OPS-REPAIR / S1"],
        ["F5", "Timer per-pipeline thresholds", "Risk Accepted (dev)", "Deferred (v0.2)", "OPS-TIMER"],
        ["F6", "Notification cooldowns", "Accepted", "Implemented (V1)", "OPS-NOTIFY / S1"],
        ["F7", "Feature flags per environment", "Accepted", "Deferred (v0.2)", "OPS-FLAGS"],
        ["F8", "Immutable audit trail", "Accepted", "Implemented (V1 minimal)", "OPS-AUDIT / S1"],
    ],
)

h3("Summary")
add_table(
    ["Status", "Count", "Findings"],
    [
        ["Accepted + Implemented in V1", "5", "F2, F3, F4, F6, F8 (minimal)"],
        ["Risk Accepted for dev", "1", "F5 \u2014 defaults only"],
        ["Accepted + Deferred to v0.2", "2", "F1, F7"],
        ["Rejected", "0", "\u2014"],
    ],
)

# ── UAT TEST SCENARIOS ──
h2("UAT Test Scenarios \u2014 Failure Path Validation")

h3("Infrastructure Failure Path")
add_table(
    ["#", "Scenario", "How to Trigger", "Expected Behavior"],
    [
        ["UAT-I1", "Database killed mid-job", "Pause Atlas cluster manually", "Rule-based \u2192 INFRA. Repair loop. Notifies human."],
        ["UAT-I2", "Connection timeout", "Rotate Atlas credentials", "Auth error \u2192 INFRA. Reconnect. If unchanged \u2192 early escalation."],
        ["UAT-I3", "Cluster stuck non-IDLE", "Resize then start job", "Timer detects stuck > 30 min. Notifies human."],
        ["UAT-I4", "Orphan cluster", "Kill Azure Function mid-job", "Timer: zero reservations \u2192 pause cluster."],
        ["UAT-I5", "Repair exhaustion", "Pause + block Atlas API", "M attempts fail, evidence unchanged \u2192 escalate."],
    ],
)

h3("ETL Code / Logic Failure Path")
add_table(
    ["#", "Scenario", "How to Trigger", "Expected Behavior"],
    [
        ["UAT-L1", "Collection missing", "Drop target collection", "KeyError \u2192 PIPELINE. Evidence to Dev Manager."],
        ["UAT-L2", "Collection empty", "Clear source documents", "Zero records \u2192 PIPELINE. Evidence to Dev Manager."],
        ["UAT-L3", "Schema mismatch", "Add required field, skip source update", "ValidationError \u2192 PIPELINE. Evidence to Dev."],
        ["UAT-L4", "Recursive loop", "Inject infinite retry", "Exceeds max \u2192 OVERDUE. Notifies human."],
        ["UAT-L5", "Job too fast", "Run against empty dataset", "Duration < min \u2192 WARNING. Dev verifies output."],
    ],
)

h3("Ambiguous / Edge Cases")
add_table(
    ["#", "Scenario", "How to Trigger", "Expected Behavior"],
    [
        ["UAT-A1", "Unclassifiable error", "Custom exception", "Rules inconclusive \u2192 GPT-4.1-mini."],
        ["UAT-A2", "Mixed infra + logic", "Drop collection AND pause cluster", "Infra first. Then second triage as PIPELINE."],
        ["UAT-A3", "Notification flood", "10+ errors rapidly", "Cooldowns prevent spam."],
        ["UAT-A4", "Concurrent failure", "3 jobs, 1 fails", "Only failed job triaged. Others continue."],
    ],
)

h3("Audit Trail Validation")
add_table(
    ["#", "Scenario", "Validation"],
    [
        ["UAT-AT1", "After any failure", "Triage, repair, escalation records exist. No updates/deletes."],
        ["UAT-AT2", "After UAT-L5 (too-fast)", "Warning record with actual/expected duration."],
    ],
)

OUT = os.path.join(os.environ['CHATHEALTHY_PROJECT_ROOT'], "brain", "BusinessArtifacts", "ops-manager-design-EPIC3-V1.docx")
doc.save(OUT)
print(f"Saved: {OUT}")
