"""Render the prompt factory design JSON as a Word doc with UML diagrams."""
import json
import os
import io
import plantuml
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

REPO_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", "..", ".."))


def render_uml(puml_text):
    """Render PlantUML text to PNG bytes."""
    p = plantuml.PlantUML(url="https://www.plantuml.com/plantuml/img/")
    return p.processes(puml_text)


def build_class_puml(design):
    """Build PlantUML class diagram from design JSON — square layout."""
    uml = design.get("uml_class_diagram", {})

    def fmt_methods(methods):
        result = []
        for m in methods:
            params = ", ".join(f"{p['name']}: {p['type']}" for p in m.get("params", []))
            result.append(f"  +{m['name']}({params}): {m.get('return_type', 'void')}")
        return "\n".join(result)

    lines = [
        "@startuml",
        "skinparam classAttributeIconSize 0",
        "skinparam shadowing false",
        "skinparam defaultFontSize 9",
        "skinparam padding 1",
        "skinparam nodesep 8",
        "skinparam ranksep 15",
        "",
    ]

    # Interface at the top
    for iface in uml.get("interfaces", []):
        lines.append(f"interface {iface['name']} {{")
        lines.append(fmt_methods(iface.get("methods", [])))
        lines.append("}\n")

    # Abstract base below interface
    for ac in uml.get("abstract_classes", []):
        lines.append(f"abstract class {ac['name']} {{")
        lines.append(fmt_methods(ac.get("methods", [])))
        lines.append("}")
        for impl in ac.get("implements", []):
            lines.append(f"{ac['name']} ..|> {impl}")
        lines.append("")

    # Epic subclasses — group together for square layout
    classes = uml.get("classes", [])
    if classes:
        lines.append("together {")
        for cls in classes:
            lines.append(f"  class {cls['name']} {{")
            for attr in cls.get("attributes", []):
                lines.append(f"    -{attr.get('name', '?')}: {attr.get('type', '?')}")
            lines.append(fmt_methods(cls.get("methods", [])))
            lines.append("  }")
        lines.append("}")
        lines.append("")
        for cls in classes:
            for inh in cls.get("inherits", []):
                lines.append(f"{cls['name']} --|> {inh}")

    # Middleware — group together below
    middleware = uml.get("middleware", [])
    if middleware:
        lines.append("")
        lines.append("together {")
        for mw in middleware:
            lines.append(f"  class {mw['name']} {{")
            lines.append(fmt_methods(mw.get("methods", [])))
            lines.append("  }")
        lines.append("}")

    lines.append("")
    lines.append("@enduml")
    return "\n".join(lines)


def build_sequence_puml(diagram):
    """Build PlantUML sequence diagram from a diagram dict."""
    title = diagram.get("title", diagram.get("name", "Sequence"))
    lines = ["@startuml", f"title {title}", ""]

    steps = diagram.get("steps", diagram.get("messages", []))
    for step in steps:
        if isinstance(step, dict):
            frm = step.get("from", step.get("actor", "?"))
            to = step.get("to", step.get("target", "?"))
            msg = step.get("message", step.get("action", step.get("description", "?")))
            ret = step.get("return", step.get("response", ""))

            # Clean names for PlantUML
            frm = frm.replace(" ", "_").replace("/", "_")
            to = to.replace(" ", "_").replace("/", "_")

            lines.append(f'"{frm}" -> "{to}": {msg}')
            if ret:
                lines.append(f'"{to}" --> "{frm}": {ret}')
        elif isinstance(step, str):
            lines.append(f"note over User: {step}")

    lines.append("@enduml")
    return "\n".join(lines)


def main():
    design_path = os.path.join(REPO_ROOT, "brain", "BusinessArtifacts", "Audits", "10_prompt_factory_design.json")
    with open(design_path, encoding="utf-8") as f:
        design = json.load(f)

    doc = Document()

    # Title
    title = doc.add_heading("ChatHealthy FindCare", level=0)
    doc.add_heading("Prompt Factory Architecture Design", level=1)
    doc.add_paragraph(f"Pattern: {design.get('pattern', 'N/A')}")
    doc.add_paragraph(f"Authors: {', '.join(design.get('authors', []))}")
    doc.add_paragraph(f"Iterations: {design.get('iterations', 'N/A')}")
    doc.add_paragraph("")

    # Class Diagram
    doc.add_heading("UML Class Diagram", level=1)
    print("Rendering class diagram...")
    try:
        class_puml = build_class_puml(design)
        img_bytes = render_uml(class_puml)
        if img_bytes and len(img_bytes) > 100:
            stream = io.BytesIO(img_bytes)
            doc.add_picture(stream, width=Inches(6))
            print(f"  Class diagram: {len(img_bytes)} bytes")
        else:
            doc.add_paragraph("[Class diagram rendering failed]")
    except Exception as e:
        doc.add_paragraph(f"[Class diagram error: {e}]")
        print(f"  Class diagram error: {e}")

    # Class descriptions
    doc.add_heading("Class Descriptions", level=2)
    for section in ("interfaces", "abstract_classes", "classes", "middleware"):
        items = design.get("uml_class_diagram", {}).get(section, [])
        for item in items:
            name = item.get("name", "?")
            doc.add_heading(name, level=3)
            if item.get("inherits"):
                doc.add_paragraph(f"Inherits: {', '.join(item['inherits'])}")
            if item.get("implements"):
                doc.add_paragraph(f"Implements: {', '.join(item['implements'])}")
            for m in item.get("methods", []):
                params = ", ".join(f"{p['name']}: {p['type']}" for p in m.get("params", []))
                doc.add_paragraph(f"{m['name']}({params}) -> {m.get('return_type', 'void')}", style="List Bullet")

    # Sequence Diagrams
    doc.add_heading("UML Sequence Diagrams", level=1)
    seq_data = design.get("uml_sequence_diagrams", [])
    if isinstance(seq_data, dict):
        seq_data = seq_data.get("diagrams", [])

    for i, diagram in enumerate(seq_data):
        title = diagram.get("title", diagram.get("name", f"Sequence {i+1}"))
        doc.add_heading(title, level=2)
        print(f"Rendering sequence diagram: {title}...")

        try:
            seq_puml = build_sequence_puml(diagram)
            img_bytes = render_uml(seq_puml)
            if img_bytes and len(img_bytes) > 100:
                stream = io.BytesIO(img_bytes)
                doc.add_picture(stream, width=Inches(6))
                print(f"  {title}: {len(img_bytes)} bytes")
            else:
                doc.add_paragraph(f"[Diagram rendering failed for {title}]")
        except Exception as e:
            doc.add_paragraph(f"[Diagram error: {e}]")
            print(f"  {title} error: {e}")

        # Steps as text
        steps = diagram.get("steps", diagram.get("messages", []))
        for step in steps:
            if isinstance(step, dict):
                frm = step.get("from", step.get("actor", "?"))
                to = step.get("to", step.get("target", "?"))
                msg = step.get("message", step.get("action", step.get("description", "?")))
                doc.add_paragraph(f"{frm} -> {to}: {msg}", style="List Bullet")

    # Backlog
    doc.add_heading("Backlog: Features, Stories, and Requirements", level=1)
    backlog = design.get("backlog", {})
    if isinstance(backlog, list):
        features = backlog
    elif isinstance(backlog, dict):
        features = backlog.get("features", backlog.get("stories", []))
    else:
        features = []

    for feat in features:
        fname = feat.get("name", feat.get("feature_id", "?"))
        doc.add_heading(fname, level=2)
        if feat.get("description"):
            doc.add_paragraph(feat["description"])

        for story in feat.get("stories", []):
            stitle = story.get("title", story.get("story_id", "?"))
            size = story.get("size", "?")
            doc.add_heading(f"{stitle} [{size}]", level=3)
            if story.get("description"):
                doc.add_paragraph(story["description"])

            for req in story.get("requirements", []):
                if isinstance(req, str):
                    rtxt = req
                elif isinstance(req, dict):
                    rtxt = req.get("requirement", req.get("description", "?"))
                else:
                    rtxt = str(req)
                doc.add_paragraph(rtxt, style="List Bullet")

    # Save
    out_path = os.path.join(REPO_ROOT, "brain", "BusinessArtifacts", "Audits", "PromptFactoryDesign.docx")
    doc.save(out_path)
    print(f"\nSaved: {out_path}")
    print(f"Size: {os.path.getsize(out_path):,} bytes")


if __name__ == "__main__":
    main()
