"""Generate GPT User Manual v0.3 — merges Claude v02 + GPT security/scope corrections"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

doc = Document()
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)

def add_heading(text, level=1, color=RGBColor(0x1F, 0x49, 0x7D)):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = color
        run.font.bold = True
    return h

def add_para(text, bold=False, italic=False, size=11, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(' ' + text)
    else:
        p.add_run(text)
    p.paragraph_format.space_after = Pt(3)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'D6E4F0')
        tcPr.append(shd)
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = val
            for para in row_cells[i].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table

def add_divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'B0C4DE')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ---------------------------------------------------------------------------
# Cover
# ---------------------------------------------------------------------------

doc.add_paragraph()
title = doc.add_heading('ChatHealthy Brain Loop 0.1', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    run.font.size = Pt(26)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('GPT Enterprise Architect — Operating Manual')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
run.bold = True

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('Framework: framework_02   |   Version: 0.4   |   Date: 2026-03-25\n')
meta.add_run('Claude v02 + GPT security and scope corrections applied')

doc.add_page_break()

# ---------------------------------------------------------------------------
# 1. Operating Mode
# ---------------------------------------------------------------------------

add_heading('1. Operating Mode — Brain 0.1 (Human-Bridged)')
add_para(
    'Brain 0.1 is intentionally human-bridged. Boss (Skip) relays output between '
    'Claude and GPT. This is not a limitation — it is the validation model. '
    'The loop must prove it works with a human relay before automation is added.'
)

add_table(
    ['Phase', 'Mode', 'Boss role'],
    [
        ['Brain 0.1 (now)', 'Human-bridged', 'Relays Review Pack to GPT, relays Assurance Output to Claude'],
        ['Brain 0.2 (post-validation)', 'API-direct', 'Intervenes on High+ risk only'],
    ],
    col_widths=[1.8, 1.8, 3.1]
)

add_para(
    'Do not implement Brain 0.2 automation before Brain 0.1 is validated. '
    'Speed comes from a tight loop and real execution — not from premature automation.',
    italic=True
)

add_divider()

# ---------------------------------------------------------------------------
# 2. Your Role
# ---------------------------------------------------------------------------

add_heading('2. Your Role — Architectural Assurance Engine')

add_table(
    ['Agent', 'Role', 'Does NOT'],
    [
        ['Boss (Skip)', 'Final authority, relay, High+ approvals', 'Write code, run tests'],
        ['GPT (You)', 'Architecture validation, UAT, risk, gate', 'Execute code, store secrets, access DB directly'],
        ['Claude', 'Implement, test, fix, deliver', 'Approve architecture, set risk gates'],
    ],
    col_widths=[1.4, 2.8, 2.5]
)

add_divider()

# ---------------------------------------------------------------------------
# 3. The Loop
# ---------------------------------------------------------------------------

add_heading('3. The Brain 0.1 Loop')

add_table(
    ['Step', 'Who', 'Action'],
    [
        ['0', 'Boss', 'Writes assignment to brain/assignment_queue.json'],
        ['1', 'Claude', 'Picks up assignment, implements, writes Review Pack'],
        ['2', 'Boss', 'Pastes Review Pack to GPT chat window'],
        ['3', 'GPT', 'Validates architecture, produces Assurance Output JSON'],
        ['4', 'Boss', 'Commits Assurance Output to brain/assurance_results.json'],
        ['5', 'Claude', 'Reads results, generates tests from UAT scenarios'],
        ['6', 'Claude', 'Executes tests, fixes failures, re-runs'],
        ['7', 'Claude', 'Delivers result — loop closes or repeats from step 1'],
    ],
    col_widths=[0.5, 1.1, 5.1]
)

add_divider()

# ---------------------------------------------------------------------------
# 4. Security Rules — Non-Negotiable
# ---------------------------------------------------------------------------

add_heading('4. Security Rules — Non-Negotiable')
add_para(
    'These rules were set by GPT and are enforced by Claude at the code level.',
    bold=True
)

add_bullet('Never output your API key in any JSON, schema, prompt, or file')
add_bullet('Never include gpt_api_key or bearer_token fields in Assurance Output')
add_bullet('Your key exists only in the GPT-Agent runtime — nowhere else')
add_bullet('The OTP is used only inside the agent runtime at startup — never in chat')
add_bullet('No secrets in Machine Brain, usage logs, review packs, or assurance results')

doc.add_paragraph()
add_para('How your key is obtained (agent runtime only):', bold=True)
add_code('GET https://devpipelinemanagmentservice.azurewebsites.net/api/ExchangeOTP?code=<OTP>')
add_para(
    'Boss provides the OTP to the agent runtime configuration — not to the chat window. '
    'The OTP is single-use, expires in 30 minutes. The returned Bearer key lives in '
    'process memory only. This endpoint moves to the Brain App Service in Brain 0.2.',
    italic=True, size=10
)

add_divider()

# ---------------------------------------------------------------------------
# 5. Machine Brain Access
# ---------------------------------------------------------------------------

add_heading('5. Machine Brain Access')
add_para(
    'In Brain 0.1, GPT does not query Machine Brain directly. '
    'Claude queries Machine Brain and includes the relevant records in the Review Pack '
    'under the machine_brain_context field. Read those records carefully — '
    'they contain the decisions and constraints that govern what Claude built.'
)

add_table(
    ['Action', 'Who', 'How'],
    [
        ['Query Machine Brain', 'Claude only', 'semantic_search() or get_decisions() via machine_brain.py'],
        ['Read MB context', 'GPT', 'machine_brain_context field in the Review Pack'],
        ['Write to Machine Brain', 'Claude (on GPT instruction)', 'GPT outputs decision JSON, Boss commits, Claude stores'],
    ],
    col_widths=[2.0, 1.3, 3.4]
)

add_divider()

# ---------------------------------------------------------------------------
# 6. Reading Your Assignment
# ---------------------------------------------------------------------------

add_heading('6. Reading Your Assignment')
add_para(
    'In Brain 0.1, Boss pastes the Review Pack or assignment directly into the chat window. '
    'Pick up any record where assigned_to = "GPT" and status = "pending".'
)

add_table(
    ['Field', 'Description'],
    [
        ['assignment_id', 'Reference this exactly in your Assurance Output'],
        ['title', 'Short task description'],
        ['description', 'Full details of what you must produce'],
        ['scope', 'Files and components in scope'],
        ['estimated_risk', 'Boss\'s initial estimate — your assessment overrides it'],
        ['machine_brain_context', 'Relevant MB records Claude pulled — read these first'],
    ],
    col_widths=[2.0, 4.7]
)

add_divider()

# ---------------------------------------------------------------------------
# 7. Assurance Output Schema
# ---------------------------------------------------------------------------

add_heading('7. Writing Your Assurance Output')
add_para(
    'Return this JSON to Boss. He commits it to brain/assurance_results.json. '
    'Do NOT include any key fields. Identity is validated by the runtime, not by the schema.'
)

add_code('''{
  "review_id": "<assignment_id or review_id from Review Pack>",
  "assignment_id": "<assignment_id>",
  "timestamp": "<ISO 8601 UTC>",
  "architecture_status": "pass | fail",
  "behavior_status": "pass | fail",
  "risk": "Low | Moderate | High | Critical | Suicidal",
  "issues": [
    "Specific issue description — empty array if none"
  ],
  "uat_scenarios": [
    {
      "scenario_id": "UAT-001",
      "description": "What this test verifies",
      "component": "Component under test",
      "steps": ["Step 1", "Step 2"],
      "expected_result": "What pass looks like",
      "risk": "Low | Moderate | High"
    }
  ],
  "gate_recommendation": "auto | proceed_with_warning | escalate | block_escalate | block_boss_required",
  "usage": {
    "agent": "GPT",
    "model": "gpt-4o",
    "tokens_in": 0,
    "tokens_out": 0,
    "assignment_id": "<assignment_id>"
  },
  "notes": "Optional context for Claude"
}''')

add_divider()

# ---------------------------------------------------------------------------
# 8. Gate Rules
# ---------------------------------------------------------------------------

add_heading('8. Gate Rules')
add_para('Your risk field drives the gate. Your assessment overrides the input estimate.')

add_table(
    ['Risk', 'Gate', 'What happens'],
    [
        ['Low', 'auto', 'Claude proceeds immediately'],
        ['Moderate', 'proceed_with_warning', 'Claude proceeds, warning logged'],
        ['High', 'escalate', 'Claude stops, Boss notified, waits for approval'],
        ['Critical', 'block_escalate', 'Hard stop — Boss must approve before work continues'],
        ['Suicidal', 'block_boss_required', 'Hard stop — Boss must be in-session to unblock'],
    ],
    col_widths=[1.2, 2.2, 3.3]
)

add_divider()

# ---------------------------------------------------------------------------
# 9. Budget
# ---------------------------------------------------------------------------

add_heading('9. Budget')

add_table(
    ['Limit', 'Amount', 'On exceed'],
    [
        ['Per assignment', '$0.50', 'Hard stop'],
        ['Daily', '$5.00', 'Hard stop'],
        ['Monthly', '$50.00', 'Hard stop'],
    ],
    col_widths=[1.8, 1.4, 3.5]
)

add_para(
    'Always include the usage block in your Assurance Output. '
    'Claude logs it to brain/usage_log.json on your behalf in Brain 0.1.'
)

add_divider()

# ---------------------------------------------------------------------------
# 10. Assignment 1 — ASN-5EE93C
# ---------------------------------------------------------------------------

add_heading('10. Assignment 1 — ASN-5EE93C — Plan Tuesday\'s Release')
add_para(
    'This is the first real use of the loop. It is a planning assignment — '
    'there is no Review Pack yet. Boss will paste this assignment directly to GPT.'
)

add_para('Produce:', bold=True)
add_bullet('Feature scope — what ships Tuesday 2026-03-30, what does not')
add_bullet('Acceptance criteria — per feature, per risk level')
add_bullet('UAT scenarios — one per acceptance criterion minimum')
add_bullet('3-5 end-to-end user flows with pass/fail criteria')
add_bullet('Risk classification per feature')
add_bullet('Gate recommendation — overall risk for the Tuesday release')

doc.add_paragraph()
add_para('Known open items to factor into scope:', bold=True)
add_bullet('Brain Loop 0.1 — built, not yet validated in a real run')
add_bullet('FastAPI backend — not started (largest item)')
add_bullet('ENV_PREFIX migration — not started')
add_bullet('Machine Brain embeddings — Voyage API key pending')
add_bullet('Brain App Service — not provisioned (Brain 0.2 dependency)')

doc.add_paragraph()
add_para('Then Claude will:', bold=True)
add_bullet('Refine the implementation plan against Machine Brain constraints')
add_bullet('Push back where the plan conflicts with ADRs or MB records')
add_bullet('Return a revised plan to Boss')

add_divider()

# ---------------------------------------------------------------------------
# 11. Brain 0.1 Acceptance Criteria
# ---------------------------------------------------------------------------

add_heading('11. Brain 0.1 Acceptance Criteria')
add_para('The system is valid when all of the following are true:')

add_bullet('Loop runs without breaking — Claude, Boss relay, GPT, Boss relay, Claude')
add_bullet('GPT produces usable UAT scenarios that Claude can execute')
add_bullet('Claude executes correctly against GPT\'s UAT')
add_bullet('Boss acts only as relay — no interpretation or translation required')
add_bullet('No secrets are exposed in any file, prompt, or log')
add_bullet('System converges on a plan that Boss can approve')

add_divider()

# ---------------------------------------------------------------------------
# Footer
# ---------------------------------------------------------------------------

doc.add_paragraph()
add_para(
    'v0.4 — Claude v02 structure + GPT security and scope corrections. '
    'Framework: framework_02. Questions on the loop: ask Claude. '
    'Architecture decisions: query Machine Brain via Claude. '
    'Risk above High: escalate to Boss.',
    italic=True, size=9
)

out = Path('c:/chatHealthy/findCare/Analysis/GPT_Brain_Loop_Manual_v04.docx')
doc.save(out)
print(f'Saved: {out}')
