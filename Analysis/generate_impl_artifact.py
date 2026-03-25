import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt

doc = Document()
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)

def h(level, text):
    doc.add_heading(text, level=level)

def p(text='', bold=False):
    para = doc.add_paragraph(text)
    if bold:
        for run in para.runs:
            run.bold = True
    return para

def b(items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

def table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for i, h_text in enumerate(headers):
        t.rows[0].cells[i].text = h_text
        for run in t.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
    for row_data in rows:
        row = t.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val
    doc.add_paragraph()

def code(text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

# ── TITLE ──────────────────────────────────────────────
doc.add_heading('ChatHealthy — Implementation Artifact v2 (framework_02)', 0)
p('Version: 2.0 — GPT Critical Issues Resolved')
p('Date: 2026-03-25')
p('Author: Claude (Solution Architect)')
p('For Review: GPT (Enterprise Architect), Skip (The Boss)')
p()

# ── EXECUTIVE SUMMARY ──────────────────────────────────
h(1, 'Executive Summary')
p(
    'ChatHealthy is a healthcare navigation platform that helps consumers find the right doctor, '
    'specialist, or care provider for their situation. The product uses conversational AI — the user '
    'describes their symptoms or needs in plain language, and the system identifies the right type of '
    'specialist, finds providers in their area, and guides them through next steps. No medical advice '
    'is given. The platform routes people to care, not away from it.'
)
p()
p(
    'This document describes the engineering blueprint for the next phase of development: '
    'replacing the current prototype (a chatbot hosted on HuggingFace) with a production-grade system '
    'built on Microsoft Azure and Cloudflare. The new system separates the user interface, the AI '
    'conversation engine, and the healthcare data pipelines into three independent applications — '
    'each deployable, scalable, and maintainable on its own.'
)
p()
p(
    'The architecture follows a formal governance model (framework_02) co-designed with GPT as '
    'enterprise architect. Every major decision is recorded, every release is traceable, and every '
    'deployment is gated by automated testing and Boss approval. The system is designed from the '
    'ground up for HIPAA compliance, patient safety, and enterprise readiness.'
)
p()
p(
    'For investors: this document represents the transition from alpha prototype to production '
    'infrastructure. The engineering investment described here is the foundation for scale — '
    'multi-state provider search, insurance eligibility, clinical trial matching, and '
    'enterprise B2B contracts (ASO, Medicare & Medicaid).'
)
p()

# ── TABLE OF CONTENTS ──────────────────────────────────
h(1, 'Table of Contents')
toc_items = [
    '1.  Executive Summary',
    '2.  Table of Contents',
    '3.  Current State',
    '4.  Target State (framework_02)',
    '5.  Git Governance',
    '6.  Environment Variables',
    '7.  MongoDB — Databases and Collections',
    '8.  FastAPI Backend — Port from Gradio',
    '9.  React Frontend',
    '10. CI/CD Pipeline',
    '11. Release Manifest',
    '12. Architecture Decision Records (ADR)',
    '13. Machine Brain Bootstrap',
    '14. Bot Collaboration Directories',
    '15. Production Data Rename Procedure',
    '16. Skip Action List',
    '17. Implementation Order',
    '18. Known Gaps (for GPT Review)',
    '19. Operating Note',
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number')
p()

# ── 1. CURRENT STATE ───────────────────────────────────
h(1, '1. Current State')
p('What exists in production today:')
table(
    ['Component', 'Technology', 'Host', 'Status'],
    [
        ['Chat App', 'Python, Gradio 5.22, OpenAI gpt-4o-mini, Claude Haiku', 'HuggingFace (hf.chathealthy.ai)', 'LIVE'],
        ['Website', 'Static HTML/CSS/JS', 'Cloudflare Pages (chathealthy.ai)', 'LIVE'],
        ['Data Pipelines', 'Python, Azure Functions, CrewAI', 'Azure Functions (US East 2)', 'LIVE'],
        ['MongoDB', 'Atlas — AboutUs.lead, AboutUs.AboutSkip', 'ChatHealthyFrontEndCluster', 'LIVE'],
        ['Dev Site', 'Static HTML + iframe injection', 'Cloudflare Pages (chathealthy-dev.pages.dev)', 'LIVE'],
    ]
)

# ── 2. TARGET STATE ────────────────────────────────────
h(1, '2. Target State (framework_02)')
p('What the system looks like when implementation is complete:')
table(
    ['Component', 'Technology', 'Host', 'Status'],
    [
        ['React Frontend', 'React 18, Vite, TypeScript', 'Cloudflare Pages — dev.chathealthy.ai / chathealthy.ai', 'TO BUILD'],
        ['FastAPI Backend', 'Python 3.12, FastAPI, Uvicorn', 'Azure App Service B1 — api-dev / api.chathealthy.ai', 'TO BUILD'],
        ['Data Pipelines', 'Python, Azure Functions, CrewAI (unchanged)', 'Azure Functions (US East 2)', 'EXISTS — update ENV_PREFIX'],
        ['MongoDB', 'dev/qa/prod databases per cluster', 'ChatHealthyFrontEndCluster + ChatHealthyPipelineCluster', 'TO CREATE'],
        ['Machine Brain', 'MongoDB — ArchitectureMemory, Decisions, Patterns', 'ChatHealthyFrontEndCluster', 'TO CREATE'],
        ['Bot Collaboration', 'claude_thoughts/ + gpt_thoughts/ + docs/', 'Repo', 'TO CREATE'],
    ]
)

# ── 3. GIT GOVERNANCE (FIX 1 + FIX 5) ─────────────────
h(1, '3. Git Governance')
p('Branching model:')
b([
    'main — production truth (renamed from master 2026-03-25 — modern GitHub standard, see ADR-0008)',
    'dev — integration branch (Cloudflare Pages dev site wired to this branch, see ADR-0011)',
    'feature/* — short-lived, merge to dev via PR',
])
p()
p('ADR-0008: main vs master', bold=True)
p('Decision: Use main as production branch.')
p('Rationale: GitHub default standard since 2020. master is legacy. Renamed 2026-03-25 before framework_02 was finalized. No operational impact.')
p('Risk: Low.')
p()
p('ADR-0011: dev vs develop', bold=True)
p('Decision: Retain dev as integration branch name.')
p('Rationale: dev branch is already wired to Cloudflare Pages dev site (chathealthy-dev.pages.dev). Renaming requires Cloudflare reconfiguration with no technical benefit.')
p('Risk: Low. Deviation from framework_02 naming convention — documented here.')
p()
p('Commit footer — mandatory on all commits:')
code('ARCH-FRAMEWORK: framework_02')
p('Release tags:')
code('release-YYYY.MM.DD-env.NN   (e.g. release-2026.03.25-prod.01)')

# ── 4. ENVIRONMENT VARIABLES ───────────────────────────
h(1, '4. Environment Variables')
table(
    ['Variable', 'Dev', 'QA', 'Prod', 'Where Set'],
    [
        ['ENV_PREFIX', 'dev', 'qa', 'prod', 'Azure App Service config'],
        ['PERSIST_MODE', 'debug', 'debug', 'hipaa', 'Azure App Service config'],
        ['MONGODB_CONNECTION_STRING', 'dev conn string', 'qa conn string', 'prod conn string', 'Azure Key Vault'],
        ['ANTHROPIC_API_KEY', 'shared', 'shared', 'shared', 'Azure App Service config'],
        ['OPENAI_API_KEY', 'shared', 'shared', 'shared', 'Azure App Service config'],
        ['VITE_API_URL', 'https://api-dev.chathealthy.ai', 'https://api-qa.chathealthy.ai', 'https://api.chathealthy.ai', 'Cloudflare Pages env var'],
    ]
)

# ── 5. MONGODB ─────────────────────────────────────────
h(1, '5. MongoDB — Databases and Collections')
table(
    ['Cluster', 'Database', 'Collections'],
    [
        ['ChatHealthyFrontEndCluster', 'dev_FindCare', 'leads, sessions, safety_incidents, build_info'],
        ['ChatHealthyFrontEndCluster', 'qa_FindCare', 'leads, sessions, safety_incidents, build_info'],
        ['ChatHealthyFrontEndCluster', 'prod_FindCare', 'leads, sessions, safety_incidents, build_info'],
        ['ChatHealthyFrontEndCluster', 'MachineBrain', 'ArchitectureMemory, Decisions, Patterns'],
        ['ChatHealthyPipelineCluster', 'dev_PublicHealthData', 'mirrors existing PublicHealthData collections'],
        ['ChatHealthyPipelineCluster', 'qa_PublicHealthData', 'mirrors existing PublicHealthData collections'],
        ['ChatHealthyPipelineCluster', 'prod_PublicHealthData', 'IS the current live PublicHealthData — rename per Section 15'],
    ]
)

# ── 6. FASTAPI BACKEND ─────────────────────────────────
h(1, '6. FastAPI Backend — Port from app.py')
p('Code path: Code/ConversationalUX/FindCareChat/backend/')
h(2, '6.1 What Gets Ported')
table(
    ['app.py Component', 'FastAPI Equivalent', 'Notes'],
    [
        ['_safety_check()', 'POST /api/chat (safety gate, first check)', 'Keyword + Haiku semantic, confidence >= 0.80'],
        ['_lock_ip_db() / _check_ip_lock_db()', 'MongoDB safety_incidents collection', 'ENV_PREFIX routes to correct DB'],
        ['_admin_unlock()', 'POST /api/admin/unlock', 'Bearer token required'],
        ['Me class + codebase', 'Loaded at startup, injected into system prompt', 'RULE 7 — no code exposure'],
        ['chat() / _chat()', 'POST /api/chat', 'Returns JSON with reply + locked flag'],
        ['HIPAA consent flow', 'Persisted per PERSIST_MODE', 'debug = log all, hipaa = consent gate'],
        ['Tools (specialty, provider search)', 'FastAPI tool definitions called by Claude', 'Same logic, new structure'],
    ]
)
h(2, '6.2 New in FastAPI')
b([
    'GET /build-info — derived directly from release manifest (Fix 4)',
    'request_id on every request — UUID, returned in response headers, logged',
    'Structured JSON logging — includes request_id, timestamp, model_version, ENV_PREFIX',
    'ENV_PREFIX routing — all DB and blob access prefixed',
    'PERSIST_MODE — controls debug vs HIPAA persistence',
])
h(2, '6.3 Skip Does')
b([
    'Create Azure App Service B1 — name: chathealthy-api-dev',
    'Set env vars: ENV_PREFIX=dev, PERSIST_MODE=debug, MONGODB_CONNECTION_STRING, ANTHROPIC_API_KEY, OPENAI_API_KEY',
    'Set custom domain: api-dev.chathealthy.ai',
    'Create GitHub secret: AZURE_APP_SERVICE_PUBLISH_PROFILE_DEV',
])

# ── 7. REACT FRONTEND ──────────────────────────────────
h(1, '7. React Frontend')
p('Code path: Code/ConversationalUX/FindCareChat/frontend/')
h(2, '7.1 Components')
b([
    'ChatWindow — message state, API calls, locked state',
    'MessageBubble — react-markdown, rehype-raw, rehype-sanitize (span.state-name only)',
    'WelcomeMessage — specialty list',
    'ProviderCard — inline provider results',
    'SafetyBanner — shown when session is locked',
])
h(2, '7.2 Skip Does')
b([
    'Add VITE_API_URL env var to chathealthy-dev Cloudflare Pages project',
    'Update Cloudflare Pages build command: npm install && npm run build',
    'Update build output directory: dist',
])

# ── 8. CI/CD PIPELINE ─────────────────────────────────
h(1, '8. CI/CD Pipeline')
table(
    ['Stage', 'Trigger', 'Gate', 'Rollback'],
    [
        ['Build', 'Push to any branch', 'Compile / lint', 'N/A'],
        ['Test', 'Push to any branch', 'Unit tests pass', 'N/A'],
        ['Dev Deploy', 'Push to dev branch', 'Tests pass', 'Revert commit'],
        ['QA Gate', 'Manual trigger', 'Schema check, smoke tests, regression (GPT designs)', 'Required'],
        ['Prod Deploy', 'Manual — Boss approves + signs manifest', 'All QA gates + manifest + Boss sign-off', 'Required'],
    ]
)
p('Rule: No stage may be skipped. No production deployment without valid manifest.')

# ── 9. RELEASE MANIFEST (FIX 2 + FIX 3) ───────────────
h(1, '9. Release Manifest')
p('A release manifest is required for every production deployment. The manifest is generated by CI, stored immutably, and must be signed by the Boss before prod deploy proceeds.', bold=False)
p()
h(2, '9.1 Manifest Schema')
code(
    '{\n'
    '  "release_id": "release-YYYY.MM.DD-prod.NN",\n'
    '  "architecture_framework": "framework_02",\n'
    '  "git_commit": "SHA",\n'
    '  "environment": "prod",\n'
    '  "component_versions": {\n'
    '    "frontend": "x.x.x",\n'
    '    "backend": "x.x.x",\n'
    '    "pipeline": "x.x.x"\n'
    '  },\n'
    '  "schema_version": "x.x",\n'
    '  "embedding_version": "x.x",\n'
    '  "risk_level": "Low|Moderate|High|Critical|Suicidal",\n'
    '  "approved_by": "Boss",\n'
    '  "approval_timestamp": "UTC ISO string",\n'
    '  "generated_by": "CI",\n'
    '  "generated_at": "UTC ISO string"\n'
    '}'
)
h(2, '9.2 Storage and Immutability')
b([
    'Stored in: docs/manifests/release-YYYY.MM.DD-prod.NN.json',
    'Committed to repo — immutable once merged to main',
    'Also attached as GitHub Actions artifact for independent audit trail',
    'CI artifact alone is NOT sufficient — must be committed to repo',
])
h(2, '9.3 Boss Approval Rule (Fix 2)')
b([
    'Release is INVALID unless manifest exists',
    'Release is INVALID unless manifest includes approved_by: Boss',
    'Release is INVALID unless risk_level is declared',
    'CI workflow blocks prod deploy if manifest is missing or unsigned',
    'Boss signs by committing the manifest with approval_timestamp filled in',
])
h(2, '9.4 Build-Info Endpoint (Fix 4)')
p('/build-info is served directly from the committed manifest. It does not generate its own data.')
code(
    'GET /build-info\n'
    '-> reads docs/manifests/[current release manifest]\n'
    '-> returns release_id, framework, git_commit, schema_version, embedding_version'
)

# ── 10. ADR REGISTRY (FIX 5) ──────────────────────────
h(1, '10. Architecture Decision Records (ADR)')
p('Path: docs/adr/')
table(
    ['ADR', 'Decision', 'Risk', 'Status'],
    [
        ['ADR-0001', 'Three-application architecture (Website / Chat / Pipelines)', 'Low', 'DECIDED'],
        ['ADR-0002', 'Cloudflare Pages for all static/frontend hosting (not Azure SWA)', 'Low', 'DECIDED'],
        ['ADR-0003', 'FastAPI on Azure App Service B1 — port from Gradio/HuggingFace', 'Moderate', 'DECIDED'],
        ['ADR-0004', 'ENV_PREFIX pattern for environment-aware data routing', 'Low', 'DECIDED'],
        ['ADR-0005', 'chat-url.txt committed per branch for iframe URL injection', 'Low', 'DECIDED'],
        ['ADR-0006', 'Two MongoDB clusters — FrontEnd (user-facing) + Pipeline (ingestion)', 'Low', 'DECIDED'],
        ['ADR-0007', 'Machine Brain — MongoDB for persistent architectural memory', 'Low', 'DECIDED'],
        ['ADR-0008', 'main instead of master — modern GitHub standard, renamed 2026-03-25', 'Low', 'DECIDED'],
        ['ADR-0009', 'PERSIST_MODE: debug (dev/QA) vs hipaa (prod)', 'High', 'DECIDED'],
        ['ADR-0010', 'Safety incidents: IP + timestamp only pending legal review', 'High', 'DECIDED'],
        ['ADR-0011', 'dev retained as integration branch — Cloudflare wiring, deviation from framework_02 naming', 'Low', 'DECIDED'],
    ]
)

# ── 11. MACHINE BRAIN ──────────────────────────────────
h(1, '11. Machine Brain Bootstrap')
p('Database: MachineBrain on ChatHealthyFrontEndCluster.')
p('Claude creates the database and seeds it with all ADR decisions above.')
p('Decision record schema:')
code(
    '{\n'
    '  "topic": string,\n'
    '  "type": "architectural|operational|security|compliance",\n'
    '  "decision": string,\n'
    '  "rationale": string,\n'
    '  "risk": "Low|Moderate|High|Critical|Suicidal",\n'
    '  "constraints": [string],\n'
    '  "components": [string],\n'
    '  "created_by": "GPT|Claude|Skip",\n'
    '  "framework": "framework_02",\n'
    '  "timestamp": "UTC ISO string",\n'
    '  "adr_ref": "ADR-XXXX"\n'
    '}'
)
p('Retrieval contract: Claude queries Machine Brain before implementing. GPT writes after decisions.')

# ── 12. BOT COLLABORATION ──────────────────────────────
h(1, '12. Bot Collaboration Directories')
code(
    'claude_thoughts/    <- Claude drops proposals, questions, implementation notes\n'
    'gpt_thoughts/       <- GPT drops architecture decisions, feedback, approvals\n'
    'docs/adr/           <- Architecture Decision Records\n'
    'docs/manifests/     <- Release manifests (immutable once committed)\n'
)
p('File naming: YYYY-MM-DD_[topic].md')

# ── 13. PRODUCTION DATA RENAME PROCEDURE (FIX 6) ───────
h(1, '13. Production Data Rename Procedure')
p('Risk: CRITICAL', bold=True)
p('Renaming PublicHealthData to prod_PublicHealthData requires a controlled procedure:')
h(2, 'Step 1 — Backup')
b([
    'Full Atlas snapshot of PublicHealthData before any change',
    'Verify snapshot is restorable',
    'Document snapshot ID in release manifest',
])
h(2, 'Step 2 — Dry Run in QA')
b([
    'Create qa_PublicHealthData as a copy of PublicHealthData',
    'Update all pipeline and frontend code to use qa_PublicHealthData',
    'Run all pipelines against qa_PublicHealthData',
    'Run all frontend queries against qa_PublicHealthData',
    'Validate record counts and spot-check data integrity',
])
h(2, 'Step 3 — Production Rename')
b([
    'Schedule maintenance window — announce to Skip',
    'Put HuggingFace space in maintenance mode (no live traffic to DB)',
    'Atlas: rename PublicHealthData to prod_PublicHealthData',
    'Deploy updated code (ENV_PREFIX=prod) to production',
    'Smoke test immediately after rename',
])
h(2, 'Step 4 — Validation')
b([
    'Run validation queries: count records, check indexes, verify last-write timestamps',
    'Confirm /build-info reports correct schema_version',
    'Monitor error rate for 30 minutes post-rename',
])
h(2, 'Step 5 — Rollback Plan')
b([
    'If validation fails: restore Atlas snapshot',
    'Redeploy previous code version',
    'RCA required before retry',
])

# ── 14. SKIP ACTION LIST ───────────────────────────────
h(1, '14. Skip Action List')
table(
    ['#', 'Action', 'Where', 'Whine Level'],
    [
        ['1', 'Create Azure App Service B1 — chathealthy-api-dev', 'Azure Portal', 'Medium'],
        ['2', 'Set env vars on App Service (ENV_PREFIX, PERSIST_MODE, keys)', 'Azure Portal', 'Medium'],
        ['3', 'Create GitHub secret: AZURE_APP_SERVICE_PUBLISH_PROFILE_DEV', 'GitHub Settings', 'Low'],
        ['4', 'Set custom domain api-dev.chathealthy.ai on App Service + Cloudflare DNS', 'Azure + Cloudflare', 'Medium'],
        ['5', 'Add VITE_API_URL to chathealthy-dev Cloudflare Pages project', 'Cloudflare Dashboard', 'Low'],
        ['6', 'Update Cloudflare Pages build command: npm install && npm run build, output: dist', 'Cloudflare Dashboard', 'Low'],
        ['7', 'Rename MongoDB cluster to ChatHealthyFrontEndCluster', 'MongoDB Atlas', 'Low'],
        ['8', 'Confirm connection strings for dev_FindCare and MachineBrain', 'MongoDB Atlas', 'Low'],
        ['9', 'Sign release manifests for prod deployments (approved_by: Boss)', 'docs/manifests/', 'Low — but mandatory'],
        ['10', 'Review and approve ADR-0001 through ADR-0011', 'docs/adr/', 'Low'],
    ]
)

# ── 15. IMPLEMENTATION ORDER ───────────────────────────
h(1, '15. Implementation Order')
table(
    ['Phase', 'What Claude Builds', 'Depends On'],
    [
        ['1', 'docs/adr/ + 11 ADR files', 'Nothing — start now'],
        ['2', 'claude_thoughts/, gpt_thoughts/, docs/manifests/ directories', 'Nothing — start now'],
        ['3', 'MongoDB init script — all databases and collections', 'Nothing — start now'],
        ['4', 'Machine Brain bootstrap — seed from ADR list', 'MongoDB init script + Atlas access'],
        ['5', 'FastAPI backend — full port + /build-info + ENV_PREFIX + request_id + structured logging', 'Azure App Service (Skip #1)'],
        ['6', 'GitHub Actions deploy-findcare-backend.yml', 'Azure publish profile secret (Skip #3)'],
        ['7', 'React frontend — all components', 'VITE_API_URL set (Skip #5)'],
        ['8', 'Release manifest generator — CI artifact + docs/manifests/ commit', 'FastAPI + React deployed'],
        ['9', 'QA dry run of prod_PublicHealthData rename', 'qa_PublicHealthData created'],
        ['10', 'chat-url.txt on dev updated to React URL', 'React deployed and smoke-tested'],
    ]
)

# ── 16. KNOWN GAPS ─────────────────────────────────────
h(1, '16. Known Gaps (for GPT Review)')
b([
    'Regression test design: GPT owns gate criteria — not yet delivered. Blocks QA → Prod promotion.',
    'Machine Brain enforcement: query-before-implement is currently manual convention. Tooling needed (Phase 2).',
    'Schema versioning: owner and timeline not yet defined. Required by framework_02.',
    'Observability: structured logging is in FastAPI Phase 5 above. Metrics/alerting is Phase 2.',
    'QA environment: Azure App Service QA slot and qa.chathealthy.ai not yet created. Phase 2.',
    'Compliance audit: PHI controls not yet formally audited. Legal review on safety incident fields pending.',
    'HuggingFace: stays live until React+FastAPI is smoke-tested. Cutover date TBD by Skip.',
    'dev branch deviation: documented in ADR-0011. GPT acknowledgment requested.',
])

# ── 17. OPERATING NOTE ────────────────────────────────
h(1, '17. Operating Note')
p(
    'This is a production healthcare system. Safety, auditability, and determinism are required '
    'at every layer. No shortcuts in prod. Claude queries Machine Brain before implementing. '
    'GPT reviews before prod promotion. Boss approves and signs every production release manifest.',
    bold=True
)

doc.save('c:/chatHealthy/findCare/Analysis/ChatHealthy-Implementation-Artifact-framework02-v2.docx')
print('Done')
