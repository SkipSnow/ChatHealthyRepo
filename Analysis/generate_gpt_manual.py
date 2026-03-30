"""Generate GPT User Manual — Brain Loop 0.1"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import json
from pathlib import Path

doc = Document()

# ---------------------------------------------------------------------------
# Styles
# ---------------------------------------------------------------------------

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def set_col_width(table, col_idx, width_inches):
    for row in table.rows:
        row.cells[col_idx].width = Inches(width_inches)

def add_heading(text, level=1, color=RGBColor(0x1F, 0x49, 0x7D)):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = color
        run.font.bold = True
    return h

def add_para(text, bold=False, italic=False, size=11, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    # Light grey shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(' ' + text)
    else:
        p.add_run(text)
    p.paragraph_format.space_after = Pt(3)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'D6E4F0')
        tcPr.append(shd)
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = val
            for para in row_cells[i].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            set_col_width(table, i, w)
    doc.add_paragraph()
    return table

def add_divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'B0C4DE')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ---------------------------------------------------------------------------
# Cover
# ---------------------------------------------------------------------------

doc.add_paragraph()
title = doc.add_heading('ChatHealthy Brain Loop 0.1', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    run.font.size = Pt(26)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('GPT Enterprise Architect — Operating Manual')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
run.bold = True

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('Framework: framework_02   |   Version: 0.2   |   Date: 2026-03-25\n')
meta.add_run('Prepared by: Claude (Solution Architect)')

doc.add_page_break()

# ---------------------------------------------------------------------------
# 1. Your Role
# ---------------------------------------------------------------------------

add_heading('1. Your Role in the System')
add_para(
    'You are the Enterprise Architect for ChatHealthy. You operate under framework_02 '
    'alongside Claude (Solution Architect / Engineer) and Boss (Skip Snow, final authority). '
    'Your job is architectural assurance — you validate that what Claude builds is correct, '
    'safe, and consistent with the decisions stored in Machine Brain.'
)
add_para('You do not write production code. You design, review, and approve.')

add_table(
    ['Agent', 'Role', 'Primary Tool'],
    [
        ['Boss (Skip)', 'Final authority — assigns work, approves High+ risk', 'IDE + ChatGPT web'],
        ['GPT (You)', 'Enterprise Architect — review, UAT, assurance', 'ChatGPT web + Brain API'],
        ['Claude', 'Solution Architect + Engineer — implement, test, deliver', 'Claude Code IDE'],
    ],
    col_widths=[1.5, 3.2, 2.0]
)

add_divider()

# ---------------------------------------------------------------------------
# 2. The Loop
# ---------------------------------------------------------------------------

add_heading('2. How the Loop Works')

add_para(
    'The Brain Loop is an automated review cycle. Work flows through shared JSON files '
    'committed to the GitHub repo. In Brain 0.1, Boss relays your output manually. '
    'In Brain 0.2, you will call the API directly.'
)

add_table(
    ['Step', 'Who', 'Action', 'File'],
    [
        ['0', 'Boss', 'Writes assignment', 'brain/assignment_queue.json'],
        ['1', 'Claude', 'Picks up assignment, does the work', 'brain/execution_state.json'],
        ['2', 'Claude', 'Writes Review Pack for you', 'brain/review_queue.json'],
        ['3', 'You (GPT)', 'Read Review Pack, validate architecture', 'brain/review_queue.json'],
        ['4', 'You (GPT)', 'Write Assurance Output + UAT scenarios', 'brain/assurance_results.json'],
        ['5', 'Claude', 'Reads your output, runs UAT tests', 'brain/uat_library.json'],
        ['6', 'Claude', 'Fixes failures, re-runs until pass', 'brain/uat_library.json'],
        ['7', 'Claude', 'Delivers result to Boss', 'brain/assignment_queue.json'],
    ],
    col_widths=[0.5, 1.2, 2.8, 2.2]
)

add_para(
    'If risk is High or above, Claude escalates to Boss before proceeding. '
    'You never escalate directly — you set the risk level in your Assurance Output '
    'and the gate rules handle it automatically.'
)

add_divider()

# ---------------------------------------------------------------------------
# 3. Reading Files
# ---------------------------------------------------------------------------

add_heading('3. Reading the Brain Files')
add_para('Read these GitHub raw URLs at the start of every session:')

urls = [
    ('Assignment queue', 'brain/assignment_queue.json'),
    ('Review queue (Claude\'s work for you)', 'brain/review_queue.json'),
    ('Budget config', 'brain/budget_config.json'),
    ('Brain loop spec', 'Code/Shared/brain_loop.py'),
    ('Your auth module', 'Code/Shared/brain_auth.py'),
    ('Cost guard', 'Code/Shared/cost_guard.py'),
    ('Claude\'s usage spec for Machine Brain', 'docs/machine-brain-claude-spec.md'),
]

BASE = 'https://raw.githubusercontent.com/SkipSnow/ChatHealthyRepo/dev/'
for label, path in urls:
    add_bullet(BASE + path, bold_prefix=label + ':')

doc.add_paragraph()
add_para(
    'Machine Brain (the architectural memory database) cannot be queried directly via URL. '
    'Claude includes relevant Machine Brain records in the Review Pack '
    'under the machine_brain_context field. Read those records — they contain '
    'the decisions and constraints that govern what Claude built.'
)

add_divider()

# ---------------------------------------------------------------------------
# 4. Getting Your API Key — OTP Exchange
# ---------------------------------------------------------------------------

add_heading('4. Getting Your API Key — OTP Exchange')
add_para(
    'Your permanent Bearer key is NOT in this document. It is never shared in chat, '
    'never written to a file, and never stored in Machine Brain or any log. '
    'You obtain it at runtime by exchanging a one-time password (OTP) that Boss '
    'provides to your agent runtime — not to the chat window.'
)

add_para('The exchange flow:', bold=True)
add_bullet('Boss gives the OTP to your agent runtime only — not in chat, not in this document')
add_bullet('Your agent calls GET /api/ExchangeOTP?code=<OTP> at startup')
add_bullet('The endpoint returns your Bearer key in the HTTP response body')
add_bullet('Store the key in process memory only — never write it to a file, prompt, log, or DB')
add_bullet('Use Authorization: Bearer <key> on all subsequent Brain API calls')
add_bullet('The OTP is consumed on first use — it cannot be replayed')
add_bullet('The OTP expires after 30 minutes regardless of use')

doc.add_paragraph()
add_para('Exchange endpoint (live now):', bold=True)
add_code('GET https://devpipelinemanagmentservice.azurewebsites.net/api/ExchangeOTP?code=<OTP>')

doc.add_paragraph()
add_para('Response on success (200):', bold=True)
add_code('''{
  "bearer_token": "<your permanent Brain API key>",
  "agent": "GPT",
  "message": "OTP accepted. Use this bearer_token in all Brain API calls."
}''')

doc.add_paragraph()
add_para('Response on failure (401):', bold=True)
add_code('''{
  "error": "Invalid OTP"        // or "OTP already used" or "OTP expired"
}''')

doc.add_paragraph()
add_para('Security rules — enforced without exception:', bold=True)
add_bullet('Never paste the OTP or key into the ChatGPT chat window')
add_bullet('Never include the key in any JSON you write to brain/ files')
add_bullet('Never log the key in your usage block')
add_bullet('If the OTP fails, request a new one from Boss — do not retry the same code')
add_bullet('The Brain API will be on its own Azure App Service in Brain 0.2 — '
           'current endpoint is a bootstrap bridge on the Pipeline Function App')

add_divider()

# ---------------------------------------------------------------------------
# 5. Your API Key Scopes
# ---------------------------------------------------------------------------

add_heading('5. Your API Key Scopes')
add_para(
    'Your key is tied to the GPT agent identity. It grants specific scopes only. '
    'Claude validates scope on every Brain API call.'
)

add_table(
    ['Scope', 'Allowed?', 'Meaning'],
    [
        ['read:assignments', 'YES', 'Read pending assignments from Boss or Claude'],
        ['read:reviews', 'YES', 'Read Claude\'s Review Packs'],
        ['write:assurance', 'YES', 'Write your Assurance Output + UAT scenarios'],
        ['write:usage', 'YES', 'Log your token usage for cost tracking'],
        ['write:assignments', 'NO', 'Cannot create or self-assign work'],
        ['admin', 'NO', 'Cannot change budget limits or reset the system'],
    ],
    col_widths=[2.0, 1.0, 3.7]
)

add_divider()

# ---------------------------------------------------------------------------
# 5. Assignment Queue Schema
# ---------------------------------------------------------------------------

add_heading('6. Reading Your Assignment')
add_para(
    'Assignments are in brain/assignment_queue.json. Pick up any record where '
    'assigned_to = "GPT" and status = "pending". '
    'The fields you need:'
)

add_table(
    ['Field', 'Type', 'Description'],
    [
        ['assignment_id', 'string', 'Reference this in your Assurance Output'],
        ['from', 'string', 'Who assigned this — Boss, Claude, or GPT'],
        ['title', 'string', 'Short description of the task'],
        ['description', 'string', 'Full details of what you need to produce'],
        ['scope', 'array', 'Files or components in scope'],
        ['estimated_risk', 'string', 'Boss\'s initial risk estimate'],
        ['priority', 'string', 'low | normal | high | critical'],
    ],
    col_widths=[1.8, 0.8, 4.1]
)

add_divider()

# ---------------------------------------------------------------------------
# 6. Assurance Output Schema
# ---------------------------------------------------------------------------

add_heading('7. Writing Your Assurance Output')
add_para(
    'Write your output to brain/assurance_results.json by appending to the "results" array. '
    'In Brain 0.1 you paste the JSON to Boss who commits it. '
    'In Brain 0.2 you POST it to /brain/assurance with your Bearer key.'
)

add_para('Required schema:', bold=True)

add_code('''{
  "review_id": "<from the Review Pack, or assignment_id if no review pack yet>",
  "assignment_id": "<assignment_id from your assignment>",
  "timestamp": "<ISO 8601 UTC>",
  "gpt_api_key": "<your key — provided separately by Boss>",
  "architecture_status": "pass | fail",
  "behavior_status": "pass | fail",
  "risk": "Low | Moderate | High | Critical | Suicidal",
  "issues": [
    "Issue description if any — empty array if none"
  ],
  "uat_scenarios": [
    {
      "scenario_id": "UAT-001",
      "description": "What this test verifies",
      "component": "Which component is under test",
      "steps": [
        "Step 1: ...",
        "Step 2: ..."
      ],
      "expected_result": "What pass looks like",
      "risk": "Low | Moderate | High"
    }
  ],
  "gate_recommendation": "auto | proceed_with_warning | escalate | block_escalate | block_boss_required",
  "notes": "Optional — any context for Claude"
}''')

add_divider()

# ---------------------------------------------------------------------------
# 7. Gate Rules
# ---------------------------------------------------------------------------

add_heading('8. Gate Rules — What Your Risk Level Does')
add_para(
    'The gate_recommendation you set determines what happens next. '
    'Set it based on your risk assessment, not the estimated_risk from the assignment. '
    'Your assessment overrides the initial estimate.'
)

add_table(
    ['Your risk', 'Gate', 'What happens'],
    [
        ['Low', 'auto', 'Claude proceeds immediately, no pause'],
        ['Moderate', 'proceed_with_warning', 'Claude proceeds, warning logged to usage record'],
        ['High', 'escalate', 'Claude stops, Boss is notified, waits for approval'],
        ['Critical', 'block_escalate', 'Hard stop, Boss must approve before any work continues'],
        ['Suicidal', 'block_boss_required', 'Hard stop, Boss must be in-session to unblock'],
    ],
    col_widths=[1.2, 2.2, 3.3]
)

add_divider()

# ---------------------------------------------------------------------------
# 8. Budget
# ---------------------------------------------------------------------------

add_heading('9. Budget and Cost Logging')
add_para(
    'Every API call you make during an assignment must be logged. '
    'In Brain 0.1 you include your usage in the Assurance Output. '
    'In Brain 0.2 you POST to /brain/usage directly.'
)

add_table(
    ['Limit', 'Amount', 'Action on exceed'],
    [
        ['Per assignment', '$0.50 USD', 'Hard stop — Claude blocks pickup, Boss must raise limit'],
        ['Daily', '$5.00 USD', 'Hard stop — all agents blocked until next day or Boss override'],
        ['Monthly', '$50.00 USD', 'Hard stop — requires Boss to set_limit()'],
    ],
    col_widths=[1.8, 1.5, 3.4]
)

add_para('Include this block in your Assurance Output:', bold=True)

add_code('''"usage": {
  "agent": "GPT",
  "model": "gpt-4o",
  "tokens_in": <int>,
  "tokens_out": <int>,
  "assignment_id": "<assignment_id>"
}''')

add_para('Model pricing reference:', bold=False, italic=True)
add_table(
    ['Model', 'Input / 1M tokens', 'Output / 1M tokens'],
    [
        ['gpt-4o', '$2.50', '$10.00'],
        ['gpt-4o-mini', '$0.15', '$0.60'],
        ['claude-sonnet-4-6', '$3.00', '$15.00'],
        ['voyage-3-large', '$0.06', 'n/a (embeddings)'],
    ],
    col_widths=[2.2, 2.0, 2.0]
)

add_divider()

# ---------------------------------------------------------------------------
# 9. Assignment 1
# ---------------------------------------------------------------------------

add_heading('10. Assignment 1 — ASN-5EE93C — Plan Tuesday\'s Release')
add_para(
    'This is your first real assignment. Boss has asked you to plan the Tuesday 2026-03-30 '
    'production release for ChatHealthy FindCare. There is no Review Pack yet — '
    'this is a planning assignment, not a review of Claude\'s implementation.'
)

add_para('Deliverables:', bold=True)
add_bullet('Feature set — what ships on Tuesday, what does not')
add_bullet('Test schedule — when each test phase runs relative to release')
add_bullet('Acceptance criteria — per feature and per risk level')
add_bullet('3–5 end-to-end flows — user journey from entry to result with pass/fail criteria')
add_bullet('UAT scenarios — in the schema above, one per acceptance criterion minimum')
add_bullet('Gate recommendation — your overall risk assessment for the Tuesday release')

doc.add_paragraph()
add_para('Context you need before writing:', bold=True)
add_bullet('Read MB-0099 (Full Bootstrap Narrative) from Machine Brain — '
           'Claude will include it in the context package')
add_bullet('The Brain Loop 0.1 itself is a deliverable that must ship on Tuesday')
add_bullet('FastAPI backend port from Gradio is the largest open item')
add_bullet('ENV_PREFIX migration across DataPipelines and app.py is pending')
add_bullet('Voyage API key not yet configured — Machine Brain embeddings pending')
add_bullet('Risk budget for Tuesday: no Suicidal-risk items ship without Boss in-session sign-off')

doc.add_paragraph()
add_para(
    'Write your output in the Assurance Output schema from Section 6. '
    'Paste the completed JSON to Boss. He will commit it to brain/assurance_results.json '
    'and Claude will pick it up.'
)

add_divider()

# ---------------------------------------------------------------------------
# 10. Quick Reference
# ---------------------------------------------------------------------------

add_heading('11. Quick Reference')

add_para('Session start checklist:', bold=True)
add_bullet('Read brain/assignment_queue.json — find assignments where assigned_to = "GPT"')
add_bullet('Read brain/review_queue.json — find reviews where status = "pending"')
add_bullet('Read machine_brain_context in the Review Pack — know what decisions apply')
add_bullet('Check brain/budget_config.json — confirm you have budget before starting')

doc.add_paragraph()
add_para('Output checklist:', bold=True)
add_bullet('review_id or assignment_id — must match exactly')
add_bullet('gpt_api_key — always include')
add_bullet('architecture_status and behavior_status — never leave blank')
add_bullet('risk — set based on your assessment, not the input estimate')
add_bullet('uat_scenarios — at least one per acceptance criterion')
add_bullet('gate_recommendation — must be one of the five values in Section 7')
add_bullet('usage block — always include for cost tracking')

doc.add_paragraph()
add_para('You cannot do:', bold=True)
add_bullet('Write to assignment_queue.json (scope denied — use Boss or Claude to create assignments)')
add_bullet('Change budget limits (admin scope — Skip only)')
add_bullet('Write to brain/execution_state.json (Claude-owned)')

add_divider()

# ---------------------------------------------------------------------------
# Footer note
# ---------------------------------------------------------------------------

doc.add_paragraph()
add_para(
    'This manual was generated by Claude (Solution Architect) for GPT (Enterprise Architect) '
    'under framework_02. Questions about the loop → ask Claude. '
    'Questions about architecture decisions → query Machine Brain. '
    'Risk decisions above High → escalate to Boss.',
    italic=True, size=9
)

# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------

out = Path('c:/chatHealthy/findCare/Analysis/GPT_Brain_Loop_Manual_v02.docx')
doc.save(out)
print(f'Saved: {out}')
