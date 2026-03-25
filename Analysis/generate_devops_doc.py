from docx import Document
from docx.shared import Pt

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    for row_data in rows:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val
    doc.add_paragraph()

def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

# Title
doc.add_heading('ChatHealthy — End-to-End DevOps & Architecture (FINAL)', 0)
doc.add_paragraph('Date: 2026-03-25')
doc.add_paragraph('Authors: Skip Snow (Product / The Boss), GPT (Enterprise Architecture), Claude (Solution Architecture)')
doc.add_paragraph()

# 1. Purpose
doc.add_heading('1. Purpose', 1)
doc.add_paragraph(
    'Define a production-grade DevOps, architecture, and system intelligence model for '
    'ChatHealthy / FindCare. This system is production-first: safe, deterministic, and auditable.'
)

# 2. Actors
doc.add_heading('2. Actors', 1)
add_table(doc,
    ['Role', 'Agent', 'Responsibilities'],
    [
        ['Product Director (The Boss)', 'Skip Snow', 'Final authority, risk acceptance, production release'],
        ['Enterprise Architect', 'GPT', 'Architecture, risk, standards, QA gates, regression test design. Writes to Machine Brain after decisions.'],
        ['Solution Architect / Engineer', 'Claude', 'Implementation, unit testing, deployment. Queries Machine Brain before implementing.'],
    ]
)
doc.add_paragraph('Roles may be fluid but responsibilities are not.')

# 3. Risk Model
doc.add_heading('3. Risk Model', 1)
doc.add_paragraph('All decisions must include a risk level:')
add_table(doc,
    ['Level', 'Description'],
    [
        ['Low', 'Routine change, easily reversible'],
        ['Moderate', 'Some impact, reversible with effort'],
        ['High', 'Significant impact, difficult to reverse'],
        ['Critical', 'Production impact, requires immediate escalation to Skip'],
        ['Suicidal', 'Catastrophic — do not proceed without full team sign-off'],
    ]
)

# 4. System Architecture
doc.add_heading('4. System Architecture', 1)
doc.add_paragraph('Three-application model with strict separation of concerns:')
add_table(doc,
    ['#', 'Application', 'Host', 'Code Path'],
    [
        ['1', 'Static Website', 'Cloudflare Pages', 'Website/'],
        ['2', 'Conversational UX (React + FastAPI)', 'Cloudflare Pages + Azure App Service', 'Code/ConversationalUX/FindCareChat/'],
        ['3', 'Data Pipelines', 'Azure Functions', 'Code/DataPipelines/'],
    ]
)
doc.add_heading('Boundary Rules', 2)
bullets(doc, [
    'App 1: static only — no business logic, no LLM calls, no DB access',
    'App 2: UX + LLM + DB reads, writes to application DBs only — no pipeline writes',
    'App 3: owns all PublicHealthData writes — ingestion, embeddings, multi-agent workflows',
    'Communication via /api/Router with Bearer token authentication',
    'No production deployment without gating and rollback capability',
])

# 5. Environments
doc.add_heading('5. Environments', 1)
add_table(doc,
    ['Environment', 'Frontend', 'Backend', 'DB Prefix'],
    [
        ['Dev', 'chathealthy-dev.pages.dev', 'Azure App Service dev slot', 'dev_'],
        ['QA', 'qa.chathealthy.ai (future)', 'Azure App Service QA slot', 'qa_'],
        ['Prod', 'chathealthy.ai', 'Azure App Service prod slot', 'prod_'],
    ]
)
doc.add_paragraph('Branch does NOT determine environment — environment is controlled by deployment target and ENV_PREFIX.')

# 6. Data Architecture
doc.add_heading('6. Data Architecture', 1)
doc.add_heading('MongoDB Clusters', 2)
add_table(doc,
    ['Cluster', 'Purpose'],
    [
        ['ChatHealthyFrontEndCluster', 'User-facing data — fast, stable, low latency'],
        ['ChatHealthyPipelineCluster', 'Pipeline data — tolerates heavy ingestion load'],
    ]
)
doc.add_heading('Database Naming Convention', 2)
add_table(doc,
    ['Environment', 'Frontend DB', 'Pipeline DB'],
    [
        ['Dev', 'dev_FindCare', 'dev_PublicHealthData'],
        ['QA', 'qa_FindCare', 'qa_PublicHealthData'],
        ['Prod', 'prod_FindCare', 'prod_PublicHealthData'],
    ]
)
bullets(doc, [
    'Collections retain exact names across all environments',
    'Schema validation: none in dev, warn mode in QA, strict in prod (future)',
    'Strict write ownership enforced — App 2 never writes to PublicHealthData',
])

# 7. Environment Variables
doc.add_heading('7. Environment Variable Pattern', 1)
doc.add_paragraph('Single ENV_PREFIX controls all data routing. Future separation of DB and blob vars may be required.')
add_code(doc,
    'ENV_PREFIX=dev    # dev_FindCare, dev_PublicHealthData, dev-[blob]\n'
    'ENV_PREFIX=qa     # qa_FindCare, qa_PublicHealthData, qa-[blob]\n'
    'ENV_PREFIX=prod   # prod_FindCare, prod_PublicHealthData, prod-[blob]'
)
doc.add_heading('Iframe URL Injection', 2)
bullets(doc, [
    'Website/chat-url.txt committed per environment branch — contains iframe src URL',
    'Build command: CHAT_URL=$(cat chat-url.txt) && sed -i "s|%%CHAT_URL%%|$CHAT_URL|g" index.html',
    'Changing iframe URL = one-line commit to chat-url.txt, no dashboard changes',
])

# 8. CI/CD
doc.add_heading('8. CI/CD Model', 1)
add_table(doc,
    ['Environment', 'Deploy Trigger', 'Gate', 'Rollback'],
    [
        ['Dev', 'Auto on push', 'None', 'Revert commit'],
        ['QA', 'Gated — manual trigger', 'Schema check, pipeline health, embedding version consistency, smoke tests', 'Required'],
        ['Prod', 'Manual approval by Skip', 'All QA gates + Skip sign-off', 'Required'],
    ]
)
doc.add_paragraph('Mandatory checks before any prod release:')
bullets(doc, [
    'Schema compatibility',
    'Pipeline health',
    'Embedding version consistency',
    'Endpoint smoke tests',
])

# 9. Data Promotion
doc.add_heading('9. Data Promotion', 1)
bullets(doc, [
    'No direct promotion from dev to prod',
    'Pipelines must be re-run in each target environment',
    'Automated regression testing required before any promotion',
    'GPT owns regression test design and gate criteria',
])

# 10. Schema Evolution
doc.add_heading('10. Schema Evolution', 1)
bullets(doc, [
    'No breaking changes in prod',
    'All schema changes must be additive',
    'Schema versions tracked and enforced',
])

# 11. Persistence Policy
doc.add_heading('11. Persistence Policy', 1)
add_table(doc,
    ['Environment', 'Normal Chat', 'Safety Incident'],
    [
        ['Dev', 'Full prompt + response (debug)', 'IP + timestamp'],
        ['QA', 'Full prompt + response (debug)', 'IP + timestamp'],
        ['Prod', 'HIPAA two-tier consent flow', 'IP + timestamp (legal review pending)'],
    ]
)
doc.add_paragraph('Prod persistence requires: request_id, timestamp, model_version, audit fields.')
doc.add_paragraph('Controlled by PERSIST_MODE env var: debug (dev/QA) or hipaa (prod).')

# 12. Observability
doc.add_heading('12. Observability', 1)
bullets(doc, [
    'Structured logging across all tiers',
    'Correlation IDs — trace a request end to end',
    'Metrics: latency, error rate, LLM cost per request',
    'Alerts on error rate thresholds and latency SLOs',
])

# 13. Machine Brain
doc.add_heading('13. Machine Brain', 1)
doc.add_paragraph('Purpose: Persist architectural decisions, patterns, and knowledge across sessions and agents.')
doc.add_heading('Collections', 2)
add_table(doc,
    ['Collection', 'Purpose'],
    [
        ['ArchitectureMemory', 'Long-term architectural knowledge'],
        ['Decisions', 'Formal decision records with risk level and rationale'],
        ['Patterns', 'Reusable implementation patterns'],
    ]
)
doc.add_heading('Decision Record Schema', 2)
add_code(doc,
    '{\n'
    '  "topic": string,\n'
    '  "type": string,\n'
    '  "decision": string,\n'
    '  "rationale": string,\n'
    '  "risk": "Low|Moderate|High|Critical|Suicidal",\n'
    '  "constraints": [string],\n'
    '  "components": [string],\n'
    '  "created_by": "GPT|Claude|Skip",\n'
    '  "timestamp": "UTC"\n'
    '}'
)
doc.add_heading('Retrieval Contract', 2)
bullets(doc, [
    'Claude MUST query Machine Brain before implementation',
    'GPT writes to Machine Brain after architectural decisions',
    'All decisions assigned a risk level per Section 3',
])

# 14. Knowledge System
doc.add_heading('14. Knowledge System', 1)
add_table(doc,
    ['Tier', 'Storage', 'Purpose'],
    [
        ['Static', 'Versioned docs in repo (Website/, Analysis/)', 'Authoritative design documents'],
        ['Working', 'claude_thoughts/ and gpt_thoughts/ directories', 'Active collaboration and handoff between agents'],
        ['Machine Brain', 'MongoDB — persistent memory', 'Long-term architectural knowledge across sessions'],
    ]
)

# 15. Bot Collaboration Protocol
doc.add_heading('15. Bot Collaboration Protocol', 1)
doc.add_heading('Directory Structure', 2)
add_code(doc,
    'claude_thoughts/    <- Claude drops questions, decisions, proposals for GPT\n'
    'gpt_thoughts/       <- GPT drops architecture feedback, concerns, approvals'
)
doc.add_heading('Flow', 2)
bullets(doc, [
    'Claude writes to claude_thoughts/ when work needs GPT review',
    'GPT writes to gpt_thoughts/ in response',
    'Skip relays between sessions (near term)',
    'Cron-driven async handoff — stretch goal',
])
doc.add_paragraph('File naming convention: YYYY-MM-DD_[topic].md')

# 16. Cost Control
doc.add_heading('16. Cost Control', 1)
bullets(doc, [
    'Embedding versioning REQUIRED — version stored with every record',
    'No re-embedding unless version mismatch detected',
    'LLM cost tracked per request in observability layer',
])

# 17. Operating Principle
doc.add_heading('17. Operating Principle', 1)
p = doc.add_paragraph(
    'This is a production healthcare system. '
    'Safety, auditability, and determinism are required at every layer. No shortcuts in prod.'
)
for run in p.runs:
    run.bold = True

doc.save('c:/chatHealthy/findCare/Analysis/ChatHealthy-DevOps-Architecture-2026-03-25.docx')
print('Done')
