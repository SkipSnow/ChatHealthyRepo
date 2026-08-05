"""Update Provider Pipeline LLD v45 - Add pipelineEditor MongoDB prerequisite.

Adds proper subsection to Prerequisites, updates TOC, version, green tracked changes.
"""
from docx import Document
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
import subprocess

INPUT_FILE = r"c:\CHATHealthyLLC\pipeline\ArchitectureDesignAndAudit\ProviderPipeline_LowLevelDesign_v44.docx"
OUTPUT_FILE = r"c:\CHATHealthyLLC\pipeline\ArchitectureDesignAndAudit\ProviderPipeline_LowLevelDesign_v45.docx"

doc = Document(INPUT_FILE)

# Step 1: Update version on cover page
for para in doc.paragraphs[:20]:
    if para.text.startswith("Version: 41"):
        para.text = "Version: 45"
        print("✓ Updated cover page version to 45")
        break

# Step 2: Find Prerequisites section and add subsection
section_index = None
for i, para in enumerate(doc.paragraphs):
    if para.text.strip() == "3.2 Prerequisites Not Managed by the Deployment Chain":
        section_index = i
        break

if section_index:
    # Find where to insert (before next Heading 1 or Heading 2 of next section)
    insert_index = section_index + 1
    for i in range(section_index + 1, len(doc.paragraphs)):
        para = doc.paragraphs[i]
        if para.style.name in ['Heading 1', 'Heading 2']:
            if '4.' in para.text or 'Topology' in para.text:
                insert_index = i
                break

    # Create Heading 3 for MongoDB subsection
    parent = doc.paragraphs[insert_index]._element

    # New Heading 3
    heading_xml = f'''<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                        xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml">
                      <w:pPr>
                        <w:pStyle w:val="Heading3"/>
                      </w:pPr>
                      <w:r>
                        <w:t>MongoDB X.509 User Setup</w:t>
                      </w:r>
                    </w:p>'''
    heading_elem = parse_xml(heading_xml)
    parent.addprevious(heading_elem)

    # Content with green color (tracked change)
    content = [
        "Before deploying the pipelineEditor identity certificates, create the corresponding X.509 user in MongoDB on the front-end cluster:",
        "",
        "db.getSiblingDB(\"$external\").createUser({",
        "  user: \"pipelineEditor\",",
        "  roles: [",
        "    { role: \"readWrite\", db: \"chathealthyfrontend\" }",
        "  ]",
        "})",
        "",
        "This creates user pipelineEditor for X.509 authentication with read/write access to the chathealthyfrontend database. MongoDB automatically maps the certificate's CN=pipelineEditor to this user at connection time."
    ]

    for line in content:
        if line == "":
            # Empty paragraph
            para_xml = f'''<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'''
        else:
            # Paragraph with green text
            para_xml = f'''<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                              xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml">
                            <w:r>
                              <w:rPr>
                                <w:color w:val="00B050"/>
                              </w:rPr>
                              <w:t xml:space="preserve">{line}</w:t>
                            </w:r>
                          </w:p>'''
        para_elem = parse_xml(para_xml)
        parent.addprevious(para_elem)

    print("✓ Added MongoDB User Setup subsection (Heading 3) in green")

# Step 3: Update TOC field
# Find TOC field and update it
toc_updated = False
for para in doc.paragraphs:
    if 'TOC' in para.text:
        # Found TOC paragraph, update the field
        for run in para.runs:
            if 'TOC' in run.text:
                # Mark for update
                toc_updated = True
                print("✓ Located TOC field")
                break

# Step 4: Save document
doc.save(OUTPUT_FILE)
print(f"✓ Saved to {OUTPUT_FILE}")

# Step 5: Use Word via COM to update TOC (if on Windows with Word installed)
try:
    import win32com.client
    word = win32com.client.Dispatch("Word.Application")
    doc_com = word.Documents.Open(OUTPUT_FILE)

    # Update all fields in the document
    for field in doc_com.Fields:
        field.Update()

    doc_com.Save()
    doc_com.Close()
    word.Quit()
    print("✓ TOC updated via Word COM interface")
except Exception as e:
    print(f"⚠ Word COM update failed (expected if Word not installed): {e}")
    print("  Manual step: Open document and right-click TOC → Update Field")

print("\n✅ Document v45 ready for review")
