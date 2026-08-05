"""Update Provider Pipeline LLD with pipelineEditor MongoDB prerequisite.

Adds tracked change (green) for MongoDB user setup in Prerequisites section.
Updates version number on cover page.
"""
from docx import Document
from docx.enum.dml import MSO_THEME_COLOR
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Pt, RGBColor
from datetime import datetime
import shutil

INPUT_FILE = r"c:\CHATHealthyLLC\pipeline\ArchitectureDesignAndAudit\ProviderPipeline_LowLevelDesign_v44.docx"
OUTPUT_FILE = r"c:\CHATHealthyLLC\pipeline\ArchitectureDesignAndAudit\ProviderPipeline_LowLevelDesign_v45.docx"

doc = Document(INPUT_FILE)

# Step 1: Update version on cover page
for para in doc.paragraphs[:20]:
    if para.text.startswith("Version: 41"):
        para.text = "Version: 45"
        print("✓ Updated cover page version to 45")
        break

# Step 2: Find Prerequisites section and add MongoDB user setup
found_prerequisites = False
mongodb_text = "MongoDB User Setup\n\nBefore deploying the pipelineEditor identity certificates, create the corresponding X.509 user in MongoDB:\n\nConnect to the front-end cluster and execute:\n\ndb.getSiblingDB(\"$external\").createUser({\n  user: \"pipelineEditor\",\n  roles: [\n    { role: \"readWrite\", db: \"chathealthyfrontend\" }\n  ]\n})\n\nThis creates user pipelineEditor for X.509 authentication with read/write access to the chathealthyfrontend database. MongoDB automatically maps the certificate's CN=pipelineEditor to this user at connection time."

# Find the Prerequisites section
for i, para in enumerate(doc.paragraphs):
    if para.text.strip() == "3.2 Prerequisites Not Managed by the Deployment Chain":
        # Found it - add new paragraph after this one
        # Insert after next paragraph (skip the heading)
        insert_index = i + 2

        # Create paragraph with tracked change (green)
        new_para = doc.paragraphs[insert_index]._element

        # Insert new paragraph
        p = OxmlElement('w:p')

        # Add MongoDB text as tracked insertion (green)
        for line in mongodb_text.split('\n'):
            run_elem = OxmlElement('w:r')

            # Add tracked change wrapper
            ins = OxmlElement('w:ins')
            ins.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id', '0')
            ins.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author', 'Claude Code')
            ins.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}date', datetime.now().isoformat())

            # Create text element
            t = OxmlElement('w:t')
            t.set('{http://schemas.openxmlformats.org/xml/space}space', 'preserve')
            t.text = line + '\n'

            # Color text green
            rPr = OxmlElement('w:rPr')
            color = OxmlElement('w:color')
            color.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '00B050')  # Green
            rPr.append(color)

            run_elem.append(rPr)
            run_elem.append(t)
            ins.append(run_elem)
            p.append(ins)

        # Insert the new paragraph into the document
        new_para.addprevious(p)
        found_prerequisites = True
        print("✓ Added MongoDB prerequisite with tracked changes (green)")
        break

if not found_prerequisites:
    print("✗ Prerequisites section not found")

# Save with new version
doc.save(OUTPUT_FILE)
print(f"✓ Saved to {OUTPUT_FILE}")
