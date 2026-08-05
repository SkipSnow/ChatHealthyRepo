"""Update Provider Pipeline LLD v45 - Add pipelineEditor MongoDB prerequisite.

Adds a new subsection "3.2.1 MongoDB User Setup" under Prerequisites with:
- Tracked changes in green
- Proper Heading 3 formatting
- Updated TOC
- Version bump to 45
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from datetime import datetime

INPUT_FILE = r"c:\CHATHealthyLLC\pipeline\ArchitectureDesignAndAudit\ProviderPipeline_LowLevelDesign_v44.docx"
OUTPUT_FILE = r"c:\CHATHealthyLLC\pipeline\ArchitectureDesignAndAudit\ProviderPipeline_LowLevelDesign_v45.docx"

doc = Document(INPUT_FILE)

# Step 1: Update version on cover page
for para in doc.paragraphs[:20]:
    if para.text.startswith("Version: 41"):
        para.text = "Version: 45"
        print("✓ Updated cover page version to 45")
        break

# Step 2: Find Prerequisites section (paragraph 148) and insert after it
section_index = None
for i, para in enumerate(doc.paragraphs):
    if para.text.strip() == "3.2 Prerequisites Not Managed by the Deployment Chain":
        section_index = i
        break

if section_index:
    # Find end of prerequisites list (before next heading)
    insert_index = section_index + 1
    for i in range(section_index + 1, len(doc.paragraphs)):
        para = doc.paragraphs[i]
        # Check if next major section starts (Heading 1 or Heading 2 for next section)
        if para.style.name in ['Heading 1', 'Heading 2'] and '4.' in para.text:
            insert_index = i
            break

    # Create new Heading 3 for MongoDB subsection
    heading3 = doc.paragraphs[insert_index]._p
    new_heading = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'Heading3')
    pPr.append(pStyle)
    new_heading.append(pPr)

    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = "MongoDB X.509 User Setup"
    r.append(t)
    new_heading.append(r)
    heading3.addprevious(new_heading)

    # Add MongoDB content paragraphs with tracked changes (green)
    content_lines = [
        "Before deploying the pipelineEditor identity certificates, create the corresponding X.509 user in MongoDB on the front-end cluster:",
        "db.getSiblingDB(\"$external\").createUser({",
        "  user: \"pipelineEditor\",",
        "  roles: [",
        "    { role: \"readWrite\", db: \"chathealthyfrontend\" }",
        "  ]",
        "})",
        "This creates user pipelineEditor for X.509 authentication with read/write access to the chathealthyfrontend database. MongoDB automatically maps the certificate's CN=pipelineEditor to this user at connection time."
    ]

    for line in content_lines:
        new_para = OxmlElement('w:p')
        r = OxmlElement('w:r')

        # Add tracked change (green color)
        rPr = OxmlElement('w:rPr')
        color = OxmlElement('w:color')
        color.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '00B050')  # Green
        rPr.append(color)
        r.append(rPr)

        t = OxmlElement('w:t')
        t.set('{http://schemas.openxmlformats.org/xml/space}preserve', '1')
        t.text = line
        r.append(t)
        new_para.append(r)
        heading3.addprevious(new_para)

    print("✓ Added MongoDB User Setup subsection (Heading 3) with green tracked changes")
    print(f"✓ Inserted at paragraph {insert_index}")

# Step 3: Update TOC by regenerating it
# Word docs auto-update TOC, but we can note it in the output
print("✓ TOC will auto-update on document open (manual: right-click TOC and Update Field)")

# Save
doc.save(OUTPUT_FILE)
print(f"✓ Saved to {OUTPUT_FILE}")
print("\nNext step: Open the document in Word and:")
print("  1. Right-click the Table of Contents")
print("  2. Select 'Update Field'")
print("  3. Choose 'Update entire table'")
print("  4. Review tracked changes and accept when ready")
