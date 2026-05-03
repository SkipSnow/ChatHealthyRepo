# One-shot generator for testing_requirements_audit_2026-05-03.docx
# Skip can re-run if needed. Read-only on every other path.

import json
import os
from collections import defaultdict
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

REPO = r"c:\CHATHealthyLLC"
BACKLOG = os.path.join(REPO, "brain", "machine_artifacts", "content", "agile_backlog.json")
OUT_DIR = os.path.join(REPO, "architecture", "Testing", "ArchitectureDesignAndAuditDocs")
OUT_PATH = os.path.join(OUT_DIR, "testing_requirements_audit_2026-05-03.docx")

os.makedirs(OUT_DIR, exist_ok=True)

with open(BACKLOG, "r", encoding="utf-8") as f:
    bk = json.load(f)

# Find EPIC-005
epic5 = None
for ep in bk.get("epics", {}).get("epic", []):
    if ep.get("epic_id") == "EPIC-005":
        epic5 = ep
        break
assert epic5 is not None, "EPIC-005 not found"

# Walk all requirements and capture context.
records = []  # list of dicts
for feat in epic5["features"]["feature"]:
    fid = feat["feature_id"]
    fname = feat["name"]
    for story in feat["stories"]["story"]:
        sid = story["story_id"]
        sname = story.get("name", "")
        sstatus = story.get("status", "")
        sdesc = story.get("description", "")
        for req in story["requirements"]["requirement"]:
            pytests = req.get("pytests", {}).get("pytest", [])
            files = [pt.get("file", "") for pt in pytests]
            ptypes = [pt.get("pytest_type", "") for pt in pytests]
            records.append({
                "fid": fid,
                "fname": fname,
                "sid": sid,
                "sname": sname,
                "sstatus": sstatus,
                "sdesc": sdesc,
                "req_id": req["req_id"],
                "name": req.get("name", ""),
                "text": req.get("requirement", ""),
                "status": req.get("status", ""),
                "approval": req.get("approval", ""),
                "pytest_files": files,
                "pytest_types": ptypes,
            })

total = len(records)

# --- Classification ---
# Each record may receive multiple labels.

CEREMONIAL = []
UNENFORCEABLE = []
ORPHANED_IMPL = []
DUPLICATE = []
META_CIRCULAR = []
STALE_DONE = []
DELETED_FEATURE_RESIDUE = []
CONTRADICTORY = []

# Pre-compute groups for duplicate detection
by_story = defaultdict(list)
for r in records:
    by_story[r["sid"]].append(r)


def all_orphan(r):
    return all(t == "ORPHAN" for t in r["pytest_types"]) if r["pytest_types"] else True


# --- F-001 S-005: Loading message + retry ---
# These are clearly UI behaviour, not really "Testing" epic; both ORPHAN; ceremonial in this epic.
for r in records:
    if r["sid"] == "EPIC-005-F-001-S-005":
        DELETED_FEATURE_RESIDUE.append((r,
            "Lives under Testing epic but is a FindCare runtime UX requirement (loading indicator + 10-min retry). Belongs in EPIC-006 (FindCare) or EPIC-009 (UX), not Testing. Both are ORPHAN with proposed approval."))

# --- F-001 S-007: every test in test_pagination.py needs Playwright equiv ---
# test_pagination.py exists with 54 unittest tests; no Playwright runner committed.
# This is meta-circular ("tests must exist for tests"). ORPHAN.
for r in records:
    if r["sid"] == "EPIC-005-F-001-S-007":
        META_CIRCULAR.append((r,
            "Meta-circular: the requirement governs OTHER tests, but specifies nothing testable about its own pytest. test_pagination.py has 54 unittest methods — none Playwrighted, no Playwright equivalent file exists (ORPHAN). If the goal is 'pagination must work in a real browser', that requirement belongs in FindCare epic and references concrete pagination behaviours, not 'every test in file X must be cloned'."))

# --- F-002 S-001: backlog-ordered regression runner (8 reqs, all ORPHAN, status not_started) ---
# These describe a hypothetical tool. Not ceremonial — they would be enforceable IF the tool
# were built — but today every one is ORPHAN and the tool itself does not exist.
# Recommendation: investigate / collapse, not delete blindly. Mark as ORPHANED_IMPL.
for r in records:
    if r["sid"] == "EPIC-005-F-002-S-001":
        ORPHANED_IMPL.append((r,
            "Specifies a regression-runner tool that does not exist in the repo. All 8 reqs are ORPHAN/not_started. Either build the runner (and have ONE pytest verify it end-to-end) or delete the story. 8 separate must-haves for a single CLI script is over-decomposition."))

# Within S-001 there are duplicates: REQ-T-001 and REQ-T-002 both describe scope ordering/args;
# REQ-T-003 and REQ-T-005 both describe results reporting (console vs JSON).
DUPLICATE.append((
    {"req_id": "EPIC-005-F-002-S-001-REQ-T-003 / REQ-T-005",
     "text": "Per-requirement reporting (T-003) and machine-readable JSON report (T-005)"},
    "Two reqs cover the same surface (output format). Collapse to one: 'Report results per requirement and emit JSON + console.'"))
DUPLICATE.append((
    {"req_id": "EPIC-005-F-002-S-001-REQ-T-001 / REQ-T-002",
     "text": "Hierarchy ordering (T-001) and scope args (T-002)"},
    "Both describe the runner's invocation contract. Consider collapsing into one CLI-contract requirement with sub-bullets, not two must-haves."))

# --- F-002 S-002: load parity ---
# test_load_parity.py EXISTS and clearly maps to T-001 and T-002, but the backlog entries are
# still tagged ORPHAN. Status is "not_started" on every req but the story is "in_progress".
for r in records:
    if r["sid"] == "EPIC-005-F-002-S-002":
        STALE_DONE.append((r,
            "Test code exists at Code/DataPipelines/tests/test_load_parity.py (TestLoadParity + TestFrontendParity classes) and matches the four requirements. Backlog still says ORPHAN/not_started — backlog is stale, not the code. Update pytest_id refs and flip status, OR delete REQ-T-003 (frontend cluster parity) and REQ-T-004 (provider_quality parity) if scope was reduced; both are claimed by the test file — verify."))

# --- F-003 S-001: TLS between local services (2 reqs, both ORPHAN) ---
# REQ-T-002 reads "https://localhost redirects to https://localhost" — that's a no-op,
# CONTRADICTORY/incoherent.
for r in records:
    if r["req_id"] == "EPIC-005-F-003-S-001-REQ-T-002":
        CONTRADICTORY.append((r,
            "Text 'https://localhost redirects to https://localhost' is incoherent — a URL cannot redirect to itself. Almost certainly a typo for 'http://localhost redirects to https://localhost', which is now redundant with EPIC-008-F-011-S-002-REQ-B-002 (HTTP→HTTPS 301 redirect, covered by httpMethodEnforcementPyTest.py added 2026-05-02). Delete or merge."))
    if r["req_id"] == "EPIC-005-F-003-S-001-REQ-T-001":
        DUPLICATE.append((r,
            "'All inter-service communication uses TLS' overlaps with EPIC-008-F-011-S-002 HTTPS-only suite + test_https_security.py. Either delete or restate as a fan-in pytest reference; do not keep ORPHAN."))

# --- F-003 S-002: architecture/roadmap pages on :443 only ---
for r in records:
    if r["sid"] == "EPIC-005-F-003-S-002":
        DUPLICATE.append((r,
            "Subsumed by EPIC-008-F-011-S-002-REQ-B-003 (static pages GET/HEAD only) plus the HTTP→HTTPS redirect rule, both covered by httpMethodEnforcementPyTest.py. Architecture/roadmap pages are not specially privileged. Delete or fold into EPIC-008."))

# --- F-004 S-001..S-007: Data-quality regression family (~25 reqs, every one ORPHAN) ---
# Several test files exist (test_data_quality.py, test_quality_gates.py, test_schema_drift.py,
# test_pipeline_compliance.py) but the backlog cites ORPHAN. Either tests cover the rules or
# they don't — but nothing has a pytest_id wired here. These are pipeline data-quality rules,
# arguably belonging under the DataPipelines epic, not "Testing".
for r in records:
    if r["fid"] == "EPIC-005-F-004" and r["sid"] != "EPIC-005-F-004-S-008":
        if all_orphan(r):
            ORPHANED_IMPL.append((r,
                "Pipeline data-quality rule with ORPHAN pytest. Likely covered (or coverable) by existing test files in Code/DataPipelines/tests/ (test_data_quality.py, test_quality_gates.py, test_schema_drift.py, test_compliance.py). Either wire the pytest_id or move the rule into EPIC-007 (Data Pipeline) where DQ belongs and delete the F-004 story shells."))

# CEREMONIAL examples within F-004:
# REQ-T-002 "NPI must be 10-digit string" — already a pipeline ingest constraint enforced upstream;
# making it a TEST EPIC requirement is documentation, not testing.
# REQ-T-003 "entity_type_code must be '1' or '2'" — same.
# REQ-T-004 "every provider must have at least one taxonomy code" — same.
# REQ-T-005 "no duplicate NPIs" — same.
for rid in ["EPIC-005-F-004-S-001-REQ-T-002",
            "EPIC-005-F-004-S-001-REQ-T-003",
            "EPIC-005-F-004-S-001-REQ-T-004",
            "EPIC-005-F-004-S-001-REQ-T-005"]:
    for r in records:
        if r["req_id"] == rid:
            CEREMONIAL.append((r,
                "Schema-level invariant; belongs in the schema definition (or pipeline ingest validator), not as a separate Testing-epic requirement. Delete from EPIC-005; if not already enforced, add to schema or pipeline write validator."))

# CV vocab reqs in S-002 and S-007 overlap heavily — flag duplicate.
DUPLICATE.append((
    {"req_id": "EPIC-005-F-004-S-002-REQ-T-001 + EPIC-005-F-004-S-007-REQ-T-001",
     "text": "CV-001 bad_data vocabulary requirements"},
    "Same vocabulary surface stated twice (one as 'values must be from CV-001', the other as 'CV-001 must exist with required members'). Collapse."))
DUPLICATE.append((
    {"req_id": "EPIC-005-F-004-S-002-REQ-T-002 + EPIC-005-F-004-S-007-REQ-T-002",
     "text": "CV-002 out_of_scope vocabulary requirements"},
    "Same duplication pattern as CV-001. Collapse."))

# REQ-T-003 in S-007 ("State filter must be mandatory") is a pipeline runtime requirement,
# not a vocabulary/structure rule — misfiled.
for r in records:
    if r["req_id"] == "EPIC-005-F-004-S-007-REQ-T-003":
        CEREMONIAL.append((r,
            "'State filter must be mandatory' is a pipeline runtime contract, not a controlled-vocabulary structure rule and not a Testing-epic concern. Move to pipeline epic or delete."))

# REQ-T-004 in S-007 ("vocabularies must have invoked_by traceability") — schema/governance,
# not testing.
for r in records:
    if r["req_id"] == "EPIC-005-F-004-S-007-REQ-T-004":
        CEREMONIAL.append((r,
            "Traceability metadata is a schema concern. Either schema enforces it (and a single schema-validator test covers all CVs) or it isn't enforced. Standalone Testing-epic requirement adds no enforcement."))

# --- F-004 S-008: Pipeline Compliance Report (DONE, all 9 reqs implemented by test_dq_compliance_report.py) ---
# The eight T-001..T-008 reqs are paraphrases of the same single deliverable: a report with hierarchy
# walking, orphan booleans, schema validation, and one output table. They are dramatically over-
# decomposed. Story is "done" so deletion candidates are weaker — but several reqs are paraphrases.
for rid in ["EPIC-005-F-004-S-008-REQ-T-002",
            "EPIC-005-F-004-S-008-REQ-T-003"]:
    for r in records:
        if r["req_id"] == rid:
            DUPLICATE.append((r,
                "Definition of 'actionable orphan' (T-002) and the EPIC-ORPHAN container rule (T-003) are implementation notes embedded as requirements. They restate the orphan-counting algorithm twice. Collapse into a single 'orphan accounting' requirement under the report deliverable."))

for rid in ["EPIC-005-F-004-S-008-REQ-T-005",
            "EPIC-005-F-004-S-008-REQ-T-006"]:
    for r in records:
        if r["req_id"] == rid:
            CEREMONIAL.append((r,
                "Schema shape requirements ('every node has orphan boolean'; 'pytest is child of requirement'). These are schema concerns — should be enforced by schema_validator and exist as one schema-test, not as Testing-epic requirements. The schema, not the report, is the source of truth."))

# Verify QA-Report cleanup completeness — already grepped, no QA-Report residue in backlog.
# This is reported in Methodology, not as a finding.

# --- Build doc ---
doc = Document()

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10)

# Title page
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = t.add_run("Testing Requirements Tree Audit\n")
run.bold = True
run.font.size = Pt(22)
run2 = t.add_run("EPIC-005 — agile_backlog.json\n\n")
run2.font.size = Pt(14)
run3 = t.add_run("2026-05-03\n")
run3.font.size = Pt(12)
run4 = t.add_run("Claude (Opus 4.7) for Skip Snow")
run4.font.size = Pt(12)
doc.add_page_break()

# --- Executive Summary ---
doc.add_heading("Executive Summary", level=1)

# class counts (records-level only; some entries are pair-style and counted once)
n_cer = len(CEREMONIAL)
n_un = len(UNENFORCEABLE)
n_orph_impl = len(ORPHANED_IMPL)
n_dup = len(DUPLICATE)
n_meta = len(META_CIRCULAR)
n_stale = len(STALE_DONE)
n_dfr = len(DELETED_FEATURE_RESIDUE)
n_contra = len(CONTRADICTORY)

summary_table = doc.add_table(rows=1, cols=2)
summary_table.style = "Light Grid Accent 1"
hdr = summary_table.rows[0].cells
hdr[0].text = "Metric"
hdr[1].text = "Count"

def row(k, v):
    r = summary_table.add_row().cells
    r[0].text = k
    r[1].text = str(v)

row("Total EPIC-005 requirements", total)
row("Stories", len({r["sid"] for r in records}))
row("Features", len({r["fid"] for r in records}))
row("Reqs with all-ORPHAN pytests", sum(1 for r in records if all_orphan(r)))
row("Reqs with at least one real pytest_id", sum(1 for r in records if not all_orphan(r)))
row("Class: CEREMONIAL", n_cer)
row("Class: UNENFORCEABLE", n_un)
row("Class: ORPHANED-IMPLEMENTATION", n_orph_impl)
row("Class: DUPLICATE (incl. cross-req pairs)", n_dup)
row("Class: META-CIRCULAR", n_meta)
row("Class: STALE-DONE", n_stale)
row("Class: DELETED-FEATURE-RESIDUE", n_dfr)
row("Class: CONTRADICTORY", n_contra)

# percent deletable: high-confidence buckets
high_conf_req_ids = set()
for r, _ in CEREMONIAL: high_conf_req_ids.add(r["req_id"])
for r, _ in DELETED_FEATURE_RESIDUE: high_conf_req_ids.add(r["req_id"])
for r, _ in META_CIRCULAR: high_conf_req_ids.add(r["req_id"])
# Add unambiguous duplicates that name a single req_id (skip 'A + B' aggregate pairs)
for r, _ in DUPLICATE:
    if "/" not in r["req_id"] and "+" not in r["req_id"]:
        high_conf_req_ids.add(r["req_id"])

pct_high = round(100.0 * len(high_conf_req_ids) / total, 1)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run(f"High-confidence delete candidates: {len(high_conf_req_ids)} of {total} ({pct_high}%).").bold = True

doc.add_paragraph(
    "Assessment: EPIC-005 is heavily ceremonial. Of "
    f"{total} requirements, only "
    f"{sum(1 for r in records if not all_orphan(r))} carry a real pytest_id, and the only story "
    "with full test wiring is the Pipeline Compliance Report (F-004-S-008), which is itself "
    "over-decomposed (eight requirements describe one report). The data-quality story family "
    "(F-004-S-001 through S-007) is pipeline-domain content misfiled under Testing — its rules "
    "either belong in the schema/pipeline validators (already partially enforced) or have ORPHAN "
    "pytests. The regression-runner story (F-002-S-001) describes a tool that does not exist in "
    "the repo. The TLS/HTTPS stories (F-003) are now subsumed by the new "
    "httpMethodEnforcementPyTest.py + existing test_https_security.py under EPIC-008. "
    "Recommendation: collapse EPIC-005 to (a) the regression runner (1 story, 1-2 reqs, real test), "
    "(b) the pipeline compliance report (already done, deduped), and (c) ONE Playwright "
    "pagination story owned by the FindCare epic — and delete the rest."
)

doc.add_page_break()

# --- Findings by Class ---
doc.add_heading("Findings by Class", level=1)

def emit_class(title, items, intro):
    doc.add_heading(title, level=2)
    if intro:
        doc.add_paragraph(intro)
    if not items:
        doc.add_paragraph("(none)")
        return
    for r, why in items:
        h = doc.add_paragraph()
        h.add_run(r["req_id"]).bold = True
        # If r is a real record, include name + status + pytest evidence
        if "fid" in r:
            doc.add_paragraph(f"Feature: {r['fid']} — {r['fname']}")
            doc.add_paragraph(f"Story: {r['sid']} — {r['sname']} (story status: {r['sstatus']})")
            doc.add_paragraph(f"Name: {r['name']}")
            doc.add_paragraph(f"Text: {r['text']}")
            doc.add_paragraph(f"Status / Approval: {r['status']} / {r['approval']}")
            doc.add_paragraph(
                "Pytest evidence: "
                + ("; ".join(f"[{t}] {f}" for t, f in zip(r["pytest_types"], r["pytest_files"]))
                   or "(none)")
            )
        else:
            doc.add_paragraph(f"Aggregate: {r.get('text','')}")
        rec = doc.add_paragraph()
        rec.add_run("Why & action: ").bold = True
        rec.add_run(why)
        doc.add_paragraph()  # spacer

emit_class(
    "CEREMONIAL",
    CEREMONIAL,
    "Documentation statements with no enforcement mechanism in the Testing epic — "
    "either schema-level invariants masquerading as test requirements, or implementation notes "
    "embedded as separate must-haves."
)
emit_class(
    "UNENFORCEABLE",
    UNENFORCEABLE,
    "No automated check could ever verify the requirement as written."
)
emit_class(
    "ORPHANED-IMPLEMENTATION",
    ORPHANED_IMPL,
    "Requirement claims to govern behaviour but no test code exists. Every entry here is "
    "ORPHAN-only at the pytest level."
)
emit_class(
    "DUPLICATE",
    DUPLICATE,
    "Paraphrased equivalents of another requirement, either inside the same story or across "
    "epics. Includes cross-requirement pairs (req_id shows 'A + B')."
)
emit_class(
    "CONTRADICTORY",
    CONTRADICTORY,
    "Mutually exclusive or self-contradictory text."
)
emit_class(
    "STALE-DONE",
    STALE_DONE,
    "Status doesn't match reality: code exists that satisfies (or refutes) the requirement, but "
    "the backlog has not been updated."
)
emit_class(
    "DELETED-FEATURE-RESIDUE",
    DELETED_FEATURE_RESIDUE,
    "Requirements that survive after the surrounding feature was reassigned/removed. Includes "
    "items misfiled under EPIC-005 that belong elsewhere."
)
emit_class(
    "META-CIRCULAR",
    META_CIRCULAR,
    "'Tests must exist for X' without specifying anything testable about the test itself."
)

doc.add_page_break()

# --- Deletion Candidates Ranked ---
doc.add_heading("Deletion Candidates Ranked", level=1)

doc.add_heading("High Confidence — DELETE", level=2)
doc.add_paragraph(
    "These can be removed from agile_backlog.json without further investigation. "
    "Copy/paste-ready req_id list:"
)
hc_table = doc.add_table(rows=1, cols=3)
hc_table.style = "Light Grid Accent 1"
hh = hc_table.rows[0].cells
hh[0].text = "req_id"
hh[1].text = "Class"
hh[2].text = "One-liner"

def class_of(rid):
    cls = []
    for r, _ in CEREMONIAL:
        if r["req_id"] == rid:
            cls.append("CEREMONIAL")
    for r, _ in DELETED_FEATURE_RESIDUE:
        if r["req_id"] == rid:
            cls.append("DELETED-FEATURE-RESIDUE")
    for r, _ in META_CIRCULAR:
        if r["req_id"] == rid:
            cls.append("META-CIRCULAR")
    for r, _ in DUPLICATE:
        if r["req_id"] == rid:
            cls.append("DUPLICATE")
    return ",".join(cls) or "?"

def oneliner(rid):
    for bucket in (CEREMONIAL, DELETED_FEATURE_RESIDUE, META_CIRCULAR, DUPLICATE,
                   STALE_DONE, ORPHANED_IMPL, CONTRADICTORY):
        for r, why in bucket:
            if r["req_id"] == rid:
                return why[:240]
    return ""

for rid in sorted(high_conf_req_ids):
    rr = hc_table.add_row().cells
    rr[0].text = rid
    rr[1].text = class_of(rid)
    rr[2].text = oneliner(rid)

doc.add_heading("Medium Confidence — DEDUPE / RESOLVE", level=2)
doc.add_paragraph(
    "Aggregate duplicates (cross-requirement) and STALE-DONE: collapse or update."
)
mc_table = doc.add_table(rows=1, cols=2)
mc_table.style = "Light Grid Accent 1"
mc_table.rows[0].cells[0].text = "req_id (or pair)"
mc_table.rows[0].cells[1].text = "Action"

agg_dups = [(r, why) for r, why in DUPLICATE if "/" in r["req_id"] or "+" in r["req_id"]]
for r, why in agg_dups:
    rr = mc_table.add_row().cells
    rr[0].text = r["req_id"]
    rr[1].text = why[:240]
for r, why in STALE_DONE:
    rr = mc_table.add_row().cells
    rr[0].text = r["req_id"]
    rr[1].text = "STALE-DONE — " + why[:240]
for r, why in CONTRADICTORY:
    rr = mc_table.add_row().cells
    rr[0].text = r["req_id"]
    rr[1].text = "CONTRADICTORY — " + why[:240]

doc.add_heading("Investigate", level=2)
doc.add_paragraph(
    "ORPHANED-IMPLEMENTATION: tool/test does not exist. Decide BUILD vs DELETE per story; "
    "do not leave in this state."
)
inv_table = doc.add_table(rows=1, cols=2)
inv_table.style = "Light Grid Accent 1"
inv_table.rows[0].cells[0].text = "req_id"
inv_table.rows[0].cells[1].text = "Story / decision"

for r, why in ORPHANED_IMPL:
    rr = inv_table.add_row().cells
    rr[0].text = r["req_id"]
    rr[1].text = why[:240]

doc.add_page_break()

# --- Methodology + Caveats ---
doc.add_heading("Methodology and Caveats", level=1)

doc.add_paragraph(
    "Method:"
)
for line in [
    "1. Loaded agile_backlog.json and isolated EPIC-005 (lines ~1121-2446).",
    "2. Walked every feature/story/requirement; recorded pytest_type and file for each.",
    "3. For each requirement, checked: is the cited test file present in the repo? "
       "Does the file's actual content satisfy the requirement text? Does another epic already "
       "cover the same surface?",
    "4. Compared sibling requirements within each story for paraphrase / overlap.",
    "5. Compared EPIC-005 surface to existing test code under Code/deploy/, "
       "Code/ConversationalUX/FindCareChat/backend/tests/, Code/DataPipelines/tests/, "
       "Code/Shared/tests/, evaluateCare/Code/, sharedServices/Code/.",
    "6. Confirmed QA-Report client-display feature cleanup: grep on agile_backlog.json for "
       "'QA-Report|QAReport|qa_report|qa-report' returned no matches. Cleanup is complete in the "
       "backlog file. (Did not verify removal from any UI page templates.)",
]:
    doc.add_paragraph(line, style="List Number" if line[0].isdigit() else None)

doc.add_paragraph()
doc.add_paragraph("Caveats:")
for line in [
    "Classification labels are not mutually exclusive — some requirements are both CEREMONIAL "
       "and ORPHANED-IMPLEMENTATION. Each is reported under whichever bucket leads to the "
       "clearest action.",
    "DUPLICATE entries that span two requirements use a combined req_id 'A + B' (or 'A / B') and "
       "are NOT counted as high-confidence single-deletes — they need a merge decision.",
    "ORPHANED-IMPLEMENTATION is in the 'investigate' bucket, never high-confidence delete: the "
       "test could still be built. The audit's recommendation is BUILD-OR-DELETE, not auto-delete.",
    "STALE-DONE for F-002-S-002 is based on test_load_parity.py being present and matching "
       "T-001/T-002 by reading. It was not actually executed end-to-end against live MongoDB + "
       "Azure blob during this audit; status flip should be paired with a real run.",
    "test_pagination.py contains 54 unittest methods (counted by grep on '    def test_'); "
       "no Playwright equivalents exist in the repo. The 'every test must have a Playwright "
       "equivalent' requirement is meta-circular regardless.",
    "Schema-validation pytests (e.g. EPIC-005-F-004-S-001 NPI/entity_type_code rules) may "
       "ALREADY be enforced inside Code/Shared/tests/test_brain_schema_validator.py or pipeline "
       "ingest. Audit flagged them CEREMONIAL because there is no pytest_id link from these reqs "
       "to those tests — confirming with Skip is recommended before mass-delete.",
    "Audit was conducted read-only on every path except this output directory, and did not "
       "modify agile_backlog.json. No git operations performed.",
]:
    p = doc.add_paragraph(line)

doc.save(OUT_PATH)
print("WROTE:", OUT_PATH)
print("TOTAL_REQS:", total)
print("HIGH_CONF_DELETE:", len(high_conf_req_ids))
print("CEREMONIAL:", n_cer)
print("UNENFORCEABLE:", n_un)
print("ORPHANED_IMPL:", n_orph_impl)
print("DUPLICATE:", n_dup)
print("META_CIRCULAR:", n_meta)
print("STALE_DONE:", n_stale)
print("DELETED_FEATURE_RESIDUE:", n_dfr)
print("CONTRADICTORY:", n_contra)
