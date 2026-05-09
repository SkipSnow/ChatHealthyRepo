"""Generate langgraph-oo-best-practices-V6.docx.

V6 is a REWRITE (not an edit) commissioned by the architect's V5 annotation.
The architect accepted V4's findings but required V6 to reconcile the
LangGraph runtime with ChatHealthy's actual target architecture: a set of
HuggingFace Spaces, each owning one capability cost center, plus a thin
parent runtime that the React application talks to.

V6's reshape:
  1. Title + lineage
  2. Skip's V5 annotation quoted verbatim (binding directive)
  3. Reconciliation -- the two-dimensional architecture and how LangGraph
     fits into it (HF Spaces x graph nodes; RemoteGraph as the official
     bridge; React <-> parent runtime boundary)
  4. Five definitive choices ChatHealthy is making (carried forward,
     two refined by the V5 comment)
  5. Per-question answers (refreshed for V6; Q9 substantially rewritten)
  6. Anti-patterns
  7. Reference implementations
  8. Citations bibliography

Tag convention in citations:
    [official]  - LangChain / LangGraph / MongoDB co-maintained
    [community] - third-party blog, tutorial, or independent repo
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor


OUT = Path(__file__).with_name("langgraph-oo-best-practices-V6.docx")


# ============================================================
# Citations dictionary -- V4 set carried forward + RemoteGraph
# additions required by V6's reconciliation section.
# ============================================================

CITES = {
    # --- LangChain / LangGraph official docs ---
    "docs_graph_api":        ("official", "Use the graph API - LangChain Docs",
                              "https://docs.langchain.com/oss/python/langgraph/use-graph-api"),
    "docs_persistence":      ("official", "Persistence - LangChain Docs",
                              "https://docs.langchain.com/oss/javascript/langgraph/persistence"),
    "docs_app_structure":    ("official", "Application structure - LangChain Docs",
                              "https://docs.langchain.com/oss/python/langgraph/application-structure"),
    "docs_tools":            ("official", "Tools - LangChain Docs",
                              "https://docs.langchain.com/oss/python/langchain/tools"),
    "docs_context":          ("official", "Context overview - LangChain Docs",
                              "https://docs.langchain.com/oss/python/concepts/context"),
    "docs_test":             ("official", "Test - LangChain Docs",
                              "https://docs.langchain.com/oss/python/langgraph/test"),
    "docs_thinking":         ("official", "Thinking in LangGraph - LangChain Docs",
                              "https://docs.langchain.com/oss/python/langgraph/thinking-in-langgraph"),
    "docs_handoffs":         ("official", "Multi-agent handoffs - LangChain Docs",
                              "https://docs.langchain.com/oss/python/langchain/multi-agent/handoffs"),
    "docs_interrupts":       ("official", "Interrupts - LangChain Docs",
                              "https://docs.langchain.com/oss/python/langgraph/interrupts"),
    "docs_add_memory":       ("official", "Memory (add-memory) - LangChain Docs",
                              "https://docs.langchain.com/oss/python/langgraph/add-memory"),
    "docs_custom_ckpt":      ("official", "How to use a custom checkpointer - LangChain Docs",
                              "https://docs.langchain.com/langsmith/custom-checkpointer"),
    "docs_subgraphs":        ("official", "Use subgraphs - LangChain Docs",
                              "https://docs.langchain.com/oss/python/langgraph/use-subgraphs"),
    "ref_checkpoints":       ("official", "checkpoints API reference - LangGraph",
                              "https://reference.langchain.com/python/langgraph/checkpoints"),

    # --- RemoteGraph (NEW in V6 -- the bridge between in-process subgraph
    # composition and ChatHealthy's HF Space deployment topology) ---
    "docs_remote_graph":     ("official", "How to interact with a deployment using RemoteGraph - LangChain Docs",
                              "https://docs.langchain.com/langsmith/use-remote-graph"),
    "ref_remote_graph":      ("official", "RemoteGraph - LangChain Reference",
                              "https://reference.langchain.com/python/langgraph/pregel/remote/RemoteGraph"),
    "pypi_langgraph_sdk":    ("official", "langgraph-sdk on PyPI",
                              "https://pypi.org/project/langgraph-sdk/"),
    "pypi_langgraph_api":    ("official", "langgraph-api on PyPI",
                              "https://pypi.org/project/langgraph-api/"),

    # --- LangChain official blog ---
    "blog_building":         ("official", "Building LangGraph: an agent runtime from first principles",
                              "https://www.langchain.com/blog/building-langgraph"),
    "blog_v1":               ("official", "LangChain and LangGraph reach v1.0",
                              "https://www.langchain.com/blog/langchain-langgraph-1dot0"),
    "blog_klarna":           ("official", "How Klarna's AI assistant redefined customer support",
                              "https://www.langchain.com/blog/customers-klarna"),
    "blog_in_prod":          ("official", "Is LangGraph used in production?",
                              "https://www.langchain.com/blog/is-langgraph-used-in-production"),
    "blog_v02_ckpt":         ("official", "LangGraph v0.2: customization with new checkpointer libraries",
                              "https://blog.langchain.com/langgraph-v0-2/"),

    # --- MongoDB / LangChain co-maintained sources ---
    "pypi_mongo_ckpt":       ("official", "langgraph-checkpoint-mongodb on PyPI (MongoDB + LangChain co-maintained)",
                              "https://pypi.org/project/langgraph-checkpoint-mongodb/"),
    "mongo_atlas_lg":        ("official", "Integrate MongoDB with LangGraph - MongoDB Atlas Docs",
                              "https://www.mongodb.com/docs/atlas/ai-integrations/langgraph/"),
    "mongo_atlas_lg_agents": ("official", "Build an AI Agent with LangGraph and MongoDB Atlas",
                              "https://www.mongodb.com/docs/atlas/ai-integrations/langgraph/build-agents/"),
    "mongo_lc_readthedocs":  ("official", "MongoDBSaver class - LangChain MongoDB Read the Docs",
                              "https://langchain-mongodb.readthedocs.io/en/latest/langgraph_checkpoint_mongodb/saver/langgraph.checkpoint.mongodb.saver.MongoDBSaver.html"),
    "mongo_blog":            ("official", "Checkpointers and Native Parent-Child Retrievers with LangChain and MongoDB",
                              "https://www.mongodb.com/company/blog/innovation/checkpointers-native-parent-child-retrievers-with-langchain-mongodb"),

    # --- Official GitHub repositories ---
    "gh_langgraph":          ("official", "langchain-ai/langgraph",
                              "https://github.com/langchain-ai/langgraph"),
    "gh_example_pyproject":  ("official", "langchain-ai/langgraph-example-pyproject",
                              "https://github.com/langchain-ai/langgraph-example-pyproject"),
    "gh_open_deep_research": ("official", "langchain-ai/open_deep_research",
                              "https://github.com/langchain-ai/open_deep_research"),
    "gh_deep_research_scratch": ("official", "langchain-ai/deep_research_from_scratch",
                              "https://github.com/langchain-ai/deep_research_from_scratch"),
    "gh_deepagents":         ("official", "langchain-ai/deepagents",
                              "https://github.com/langchain-ai/deepagents"),
    "gh_swarm":              ("official", "langchain-ai/langgraph-swarm-py",
                              "https://github.com/langchain-ai/langgraph-swarm-py"),
    "gh_supervisor":         ("official", "langchain-ai/langgraph-supervisor-py",
                              "https://github.com/langchain-ai/langgraph-supervisor-py"),
    "gh_issue_1950":         ("official", "Issue #1950: Passing private state with Class Node",
                              "https://github.com/langchain-ai/langgraph/issues/1950"),

    # --- Community repositories ---
    "gh_awesome":            ("community", "von-development/awesome-LangGraph",
                              "https://github.com/von-development/awesome-LangGraph"),
    "gh_fastapi_template":   ("community", "wassim249/fastapi-langgraph-agent-production-ready-template",
                              "https://github.com/wassim249/fastapi-langgraph-agent-production-ready-template"),

    # --- Community blog posts ---
    "blog_swarnendu":        ("community", "Swarnendu De - LangGraph Best Practices",
                              "https://www.swarnendu.de/blog/langgraph-best-practices/"),
    "blog_yasin_modular":    ("community", "Yassin Hashem - Scaling AI Agents Beyond Notebooks: A Modular Architecture",
                              "https://medium.com/@yasin162001/scaling-ai-agents-beyond-notebooks-a-modular-architecture-for-langgraph-in-production-4711764de464"),
    "blog_shaza_typing":     ("community", "Shaza Ali - Type Safety in LangGraph: TypedDict vs Pydantic",
                              "https://shazaali.substack.com/p/type-safety-in-langgraph-when-to"),
    "blog_pankaj_pyd":       ("community", "Pankaj Chandravanshi - LangGraph: State with Pydantic BaseModel",
                              "https://medium.com/fundamentals-of-artificial-intelligence/langgraph-state-with-pydantic-basemodel-023a2158ab00"),
    "blog_easton_2026":      ("community", "BetterLink Blog - LangGraph State Management in Practice (2026)",
                              "https://eastondev.com/blog/en/posts/ai/20260424-langgraph-agent-architecture/"),
    "blog_markaicode":       ("community", "Markaicode - Production Multi-Agent System with LangGraph",
                              "https://markaicode.com/langgraph-production-agent/"),
    "blog_clean_state":      ("community", "Vishal Lad - Clean State Architecture in LangGraph (Mar 2026)",
                              "https://medium.com/@ladvishal1985/everything-ive-learned-about-clean-state-architecture-in-langgraph-6d1352b0c00c"),
    "blog_unit_testing":     ("community", "Anirudh Sharma - Unit Testing LangGraph: Nodes and Flow Paths",
                              "https://medium.com/@anirudhsharmakr76/unit-testing-langgraph-testing-nodes-and-flow-paths-the-right-way-34c81b445cd6"),
    "blog_mock_llm":         ("community", "Matt Carvalho - How to Properly Mock LangChain LLM Execution",
                              "https://medium.com/@matgmc/how-to-properly-mock-langchain-llm-execution-in-unit-tests-python-76efe1b8707e"),
    "blog_tools_first":      ("community", "SitePoint - Implementing the Tools-First Pattern in LangGraph",
                              "https://www.sitepoint.com/implementing-the-tools-first-pattern-in-lang-graph/"),
    "deepwiki_subgraphs":    ("community", "Graph composition and nested graphs - DeepWiki (langchain-ai/langgraph)",
                              "https://deepwiki.com/langchain-ai/langgraph/3.6-graph-composition-and-nested-graphs"),
    "deepwiki_hitl":         ("community", "Human-in-the-loop and interrupts - DeepWiki (langchain-ai/langgraph)",
                              "https://deepwiki.com/langchain-ai/langgraph/3.7-human-in-the-loop-and-interrupts"),
    "blog_send_api":         ("community", "Sreeni - Leveraging LangGraph's Send API for Dynamic and Parallel Workflow Execution",
                              "https://dev.to/sreeni5018/leveraging-langgraphs-send-api-for-dynamic-and-parallel-workflow-execution-4pgd"),
    "blog_handoffs_tds":     ("community", "Towards Data Science - How Agent Handoffs Work in Multi-Agent Systems",
                              "https://towardsdatascience.com/how-agent-handoffs-work-in-multi-agent-systems/"),
    "blog_multiagent_2026":  ("community", "Mager - LangGraph: Build Stateful Multi-Agent Systems That Don't Crash (2026)",
                              "https://www.mager.co/blog/2026-03-12-langgraph-deep-dive/"),
    "blog_scaling_lg":       ("community", "AI Practitioner - Scaling LangGraph Agents: Parallelization, Subgraphs, and Map-Reduce Trade-Offs",
                              "https://aipractitioner.substack.com/p/scaling-langgraph-agents-parallelization"),
    "dev_mongo_ltm":         ("community", "DEV (MongoDB) - LangGraph With MongoDB: Building Conversational Long-Term Memory",
                              "https://dev.to/mongodb/langgraph-with-mongodb-building-conversational-long-term-memory-for-intelligent-ai-agents-2pcn"),
    "blog_bentoml_deploy":   ("community", "BentoML - Deploying a LangGraph Agent Application",
                              "https://www.bentoml.com/blog/deploying-a-langgraph-agent-application-with-an-open-source-model"),
}


# ============================================================
# Helpers
# ============================================================

def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p


def add_do(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    lead = p.add_run("Do: ")
    lead.bold = True
    lead.font.size = Pt(11)
    body = p.add_run(text)
    body.font.size = Pt(11)
    return p


def add_dont(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    lead = p.add_run("Do NOT: ")
    lead.bold = True
    lead.font.size = Pt(11)
    body = p.add_run(text)
    body.font.size = Pt(11)
    return p


def add_quote_para(doc, text):
    """A blockquote-style indented italic paragraph for the architect's note."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(24)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(11)
    return p


def add_citation(doc, key):
    tag, title, url = CITES[key]
    p = doc.add_paragraph(style="List Bullet 2")
    run = p.add_run(f"[{tag}] {title} - ")
    run.font.size = Pt(10)
    run.italic = True
    link = p.add_run(url)
    link.font.size = Pt(10)
    link.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p


def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(24)
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    return p


# ============================================================
# Build
# ============================================================

def build():
    doc = Document()

    # ---------------- Title block ----------------
    title = doc.add_heading("LangChain / LangGraph OO Best-Practices Guide - V6", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sr = sub.add_run("Standalone reference. Rewrite (not edit) commissioned by "
                     "the architect's V5 annotation. Reconciles the LangGraph "
                     "runtime with ChatHealthy's HuggingFace-Space deployment "
                     "topology.")
    sr.italic = True
    sr.font.size = Pt(11)
    meta = doc.add_paragraph()
    meta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    mr = meta.add_run("Compiled 2026-04-28  -  Skip Snow / ChatHealthy  -  "
                      "Sources tagged [official] = LangChain / LangGraph or "
                      "MongoDB co-maintained, [community] = third-party")
    mr.font.size = Pt(9)
    mr.italic = True

    add_para(doc,
        "Lineage: V1 (research) -> V2 (architect annotations) -> V3 "
        "(annotation resolution) -> V4 (definitive synthesis) -> V5 "
        "(architect's reconciliation directive) -> this V6 (rewrite, "
        "standalone). A reader new to ChatHealthy can read V6 alone and "
        "understand the LangGraph implementation philosophy without "
        "reading V1-V5 first.",
        italic=True)

    # ============================================================
    # 1. Executive summary
    # ============================================================
    add_heading(doc, "1. Executive Summary", level=1)
    add_para(doc,
        "ChatHealthy implements its conversational user experience as a thin "
        "parent LangGraph runtime (one process, one MongoDBSaver, one "
        "checkpoint thread per session) that composes a small number of "
        "capability subgraphs. Each capability subgraph corresponds 1:1 to a "
        "HuggingFace Space - one Space hosts the FindCare capability, another "
        "hosts EvaluateCare, another hosts Security. The parent runtime "
        "binds those Spaces in via RemoteGraph, the official LangGraph "
        "primitive that exposes a deployed graph behind the same "
        "PregelProtocol interface as a local CompiledStateGraph. The result "
        "is a single graph the parent runtime sees, one set of checkpoints, "
        "one place where Command(graph=Command.PARENT) and interrupts work, "
        "but a deployment topology that mirrors the cost-center / org-chart "
        "axiom: each Space is one capability cost center, one deploy "
        "pipeline, one observability surface.")

    add_para(doc,
        "Within a Space, the implementation follows the production pattern "
        "the late-2025 / 2026 ecosystem has converged on: state is a "
        "Pydantic BaseModel in its own state.py module; conversation is "
        "messages: Annotated[list[AnyMessage], add_messages] with tool I/O "
        "as ToolMessage entries on that channel; nodes are pure functions "
        "or callable classes that return partial state updates; runtime "
        "dependencies are injected via Runtime[ContextSchema] / ToolRuntime; "
        "persistence is MongoDBSaver from the co-maintained "
        "langgraph-checkpoint-mongodb package.")

    add_para(doc,
        "The React application talks to one URL: the parent runtime's HTTP "
        "endpoint. The React side does NOT see the capability Spaces "
        "directly. The parent runtime owns session state, owns the "
        "checkpoint thread, owns interrupt / resume, and owns routing "
        "between epics; capability Spaces own their bespoke business logic "
        "and own their own /health endpoint for operations.",
        italic=False)

    # ============================================================
    # 2. The architect's V5 annotation (binding)
    # ============================================================
    add_heading(doc, "2. The Architect's V5 Annotation (Binding)", level=1)
    add_para(doc,
        "V5 carried one large architect comment on Q9 (Graph decomposition "
        "and OO principles). It is reproduced verbatim below. It is the "
        "binding directive that drove this V6 rewrite. References to "
        "'this document' are to V5; the response is in section 3.",
        italic=True)
    add_quote_para(doc,
        "I understand how deep research is a sub graph, or a set of "
        "subgraphs but I do not think it is an epic or needs its own "
        "Hugging face server. We have two architectural considerations "
        "that are getting mixed up a bit. For me it is pictorially a two "
        "dimensional table where the columns are the individual Hugging "
        "Face servers that realize the work that the nodes ask for. And "
        "the columns are the nodes on the graphs. By Working on one cell "
        "at a time we can radically control the blast area of harm that "
        "can be done and strictly encapsulate capabilities as nodes on a "
        "graph within a server. But we seem to be interlacing the "
        "concepts. What I see in this document as a whole is a clear "
        "directive to have a master graph that delegates to sub graphs. "
        "For me each sub graph will be associated with one Hugging face "
        "server, because that server encapsulates a major business "
        "capability. There are two epics that are on the front end "
        "(LangGraph is only a concept for the runtime application and "
        "not the pipeline or dev ops (operations)) that are utilized by "
        "the core biz epics: they are security and shared services. I'm "
        "sort of migrating to the notion that security must be in its own "
        "hugging face container. So the model is in fact consistent. What "
        "I want out of the last version of this doc is a reconciliation "
        "between the application's overall architecture and how we will "
        "use LangGraph to orchestrate the activities within that "
        "architecture. Of particular concern is how the parent graph and "
        "the React application interface with one another and how they "
        "share or divide responsibilities. I.e. I accept all the "
        "findings, and as we migrate to design, we need to put this into "
        "the context of the overall architecture of the application "
        "which is best understood by looking at our deployment scripts. "
        "I'm not asking for a design doc. I just need this doc to be "
        "aware of and take account of the actual target architecture "
        "that we are moving towards. This doc will be, combined with the "
        "work we have been doing on rationalizing the file structures, "
        "and some guidance I have yet given, one of the main inputs to "
        "the design we must complete before we implement LangGraph and "
        "refactor our code base to fit within the target architecture's "
        "constraints.")

    # ============================================================
    # 3. Reconciliation -- the heart of V6
    # ============================================================
    add_heading(doc, "3. Reconciliation: Architecture-Aware LangGraph Design", level=1)

    add_para(doc,
        "The architect's note resolves into a precise architectural model. "
        "V6 adopts it as the framing for every recommendation that follows.")

    add_heading(doc, "3.1 The two-dimensional model", level=2)
    add_para(doc,
        "Two orthogonal axes:")
    add_bullet(doc,
        "Deployment axis (rows) - HuggingFace Spaces. Each Space is one "
        "container, one deploy pipeline, one /health endpoint, one cost-"
        "center line on the GL, one line on the org chart. ChatHealthy "
        "today has the Spaces visible in Code/deploy/_start_*.py: "
        "_start_findcare, _start_evalcare, _start_shared. The "
        "codebase-reorganization-V5 recommendation adds a Security Space "
        "(EPIC-002 carved out of shared_services).")
    add_bullet(doc,
        "Runtime axis (columns) - graph nodes. Each node is one step in a "
        "single conversation. A node may execute locally (in the parent "
        "runtime process) or it may execute remotely (in a capability "
        "Space, behind RemoteGraph). The parent runtime cannot tell the "
        "difference at compile time - that is the point of the "
        "PregelProtocol abstraction.")
    add_para(doc,
        "Each cell in the table is one capability-step pair: 'the part of "
        "FindCare that runs as a node in the conversation graph,' 'the "
        "part of Security that runs as a node in the conversation graph.' "
        "Working one cell at a time is exactly the architect's blast-area "
        "containment goal: a bug in the FindCare specialty classifier "
        "node cannot brick the EvaluateCare scoring node because they run "
        "in different containers behind the parent runtime's RemoteGraph "
        "wrappers.")

    add_heading(doc, "3.2 RemoteGraph: the official bridge", level=2)
    add_para(doc,
        "The mechanism the V4 doc was missing is RemoteGraph. The "
        "official reference: 'RemoteGraph behaves the same way as a Graph "
        "and can be used directly as a node in another Graph,' implementing "
        "the PregelProtocol. Constructor signature:")
    add_code_block(doc,
        "from langgraph.pregel.remote import RemoteGraph\n"
        "\n"
        "find_care = RemoteGraph(\n"
        "    'find_care',\n"
        "    url='https://skipsnow-dev-chathealthyspace.hf.space',\n"
        "    api_key=...,\n"
        ")")
    add_para(doc,
        "The official compose example:")
    add_code_block(doc,
        "from langgraph.graph import StateGraph, MessagesState, START\n"
        "from langgraph.pregel.remote import RemoteGraph\n"
        "\n"
        "find_care = RemoteGraph('find_care', url=FINDCARE_HF_URL)\n"
        "evaluate = RemoteGraph('evaluate_care', url=EVALCARE_HF_URL)\n"
        "security = RemoteGraph('security', url=SECURITY_HF_URL)\n"
        "\n"
        "parent = StateGraph(ParentState)\n"
        "parent.add_node('find_care', find_care)\n"
        "parent.add_node('evaluate_care', evaluate)\n"
        "parent.add_node('security', security)\n"
        "parent.add_edge(START, 'security')\n"
        "# ...routing edges...\n"
        "graph = parent.compile(checkpointer=mongo_saver)")
    add_para(doc,
        "The same code, with a stub URL pointing at localhost during dev, "
        "is exactly what an in-process subgraph composition would look "
        "like - that is the abstraction the parent runtime needs.")
    add_para(doc,
        "Two binding constraints from the official reference: "
        "(a) RemoteGraph must NOT be used to call back into the deployment "
        "that is currently executing - that causes deadlocks and resource "
        "exhaustion; for graphs in the same deployment use local subgraph "
        "composition. (b) RemoteGraph requires the target deployment to be "
        "running an Agent Server (langgraph-api or LangGraph Platform) that "
        "exposes the standard endpoints. ChatHealthy's HF Spaces today run "
        "FastAPI with bespoke endpoints, not langgraph-api. Adopting "
        "RemoteGraph requires either standing up langgraph-api inside each "
        "capability Space, or implementing the same protocol behind the "
        "existing FastAPI apps. This is a real cost and a real design "
        "decision - V6 surfaces it; the design doc the architect is "
        "moving towards must resolve it.")
    for k in ("docs_remote_graph", "ref_remote_graph", "pypi_langgraph_sdk",
              "pypi_langgraph_api", "deepwiki_subgraphs"):
        add_citation(doc, k)

    add_heading(doc, "3.3 The parent runtime and the React application", level=2)
    add_para(doc,
        "The architect's question 'how the parent graph and the React "
        "application interface with one another and how they share or "
        "divide responsibilities' has a precise answer in this model.")
    add_para(doc, "What the parent runtime owns:")
    add_bullet(doc,
        "Session state. One MongoDBSaver instance, one thread_id per user "
        "session, one checkpoint history. Capability Spaces do NOT own "
        "session checkpoints - they receive a focused projection of the "
        "parent state on each invocation.")
    add_bullet(doc,
        "Routing between epics. The parent's routing edges decide which "
        "capability node runs next based on the conversation so far.")
    add_bullet(doc,
        "Interrupts and human-in-the-loop. interrupt() / "
        "Command(resume=...) requires a checkpointer; the only "
        "checkpointer is the parent's. Long-running deep-research runs "
        "(Q9) call interrupt() in the parent's flow and resume there.")
    add_bullet(doc,
        "The HTTP endpoint that React talks to. One URL. The React "
        "client does NOT call FindCare or EvaluateCare directly for "
        "conversational flow - those calls go through the parent.")
    add_para(doc, "What the capability Space owns:")
    add_bullet(doc,
        "Bespoke business logic for that capability. Specialty "
        "classifier, provider scoring, trial matching for FindCare; "
        "weighted-composite scoring for EvaluateCare; OAuth, session-token "
        "verification, entitlements for Security.")
    add_bullet(doc,
        "Its own MongoDB collections for capability-specific reads and "
        "writes (the Space's data, not session checkpoints).")
    add_bullet(doc,
        "Its own deployment lifecycle. Build, deploy, observability, "
        "/health endpoint, on-call rotation if any.")
    add_bullet(doc,
        "Its own non-conversational HTTP API for the cases where React "
        "DOES call it directly - non-conversational read endpoints "
        "(load a provider detail page, render an evaluation report) "
        "bypass the parent runtime by design. The parent runtime is for "
        "conversational flow; static data fetches are not.")
    add_para(doc,
        "What stays out of LangGraph entirely (per the architect's note):")
    add_bullet(doc,
        "The data-management pipelines (EPIC-012). Ingest, dedup, "
        "embedding generation - these are scheduled jobs, not "
        "conversational steps. They run under Operations, not under the "
        "parent runtime, and they should NOT be modeled as a graph.")
    add_bullet(doc,
        "Operations / DevOps. CI, deploy scripts, build counters, smoke "
        "tests - those are EPIC-010 territory, not LangGraph territory.")

    add_heading(doc, "3.4 Same-Space subgraph composition vs cross-Space RemoteGraph", level=2)
    add_para(doc,
        "Inside one Space, the conventional in-process subgraph "
        "composition (compile a child graph, add it as a node) remains the "
        "right pattern. RemoteGraph is reserved for the cross-Space case. "
        "Concrete heuristic:")
    add_bullet(doc,
        "Two nodes that always run together, share a database client, and "
        "deploy together -> same Space, in-process composition (or simply "
        "two nodes in the same subgraph).")
    add_bullet(doc,
        "Two nodes that have distinct cost-center lines on the GL or "
        "distinct deploy pipelines -> different Spaces, RemoteGraph.")
    add_bullet(doc,
        "Deep research is the cleanest example of the first kind. Per the "
        "architect: deep research is a subgraph (or set of subgraphs), "
        "not its own Space. It runs inside whichever capability Space "
        "needs it - the FindCare Space or the EvaluateCare Space - "
        "imported as a Python module, composed in-process. This corrects "
        "V4/V5's implicit suggestion that deep research could be its own "
        "epic.")
    for k in ("docs_subgraphs", "docs_handoffs", "deepwiki_subgraphs",
              "docs_remote_graph"):
        add_citation(doc, k)

    add_heading(doc, "3.5 Mapping to ChatHealthy's deployment scripts", level=2)
    add_para(doc,
        "The deployment topology is verifiable today in Code/deploy/. "
        "_start_findcare.py, _start_evalcare.py, and _start_shared.py "
        "each spawn a separate uvicorn process on a distinct port behind "
        "TLS. The codebase-reorganization-V5 doc adds Code/Security/ as a "
        "fourth Space (EPIC-002) and renames Code/shared_services/ to "
        "host narrower 'shared services' scope (Tool Router, LLM client, "
        "embeddings). Under V6's model:")
    add_bullet(doc,
        "FindCare HF Space - hosts the FindCare capability subgraph "
        "(provider search, specialty classification, provider scoring, "
        "trial matching). Exposes the LangGraph Agent Server protocol "
        "OR an equivalent protocol the parent runtime calls via "
        "RemoteGraph.")
    add_bullet(doc,
        "EvaluateCare HF Space - hosts the EvaluateCare scoring "
        "capability (weighted-composite scoring, normalization, "
        "explanation assembly).")
    add_bullet(doc,
        "Security HF Space (Option A from codebase-reorg V5) - hosts "
        "the Security capability subgraph: pre-conversation token "
        "verification, entitlement checks, audit emission.")
    add_bullet(doc,
        "Shared Services HF Space - hosts non-LangGraph utilities "
        "(Tool Router endpoint, LLM client, embeddings). These are "
        "called as TOOLS from inside capability subgraphs, not "
        "composed as nodes.")
    add_bullet(doc,
        "Parent runtime - lives wherever the React-facing endpoint "
        "lives. Today that is the FindCare Space's main.py; under V6's "
        "model the parent runtime is logically separate from any one "
        "capability Space, even if it is co-located with one of them "
        "for now. The recommended end state is a thin parent Space "
        "(Code/Parent/ or co-located in Operations) that hosts only "
        "the parent graph + the MongoDBSaver instance + the React-facing "
        "endpoint.")

    # ============================================================
    # 4. Definitive choices
    # ============================================================
    add_heading(doc, "4. Definitive Choices ChatHealthy Is Making", level=1)
    add_para(doc,
        "Five resolved decisions. Items marked (refined V6) changed in "
        "scope or implementation because of the architect's V5 comment. "
        "Items marked (carried V5) are unchanged in substance.",
        italic=True)
    add_bullet(doc,
        "Persistence: MongoDBSaver from langgraph-checkpoint-mongodb, "
        "instantiated once in the parent runtime. NOT PostgresSaver. "
        "(Carried V5; see Q2.)")
    add_bullet(doc,
        "Graph decomposition: one capability subgraph per HuggingFace "
        "Space, composed under the parent runtime via RemoteGraph for "
        "cross-Space and in-process composition for same-Space. NOT a "
        "monolithic graph, NOT separately invoked top-level graphs that "
        "React calls directly. Deep research is a same-Space subgraph, "
        "not its own epic and not its own Space. (Refined V6; see Q9.)")
    add_bullet(doc,
        "State home: Pydantic BaseModel in its own state.py file in each "
        "capability Space; the parent state is a small Pydantic schema "
        "that contains messages plus a few routing-control fields. NOT "
        "TypedDict in the orchestration script. (Carried V5; see Q1.)")
    add_bullet(doc,
        "Dependency injection: nodes and tools receive dependencies via "
        "Runtime[ContextSchema] / ToolRuntime. The parent passes any "
        "session-scoped context (user_id, thread_id, locale) to "
        "RemoteGraph invocations through the standard context= argument; "
        "capability Spaces internally use Runtime[ContextSchema] for "
        "their own DB clients. NOT module-level imports. (Carried V5 "
        "with a cross-Space refinement; see Q5.)")
    add_bullet(doc,
        "Conversation modeling: messages: Annotated[list[AnyMessage], "
        "add_messages] on the parent state; tool calls and tool results "
        "are ToolMessage entries on that channel. Capability subgraphs "
        "return ToolMessage entries (not flat fields, not tool_outputs "
        "dict) when they need to surface tool exchanges to the LLM. "
        "(Carried V5; see Q4 + Q10.)")

    add_para(doc, "Two architectural choices new in V6:", italic=True)
    add_bullet(doc,
        "Cross-Space composition primitive: RemoteGraph from "
        "langgraph.pregel.remote. Each capability Space exposes the "
        "standard Agent Server protocol; the parent runtime composes "
        "Spaces as RemoteGraph nodes. (New V6; see section 3.2.)")
    add_bullet(doc,
        "React <-> parent boundary: React talks to the parent runtime "
        "for conversational flow. React talks to capability Spaces "
        "directly only for non-conversational read endpoints. (New V6; "
        "see section 3.3.)")

    add_para(doc, "Selected supporting sources:", italic=True)
    for k in ("docs_graph_api", "docs_app_structure", "docs_persistence",
              "docs_context", "docs_remote_graph", "ref_remote_graph",
              "blog_building", "blog_clean_state", "pypi_mongo_ckpt",
              "mongo_atlas_lg", "docs_handoffs", "docs_subgraphs",
              "docs_interrupts", "deepwiki_subgraphs",
              "gh_deep_research_scratch", "gh_example_pyproject"):
        add_citation(doc, k)

    # ============================================================
    # 5. Per-question findings (10 questions from V1, refreshed)
    # ============================================================
    add_heading(doc, "5. Per-Question Findings (Refreshed)", level=1)
    add_para(doc,
        "Each Q below is an answer to one of the 10 V1 research "
        "questions. Q9 is substantially rewritten because the architect's "
        "V5 comment lands directly on it; the others are refreshed to "
        "reflect the parent + RemoteGraph topology.",
        italic=True)

    # ---- Q1 ----
    add_heading(doc, "Q1. State schema home", level=2)
    add_do(doc,
        "Define state as a Pydantic BaseModel in a dedicated state.py "
        "module, one per capability Space and one for the parent runtime.")
    add_do(doc,
        "Define an input_schema and an output_schema per capability "
        "subgraph - these are the only fields that cross the Space "
        "boundary. Internal-only fields stay private to the Space.")
    add_dont(doc,
        "Define a 15-field flat TypedDict inline at the top of the graph "
        "build script. (This is the POC's current shape and an explicit "
        "anti-pattern; see AP-1.)")
    add_dont(doc,
        "Use a bare dataclass unless Pydantic's recursive validation "
        "cost has been measured and is unacceptable - and document the "
        "measurement.")
    add_para(doc,
        "Rationale. The official graph-API doc states 'State in LangGraph "
        "can be a TypedDict, Pydantic model, or dataclass' and presents "
        "TypedDict as the introductory example only. Markaicode's "
        "production write-up: 'Use a Pydantic BaseModel the moment you "
        "step into production.' Shaza Ali, Pankaj Chandravanshi, and "
        "BetterLink 2026 converge: TypedDict may stay for purely-internal "
        "subgraph state, Pydantic is mandatory at boundaries. ChatHealthy's "
        "every capability subgraph crosses an HTTP boundary (RemoteGraph), "
        "so Pydantic is mandatory throughout.")
    for k in ("docs_graph_api", "blog_shaza_typing", "blog_pankaj_pyd",
              "blog_easton_2026", "blog_markaicode", "docs_app_structure",
              "gh_deep_research_scratch"):
        add_citation(doc, k)

    # ---- Q2 ----
    add_heading(doc, "Q2. State persistence (checkpointer)", level=2)
    add_do(doc,
        "Use MongoDBSaver from langgraph-checkpoint-mongodb (co-maintained "
        "by MongoDB and LangChain), with the async API. Instantiate "
        "ONCE, in the parent runtime. Capability Spaces do NOT instantiate "
        "their own checkpointers.")
    add_do(doc,
        "Use InMemorySaver in unit and integration tests so each test "
        "gets a fresh checkpoint store.")
    add_dont(doc,
        "Use PostgresSaver. ChatHealthy already runs MongoDB; introducing "
        "Postgres only to host the LangGraph checkpointer adds a database "
        "for no concrete requirement.")
    add_dont(doc,
        "Use SqliteSaver in production - the official doc labels it "
        "'ideal for experimentation and local workflows.'")
    add_dont(doc,
        "Build a custom MongoDB saver - the official co-maintained one "
        "already covers the entire BaseCheckpointSaver contract.")
    add_dont(doc,
        "Stand up one MongoDBSaver per Space. Every capability Space "
        "writing its own checkpoints fragments the session - interrupt() "
        "and Command(resume=...) lose coherence. The single parent "
        "checkpointer is the consistency anchor.")
    add_para(doc,
        "Rationale. PyPI lists langgraph-checkpoint-mongodb as MongoDB- "
        "and LangChain-maintained (not community); the MongoDB Atlas "
        "integration page documents MongoDBSaver as a first-party feature "
        "that 'persists agent state in MongoDB' and 'enables human-in-the-"
        "loop, memory, time travel, and fault-tolerance.' The full "
        "BaseCheckpointSaver contract (sync + async) is implemented and "
        "is a published Read-the-Docs class. The single-checkpointer "
        "constraint is a direct consequence of section 3.3: only the "
        "parent owns session state.")
    for k in ("docs_persistence", "blog_v02_ckpt", "blog_easton_2026",
              "blog_klarna", "pypi_mongo_ckpt", "mongo_atlas_lg",
              "mongo_atlas_lg_agents", "mongo_lc_readthedocs",
              "mongo_blog", "ref_checkpoints", "docs_custom_ckpt",
              "dev_mongo_ltm"):
        add_citation(doc, k)

    # ---- Q3 ----
    add_heading(doc, "Q3. Node organization", level=2)
    add_do(doc,
        "Default to small pure functions returning partial state updates, "
        "declared in a per-subgraph nodes.py.")
    add_do(doc,
        "Promote a node to a class with __call__ when it needs "
        "constructor-injected dependencies (LLM client, DB handle, "
        "config) or when several closely-related nodes share state through "
        "instance attributes.")
    add_dont(doc,
        "Use the procedural top-level-function-per-step pattern beyond an "
        "exploratory prototype. (This is the POC's current shape; see AP-4.)")
    add_dont(doc,
        "Hide mutable state inside a class-node - issue #1950 documents "
        "broken behaviour for class-nodes that try to carry hidden private "
        "state. Class-nodes are for dependency injection only.")
    add_dont(doc,
        "Mutate the incoming state dict. Always return a partial-update "
        "dict.")
    add_para(doc,
        "Rationale. The official graph-API doc defines a node as 'just a "
        "Python function that reads our graph's state and makes updates "
        "to it' and warns: 'Nodes should return updates to the state "
        "directly, instead of mutating the state.' The 'Thinking in "
        "LangGraph' doc reinforces fine-grained nodes because 'durable "
        "execution creates checkpoints at node boundaries' - every "
        "checkpoint is a chance to resume.")
    for k in ("docs_graph_api", "docs_thinking", "gh_issue_1950",
              "docs_app_structure"):
        add_citation(doc, k)

    # ---- Q4 ----
    add_heading(doc, "Q4. Tool design", level=2)
    add_do(doc,
        "Use the @tool decorator for the common case (typed args; "
        "docstring becomes the tool description).")
    add_do(doc,
        "Drop to a BaseTool subclass when you need explicit args_schema, "
        "paired sync+async behaviour, or tool-level state.")
    add_do(doc,
        "Wire tools into the graph via the prebuilt ToolNode.")
    add_do(doc,
        "Model tool calls and tool results as ToolMessage entries on the "
        "messages channel. This is true inside a capability Space and "
        "true on the parent runtime - the channel is the same shape on "
        "both sides of the RemoteGraph boundary.")
    add_dont(doc,
        "Model tool outputs as flat top-level state fields per tool. "
        "(POC has six such fields: location, specialty_query, specialties, "
        "homeopathic_expansion, providers, trials. This is AP-3.)")
    add_dont(doc,
        "Model tool outputs as a tool_outputs: dict[str, Any] catch-all. "
        "The messages channel IS the right home for tool I/O; structured "
        "artefacts the LLM does not need to read get their own typed "
        "Pydantic field on the subgraph state.")
    add_dont(doc,
        "Build tools after the prompt is finalised. SitePoint names this "
        "the dominant cause of 'hallucinated parameters and silent "
        "failures.'")
    add_para(doc,
        "Rationale. The official tools doc states 'The simplest way to "
        "create a tool is with the @tool decorator. By default, the "
        "function's docstring becomes the tool's description.' Reserved "
        "parameter names config and runtime imply tools must use "
        "ToolRuntime for dependency injection (Q5). The 'Clean State "
        "Architecture' post warns against overloading messages with non-"
        "conversational artefacts AND against multiplying top-level "
        "fields - the resolution is: ToolMessage on the messages channel "
        "for the LLM-visible exchange, named typed Pydantic fields for "
        "structured artefacts.")
    for k in ("docs_tools", "blog_tools_first", "blog_clean_state",
              "docs_graph_api"):
        add_citation(doc, k)

    # ---- Q5 ----
    add_heading(doc, "Q5. Dependency injection", level=2)
    add_do(doc,
        "Pass dependencies via the LangGraph runtime context: "
        "Runtime[ContextSchema] for nodes, ToolRuntime[ContextSchema] for "
        "tools. Supply the dependencies at graph.invoke(...) time via the "
        "context= argument.")
    add_do(doc,
        "For class-based nodes, accept the dependency in __init__ and "
        "bind the constructed instance into the graph.")
    add_do(doc,
        "For RemoteGraph invocations, pass session-scoped context "
        "(user_id, thread_id, locale, request_id for tracing) through "
        "the context= argument; the receiving capability Space reads it "
        "via its own Runtime[ContextSchema].")
    add_dont(doc,
        "Import database clients or LLM clients at module scope inside "
        "node files. This is the most common anti-pattern in early "
        "LangGraph code and breaks testability.")
    add_dont(doc,
        "Pass dependencies through state. State is data that "
        "checkpoints; clients are not data, they are runtime context.")
    add_para(doc,
        "Rationale. The official Context overview is explicit: 'Runtime "
        "context is a form of dependency injection... optimize the LLM "
        "context by providing dependencies (like database connections, "
        "user IDs, or API clients) to your tools and nodes at runtime "
        "rather than hardcoding them.' The canonical pattern is "
        "def node(state, runtime: Runtime[ContextSchema]). Swarnendu De "
        "and the fastapi-langgraph production template demonstrate "
        "isolating clients in a services layer that nodes consume.")
    for k in ("docs_context", "blog_swarnendu", "gh_fastapi_template",
              "docs_tools"):
        add_citation(doc, k)

    # ---- Q6 ----
    add_heading(doc, "Q6. Module / file structure", level=2)
    add_do(doc,
        "Adopt the official 4-file layout per subgraph: state.py "
        "(Pydantic schemas), nodes.py (node functions / class-callables), "
        "tools.py (tool definitions), agent.py (graph construction). "
        "Include a langgraph.json manifest at each Space's root per the "
        "langgraph-example-pyproject template.")
    add_do(doc,
        "Use core/ as the layer name for capability business logic "
        "(provider_search_service.py, scoring_engine.py, normalization.py). "
        "core/ passes the cost-center axiom (the FindCare core, the "
        "EvaluateCare scoring core) and has no overlap with the "
        "engineering-rules controlled vocabulary.")
    add_do(doc,
        "Decompose by capability - one folder per Space (Code/FindCare/, "
        "Code/EvaluateCare/, Code/Security/, Code/SharedServices/), each "
        "containing a backend/ tree with the 4-file layout under "
        "backend/agent/ and the business logic under backend/core/.")
    add_dont(doc,
        "Use 'domain' as a structural noun. The hexagonal-architecture "
        "term fails the operating-model axiom (no HR, GL, or capability-"
        "map line for it). Where it appears in legacy paths, plan to "
        "rename to core/ as part of the refactor.")
    add_dont(doc,
        "Put 14 nodes and the graph builder in one 336-line file. (POC "
        "user_journey.py - this is AP-2.)")
    add_para(doc,
        "Rationale. The official Application Structure doc presents "
        "state.py / nodes.py / tools.py / agent.py inside my_agent/utils/ "
        "+ my_agent/agent.py. The langgraph-example-pyproject template "
        "ships exactly this layout. Yassin Hashem's 'Scaling AI Agents "
        "Beyond Notebooks' extends it for multi-agent production. The "
        "fastapi-langgraph production template shows how this scales "
        "into a real service. The codebase-reorganization-V5 doc settled "
        "core/ as the layer name (over 'domain' and 'rules/') for "
        "ChatHealthy's particular constraints.")
    for k in ("docs_app_structure", "gh_example_pyproject",
              "blog_yasin_modular", "gh_fastapi_template",
              "blog_multiagent_2026"):
        add_citation(doc, k)

    # ---- Q7 ----
    add_heading(doc, "Q7. State immutability and reducers", level=2)
    add_do(doc,
        "Return partial-update dicts from every node. Never mutate "
        "incoming state.")
    add_do(doc,
        "Annotate every list / dict field with an explicit reducer at "
        "the moment you declare the field. Use add_messages for the "
        "messages channel, operator.add for simple list accumulation, and "
        "a custom reducer when you need de-duplication or capping.")
    add_do(doc,
        "Cap unbounded growth via a custom reducer when a list channel "
        "is written by many nodes. Otherwise the checkpoint grows "
        "unboundedly.")
    add_dont(doc,
        "Leave list fields un-annotated. LangGraph's default is "
        "last-write-wins; in any fan-out topology (e.g. the POC's "
        "parallel safety + classifier path) this silently loses data.")
    add_dont(doc,
        "Mutate state in place inside a node body, even for convenience. "
        "The Pregel/BSP merge model assumes immutable returns.")
    add_para(doc,
        "Rationale. The official graph-API doc: 'Nodes should return "
        "updates to the state directly, instead of mutating the state' "
        "and 'If no reducer function is explicitly specified then it is "
        "assumed that all updates to the key should override it.' The "
        "'Building LangGraph' post explains the underlying Pregel/BSP "
        "model: channels are versioned and updates are merged "
        "deterministically, which depends on immutable returns. The "
        "'Clean State Architecture' post names 'missing reducers on list "
        "fields' as an explicit anti-pattern: 'Always give it a reducer, "
        "even if only one node writes to it now.'")
    for k in ("docs_graph_api", "blog_building", "blog_clean_state"):
        add_citation(doc, k)

    # ---- Q8 ----
    add_heading(doc, "Q8. Testing", level=2)
    add_do(doc,
        "Run a two-layer split. Layer 1: unit-test each node by calling "
        "the function (or class instance) directly with a synthetic "
        "state dict and asserting on the returned partial update - no "
        "graph compilation needed. Layer 2: integration-test the "
        "compiled subgraph with InMemorySaver, mocked LLM (FakeListLLM), "
        "and stubbed tools.")
    add_do(doc,
        "For the parent runtime, replace each RemoteGraph with a stub "
        "CompiledStateGraph in tests - the PregelProtocol abstraction is "
        "exactly what makes this swap clean. The parent's routing logic "
        "is unit-testable without any HF Space being up.")
    add_do(doc,
        "Use graph.nodes['name'] to access individual nodes inside a "
        "compiled graph for targeted tests.")
    add_dont(doc,
        "Test the production graph against the production checkpointer "
        "(MongoDBSaver). Tests use InMemorySaver. MongoDBSaver is "
        "verified by integration env, not by unit tests.")
    add_dont(doc,
        "Mock at .invoke() on the LLM directly - Matt Carvalho's post "
        "documents why this is brittle. Use FakeListLLM or pre-built "
        "LangChain fakes.")
    add_para(doc,
        "Rationale. The official Test doc: 'Compiled LangGraph agents "
        "expose references to each individual node as graph.nodes. You "
        "can take advantage of this to test individual nodes' and "
        "recommends per-test compilation with a fresh checkpointer. "
        "Anirudh Sharma: 'For node-level unit tests, you don't need "
        "pytest - you can call the node function directly, which is "
        "fast and deterministic.' The RemoteGraph-as-stub pattern is a "
        "direct consequence of the PregelProtocol abstraction.")
    for k in ("docs_test", "blog_unit_testing", "blog_mock_llm",
              "ref_remote_graph"):
        add_citation(doc, k)

    # ---- Q9 (substantially rewritten) ----
    add_heading(doc, "Q9. Graph decomposition - architecture-aware", level=2)
    add_para(doc,
        "This question is where the architect's V5 comment lands. The "
        "rewrite below is the substantive change; section 3 above is the "
        "extended reasoning.")
    add_para(doc, "Decomposition rules:", italic=True)
    add_do(doc,
        "Build one subgraph per CAPABILITY Space (FindCare, EvaluateCare, "
        "Security). Each subgraph compiles in its own Space and exposes "
        "the standard Agent Server protocol. The parent runtime composes "
        "them with RemoteGraph.")
    add_do(doc,
        "Build one PARENT graph that lives in (or alongside) the runtime "
        "Space React talks to. The parent owns the messages channel, the "
        "MongoDBSaver, and the routing edges between capability Spaces.")
    add_do(doc,
        "Define each subgraph's wire interface as input_schema and "
        "output_schema (Pydantic). Keys not in the schema do not cross "
        "the boundary - this matters more across Spaces than within.")
    add_do(doc,
        "Hand off between epics with Command(graph=Command.PARENT). The "
        "parent's routing table picks the next epic; the source epic does "
        "not need to know what comes next.")
    add_do(doc,
        "Implement deep_research as a same-Space subgraph imported as a "
        "Python module by whichever capability needs it. NOT its own "
        "Space, NOT its own epic.")
    add_do(doc,
        "Implement shared per-utterance pre-processing (emergency_detector, "
        "ip_lock_check, repetitive_detector, classifier) once - either as "
        "a small subgraph in the parent runtime or as a node-shaped call "
        "into the Security Space (recommended for emergency_detector and "
        "ip_lock_check, which are entitlement-adjacent).")
    add_dont(doc,
        "Build one monolithic graph that contains every node for every "
        "capability. (POC AP-2.)")
    add_dont(doc,
        "Build separately invoked top-level graphs per epic that React "
        "calls directly for conversational flow. The parent runtime is "
        "the single conversational entry point.")
    add_dont(doc,
        "Spin up a graph for a deterministic pipeline that has no LLM "
        "decisions. That should be a tool. The 'Building LangGraph' post: "
        "'little to no abstraction... we focused on control and "
        "durability.' Graphs are for control flow across LLM calls.")
    add_dont(doc,
        "Promote deep_research, the safety subgraph, or the data-pipeline "
        "jobs to their own HF Spaces. The cost-center axiom does not "
        "support a Space line for them.")
    add_dont(doc,
        "Use RemoteGraph to call back into the same deployment - the "
        "official reference warns 'this can lead to deadlocks and "
        "resource exhaustion.' Same-Space composition uses local "
        "subgraph add_node().")
    add_para(doc, "Long-running async work (deep research):", italic=True)
    add_do(doc,
        "Implement deep_research as a reusable subgraph (its own state "
        "schema, its own input_schema = {ResearchQuestion}, output_schema "
        "= {ResearchReport}). Whichever capability Space needs it "
        "imports it as a Python module.")
    add_do(doc,
        "Inside deep_research, fan out to per-source workers using the "
        "Send API; workers run concurrently under asyncio.gather.")
    add_do(doc,
        "Place a wait_for_research node in the parent flow that calls "
        "interrupt() and yields. When research completes, resume with "
        "Command(resume=research_payload). The conversation transcript "
        "receives the consolidated ResearchReport as a system ToolMessage. "
        "The interrupt and resume happen on the parent's checkpointer "
        "(Q2).")
    add_dont(doc,
        "Block the meeting transcript while research runs (Send fan-out "
        "fully inline with await). The architect's requirement is that "
        "meeting participants continue to talk while research runs - "
        "inline blocking violates this.")
    add_para(doc, "Routing decisions as Pydantic schemas:", italic=True)
    add_do(doc,
        "Adopt the deep_research_from_scratch pattern: ClarifyWithUser, "
        "ResearchQuestion, ConductResearch, ResearchComplete are Pydantic "
        "structured outputs that the LLM emits and the graph routes on. "
        "Each is testable in isolation.")
    add_dont(doc,
        "Encode routing decisions as free-form strings or untyped dicts "
        "that the next node has to re-parse. That is precisely the "
        "brittleness Pydantic schemas exist to remove.")
    add_para(doc, "OO principles applied:", italic=True)
    add_do(doc,
        "Apply Single-Responsibility at the Space level - each Space "
        "owns one capability. And at the subgraph level inside a Space - "
        "each subgraph owns a focused slice of state.")
    add_do(doc,
        "Apply Encapsulation via class-based nodes whose dependencies "
        "are constructor-injected.")
    add_do(doc,
        "Prefer Composition over Inheritance - assemble graphs and "
        "subgraphs rather than subclassing node base classes.")
    add_do(doc,
        "Use Dependency Injection via the runtime context (Q5).")
    add_dont(doc,
        "Subclass a custom NodeBase. There is no canonical framework "
        "class to subclass - just classes with __call__.")
    add_para(doc,
        "Rationale. The official handoff doc: 'Use single agent with "
        "middleware for most handoff use cases - it's simpler. Only use "
        "multiple agent subgraphs when you need bespoke agent "
        "implementations.' ChatHealthy's epics ARE bespoke agent "
        "implementations, so the multi-subgraph pattern applies. DeepWiki "
        "on graph composition: a CompiledStateGraph 'implements the "
        "PregelProtocol and can be passed to StateGraph.add_node()' - "
        "i.e., a compiled subgraph IS a node. RemoteGraph extends the "
        "same protocol across the network. State passes by 'selective "
        "projection': downward 'the parent passes only state keys that "
        "exist in the child's input_schema' and upward 'the subgraph "
        "returns values from its output_schema.' Command(graph="
        "Command.PARENT) lets a subgraph direct the parent to the next "
        "epic without the subgraph knowing the parent's routing table. "
        "Long-running work uses interrupt() / Command(resume=value), "
        "requiring 'a checkpointer to persist the graph state.'")
    for k in ("docs_handoffs", "deepwiki_subgraphs", "deepwiki_hitl",
              "docs_interrupts", "docs_subgraphs", "docs_remote_graph",
              "ref_remote_graph", "blog_send_api", "blog_scaling_lg",
              "blog_easton_2026", "blog_handoffs_tds",
              "blog_multiagent_2026", "blog_building",
              "gh_open_deep_research", "gh_deep_research_scratch",
              "gh_deepagents", "gh_swarm", "gh_supervisor",
              "docs_thinking", "docs_add_memory"):
        add_citation(doc, k)

    # ---- Q10 ----
    add_heading(doc, "Q10. Dialogue / message / history modeling", level=2)
    add_do(doc,
        "Model conversation as messages: Annotated[list[AnyMessage], "
        "add_messages] - either by extending the prebuilt MessagesState "
        "or by declaring the field manually on the parent state.")
    add_do(doc,
        "The PARENT state owns the messages channel. Capability subgraphs "
        "see the messages projection that their input_schema declares; "
        "they return ToolMessage entries that flow back through the "
        "output_schema.")
    add_do(doc,
        "Carry tool calls and tool results as ToolMessage entries on "
        "that channel. The LLM sees the tool exchange exactly as it "
        "happened.")
    add_do(doc,
        "Keep non-conversational artefacts (retrieved docs, structured "
        "tool payloads, research reports) in separate typed Pydantic "
        "fields on the subgraph state. They do not need to ride the "
        "messages channel.")
    add_dont(doc,
        "Overload the messages list with retrieved documents, "
        "intermediate reasoning chains, or raw tool payloads. The "
        "'Clean State Architecture' post names this an anti-pattern.")
    add_dont(doc,
        "Model tool outputs as flat top-level fields per tool, or as a "
        "tool_outputs: dict[str, Any]. (See Q4.)")
    add_para(doc,
        "Rationale. The official graph-API doc introduces the pattern: "
        "'LangGraph includes a built-in reducer add_messages' and the "
        "prebuilt MessagesState 'so that we can have' the simpler "
        "pattern. The 'Clean State Architecture' post names 'Overloading "
        "Messages Field' as an anti-pattern: 'The messages list is easy "
        "to append to, and before long it can end up holding retrieved "
        "documents, intermediate reasoning, tool outputs, and more' - "
        "prescribing dedicated fields like retrieved_docs instead.")
    for k in ("docs_graph_api", "blog_clean_state", "docs_tools"):
        add_citation(doc, k)

    # ============================================================
    # 6. Anti-patterns
    # ============================================================
    add_heading(doc, "6. Anti-Patterns to Avoid", level=1)
    add_para(doc,
        "Each anti-pattern below is named in a cited source. Items "
        "marked POC are concretely present in "
        "LangGraph/poc/user_journey.py. Items marked (V5-resolved) came "
        "out of the V2 / V3 / V5 architect dialogue. Items marked "
        "(V6-new) surfaced during the V6 reconciliation and concern the "
        "deployment topology.",
        italic=True)

    add_heading(doc, "AP-1. The God State Object  (POC)", level=3)
    add_para(doc,
        "Source: 'It's tempting to create one big state class that every "
        "node in your graph can read and write... State that belongs to "
        "everyone belongs to no one.' POC reality: JourneyState has 15 "
        "unrelated fields all flat in one TypedDict. Resolution: "
        "Pydantic state per subgraph (Q1, Q9).")
    add_citation(doc, "blog_clean_state")

    add_heading(doc, "AP-2. Monolithic Graph (no subgraph decomposition)  (POC)", level=3)
    add_para(doc,
        "Source: 'I had a single LangGraph graph with 22 nodes and one "
        "huge state object... essentially building a monolith.' POC "
        "reality: 14 nodes and the build_user_journey() function in one "
        "336-line file. Resolution: one capability subgraph per Space, "
        "composed under the parent runtime via RemoteGraph + in-process "
        "subgraph add_node() (Q9).")
    add_citation(doc, "blog_clean_state")

    add_heading(doc, "AP-3. Tool outputs as flat top-level state fields  (POC, V5-resolved)", level=3)
    add_para(doc,
        "Source guidance: tool I/O rides the messages channel as "
        "ToolMessage entries; structured artefacts go into named typed "
        "Pydantic fields on the subgraph state. POC reality: location, "
        "specialty_query, specialties, homeopathic_expansion, providers, "
        "trials are all flat top-level keys. Resolution (V5): "
        "ToolMessage on messages for the conversation; Pydantic fields "
        "on subgraph state for the structured artefacts; NOT a "
        "tool_outputs: dict[str, Any] catch-all.")
    add_citation(doc, "docs_graph_api")
    add_citation(doc, "blog_clean_state")

    add_heading(doc, "AP-4. Procedural top-level node functions with no encapsulation  (POC)", level=3)
    add_para(doc,
        "Source guidance: nodes are pure functions OR classes-with-__call__ "
        "when they have dependencies. POC reality: every node is a "
        "top-level def in user_journey.py with no module boundary, no "
        "class wrappers, and stub bodies. Resolution: Q3 + Q6 file "
        "layout.")
    add_citation(doc, "docs_app_structure")
    add_citation(doc, "blog_yasin_modular")

    add_heading(doc, "AP-5. Module-level imports of clients instead of dependency injection", level=3)
    add_para(doc,
        "Source guidance: 'Runtime context is a form of dependency "
        "injection... optimize the LLM context by providing dependencies "
        "(like database connections, user IDs, or API clients) to your "
        "tools and nodes at runtime rather than hardcoding them.' "
        "Resolution: Q5 - Runtime[ContextSchema] / ToolRuntime.")
    add_citation(doc, "docs_context")

    add_heading(doc, "AP-6. Missing reducers on list fields  (POC risk)", level=3)
    add_para(doc,
        "Source: 'If you don't set a reducer, LangGraph uses last-write-"
        "wins by default... Always give it a reducer, even if only one "
        "node writes to it now.' POC reality: history, specialties, "
        "providers, trials, and homeopathic_expansion are all bare list "
        "with no annotation; the parallel safety + classifier fan-out is "
        "exactly the topology where last-write-wins silently loses data. "
        "Resolution: Q7.")
    add_citation(doc, "blog_clean_state")
    add_citation(doc, "docs_graph_api")

    add_heading(doc, "AP-7. Overloaded messages field", level=3)
    add_para(doc,
        "Source: 'The messages list is easy to append to, and before "
        "long it can end up holding retrieved documents, intermediate "
        "reasoning, tool outputs, and more.' Fix: keep messages for the "
        "LLM-visible conversation (including ToolMessage); store "
        "retrieval / structured tool artefacts in named typed Pydantic "
        "fields. Resolution: Q4 + Q10.")
    add_citation(doc, "blog_clean_state")

    add_heading(doc, "AP-8. Dev-only checkpointer in production  (V5-resolved)", level=3)
    add_para(doc,
        "Source quote (official): 'InMemorySaver: suitable for "
        "development and testing... SqliteSaver: ideal for "
        "experimentation and local workflows... PostgresSaver: ideal for "
        "using in production.' Resolution (V5): the actual anti-pattern "
        "is running on a dev-only checkpointer in production. "
        "PostgresSaver is the LangChain reference for the generic case. "
        "For ChatHealthy, MongoDBSaver from langgraph-checkpoint-mongodb "
        "is the production target because Mongo is already in our stack "
        "(Q2).")
    add_citation(doc, "docs_persistence")
    add_citation(doc, "blog_easton_2026")
    add_citation(doc, "pypi_mongo_ckpt")
    add_citation(doc, "mongo_atlas_lg")

    add_heading(doc, "AP-9. Introducing Postgres only to host the LangGraph checkpointer  (V5-resolved)", level=3)
    add_para(doc,
        "Resolution (V5): do NOT use PostgresSaver when MongoDB is the "
        "existing infrastructure. Adding Postgres for the checkpointer "
        "alone introduces a new database tier without a concrete "
        "requirement. Reserve a possible Postgres introduction for an "
        "explicit future need (e.g. multi-row ACID join with another "
        "Postgres-only system) and document it as a separate decision.")
    add_citation(doc, "pypi_mongo_ckpt")
    add_citation(doc, "mongo_atlas_lg")
    add_citation(doc, "mongo_lc_readthedocs")
    add_citation(doc, "ref_checkpoints")

    add_heading(doc, "AP-10. tool_outputs: dict[str, Any] catch-all on state  (V5-resolved)", level=3)
    add_para(doc,
        "Resolution (V5): do NOT model tool outputs as a single "
        "tool_outputs dict on top-level state - and do NOT model them as "
        "one flat field per tool either. ToolMessage entries on the "
        "messages channel ARE the home for tool calls and tool results.")
    add_citation(doc, "blog_clean_state")
    add_citation(doc, "docs_graph_api")
    add_citation(doc, "docs_tools")

    add_heading(doc, "AP-11. Epics as separately invoked top-level graphs that React calls directly  (V5-resolved)", level=3)
    add_para(doc,
        "Resolution (V5): do NOT have React invoke each capability epic "
        "as its own top-level graph for conversational flow. Compose the "
        "epics under the parent runtime. Separately invoked top-level "
        "graphs are appropriate ONLY when an epic must run on a "
        "different process / deployment / SLA - which is exactly the "
        "case for ChatHealthy, and which V6 resolves with RemoteGraph "
        "rather than direct React-to-Space calls. The parent runtime is "
        "the single conversational entry point; non-conversational read "
        "endpoints (provider detail page) remain direct React-to-Space.")
    add_citation(doc, "docs_handoffs")
    add_citation(doc, "deepwiki_subgraphs")
    add_citation(doc, "blog_easton_2026")

    add_heading(doc, "AP-12. Building a custom MongoDB saver  (V5-resolved)", level=3)
    add_para(doc,
        "Resolution (V5): do NOT build a parallel MongoDB saver. The "
        "officially-supported langgraph-checkpoint-mongodb package "
        "already exists, is co-maintained by MongoDB and LangChain, and "
        "implements the full BaseCheckpointSaver contract. Building a "
        "parallel implementation duplicates work and imposes a tracking "
        "burden against LangGraph's interface drift.")
    add_citation(doc, "pypi_mongo_ckpt")
    add_citation(doc, "ref_checkpoints")
    add_citation(doc, "docs_custom_ckpt")

    # --- V6-new anti-patterns from the reconciliation ---

    add_heading(doc, "AP-13. Promoting deep research to its own HF Space  (V6-new)", level=3)
    add_para(doc,
        "Resolution (V6): per the architect's V5 directive, deep research "
        "is a SUBGRAPH (or a small set of subgraphs) that runs in "
        "whichever capability Space needs it. It is NOT its own "
        "capability cost center, NOT its own HF Space, NOT its own "
        "epic. The cost-center axiom does not support a Space line for "
        "deep research. V4 and V5 implicitly tolerated framing it as a "
        "co-equal capability; V6 rules that framing out.")
    add_citation(doc, "docs_subgraphs")
    add_citation(doc, "deepwiki_subgraphs")

    add_heading(doc, "AP-14. Per-Space MongoDBSaver instances  (V6-new)", level=3)
    add_para(doc,
        "Resolution (V6): do NOT instantiate a MongoDBSaver in each "
        "capability Space. The session checkpoint lives in ONE place - "
        "the parent runtime. Capability Spaces are stateless with "
        "respect to session checkpoints (they may, of course, have "
        "their own MongoDB collections for capability-specific data). "
        "Multiple checkpointers fragment interrupt() / "
        "Command(resume=...) coherence and break the single-thread "
        "guarantee LangGraph's persistence model assumes.")
    add_citation(doc, "docs_persistence")
    add_citation(doc, "docs_interrupts")
    add_citation(doc, "ref_checkpoints")

    add_heading(doc, "AP-15. RemoteGraph calling back into the same deployment  (V6-new)", level=3)
    add_para(doc,
        "Resolution (V6): the official RemoteGraph reference is "
        "explicit: 'Do not use RemoteGraph to call itself or another "
        "graph on the same deployment, as this can lead to deadlocks "
        "and resource exhaustion.' Same-Space composition is in-process "
        "subgraph add_node(); cross-Space composition is RemoteGraph. "
        "Do not mix them.")
    add_citation(doc, "docs_remote_graph")
    add_citation(doc, "ref_remote_graph")

    add_heading(doc, "AP-16. React talking to capability Spaces directly for conversational flow  (V6-new)", level=3)
    add_para(doc,
        "Resolution (V6): React calls the parent runtime for "
        "conversational flow. Conversational flow includes everything "
        "that touches the messages channel, the session checkpoint, the "
        "deep-research interrupt, or epic-to-epic routing. React MAY "
        "call capability Spaces directly for non-conversational read "
        "endpoints (load a provider detail page, render a static "
        "evaluation report, fetch a configuration document) - those are "
        "REST reads, not conversation. Mixing the two paths in React "
        "fragments responsibility and breaks the single-thread "
        "guarantee.")
    add_citation(doc, "docs_handoffs")
    add_citation(doc, "docs_remote_graph")

    add_heading(doc, "AP-17. Modeling data pipelines or operations as LangGraph graphs  (V6-new)", level=3)
    add_para(doc,
        "Resolution (V6): per the architect's V5 directive, 'LangGraph "
        "is only a concept for the runtime application and not the "
        "pipeline or dev ops.' EPIC-012 (Data Management Pipelines) and "
        "EPIC-010 (Operations) are NOT LangGraph territory. Pipelines "
        "are scheduled jobs (cron, GitHub Actions, dedicated workers); "
        "ops is shell scripts and CI configuration. Wrapping them in a "
        "graph adds the BSP / checkpoint / Pregel machinery for no LLM "
        "decision-making.")
    add_citation(doc, "blog_building")
    add_citation(doc, "docs_app_structure")

    # ============================================================
    # 7. Reference implementations
    # ============================================================
    add_heading(doc, "7. Reference Implementations", level=1)
    add_para(doc,
        "Six real GitHub repositories that exemplify the patterns in "
        "this guide. Star counts are from the V4 research pass on "
        "2026-04-28. Grouped as 'primary' (study first) vs 'supporting' "
        "(reference when needed).",
        italic=True)

    add_heading(doc, "Primary references", level=2)

    add_heading(doc, "R-1. langchain-ai/langgraph-example-pyproject  (PRIMARY)", level=3)
    add_para(doc,
        "52 stars but officially curated. The canonical layout: "
        "my_agent/utils/{state.py, nodes.py, tools.py} + my_agent/agent.py "
        "+ langgraph.json. This is the skeleton each ChatHealthy "
        "capability Space should mirror (Q6). Read first because it is "
        "the smallest complete example of the production layout.")
    add_citation(doc, "gh_example_pyproject")
    add_citation(doc, "docs_app_structure")

    add_heading(doc, "R-2. langchain-ai/deep_research_from_scratch  (PRIMARY)", level=3)
    add_para(doc,
        "708 stars. Tutorial repo (5 notebooks: scoping, research_agent, "
        "research_agent_mcp, research_supervisor, full_agent). THE "
        "reference for Pydantic-schemas-as-routing-decisions "
        "(ClarifyWithUser, ResearchQuestion, ConductResearch, "
        "ResearchComplete) and for the supervisor / worker split inside "
        "the deep-research subgraph. Read second because it is the "
        "schema-discipline reference for Q9's deep-research subgraph.")
    add_citation(doc, "gh_deep_research_scratch")

    add_heading(doc, "R-3. langchain-ai/open_deep_research  (PRIMARY)", level=3)
    add_para(doc,
        "11.3k stars. Production-grade research agent. src/ layout, "
        "configurable models for distinct LLM roles (summariser, "
        "researcher, compressor, writer), legacy/ folder preserves an "
        "older workflow + multi-agent supervisor implementation. Read "
        "third because it is the production-hardening reference for the "
        "deep-research subgraph scaffolded from R-2.")
    add_citation(doc, "gh_open_deep_research")

    add_heading(doc, "Supporting references", level=2)

    add_heading(doc, "R-4. langchain-ai/langgraph (the framework itself)", level=3)
    add_para(doc,
        "30.7k stars. Authoritative source for the graph runtime, "
        "checkpointer interfaces, RemoteGraph implementation, and "
        "example notebooks under examples/. Read for: Pregel/BSP model, "
        "channel implementation, BaseCheckpointSaver interface, "
        "PregelProtocol.")
    add_citation(doc, "gh_langgraph")

    add_heading(doc, "R-5. langchain-ai/deepagents", level=3)
    add_para(doc,
        "21.9k stars; latest release 2026-04-15. Agent harness with "
        "planning, filesystem, sub-agent delegation. LangGraph-native; "
        "compiled graphs are returned for downstream composition. Read "
        "for: how a complex agent stays composable.")
    add_citation(doc, "gh_deepagents")

    add_heading(doc, "R-6. wassim249/fastapi-langgraph-agent-production-ready-template", level=3)
    add_para(doc,
        "2.2k stars (community). Realistic FastAPI service skeleton: "
        "app/api/v1/ + app/core/langgraph/ + app/services/ + app/schemas/ "
        "+ app/models/. Demonstrates JWT, rate limiting, Langfuse "
        "observability, mem0+pgvector long-term memory. Read for: how "
        "the 4-file template scales into a deployable service. Note: "
        "ChatHealthy uses MongoDBStore, not pgvector, for long-term "
        "memory; this template's schema is otherwise applicable.")
    add_citation(doc, "gh_fastapi_template")

    add_para(doc,
        "Index of the wider ecosystem (community-curated):",
        italic=True)
    add_citation(doc, "gh_awesome")

    # ============================================================
    # 8. Citations bibliography
    # ============================================================
    add_heading(doc, "8. Citations Bibliography", level=1)

    add_heading(doc, "Official LangChain / LangGraph documentation", level=2)
    for k in ("docs_graph_api", "docs_persistence", "docs_app_structure",
              "docs_tools", "docs_context", "docs_test", "docs_thinking",
              "docs_handoffs", "docs_interrupts", "docs_add_memory",
              "docs_custom_ckpt", "docs_subgraphs", "docs_remote_graph",
              "ref_checkpoints", "ref_remote_graph",
              "pypi_langgraph_sdk", "pypi_langgraph_api"):
        add_citation(doc, k)

    add_heading(doc, "Official LangChain blog", level=2)
    for k in ("blog_building", "blog_v1", "blog_klarna", "blog_in_prod",
              "blog_v02_ckpt"):
        add_citation(doc, k)

    add_heading(doc, "Official MongoDB / LangChain co-maintained sources", level=2)
    for k in ("pypi_mongo_ckpt", "mongo_atlas_lg", "mongo_atlas_lg_agents",
              "mongo_lc_readthedocs", "mongo_blog"):
        add_citation(doc, k)

    add_heading(doc, "Official GitHub repositories", level=2)
    for k in ("gh_langgraph", "gh_example_pyproject", "gh_open_deep_research",
              "gh_deep_research_scratch", "gh_deepagents", "gh_swarm",
              "gh_supervisor", "gh_issue_1950"):
        add_citation(doc, k)

    add_heading(doc, "Community blog posts and tutorials", level=2)
    for k in ("blog_swarnendu", "blog_yasin_modular", "blog_shaza_typing",
              "blog_pankaj_pyd", "blog_easton_2026", "blog_markaicode",
              "blog_clean_state", "blog_unit_testing", "blog_mock_llm",
              "blog_tools_first", "deepwiki_subgraphs", "deepwiki_hitl",
              "blog_send_api", "blog_handoffs_tds", "blog_multiagent_2026",
              "blog_scaling_lg", "dev_mongo_ltm", "blog_bentoml_deploy"):
        add_citation(doc, k)

    add_heading(doc, "Community repositories", level=2)
    for k in ("gh_awesome", "gh_fastapi_template"):
        add_citation(doc, k)

    add_heading(doc, "Recency note", level=2)
    add_para(doc,
        "All cited sources were retrieved on 2026-04-28. None of the "
        "URLs cited here are older than 2024; the official LangChain "
        "blog post 'LangGraph v0.2 checkpointers' is from 2024 and is "
        "the oldest source. Three community pieces (BetterLink Apr-2026, "
        "Vishal Lad Mar-2026, Markaicode 2025) are explicitly 2025-2026, "
        "as is the Mager 2026 multi-agent piece used in Q9. RemoteGraph "
        "documentation (new in V6) is current as of the 2026-04-28 "
        "research pass.")

    add_heading(doc, "Sources searched but no firm answer found", level=2)
    add_para(doc,
        "Klarna's published case study confirms LangGraph use and "
        "checkpointing posture but does not disclose state-schema choice "
        "or directory layout. The LangGraph 1.0 announcement asserts "
        "stability but does not pick TypedDict vs Pydantic. The official "
        "subgraph and handoff docs do not address remote / cross-process "
        "subgraph composition - that gap is filled by the RemoteGraph "
        "reference (use-remote-graph) and the langgraph-api / langgraph-"
        "sdk packages. No claim in this document is sourced from those "
        "gaps.")
    for k in ("blog_klarna", "blog_v1", "docs_subgraphs", "docs_handoffs",
              "docs_remote_graph", "pypi_langgraph_api"):
        add_citation(doc, k)

    # ---------------- Save ----------------
    doc.save(OUT)
    print(f"Wrote: {OUT}")


if __name__ == "__main__":
    build()
