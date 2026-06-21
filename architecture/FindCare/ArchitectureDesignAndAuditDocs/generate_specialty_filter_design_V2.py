"""Generate SpecialtyFilter Design (V1) — end-to-end OO design.

Sections (per Skip's directive):
  1. What SpecialtyFilter is
  2. Source inventory
  3. Why each file moves where
  4. Object model
  5. Flow
  6. RAG-based AI prompt design (collapses V6 S-002 two-stage into one)
  7. Playwright test design (each master business REQ → explicit test classes)
  8. Deployment script changes — exact diffs to deploy YAML, Dockerfile,
     wrapper page; opportunistic Code/-removal migration impact

Substance source: agile_backlog.json EPIC-006-F-002 (V6). Body text below
is design substance authored against V6; V6 docx is now relic.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_ALIGN_VERTICAL

HERE = Path(__file__).resolve().parent
OUT = HERE / "SpecialtyFilterDesign_V2.docx"


def H(doc, text, level=1):
    return doc.add_heading(text, level=level)


def P(doc, text):
    return doc.add_paragraph(text)


def BULLET(doc, text):
    return doc.add_paragraph(text, style="List Bullet")


def NUMBER(doc, text):
    return doc.add_paragraph(text, style="List Number")


def CODE(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    return p


# ── Source-inventory rows (file_name, target, current, what, pattern) ─────────

_DELETING = "DELETING"
_NA = "N/A (new file)"

INVENTORY_ROWS = [
    (
        "filter.py",
        "findCare/specialty_filter/filter.py",
        "Code/ConversationalUX/FindCareChat/backend/domain/find_care/specialty_filter/filter.py",
        "SpecialtyFilter Python class. The RAG-based AI flow (Stage 1 retrieve "
        "context + Stage 2 single LLM call). Loaded once at FindCare backend "
        "startup; called per user query. Implements V6 S-002 REQ-T-001..T-006 "
        "(AI flow, collapsed under the RAG redesign — see section 6).",
        "Has a class (SpecialtyFilter), no __main__. Pure library module.",
    ),
    (
        "__init__.py",
        "findCare/specialty_filter/__init__.py",
        "Code/ConversationalUX/FindCareChat/backend/domain/find_care/specialty_filter/__init__.py",
        "Python package marker. Empty.",
        "Empty.",
    ),
    (
        "specialty_filter.css",
        "findCare/specialty_filter/specialty_filter.css",
        _NA,
        "Visual theme for the SpecialtyFilter wrapper-DOM widget. Owns all "
        ".specialty-filter-* selectors. Greenfield; no inheritance from the "
        "current Website/index.html inline styles. Served by the FindCare HF "
        "Space at /styles/specialty_filter.css.",
        "Pure CSS. No build step. Selectors scoped under .specialty-filter root.",
    ),
    (
        "package.json",
        "sharedServices/widgets/specialty_filter/package.json",
        _NA,
        "npm manifest for the widget bundle build (vite, typescript). "
        "First TS-build entry-point in SharedServices.",
        "JSON manifest.",
    ),
    (
        "vite.config.ts",
        "sharedServices/widgets/specialty_filter/vite.config.ts",
        _NA,
        "Vite build config. Produces dist/specialty_filter.js — a single "
        "ES-module bundle suitable for <script type='module'> loading from the wrapper.",
        "Configuration object. No runtime classes.",
    ),
    (
        "tsconfig.json",
        "sharedServices/widgets/specialty_filter/tsconfig.json",
        _NA,
        "TypeScript compiler options. Strict, ES2022 target, bundler resolution.",
        "Configuration JSON.",
    ),
    (
        "abstract_widget.ts",
        "sharedServices/widgets/specialty_filter/src/abstract_widget.ts",
        _NA,
        "IWidget TypeScript interface + AbstractWidget base class. Future widgets "
        "in SharedServices extend AbstractWidget.",
        "Interface + abstract class. No top-level execution.",
    ),
    (
        "event_bus.ts",
        "sharedServices/widgets/specialty_filter/src/event_bus.ts",
        _NA,
        "EventBus class — typed pub/sub. Used by AbstractWidget.emit() and by "
        "WrapperHost.",
        "One class.",
    ),
    (
        "specialty_filter_widget.ts",
        "sharedServices/widgets/specialty_filter/src/specialty_filter_widget.ts",
        _NA,
        "SpecialtyFilterWidget — the concrete widget class. Extends "
        "AbstractWidget. Owns SpecialtyFilterState + CountsView + "
        "SpecialtyListView + MacroControlBar. Renders into the wrapper's "
        "#specialty-filter-mount node.",
        "One class. Bundle exports it for the wrapper to instantiate.",
    ),
    (
        "specialty_filter_state.ts",
        "sharedServices/widgets/specialty_filter/src/specialty_filter_state.ts",
        _NA,
        "SpecialtyFilterState class. Owns the widget's state: allPossible[], "
        "checked Set<NUCC>, derived counts, apply-button enabled flag. Pure data + "
        "state-mutation methods; no DOM access.",
        "One class.",
    ),
    (
        "counts_view.ts",
        "sharedServices/widgets/specialty_filter/src/counts_view.ts",
        _NA,
        "CountsView class. Renders the three header counts as DOM. Re-renders on "
        "state change.",
        "One class.",
    ),
    (
        "specialty_list_view.ts",
        "sharedServices/widgets/specialty_filter/src/specialty_list_view.ts",
        _NA,
        "SpecialtyListView class. Renders the scrollable specialty rows with "
        "checkboxes. Emits per-row click events upward.",
        "One class.",
    ),
    (
        "macro_control_bar.ts",
        "sharedServices/widgets/specialty_filter/src/macro_control_bar.ts",
        _NA,
        "MacroControlBar class. Renders the Uncheck All / Check All toggle + "
        "Prescribers checkbox + Homeopathic checkbox. Emits macro-toggle events.",
        "One class.",
    ),
    (
        "types.ts",
        "sharedServices/widgets/specialty_filter/src/types.ts",
        _NA,
        "TypeScript interfaces for V6 S-005 API contracts: "
        "SpecialtyClassifyRequest, SpecialtyClassifyResponse, ProviderSearchHandoff, "
        "GeoQualifier, NUCC.",
        "Type-only module — interfaces only.",
    ),
    (
        "index.html",
        "Website/index.html (same path; surgical edits)",
        "Website/index.html",
        "Wrapper page hosted at chathealthy.ai. Surgical edits: REMOVE inline "
        "SpecialtyFilter HTML/JS/CSS (filter-state, count math, toggle logic, "
        "render code). ADD <link href='…findcare…/styles/specialty_filter.css'>, "
        "<script type='module' src='…sharedservices…/widgets/specialty_filter.js'>, "
        "<div id='specialty-filter-mount'></div>. Non-SpecialtyFilter inline code "
        "untouched.",
        "Static HTML with inline JS. Procedural script blocks. SpecialtyFilter "
        "removal subtracts ~hundreds of inline lines; replacement is three small "
        "additive elements.",
    ),
    (
        "app.py",
        "findCare/backend/app.py",
        "Code/ConversationalUX/FindCareChat/backend/app.py",
        "FindCare backend FastAPI entry-point. Wires SpecialtyFilter into "
        "FindCareService. After Code/-removal moves to findCare/backend/. Already "
        "imports via normal package path; no sys.path hacks.",
        "Has __main__ (uvicorn launch) + module-level FastAPI app + service "
        "instantiation.",
    ),
    (
        "specialty_filter.spec.ts",
        "findCare/tests/specialty_filter.spec.ts",
        "Code/ConversationalUX/FindCareChat/tests/specialty_filter.spec.ts",
        "Playwright SIT tests — see Section 7 for the test design tied to each "
        "master business REQ.",
        "Procedural Playwright test() declarations.",
    ),
    (
        "Dockerfile",
        "findCare/backend/Dockerfile",
        "Code/ConversationalUX/FindCareChat/backend/Dockerfile",
        "Container image for FindCare HF Space. Paths shift to findCare/backend/ "
        "post-Code/-removal. The findCare/specialty_filter/ directory is COPYed "
        "in alongside the backend tree.",
        "Dockerfile — declarative build.",
    ),
    (
        "prompts.json",
        "brain/machine_artifacts/content/prompts.json",
        "brain/machine_artifacts/content/prompts.json",
        "Authoritative LLM prompt store. Two existing SpecialtyFilter records "
        "(specialty_normalize_system_prompt, specialty_filter_system_prompt) "
        "are SUPERSEDED by a single new record under the RAG redesign — see "
        "Section 6. Old records can be retired in the same edit; substance of "
        "the new record is in this design.",
        "Single-document JSON.",
    ),
    (
        "agile_backlog.json",
        "brain/machine_artifacts/content/agile_backlog.json",
        "brain/machine_artifacts/content/agile_backlog.json",
        "Authoritative agile backlog. EPIC-006-F-002 V6 substance currently here. "
        "Per the wider refactor: 6 stories will be moved to a new EPIC-009-F-NNN "
        "'Shared front-end widgets' under SharedServices, leaving EPIC-006-F-002 "
        "with the single CSS-theme REQ.",
        "Single-document JSON.",
    ),
    (
        "Code/ (top-level)",
        _DELETING,
        "Code/",
        "After this refactor + the broader Code/-removal cleanup, the top-level "
        "Code/ namespace ceases to exist. SpecialtyFilter-related contents that "
        "currently live under Code/ConversationalUX/FindCareChat/* relocate to "
        "findCare/* per the rules above.",
        "Directory.",
    ),
    (
        "findCare/Code/",
        _DELETING,
        "findCare/Code/ (already deleted in the SpecialtyFilter session)",
        "Defunct shadow namespace. Was a single-child path; removed as part of "
        "the same anomaly cleanup.",
        "Directory.",
    ),
]


def _add_inventory_table(doc):
    table = doc.add_table(rows=1 + len(INVENTORY_ROWS), cols=5)
    table.style = "Light Grid Accent 1"

    headers = ["File name", "Target location", "Current location",
               "What it does", "Pattern of implementation"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

    for r, row in enumerate(INVENTORY_ROWS, start=1):
        for c, val in enumerate(row):
            cell = table.rows[r].cells[c]
            cell.text = val
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    widths_in = [1.4, 2.2, 2.2, 3.2, 2.0]
    for col_idx, w in enumerate(widths_in):
        for cell in table.columns[col_idx].cells:
            cell.width = Inches(w)


def main():
    doc = Document()

    # ── Title and intro ───────────────────────────────────────────────────────
    H(doc, "SpecialtyFilter Design (V2)", level=0)
    P(doc, "Author: Skip Snow + Claude (proposal). V2 — V1 + Section 8 "
           "(deployment script changes) + opportunistic Code/-removal migration "
           "coverage in the file inventory.")
    P(doc, "Substance source: EPIC-006-F-002 V6 (now resident in agile_backlog.json "
           "after the V6 merge). This document is the OO design that implements V6 "
           "end-to-end: class model, runtime topology, sequence flows, state "
           "machine, API contracts, source-tree placement, deploy story, test "
           "approach.")
    P(doc, "Notation: UML class, sequence, state, and component diagrams in ASCII. "
           "OO discipline throughout — interfaces, abstract base classes, concrete "
           "subclasses, dependency direction by composition; no procedural "
           "mount-as-function patterns.")

    # ── 1. What SpecialtyFilter is ─────────────────────────────────────────────
    H(doc, "1. What SpecialtyFilter is", level=1)
    P(doc, "A user-facing widget that lets the user narrow the set of medical "
           "specialties applied to a FindCare provider search. The widget renders "
           "three header counts (All Possible, Prescribers, Your Choices), a "
           "scrollable list of specialty rows with checkboxes, three macro "
           "controls (Uncheck All / Check All toggle, Prescribers checkbox, "
           "Homeopathic checkbox), and an Apply Filter button. The user toggles "
           "checkboxes; counts update reactively; on Apply the wrapper issues a "
           "new provider search using the selected NUCC codes.")
    P(doc, "Architectural placement: SpecialtyFilter is a wrapper-DOM widget "
           "(renders in chathealthy.ai's parent page, not inside any service "
           "iframe). It survives iframe transfers between services. Its substance "
           "is owned by SharedServices; its visual theme is owned by FindCare; "
           "its backend AI flow lives in the FindCare service.")

    # ── 2. Source inventory ────────────────────────────────────────────────────
    H(doc, "2. Source inventory", level=1)
    P(doc, "Every file this refactor touches, creates, or deletes. Target locations "
           "reflect the post-refactor topology (no top-level Code/ namespace; "
           "business-unit dirs at root mirror agile_backlog ownership). For new "
           "files the current location is N/A. For deleted files the target column "
           "shows DELETING and the explanation states where the functionality "
           "relocates.")
    _add_inventory_table(doc)

    # ── 3. Why each file moves where ──────────────────────────────────────────
    H(doc, "3. Why each file moves where", level=1)
    P(doc, "Two principles drive every placement: (a) the source tree mirrors "
           "agile_backlog substance ownership; (b) every file traces to a feature "
           "in the backlog. Where principles (a) and (b) collide with the existing "
           "monolithic structure, the move resolves the collision in favor of "
           "principle (a).")

    H(doc, "3.1 Backend AI flow → findCare/specialty_filter/", level=2)
    P(doc, "The Python class implementing the specialty-matching pipeline (V6 "
           "S-002 REQ-T-001..T-006) is FindCare-domain logic — it answers \"what "
           "specialties are clinically relevant to this user's question?\" That "
           "is FindCare's substance. It runs server-side inside the FindCare HF "
           "Space, called by FindCareService for each user query.")
    P(doc, "Today it lives at "
           "Code/ConversationalUX/FindCareChat/backend/domain/find_care/specialty_filter/filter.py — "
           "an artifact of the broader Code/ shadow namespace. Target: "
           "findCare/specialty_filter/filter.py. After Code/ is dissolved, the "
           "FindCare backend lives at findCare/backend/ and reaches the filter "
           "via a normal package import (no sys.path hacks).")

    H(doc, "3.2 Visual theme (CSS) → findCare/specialty_filter/", level=2)
    P(doc, "The CSS owns the widget's visual identity, which is themed to match "
           "FindCare. If a future widget wears EvaluateCare's theme, that widget's "
           "CSS would ship from EvaluateCare. The CSS file colocates with the "
           "FindCare backend filter so a FindCare engineer changing visual + "
           "behavior touches one directory.")

    H(doc, "3.3 Wrapper-DOM widget → sharedServices/widgets/specialty_filter/", level=2)
    P(doc, "The TypeScript widget classes (SpecialtyFilterWidget, "
           "SpecialtyFilterState, sub-views) render in the parent wrapper DOM and "
           "persist across iframe transfers between FindCare, EvaluateCare, and "
           "SharedServices. They are not service-specific — no service owns the "
           "widget's behavior alone. The closest existing owner of cross-service "
           "wrapper-level concerns is SharedServices. The SpecialtyFilter widget "
           "is the first instance under sharedServices/widgets/.")
    P(doc, "First-time SharedServices gets a TS toolchain (package.json, vite, "
           "tsconfig). The SharedServices HF Space FastAPI app gains a StaticFiles "
           "mount for /widgets serving the built bundle. SharedServices' Python "
           "endpoints stay; the widgets/ subtree is additive.")

    H(doc, "3.4 Wrapper changes → Website/index.html (in place)", level=2)
    P(doc, "The wrapper page itself remains at Website/index.html. The "
           "SpecialtyFilter inline state/render/toggle code is excised "
           "surgically; the file gains a <link>, a <script type='module'>, and "
           "a <div id='specialty-filter-mount'> element. Other inline code in "
           "index.html (non-SpecialtyFilter features) is untouched. Cleaning "
           "the rest of index.html is a separate refactor.")

    H(doc, "3.5 Test stubs → findCare/tests/", level=2)
    P(doc, "Playwright SIT tests follow FindCare-iframe behavior end-to-end. "
           "They live with FindCare for the same colocation reason as backend + "
           "CSS. Pytest is not used for SpecialtyFilter — the substance is 90% "
           "frontend behavioral, and pytest is the wrong tool for that.")

    H(doc, "3.6 Deletions: Code/ and findCare/Code/", level=2)
    P(doc, "After this refactor, Code/ at the repo root no longer exists. Every "
           "file currently under Code/ that this refactor touches relocates to "
           "the appropriate business-unit directory; the rest of Code/ moves in "
           "the broader Code/-removal sweep that runs alongside this refactor. "
           "findCare/Code/ was a single-child shadow path created by an earlier "
           "SpecialtyFilter-move commit; it was removed as the architectural "
           "anomaly was caught.")

    # ── 4. Object model ───────────────────────────────────────────────────────
    H(doc, "4. Object model", level=1)
    P(doc, "Three layers, each with explicit ownership.")

    H(doc, "4.1 Widget hierarchy (SharedServices)", level=2)
    CODE(doc,
         "  ┌─────────────────────── «interface» ────────────────────────┐\n"
         "  │                       IWidget                              │\n"
         "  │ ────────────────────────────────────────────────────────── │\n"
         "  │ + mount(root: HTMLElement, host: IWrapperHost): void       │\n"
         "  │ + update(payload: WidgetUpdate): void                      │\n"
         "  │ + destroy(): void                                          │\n"
         "  └────────────────────────────┬───────────────────────────────┘\n"
         "                               △ implements\n"
         "                               │\n"
         "  ┌────────────────────────────┴──────────────────────────────┐\n"
         "  │ «abstract class» AbstractWidget                            │\n"
         "  │ ───────────────────────────────────────────────────────── │\n"
         "  │ # root: HTMLElement                                       │\n"
         "  │ # host: IWrapperHost                                      │\n"
         "  │ ───────────────────────────────────────────────────────── │\n"
         "  │ + mount(...)        + destroy()                            │\n"
         "  │ # render(): void              ← abstract                   │\n"
         "  │ # emit(evt: WidgetEvent): void                             │\n"
         "  └────────────────────────────┬──────────────────────────────┘\n"
         "                               △ extends\n"
         "                               │\n"
         "  ┌────────────────────────────┴──────────────────────────────┐\n"
         "  │ «class» SpecialtyFilterWidget                              │\n"
         "  │ ────────────────────────────────────────────────────────── │\n"
         "  │ - state:      SpecialtyFilterState                         │\n"
         "  │ - countsView: CountsView                                   │\n"
         "  │ - listView:   SpecialtyListView                            │\n"
         "  │ - macroBar:   MacroControlBar                              │\n"
         "  │ ────────────────────────────────────────────────────────── │\n"
         "  │ + update(payload: SpecialtyResult): void                   │\n"
         "  │ # render(): void                                           │\n"
         "  │ - onUserToggle(code: NUCC): void                           │\n"
         "  │ - onMacroToggle(macro: MacroKind): void                    │\n"
         "  │ - onApplyClick(): void                                     │\n"
         "  └──────────────────────────────────────────────────────────┬─┘\n"
         "       │ owns               │ owns                │ owns      │\n"
         "       ▼                    ▼                     ▼           ▼\n"
         "  SpecialtyFilterState  CountsView         SpecialtyListView  MacroControlBar\n")

    H(doc, "4.2 Wrapper host (Website)", level=2)
    CODE(doc,
         "  ┌──────────────────────── «interface» ────────────────────────┐\n"
         "  │                        IWrapperHost                         │\n"
         "  │ ─────────────────────────────────────────────────────────── │\n"
         "  │ + routeApiCall(req: ServiceRequest): Promise<ServiceResp>   │\n"
         "  │ + on(evt: string, handler): void                            │\n"
         "  │ + off(evt: string, handler): void                           │\n"
         "  └─────────────────────────────┬───────────────────────────────┘\n"
         "                                △ implements\n"
         "  ┌─────────────────────────────┴───────────────────────────────┐\n"
         "  │ «class» WrapperHost                                          │\n"
         "  │ ──────────────────────────────────────────────────────────── │\n"
         "  │ - widgets:       Map<string, IWidget>                        │\n"
         "  │ - serviceRouter: ServiceRouter                               │\n"
         "  │ - eventBus:      EventBus                                    │\n"
         "  │ ──────────────────────────────────────────────────────────── │\n"
         "  │ + register(id: string, w: IWidget): void                     │\n"
         "  │ + mountAll(): void                                           │\n"
         "  │ + routeApiCall(req): Promise<resp>                           │\n"
         "  │ + on / off                                                   │\n"
         "  └──────────────────────────────────────────────────────────────┘\n"
         "       │ owns                          │ owns\n"
         "       ▼                               ▼\n"
         "  ServiceRouter                 EventBus\n")

    H(doc, "4.3 FindCare backend (Python)", level=2)
    CODE(doc,
         "  ┌─────────────────────────────────────────────────────────────┐\n"
         "  │ «class» SpecialtyFilter                                     │\n"
         "  │ ─────────────────────────────────────────────────────────── │\n"
         "  │ - rag:        RAGRetriever                                  │\n"
         "  │ - prompt:     str                  (loaded from prompts.json)│\n"
         "  │ - llm:        AnthropicClient                                │\n"
         "  │ - cache:      ResultCache                                   │\n"
         "  │ ─────────────────────────────────────────────────────────── │\n"
         "  │ + classify(query: str, geo: GeoQualifier): SpecialtyResult  │\n"
         "  │ - retrieve(query: str): list[NuccCandidate]                 │\n"
         "  │ - rank(query: str, candidates: list): list[NuccCandidate]   │\n"
         "  │ - cache_key(query: str, geo: GeoQualifier): str             │\n"
         "  └─────────────────────────────────────────────────────────────┘\n"
         "       │ owns                       │ uses (lazy)\n"
         "       ▼                            ▼\n"
         "  ┌──────────────────────┐    ┌──────────────────────┐\n"
         "  │ «class»              │    │ «class»              │\n"
         "  │ RAGRetriever         │    │ AnthropicClient      │\n"
         "  │ ──────────────────── │    │ ──────────────────── │\n"
         "  │ - vectorIndex        │    │ + complete(prompt)   │\n"
         "  │ - embedClient        │    └──────────────────────┘\n"
         "  │ ──────────────────── │\n"
         "  │ + topK(query, K)     │\n"
         "  └──────────────────────┘\n")

    # ── 5. Flow ────────────────────────────────────────────────────────────────
    H(doc, "5. Flow", level=1)

    H(doc, "5.1 Page load and widget bootstrap", level=2)
    CODE(doc,
         "  Browser   Wrapper.html  WrapperHost  SpecialtyFilterWidget   FindCare CSS    SharedServices JS\n"
         "    │           │              │              │                     │                  │\n"
         "    │ load page │              │              │                     │                  │\n"
         "    ├──────────►│              │              │                     │                  │\n"
         "    │ <link href=…/styles/specialty_filter.css>                     │                  │\n"
         "    ├────────────────────────────────────────────────────────────►│                  │\n"
         "    │ <script type=module src=…/widgets/specialty_filter.js>       │                  │\n"
         "    ├──────────────────────────────────────────────────────────────────────────────►│\n"
         "    │           │ host = new WrapperHost()  │                     │                  │\n"
         "    │           │ w = new SpecialtyFilterWidget()                  │                  │\n"
         "    │           │ host.register('specialty_filter', w)             │                  │\n"
         "    │           │ host.mountAll() → w.mount(rootEl, host)          │                  │\n"
         "    │ blank widget visible       │              │                  │                  │\n"
         "    │◄──────────┤              │              │                     │                  │\n")

    H(doc, "5.2 Initial filter populate (S-002)", level=2)
    CODE(doc,
         "  user query   SpecialtyFilterWidget   WrapperHost     ServiceRouter   FindCare HF (SpecialtyFilter)\n"
         "       │                  │                 │                │             │\n"
         "       │ user submits     │                 │                │             │\n"
         "       │ search question  │                 │                │             │\n"
         "       │ in iframe;       │                 │                │             │\n"
         "       │ iframe postMessages wrapper                          │             │\n"
         "       ├──────────────────┼────────────────►│                │             │\n"
         "       │                  │                 │ POST /specialty │             │\n"
         "       │                  │                 │ /classify       │             │\n"
         "       │                  │                 ├───────────────►│             │\n"
         "       │                  │                 │                │ retrieve(q) │\n"
         "       │                  │                 │                │ → topK NUCC │\n"
         "       │                  │                 │                │ rank(q, K)  │\n"
         "       │                  │                 │                │ → ranked    │\n"
         "       │                  │                 │                │ resp        │\n"
         "       │                  │                 │ resp           │             │\n"
         "       │                  │                 │◄───────────────┤             │\n"
         "       │                  │ widget.update(resp)               │             │\n"
         "       │                  │◄────────────────┤                │             │\n"
         "       │                  │ state.allPossible=resp.items     │             │\n"
         "       │                  │ state.checked = new Set(all)      │             │\n"
         "       │                  │ render()                          │             │\n"
         "       │  populated list  │                 │                │             │\n"
         "       │◄─────────────────┤                 │                │             │\n")

    H(doc, "5.3 User manages filter (S-003)", level=2)
    CODE(doc,
         "  user    SpecialtyFilterWidget     SpecialtyFilterState\n"
         "    │              │                       │\n"
         "    │ click row    │                       │\n"
         "    ├─────────────►│ state.toggle(code)    │\n"
         "    │              ├──────────────────────►│ checked.toggle(code)\n"
         "    │              │ render()              │\n"
         "    │  counts      │ ▲                     │\n"
         "    │  reactive    │ │ recompute counts    │\n"
         "    │              │ │ from state          │\n"
         "    │◄─────────────┤ │                     │\n"
         "    │              │                       │\n"
         "    │ click Macro  │ state.toggleMacro(M)  │\n"
         "    │ (Prescribers/Homeopathic/Uncheck-All)│\n"
         "    ├─────────────►├──────────────────────►│\n"
         "    │              │ render()              │\n"
         "    │◄─────────────┤                       │\n"
         "    │              │                       │\n"
         "    NO backend round-trip during S-003 — all changes operate on cached state.\n")

    H(doc, "5.4 Apply filter — submit (S-004)", level=2)
    CODE(doc,
         "  user    SpecialtyFilterWidget    WrapperHost   ServiceRouter   FindCare HF (Provider Search)\n"
         "    │              │                    │              │             │\n"
         "    │ click Apply  │                    │              │             │\n"
         "    ├─────────────►│ onApplyClick()     │              │             │\n"
         "    │              │ build req from     │              │             │\n"
         "    │              │  state.selected()  │              │             │\n"
         "    │              │ host.routeApi      │              │             │\n"
         "    │              │  Call(req)         │              │             │\n"
         "    │              ├───────────────────►│              │             │\n"
         "    │              │                    │ POST         │             │\n"
         "    │              │                    │ /provider    │             │\n"
         "    │              │                    │ /search      │             │\n"
         "    │              │                    │ {nucc[], geo}│             │\n"
         "    │              │                    ├─────────────►│             │\n"
         "    │              │                    │              │ search runs │\n"
         "    │              │                    │ resp         │             │\n"
         "    │              │                    │◄─────────────┤             │\n"
         "    │              │ resp               │              │             │\n"
         "    │              │◄───────────────────┤              │             │\n"
         "    │              │ wrapper updates iframe (provider list refreshed)│\n"
         "    │  new list    │                    │              │             │\n"
         "    │◄─────────────┤                    │              │             │\n")

    # ── 6. RAG-based AI prompt ────────────────────────────────────────────────
    H(doc, "6. RAG-based AI prompt design", level=1)
    P(doc, "V6 S-002 currently specifies a 7-stage pipeline: AI normalize → "
           "embed → vector search → cosine floor → AI filter → cache. Two "
           "separate LLM calls (normalize + filter), one vector search step, "
           "one floor cutoff. This design REPLACES that with a single retrieval "
           "step plus a single LLM call grounded by RAG.")

    H(doc, "6.1 Hypothesis", level=2)
    P(doc, "If the entire NUCC taxonomy plus our SpecialtyMetaData (which carries "
           "the prescriber and homeopathic flags ChatHealthy authored) lives in a "
           "RAG vector store, then a single retrieval-augmented LLM call can "
           "produce the ranked, flag-aware NUCC list directly. The first stage "
           "(normalize) becomes implicit in the retrieval embedding; the second "
           "stage (filter) becomes explicit in the prompt's output format.")

    H(doc, "6.2 RAG knowledge base", level=2)
    P(doc, "One vector index, populated at build time:")
    BULLET(doc, "Document type 1 — NUCC code: { nucc, displayName, definition, "
                "grouping, classification }. Source: nucc.org provider-taxonomy CSV.")
    BULLET(doc, "Document type 2 — ChatHealthy SpecialtyMetaData: { nucc, "
                "is_prescriber, is_homeopathic, licensure_states[] }. Source: "
                "internal classification job output.")
    P(doc, "Documents are joined on `nucc` so retrieval returns NUCC codes already "
           "carrying their flags. Each document is embedded with the canonical "
           "embedding model (per EPIC-009-F-001-S-012-REQ-T-002 — get_specialty_vector). "
           "The vector store is rebuilt on every NUCC dissemination + every "
           "SpecialtyMetaData refresh; it is otherwise immutable.")

    H(doc, "6.3 Single-call flow", level=2)
    CODE(doc,
         "  user query (sanitized, geo-stripped)\n"
         "      │\n"
         "      ▼\n"
         "  embed(query)  ──►  vector_search(query_vec, K=top_k)\n"
         "                              │\n"
         "                              ▼\n"
         "                  retrieved: list[NuccDocument]  (nucc + flags)\n"
         "                              │\n"
         "                              ▼\n"
         "                  build prompt:\n"
         "                    system_prompt + retrieved_context + user_query\n"
         "                              │\n"
         "                              ▼\n"
         "                  LLM call (Claude / GPT)\n"
         "                              │\n"
         "                              ▼\n"
         "                  parse JSON → ranked NUCC list with flags\n"
         "                              │\n"
         "                              ▼\n"
         "                  cache (key = sha256(query, geo))\n")

    H(doc, "6.4 The prompt (system_prompt for prompts.json)", level=2)
    CODE(doc,
         "  You are a healthcare specialty matching expert.\n"
         "  Given a patient's natural-language description of their healthcare\n"
         "  need and a list of candidate NUCC specialty codes (each with\n"
         "  taxonomy metadata and ChatHealthy flags is_prescriber and\n"
         "  is_homeopathic), return a ranked list of the NUCC codes that\n"
         "  could appropriately treat the described condition.\n"
         "\n"
         "  RULES\n"
         "  1. Inclusion: a code is appropriate if any practitioner with that\n"
         "     NUCC classification could clinically address the user's request.\n"
         "  2. CAM inclusion: complementary, alternative, and integrative\n"
         "     medicine practitioners (Acupuncturist, Naturopath, Naprapath,\n"
         "     Homeopath, Mechanotherapist, Chiropractor, Massage Therapist,\n"
         "     and any other CAM type) MUST be included whenever clinically\n"
         "     relevant. ChatHealthy is CAM-inclusive by editorial policy.\n"
         "  3. Exclusion: do NOT return codes where the role is unlicensed\n"
         "     (e.g., Homemaker, Driver, Lay Midwife, Religious Practitioner)\n"
         "     or where the role is licensed in zero US states/PR.\n"
         "  4. Ranking: order by clinical relevance to the user's specific\n"
         "     request — most natural fit first.\n"
         "  5. Output format: JSON array of objects:\n"
         "       [\n"
         "         { \"nucc\": \"<10-char code>\",\n"
         "           \"displayName\": \"<NUCC classification name>\",\n"
         "           \"is_prescriber\": <bool>,\n"
         "           \"is_homeopathic\": <bool>,\n"
         "           \"rank\": <integer, 1 = best fit> }\n"
         "       ]\n"
         "     No prose, no preamble, no trailing commentary.\n"
         "  6. Use only the candidate codes provided in the retrieval context.\n"
         "     Do not invent codes. Do not return codes outside the candidate\n"
         "     list. If the candidate list contains zero clinically relevant\n"
         "     codes, return an empty array.\n"
         "\n"
         "  RETRIEVAL CONTEXT (top-K NUCC candidates with flags):\n"
         "  {{retrieved_context_block}}\n"
         "\n"
         "  USER REQUEST:\n"
         "  {{user_query}}\n")

    H(doc, "6.5 What collapses from V6 S-002", level=2)
    BULLET(doc, "REQ-T-001 (AI normalize step) — absorbed into the embedding step. "
                "The retrieval embedding does the normalization implicitly.")
    BULLET(doc, "REQ-T-002 (embed normalize output) — replaced by embed(user_query) "
                "directly via the canonical model.")
    BULLET(doc, "REQ-T-003 (vector search SpecialtyMetaData with Section=Individual) — "
                "becomes vector_search against the joined RAG index. Section filter "
                "applied at build time of the index, not query time.")
    BULLET(doc, "REQ-T-004 (cosine candidate floor) — replaced by top-K. K is the "
                "design knob; tune for recall.")
    BULLET(doc, "REQ-T-005 (AI filter step uses prompts.json) — kept, but with the "
                "single combined system prompt above.")
    BULLET(doc, "REQ-T-006 (cache result) — kept verbatim; cache key = "
                "sha256(query + geo).")

    H(doc, "6.6 Trade-offs", level=2)
    BULLET(doc, "Pro: one LLM call instead of two; lower latency, lower cost.")
    BULLET(doc, "Pro: the LLM is grounded in actual NUCC + ChatHealthy data via the "
                "retrieval context; less hallucination risk than the prior normalize+filter chain.")
    BULLET(doc, "Pro: maintenance is simpler — one prompt, one retrieval index, one "
                "cache. The two prompt records in prompts.json collapse to one.")
    BULLET(doc, "Risk: if top-K misses a relevant code, the LLM cannot recover it. "
                "Mitigation: K is tunable; floor recall in the test plan asserts a "
                "minimum recall on a fixed eval set before the design is approved "
                "for production.")
    BULLET(doc, "Risk: retrieval context size scales with K. At K=50 with full "
                "metadata per record, the prompt is ~5 KB — well within model context.")

    H(doc, "6.7 prompts.json edits", level=2)
    BULLET(doc, "ADD record `specialty_rag_system_prompt` with the prompt text "
                "above. Status: pending Skip's approval before write.")
    BULLET(doc, "RETIRE records `specialty_normalize_system_prompt` and "
                "`specialty_filter_system_prompt` (set status=retired; do not "
                "delete — preserve audit trail).")

    # ── 7. Playwright test design ─────────────────────────────────────────────
    H(doc, "7. Playwright test design", level=1)
    P(doc, "Each master business REQ-B in V6 maps to a Playwright test class with "
           "explicit assertions tied verbatim to the REQ. Tests run against the "
           "deployed wrapper at chathealthy.ai (env-specific URL) with the "
           "FindCare iframe live. Visual pixel-comparable assertions are out of "
           "scope here — those are UAT per V6 S-006-REQ-T-004.")

    H(doc, "7.1 Master business REQs (the targets)", level=2)
    BULLET(doc, "EPIC-006-F-002-S-001-REQ-B-001 — Comprehensive UX (the master). "
                "Specifies layout, header rows, count format, body scrollbar, "
                "macro controls, Apply Filter button placement, and reactivity.")
    BULLET(doc, "EPIC-006-F-002-S-002-REQ-B-001 — Present-Initial UX MUST comply "
                "with S-001. Adds: filter populates from the AI flow result; all "
                "boxes default-checked; CAM inclusions present when clinically "
                "relevant; unlicensed roles absent.")
    BULLET(doc, "EPIC-006-F-002-S-003-REQ-B-001 — Manage-Filter UX MUST comply "
                "with S-001. Adds: per-row toggle reactivity; macro toggles "
                "operate on cached state with no network; counts react in real "
                "time; persists across iframe transfers.")
    BULLET(doc, "EPIC-006-F-002-S-004-REQ-B-001 — Submit-Filter UX MUST comply "
                "with S-001. Adds: Apply Filter behaves per the master spec "
                "(greyed-with-popup when no change; on change clears center "
                "panel and triggers new query with selected codes).")

    H(doc, "7.2 Playwright test classes (one per master REQ-B)", level=2)

    H(doc, "Class A: SpecialtyFilterLayout (S-001-REQ-B-001)", level=3)
    BULLET(doc, "Panel renders on the LEFT of the desktop viewport.")
    BULLET(doc, "Panel uses green styling and occupies full vertical height of "
                "the left-panel container.")
    BULLET(doc, "Panel header has THREE rows: title row, count row, controls row.")
    BULLET(doc, "Title row text equals 'Choose Specialties'.")
    BULLET(doc, "Count row contains exactly THREE count cells, each with label "
                "above the value: 'All possible', 'All prescribers', 'Your choices'.")
    BULLET(doc, "Each count value is an integer rendered as 'Label: N'.")
    BULLET(doc, "Controls row contains ONE button ('Uncheck all' / 'Check all') + "
                "TWO checkboxes ('Prescribers', 'Homeopathic') stacked vertically.")
    BULLET(doc, "Body renders specialty rows; each row = one checkbox + one "
                "Display Name on the same horizontal line.")
    BULLET(doc, "Body shows up to 12 rows visible at once; vertical scroll engages "
                "when the list exceeds 12.")
    BULLET(doc, "Apply Filter button is positioned at the bottom of the screen.")

    H(doc, "Class B: SpecialtyFilterPresentInitial (S-002-REQ-B-001)", level=3)
    BULLET(doc, "After a search query, the filter populates with at least one "
                "specialty.")
    BULLET(doc, "Every specialty row defaults to CHECKED state on initial render.")
    BULLET(doc, "The 'All possible' count equals the number of returned rows.")
    BULLET(doc, "The 'Your choices' count equals 'All possible' on initial render "
                "(all default-checked).")
    BULLET(doc, "Submitting an injected query 'foot pain' returns at least the "
                "Acupuncturist, Chiropractor, and Podiatrist NUCC codes "
                "(CAM inclusion for clinically relevant queries).")
    BULLET(doc, "Submitting an injected query 'foot pain' does NOT return "
                "Homemaker, Driver, Religious Nonmedical Practitioner, or Lay "
                "Midwife (unlicensed-role exclusion).")
    BULLET(doc, "Hospitals, labs, agencies, and other non-individual entries are "
                "absent from the result list.")
    BULLET(doc, "The filter never makes a public/protected HTTP call from the "
                "browser. Network panel inspection: zero direct requests to NUCC "
                "or backend public APIs from the widget.")

    H(doc, "Class C: SpecialtyFilterManage (S-003-REQ-B-001)", level=3)
    BULLET(doc, "Clicking a checked row's checkbox makes it unchecked.")
    BULLET(doc, "Clicking the same checkbox a second time makes it checked again.")
    BULLET(doc, "'Your choices' count decrements by exactly 1 when one row is "
                "unchecked; increments by 1 when re-checked.")
    BULLET(doc, "'Prescribers' count tracks the prescriber-flagged subset of "
                "currently-showing items, reactively.")
    BULLET(doc, "'All possible' count is constant during management — it equals "
                "the initial returned count and never changes during manage operations.")
    BULLET(doc, "Toggling 'Prescribers' checkbox in the controls row changes the "
                "set of visible/checked specialties per the master spec.")
    BULLET(doc, "Toggling 'Homeopathic' checkbox does the same per its spec.")
    BULLET(doc, "Toggling EITHER macro checkbox triggers ZERO network requests. "
                "Verified via network panel monitor.")
    BULLET(doc, "When fewer than 50% of the rows are unchecked, the "
                "Uncheck-All / Check-All button label reads 'Check All' and "
                "clicking it checks every row.")
    BULLET(doc, "When ≥50% are unchecked, the button label reads 'Uncheck All' "
                "and clicking it unchecks every row.")
    BULLET(doc, "When zero rows are checked, the button is greyed out and shows "
                "tooltip 'No specialties are checked'.")
    BULLET(doc, "After triggering a service-control transfer (FindCare → "
                "EvaluateCare → SharedServices), the SpecialtyFilter panel is "
                "still rendered with its prior state intact.")

    H(doc, "Class D: SpecialtyFilterSubmit (S-004-REQ-B-001)", level=3)
    BULLET(doc, "When the user has made no changes since render, Apply Filter is "
                "either invisible or greyed out and a click produces a non-modal "
                "popup with text 'There are no changes to apply, try a new prompt "
                "if this is the wrong list. The AI agent will help you craft the "
                "right question'. The popup has an OK button.")
    BULLET(doc, "When the user has made at least one change, Apply Filter is "
                "active. Clicking it: (a) clears the visible provider list to "
                "white/empty in the center panel; (b) shows a timer in the bottom "
                "control frame in the same position as the original-search timer; "
                "(c) issues exactly one network request to the Provider Search "
                "API with payload {selected_nucc_codes: list[NUCC], geo: "
                "GeoQualifier}.")
    BULLET(doc, "After the response arrives, the provider list re-renders with "
                "the new result set.")
    BULLET(doc, "The selected-provider window in the bottom panel is unchanged "
                "by Apply Filter (preserved across the click).")
    BULLET(doc, "The chat prompt input in the bottom panel remains available "
                "during and after the Apply.")

    H(doc, "7.3 Test infrastructure", level=2)
    BULLET(doc, "Spec file: findCare/tests/specialty_filter.spec.ts. Single file "
                "carries all four classes A–D as Playwright `describe` blocks.")
    BULLET(doc, "Each test_id in agile_backlog points at a single test() in the "
                "spec, format `<REQ-id>-TEST-1`. The agile_backlog REQ → test "
                "binding is 1:1.")
    BULLET(doc, "Tests run against env-specific wrapper URLs (localhost / "
                "dev.chathealthy.ai / qa.chathealthy.ai / chathealthy.ai). The "
                "test plan inherits S-006 macro-variance — chrome assertions "
                "skip when env=prod per S-004 of the SpecialtyFilter parent epic.")
    BULLET(doc, "AI-flow integration assertions (Class B 'foot pain' returns "
                "Acupuncturist + Chiropractor; excludes Homemaker) are tied to "
                "the RAG single-prompt flow (Section 6). The test fixture "
                "submits a known query and asserts membership in the returned "
                "list — does NOT assert ranking position (rank is non-deterministic "
                "by design until the eval set freezes).")

    H(doc, "7.4 What is intentionally NOT tested by Playwright", level=2)
    BULLET(doc, "Pixel-comparable visual screenshots — UAT per V6 S-006-REQ-T-004; "
                "human eyes at this maturity, not Playwright pixel diff.")
    BULLET(doc, "Backend internals of SpecialtyFilter.classify (RAG retrieval "
                "mechanics, LLM prompt structure) — unit-tested separately if "
                "needed, not via Playwright.")
    BULLET(doc, "Schema validation of agile_backlog / bugs / prompts JSONs — "
                "covered by Rule-008 pre-commit hook, not Playwright.")
    BULLET(doc, "Rank-order stability across LLM calls — non-deterministic by "
                "design; tests assert membership, not order.")

    # ── 8. Deployment script changes ──────────────────────────────────────────
    H(doc, "8. Deployment script changes", level=1)
    P(doc, "This refactor changes the source-tree topology in two ways: (a) the "
           "FindCare backend moves from Code/ConversationalUX/FindCareChat/backend/ "
           "to findCare/backend/ and is joined by a new sibling findCare/specialty_filter/; "
           "(b) SharedServices gains a TypeScript widget bundle subtree at "
           "sharedServices/widgets/specialty_filter/. Both changes propagate "
           "into the deploy workflows, the Dockerfile, and the wrapper page. "
           "This section enumerates every edit needed in deploy artifacts so "
           "that Engineering can apply them in a single CI cycle.")
    P(doc, "Convention reaffirmed (per Skip's deploy-script directive): deploy "
           "scripts are bare-bones gospel truth. They MUST fail loudly when the "
           "file tree shifts under them. No fallback branches, no try-the-old-"
           "path-then-the-new-path. When the topology moves, the YAML moves with "
           "it in the same commit; CI tells the operator immediately if a path "
           "is wrong.")

    H(doc, "8.1 .github/workflows/deploy-findcare-backend.yml", level=2)
    P(doc, "Five precise edits, all under existing steps.")

    H(doc, "8.1.1 Path triggers (lines 10-17)", level=3)
    CODE(doc,
         "  # BEFORE\n"
         "    paths:\n"
         "      - 'Code/ConversationalUX/FindCareChat/backend/**'\n"
         "      - 'Code/ConversationalUX/FindCareChat/frontend/**'\n"
         "      - 'Code/ConversationalUX/ChatHealthyWhoAmIChat/me/**'\n"
         "      - 'Code/Shared/ChatHealthyMongoUtilities.py'\n"
         "      - 'Code/Shared/prompt_system_maker.py'\n"
         "      - 'brain/machine_artifacts/content/emergency_keywords.json'\n"
         "      - '.github/workflows/deploy-findcare-backend.yml'\n"
         "\n"
         "  # AFTER\n"
         "    paths:\n"
         "      - 'findCare/backend/**'\n"
         "      - 'findCare/frontend/**'\n"
         "      - 'findCare/specialty_filter/**'        # NEW - backend filter + theme CSS\n"
         "      - 'shared/me/**'                        # was Code/ConversationalUX/ChatHealthyWhoAmIChat/me/\n"
         "      - 'shared/mongo/ChatHealthyMongoUtilities.py'\n"
         "      - 'shared/prompts/prompt_system_maker.py'\n"
         "      - 'brain/machine_artifacts/content/emergency_keywords.json'\n"
         "      - 'brain/machine_artifacts/content/prompts.json'  # NEW - RAG prompt\n"
         "      - '.github/workflows/deploy-findcare-backend.yml'\n")
    P(doc, "Code/Shared/ contents move under shared/* during the Code/-removal "
           "sweep. The sweep is the same commit as this refactor's first deploy; "
           "until then the path triggers reference the future locations and "
           "Engineering must move the files before the deploy fires.")

    H(doc, "8.1.2 Frontend build working-directory (line 57)", level=3)
    CODE(doc,
         "  # BEFORE\n"
         "        working-directory: Code/ConversationalUX/FindCareChat/frontend\n"
         "\n"
         "  # AFTER\n"
         "        working-directory: findCare/frontend\n")

    H(doc, "8.1.3 Cert read paths (lines 150-153)", level=3)
    CODE(doc,
         "  # BEFORE\n"
         "          FINDCARE_CERT_PEM_B64=$(base64 -w 0 Code/Shared/ops/certs/findcare.crt)\n"
         "          SHARED_CERT_PEM_B64=$(base64 -w 0 Code/Shared/ops/certs/shared.crt)\n"
         "          EVALCARE_CERT_PEM_B64=$(base64 -w 0 Code/Shared/ops/certs/evalcare.crt)\n"
         "          CA_CERT_PEM_B64=$(base64 -w 0 Code/Shared/ops/certs/ca.crt)\n"
         "\n"
         "  # AFTER\n"
         "          FINDCARE_CERT_PEM_B64=$(base64 -w 0 shared/ops/certs/findcare.crt)\n"
         "          SHARED_CERT_PEM_B64=$(base64 -w 0 shared/ops/certs/shared.crt)\n"
         "          EVALCARE_CERT_PEM_B64=$(base64 -w 0 shared/ops/certs/evalcare.crt)\n"
         "          CA_CERT_PEM_B64=$(base64 -w 0 shared/ops/certs/ca.crt)\n")

    H(doc, "8.1.4 rsync stages (lines 192-230)", level=3)
    CODE(doc,
         "  # BEFORE - flat under hf_space/Code/...\n"
         "          mkdir -p hf_space/Code/ConversationalUX/FindCareChat/backend\n"
         "          rsync -av Code/ConversationalUX/FindCareChat/backend/ \\\n"
         "                    hf_space/Code/ConversationalUX/FindCareChat/backend/\n"
         "          cp -r Code/ConversationalUX/FindCareChat/frontend/dist \\\n"
         "                hf_space/Code/ConversationalUX/FindCareChat/backend/static\n"
         "          mkdir -p hf_space/Code/Shared\n"
         "          rsync -av Code/Shared/ hf_space/Code/Shared/\n"
         "          mkdir -p hf_space/Code/ConversationalUX/ChatHealthyWhoAmIChat/me\n"
         "          rsync -av Code/ConversationalUX/ChatHealthyWhoAmIChat/me/ \\\n"
         "                    hf_space/Code/ConversationalUX/ChatHealthyWhoAmIChat/me/\n"
         "\n"
         "  # AFTER - flat under hf_space/findCare/* + hf_space/shared/*\n"
         "          mkdir -p hf_space/findCare/backend\n"
         "          rsync -av findCare/backend/ hf_space/findCare/backend/\n"
         "\n"
         "          mkdir -p hf_space/findCare/specialty_filter         # NEW\n"
         "          rsync -av findCare/specialty_filter/ \\\n"
         "                    hf_space/findCare/specialty_filter/        # NEW\n"
         "\n"
         "          cp -r findCare/frontend/dist hf_space/findCare/backend/static\n"
         "\n"
         "          mkdir -p hf_space/shared\n"
         "          rsync -av shared/ hf_space/shared/\n")
    P(doc, "Brain artifacts rsync (lines 217-218) is unchanged - brain/ is "
           "already a top-level directory and the path triggers already "
           "reference brain/machine_artifacts/content/prompts.json so a "
           "prompt-only edit triggers a redeploy.")

    H(doc, "8.1.5 Dockerfile copy + smoke path (lines 232-233, 318)", level=3)
    CODE(doc,
         "  # BEFORE\n"
         "          cp Code/ConversationalUX/FindCareChat/backend/Dockerfile \\\n"
         "             hf_space/Dockerfile\n"
         "  ...\n"
         "          python -m pytest Code/deploy/localSmokeTestPyTest.py -v\n"
         "\n"
         "  # AFTER\n"
         "          cp findCare/backend/Dockerfile hf_space/Dockerfile\n"
         "  ...\n"
         "          python -m pytest deploy/localSmokeTestPyTest.py -v\n")

    H(doc, "8.2 .github/workflows/deploy-shared-services.yml", level=2)
    P(doc, "Three edits: path triggers, new TypeScript build step, rsync gains "
           "the widget bundle.")

    H(doc, "8.2.1 Path triggers (lines 10-12)", level=3)
    CODE(doc,
         "  # BEFORE\n"
         "    paths:\n"
         "      - 'sharedServices/Code/**'\n"
         "      - '.github/workflows/deploy-shared-services.yml'\n"
         "\n"
         "  # AFTER\n"
         "    paths:\n"
         "      - 'sharedServices/Code/**'\n"
         "      - 'sharedServices/widgets/**'                       # NEW - TS widgets\n"
         "      - '.github/workflows/deploy-shared-services.yml'\n")

    H(doc, "8.2.2 New TypeScript build step (insert after Set up Python, line 45)", level=3)
    CODE(doc,
         "      - name: Set up Node\n"
         "        uses: actions/setup-node@v4\n"
         "        with:\n"
         "          node-version: 20\n"
         "\n"
         "      - name: Build SpecialtyFilter widget bundle\n"
         "        working-directory: sharedServices/widgets/specialty_filter\n"
         "        run: |\n"
         "          npm ci\n"
         "          npm run build\n"
         "          # Output: dist/specialty_filter.js (single ES-module bundle).\n"
         "          # Vite copies it into dist/; rsync below stages dist/ into\n"
         "          # the HF build context under widgets/specialty_filter/.\n")

    H(doc, "8.2.3 rsync stage for the widget bundle (after line 148)", level=3)
    CODE(doc,
         "  # BEFORE\n"
         "          rsync -av --exclude='.env' --exclude='__pycache__' \\\n"
         "                    --exclude='*.pyc' \\\n"
         "                    sharedServices/Code/ hf_space/\n"
         "\n"
         "  # AFTER\n"
         "          rsync -av --exclude='.env' --exclude='__pycache__' \\\n"
         "                    --exclude='*.pyc' \\\n"
         "                    sharedServices/Code/ hf_space/\n"
         "\n"
         "          # NEW - stage built widget bundle to /widgets/specialty_filter/\n"
         "          mkdir -p hf_space/widgets/specialty_filter\n"
         "          rsync -av sharedServices/widgets/specialty_filter/dist/ \\\n"
         "                    hf_space/widgets/specialty_filter/\n")

    P(doc, "SharedServices' app.py gains a StaticFiles mount alongside its "
           "existing Python routes:")
    CODE(doc,
         "  # SharedServices app.py - add at app-bootstrap time\n"
         "  from fastapi.staticfiles import StaticFiles\n"
         "  app.mount(\"/widgets\",\n"
         "            StaticFiles(directory=\"widgets\"),\n"
         "            name=\"widgets\")\n"
         "  # Wrapper now resolves\n"
         "  #   https://skipsnow-{env-}sharedservicesspace.hf.space\n"
         "  #     /widgets/specialty_filter/specialty_filter.js\n")

    H(doc, "8.3 findCare/backend/Dockerfile", level=2)
    P(doc, "Build context stays the repo root (per V11 EPIC-008-F-004-S-002-"
           "REQ-T-001 + BUG-003 #1). COPY paths shift from Code/* to "
           "findCare/* and shared/*. The findCare/specialty_filter/ subtree "
           "joins the COPY list so the runtime backend sees the filter Python "
           "module and CSS file.")
    CODE(doc,
         "  # BEFORE - pre-refactor Dockerfile\n"
         "  COPY Code/ConversationalUX/FindCareChat/backend/requirements.txt \\\n"
         "       /app/requirements.txt\n"
         "  RUN pip install --no-cache-dir -r /app/requirements.txt\n"
         "\n"
         "  COPY Code/ConversationalUX/FindCareChat/backend \\\n"
         "       /app/Code/ConversationalUX/FindCareChat/backend\n"
         "  COPY Code/Shared /app/Code/Shared\n"
         "  COPY Code/ConversationalUX/ChatHealthyWhoAmIChat/me \\\n"
         "       /app/Code/ConversationalUX/ChatHealthyWhoAmIChat/me\n"
         "  COPY brain /app/brain\n"
         "  ...\n"
         "  WORKDIR /app/Code/ConversationalUX/FindCareChat/backend\n"
         "\n"
         "  # AFTER - post-refactor\n"
         "  COPY findCare/backend/requirements.txt /app/requirements.txt\n"
         "  RUN pip install --no-cache-dir -r /app/requirements.txt\n"
         "\n"
         "  COPY findCare/backend          /app/findCare/backend\n"
         "  COPY findCare/specialty_filter /app/findCare/specialty_filter   # NEW\n"
         "  COPY shared                    /app/shared\n"
         "  COPY brain                     /app/brain\n"
         "  COPY FrontEndApplicationLib    /app/FrontEndApplicationLib\n"
         "  RUN pip install --no-cache-dir /app/FrontEndApplicationLib\n"
         "  ...\n"
         "  WORKDIR /app/findCare/backend\n")
    P(doc, "Inside app.py the corresponding sys.path.insert(...) removals + "
           "import-path updates are described in Section 3.1; they are not "
           "deploy-script changes but are coupled to the Dockerfile move.")

    H(doc, "8.4 Wrapper page deploy (Website/index.html via deploy-findcare-website-*.yml)", level=2)
    P(doc, "There are three website workflows (dev/qa/prod). Each one stages "
           "Website/index.html into the wrapper deploy target. Two changes: "
           "(a) env-derived URLs for the new <link> and <script> tags so each "
           "env points at its matching FindCare and SharedServices HF Space; "
           "(b) inline SpecialtyFilter HTML/JS/CSS removed in source, replaced "
           "by the three additive elements in Section 2.")
    CODE(doc,
         "  # Pattern (applies in all three deploy-findcare-website-*.yml)\n"
         "  # During the staging step that templates Website/index.html:\n"
         "  ENV=\"${{ steps.env.outputs.env }}\"   # dev | qa | prod\n"
         "  if [ \"$ENV\" = \"prod\" ]; then PREFIX=\"\"; else PREFIX=\"${ENV}-\"; fi\n"
         "  FINDCARE_URL=\"https://skipsnow-${PREFIX}chathealthyspace.hf.space\"\n"
         "  SHARED_URL=\"https://skipsnow-${PREFIX}sharedservicesspace.hf.space\"\n"
         "\n"
         "  # Substitute placeholders in index.html\n"
         "  sed -i \"s|@@FINDCARE_URL@@|${FINDCARE_URL}|g\" Website/index.html\n"
         "  sed -i \"s|@@SHARED_URL@@|${SHARED_URL}|g\"   Website/index.html\n"
         "\n"
         "  # index.html source contains:\n"
         "  #   <link rel='stylesheet'\n"
         "  #         href='@@FINDCARE_URL@@/styles/specialty_filter.css'>\n"
         "  #   <script type='module'\n"
         "  #     src='@@SHARED_URL@@/widgets/specialty_filter/specialty_filter.js'>\n"
         "  #   </script>\n"
         "  #   <div id='specialty-filter-mount'></div>\n")
    P(doc, "Substitution at deploy time keeps the source HTML free of "
           "env-specific URLs while letting each website deploy land at the "
           "right pair of HF Space URLs. No client-side runtime URL discovery; "
           "the wrapper is static after deploy.")

    H(doc, "8.5 FindCare app.py StaticFiles mount", level=2)
    P(doc, "FindCare backend gains one StaticFiles mount so the wrapper's "
           "<link> tag resolves. The CSS lives in findCare/specialty_filter/ "
           "(colocated with the backend filter); the mount path /styles serves "
           "the entire specialty_filter directory's static assets.")
    CODE(doc,
         "  # findCare/backend/app.py - add at app-bootstrap time\n"
         "  from fastapi.staticfiles import StaticFiles\n"
         "  app.mount(\"/styles\",\n"
         "            StaticFiles(directory=\"../specialty_filter\"),\n"
         "            name=\"styles\")\n"
         "  # WORKDIR is /app/findCare/backend; ../specialty_filter resolves\n"
         "  # to /app/findCare/specialty_filter (Section 8.3 COPY).\n"
         "  # Wrapper resolves\n"
         "  #   https://skipsnow-{env-}chathealthyspace.hf.space\n"
         "  #     /styles/specialty_filter.css\n")

    H(doc, "8.6 Code/-removal sweep impact", level=2)
    P(doc, "The opportunistic Code/-removal sweep that runs alongside this "
           "refactor breaks the existing pipelines if not coordinated. The "
           "complete inventory of paths to update across the workflow suite:")
    BULLET(doc, "deploy-evaluatecare-backend.yml - same shape as deploy-findcare-"
                "backend.yml; will need the same Code/* -> evaluateCare/* + "
                "shared/* migration when EvaluateCare moves.")
    BULLET(doc, "deploy-pipelines.yml - already migrated from Code/DataPipelines/ "
                "to pipeline/Code/ in commit 59f4c7f; no SpecialtyFilter impact.")
    BULLET(doc, "promote-build.yml + promote-data.yml - Code/-touching paths "
                "are limited to the smoke-test invocation; switch to deploy/* "
                "in the same edit.")
    BULLET(doc, "test.yml - Code/-rooted pytest paths follow the file moves; "
                "pytest discovery roots update from Code/ to the business-unit "
                "directories.")
    P(doc, "Single-commit discipline: every workflow file that references a "
           "moved path is edited in the same commit as the file move. CI red "
           "across multiple workflows immediately tells the operator a path "
           "was missed; staggered moves create a window where some deploys "
           "succeed and some fail silently. Per the deploy-scripts-bare-bones "
           "directive, that window is forbidden.")

    H(doc, "8.7 Order of operations for Engineering", level=2)
    NUMBER(doc, "Land the file moves on a feature branch off dev. Run the full "
                "workflow suite locally (act / yamllint) before pushing.")
    NUMBER(doc, "Update each YAML in the same diff as the corresponding file "
                "move. Do NOT split moves and YAML edits across commits.")
    NUMBER(doc, "Push to dev. The dev deploy fires for FindCare backend, "
                "SharedServices, and the dev FindCare website. All three must "
                "go green before merging onward.")
    NUMBER(doc, "Verify wrapper at dev.chathealthy.ai loads the SpecialtyFilter "
                "widget end-to-end (style + behavior + AI flow). This is the "
                "smoke equivalent of Class A-D Playwright runs (Section 7).")
    NUMBER(doc, "Promote dev -> qa -> prod the same EOD per the enterprise "
                "promote chain. Each promotion is the same build artifact; no "
                "rebuilds, no per-env source forks.")

    doc.save(str(OUT))
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
