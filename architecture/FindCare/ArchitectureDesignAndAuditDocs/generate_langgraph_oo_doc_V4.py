"""Generate langgraph-oo-best-practices-V4.docx.

V4 is the FINAL consolidated synthesis of V1 (research), V2 (architect
annotations), and V3 (annotation resolution). It uses V1's structural shape
(executive summary -> 10 per-question findings -> anti-patterns -> reference
implementations -> bibliography) but folds V2/V3's resolved dialogue back into
the relevant sections. Every recommendation is now definitive: "Do this" /
"Do NOT do this" with a citation. No "you could consider" or "options
include" - the dialogue is closed.

Tag convention in citations:
    [official]  - LangChain / LangGraph / MongoDB co-maintained
    [community] - third-party blog, tutorial, or independent repo
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor


OUT = Path(__file__).with_name("langgraph-oo-best-practices-V4.docx")


# ============================================================
# Citations dictionary - full V1+V3 set carried forward.
# ============================================================

CITES = {
    # --- LangChain / LangGraph official docs ---
    "docs_graph_api":        ("official", "Use the graph API - LangChain Docs",
                              "https://docs.langchain.com/oss/python/langgraph/use-graph-api"),
    "docs_persistence":      ("official", "Persistence - LangChain Docs",
                              "https://docs.langchain.com/oss/javascript/langgraph/persistence"),
    "docs_app_structure":    ("official", "Application structure - LangChain Docs",
                              "https://docs.langchain.com/oss/python/langgraph/application-structure"),
    "docs_tools":            ("official", "Tools - LangChain Docs",
                              "https://docs.langchain.com/oss/python/langchain/tools"),
    "docs_context":          ("official", "Context overview - LangChain Docs",
                              "https://docs.langchain.com/oss/python/concepts/context"),
    "docs_test":             ("official", "Test - LangChain Docs",
                              "https://docs.langchain.com/oss/python/langgraph/test"),
    "docs_thinking":         ("official", "Thinking in LangGraph - LangChain Docs",
                              "https://docs.langchain.com/oss/python/langgraph/thinking-in-langgraph"),
    "docs_handoffs":         ("official", "Multi-agent handoffs - LangChain Docs",
                              "https://docs.langchain.com/oss/python/langchain/multi-agent/handoffs"),
    "docs_interrupts":       ("official", "Interrupts - LangChain Docs",
                              "https://docs.langchain.com/oss/python/langgraph/interrupts"),
    "docs_add_memory":       ("official", "Memory (add-memory) - LangChain Docs",
                              "https://docs.langchain.com/oss/python/langgraph/add-memory"),
    "docs_custom_ckpt":      ("official", "How to use a custom checkpointer - LangChain Docs",
                              "https://docs.langchain.com/langsmith/custom-checkpointer"),
    "ref_checkpoints":       ("official", "checkpoints API reference - LangGraph",
                              "https://reference.langchain.com/python/langgraph/checkpoints"),

    # --- LangChain official blog ---
    "blog_building":         ("official", "Building LangGraph: an agent runtime from first principles",
                              "https://www.langchain.com/blog/building-langgraph"),
    "blog_v1":               ("official", "LangChain and LangGraph reach v1.0",
                              "https://www.langchain.com/blog/langchain-langgraph-1dot0"),
    "blog_klarna":           ("official", "How Klarna's AI assistant redefined customer support",
                              "https://www.langchain.com/blog/customers-klarna"),
    "blog_in_prod":          ("official", "Is LangGraph used in production?",
                              "https://www.langchain.com/blog/is-langgraph-used-in-production"),
    "blog_v02_ckpt":         ("official", "LangGraph v0.2: customization with new checkpointer libraries",
                              "https://blog.langchain.com/langgraph-v0-2/"),

    # --- MongoDB / LangChain co-maintained sources ---
    "pypi_mongo_ckpt":       ("official", "langgraph-checkpoint-mongodb on PyPI (MongoDB + LangChain co-maintained)",
                              "https://pypi.org/project/langgraph-checkpoint-mongodb/"),
    "mongo_atlas_lg":        ("official", "Integrate MongoDB with LangGraph - MongoDB Atlas Docs",
                              "https://www.mongodb.com/docs/atlas/ai-integrations/langgraph/"),
    "mongo_atlas_lg_agents": ("official", "Build an AI Agent with LangGraph and MongoDB Atlas",
                              "https://www.mongodb.com/docs/atlas/ai-integrations/langgraph/build-agents/"),
    "mongo_lc_readthedocs":  ("official", "MongoDBSaver class - LangChain MongoDB Read the Docs",
                              "https://langchain-mongodb.readthedocs.io/en/latest/langgraph_checkpoint_mongodb/saver/langgraph.checkpoint.mongodb.saver.MongoDBSaver.html"),
    "mongo_blog":            ("official", "Checkpointers and Native Parent-Child Retrievers with LangChain and MongoDB",
                              "https://www.mongodb.com/company/blog/innovation/checkpointers-native-parent-child-retrievers-with-langchain-mongodb"),

    # --- Official GitHub repositories ---
    "gh_langgraph":          ("official", "langchain-ai/langgraph",
                              "https://github.com/langchain-ai/langgraph"),
    "gh_example_pyproject":  ("official", "langchain-ai/langgraph-example-pyproject",
                              "https://github.com/langchain-ai/langgraph-example-pyproject"),
    "gh_open_deep_research": ("official", "langchain-ai/open_deep_research",
                              "https://github.com/langchain-ai/open_deep_research"),
    "gh_deep_research_scratch": ("official", "langchain-ai/deep_research_from_scratch",
                              "https://github.com/langchain-ai/deep_research_from_scratch"),
    "gh_deepagents":         ("official", "langchain-ai/deepagents",
                              "https://github.com/langchain-ai/deepagents"),
    "gh_swarm":              ("official", "langchain-ai/langgraph-swarm-py",
                              "https://github.com/langchain-ai/langgraph-swarm-py"),
    "gh_supervisor":         ("official", "langchain-ai/langgraph-supervisor-py",
                              "https://github.com/langchain-ai/langgraph-supervisor-py"),
    "gh_issue_1950":         ("official", "Issue #1950: Passing private state with Class Node",
                              "https://github.com/langchain-ai/langgraph/issues/1950"),

    # --- Community repositories ---
    "gh_awesome":            ("community", "von-development/awesome-LangGraph",
                              "https://github.com/von-development/awesome-LangGraph"),
    "gh_fastapi_template":   ("community", "wassim249/fastapi-langgraph-agent-production-ready-template",
                              "https://github.com/wassim249/fastapi-langgraph-agent-production-ready-template"),

    # --- Community blog posts ---
    "blog_swarnendu":        ("community", "Swarnendu De - LangGraph Best Practices",
                              "https://www.swarnendu.de/blog/langgraph-best-practices/"),
    "blog_yasin_modular":    ("community", "Yassin Hashem - Scaling AI Agents Beyond Notebooks: A Modular Architecture",
                              "https://medium.com/@yasin162001/scaling-ai-agents-beyond-notebooks-a-modular-architecture-for-langgraph-in-production-4711764de464"),
    "blog_shaza_typing":     ("community", "Shaza Ali - Type Safety in LangGraph: TypedDict vs Pydantic",
                              "https://shazaali.substack.com/p/type-safety-in-langgraph-when-to"),
    "blog_pankaj_pyd":       ("community", "Pankaj Chandravanshi - LangGraph: State with Pydantic BaseModel",
                              "https://medium.com/fundamentals-of-artificial-intelligence/langgraph-state-with-pydantic-basemodel-023a2158ab00"),
    "blog_easton_2026":      ("community", "BetterLink Blog - LangGraph State Management in Practice (2026)",
                              "https://eastondev.com/blog/en/posts/ai/20260424-langgraph-agent-architecture/"),
    "blog_markaicode":       ("community", "Markaicode - Production Multi-Agent System with LangGraph",
                              "https://markaicode.com/langgraph-production-agent/"),
    "blog_clean_state":      ("community", "Vishal Lad - Clean State Architecture in LangGraph (Mar 2026)",
                              "https://medium.com/@ladvishal1985/everything-ive-learned-about-clean-state-architecture-in-langgraph-6d1352b0c00c"),
    "blog_unit_testing":     ("community", "Anirudh Sharma - Unit Testing LangGraph: Nodes and Flow Paths",
                              "https://medium.com/@anirudhsharmakr76/unit-testing-langgraph-testing-nodes-and-flow-paths-the-right-way-34c81b445cd6"),
    "blog_mock_llm":         ("community", "Matt Carvalho - How to Properly Mock LangChain LLM Execution",
                              "https://medium.com/@matgmc/how-to-properly-mock-langchain-llm-execution-in-unit-tests-python-76efe1b8707e"),
    "blog_tools_first":      ("community", "SitePoint - Implementing the Tools-First Pattern in LangGraph",
                              "https://www.sitepoint.com/implementing-the-tools-first-pattern-in-lang-graph/"),
    "deepwiki_subgraphs":    ("community", "Graph composition and nested graphs - DeepWiki (langchain-ai/langgraph)",
                              "https://deepwiki.com/langchain-ai/langgraph/3.6-graph-composition-and-nested-graphs"),
    "deepwiki_hitl":         ("community", "Human-in-the-loop and interrupts - DeepWiki (langchain-ai/langgraph)",
                              "https://deepwiki.com/langchain-ai/langgraph/3.7-human-in-the-loop-and-interrupts"),
    "blog_send_api":         ("community", "Sreeni - Leveraging LangGraph's Send API for Dynamic and Parallel Workflow Execution",
                              "https://dev.to/sreeni5018/leveraging-langgraphs-send-api-for-dynamic-and-parallel-workflow-execution-4pgd"),
    "blog_handoffs_tds":     ("community", "Towards Data Science - How Agent Handoffs Work in Multi-Agent Systems",
                              "https://towardsdatascience.com/how-agent-handoffs-work-in-multi-agent-systems/"),
    "blog_multiagent_2026":  ("community", "Mager - LangGraph: Build Stateful Multi-Agent Systems That Don't Crash (2026)",
                              "https://www.mager.co/blog/2026-03-12-langgraph-deep-dive/"),
    "blog_scaling_lg":       ("community", "AI Practitioner - Scaling LangGraph Agents: Parallelization, Subgraphs, and Map-Reduce Trade-Offs",
                              "https://aipractitioner.substack.com/p/scaling-langgraph-agents-parallelization"),
    "dev_mongo_ltm":         ("community", "DEV (MongoDB) - LangGraph With MongoDB: Building Conversational Long-Term Memory",
                              "https://dev.to/mongodb/langgraph-with-mongodb-building-conversational-long-term-memory-for-intelligent-ai-agents-2pcn"),
}


# ============================================================
# Helpers
# ============================================================

def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p


def add_do(doc, text):
    """A 'Do this' bullet - bold leading verb."""
    p = doc.add_paragraph(style="List Bullet")
    lead = p.add_run("Do: ")
    lead.bold = True
    lead.font.size = Pt(11)
    body = p.add_run(text)
    body.font.size = Pt(11)
    return p


def add_dont(doc, text):
    """A 'Do NOT do this' bullet - bold leading negation."""
    p = doc.add_paragraph(style="List Bullet")
    lead = p.add_run("Do NOT: ")
    lead.bold = True
    lead.font.size = Pt(11)
    body = p.add_run(text)
    body.font.size = Pt(11)
    return p


def add_citation(doc, key):
    tag, title, url = CITES[key]
    p = doc.add_paragraph(style="List Bullet 2")
    run = p.add_run(f"[{tag}] {title} - ")
    run.font.size = Pt(10)
    run.italic = True
    link = p.add_run(url)
    link.font.size = Pt(10)
    link.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p


# ============================================================
# Build
# ============================================================

def build():
    doc = Document()

    # Title block
    title = doc.add_heading("LangChain / LangGraph OO Best-Practices Guide - V4", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sr = sub.add_run("Final consolidated reference - synthesis of V1 (research), "
                     "V2 (architect annotations), and V3 (resolution). "
                     "Every recommendation is definitive.")
    sr.italic = True
    sr.font.size = Pt(11)
    meta = doc.add_paragraph()
    meta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    mr = meta.add_run("Compiled 2026-04-28  -  Skip Snow / ChatHealthy  -  "
                      "Sources tagged [official] = LangChain/LangGraph or MongoDB "
                      "co-maintained, [community] = third-party")
    mr.font.size = Pt(9)
    mr.italic = True

    # ============================================================
    # 1. Executive Summary
    # ============================================================
    add_heading(doc, "1. Executive Summary", level=1)

    add_para(doc,
        "The dominant production pattern in late-2025 / 2026 LangGraph is: "
        "(a) define state as a Pydantic BaseModel in its own state.py module - "
        "TypedDict is acceptable in a notebook prototype but is not the "
        "production target; (b) carry conversation as messages: "
        "Annotated[list[AnyMessage], add_messages], with tool calls and tool "
        "results living as ToolMessage entries on the same channel - never as "
        "flat top-level state fields per tool; (c) decompose the graph into "
        "subgraphs each owning a focused state schema (Single-Responsibility "
        "applied to graphs), composed under one parent runtime; (d) implement "
        "nodes as small pure functions or as callable classes with __call__ "
        "when they need encapsulated dependencies, returning partial state "
        "updates and never mutating; (e) inject runtime dependencies (DB "
        "clients, LLM clients, config) via Runtime[ContextSchema] / "
        "ToolRuntime, NOT module-level imports; (f) persist with a "
        "production-grade checkpointer that matches the existing data tier - "
        "for ChatHealthy that means MongoDBSaver from "
        "langgraph-checkpoint-mongodb (co-maintained by MongoDB and "
        "LangChain), not PostgresSaver; (g) lay the project out as state.py / "
        "nodes.py / tools.py / agent.py per the official "
        "langgraph-example-pyproject template, with subgraphs as the unit of "
        "decomposition.")

    add_heading(doc, "Definitive choices ChatHealthy is making", level=2)
    add_para(doc, "Five resolved decisions that supersede any earlier "
                  "exploration. Anything in this guide that conflicts with "
                  "one of these has been corrected; rationale lives in the "
                  "relevant per-question section.", italic=True)
    add_bullet(doc, "Persistence: MongoDBSaver from langgraph-checkpoint-mongodb. "
                    "NOT PostgresSaver. (Q2)")
    add_bullet(doc, "Graph decomposition: each capability epic (Find Care, "
                    "Talk About Care, Evaluate Care) is its own subgraph "
                    "composed under one parent runtime; epics communicate "
                    "via Command(graph=Command.PARENT). NOT separately "
                    "invoked top-level graphs; NOT one monolithic graph. (Q9)")
    add_bullet(doc, "State home: Pydantic BaseModel in its own state.py "
                    "file. NOT TypedDict in the orchestration script. (Q1)")
    add_bullet(doc, "Dependency injection: nodes and tools receive "
                    "dependencies via Runtime[ContextSchema] / ToolRuntime. "
                    "NOT module-level imports. (Q5)")
    add_bullet(doc, "Conversation modeling: messages: "
                    "Annotated[list[AnyMessage], add_messages]; tool calls "
                    "and tool results are ToolMessage entries on that "
                    "channel. NOT flat top-level fields per tool; NOT a "
                    "tool_outputs: dict. (Q4, Q10)")

    add_para(doc, "Key supporting sources for this summary:", italic=True)
    for k in ("docs_graph_api", "docs_app_structure", "docs_persistence",
              "docs_context", "blog_building", "blog_clean_state",
              "blog_easton_2026", "blog_markaicode",
              "pypi_mongo_ckpt", "mongo_atlas_lg",
              "docs_handoffs", "docs_interrupts",
              "deepwiki_subgraphs", "gh_deep_research_scratch",
              "gh_example_pyproject"):
        add_citation(doc, k)

    # ============================================================
    # 2. Per-question findings (definitive answers)
    # ============================================================
    add_heading(doc, "2. Per-Question Findings", level=1)

    # ---- Q1 ----
    add_heading(doc, "Q1. State schema home - TypedDict, Pydantic BaseModel, "
                     "dataclass, or JSON Schema?", level=2)
    add_do(doc, "Define state as a Pydantic BaseModel in a dedicated "
                "state.py module. Use this for production graphs and for "
                "any graph that crosses an HTTP boundary.")
    add_dont(doc, "Define a 15-field flat TypedDict inline at the top of "
                  "the graph build script. (This is the POC's current "
                  "shape and an explicit anti-pattern.)")
    add_dont(doc, "Use a bare dataclass unless Pydantic's recursive "
                  "validation cost has been measured and is unacceptable - "
                  "and document the measurement.")
    add_para(doc, "Rationale. The official graph-API doc states 'State in "
                  "LangGraph can be a TypedDict, Pydantic model, or "
                  "dataclass' and presents TypedDict as the introductory "
                  "example. Markaicode's production write-up is "
                  "unambiguous: 'Use a Pydantic BaseModel the moment you "
                  "step into production. It provides validation, "
                  "serialisation for checkpointing, and IDE "
                  "autocompletion.' Shaza Ali, Pankaj Chandravanshi, and "
                  "BetterLink 2026 converge on the same hybrid: TypedDict "
                  "may stay for purely-internal subgraph state; Pydantic "
                  "is mandatory at boundaries. ChatHealthy's epic "
                  "subgraphs all cross HTTP and Kafka boundaries, so "
                  "Pydantic is mandatory throughout. V1 left this as a "
                  "graded recommendation; V3 dialogue and the architect's "
                  "reuse goals close it to Pydantic-only.")
    for k in ("docs_graph_api", "blog_shaza_typing", "blog_pankaj_pyd",
              "blog_easton_2026", "blog_markaicode", "docs_app_structure",
              "gh_deep_research_scratch"):
        add_citation(doc, k)

    # ---- Q2 ----
    add_heading(doc, "Q2. State persistence - which checkpointer in production?", level=2)
    add_do(doc, "Use MongoDBSaver from the langgraph-checkpoint-mongodb "
                "package (co-maintained by MongoDB and LangChain), with "
                "the async API (aput / aget_tuple / alist / aput_writes / "
                "adelete_thread) for all production graphs.")
    add_do(doc, "Use InMemorySaver in unit and integration tests so each "
                "test gets a fresh checkpoint store.")
    add_dont(doc, "Use PostgresSaver. ChatHealthy already runs MongoDB; "
                  "introducing Postgres only to host the LangGraph "
                  "checkpointer would add a database for no concrete "
                  "requirement.")
    add_dont(doc, "Use SqliteSaver in production - the official doc "
                  "labels it 'ideal for experimentation and local "
                  "workflows' and BetterLink 2026 says 'skip SqliteSaver "
                  "entirely.'")
    add_dont(doc, "Build a custom MongoDB saver - the official "
                  "co-maintained one already exists and covers the entire "
                  "BaseCheckpointSaver contract.")
    add_para(doc, "Rationale. The official persistence doc names "
                  "PostgresSaver as the production target in the generic "
                  "case. V1's research recommended PostgresSaver on that "
                  "basis. V3 dialogue refined the answer to MongoDBSaver "
                  "for ChatHealthy specifically because MongoDB is already "
                  "in our stack and a co-maintained MongoDBSaver package "
                  "exists. PyPI lists langgraph-checkpoint-mongodb as "
                  "MongoDB- and LangChain-maintained (not community); the "
                  "MongoDB Atlas integration page documents MongoDBSaver "
                  "as a first-party feature that 'persists agent state in "
                  "MongoDB' and 'enables human-in-the-loop, memory, time "
                  "travel, and fault-tolerance for your LangGraph agents.' "
                  "The full BaseCheckpointSaver contract (sync + async) is "
                  "implemented and is a published Read-the-Docs class. "
                  "Subgraph composition (Q9) under one parent runtime "
                  "means a single MongoDBSaver instance covers every epic.")
    for k in ("docs_persistence", "blog_v02_ckpt", "blog_easton_2026",
              "blog_klarna", "pypi_mongo_ckpt", "mongo_atlas_lg",
              "mongo_atlas_lg_agents", "mongo_lc_readthedocs",
              "mongo_blog", "ref_checkpoints", "docs_custom_ckpt",
              "dev_mongo_ltm"):
        add_citation(doc, k)

    # ---- Q3 ----
    add_heading(doc, "Q3. Node organization - top-level functions, classes, or class-based callables?", level=2)
    add_do(doc, "Default to small pure functions returning partial state "
                "updates, declared in a per-subgraph nodes.py.")
    add_do(doc, "Promote a node to a class with __call__ when it needs "
                "constructor-injected dependencies (LLM client, DB "
                "handle, config) or when several closely-related nodes "
                "share state through instance attributes.")
    add_dont(doc, "Use the procedural top-level-function-per-step pattern "
                  "beyond an exploratory prototype. (This is the POC's "
                  "current shape.)")
    add_dont(doc, "Hide mutable state inside a class-node - issue #1950 "
                  "documents broken behaviour for class-nodes that try to "
                  "carry hidden private state. Class-nodes are for "
                  "dependency injection only.")
    add_dont(doc, "Mutate the incoming state dict. Always return a "
                  "partial-update dict.")
    add_para(doc, "Rationale. The official graph-API doc defines a node "
                  "as 'just a Python function that reads our graph's "
                  "state and makes updates to it' and warns: 'Nodes should "
                  "return updates to the state directly, instead of "
                  "mutating the state.' The 'Thinking in LangGraph' doc "
                  "reinforces fine-grained nodes because 'durable "
                  "execution creates checkpoints at node boundaries' - "
                  "every checkpoint is a chance to resume, so node "
                  "granularity directly affects how interrupts (Q9, deep "
                  "research) work.")
    for k in ("docs_graph_api", "docs_thinking", "gh_issue_1950",
              "docs_app_structure"):
        add_citation(doc, k)

    # ---- Q4 ----
    add_heading(doc, "Q4. Tool design - @tool, BaseTool, and how tool I/O lives in state", level=2)
    add_do(doc, "Use the @tool decorator for the common case (typed "
                "args; docstring becomes the tool description).")
    add_do(doc, "Drop to a BaseTool subclass when you need explicit "
                "args_schema, paired sync+async behaviour, or tool-level "
                "state.")
    add_do(doc, "Wire tools into the graph via the prebuilt ToolNode.")
    add_do(doc, "Model tool calls and tool results as ToolMessage "
                "entries on the messages channel.")
    add_dont(doc, "Model tool outputs as flat top-level state fields per "
                  "tool. (POC has six such fields: location, "
                  "specialty_query, specialties, homeopathic_expansion, "
                  "providers, trials.) This is an explicit V2/V3-resolved "
                  "anti-pattern.")
    add_dont(doc, "Model tool outputs as a tool_outputs: dict[str, Any] "
                  "catch-all. The messages channel IS the right home for "
                  "tool I/O; structured artefacts that the LLM does not "
                  "need to read get their own typed Pydantic field on the "
                  "subgraph state.")
    add_dont(doc, "Build tools after the prompt is finalised. SitePoint "
                  "names this the dominant cause of 'hallucinated "
                  "parameters and silent failures.'")
    add_para(doc, "Rationale. The official tools doc states 'The "
                  "simplest way to create a tool is with the @tool "
                  "decorator. By default, the function's docstring "
                  "becomes the tool's description.' It identifies "
                  "config and runtime as reserved parameter names, "
                  "implying tools must use ToolRuntime for dependency "
                  "injection (Q5). The 'Clean State Architecture' post "
                  "explicitly warns against overloading messages with "
                  "non-conversational artefacts AND against multiplying "
                  "top-level fields - the resolution in V3 is: "
                  "ToolMessage on the messages channel for the LLM-"
                  "visible exchange, named typed Pydantic fields for the "
                  "structured artefacts the LLM does not need to "
                  "re-read.")
    for k in ("docs_tools", "blog_tools_first", "blog_clean_state",
              "docs_graph_api"):
        add_citation(doc, k)

    # ---- Q5 ----
    add_heading(doc, "Q5. Dependency injection - how do clients/config reach nodes and tools?", level=2)
    add_do(doc, "Pass dependencies via the LangGraph runtime context: "
                "Runtime[ContextSchema] for nodes, "
                "ToolRuntime[ContextSchema] for tools. Supply the "
                "dependencies at graph.invoke(...) time via the "
                "context= argument.")
    add_do(doc, "For class-based nodes, accept the dependency in "
                "__init__ and bind the constructed instance into the "
                "graph.")
    add_dont(doc, "Import database clients or LLM clients at module "
                  "scope inside node files. This is the most common "
                  "anti-pattern in early LangGraph code and breaks "
                  "testability.")
    add_dont(doc, "Pass dependencies through state. State is data that "
                  "checkpoints; clients are not data, they are runtime "
                  "context.")
    add_para(doc, "Rationale. The official Context overview is explicit: "
                  "'Runtime context is a form of dependency injection... "
                  "optimize the LLM context by providing dependencies "
                  "(like database connections, user IDs, or API clients) "
                  "to your tools and nodes at runtime rather than "
                  "hardcoding them.' The canonical pattern is "
                  "def node(state, runtime: Runtime[ContextSchema]) "
                  "with graph.invoke({...}, context={'user_name': "
                  "'John Smith'}). Swarnendu De: 'Move provider choices, "
                  "model names, and feature flags to config, and inject "
                  "them with configurable inputs at runtime.' The "
                  "fastapi-langgraph production template demonstrates "
                  "the pattern at scale by isolating clients in a "
                  "services/ layer that nodes consume.")
    for k in ("docs_context", "blog_swarnendu", "gh_fastapi_template",
              "docs_tools"):
        add_citation(doc, k)

    # ---- Q6 ----
    add_heading(doc, "Q6. Module / file structure", level=2)
    add_do(doc, "Adopt the official 4-file layout per subgraph: "
                "state.py (Pydantic schemas), nodes.py (node functions / "
                "class-callables), tools.py (tool definitions), agent.py "
                "(graph construction). Include a langgraph.json manifest "
                "at the project root per the langgraph-example-pyproject "
                "template.")
    add_do(doc, "Decompose by capability. One folder per epic (find_care/, "
                "talk_about_care/, evaluate_care/, deep_research/), each "
                "containing the 4-file layout. Shared utilities live in a "
                "base/ folder that all epics import.")
    add_do(doc, "Make subgraphs the unit of decomposition. Each "
                "capability epic is a compiled subgraph (Q9).")
    add_dont(doc, "Put 14 nodes and the graph builder in one 336-line "
                  "file. (POC user_journey.py.) This is anti-pattern AP-2 "
                  "and is the single largest predictor of state-"
                  "management bugs in production according to the 2026 "
                  "Mager piece ('over 60% of agent production incidents "
                  "relate to state management').")
    add_para(doc, "Rationale. The official Application Structure doc "
                  "presents state.py / nodes.py / tools.py / agent.py "
                  "inside my_agent/utils/ + my_agent/agent.py and states "
                  "'all project code lies within' the package directory. "
                  "The langgraph-example-pyproject template ships exactly "
                  "this layout. Yassin Hashem's 'Scaling AI Agents Beyond "
                  "Notebooks' extends it for multi-agent production: "
                  "per-agent models/, node/, workflow/ folders + a base/ "
                  "for shared code, justified as 'Each component lives in "
                  "its own place,' enabling reuse across workflows. The "
                  "fastapi-langgraph production template (2.2k stars) "
                  "uses an even fuller layout (api/ + core/langgraph/ + "
                  "services/ + schemas/ + models/) showing how this "
                  "scales into a real service.")
    for k in ("docs_app_structure", "gh_example_pyproject",
              "blog_yasin_modular", "gh_fastapi_template",
              "blog_multiagent_2026"):
        add_citation(doc, k)

    # ---- Q7 ----
    add_heading(doc, "Q7. State immutability and reducers", level=2)
    add_do(doc, "Return partial-update dicts from every node. Never "
                "mutate incoming state.")
    add_do(doc, "Annotate every list / dict field with an explicit "
                "reducer at the moment you declare the field. Use "
                "add_messages for the messages channel, operator.add for "
                "simple list accumulation, and a custom reducer when you "
                "need de-duplication or capping.")
    add_do(doc, "Cap unbounded growth via a custom reducer when a list "
                "channel is written by many nodes (e.g. messages, "
                "research notes). Otherwise the checkpoint grows "
                "unboundedly.")
    add_dont(doc, "Leave list fields un-annotated. LangGraph's default "
                  "is last-write-wins; in any fan-out topology (e.g. "
                  "the POC's parallel safety + classifier path) this "
                  "silently loses data.")
    add_dont(doc, "Mutate state in place inside a node body, even for "
                  "convenience. The Pregel/BSP merge model assumes "
                  "immutable returns.")
    add_para(doc, "Rationale. The official graph-API doc: 'Nodes should "
                  "return updates to the state directly, instead of "
                  "mutating the state' and 'If no reducer function is "
                  "explicitly specified then it is assumed that all "
                  "updates to the key should override it.' The 'Building "
                  "LangGraph' design post explains the underlying "
                  "Pregel/BSP model: channels are versioned and updates "
                  "are merged deterministically, which depends on "
                  "immutable returns. The 'Clean State Architecture' "
                  "post names 'missing reducers on list fields' and "
                  "'unbounded state growth' as explicit anti-patterns: "
                  "'Always give it a reducer, even if only one node "
                  "writes to it now.'")
    for k in ("docs_graph_api", "blog_building", "blog_clean_state"):
        add_citation(doc, k)

    # ---- Q8 ----
    add_heading(doc, "Q8. Testing - node isolation, graph integration, LLM mocking", level=2)
    add_do(doc, "Run a two-layer split. Layer 1: unit-test each node by "
                "calling the function (or class instance) directly with "
                "a synthetic state dict and asserting on the returned "
                "partial update - no graph compilation needed. Layer 2: "
                "integration-test the compiled graph with InMemorySaver, "
                "mocked LLM (FakeListLLM or a custom fake), and stubbed "
                "tools; assert on final state and the chosen edges.")
    add_do(doc, "Use graph.nodes['name'] to access individual nodes "
                "inside a compiled graph for targeted tests.")
    add_dont(doc, "Test the production graph against the production "
                  "checkpointer (MongoDBSaver). Tests use InMemorySaver. "
                  "MongoDBSaver is verified by integration env, not by "
                  "unit tests.")
    add_dont(doc, "Mock at .invoke() on the LLM directly - Matt "
                  "Carvalho's post documents why this is brittle. Use "
                  "FakeListLLM or pre-built LangChain fakes.")
    add_para(doc, "Rationale. The official Test doc: 'Compiled LangGraph "
                  "agents expose references to each individual node as "
                  "graph.nodes. You can take advantage of this to test "
                  "individual nodes' and recommends per-test compilation "
                  "with a fresh checkpointer. Anirudh Sharma: 'For node-"
                  "level unit tests, you don't need pytest - you can "
                  "call the node function directly, which is fast and "
                  "deterministic.' Matt Carvalho documents the "
                  "FakeListLLM technique and the reasons mocking "
                  ".invoke() directly is brittle.")
    for k in ("docs_test", "blog_unit_testing", "blog_mock_llm"):
        add_citation(doc, k)

    # ---- Q9 ----
    add_heading(doc, "Q9. Graph decomposition and OO principles - epic "
                     "subgraphs, handoffs, and long-running async work", level=2)
    add_para(doc, "This question carries the resolution of the "
                  "architect's V2 annotation on graph granularity and the "
                  "deep-research use case. Both are settled.", italic=True)

    add_heading(doc, "Graph decomposition", level=3)
    add_do(doc, "Build one subgraph per capability epic (Find Care, Talk "
                "About Care, Evaluate Care, Deep Research). Compose them "
                "as nodes in a single thin parent graph compiled once.")
    add_do(doc, "Define each subgraph's wire interface as input_schema "
                "and output_schema (Pydantic). Keys not in the schema do "
                "not cross the boundary.")
    add_do(doc, "Hand off between epics with "
                "Command(graph=Command.PARENT). The parent's routing "
                "table picks the next epic; the source epic does not "
                "need to know what comes next.")
    add_do(doc, "Implement shared per-utterance pre-processing "
                "(emergency_detector, ip_lock_check, "
                "repetitive_detector, classifier) as a single small "
                "subgraph that every epic invokes, not duplicated "
                "logic per epic.")
    add_dont(doc, "Build one monolithic graph that contains every node "
                  "for every epic. (POC AP-2.)")
    add_dont(doc, "Build separately invoked top-level graphs per epic. "
                  "Use subgraph composition under one parent runtime "
                  "instead. The exception: an epic that must run on a "
                  "different process / deployment / SLA can be a "
                  "separate top-level graph - document why.")
    add_dont(doc, "Spin up a graph for a deterministic pipeline that "
                  "has no LLM decisions. That should be a tool. The "
                  "'Building LangGraph' post is explicit: the framework "
                  "chose 'little to no abstraction... we focused on "
                  "control and durability.' Graphs are for control flow "
                  "across LLM calls.")

    add_heading(doc, "Long-running async work (deep research during a meeting)", level=3)
    add_do(doc, "Implement deep_research as a reusable subgraph (its "
                "own state schema, its own input_schema = "
                "{ResearchQuestion}, output_schema = {ResearchReport}). "
                "Talk About Care, Find Care, and Evaluate Care all "
                "import the same subgraph as a node.")
    add_do(doc, "Inside deep_research, fan out to per-source workers "
                "(clinicaltrials.gov, PubMed, preprint servers) using "
                "the Send API. The supervisor emits one Send per "
                "research target; workers run concurrently under "
                "asyncio.gather.")
    add_do(doc, "In the meeting subgraph, after dispatching the "
                "research, place a wait_for_research node that calls "
                "interrupt() with a structured payload and yields. "
                "When research completes, resume with "
                "Command(resume=research_payload). The meeting "
                "transcript receives the consolidated ResearchReport "
                "as a system ToolMessage.")
    add_do(doc, "Use MongoDBSaver async API throughout. interrupt() "
                "requires a checkpointer; concurrent Send fan-out "
                "requires async aput / aput_writes / aget_tuple / alist "
                "- all implemented by MongoDBSaver per Q2.")
    add_dont(doc, "Block the meeting transcript while research runs "
                  "(Send fan-out fully inline with await). The "
                  "architect's requirement is that meeting participants "
                  "continue to talk while research runs - inline "
                  "blocking violates this.")
    add_dont(doc, "Treat open_deep_research as an anti-pattern. It is "
                  "the canonical reference (R-3) for this exact use "
                  "case; deep_research_from_scratch (R-4) is the "
                  "scaffold for the schema discipline.")

    add_heading(doc, "Routing decisions as Pydantic schemas", level=3)
    add_do(doc, "Adopt the deep_research_from_scratch pattern: "
                "ClarifyWithUser, ResearchQuestion, ConductResearch, "
                "ResearchComplete are Pydantic structured outputs that "
                "the LLM emits and the graph routes on. Each is "
                "testable in isolation.")
    add_dont(doc, "Encode routing decisions as free-form strings or "
                  "untyped dicts that the next node has to re-parse. "
                  "That is precisely the brittleness Pydantic schemas "
                  "exist to remove.")

    add_heading(doc, "OO principles applied", level=3)
    add_do(doc, "Apply Single-Responsibility at the subgraph level - "
                "each subgraph owns a focused slice of state.")
    add_do(doc, "Apply Encapsulation via class-based nodes whose "
                "dependencies are constructor-injected.")
    add_do(doc, "Prefer Composition over Inheritance - assemble graphs "
                "and subgraphs rather than subclassing node base "
                "classes. The 'Building LangGraph' post: the framework "
                "chose 'little to no abstraction at all... we focused "
                "on control and durability.' Your OO discipline comes "
                "from your code, not from a framework base class.")
    add_do(doc, "Use Dependency Injection via the runtime context "
                "(Q5).")
    add_dont(doc, "Subclass a custom NodeBase. There is no canonical "
                  "framework class to subclass - just classes with "
                  "__call__.")

    add_para(doc, "Rationale. The official handoff doc presents two "
                  "patterns: 'single agent with middleware' and "
                  "'multiple agent subgraphs,' guiding 'Use single agent "
                  "with middleware for most handoff use cases - it's "
                  "simpler. Only use multiple agent subgraphs when you "
                  "need bespoke agent implementations.' ChatHealthy's "
                  "epics ARE bespoke agent implementations, so the "
                  "multi-subgraph pattern is the appropriate one. "
                  "DeepWiki on graph composition: a CompiledStateGraph "
                  "'implements the PregelProtocol and can be passed to "
                  "StateGraph.add_node()' - i.e., a compiled subgraph "
                  "IS a node. State passes by 'selective projection': "
                  "downward 'the parent passes only state keys that "
                  "exist in the child's input_schema' and upward 'the "
                  "subgraph returns values from its output_schema.' "
                  "Command(graph=Command.PARENT) lets a subgraph "
                  "direct the parent to the next epic without the "
                  "subgraph knowing the parent's routing table. For "
                  "long-running work: the official Interrupts doc "
                  "documents interrupt() / Command(resume=value) as "
                  "the canonical mechanism, requiring 'a checkpointer "
                  "to persist the graph state.' V1 grouped this under "
                  "'OO principles' generically; V2/V3 dialogue refined "
                  "it to a concrete decomposition recipe.")
    for k in ("docs_handoffs", "deepwiki_subgraphs", "deepwiki_hitl",
              "docs_interrupts", "blog_send_api", "blog_scaling_lg",
              "blog_easton_2026", "blog_handoffs_tds",
              "blog_multiagent_2026", "blog_building",
              "gh_open_deep_research", "gh_deep_research_scratch",
              "gh_deepagents", "gh_swarm", "gh_supervisor",
              "docs_thinking", "docs_add_memory"):
        add_citation(doc, k)

    # ---- Q10 ----
    add_heading(doc, "Q10. Dialogue / message / history modeling", level=2)
    add_do(doc, "Model conversation as messages: "
                "Annotated[list[AnyMessage], add_messages] - either by "
                "extending the prebuilt MessagesState or by declaring "
                "the field manually on the subgraph state.")
    add_do(doc, "Carry tool calls and tool results as ToolMessage "
                "entries on the same messages channel. The LLM sees the "
                "tool exchange in its context exactly as it happened.")
    add_do(doc, "Keep non-conversational artefacts (retrieved docs, "
                "structured tool payloads, research reports) in "
                "separate typed Pydantic fields on the subgraph state.")
    add_dont(doc, "Overload the messages list with retrieved documents, "
                  "intermediate reasoning chains, or raw tool payloads. "
                  "The 'Clean State Architecture' post names this an "
                  "anti-pattern.")
    add_dont(doc, "Model tool outputs as flat top-level fields per "
                  "tool, or as a tool_outputs: dict[str, Any]. (See "
                  "Q4.)")
    add_para(doc, "Rationale. The official graph-API doc introduces the "
                  "pattern: 'LangGraph includes a built-in reducer "
                  "add_messages' and the prebuilt MessagesState 'so "
                  "that we can have' the simpler pattern. The 'Clean "
                  "State Architecture' post names 'Overloading Messages "
                  "Field' as an anti-pattern: 'The messages list is "
                  "easy to append to, and before long it can end up "
                  "holding retrieved documents, intermediate reasoning, "
                  "tool outputs, and more' - prescribing dedicated "
                  "fields like retrieved_docs instead. V3 dialogue "
                  "closed the residual ambiguity from V1 by being "
                  "explicit that ToolMessage on the messages channel "
                  "IS the home for tool I/O.")
    for k in ("docs_graph_api", "blog_clean_state", "docs_tools"):
        add_citation(doc, k)

    # ============================================================
    # 3. Anti-patterns (definitive, expanded with V2/V3 resolutions)
    # ============================================================
    add_heading(doc, "3. Anti-Patterns to Avoid", level=1)

    add_para(doc, "Each anti-pattern below is named in a cited source. "
                  "Items marked POC are concretely present in "
                  "LangGraph/poc/user_journey.py. Items marked "
                  "(V2/V3-resolved) came out of the architect dialogue "
                  "and supersede any earlier softer guidance.",
                  italic=True)

    # AP-1
    add_heading(doc, "AP-1. The God State Object  (POC)", level=3)
    add_para(doc, "Quote from source: 'It's tempting to create one big "
                  "state class that every node in your graph can read "
                  "and write... State that belongs to everyone belongs "
                  "to no one.' POC reality: JourneyState has 15 "
                  "unrelated fields (input, session flags, safety "
                  "flags, classifier output, six tool-output fields, "
                  "dispatch metadata) all flat in one TypedDict. "
                  "Resolution: Pydantic state per subgraph, with each "
                  "subgraph owning a focused slice (Q1, Q9).")
    add_citation(doc, "blog_clean_state")

    # AP-2
    add_heading(doc, "AP-2. Monolithic Graph (no subgraph decomposition)  (POC)", level=3)
    add_para(doc, "Quote from source: 'I had a single LangGraph graph "
                  "with 22 nodes and one huge state object... essentially "
                  "building a monolith.' POC reality: 14 nodes and the "
                  "build_user_journey() function all live in one "
                  "336-line file. Resolution: one subgraph per epic, "
                  "composed under one parent runtime via "
                  "Command(graph=Command.PARENT) (Q9).")
    add_citation(doc, "blog_clean_state")

    # AP-3
    add_heading(doc, "AP-3. Tool outputs as flat top-level state fields  (POC, V2/V3-resolved)", level=3)
    add_para(doc, "Source guidance: tool I/O rides the messages channel "
                  "as ToolMessage entries; structured artefacts go into "
                  "named typed Pydantic fields on the subgraph state. "
                  "POC reality: location, specialty_query, specialties, "
                  "homeopathic_expansion, providers, trials are all "
                  "flat top-level keys. Resolution (V2/V3): "
                  "ToolMessage on messages for the conversation; "
                  "Pydantic fields on subgraph state for the structured "
                  "artefacts the LLM does not need to re-read; NOT a "
                  "tool_outputs: dict[str, Any] catch-all.")
    add_citation(doc, "docs_graph_api")
    add_citation(doc, "blog_clean_state")

    # AP-4
    add_heading(doc, "AP-4. Procedural top-level node functions with no encapsulation  (POC)", level=3)
    add_para(doc, "Source guidance: nodes are pure functions OR "
                  "classes-with-__call__ when they have dependencies. "
                  "The official Application Structure puts node code "
                  "under a nodes.py module; the modular-architecture "
                  "Medium piece pushes further to per-capability "
                  "folders. POC reality: every node is a top-level def "
                  "in user_journey.py with no module boundary, no "
                  "class wrappers, and stub bodies. Resolution: Q3 + "
                  "Q6 file layout.")
    add_citation(doc, "docs_app_structure")
    add_citation(doc, "blog_yasin_modular")

    # AP-5
    add_heading(doc, "AP-5. Module-level imports of clients instead of dependency injection", level=3)
    add_para(doc, "Source guidance: 'Runtime context is a form of "
                  "dependency injection... optimize the LLM context by "
                  "providing dependencies (like database connections, "
                  "user IDs, or API clients) to your tools and nodes "
                  "at runtime rather than hardcoding them.' Anti-"
                  "pattern: a top-of-file import of a singleton DB "
                  "client referenced inside node bodies. Resolution: "
                  "Q5 - Runtime[ContextSchema] / ToolRuntime.")
    add_citation(doc, "docs_context")

    # AP-6
    add_heading(doc, "AP-6. Missing reducers on list fields  (POC risk)", level=3)
    add_para(doc, "Quote from source: 'If you don't set a reducer, "
                  "LangGraph uses last-write-wins by default... Always "
                  "give it a reducer, even if only one node writes to "
                  "it now.' POC reality: history, specialties, "
                  "providers, trials, and homeopathic_expansion are "
                  "all bare list with no annotation; the parallel "
                  "safety + classifier fan-out is exactly the topology "
                  "where last-write-wins silently loses data. "
                  "Resolution: Q7.")
    add_citation(doc, "blog_clean_state")
    add_citation(doc, "docs_graph_api")

    # AP-7
    add_heading(doc, "AP-7. Overloaded messages field", level=3)
    add_para(doc, "Quote: 'The messages list is easy to append to, and "
                  "before long it can end up holding retrieved "
                  "documents, intermediate reasoning, tool outputs, and "
                  "more.' Fix: keep messages for the LLM-visible "
                  "conversation (including ToolMessage); store retrieval "
                  "/ structured tool artefacts in named typed Pydantic "
                  "fields on the subgraph state. Resolution: Q4 + Q10.")
    add_citation(doc, "blog_clean_state")

    # AP-8 - rewritten for V4
    add_heading(doc, "AP-8. Dev-only checkpointer in production  (V2/V3-resolved)", level=3)
    add_para(doc, "Source quote (official): 'InMemorySaver: suitable "
                  "for development and testing... SqliteSaver: ideal "
                  "for experimentation and local workflows... "
                  "PostgresSaver: ideal for using in production.' "
                  "Resolution (V2/V3): the actual anti-pattern is "
                  "running on a dev-only checkpointer in production. "
                  "PostgresSaver is the LangChain reference for the "
                  "generic case. For ChatHealthy specifically, "
                  "MongoDBSaver from langgraph-checkpoint-mongodb (co-"
                  "maintained by MongoDB and LangChain) is the "
                  "production target because Mongo is already in our "
                  "stack. V1's research recommended PostgresSaver; "
                  "refined in V3 dialogue to MongoDBSaver because Mongo "
                  "is already in ChatHealthy's stack and a co-maintained "
                  "MongoDBSaver package exists. See AP-9 for the "
                  "do-not-introduce-Postgres rule.")
    add_citation(doc, "docs_persistence")
    add_citation(doc, "blog_easton_2026")
    add_citation(doc, "pypi_mongo_ckpt")
    add_citation(doc, "mongo_atlas_lg")

    # AP-9 (new in V4)
    add_heading(doc, "AP-9. Introducing Postgres only to host the LangGraph checkpointer  (V2/V3-resolved)", level=3)
    add_para(doc, "Resolution (V2/V3): do NOT use PostgresSaver when "
                  "MongoDB is the existing infrastructure. The "
                  "officially-supported langgraph-checkpoint-mongodb "
                  "package implements the full BaseCheckpointSaver "
                  "contract (sync + async). Adding Postgres for the "
                  "checkpointer alone introduces a new database tier "
                  "without a concrete requirement and violates the "
                  "engineering principle of don't-introduce-"
                  "infrastructure-without-a-driver. Reserve a possible "
                  "Postgres introduction for an explicit future need "
                  "(e.g. multi-row ACID join with another Postgres-only "
                  "system) and document it as a separate decision.")
    add_citation(doc, "pypi_mongo_ckpt")
    add_citation(doc, "mongo_atlas_lg")
    add_citation(doc, "mongo_lc_readthedocs")
    add_citation(doc, "ref_checkpoints")

    # AP-10 (new in V4)
    add_heading(doc, "AP-10. tool_outputs: dict[str, Any] catch-all on state  (V2/V3-resolved)", level=3)
    add_para(doc, "Resolution (V2/V3): do NOT model tool outputs as a "
                  "single tool_outputs dict on top-level state - and "
                  "do NOT model them as one flat field per tool either. "
                  "ToolMessage entries on the messages channel ARE the "
                  "home for tool calls and tool results; structured "
                  "artefacts the LLM does not need to re-read get their "
                  "own typed Pydantic field on the subgraph state. "
                  "This anti-pattern was named through the V2/V3 "
                  "dialogue clarifying Q4 / Q10.")
    add_citation(doc, "blog_clean_state")
    add_citation(doc, "docs_graph_api")
    add_citation(doc, "docs_tools")

    # AP-11 (new in V4)
    add_heading(doc, "AP-11. Epics as separately invoked top-level graphs  (V2/V3-resolved)", level=3)
    add_para(doc, "Resolution (V2/V3): do NOT invoke each capability "
                  "epic as its own top-level graph. Compose the epics "
                  "as compiled subgraphs under one parent runtime. "
                  "Separately invoked top-level graphs are appropriate "
                  "ONLY when an epic must run on a different process / "
                  "deployment / SLA - and that justification must be "
                  "documented. Subgraph composition gives you: a "
                  "single BaseCheckpointSaver instance per session, "
                  "automatic interrupt / resume across epics, "
                  "Command(graph=Command.PARENT) for handoffs, and "
                  "selective state projection via input_schema / "
                  "output_schema.")
    add_citation(doc, "docs_handoffs")
    add_citation(doc, "deepwiki_subgraphs")
    add_citation(doc, "blog_easton_2026")

    # AP-12 (new in V4)
    add_heading(doc, "AP-12. Building a custom MongoDB saver  (V2/V3-resolved)", level=3)
    add_para(doc, "Resolution (V2/V3): do NOT build a parallel MongoDB "
                  "saver. The officially-supported "
                  "langgraph-checkpoint-mongodb package already exists "
                  "and is co-maintained by MongoDB and LangChain. "
                  "Building a parallel implementation duplicates work, "
                  "imposes a tracking burden against LangGraph's "
                  "BaseCheckpointSaver interface drift, and violates "
                  "don't-reinvent-the-wheel.")
    add_citation(doc, "pypi_mongo_ckpt")
    add_citation(doc, "ref_checkpoints")
    add_citation(doc, "docs_custom_ckpt")

    # ============================================================
    # 4. Reference implementations
    # ============================================================
    add_heading(doc, "4. Reference Implementations", level=1)

    add_para(doc, "Six real GitHub repositories that exemplify the "
                  "patterns in this guide. Star counts and dates were "
                  "captured during the research pass on 2026-04-28. "
                  "Grouped as 'primary' (study first) vs 'supporting' "
                  "(reference when needed).")

    add_heading(doc, "Primary references (study first, in this order)", level=2)

    # R1
    add_heading(doc, "R-1. langchain-ai/langgraph-example-pyproject  (PRIMARY)", level=3)
    add_para(doc, "52 stars but officially curated. The canonical "
                  "layout: my_agent/utils/{state.py, nodes.py, tools.py} "
                  "+ my_agent/agent.py + langgraph.json. This is the "
                  "skeleton the ChatHealthy refactor should mirror "
                  "(Q6). Read first because it is the smallest "
                  "complete example of the production layout.")
    add_citation(doc, "gh_example_pyproject")
    add_citation(doc, "docs_app_structure")

    # R2
    add_heading(doc, "R-2. langchain-ai/deep_research_from_scratch  (PRIMARY)", level=3)
    add_para(doc, "708 stars. Tutorial repo (5 notebooks: scoping, "
                  "research_agent, research_agent_mcp, "
                  "research_supervisor, full_agent). This is THE "
                  "reference for Pydantic-schemas-as-routing-decisions "
                  "(ClarifyWithUser, ResearchQuestion, ConductResearch, "
                  "ResearchComplete) and for the supervisor / worker "
                  "split inside the deep-research subgraph. Read second "
                  "because it is the schema-discipline reference for "
                  "Q9's deep-research subgraph.")
    add_citation(doc, "gh_deep_research_scratch")

    # R3
    add_heading(doc, "R-3. langchain-ai/open_deep_research  (PRIMARY)", level=3)
    add_para(doc, "11.3k stars. Production-grade research agent. src/ "
                  "layout, configurable models for distinct LLM roles "
                  "(summariser, researcher, compressor, writer), "
                  "legacy/ folder preserves an older workflow + "
                  "multi-agent supervisor implementation for reference. "
                  "Read third because it is the production-hardening "
                  "reference (retries, observability, model "
                  "configuration) for the deep-research subgraph "
                  "scaffolded from R-2.")
    add_citation(doc, "gh_open_deep_research")

    add_heading(doc, "Supporting references", level=2)

    # R4
    add_heading(doc, "R-4. langchain-ai/langgraph (the framework itself)", level=3)
    add_para(doc, "30.7k stars. Authoritative source for the graph "
                  "runtime, checkpointer interfaces, and example "
                  "notebooks under examples/. Read for: Pregel/BSP "
                  "model, channel implementation, BaseCheckpointSaver "
                  "interface.")
    add_citation(doc, "gh_langgraph")

    # R5
    add_heading(doc, "R-5. langchain-ai/deepagents", level=3)
    add_para(doc, "21.9k stars; latest release 2026-04-15. Agent "
                  "harness with planning, filesystem, sub-agent "
                  "delegation. LangGraph-native; compiled graphs are "
                  "returned for downstream composition. Read for: how "
                  "a complex agent stays composable.")
    add_citation(doc, "gh_deepagents")

    # R6
    add_heading(doc, "R-6. wassim249/fastapi-langgraph-agent-production-ready-template", level=3)
    add_para(doc, "2.2k stars (community). Realistic FastAPI service "
                  "skeleton: app/api/v1/ + app/core/langgraph/ + "
                  "app/services/ + app/schemas/ + app/models/. "
                  "Demonstrates JWT, rate limiting, Langfuse "
                  "observability, mem0+pgvector long-term memory. Read "
                  "for: how the 4-file template scales into a "
                  "deployable service. Note: ChatHealthy uses "
                  "MongoDBStore, not pgvector, for long-term memory "
                  "(Q2); this template's schema is otherwise "
                  "applicable.")
    add_citation(doc, "gh_fastapi_template")

    add_para(doc, "Index of the wider ecosystem (community-curated):")
    add_citation(doc, "gh_awesome")

    # ============================================================
    # 5. Citations Bibliography
    # ============================================================
    add_heading(doc, "5. Citations Bibliography", level=1)

    def section(label, keys):
        add_heading(doc, label, level=2)
        for k in keys:
            tag, title, url = CITES[k]
            p = doc.add_paragraph()
            r = p.add_run(f"[{tag}] {title}\n")
            r.font.size = Pt(10)
            r.bold = True
            r2 = p.add_run(url)
            r2.font.size = Pt(10)
            r2.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    section("Official LangChain / LangGraph documentation", [
        "docs_graph_api", "docs_persistence", "docs_app_structure",
        "docs_tools", "docs_context", "docs_test", "docs_thinking",
        "docs_handoffs", "docs_interrupts", "docs_add_memory",
        "docs_custom_ckpt", "ref_checkpoints",
    ])
    section("Official LangChain blog", [
        "blog_building", "blog_v1", "blog_klarna", "blog_in_prod",
        "blog_v02_ckpt",
    ])
    section("Official MongoDB / LangChain co-maintained sources", [
        "pypi_mongo_ckpt", "mongo_atlas_lg", "mongo_atlas_lg_agents",
        "mongo_lc_readthedocs", "mongo_blog",
    ])
    section("Official GitHub repositories", [
        "gh_langgraph", "gh_example_pyproject", "gh_open_deep_research",
        "gh_deep_research_scratch", "gh_deepagents", "gh_swarm",
        "gh_supervisor", "gh_issue_1950",
    ])
    section("Community blog posts and tutorials", [
        "blog_swarnendu", "blog_yasin_modular", "blog_shaza_typing",
        "blog_pankaj_pyd", "blog_easton_2026", "blog_markaicode",
        "blog_clean_state", "blog_unit_testing", "blog_mock_llm",
        "blog_tools_first", "deepwiki_subgraphs", "deepwiki_hitl",
        "blog_send_api", "blog_handoffs_tds", "blog_multiagent_2026",
        "blog_scaling_lg", "dev_mongo_ltm",
    ])
    section("Community repositories", [
        "gh_awesome", "gh_fastapi_template",
    ])

    add_heading(doc, "Recency note", level=2)
    add_para(doc, "All cited sources were retrieved on 2026-04-28. None "
                  "of the URLs cited here are older than 2024; the "
                  "official LangChain blog post 'LangGraph v0.2 "
                  "checkpointers' is from 2024 and is the oldest "
                  "source. Three community pieces (BetterLink Apr-2026, "
                  "Vishal Lad Mar-2026, Markaicode 2025) are "
                  "explicitly 2025-2026, as is the Mager 2026 "
                  "multi-agent piece used in Q9.")

    add_heading(doc, "Sources searched but no firm answer found", level=2)
    add_para(doc, "Klarna's published case study confirms LangGraph use "
                  "and checkpointing posture but does not disclose "
                  "state-schema choice or directory layout. The "
                  "LangGraph 1.0 announcement asserts stability but "
                  "does not pick TypedDict vs Pydantic. No claim in "
                  "this document is sourced from those gaps.")
    for k in ("blog_klarna", "blog_v1"):
        add_citation(doc, k)

    # Save
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
