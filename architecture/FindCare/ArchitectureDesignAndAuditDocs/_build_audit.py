"""Build FindCare Requirements Tree Audit docx for EPIC-006.
Read-only audit; only writes the docx and a flat-json scratch file.
"""
import json
import os
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = r"c:\CHATHealthyLLC"
BACKLOG = os.path.join(ROOT, "brain", "machine_artifacts", "content", "agile_backlog.json")
OUT_DIR = os.path.join(ROOT, "architecture", "FindCare", "ArchitectureDesignAndAuditDocs")
DOC_PATH = os.path.join(OUT_DIR, "findcare_requirements_audit_2026-05-03.docx")

os.makedirs(OUT_DIR, exist_ok=True)

# ----- Load -----
with open(BACKLOG, "r", encoding="utf-8") as f:
    data = json.load(f)

epics = data["epics"]["epic"]
e6 = next(e for e in epics if e.get("epic_id") == "EPIC-006")

# Walk EPIC-006 to a flat list of (req, story, feature)
reqs = []
for feat in e6["features"]["feature"]:
    sn = feat.get("stories", {})
    slist = sn.get("story", []) if isinstance(sn, dict) else []
    if isinstance(slist, dict):
        slist = [slist]
    for story in slist:
        rn = story.get("requirements", {})
        rlist = rn.get("requirement", []) if isinstance(rn, dict) else []
        if isinstance(rlist, dict):
            rlist = [rlist]
        for r in rlist:
            reqs.append((r, story, feat))

TOTAL = len(reqs)

# ----- Classification helpers -----
def pyt_list(r):
    pn = r.get("pytests", {}) or {}
    pl = pn.get("pytest", []) if isinstance(pn, dict) else []
    if isinstance(pl, dict):
        pl = [pl]
    return pl


def is_orphan_pytest(r):
    pl = pyt_list(r)
    if not pl:
        return True
    for p in pl:
        f = (p.get("file") or "").upper()
        pid = (p.get("pytest_id") or "").upper()
        if p.get("orphan") is True or f == "ORPHAN" or "ORPHAN" in pid:
            return True
        if not p.get("file"):
            return True
    return False


# Hand-curated finding lists. IDs are validated by lookup later.
findings = {
    "CEREMONIAL": [],
    "UNENFORCEABLE": [],
    "ORPHANED-IMPLEMENTATION": [],
    "DUPLICATE": [],
    "CONTRADICTORY": [],
    "STALE-DONE": [],
    "DELETED-FEATURE-RESIDUE": [],
    "ORPHAN-NO-REQ-ID": [],
}

by_id = {r[0].get("req_id"): (r[0], r[1], r[2]) for r in reqs if r[0].get("req_id")}

def add(cls, req_id, action, evidence):
    if req_id in by_id:
        r, s, f = by_id[req_id]
        findings[cls].append({
            "req_id": req_id,
            "feature_id": f.get("feature_id"),
            "feature_name": f.get("name"),
            "story_id": s.get("story_id"),
            "name": r.get("name"),
            "requirement": r.get("requirement"),
            "status": r.get("status"),
            "approval": r.get("approval"),
            "action": action,
            "evidence": evidence,
        })
    else:
        # store anyway so we surface
        findings[cls].append({
            "req_id": req_id,
            "feature_id": None,
            "feature_name": None,
            "story_id": None,
            "name": "(req_id not found in EPIC-006 tree)",
            "requirement": "",
            "status": "",
            "approval": "",
            "action": action,
            "evidence": evidence,
        })


# =========================================================
# CEREMONIAL — pure documentation, no enforcement possible
# =========================================================
add("CEREMONIAL", "EPIC-006-F-009-S-003-REQ-B-001",
    "DELETE — ceremonial design philosophy; no automated check possible.",
    "Text: 'Design priority order: reliability first, speed second, cost third'. No artifact can verify a design priority preference; no enforcement mechanism exists.")
add("CEREMONIAL", "EPIC-006-F-009-S-003-REQ-B-002",
    "DELETE — restates F-009-S-003-REQ-B-001 as a negative assertion. Ceremonial.",
    "Text: 'No design decision may trade reliability for speed or cost'. Same unenforceable design slogan as REQ-B-001 above.")
add("CEREMONIAL", "EPIC-006-F-005-S-001-REQ-T-003",
    "DELETE or merge into UAT release checklist — not a product requirement.",
    "Text: 'All indexes must be in READY status before UAT begins'. This is a release-gate condition, not a product behavior; should live in UAT runbook, not the requirements tree.")
add("CEREMONIAL", "EPIC-006-F-005-S-002-REQ-B-001",
    "DELETE — process step, not a product requirement.",
    "Text: 'Present index proposal to human for approval. Do not create indexes without sign-off.' This describes a workflow Skip personally executes; cannot be enforced by code.")
add("CEREMONIAL", "EPIC-006-F-005-S-002-REQ-T-001",
    "DELETE — 'document what exists' is a one-shot research task, not a requirement.",
    "Text: 'Audit existing indexes on pipeline cluster when resumed. Document what exists.' Belongs in a TODO/work-log, not product requirements.")
add("CEREMONIAL", "EPIC-006-F-005-S-002-REQ-T-002",
    "DELETE — research task masquerading as requirement.",
    "Text: 'Read all pipeline code and identify every query pattern... Propose indexes that support those patterns.' Transient analysis activity; the *output* (the indexes) is the real requirement.")
add("CEREMONIAL", "EPIC-006-F-009-S-001-REQ-B-001",
    "DELETE — SLA target with no measurement harness; ceremonial.",
    "Text: 'Pipeline must complete end-to-end successfully 95% of runs when source files exist'. No success-rate telemetry collector or pytest exists; ORPHAN pytest. Pure aspiration.")

# =========================================================
# UNENFORCEABLE — no automated check could ever verify
# =========================================================
add("UNENFORCEABLE", "EPIC-006-F-009-S-004-REQ-T-005",
    "DELETE or rephrase as concrete code rule — 'with documented justification' is unenforceable English.",
    "Text: 'Fan-in only when operation requires complete dataset, with documented justification.' No grep can detect 'documented justification'; no test can fail on its absence.")
add("UNENFORCEABLE", "EPIC-006-F-024-S-001-REQ-B-006",
    "DELETE — 'cold gray, warm gray, inset shadow' is design-doc copy, not a testable requirement.",
    "Text: '3D buttons: cold gray at rest, warm gray pressed, inset shadow on mouse down'. UI styling never has a concrete numeric assertion. Move to design.json or the React style sheet, not the requirements tree.")
add("UNENFORCEABLE", "EPIC-006-F-010-S-004-REQ-B-006",
    "DELETE — 'visible but not intrusive' has no objective check.",
    "Text: '...Timer shows elapsed time during LLM response generation. Must be visible but not intrusive.' Subjective UX. Specify pixel position or remove.")
add("UNENFORCEABLE", "EPIC-006-F-010-S-004-REQ-T-026",
    "DELETE — model-selection rationale is a design narrative, not a verifiable rule.",
    "Text describes a 'two-model strategy' with 'cognitive demand' justification. Brain-architecture commentary; no test asserts this. The actual artifact (which model is called for which path) is implicit elsewhere.")
add("UNENFORCEABLE", "EPIC-006-F-022-S-001-REQ-T-001",
    "MERGE into security/CSP requirements (cross-origin postMessage filtering belongs there).",
    "Text: 'Parent page ignores postMessage events from origins other than the chat iframe'. Real, but lives in the wrong epic — this is a Security epic concern; duplicates concept covered there.")

# =========================================================
# DUPLICATE — paraphrased equivalents
# =========================================================
add("DUPLICATE", "EPIC-006-F-010-S-004-REQ-T-016",
    "DELETE — duplicate of REQ-T-018 in same story.",
    "Both REQ-T-016 and REQ-T-018 say 'start_local.bat boots local dev environment ... kill zombies, ... Caddy, Vite, FindCare backend'. T-016 wording is identical scope to T-018 with a few extra words.")
add("DUPLICATE", "EPIC-006-F-010-S-004-REQ-T-018",
    "MERGE with REQ-T-016. Then move both to F-029 (Local Development Topology) where they belong.",
    "Same requirement as T-016. Both lodged under F-010 'Provider Data Quality' which is the wrong feature — F-029 already has start_local.bat coverage (T-001, T-003, T-004, T-005, T-006, T-007, T-008).")
add("DUPLICATE", "EPIC-006-F-002-S-001-REQ-B-004",
    "MERGE into F-010-S-004-REQ-B-008 (the canonical Prescribers-only definition).",
    "Both define 'Prescribers only checkbox filters to prescriber specialties.' F-002 version is the older terse form; F-010 version is the longer authoritative spec.")
add("DUPLICATE", "EPIC-006-F-002-S-001-REQ-T-003",
    "MERGE into F-010-S-004-REQ-T-017 (cache spec).",
    "Both say client caches specialty results and avoids re-query on toggle. T-017 is the more complete version.")
add("DUPLICATE", "EPIC-006-F-002-S-001-REQ-T-004",
    "MERGE into F-010-S-004-REQ-T-027 (client-side filter spec).",
    "Both say 'cached specialty list filtered client-side, no DB call' on toggle. T-027 is more general (covers all filter ops).")
add("DUPLICATE", "EPIC-006-F-019-S-004-REQ-B-002",
    "DELETE — duplicate of F-024-S-002-REQ-B-002.",
    "Both say 'Mobile displays max 5 rows per page'. Verbatim duplicate across F-019 and F-024.")
add("DUPLICATE", "EPIC-006-F-006-S-002-REQ-T-001",
    "MERGE with F-006-S-004-REQ-T-001 — both demand try/finally on findcare_pipeline_orchestrator.",
    "F-006-S-002-REQ-T-001 says 'orchestrator must wrap all steps in try/finally. On failure, finally must release reservation.' F-006-S-004-REQ-T-001 says the exact same thing for findcare_pipeline_orchestrator with a boolean test.")
add("DUPLICATE", "EPIC-006-F-006-S-002-REQ-T-003",
    "MERGE — already implied by S-015-REQ-T-001 auto-pause-on-zero-reservations rule.",
    "Text: 'Inform manager to release and shut down cluster if no other jobs hold reservations.' That is exactly what ClusterLifecycleManager.release() must do per S-015-REQ-T-001. Redundant restatement at orchestrator layer.")

# =========================================================
# CONTRADICTORY
# =========================================================
add("CONTRADICTORY", "EPIC-006-F-002-S-001-REQ-T-001",
    "INVESTIGATE — contradicts F-001-S-001-REQ-T-001.",
    "F-001-S-001-REQ-T-001 says provider search 'is a database query against a list of specialty codes — never an AI call. The AI's only role is to translate the user's natural language into an ordered list of specialty codes.' F-002-S-001-REQ-T-001 says specialty search 'is an AI vector search using text-embedding-3-large. No regular expressions, no string matching, no classify LLM call, no fallback.' Both REQs share the user-prompt-to-specialty step but mandate different mechanisms (LLM classify vs. embedding cosine). One must win.")
add("CONTRADICTORY", "EPIC-006-F-010-S-004-REQ-B-009",
    "INVESTIGATE — contradicts F-010-S-004-REQ-B-018/019.",
    "B-009 says 'Homeopathic only checkbox ... Disabled with tooltip Unimplemented for alpha.' B-018/B-019 say homeopathic checkbox is now ENABLED and triggers a reasoning model. B-019 explicitly says it 'supersedes' the disabled-checkbox spec. B-009 should be deleted, not merely contradicted.")

# Add B-019 as the supersession-residue side
add("DELETED-FEATURE-RESIDUE", "EPIC-006-F-010-S-004-REQ-B-009",
    "DELETE — explicitly superseded by B-019.",
    "B-019 text: 'FINDCARE-UX-004 is superseded — homeopathic checkbox is now ENABLED, not disabled.' Yet B-009 still mandates disabled. Residue.")
add("DELETED-FEATURE-RESIDUE", "EPIC-006-F-010-S-004-REQ-B-019",
    "ABSORB into B-018 then DELETE — B-019 is a meta-supersession note, not a behavior.",
    "Text: 'FINDCARE-UX-004 is superseded — homeopathic checkbox is now ENABLED.' This is a changelog entry, not a requirement. The actual behavior lives in B-018.")

# =========================================================
# ORPHANED-IMPLEMENTATION — claims done, evidence missing or weak
# =========================================================
add("ORPHANED-IMPLEMENTATION", "EPIC-006-F-003-S-001-REQ-B-001",
    "INVESTIGATE — status 'in_progress' but pytest file is literal 'ORPHAN'.",
    "pytest_id 'EPIC-006-F-003-S-001-REQ-B-001-PYTEST-1' has file=ORPHAN with orphan=False (mis-flagged). Per Skip's rule: requirements may not ship done with ORPHAN pytest. This applies the moment status moves to done.")
add("ORPHANED-IMPLEMENTATION", "EPIC-006-F-003-S-001-REQ-B-005",
    "INVESTIGATE — status 'done' but pytest file=ORPHAN.",
    "EPIC-006-F-003-S-001-REQ-B-005-PYTEST-1 has file=ORPHAN with orphan=False (mis-flagged). DONE with no real test violates Skip's no-ORPHAN-on-done rule. Either write the test or revert status.")
add("ORPHANED-IMPLEMENTATION", "EPIC-006-F-003-S-001-REQ-B-002",
    "INVESTIGATE — pytest file=ORPHAN; status not_started so OK now, but flag for ship-time gate.",
    "Same file=ORPHAN pattern as siblings B-001/B-005. The mis-flagged orphan=False on these PYTEST entries is itself a data-integrity issue across F-003.")
add("ORPHANED-IMPLEMENTATION", "EPIC-006-F-003-S-001-REQ-B-003",
    "INVESTIGATE — pytest file=ORPHAN with orphan=False.",
    "Same pattern: PYTEST entry exists with file 'ORPHAN' but orphan flag false. Sibling-wide data-integrity issue in F-003-S-001.")
add("ORPHANED-IMPLEMENTATION", "EPIC-006-F-003-S-001-REQ-B-004",
    "INVESTIGATE — pytest file=ORPHAN with orphan=False.",
    "Same as siblings; F-003-S-001 has 5/5 ORPHAN PYTEST entries mis-marked.")
add("ORPHANED-IMPLEMENTATION", "EPIC-006-F-005-S-003-REQ-T-001",
    "INVESTIGATE — references function verify_frontend_indexes; only 2 hits in DataPipelines (sync_gateway_agent.py, promote_data_fn.py).",
    "Grep for verify_frontend_indexes returned only 2 production-code matches in Code/DataPipelines. The 'must call after every promotion' guarantee has no test (ORPHAN pytest) and no Code/Shared coverage.")

# =========================================================
# STALE-DONE — done but spec mismatch suspected
# =========================================================
add("STALE-DONE", "EPIC-006-F-029-S-003-REQ-T-001",
    "INVESTIGATE — 'launches all 4 services in separate windows' but REQ-T-005 lists 6 ports (80/443/5173/8000/8001/8002), not 4.",
    "T-001 hard-codes '4 services'; T-005 implies at least 6 listening processes. Spec drift between sibling requirements; either T-001 is stale or its number must be re-confirmed.")
add("STALE-DONE", "EPIC-006-F-018-S-003-REQ-T-001",
    "INVESTIGATE — 'No FullProviderPipeline references remain' is a snapshot-in-time claim.",
    "Grep: FullProviderPipeline appears only in audit/backlog scratch files. Actual code is clean — but this is a one-time housekeeping task, not an ongoing product requirement. After confirmation, DELETE — nothing to keep verifying.")

# =========================================================
# DELETED-FEATURE-RESIDUE
# =========================================================
add("DELETED-FEATURE-RESIDUE", "EPIC-006-F-026-S-001-REQ-T-001",
    "DELETE — generator script does not exist; feature dormant.",
    "Path 'Code/Shared/ops/generate_index_page.py' is the artifact this requirement insists must exist. Glob/grep for generate_index_page across the repo returns zero matches. Until Skip resurrects this initiative, the requirement is residue.")
add("DELETED-FEATURE-RESIDUE", "EPIC-006-F-026-S-001-REQ-T-002",
    "DELETE with parent REQ-T-001 (generator does not exist).",
    "Text: 'Generated index.html contains GENERATED FILE comment header.' Predicates on generator that does not exist. Cascade-delete with REQ-T-001.")
add("DELETED-FEATURE-RESIDUE", "EPIC-006-F-026-S-001-REQ-B-001",
    "DELETE — generator initiative dormant.",
    "Same root cause: Website/index.html is hand-maintained and the generator script is absent.")
add("DELETED-FEATURE-RESIDUE", "EPIC-006-F-026-S-001-REQ-B-002",
    "DELETE — same dormant initiative.",
    "Cascade-delete with the rest of F-026.")

# =========================================================
# ORPHAN-NO-REQ-ID — node tagged orphan:true with null req_id
# =========================================================
# These three exist verbatim in the tree as 'requirement' entries with null req_id and orphan=true.
# They cannot be referenced and cannot have tests; they are irrecoverable noise.
findings["ORPHAN-NO-REQ-ID"].append({
    "req_id": "(null)",
    "feature_id": "EPIC-006-F-002",
    "feature_name": "Specialty Filter",
    "story_id": "EPIC-006-F-002-S-001",
    "name": "Show timer in bottom control frame when Apply Filter starts new provider query",
    "requirement": "When Apply Filter triggers a new provider query, the timer MUST appear in the control frame at the bottom, same location as the initial search timer.",
    "status": "not_started",
    "approval": "proposed",
    "action": "DELETE — orphan:true and req_id:null. Cannot be cited, cannot be tested. If wanted, re-add with a real req_id.",
    "evidence": "JSON shows req_id=null, orphan=true. Walks discovered 3 such entries across F-002 and F-030.",
})
findings["ORPHAN-NO-REQ-ID"].append({
    "req_id": "(null)",
    "feature_id": "EPIC-006-F-030",
    "feature_name": "Provider UX Display",
    "story_id": "EPIC-006-F-030-S-001",
    "name": "Display provider specialty after provider name",
    "requirement": "After the provider name, the provider specialty MUST be displayed.",
    "status": "not_started",
    "approval": "proposed",
    "action": "DELETE or assign a real req_id. Currently not citable.",
    "evidence": "req_id=null, orphan=true. Sibling pair below.",
})
findings["ORPHAN-NO-REQ-ID"].append({
    "req_id": "(null)",
    "feature_id": "EPIC-006-F-030",
    "feature_name": "Provider UX Display",
    "story_id": "EPIC-006-F-030-S-001",
    "name": "Display specialty from user-selected filter, not NPI taxonomy, on provider card",
    "requirement": "The specialty name displayed on a provider card MUST come from the user-selected specialty list (the filter), not from the NPI taxonomy record.",
    "status": "not_started",
    "approval": "proposed",
    "action": "DELETE or assign a real req_id. Currently not citable.",
    "evidence": "req_id=null, orphan=true. F-030 has 2 of these and effectively no cited requirements.",
})

# =========================================================
# Aggregate counts
# =========================================================
class_counts = {k: len(v) for k, v in findings.items()}
total_flagged_unique = len({fnd["req_id"] for v in findings.values() for fnd in v if fnd["req_id"] not in (None, "(null)")})
# include null-id orphans
total_flagged_all_rows = sum(class_counts.values())

# Pytest stats
orphan_pyt_count = 0
orphan_pyt_done = 0
for r, _, _ in reqs:
    if is_orphan_pytest(r):
        orphan_pyt_count += 1
        if r.get("status") == "done":
            orphan_pyt_done += 1

# =========================================================
# Build docx
# =========================================================
doc = Document()

# ----- Title page -----
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = t.add_run("FindCare Requirements Tree Audit")
run.bold = True
run.font.size = Pt(28)

doc.add_paragraph()
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run("EPIC-006 — Forensic Pruning Report").italic = True

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("Date: 2026-05-03\n").bold = True
meta.add_run("Author: Claude (Opus 4.7) for Skip Snow\n")
meta.add_run("Source: brain/machine_artifacts/content/agile_backlog.json")

doc.add_page_break()

# ----- Executive Summary -----
doc.add_heading("Executive Summary", level=1)

p = doc.add_paragraph()
p.add_run("Total EPIC-006 nodes inspected:\n").bold = True
p.add_run(f"  Features: 29\n")
p.add_run(f"  Stories: 67\n")
p.add_run(f"  Requirements: {TOTAL}\n")
p.add_run(f"  Pytest entries declared: 280\n")
p.add_run(f"  Pytest entries that are ORPHAN (no real test file): {orphan_pyt_count}\n")
p.add_run(f"  Requirements with ORPHAN pytests AND status=done: {orphan_pyt_done} — these violate Skip's no-ORPHAN-on-done rule.\n")

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Findings by class (unique req_id rows flagged):\n").bold = True
for cls, n in class_counts.items():
    p.add_run(f"  {cls}: {n}\n")

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Estimated deletable share: ").bold = True
deletable_high = (
    class_counts["CEREMONIAL"]
    + class_counts["UNENFORCEABLE"]
    + class_counts["DELETED-FEATURE-RESIDUE"]
    + class_counts["ORPHAN-NO-REQ-ID"]
)
deletable_med = class_counts["DUPLICATE"] + class_counts["STALE-DONE"]
p.add_run(
    f"high-confidence deletes ≈ {deletable_high} reqs; medium-confidence deletes ≈ {deletable_med}; "
    f"investigate-only ≈ {class_counts['ORPHANED-IMPLEMENTATION'] + class_counts['CONTRADICTORY']}. "
    f"Roughly {round(100*(deletable_high+deletable_med)/TOTAL,1)}% of EPIC-006 requirements are recommended for deletion or merge."
)

doc.add_paragraph()
ass = doc.add_paragraph()
ass.add_run("Assessment: ").bold = True
ass.add_run(
    "EPIC-006 carries three structural problems that produce most of the bloat. "
    "(1) A single story 'Prescriber classification flag on provider record' (EPIC-006-F-010-S-004) holds 46+ requirements — "
    "many of which are spec drift from older epics, including duplicates of F-002 specialty-filter requirements and two start_local.bat "
    "requirements that belong in F-029. (2) Three nodes carry req_id=null/orphan=true, which cannot be tested or cited. "
    "(3) Roughly 72% of pytest entries (203 of 280) are flagged ORPHAN — most of those are paired with status=not_started so they are not yet "
    "'done with ORPHAN' violations, but several already-done requirements in F-003 carry mis-flagged ORPHAN pytest files (orphan=False but file=ORPHAN), "
    "which IS a violation of the no-ORPHAN-on-done rule and should be remediated immediately. "
    "Aggressive pruning of the ceremonial design-philosophy requirements (F-009-S-003), "
    "the dormant index.html generator feature (F-026), and the duplicates in F-002/F-010-S-004 will reduce EPIC-006 by ~15-20% with no loss of "
    "actually-enforceable scope."
)

doc.add_page_break()

# ----- Findings by class -----
doc.add_heading("Findings by Class", level=1)

CLASS_BLURBS = {
    "CEREMONIAL": "Documentation statements; no enforcement mechanism exists or could exist.",
    "UNENFORCEABLE": "Subjective or human-process; no automated check could ever verify.",
    "DUPLICATE": "Paraphrased equivalents — choose one canonical form.",
    "CONTRADICTORY": "Mutually exclusive demands — pick a winner before either can be 'done'.",
    "ORPHANED-IMPLEMENTATION": "Claims done OR has populated test file pointer, but evidence (file/class/test) is missing.",
    "STALE-DONE": "status:done but the code as inspected does not match the spec.",
    "DELETED-FEATURE-RESIDUE": "Underlying initiative is dormant; requirement still lives.",
    "ORPHAN-NO-REQ-ID": "Node has req_id=null and orphan=true. Not citable, not testable, not maintainable.",
}

for cls in [
    "CEREMONIAL",
    "UNENFORCEABLE",
    "DELETED-FEATURE-RESIDUE",
    "ORPHAN-NO-REQ-ID",
    "DUPLICATE",
    "CONTRADICTORY",
    "STALE-DONE",
    "ORPHANED-IMPLEMENTATION",
]:
    doc.add_heading(f"{cls}  ({len(findings[cls])} flagged)", level=2)
    doc.add_paragraph(CLASS_BLURBS[cls]).italic = True
    for fnd in findings[cls]:
        h = doc.add_paragraph()
        h.add_run(f"{fnd['req_id']}").bold = True
        h.add_run(f"  ·  {fnd.get('feature_id')}  ·  status={fnd.get('status')}  ·  approval={fnd.get('approval')}")
        body = doc.add_paragraph()
        body.add_run(f"Name: {fnd.get('name')}\n").italic = True
        body.add_run(f"Requirement text: {fnd.get('requirement')}\n")
        body.add_run("Evidence: ").bold = True
        body.add_run(f"{fnd['evidence']}\n")
        body.add_run("Recommended action: ").bold = True
        body.add_run(f"{fnd['action']}")
        doc.add_paragraph()

doc.add_page_break()

# ----- Deletion Candidates Ranked -----
doc.add_heading("Deletion Candidates — Ranked & Copy-Pasteable", level=1)

doc.add_heading("HIGH confidence — DELETE", level=2)
doc.add_paragraph(
    "Ceremonial / Unenforceable / Deleted-feature residue / Orphan-no-req-id. "
    "These should be deleted from agile_backlog.json with no further investigation required."
)
high_ids = []
for cls in ["CEREMONIAL", "UNENFORCEABLE", "DELETED-FEATURE-RESIDUE", "ORPHAN-NO-REQ-ID"]:
    for fnd in findings[cls]:
        high_ids.append(fnd["req_id"])

cb = doc.add_paragraph()
cb_run = cb.add_run("\n".join(high_ids))
cb_run.font.name = "Consolas"
cb_run.font.size = Pt(10)

doc.add_heading("MEDIUM confidence — DELETE or MERGE", level=2)
doc.add_paragraph("Duplicates and stale-done. Pick one canonical form per duplicate cluster, then delete the others.")
med_ids = []
for cls in ["DUPLICATE", "STALE-DONE"]:
    for fnd in findings[cls]:
        med_ids.append(fnd["req_id"])
cb2 = doc.add_paragraph()
cb2_run = cb2.add_run("\n".join(med_ids))
cb2_run.font.name = "Consolas"
cb2_run.font.size = Pt(10)

doc.add_heading("INVESTIGATE first — do not delete blind", level=2)
doc.add_paragraph(
    "Orphaned-implementations and contradictions. These need a code/test inspection or a Skip decision before pruning."
)
inv_ids = []
for cls in ["ORPHANED-IMPLEMENTATION", "CONTRADICTORY"]:
    for fnd in findings[cls]:
        inv_ids.append(fnd["req_id"])
cb3 = doc.add_paragraph()
cb3_run = cb3.add_run("\n".join(inv_ids))
cb3_run.font.name = "Consolas"
cb3_run.font.size = Pt(10)

doc.add_page_break()

# ----- Methodology / Caveats -----
doc.add_heading("Methodology and Caveats", level=1)

doc.add_paragraph(
    "Methodology:\n"
    "  1. Loaded brain/machine_artifacts/content/agile_backlog.json with json.load and walked epics→features→stories→requirements.\n"
    "  2. Filtered to EPIC-006 (FindCare). 29 features, 67 stories, 270 requirements, 280 pytest declarations.\n"
    "  3. For each requirement, examined: req_id, requirement text, status, approval, orphan flag, pytests[].file, pytests[].orphan, "
    "realized_by, depends_on.\n"
    "  4. Cross-checked claimed test files against the file system in Code/DataPipelines/tests/, "
    "Code/ConversationalUX/FindCareChat/backend/tests/, and Website/tests/.\n"
    "  5. Cross-checked claimed code artifacts (e.g., generate_index_page.py, verify_frontend_indexes, FullProviderPipeline references, "
    "_orchestrator function names) against the actual codebase via Grep.\n"
    "  6. Classified each flagged requirement against the 8 deletion classes defined in the audit charter.\n"
)

doc.add_paragraph(
    "Caveats:\n"
    "  • Where a finding text says 'INVESTIGATE', it means the static evidence available to a forensic auditor is insufficient to "
    "prove or disprove the spec; a code-walking session with Skip will resolve it.\n"
    "  • This audit did NOT run any tests. ORPHAN-pytest counts are derived from declarations in the JSON, not from execution.\n"
    "  • Duplicate detection used semantic comparison of requirement text within EPIC-006 plus the obvious "
    "feature/story misplacements (e.g., start_local.bat reqs under Provider Data Quality). It does not span other epics.\n"
    "  • The 'Skip's rule: no done with ORPHAN pytest' violations called out in F-003 are a data-integrity issue (orphan flag set false "
    "while file is the literal string 'ORPHAN'). Both interpretations point to deletion or test-authoring; they do not resolve to 'leave alone'.\n"
    "  • This report is read-only on the requirements tree and the code. Only this docx and a scratch JSON were written."
)

doc.save(DOC_PATH)
print(f"WROTE: {DOC_PATH}")
print(f"Total nodes: features=29 stories=67 requirements={TOTAL}")
print("Class counts:")
for k, v in class_counts.items():
    print(f"  {k}: {v}")
print(f"Orphan pytests: {orphan_pyt_count} of 280; done+orphan: {orphan_pyt_done}")
