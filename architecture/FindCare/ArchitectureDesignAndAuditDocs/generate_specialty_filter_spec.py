"""Generate the approved Requirements Specification Word doc for
EPIC-006-F-002 (Specialty Filter).

Output: architecture/FindCare/ArchitectureDesignAndAuditDocs/
        EPIC-006-F-002_SpecialtyFilter_RequirementsSpec.docx
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor


HERE = Path(__file__).resolve().parent
OUT = HERE / "EPIC-006-F-002_SpecialtyFilter_RequirementsSpec.docx"


# ============================================================================
# FINAL approved tree (5 stories, 41 business + 11 technical = 52 REQs)
# ============================================================================
TREE = {
    "EPIC-006-F-002-S-001": {
        "name": "Display Specification",
        "kind": "cross-cutting visual",
        "description": (
            "How the SpecialtyFilter LOOKS — appearance, layout, frame "
            "placement (visual), colors, header structure, body, scroll, "
            "counts, button placement. Cross-cutting: applies uniformly "
            "across the user-journey stories. Behavior lives in "
            "S-002/S-003/S-004; visual rules live here."
        ),
        "reqs": [
            # PANEL-LEVEL
            ("REQ-B-001", "Reference screenshot is the visual ground truth",
             "The filter panel's rendered appearance MUST match the reference screenshot at {path/to/specialty_filter_mockup.png}. Implementation MUST be pixel-comparable on header structure, checkbox row spacing, font, and color."),
            ("REQ-B-002", "Panel positioned on the LEFT side of the screen on desktop",
             "On desktop displays, the SpecialtyFilter panel is positioned on the LEFT side of the screen."),
            ("REQ-B-003", "Panel green styling, occupies full vertical height of left-panel container",
             "The SpecialtyFilter panel uses green styling and occupies the full vertical height of the left-panel container."),
            ("REQ-B-004", "Responsive — must support Galaxy + iPhone 14+",
             "The SpecialtyFilter MUST be reactive and support the Galaxy and iPhone 14+ devices specified in EPIC-006-F-019."),
            ("REQ-B-005", "Panel occupies 58% vertical (header 18% + body 40%); 42% reserved for other widgets",
             "The SpecialtyFilter occupies 58% of the left-panel vertical space, divided into a header (18%) and a body (40%). The remaining 42% of the left-panel vertical space is reserved for other widgets (not part of the SpecialtyFilter feature display)."),

            # CELL 1 — HEADER
            ("REQ-B-006", "Header element 1 — 'Filter by specialty' label",
             "The first (left-most) element in the panel header is a text label reading 'Filter by specialty'."),
            ("REQ-B-007", "Header element 2 — 'All Possible' count display",
             "The second element in the panel header is a count display labeled 'All Possible'. The value is an integer."),
            ("REQ-B-008", "Header element 3 — 'Prescribers' count display",
             "The third element in the panel header is a count display labeled 'Prescribers'. The value is an integer."),
            ("REQ-B-009", "Header element 4 — 'Your Choices' count display",
             "The fourth element in the panel header is a count display labeled 'Your Choices'. The value is an integer."),
            ("REQ-B-010", "Header element 5 — Uncheck All / Check All toggle button",
             "The fifth element in the panel header is a single toggle button whose label flips between 'Uncheck All' and 'Check All' per the majority state of currently-visible checkboxes (label behavior in S-003-REQ-B-006)."),
            ("REQ-B-011", "Header element 6 — Prescribers checkbox (right-aligned)",
             "The sixth element in the panel header is a Prescribers checkbox, right-aligned within the header."),
            ("REQ-B-012", "Header element 7 — Homeopathic checkbox (right-aligned, below Prescribers)",
             "The seventh element in the panel header is a Homeopathic checkbox, right-aligned and positioned below the Prescribers checkbox."),
            ("REQ-B-013", "Header layout constraint — 7 elements fit, no truncation",
             "All seven header elements (REQ-B-006 through REQ-B-012) MUST fit inside the panel header with reactive percentage widths and no truncation. The Uncheck All / Check All toggle button MUST NOT appear anywhere other than the SpecialtyFilter header."),

            # CELL 2 — BODY
            ("REQ-B-014", "Body — each specialty rendered as a single horizontal checkbox row",
             "Below the header, each specialty in the displayed list is rendered as a single horizontal row containing the checkbox itself and the specialty Display Name."),
            ("REQ-B-015", "Body — max 12 visible specialty rows; vertical scroll engages beyond",
             "The body of the panel accommodates up to 12 visible specialty rows before vertical scrolling engages."),
            ("REQ-B-016", "Body — total list size communicated by 'All Possible' count in header",
             "The total size of the underlying specialty list is communicated to the user via the 'All Possible' count in the header. There is no separate 'showing N of M' affordance in the body."),

            # COUNT FORMATTING & REACTIVITY
            ("REQ-B-017", "Count format — each header count displays as 'Label: N'",
             "Each of the three header counts displays in the format 'Label: N' where N is an integer. Formatting is consistent across all three counts."),
            ("REQ-B-018", "Counts in header update reactively (no manual refresh)",
             "The 'All Possible', 'Prescribers', and 'Your Choices' counts update on screen reactively whenever the underlying counts change. No manual refresh is required."),

            # APPLY FILTER BUTTON PLACEMENT
            ("REQ-B-019", "Apply Filter button placed at the bottom of the screen",
             "The Apply Filter button is placed at the bottom of the screen, in the position it currently occupies in the implementation. Visibility logic for the button is governed by EPIC-006-F-002-S-004-REQ-B-005."),
        ],
    },

    "EPIC-006-F-002-S-002": {
        "name": "Present Initial Specialty Filter",
        "kind": "user journey",
        "description": (
            "The system produces and displays the relevant specialty list "
            "in response to the user's healthcare request. WHAT shows up "
            "— quality, scope, and inclusion — plus the technical AI flow "
            "that produces the list."
        ),
        "reqs": [
            ("REQ-B-001", "Display MUST comply with EPIC-006-F-002-S-001",
             "The visual presentation of the SpecialtyFilter panel in this story MUST comply with every REQ in EPIC-006-F-002-S-001 (Display Specification)."),
            ("REQ-B-002", "API surface and contracts MUST comply with EPIC-006-F-002-S-005",
             "The API surface and any contract-level behavior in this story MUST comply with every REQ in EPIC-006-F-002-S-005 (API Surface and Contracts)."),
            ("REQ-B-003", "Filter shows specialty types relevant to the user's request",
             "When the user submits a healthcare request, the SpecialtyFilter panel is populated with specialty types relevant to that request. Relevance is determined by the system, not by the user."),
            ("REQ-B-004", "Default checkbox state for every returned specialty is CHECKED",
             "When the panel is initially populated, every returned specialty's checkbox is in the CHECKED state. The user uncheckes from a fully-selected starting point."),
            ("REQ-B-005", "List MUST include licensed CAM practitioners when clinically relevant",
             "The specialty list MUST include complementary, alternative, and integrative medicine (CAM) practitioners — Acupuncturist, Naturopath, Naprapath, Homeopath, Mechanotherapist, Chiropractor, Massage Therapist, and any other CAM type — whenever clinically relevant to the user's request. CAM practitioners MUST NOT be deprioritized, downweighted, or excluded on the basis of being CAM."),
            ("REQ-B-006", "List MUST NOT include unlicensed roles or specialties not licensed in any state or Puerto Rico",
             "The specialty list MUST NOT include unlicensed roles such as Homemaker, Driver, Religious Nonmedical Practitioner, Lay Midwife, Doula, Health & Wellness Coach, Health Educator, Lactation Consultant Non-RN, Peer Specialist, Community Health Worker, Case Manager, or Reflexologist. The list MUST NOT include any specialty that is not licensed in at least one US state or Puerto Rico."),
            ("REQ-B-007", "List MUST NOT include institutions, facilities, agencies, or non-individual entries",
             "The specialty list MUST NOT include hospitals, laboratories, pharmacies, supplier organizations, transportation services, residential treatment facilities, managed care organizations, agencies, group practices, or any other non-individual entity. Only individual licensed human practitioners belong in the list."),

            ("REQ-T-001", "AI normalize step uses prompts.json record",
             "The user's vernacular request is passed to an LLM normalize call using the system prompt registered at prompts.json[_record_id = 'specialty_normalize_system_prompt']."),
            ("REQ-T-002", "Embed normalize output with the canonical embedding model",
             "The normalize step's output is embedded using the canonical embedding model declared at EPIC-008-F-011-S-004-REQ-B-001."),
            ("REQ-T-003", "Vector search against SpecialtyMetaData with Section=Individual filter",
             "The embedded vector queries SpecialtyMetaData via MongoDB Atlas $vectorSearch with the filter Section='Individual'."),
            ("REQ-T-004", "Candidate pool above the cosine candidate floor passes to the AI filter step",
             "Vector-search results scoring at or above the cosine candidate floor are passed to the AI filter step. The cosine floor value itself is a design choice and is not specified at the requirement level."),
            ("REQ-T-005", "AI filter step uses prompts.json record; returns final NUCC code list",
             "The candidate pool is passed to an LLM filter call using the system prompt registered at prompts.json[_record_id = 'specialty_filter_system_prompt']. The filter call returns the final NUCC code list."),
            ("REQ-T-006", "Pipeline executes once per query; result list cached for downstream use",
             "The AI flow above runs exactly ONCE per user query. The result list is cached for downstream client-side filtering per EPIC-006-F-002-S-003-REQ-T-001."),
        ],
    },

    "EPIC-006-F-002-S-003": {
        "name": "User Manages Filter",
        "kind": "user journey",
        "description": (
            "The user refines the displayed specialty list — toggles "
            "prescriber/homeopathic categories, unchecks individual "
            "specialties, watches the counts update. The panel persists "
            "across service-control transfers."
        ),
        "reqs": [
            ("REQ-B-001", "Display MUST comply with EPIC-006-F-002-S-001",
             "The visual presentation of the SpecialtyFilter panel in this story MUST comply with every REQ in EPIC-006-F-002-S-001 (Display Specification)."),
            ("REQ-B-002", "User can uncheck individual specialties to narrow the selection",
             "The user can uncheck any individual specialty checkbox to remove that specialty from the active filter selection."),
            ("REQ-B-003", "Prescribers-Only checkbox filters visible list to prescriber-flagged specialties",
             "When the Prescribers-Only checkbox is checked, the displayed specialty list is restricted to specialties flagged as can_prescribe=true. When unchecked, all specialties are shown again."),
            ("REQ-B-004", "Homeopathic-Only checkbox filters visible list to homeopathic-flagged specialties",
             "When the Homeopathic-Only checkbox is checked, the displayed specialty list is restricted to specialties flagged as homeopathic=true. When unchecked, all specialties are shown again."),
            ("REQ-B-005", "Toggling Prescribers or Homeopathic MUST NOT trigger a backend query",
             "All effects of the Prescribers-Only and Homeopathic-Only checkboxes are realized client-side against the cached specialty list. No round-trip to the backend, no DB call, no observable latency."),
            ("REQ-B-006", "Check-all / uncheck-all toggle (label flips per majority state)",
             "A single toggle button in the panel header checks all currently-visible specialties or unchecks all of them. The button's label reads 'Uncheck All' when the majority of visible checkboxes are checked, and 'Check All' when the majority are unchecked."),
            ("REQ-B-007", "COUNT 'All Possible' — total NUCC specialties returned",
             "The 'All Possible' count displays the total number of NUCC specialties the system returned for the current request. This count does not change as the user toggles checkboxes — it reflects the size of the original returned set, constant for the lifetime of one query result."),
            ("REQ-B-008", "COUNT 'Prescribers' — count of prescriber-flagged specialties currently showing",
             "The 'Prescribers' count displays the number of prescriber-flagged specialties present in the currently-showing list (after any Prescribers-Only or Homeopathic-Only toggle filters)."),
            ("REQ-B-009", "COUNT 'Your Choices' — count of specialties currently checked",
             "The 'Your Choices' count displays the number of specialty checkboxes that are currently checked. This is the set that will be sent to Provider Search when the user hits Apply Filter."),
            ("REQ-B-010", "SpecialtyFilter is FindCare territory; persists across every service-control transfer",
             "The SpecialtyFilter MUST remain on the screen at all times when the wrapper is active, including across every service-control transfer (FindCare ↔ EvaluateCare ↔ SharedServices). The SpecialtyFilter is FindCare territory — it MUST be served by FindCare and no other service can or may serve this panel."),

            ("REQ-T-001", "Cache + client-side filtering — explicit exception to no-biz-logic-on-client",
             "The cached specialty list from EPIC-006-F-002-S-002-REQ-T-006 is the operand for all toggle/checkbox changes in this story. No backend round-trip occurs for any interaction. EXPLICIT EXCEPTION: this is an authorized exception to the project-wide rule 'no business logic on the client,' justified by interactivity latency requirements (toggle response must feel instantaneous). Only the final Apply Filter submission traverses to the backend per EPIC-006-F-002-S-004-REQ-B-003."),
        ],
    },

    "EPIC-006-F-002-S-004": {
        "name": "User Submits Filter to Provider Tab",
        "kind": "user journey",
        "description": (
            "The user commits the current selection by hitting Apply "
            "Filter. The Providers tab refreshes to reflect the new "
            "selection. Hand-off contract details live in S-005."
        ),
        "reqs": [
            ("REQ-B-001", "Display MUST comply with EPIC-006-F-002-S-001",
             "The visual presentation in this story MUST comply with every REQ in EPIC-006-F-002-S-001 (Display Specification)."),
            ("REQ-B-002", "API surface and contracts MUST comply with EPIC-006-F-002-S-005",
             "The API surface and any contract-level behavior in this story MUST comply with every REQ in EPIC-006-F-002-S-005 (API Surface and Contracts)."),
            ("REQ-B-003", "Apply Filter detects changes, clears unpicked providers, triggers new provider query",
             "When the user hits Apply Filter and the checkbox selection has changed since the last query, the unpicked provider list is emptied and a new provider query is triggered with the currently-checked NUCC codes plus the user's geo qualifier."),
            ("REQ-B-004", "Timer shown in bottom control frame on Apply Filter",
             "When Apply Filter triggers a new provider query, the timer MUST appear in the control frame at the bottom of the screen, in the same location as the initial search timer."),
            ("REQ-B-005", "Apply Filter button visibility — invisible until a change, disappears on submit",
             "The Apply Filter button is INVISIBLE by default. It becomes VISIBLE only when the user has made an actual change to the filter selection that has not yet been applied. When the user clicks Apply Filter, the button disappears until another change is made. (Placement governed by EPIC-006-F-002-S-001-REQ-B-019.)"),
        ],
    },

    "EPIC-006-F-002-S-005": {
        "name": "API Surface and Contracts",
        "kind": "cross-cutting technical",
        "description": (
            "The technical contracts that bound the Specialty Filter "
            "feature: API visibility model, input contract, hand-off "
            "contract to Provider Search, and DOM frame placement. "
            "Cross-cutting: applies uniformly across the user-journey "
            "stories."
        ),
        "reqs": [
            ("REQ-T-001", "API visibility — all Specialty Filter APIs are PRIVATE",
             "All Specialty Filter APIs are PRIVATE — in-process function calls within the FindCare service. The feature exposes ZERO public HTTP endpoints (no browser-callable routes) and ZERO protected HTTP endpoints (no peer-service-callable routes). All Specialty Filter functionality is reachable only by other code running inside the FindCare process. The folder layout MUST reflect this — Specialty Filter modules live under `private/` (or equivalent per the FindCare API folder convention)."),
            ("REQ-T-002", "Input contract — geographic qualifier supplied as a separate request field",
             "The user's geographic qualifier (ZIP code, county, city, or state) is supplied to the system as a SEPARATE request field. The system MUST NOT parse the user's vernacular utterance for geographic terms; geography never reaches the specialty-matching mechanism."),
            ("REQ-T-003", "Handoff contract — Apply Filter calls Provider Search API with {selected_nucc_codes, geo}",
             "On Apply Filter, the Specialty Filter calls the Provider Search API endpoint with a payload containing: selected_nucc_codes (array of NUCC code strings) and geo (object carrying the user's geographic qualifier). The endpoint name and canonical JSON shape live in the FindCare service-API contract (BUG-004 / FindCare API isolation refactor). The Provider Search feature's behavior on receipt is governed by EPIC-006-F-001 and is NOT specified here."),
            ("REQ-T-004", "DOM placement — rendered in parent wrapper DOM, NOT inside any service iframe",
             "The SpecialtyFilter MUST be rendered in the parent wrapper's DOM tree, NOT inside any service iframe (FindCare, EvaluateCare, or SharedServices). This is the technical mechanism that enables EPIC-006-F-002-S-003-REQ-B-010 (persistence across service-control transfers)."),
        ],
    },

    "EPIC-006-F-002-S-006": {
        "name": "Test Coverage",
        "kind": "cross-cutting test infrastructure",
        "description": (
            "Playwright UX test infrastructure for the SpecialtyFilter "
            "feature. Specifies where the tests live, how each REQ maps "
            "to a test_id, and the current implementation status of the "
            "test file. The tests themselves are placeholders awaiting "
            "Skip-approved test plans."
        ),
        "reqs": [
            ("REQ-T-001", "Test file location — uniform per-service tests/ directory",
             "All Playwright UX tests for the Specialty Filter feature live in a single TypeScript test file at Code/ConversationalUX/FindCareChat/tests/specialty_filter.spec.ts. This location follows the uniform per-service tests/ directory convention (parallel to backend/ and frontend/), mirroring the same convention used by SharedServices and EvaluateCare."),
            ("REQ-T-002", "Test naming — one stub per REQ, named by REQ id",
             "The test file contains one Playwright test stub per REQ in this feature (S-001 through S-005). Test ids follow the convention {full-REQ-id}-TEST-1 (e.g., EPIC-006-F-002-S-001-REQ-B-001-TEST-1). Each REQ in this feature references its corresponding test_id in its own `tests` field."),
            ("REQ-T-003", "Test file status — NOT STARTED",
             "The test file is currently NOT STARTED. The file exists at the specified location and contains a stub function for every REQ, but every test body is a placeholder pending a Skip-approved test plan (per the no-Claude-authored-pytests rule). This story's status remains not_started until each test body is implemented and Skip approves the test plan."),
        ],
    },
}


def set_run(run, *, bold=False, italic=False, size=11, color=None):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_field(doc, label, value, *, mono_value=False):
    """Add a single 'field: value' paragraph with bold label."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(18)
    set_run(p.add_run(f"{label:<14} : "), bold=True, size=10)
    if mono_value:
        run = p.add_run(value)
        run.font.name = "Consolas"
        run.font.size = Pt(10)
    else:
        set_run(p.add_run(value), size=10)


def main() -> int:
    doc = Document()

    doc.add_heading("EPIC-006-F-002 Specialty Filter — Requirements Specification", level=0)

    p = doc.add_paragraph()
    set_run(p.add_run("Approved by Skip on 2026-05-04. "), bold=True)
    p.add_run(
        "Five stories: two cross-cutting (Display Specification S-001, "
        "API Surface and Contracts S-005) plus three user-journey "
        "(Present S-002, Manage S-003, Submit S-004). Within each "
        "story, business REQs precede technical REQs. All REQs are "
        "user-visible or contract-level. Implementation choices "
        "(pipeline stages, models, prompts, cosine thresholds) are "
        "design and live in design docs, not REQs. Two new "
        "prompts.json records — `specialty_normalize_system_prompt` "
        "and `specialty_filter_system_prompt` — are referenced by "
        "S-002-T-001 and S-002-T-005."
    )
    doc.add_paragraph()
    p = doc.add_paragraph()
    set_run(p.add_run(
        "This document is laid out to mirror the agile_backlog.json schema. "
        "Each REQ shows every schema field (req_id, name, priority, status, "
        "approval, orphan, depends_on, tests, requirement body) so the "
        "Word doc and the JSON record carry identical substance."
    ), italic=True, size=10)
    doc.add_paragraph()

    # Feature header (mirrors feature schema)
    doc.add_heading("Feature", level=1)
    add_field(doc, "feature_id", "EPIC-006-F-002", mono_value=True)
    add_field(doc, "name", "Specialty Filter")
    add_field(doc, "status", "in_progress")
    add_field(doc, "approval", "approved")
    add_field(doc, "description",
              "Filter search results by specialty type, prescriber status, "
              "and homeopathic status. Five-story decomposition: Display "
              "Specification (cross-cutting visual), Present Initial Filter, "
              "User Manages Filter, User Submits Filter, API Surface and "
              "Contracts (cross-cutting technical).")
    doc.add_paragraph()

    n_b = 0
    n_t = 0

    for story_id, story in TREE.items():
        doc.add_heading(f"{story_id} — {story['name']}", level=1)

        # Story header (mirrors story schema)
        add_field(doc, "story_id", story_id, mono_value=True)
        add_field(doc, "name", story["name"])
        add_field(doc, "kind", story["kind"])
        add_field(doc, "status", "in_progress")
        add_field(doc, "approval", "approved")
        add_field(doc, "description", story["description"])
        doc.add_paragraph()

        for req_id, req_name, req_body in story["reqs"]:
            full_id = f"{story_id}-{req_id}"
            doc.add_heading(full_id, level=2)

            # All schema fields per REQ
            add_field(doc, "req_id", full_id, mono_value=True)
            add_field(doc, "name", req_name)
            add_field(doc, "priority", "must-have")
            add_field(doc, "status", "not_started")
            add_field(doc, "approval", "pending")
            add_field(doc, "orphan", "false")
            add_field(doc, "depends_on", "[]")
            # tests block — schema-shaped (object with `test` array)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            set_run(p.add_run("tests          : "), bold=True, size=10)
            p2 = doc.add_paragraph("test:")
            p2.paragraph_format.left_indent = Pt(36)
            p2.runs[0].font.name = "Consolas"
            p2.runs[0].font.size = Pt(10)
            for line in [
                f"- test_id   : {full_id}-TEST-1",
                f"  test_type : playwright",
                f"  orphan    : false",
                f"  file      : Code/ConversationalUX/FindCareChat/tests/specialty_filter.spec.ts",
            ]:
                p3 = doc.add_paragraph(line)
                p3.paragraph_format.left_indent = Pt(54)
                p3.runs[0].font.name = "Consolas"
                p3.runs[0].font.size = Pt(10)

            # Requirement body — full text
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            set_run(p.add_run("requirement    : "), bold=True, size=10)
            doc.add_paragraph()
            p = doc.add_paragraph(req_body)
            p.paragraph_format.left_indent = Pt(36)

            doc.add_paragraph()

            if "REQ-B" in req_id:
                n_b += 1
            else:
                n_t += 1

    # Summary
    doc.add_heading("Summary", level=1)
    doc.add_paragraph(
        f"Stories: {len(TREE)}\n"
        f"Total REQs: {n_b + n_t} ({n_b} business + {n_t} technical)\n"
    )

    # Compliance gaps
    doc.add_heading("Known Compliance Gaps (current implementation vs. spec)", level=1)
    doc.add_paragraph(
        "The following gaps were identified by reverse-engineering the "
        "current SpecialtyFilter implementation against this spec. They "
        "constitute the implementation backlog flowing from this spec:"
    )
    gaps = [
        "S-002-T-005 (AI filter step) — does not exist in production. Backend has only the cosine 0.66 cut, no LLM filter call.",
        "S-002-B-005 (Holistic CAM inclusion) — production normalize prompt does not encode Rule #2; the rule lives only in the test harness. Phase C data shows Acupuncturist dropped 100% of the time on foot queries.",
        "S-002-B-006 (Licensed-only) — no enforcement in code; relies on BUG-005 fix for upstream data cleanliness.",
        "S-003-T-001 (Cache + client-side filtering) — only `window._homeoGeneralists` is cached; the full specialty list is re-rendered on every toggle.",
        "S-001 + S-004-B-005 (Apply Filter button visibility) — current implementation always shows the button and displays a 'No filter changes to apply' message when clicked without changes; the spec eliminates that message-on-click pattern in favor of hiding the button entirely until there is something to apply.",
        "S-002-T-001 / T-005 (Prompts.json registry) — normalize prompt lives inline in `specialty_service.py`, not registered as a brain record. The two new records `specialty_normalize_system_prompt` and `specialty_filter_system_prompt` will be authored alongside this spec.",
        "S-005-T-001 (Private API folder layout) — current FindCare main.py mixes public/protected/private endpoints; the planned strangler-fig refactor (BUG-004) will isolate Specialty Filter modules under `private/`.",
    ]
    for g in gaps:
        doc.add_paragraph(g, style="List Bullet")

    doc.save(OUT)
    print(f"Wrote {OUT}")
    print(f"  Stories: {len(TREE)}")
    print(f"  REQs: {n_b + n_t} ({n_b} business + {n_t} technical)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
