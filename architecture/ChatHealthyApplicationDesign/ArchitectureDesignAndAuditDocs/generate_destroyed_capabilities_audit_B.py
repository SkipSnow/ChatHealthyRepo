# Copyright (c) 2026 ChatHealthy.ai LLC. All rights reserved.
#
# generate_destroyed_capabilities_audit_B.py
# Builds destroyed-capabilities-audit-B.docx
# Read-only capability audit. No commits, no code changes.
#
# Scope: front-end capabilities that existed at the latest pre-2026-04-03 commit
# but no longer exist today AND are NOT recovered by any V2 BR in
# design-V2.docx. One row per user-observable behavior.

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUT = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "destroyed-capabilities-audit-B.docx",
)


def shade_cell(cell, fill_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def set_col_widths(table, widths_inches):
    for i, w in enumerate(widths_inches):
        for cell in table.columns[i].cells:
            cell.width = Inches(w)


def add_table(doc, headers, rows, col_widths=None, header_fill="1F4E79"):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(10)
        shade_cell(hdr[i], header_fill)
    for row in rows:
        rc = t.add_row().cells
        for i, val in enumerate(row):
            rc[i].text = str(val) if val is not None else ""
            for p in rc[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    if col_widths:
        set_col_widths(t, col_widths)
    return t


# ---------------------------------------------------------------------------
# Inputs
# ---------------------------------------------------------------------------
PRE_COMMIT = "7c01b8cc7ab947bbeccb5b72d13e85a47311aa94"
PRE_DATE   = "2026-04-03 11:37:33 -0700"
PRE_BUILD  = "431"


# Each row: (#, capability description, pre-event evidence, current evidence of absence)
# Front-end only. User-observable. NOT covered by any V2 BR.
ROWS = [
    (
        "1",
        "System-built summary message rendered as its own chat bubble after a "
        "provider search (FC-RESULT-MSG / GOV-011-STD-001). The bubble told the "
        "user how many providers matched, named the searched-for term in the "
        "user's own words, named the matched state and any city/county "
        "narrowing, and offered next-step affordances. After 4/3 the bubble "
        "is no longer pushed to chat for provider results: chat now shows a "
        "one-liner pointing the user to a separate selection panel.",
        "Code/ConversationalUX/FindCareChat/frontend/src/components/ChatWindow.tsx:389-395 "
        "(at " + PRE_COMMIT[:10] + ") — pushes data.pagination.summary_message "
        "as an assistant message with isSummary:true. Backend producer: "
        "Code/ConversationalUX/FindCareChat/backend/main.py:469-489.",
        "Code/ConversationalUX/FindCareChat/frontend/src/components/ChatWindow.tsx:559-561 "
        "(current) — explicit comment 'system-built summary — suppressed when "
        "selection panel active … BUG-UX-010'. The provider-summary push is "
        "removed; only data.trials.summary_message still pushes a bubble "
        "(line 564-569).",
    ),
    (
        "2",
        "Click-to-act action links inside the system summary bubble: "
        "[next page](#action:next-page) advanced pagination without typing, "
        "and [refine by specialty](#action:filter) flashed/highlighted the "
        "left-panel filter. The handlers were wired in MessageBubble; the "
        "links were authored by the backend summary builder. Today the "
        "handlers are still defined but the bubble that contained the links "
        "is no longer emitted, so the user can never click them.",
        "Code/ConversationalUX/FindCareChat/frontend/src/components/MessageBubble.tsx:39-47 "
        "(at " + PRE_COMMIT[:10] + ") — handleActionLink switches on "
        "'#action:filter' and '#action:next-page'. Bubble emission: "
        "frontend ChatWindow.tsx:389-395 (same commit).",
        "Code/ConversationalUX/FindCareChat/frontend/src/components/MessageBubble.tsx:39-47 "
        "(current) — handlers preserved but unreachable; the host bubble is "
        "no longer pushed (ChatWindow.tsx:559-561 suppresses it).",
    ),
    (
        "3",
        "Full provider listing rendered inline in the chat conversation on "
        "the first results page. Pre-4/3 the LLM-authored markdown listing "
        "(name / specialty / address / phone / NPI per provider) appeared "
        "in a chat bubble alongside the rest of the assistant's prose. "
        "Today the chat bubble for the first page is replaced by a fixed "
        "string 'Found N providers. Use the panel below to select up to 5 "
        "for evaluation.' — the listing itself moved to a separate "
        "selection panel (BR-24 names that panel but does NOT mandate the "
        "in-chat listing).",
        "Code/ConversationalUX/FindCareChat/frontend/src/components/ChatWindow.tsx:362-367 "
        "(at " + PRE_COMMIT[:10] + ") — pushes data.response (LLM-authored "
        "listing) verbatim as an assistant bubble.",
        "Code/ConversationalUX/FindCareChat/frontend/src/components/ChatWindow.tsx:493-502 "
        "(current) — when data.pagination.total_count > 0 the assistant "
        "bubble is replaced with the literal 'Found ${data.pagination.total_count} "
        "providers. Use the panel below to select up to 5 for evaluation.' "
        "(BUG-UX-010 / FC-SELECT-001).",
    ),
    (
        "4",
        "LLM-authored narrative around provider results — introductions "
        "('Here are the first…'), location-narrowing prompts ('Are you "
        "looking in a specific city…'), filter suggestions, follow-up "
        "questions every 5 user turns, etc. These were spoken by the LLM "
        "in the same response as the listing. Today the LLM's text is "
        "discarded entirely whenever providers are returned (replaced by "
        "the one-liner above), so the conversational shell goes silent on "
        "results turns. No V2 BR re-mandates the narrative.",
        "Code/ConversationalUX/FindCareChat/backend/main.py:415, 456-459 "
        "(at " + PRE_COMMIT[:10] + ") — system prompt asks for the prose; "
        "_strip_redundant_summary keeps the listing AND the prose around it; "
        "main.py returns text=…full LLM output. Frontend pushes it (ChatWindow.tsx:362-367).",
        "Code/ConversationalUX/FindCareChat/frontend/src/components/ChatWindow.tsx:493-497 "
        "(current) — `const chatContent = hasProviders ? \"Found …\" : data.response`. "
        "When hasProviders is true the LLM's response is dropped on the floor.",
    ),
    (
        "5",
        "HUMAN_TESTING-mode UAT-rich welcome message. When the env var "
        "HUMAN_TESTING was truthy, GET /welcome returned a build/version-rich "
        "UAT banner produced by build_uat_welcome (it surfaced the live "
        "build number, environment, and a UAT navigation hint). Today the "
        "endpoint always returns the static WELCOME_MESSAGE regardless of "
        "HUMAN_TESTING; build_uat_welcome is still imported and "
        "_build_test_welcome is still defined but no longer called.",
        "Code/ConversationalUX/FindCareChat/backend/main.py:278-280 "
        "(at " + PRE_COMMIT[:10] + ") — `def welcome(): return {\"message\": "
        "_build_test_welcome() if _HUMAN_TESTING else WELCOME_MESSAGE}`.",
        "Code/ConversationalUX/FindCareChat/backend/main.py:572-574 "
        "(current) — `def welcome(): return {\"message\": WELCOME_MESSAGE}`. "
        "The HUMAN_TESTING branch is gone; _build_test_welcome (line 165) is "
        "dead code.",
    ),
]


# ---------------------------------------------------------------------------
# Build doc
# ---------------------------------------------------------------------------
def build():
    doc = Document()

    # Title
    title = doc.add_heading("Destroyed Front-End Capabilities Audit (Run B)", level=0)

    # 3-line executive summary
    p = doc.add_paragraph()
    p.add_run("Pre-event commit: ").bold = True
    p.add_run(PRE_COMMIT + " (" + PRE_DATE + ", build " + PRE_BUILD + ").")

    p = doc.add_paragraph()
    p.add_run("Destroyed front-end capabilities NOT recovered by any V2 BR: ").bold = True
    p.add_run(str(len(ROWS)) + ".")

    p = doc.add_paragraph()
    p.add_run("Scope: ").bold = True
    p.add_run(
        "front-end surfaces only — Website/index.html (static parent page + "
        "iframe wrapper), the React app under "
        "Code/ConversationalUX/FindCareChat/frontend/, and the FastAPI routes "
        "in backend/main.py that the React app consumes. Each row below is a "
        "behavior the user could observe pre-event and can no longer observe "
        "today; cross-checked against design-V2.docx Table 1 (BR-01..BR-35) — "
        "no BR mandates the listed behavior's recovery."
    )

    doc.add_paragraph()  # spacer

    # The list
    add_table(
        doc,
        ["#", "Destroyed capability (user-facing)",
         "Pre-event evidence (file:line at " + PRE_COMMIT[:10] + ")",
         "Current evidence of absence"],
        ROWS,
        col_widths=[0.4, 2.6, 2.4, 2.1],
    )

    doc.save(OUT)
    print("Wrote", OUT)


if __name__ == "__main__":
    build()
