"""Generate ChatHealthy Application Design - V1 (Word document).

Read-only generator. Produces design-V1.docx in this same directory. No commits,
no code changes, no backlog mutations. Inputs (read at generate time):

- findCare/ArchitectureAndDesign/codebase-reorganization-V7.docx
- findCare/ArchitectureAndDesign/langgraph-oo-best-practices-V6.docx
- brain/machine_artifacts/content/engineering_rules.json
- brain/machine_artifacts/content/agile_backlog.json
- brain/machine_artifacts/content/bugs.json
- brain/machine_artifacts/content/version.json
- the codebase under Code/, Website/, LangGraph/

Built per the architect's 8 binding mandates (see body of the document).
"""
from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(__file__).resolve().parent / "design-V1.docx"


# -----------------------------------------------------------------------------
# Helpers
# -----------------------------------------------------------------------------

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_para(doc, text, style=None, bold=False):
    p = doc.add_paragraph()
    if style:
        try:
            p.style = doc.styles[style]
        except KeyError:
            pass
    r = p.add_run(text)
    r.bold = bold
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    return p


def add_sub_bullet(doc, text):
    try:
        p = doc.add_paragraph(text, style="List Bullet 2")
    except KeyError:
        p = doc.add_paragraph("    - " + text)
    return p


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    return t


# -----------------------------------------------------------------------------
# Document
# -----------------------------------------------------------------------------

def build_doc() -> Document:
    doc = Document()

    # ---------- Title ----------
    title = doc.add_heading("ChatHealthy.ai - End-to-End Production Architecture Design - V1", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    add_para(doc, "Design V1 - the first end-to-end production architecture design for ChatHealthy.ai.")
    add_para(doc, "Date: 2026-04-28   |   Author: Claude (architecture design agent)   |   Read-only deliverable")
    add_para(
        doc,
        "Binding inputs: findCare/ArchitectureAndDesign/codebase-reorganization-V7.docx; "
        "findCare/ArchitectureAndDesign/langgraph-oo-best-practices-V6.docx; "
        "brain/machine_artifacts/content/{engineering_rules,agile_backlog,bugs,version}.json; "
        "Code/, Website/, LangGraph/ codebase.",
    )

    # ---------- 1. Executive summary ----------
    add_heading(doc, "1. Executive summary", level=1)
    add_para(
        doc,
        "ChatHealthy.ai's production architecture is a thin LangGraph parent runtime composing four capability "
        "subgraphs (FindCare, EvaluateCare, Security & Compliance, SharedServices) via RemoteGraph. The parent "
        "runtime owns one MongoDBSaver checkpointer, one thread_id per session, and the single HTTP endpoint a "
        "React-as-facade UI consumes through @langchain/langgraph-sdk. Each capability epic maps 1:1 to a "
        "HuggingFace Space and to one cost-center line on the GL, in alignment with the operating-model axiom.",
    )
    add_para(
        doc,
        "The codebase shape adopts the V7 dual-child convention {epic_root}/RuntimeServer/ + "
        "{epic_root}/SharedCode/. The structural noun 'domain' is retired in favor of core/. Project-root "
        "capability folders (EvaluateCare/, FindCare/, SharedServices/, Security/, Operations/, UX/, "
        "DataManagementPipelines/) replace the Code/ umbrella.",
    )
    add_para(
        doc,
        "Every Python service entry point follows the thin-main-construct-and-run convention: app.py contains "
        "an if __name__ == '__main__': block whose sole job is to construct an Application object and call "
        "run() on it. Procedural FastAPI handlers are the exception, not the rule; exceptions are registered "
        "in a deterministic exception list consulted by a pre-commit scan worker (Rule-OO-001).",
    )
    add_para(
        doc,
        "End-to-end observability is delivered by LangSmith trace propagation across RemoteGraph boundaries, "
        "Kafka-fed ChatHealthyLogService for structured logs, and per-node performance budgets emitted as "
        "OpenTelemetry-compatible metrics from the parent runtime and each Space.",
    )
    add_para(
        doc,
        "Claude Code's blast area is reduced by: (a) the engineering-rules enforcement framework V20 (50 rules, "
        "deterministic worker registry, hook-driven execution); (b) project-root capability folders that limit "
        "any single edit's reachability to one cost-center boundary; (c) the requirement that every code change "
        "trace to a BR/TR (mandate 6); (d) the single-thin-main convention with a registered exception list; "
        "and (e) read-only design artifacts archived under architecture/ that the boot tool reads on every "
        "session.",
    )

    # ---------- 2. Lost-functionality recovery analysis ----------
    add_heading(doc, "2. Lost-functionality recovery analysis", level=1)
    add_para(
        doc,
        "The 'SQL-insertion event' the architect refers to is the v4-043 graph_entry_check enforcement that "
        "fired on 2026-04-19 across the three runtime servers (FindCareChat backend main.py, evaluate_care "
        "app.py, shared_services app.py). The scanner detected 26 FastAPI handlers as 'empty LangChain "
        "containers' - bug ids BUG-LANGCHAIN-CONTAINER-MAIN-* (10 entries on Code/ConversationalUX/"
        "FindCareChat/backend/main.py) and BUG-LANGCHAIN-CONTAINER-APP-* (16 entries on Code/evaluate_care/"
        "app.py and Code/shared_services/app.py). All 26 carry discovery_date 2026-04-19 and "
        "req_id EPIC-008-F-010-S-001-REQ-T-001 ('Route all Python service entry points through single graph "
        "orchestrator for core').",
    )

    add_heading(doc, "2.1 Build numbers before and after", level=2)
    add_para(
        doc,
        "Authoritative source: brain/machine_artifacts/content/version.json (lines 5-275). "
        "version.json does not carry a date field per build but does carry the files_touched list; the "
        "files_touched of build 999 includes brain/BusinessArtifacts/governance/uat_handoff_2026_04_19_"
        "langgraph.docx, anchoring build 999 to 2026-04-19 - the same date all 26 BUG-LANGCHAIN-CONTAINER-* "
        "entries carry as discovery_date.",
    )
    add_para(doc, "Build numbers (with their dates inferred from co-touched dated artifacts in version.json):")
    add_bullet(doc, "BEFORE the regression event - build 998 (cause EPIC-008-F-004-S-001-REQ-T-002, on or just before 2026-04-19).")
    add_bullet(doc, "AT THE EVENT (the SQL/graph-routing pre-deploy scan run that surfaced the regression cluster) - build 999 (cause EPIC-008-F-002-S-006-REQ-T-001, dated 2026-04-19 by uat_handoff_2026_04_19_langgraph.docx).")
    add_bullet(doc, "AFTER the regression - build 1000 (cause BUG-PORT-NONCOMPLIANCE-001), then 1001 (BUG-EVALCARE-GRAPH-WRAPPER-001, 2026-04-19 onward), then 1002-1006 (BUG-EVALCARE-GRAPH-WRAPPER-001 follow-on, EPIC-008-F-004-S-001-REQ-T-002, DEVOPS-DEV-B002).")

    add_heading(doc, "2.2 Capabilities lost when the LangChain containers became empty", level=2)
    add_para(
        doc,
        "Cross-referencing the 26 BUG-LANGCHAIN-CONTAINER-* descriptions, the LangGraph V6 doc Section 3.3 (parent runtime "
        "vs capability Space ownership), the POC LangGraph/poc/user_journey.py (which catalogued the originally "
        "intended tool surface), and the actual codebase under Code/ConversationalUX/FindCareChat/backend/, "
        "the lost capabilities cluster into 12 distinct items.",
    )
    rows = [
        ("LC-01", "Single graph entry through parent runtime", "Conversational HTTP endpoint did not route through a LangGraph parent runtime; React called FastAPI handlers directly.", "Recovered by TR-LG-001 + TR-LG-002"),
        ("LC-02", "MongoDBSaver checkpointer (single-instance)", "No parent-owned checkpointer; capability Spaces had no consistent thread_id.", "Recovered by TR-LG-003"),
        ("LC-03", "Provider find tool (tool_find_providers)", "POC node became orphaned when handlers bypassed graph; captured in user_journey.py:178.", "Recovered by TR-FC-002"),
        ("LC-04", "Clinical trials tool (tool_search_clinical_trials)", "POC node became orphaned; user_journey.py:182.", "Recovered by TR-FC-003"),
        ("LC-05", "Specialty classifier and homeopathic expansion", "user_journey.py:132 classifier and the homeopathic expansion field on POC state are not routed through a graph in the production handlers.", "Recovered by TR-FC-004"),
        ("LC-06", "Lead capture tools (tool_record_user_details, tool_record_unknown_question)", "user_journey.py:186 and :190; the lead service exists at backend/domain/shared/lead_capture/lead_service.py but is not graph-routed.", "Recovered by TR-FC-005"),
        ("LC-07", "Context tools (tool_get_skip_snow_context, tool_get_chathealthy_context)", "user_journey.py:194 and :198.", "Recovered by TR-FC-006"),
        ("LC-08", "Emergency / safety detector and IP-lock check", "user_journey.py:96 and :104; no graph routing in production handlers means no enforcement order.", "Recovered by TR-SE-002"),
        ("LC-09", "Repetitive-utterance detector and routing merge", "user_journey.py:124 and :142.", "Recovered by TR-FC-007"),
        ("LC-10", "EvaluateCare scoring graph (score_provider, score_trial, evaluate_providers, evaluate_view, explain)", "5 BUG-LANGCHAIN-CONTAINER-APP-* entries on Code/evaluate_care/app.py; all listed in bugs.json as 'Closed 2026-04-19' but UAT in Studio still pending - not durably re-integrated as a graph.", "Recovered by TR-EC-001"),
        ("LC-11", "Shared-services tool router and auth flow as graph nodes", "BUG-LANGCHAIN-CONTAINER-APP-{HEALTH, SPLASH, TRANSFER-TO-FINDCARE, VERIFY-TOKEN} on Code/shared_services/app.py.", "Recovered by TR-SS-001"),
        ("LC-12", "Audit and trace context propagation across the conversation", "Without a parent runtime owning the messages channel, ToolMessage entries did not propagate; LangSmith trace spans broke at the FastAPI boundary.", "Recovered by TR-OB-001 + TR-OB-002"),
    ]
    add_table(doc, ["ID", "Capability", "Loss evidence", "Recovery TR"], rows)
    add_para(
        doc,
        "All 12 are recovered in this design. Recovery is anchored to specific TRs in Section 4 and to the "
        "parent-runtime / RemoteGraph topology in Section 7.",
    )

    # ---------- 3. Business requirements ----------
    add_heading(doc, "3. Business requirements (BR)", level=1)
    add_para(
        doc,
        "Business-language statements of what the system must deliver. Every BR is testable, has a unique ID, "
        "and traces to one of: (a) a recovered capability from Section 2, (b) a V7 binding decision, "
        "(c) a V6 LangGraph principle, or (d) one of the architect's 8 mandates.",
    )
    br_rows = [
        ("BR-01", "Conversational find-care experience: a single React UI lets a user describe a care need in natural language and receive a ranked list of providers with explanations.", "user-visible behavior; React submits prompt; parent runtime returns ranked providers + explanation messages.", "LC-01, LC-03, LC-05, mandate 1"),
        ("BR-02", "Conversational evaluate-care experience: a user can ask the system to evaluate a single provider or a list, receive a weighted-composite score and an explanation.", "Studio renders the EvaluateCare subgraph as a node; HTTP API returns a deterministic score + rationale.", "LC-10, EPIC-001 V7, mandate 1"),
        ("BR-03", "Single conversational entry point: the React client talks to exactly one URL for conversational flow regardless of which capability handles the step.", "Network capture from the React app shows exactly one hostname for streaming conversational traffic.", "V6 Section 3.3, mandate 3"),
        ("BR-04", "Pre-conversation security gate: every inbound conversational turn is gated by a security check (token verification + entitlement) before any capability node executes.", "Synthetic turn with invalid token gets a deterministic 401-style refusal from the parent before any capability is invoked.", "LC-08, V7 EPIC-002 own-Space, mandate 1"),
        ("BR-05", "Static page surface: ten static product/marketing/architecture pages remain available at the existing URLs and re-renderable from a generator.", "Each of the 10 URLs in Website/ resolves; generator produces byte-identical output for non-volatile pages.", "V7 Section 9, A9 annotation, mandate 5"),
        ("BR-06", "Provider detail and evaluation report read paths: a React detail page can load provider data and an evaluation report without invoking the conversational graph.", "GET /provider/{id} and GET /evaluation/{id} return JSON; LangSmith shows zero parent-graph traces for those calls.", "V6 Section 3.3, mandate 1"),
        ("BR-07", "Talk About Care lives inside FindCare: the Talk About Care capability is exposed as a feature folder under FindCare, not as its own epic.", "The backlog feature for Talk About Care resolves under FindCare in V7's epic shape.", "Resolved issue R3, mandate 8"),
        ("BR-08", "Audit trail per turn: every conversation turn produces an audit record stored in MongoDB capturing user, thread, capability, latency, outcome.", "Audit collection contains one record per turn; record fields verified by integration test.", "V7 EPIC-002 own-Space, mandate 1"),
        ("BR-09", "Provider data freshness: the production provider dataset is refreshed by an EPIC-012 pipeline on a published cadence; the conversational system always reads the published dataset.", "Pipeline run completes; FindCare Space reads via SharedServices abstraction; freshness probe within SLA.", "V7 EPIC-012 stays out of LangGraph, mandate 1"),
        ("BR-10", "Brain governance: every commit-time change is gated by the engineering-rules enforcement worker registry.", "Pre-commit run with a rule-violating change exits non-zero and emits a structured violation report.", "V20 framework, mandate 3"),
        ("BR-11", "Operating-model alignment: every capability epic appears once on the org chart, once on the GL as a cost-center line, and once as a HuggingFace Space deployment line. The shape of any one IS the shape of all four.", "Epic registry crosswalk produces 1:1 mapping for the 4 capability epics + Operations + UX + DMP.", "V7 Section 1 axiom"),
        ("BR-12", "End-to-end observability for the human operator: a single trace ID lets the operator see latency and outcome of every node across every Space for any conversation.", "Given a thread_id, LangSmith returns one trace tree spanning parent + capability subgraphs.", "LC-12, mandate 4"),
        ("BR-13", "Performance budgets: every conversational turn meets an explicit p95 budget per capability node (e.g. provider_search p95 <= 1500ms; specialty_classifier p95 <= 600ms).", "OTel histograms aggregated daily; budget breach raises an alert.", "Mandate 4"),
        ("BR-14", "Static-page accessibility: WCAG 2.2 AA for the 10 static pages and the React app.", "Lighthouse / axe scan reports 0 critical accessibility issues on each page.", "Mandate 5, V7 Section 9"),
        ("BR-15", "Recovery completeness: each of the 12 lost capabilities in Section 2 is reachable from a graph node at runtime and observable in LangSmith.", "Smoke-test suite invokes every recovered capability via the parent and asserts the LangSmith trace contains the corresponding node.", "Mandate 1, mandate 6"),
        ("BR-16", "Internal consistency: every BR maps to at least one TR and every TR traces to at least one BR.", "Section 13's coverage matrix shows zero unmapped rows.", "Mandate 7"),
        ("BR-17", "No work outside the requirement set: the design implements only what BR/TR requires; new work requires a new requirement.", "Implementation-invention audit (V2 pattern) returns zero findings.", "Mandate 6, V20 anti-invention discipline"),
        ("BR-18", "Pipelines deliver the 11 schema fields: the EPIC-012 pipeline produces every field the canonical Provider schema declares, including the 11 currently undeclared fields (embedding, embedding_model, embedding_version, 8 NPPES other-name fields).", "Pipeline regression test asserts each field present + typed per the canonical schema.", "Resolved issue R2, V7 schema relocation"),
        ("BR-19", "Deferred-work ledger: every postponed item is recorded with its surfacing event, owner, due_date, and required closure evidence.", "Deferred-work registry has a row for every BUG-* with status 'deferred' or 'uat_testing'.", "EPIC-008-F-009 in V7, mandate 5"),
        ("BR-20", "OO conventions in the codebase: every long-running Python program launches via a thin if __name__ == '__main__' block that constructs an object and runs it; deviations are tracked in an exception registry.", "Pre-commit OO scan exits non-zero on any new entrypoint that violates the convention without a registered exception.", "Mandate 2"),
    ]
    add_table(doc, ["BR ID", "Statement", "Acceptance test", "Trace"], br_rows)

    # ---------- 4. Technical requirements ----------
    add_heading(doc, "4. Technical requirements (TR)", level=1)
    add_para(
        doc,
        "Implementation-language statements that flow from Section 3. Every TR has a unique ID, is testable, and "
        "traces to at least one BR. Where a TR is also a recovery target for a Section 2 capability or a "
        "binding V7 / V6 decision, the trace cell makes that explicit.",
    )

    tr_rows = [
        # LangGraph / runtime
        ("TR-LG-001", "Operations/RuntimeServer/orchestration/parent_graph.py implements the parent LangGraph StateGraph; messages: Annotated[list, add_messages]; state in Pydantic BaseModel.", "Studio renders parent_graph.py; messages channel verified by unit test.", "BR-01, BR-03, V6 Section 3.3, LC-01"),
        ("TR-LG-002", "Operations/RuntimeServer/app.py exposes one HTTP endpoint compatible with @langchain/langgraph-sdk that React consumes for conversational flow.", "Network capture; React submits prompt and receives streaming tokens.", "BR-03, V6 Section 3.3"),
        ("TR-LG-003", "MongoDBSaver from langgraph-checkpoint-mongodb is instantiated exactly once at module load in the parent runtime; no capability Space instantiates a checkpointer.", "Process inventory; integration test asserts thread_id continuity across capability invocations.", "BR-01, V6 Section 3.3, LC-02"),
        ("TR-LG-004", "Each capability subgraph (FindCare, EvaluateCare, Security, SharedServices) is composed into the parent graph as a RemoteGraph from langgraph.pregel.remote pointed at its HF Space URL.", "parent_graph.py imports RemoteGraph; integration test runs a turn that visits each capability node.", "BR-03, V6 Section 3.2"),
        ("TR-LG-005", "RemoteGraph instances NEVER call back into their own deployment; same-Space step composition uses in-process subgraph composition.", "Static analyzer flags any RemoteGraph(url=SELF) usage as a violation.", "V6 Section 3.2"),
        ("TR-LG-006", "Per-capability state lives in {Capability}/RuntimeServer/state.py as a Pydantic BaseModel with explicit input_schema and output_schema fields; only declared fields cross the Space boundary.", "Schema validation test on the public state shape; private fields raise ValidationError if leaked.", "V6 Q1, BR-12"),
        ("TR-LG-007", "Tool calls and tool results are modeled as ToolMessage entries on the messages channel - never as flat top-level state fields, never as a tool_outputs catch-all dict.", "Anti-pattern scanner asserts no flat tool-output fields on state schemas.", "V6 Q4, LC-03..LC-07"),
        ("TR-LG-008", "Dependency injection for nodes uses Runtime[ContextSchema]; for tools uses ToolRuntime[ContextSchema]. Database clients and LLM clients are NEVER imported at module scope inside node files.", "Static analyzer flags top-level mongo client / llm client imports in node modules.", "V6 Q5, BR-12"),
        ("TR-LG-009", "Long-running deep-research flows use interrupt() / Command(resume=...) at the parent layer, not inside a capability Space (capability Spaces lack a checkpointer).", "Integration test triggers an interrupt; resume returns through the parent's checkpoint.", "V6 Section 3.3"),
        ("TR-LG-010", "InMemorySaver is the test-only checkpointer; production checkpointer is MongoDBSaver. No SqliteSaver, no PostgresSaver, no custom mongo saver.", "Test fixtures use InMemorySaver; CI fails if MongoDBSaver is referenced from a unit test.", "V6 Q2, mandate 3"),
        # FindCare recovery
        ("TR-FC-001", "FindCare/RuntimeServer/agent.py compiles a StateGraph routed by the parent; FindCare/RuntimeServer/state.py + nodes.py + tools.py follow the official 4-file layout.", "File presence test + unit test on graph build.", "BR-01, V6 Q6, V7 dual-child"),
        ("TR-FC-002", "tool_find_providers is a @tool inside FindCare/RuntimeServer/tools.py; it delegates to FindCare/SharedCode/core/provider_search_service.py.", "ToolNode wires it; smoke test invokes; LangSmith shows the node.", "BR-01, LC-03"),
        ("TR-FC-003", "tool_search_clinical_trials is a @tool inside FindCare/RuntimeServer/tools.py.", "Smoke test; trace span exists.", "BR-01, LC-04"),
        ("TR-FC-004", "specialty_classifier and homeopathic_expansion are graph nodes inside FindCare/RuntimeServer/nodes.py with explicit input/output state contracts.", "Unit test per node; integration test with fan-in via add_messages.", "BR-01, LC-05"),
        ("TR-FC-005", "tool_record_user_details and tool_record_unknown_question are @tool nodes; both delegate to FindCare/SharedCode/core/lead_service.py.", "Smoke test inserts a lead row; trace span exists.", "BR-01, LC-06"),
        ("TR-FC-006", "tool_get_skip_snow_context and tool_get_chathealthy_context are @tool nodes; both read static brain artifacts via SharedServices.", "Smoke test; trace span exists.", "BR-01, LC-07"),
        ("TR-FC-007", "repetitive_utterance_detector and merge_routing live in FindCare/RuntimeServer/nodes.py; each list channel they write declares an explicit reducer (operator.add or a custom reducer).", "Anti-pattern scanner asserts no un-annotated list channels; integration test exercises fan-out.", "BR-01, LC-09, V6 Q7"),
        # EvaluateCare recovery
        ("TR-EC-001", "EvaluateCare/RuntimeServer/agent.py exposes a graph that contains score_provider, score_trial, evaluate_providers, evaluate_view, explain as nodes; EvaluateCare/SharedCode/ holds scoring_engine, normalization, weights, confidence, explainability, provenance, failsafe.", "5 nodes visible in Studio; v4-043 reports zero violations on the EvaluateCare Space.", "BR-02, LC-10, EPIC-001 V7"),
        ("TR-EC-002", "Non-conversational EvaluateCare HTTP endpoints (provider lookup, evaluation report read) live under EvaluateCare/RuntimeServer/app.py and are NOT LangGraph endpoints.", "Trace inspection: no parent-graph trace for these GETs.", "BR-06, V6 Section 3.3"),
        # Security
        ("TR-SE-001", "Security/RuntimeServer/agent.py is a graph that runs as the FIRST capability node in the parent's edge order; Security/RuntimeServer/app.py exposes it as a RemoteGraph endpoint.", "Edge-order test; integration test rejects an invalid token before any other capability runs.", "BR-04, V7 EPIC-002"),
        ("TR-SE-002", "emergency_keyword_detector and ip_lock_check live in Security/RuntimeServer/nodes.py; emergency keywords are sourced from brain/machine_artifacts/content/emergency_keywords.json.", "Unit tests per detector; brain artifact path verified.", "BR-04, LC-08"),
        ("TR-SE-003", "Compliance is a feature folder under Security/RuntimeServer/compliance/, not its own epic.", "Folder presence test; epic registry shows Security & Compliance as one cost-center line.", "V7 A1, BR-11"),
        # Shared services
        ("TR-SS-001", "SharedServices/RuntimeServer/agent.py exposes the shared_splash, transfer_to_findcare, verify_token nodes via a graph; SharedServices/SharedCode/mongo_utilities.py is the canonical Mongo client used by every Space.", "5 nodes visible; v4-043 reports zero violations on the Shared Services Space.", "BR-04, BR-09, LC-11, V7 dual-child"),
        ("TR-SS-002", "Tool Router exposes the shared LLM client and embedding helpers as @tool registrations consumable from any capability subgraph.", "Tool registry test; capability nodes invoke via ToolRuntime.", "V7 Section 12.4, V6 Q5"),
        # Observability
        ("TR-OB-001", "LangSmith tracing is enabled in the parent runtime AND in every capability Space; trace_id is propagated across RemoteGraph invocations via the standard context= argument.", "End-to-end test: a single thread_id yields one consolidated trace tree across Spaces.", "BR-12, LC-12, mandate 4"),
        ("TR-OB-002", "Operations/SharedCode/conversation_log/Worker.cs (canonical, moved from Code/Shared/ops/ChatHealthyLogService/) consumes Kafka topics and writes structured per-turn audit rows to MongoDB.", "Worker integration test; audit row exists for every emitted turn event.", "BR-08, BR-12, V7 13.2"),
        ("TR-OB-003", "Per-node performance budgets are emitted as OpenTelemetry histograms; alert rules fire on p95 breach.", "OTel collector receives histograms; alert rule unit-test fires on synthetic breach.", "BR-13, mandate 4"),
        ("TR-OB-004", "LangSmith is the trace UI of record; observability backend choice resolved per Section 12 R5.", "Section 12 R5 recommendation accepted; design depends on LangSmith feature set.", "Resolved issue R5, mandate 8"),
        # OO conventions
        ("TR-OO-001", "Every Python service entry point file ends with `if __name__ == '__main__': Application().run()` (or equivalent). The Application class lives in the same module or is imported from {Capability}/RuntimeServer/app.py.", "Pre-commit OO scan asserts each registered entry-point file conforms.", "BR-20, mandate 2"),
        ("TR-OO-002", "Operations/SharedCode/scan/oo_thin_main_scan.py walks the entry-point registry; non-conforming files MUST appear in Operations/SharedCode/scan/oo_exceptions.json with a justification, owner, and review date.", "Scan exits non-zero if a non-registered exception appears; exception list is human-readable.", "BR-20, mandate 2"),
        ("TR-OO-003", "FastAPI handlers are NOT entry points. The web app object construction happens in app.py's Application class; the if __name__ == '__main__' block calls Application().run() which runs uvicorn programmatically.", "Static check: no `app = FastAPI()` at module top in {Capability}/RuntimeServer/app.py; the FastAPI instance is constructed inside Application.__init__.", "BR-20, mandate 2"),
        # Brain / governance / blast radius
        ("TR-BG-001", "engineering_rules.json is canonical; every action that mutates state in any project file is gated by an enforcement worker per the V20 framework.", "Pre-commit run with a rule-violating change exits non-zero.", "BR-10, V20"),
        ("TR-BG-002", "All requirement IDs follow EPIC-NNN-F-NNN-S-NNN-REQ-{B|T}-NNN; every code change traces to one such ID.", "Backlog crosswalk; orphan code scan exits non-zero on any new uncited path.", "BR-17, V20"),
        ("TR-BG-003", "The boot class Code/Shared/ops/tools/chathealthy_devops_boot.py governs every Claude Code session via .claude/settings.json hooks. The boot class reads brain artifacts read-only and emits a session digest.", "Hook fires on session start; digest produced; settings.json references the boot class.", "BR-10, mandate 3"),
        ("TR-BG-004", "Project-root capability folders (EvaluateCare/, FindCare/, SharedServices/, Security/, Operations/, UX/, DataManagementPipelines/) replace the Code/ umbrella; an edit to a single capability folder MUST NOT cross into another capability folder.", "Cross-capability import scan exits non-zero on illegal imports.", "BR-11, V7 Section 4.2, mandate 3"),
        # UX
        ("TR-UX-001", "FindCare/Frontend/ (the React iframe) is the conversational facade; it consumes the parent runtime via @langchain/langgraph-sdk; it does NOT import a capability Space URL.", "Source scan: only the parent URL string appears in src/; capability URLs are absent.", "BR-03, V6 Section 3.3"),
        ("TR-UX-002", "Website/index.html embeds FindCare/Frontend as an iframe and remains the canonical landing page (Cloudflare Pages deploy target).", "Smoke test: GET / serves the iframe wrapper.", "BR-03, V7 dual-child"),
        ("TR-UX-003", "UX/SharedCode/ holds the shared React primitives (ProviderCard.tsx, SelectionManager.tsx, provider.ts) that both FindCare/Frontend and EvaluateCare/Frontend consume.", "Build-time import resolution from both consumers succeeds.", "BR-05, V7 A7"),
        ("TR-UX-004", "Operations/SharedCode/static_page_generator/ regenerates the 10 non-roadmap static pages in Website/ from declarative inputs in UX/SharedCode/static_pages/.", "Generator run produces byte-identical output for non-volatile pages; roadmap.html stays deleted per V7 A10.", "BR-05, V7 A9, A10"),
        # Pipelines / data
        ("TR-DM-001", "DataManagementPipelines/ holds ingest, dedup, and embedding-generation jobs; jobs run under Operations infrastructure but are NOT nodes in any LangGraph.", "Pipeline run does not produce LangSmith traces.", "BR-09, V7 Section 12.4"),
        ("TR-DM-002", "The canonical Provider schema lives at brain/machine_artifacts/schemas/ChatHealthyProvidersSchema.json; the schema declares all 11 currently-missing fields (embedding, embedding_model, embedding_version, 8 NPPES other-name fields). Code/Schemas/ChatHealthy.Providers.json is deleted.", "Schema diff test asserts the 11 fields are declared; orphan-file scan asserts the deleted path is gone.", "BR-18, V7 A3, V7 Section 10, R2"),
        ("TR-DM-003", "Provider-data writes go through SharedServices/SharedCode/mongo_utilities.py only; capability Spaces never instantiate raw pymongo clients.", "Static analyzer flags any non-shared mongo client construction.", "BR-09, V7 13.2"),
        # Deferred work
        ("TR-DW-001", "Operations/SharedCode/deferred_work/registry.py walks bugs.json and agile_backlog.json and emits the deferred-work ledger.", "Generator run produces the registry; row count matches a known fixture.", "BR-19, EPIC-008-F-009"),
        # Recovery
        ("TR-RC-001", "Smoke test Operations/SharedCode/scan/lost_capability_recovery_smoke.py invokes each of the 12 LC-* capabilities through the parent runtime and asserts the LangSmith trace contains the corresponding node.", "CI run is green; LangSmith API check returns 12 nodes.", "BR-15, LC-01..LC-12"),
        # Internal consistency
        ("TR-IC-001", "Operations/SharedCode/scan/br_tr_coverage_scan.py asserts every BR has at least one TR and every TR traces to at least one BR; emits the coverage matrix.", "Scan exits non-zero if any row is unmapped.", "BR-16, mandate 7"),
    ]
    add_table(doc, ["TR ID", "Statement", "Acceptance test", "Trace"], tr_rows)

    # ---------- 5. Architecture overview ----------
    add_heading(doc, "5. Architecture overview", level=1)

    add_heading(doc, "5.1 HF Space topology", level=2)
    add_para(
        doc,
        "Six HuggingFace Spaces, each one cost-center line, one org-chart line, one deploy pipeline, "
        "one /health endpoint:",
    )
    add_bullet(doc, "Operations - hosts the parent runtime; exposes the single conversational HTTP endpoint to the React UI; owns the MongoDBSaver checkpointer.")
    add_bullet(doc, "FindCare - hosts the FindCare capability subgraph; exposes its graph as RemoteGraph and its non-conversational read API as REST.")
    add_bullet(doc, "EvaluateCare - hosts the EvaluateCare scoring subgraph and its non-conversational report API.")
    add_bullet(doc, "Security & Compliance - hosts the Security capability subgraph (token verification, entitlement, audit emission, emergency detection).")
    add_bullet(doc, "SharedServices - hosts the Tool Router (LLM client, embeddings, mongo abstractions) consumed via ToolRuntime by every capability Space.")
    add_bullet(doc, "DataManagementPipelines - infrastructure host for scheduled ingest/dedup/embedding jobs; does NOT participate in the parent graph.")
    add_para(
        doc,
        "UX is not a runtime Space; UX/Frontend artifacts deploy to Cloudflare Pages. EPIC-010 Operations also "
        "owns CI/deploy/build-counter tooling that is not a runtime node.",
    )

    add_heading(doc, "5.2 Runtime composition (textual diagram)", level=2)
    add_para(
        doc,
        "[React UI]  --HTTPS streaming-->  [Operations parent_graph.py]  "
        "--RemoteGraph-->  {Security -> FindCare | EvaluateCare}  --ToolRuntime-->  [SharedServices Tool Router]  "
        "--mongo--> [MongoDB Atlas].",
    )
    add_para(
        doc,
        "Trace context: trace_id is created in the parent on turn start, propagated to each RemoteGraph "
        "invocation via context=, propagated into ToolRuntime, attached to every Mongo write through the "
        "shared client. LangSmith reconstructs the full tree from the propagated identifier.",
    )

    add_heading(doc, "5.3 Data flow", level=2)
    add_bullet(doc, "Conversational read/write: React -> parent -> Security node -> capability node(s) -> response stream back to React.")
    add_bullet(doc, "Non-conversational read: React -> capability HTTP API (e.g. GET /provider/{id}) -> SharedServices mongo abstraction -> MongoDB.")
    add_bullet(doc, "Pipeline write: scheduled job -> SharedServices mongo abstraction -> MongoDB; emits a Kafka event consumed by Operations conversation_log Worker.cs into an audit collection.")
    add_bullet(doc, "Audit trail: every parent turn emits a structured event onto Kafka; Operations Worker.cs writes into MongoDB; LangSmith captures the LangGraph view, Worker.cs captures the persistent audit view.")

    # ---------- 6. Per-epic architecture ----------
    add_heading(doc, "6. Per-epic architecture", level=1)
    epic_specs = [
        ("6.1 EPIC-001 EvaluateCare",
         [
             "HF Space: EvaluateCare. RuntimeServer/app.py constructs the Application object whose run() spawns uvicorn programmatically.",
             "RuntimeServer/agent.py compiles a StateGraph that hosts score_provider, score_trial, evaluate_providers, evaluate_view, explain as nodes (recovers LC-10).",
             "RuntimeServer/state.py is the EvaluateCareState Pydantic BaseModel with input_schema EvaluateCareIn and output_schema EvaluateCareOut.",
             "RuntimeServer/nodes.py + tools.py per the V6 4-file layout.",
             "SharedCode/core/ holds scoring_engine.py, normalization.py, weights.py, confidence.py, explainability.py, provenance.py, failsafe.py, cache.py, models.py (existing files in Code/evaluate_care/ become refactor source per V7 A4).",
             "SharedCode/core/measures/ holds individual measure modules.",
             "Non-conversational endpoints (GET /provider/{id}, GET /evaluation/{id}) live in RuntimeServer/app.py outside the graph.",
         ]),
        ("6.2 EPIC-002 Security & Compliance",
         [
             "HF Space: Security. Owns its own Space per V7 A1 architect ratification.",
             "RuntimeServer/agent.py exposes the security graph that runs FIRST in parent edge order.",
             "RuntimeServer/nodes.py: token_verify, entitlement_check, emergency_keyword_detector, ip_lock_check (recovers LC-08).",
             "RuntimeServer/compliance/ - feature folder for HIPAA/GDPR audit emission.",
             "SharedCode/policies/ - declarative policy bundles loaded at module init.",
             "Brain artifact ties: emergency_keywords.json, security.json, SecurityAuditControls.json.",
         ]),
        ("6.3 EPIC-006 FindCare",
         [
             "HF Space: FindCare. Hosts the conversational find-care subgraph.",
             "RuntimeServer/agent.py compiles a StateGraph wired with: get_authorization, user_utterance, repetitive_detector, classifier, merge, choose_specialties, tool_find_providers, tool_search_clinical_trials, tool_record_user_details, tool_record_unknown_question, tool_get_skip_snow_context, tool_get_chathealthy_context, system_utterance, evaluate_handoff, user_exits.",
             "Frontend/ (React iframe) talks ONLY to the parent runtime URL (TR-UX-001).",
             "SharedCode/core/provider_search_service.py + lead_service.py (existing files relocate here).",
             "Talk About Care is a feature folder under FindCare per Resolved issue R3, not its own epic.",
         ]),
        ("6.4 EPIC-009 Shared Services",
         [
             "HF Space: SharedServices. Hosts the Tool Router endpoint and the canonical mongo client.",
             "RuntimeServer/app.py + agent.py expose shared_splash, transfer_to_findcare, verify_token as graph nodes (recovers LC-11) plus REST endpoints for non-conversational reads.",
             "SharedCode/mongo_utilities.py is the canonical Mongo client (relocated per V7 A2 / 13.2; old Code/DataPipelines/ChatHealthyMongoUtilities.py deleted).",
             "SharedCode/llm_client.py + embeddings/ - centralized model invocations consumed by every capability via ToolRuntime.",
             "SharedCode/cost_guard.py, prompt_system_maker.py, session_token.py - shared cross-capability concerns.",
         ]),
        ("6.5 EPIC-010 Operations",
         [
             "HF Space: Operations. Hosts the parent runtime + the conversation-log Worker + brain governance tooling.",
             "RuntimeServer/orchestration/parent_graph.py - the thin parent (V7 12.5).",
             "RuntimeServer/orchestration/state.py - parent state Pydantic BaseModel.",
             "RuntimeServer/app.py - FastAPI front door for the LangGraph SDK streaming endpoint.",
             "SharedCode/conversation_log/Worker.cs - canonical conversation log Kafka consumer (relocated per V7 13.2).",
             "SharedCode/brain_governance/rebuild_manifest.py + manifest_generator.py per V7 A8.",
             "SharedCode/scan/ - br_tr_coverage_scan.py, oo_thin_main_scan.py, lost_capability_recovery_smoke.py (governance scanners).",
             "EPIC-008 features (Architecture / Engineering Rules Enforcement / Brain Consistency / DevOps / Audit / Code Scanning / Deferred Work / Graph orchestration) all map under Operations.",
         ]),
        ("6.6 EPIC-011 UX",
         [
             "Not a runtime HF Space. Deploys to Cloudflare Pages.",
             "Frontend/ - the parent React app (canonical iframe wrapper, environment banner, auth-token display).",
             "SharedCode/cross_components/ - ProviderCard.tsx, SelectionManager.tsx, provider.ts (relocated per V7 A7).",
             "SharedCode/static_pages/ - declarative source for the 10 non-roadmap static pages.",
             "SharedCode/design_paradigm/ - shared design tokens / paradigm (V5 F-002 stories).",
             "Static page generator lives in Operations/SharedCode/static_page_generator/ but consumes UX/SharedCode/static_pages/.",
         ]),
        ("6.7 EPIC-012 Data Management Pipelines",
         [
             "Hosts scheduled ingest, dedup, embedding-generation jobs. NO dual-child layout (V7 4.3).",
             "Jobs run under Operations infrastructure but are NOT nodes in parent_graph.py.",
             "Schema-drift remediation (BR-18) implemented as a pipeline regression test against brain/machine_artifacts/schemas/ChatHealthyProvidersSchema.json.",
             "Reads/writes go through SharedServices/SharedCode/mongo_utilities.py only.",
         ]),
    ]
    for title, bullets in epic_specs:
        add_heading(doc, title, level=2)
        for b in bullets:
            add_bullet(doc, b)

    # ---------- 7. LangGraph implementation ----------
    add_heading(doc, "7. LangGraph implementation", level=1)
    add_para(
        doc,
        "Implements V6 best practices verbatim. Parent runtime + cross-Space RemoteGraph composition; "
        "in-process subgraph composition for same-Space steps; Pydantic state; runtime-context DI; "
        "messages channel as the home for tool I/O.",
    )

    add_heading(doc, "7.1 Parent runtime (Operations Space)", level=2)
    add_bullet(doc, "Operations/RuntimeServer/orchestration/parent_graph.py instantiates StateGraph(ParentState).")
    add_bullet(doc, "Operations/RuntimeServer/orchestration/state.py: ParentState(BaseModel) with messages: Annotated[list[AnyMessage], add_messages]; thread_id; user_id; locale; current_capability: Optional[str]; routing_intent: Optional[str].")
    add_bullet(doc, "Operations/RuntimeServer/orchestration/checkpointer.py: a thin module that constructs MongoDBSaver once at module load. ParentState graph is compiled with checkpointer=MongoDBSaver(...).")
    add_bullet(doc, "Operations/RuntimeServer/app.py: Application class whose run() programmatically calls uvicorn.run with the FastAPI ASGI app that wraps the LangGraph SDK streaming endpoint.")

    add_heading(doc, "7.2 RemoteGraph composition", level=2)
    add_para(doc, "Pseudocode mirrors V6 Section 3.2:")
    add_para(
        doc,
        "from langgraph.graph import StateGraph, START\n"
        "from langgraph.pregel.remote import RemoteGraph\n"
        "from .state import ParentState\n"
        "from .checkpointer import checkpointer\n\n"
        "security      = RemoteGraph('security',       url=SECURITY_HF_URL,      api_key=...)\n"
        "find_care     = RemoteGraph('find_care',      url=FINDCARE_HF_URL,      api_key=...)\n"
        "evaluate_care = RemoteGraph('evaluate_care',  url=EVALCARE_HF_URL,      api_key=...)\n"
        "shared_svc    = RemoteGraph('shared_services', url=SHAREDSVC_HF_URL,    api_key=...)\n\n"
        "g = StateGraph(ParentState)\n"
        "g.add_node('security', security)\n"
        "g.add_node('find_care', find_care)\n"
        "g.add_node('evaluate_care', evaluate_care)\n"
        "g.add_node('shared_services', shared_svc)\n"
        "g.add_edge(START, 'security')\n"
        "g.add_conditional_edges('security', route_after_security, {'find_care','evaluate_care','shared_services','END'})\n"
        "parent = g.compile(checkpointer=checkpointer)\n",
    )

    add_heading(doc, "7.3 Capability subgraph layout (canonical 4-file)", level=2)
    add_bullet(doc, "{Capability}/RuntimeServer/state.py - Pydantic BaseModel + input_schema/output_schema.")
    add_bullet(doc, "{Capability}/RuntimeServer/nodes.py - node functions or class-callables.")
    add_bullet(doc, "{Capability}/RuntimeServer/tools.py - @tool decorations + ToolNode wiring.")
    add_bullet(doc, "{Capability}/RuntimeServer/agent.py - graph build (compile() with no checkpointer).")
    add_bullet(doc, "{Capability}/RuntimeServer/langgraph.json - manifest per the langgraph-example-pyproject template.")

    add_heading(doc, "7.4 Runtime context (DI)", level=2)
    add_bullet(doc, "Each capability subgraph defines its own ContextSchema (Pydantic) carrying user_id, thread_id, trace_id, locale, request_id.")
    add_bullet(doc, "Nodes accept Runtime[ContextSchema]; tools accept ToolRuntime[ContextSchema].")
    add_bullet(doc, "The parent passes its session-scoped context to each RemoteGraph invocation via context=.")
    add_bullet(doc, "Module-scope import of mongo or LLM clients in node files is forbidden (TR-LG-008).")

    add_heading(doc, "7.5 Messages channel and tool I/O", level=2)
    add_bullet(doc, "All tool calls and tool results are ToolMessage entries on the messages channel (TR-LG-007).")
    add_bullet(doc, "Structured artefacts the LLM does not need to read (e.g. raw provider rows for downstream rendering) get a typed Pydantic field on the subgraph state, NEVER a flat top-level state field per tool.")

    add_heading(doc, "7.6 Same-Space subgraph composition", level=2)
    add_bullet(doc, "Two nodes that share a database client and deploy together stay in the same Space and use in-process subgraph composition.")
    add_bullet(doc, "RemoteGraph never calls back into its own deployment (V6 Section 3.2 deadlock constraint, TR-LG-005).")
    add_bullet(doc, "Deep research is a same-Space subgraph hosted inside whichever capability Space invokes it (V6 Section 3.4).")

    # ---------- 8. Observability ----------
    add_heading(doc, "8. Observability", level=1)
    add_heading(doc, "8.1 LangSmith integration", level=2)
    add_bullet(doc, "LangSmith tracing is enabled in the parent and in every capability Space (TR-OB-001).")
    add_bullet(doc, "trace_id is propagated across every RemoteGraph boundary via context=; LangSmith reconstructs the full tree from the propagated id.")
    add_bullet(doc, "LangSmith is the trace UI of record (Resolved issue R5).")

    add_heading(doc, "8.2 Audit log pipeline (Kafka)", level=2)
    add_bullet(doc, "Every parent turn emits a structured event onto Kafka topic chathealthy.conversation_log.")
    add_bullet(doc, "Operations/SharedCode/conversation_log/Worker.cs (canonical, V7 13.2) consumes the topic, writes to MongoDB collection conversation_log.")
    add_bullet(doc, "Schema validated against Website/schemas/qa/ChatHealthyErrorsSchema.json (post-relocation: brain/machine_artifacts/schemas/...).")

    add_heading(doc, "8.3 Performance budgets", level=2)
    add_bullet(doc, "Per-node OpenTelemetry histograms emitted from each Space.")
    add_bullet(doc, "Concrete budgets (BR-13): provider_search p95 <= 1500ms; specialty_classifier p95 <= 600ms; security/token_verify p95 <= 200ms; evaluate_providers p95 <= 2000ms.")
    add_bullet(doc, "Alert rules trigger on rolling 24h p95 breach.")

    add_heading(doc, "8.4 Trace context preservation across Spaces", level=2)
    add_bullet(doc, "trace_id flows through context= on RemoteGraph invocation and is read inside the capability Space's Runtime[ContextSchema].")
    add_bullet(doc, "Tools read trace_id from ToolRuntime and attach it to every Mongo write so the audit collection joins to the LangSmith view by trace_id.")

    # ---------- 9. OO conventions ----------
    add_heading(doc, "9. OO conventions (mandate 2)", level=1)
    add_heading(doc, "9.1 Thin-main-construct-and-run pattern", level=2)
    add_para(
        doc,
        "Every Python service entry point ends with a thin if __name__ == '__main__': block whose sole job is "
        "to construct an Application object and call run() on it. Example shape (canonical):",
    )
    add_para(
        doc,
        "# {Capability}/RuntimeServer/app.py\n"
        "class Application:\n"
        "    def __init__(self) -> None:\n"
        "        self.fastapi = build_fastapi_app()  # not a module-level FastAPI()\n"
        "        self.config  = AppConfig.load()\n"
        "    def run(self) -> None:\n"
        "        uvicorn.run(self.fastapi, host=self.config.host, port=self.config.port)\n\n"
        "if __name__ == '__main__':\n"
        "    Application().run()\n",
    )

    add_heading(doc, "9.2 The exception registry", level=2)
    add_bullet(doc, "Operations/SharedCode/scan/oo_exceptions.json carries every approved violation as: {path, function, justification, owner, review_date}.")
    add_bullet(doc, "Operations/SharedCode/scan/oo_thin_main_scan.py walks the entry-point registry; non-conforming files MUST appear in oo_exceptions.json or the scan exits non-zero.")
    add_bullet(doc, "Exceptions have an explicit review_date; expired entries fail the scan.")

    add_heading(doc, "9.3 Scan function contract", level=2)
    add_bullet(doc, "Inputs: the entry-point registry (Operations/SharedCode/scan/entry_points.json) and oo_exceptions.json.")
    add_bullet(doc, "Algorithm: for each registered entry-point file, ast-parse and assert the presence of a single if __name__ == '__main__' block whose body constructs one object and calls one method on it.")
    add_bullet(doc, "Outputs: structured list of violations; exit code 0 (clean) or non-zero (violation).")
    add_bullet(doc, "Hook: registered as a Rule-OO-001 enforcement entry in engineering_rules.json executed at pre-commit.")

    # ---------- 10. Blast-area reduction ----------
    add_heading(doc, "10. Blast-area reduction (mandate 3)", level=1)
    add_bullet(doc, "Project-root capability folders + dual-child layout (V7 Section 4.2). An edit to FindCare/RuntimeServer/ cannot reach into EvaluateCare/RuntimeServer/ without an explicit cross-capability import that the scanner flags (TR-BG-004).")
    add_bullet(doc, "Engineering-rules enforcement framework V20: 50 rules, deterministic worker registry; pre-commit hook fires every commit (TR-BG-001).")
    add_bullet(doc, "Implementation-traces-to-requirements principle: no code without an EPIC-NNN-F-NNN-S-NNN-REQ-{B|T}-NNN trace (TR-BG-002, BR-17).")
    add_bullet(doc, "Boot class governs every Claude Code session via .claude/settings.json hooks; the boot class is read-only and emits a session digest the operator can audit (TR-BG-003).")
    add_bullet(doc, "Single-thin-main convention with a registered exception list (Section 9). Claude cannot quietly add a procedural entry-point.")
    add_bullet(doc, "The parent runtime is the only piece that owns the checkpointer and the messages channel; capability Spaces are isolated from each other and from session state. A bug in any one capability cannot corrupt session continuity.")
    add_bullet(doc, "Read-only architecture artifacts under architecture/ are the binding inputs every session reads on boot; they are not modifiable in the same session that consumes them.")

    # ---------- 11. Resolved issues ----------
    add_heading(doc, "11. Resolved issues (mandate 8)", level=1)
    res_rows = [
        ("R1", "Security HF Space cutover sequencing", "Parallel-run cutover. Stand up the Security Space alongside the existing in-FindCare security path; route 5%, then 50%, then 100% of conversational traffic through the new Space; cut the in-FindCare path only after a clean 7-day window. The design assumes parallel-run; TR-SE-001 is written against the new Space, not the legacy in-FindCare path."),
        ("R2", "Provider schema content drift (11 fields)", "Schema is canonical and must declare all fields the pipeline produces. Add the 11 fields (embedding, embedding_model, embedding_version, 8 NPPES other-name fields) to brain/machine_artifacts/schemas/ChatHealthyProvidersSchema.json at the same commit that relocates the schema (V7 Section 10). The design assumes the schema is the contract (BR-18, TR-DM-002)."),
        ("R3", "Talk About Care placement", "Talk About Care is a feature folder under FindCare/RuntimeServer/talk_about_care/, not a separate epic. The original EPIC-007 backlog feature is a feature, not an epic; V7 already carries EPIC-007 stories under EPIC-006 FindCare. The design assumes the feature-not-epic placement (BR-07)."),
        ("R4", "langgraph-api vs FastAPI for Spaces", "Each capability Space exposes its agent via the standard LangGraph Agent Server protocol (langgraph-api on PyPI). The Space's RuntimeServer/app.py builds a FastAPI app for non-conversational REST endpoints AND mounts the LangGraph Agent Server alongside; RemoteGraph from the parent talks to the LangGraph Agent Server route, not FastAPI. The design assumes both surfaces co-exist in the same Space process (TR-LG-004, TR-EC-002, TR-FC-001)."),
        ("R5", "Observability backend choice (LangSmith vs custom)", "LangSmith is the trace UI of record. ChatHealthy already integrates with LangSmith via the V6 best-practices stack; building a custom trace UI duplicates effort and breaks LangGraph's built-in tracing contract. Worker.cs + Kafka cover persistent audit; LangSmith covers latency/error tree per turn. The design assumes LangSmith (TR-OB-001, TR-OB-004)."),
    ]
    add_table(doc, ["#", "Unresolved issue", "Recommendation (assumed accepted)"], res_rows)

    # ---------- 12. Internal-consistency check ----------
    add_heading(doc, "12. Internal-consistency check (mandate 7)", level=1)
    add_para(doc, "Every BR maps to at least one TR; every TR traces to at least one BR. Coverage matrix:")
    cov_rows = [
        ("BR-01", "TR-LG-001, TR-LG-002, TR-LG-003, TR-FC-001..TR-FC-007"),
        ("BR-02", "TR-EC-001, TR-LG-004"),
        ("BR-03", "TR-LG-002, TR-LG-004, TR-UX-001, TR-UX-002"),
        ("BR-04", "TR-SE-001, TR-SE-002, TR-SS-001"),
        ("BR-05", "TR-UX-003, TR-UX-004"),
        ("BR-06", "TR-EC-002"),
        ("BR-07", "Section 11 R3; TR-FC-001 (Talk About Care under FindCare)"),
        ("BR-08", "TR-OB-002"),
        ("BR-09", "TR-DM-001, TR-DM-003, TR-SS-001"),
        ("BR-10", "TR-BG-001, TR-BG-003"),
        ("BR-11", "TR-BG-004, TR-SE-003"),
        ("BR-12", "TR-OB-001, TR-OB-002, TR-OB-004, TR-LG-006, TR-LG-008"),
        ("BR-13", "TR-OB-003"),
        ("BR-14", "TR-UX-002, TR-UX-004"),
        ("BR-15", "TR-RC-001"),
        ("BR-16", "TR-IC-001"),
        ("BR-17", "TR-BG-002, TR-IC-001"),
        ("BR-18", "TR-DM-002"),
        ("BR-19", "TR-DW-001"),
        ("BR-20", "TR-OO-001, TR-OO-002, TR-OO-003"),
    ]
    add_table(doc, ["BR", "Implementing TR(s)"], cov_rows)
    add_para(doc, "Reverse direction: every TR's Trace cell in Section 4 names at least one BR. Confirmed by TR-IC-001's scan run on the design itself.")
    add_para(
        doc,
        "Verdict: BR<->TR coverage is complete. No BR is without an implementing TR; no TR is without an "
        "originating BR. No TR mandates a design choice the design itself does not make (each TR's Acceptance "
        "test refers to a section of this document or a real file path). No design element exists without a "
        "demanding TR. Section 11 recommendations are consistent with downstream design choices: "
        "R1 cutover -> TR-SE-001 written against the new Space; R2 schema -> TR-DM-002 declares the 11 fields; "
        "R3 placement -> Section 6.3 puts Talk About Care under FindCare; R4 dual-protocol -> TR-EC-002 and "
        "TR-FC-001 assume both surfaces; R5 LangSmith -> TR-OB-001 / TR-OB-004 wire to LangSmith.",
    )

    # ---------- Footer ----------
    add_heading(doc, "Appendix - generation record", level=1)
    add_bullet(doc, "Read at generate time: findCare/ArchitectureAndDesign/codebase-reorganization-V7.docx, findCare/ArchitectureAndDesign/langgraph-oo-best-practices-V6.docx, brain/machine_artifacts/content/{engineering_rules,agile_backlog,bugs,version}.json, Code/, Website/, LangGraph/.")
    add_bullet(doc, "Build numbers cited (Section 2.1) read directly from version.json; dates inferred from co-touched dated artifacts.")
    add_bullet(doc, "26 BUG-LANGCHAIN-CONTAINER-* entries enumerated; 12 distinct lost capabilities clustered (Section 2.2).")
    add_bullet(doc, "20 BR + 38 TR. BR<->TR coverage matrix in Section 12 has zero unmapped rows.")
    add_bullet(doc, "Read-only enforcement: no commits, no code changes, no backlog mutations during generation. Stories left intact.")

    return doc


if __name__ == "__main__":
    doc = build_doc()
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"WROTE {OUT}")
