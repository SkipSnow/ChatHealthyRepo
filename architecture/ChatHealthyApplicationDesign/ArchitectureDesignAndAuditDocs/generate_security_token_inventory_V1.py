"""Generate Security / Session Token Inventory (V1).

Substance: produced by an agent on 2026-05-08 mapping every reference
to four concepts (security/session token, GUID, display panels, NONCE)
across code, wrapper, library, agile_backlog, bugs.json, and design
docs. Intended as the seed reference for the EPIC-002 (Security And
Shared Services) Authentication feature work scheduled for 2026-05-09.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt

HERE = Path(__file__).resolve().parent
OUT = HERE / "SecurityTokenInventory_V1.docx"


def H(doc, text, level=1):
    return doc.add_heading(text, level=level)


def P(doc, text):
    return doc.add_paragraph(text)


def BULLET(doc, text):
    return doc.add_paragraph(text, style="List Bullet")


def CODE(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    return p


# ── Inventory data — one tuple per concept section ─────────────────────────
# Each section is (title, [ (subsection_label, [bullet_lines]) ]).

CONCEPT_1_TOKEN = ("1. Security / Session Token", [
    ("Backend (Python) references", [
        "Code\\ConversationalUX\\FindCareChat\\backend\\app.py:247 — comment \"/session depends on session_token.generate_session_token producing a\"",
        "Code\\ConversationalUX\\FindCareChat\\backend\\app.py:270 — \"FINDCARE_SIGNING_KEY_PEM\": \"findcare.key\" cert bootstrap mapping",
        "Code\\ConversationalUX\\FindCareChat\\backend\\app.py:306-328 — startup security-primitive verification block; imports generate_session_token, calls it, exits 78/77/70 on failure",
        "Code\\ConversationalUX\\FindCareChat\\backend\\app.py:594-609 — @app.post(\"/session\") mints FindCare-signed token via generate_session_token(\"FindCare\")",
        "Code\\ConversationalUX\\FindCareChat\\backend\\app.py:616-665 — VerifyTokenRequest model + @app.post(\"/verify-token\") calls verify_session_token, returns audit response",
        "Code\\ConversationalUX\\FindCareChat\\backend\\tests\\test_https_security.py:75-133 — HTTPS smoke for /session and providers calls with session_token",
        "Code\\ConversationalUX\\FindCareChat\\backend\\tests\\local_environment_user_ready.py:199-278 — test_session_token_in_left_panel, test_session_token_visible",
        "Code\\ConversationalUX\\FindCareChat\\backend\\tests\\dev_environment_user_ready.py:247-255 — test_session_token_visible",
        "sharedServices\\Code\\app.py:105-133 — imports SessionToken, mounts @app.post(\"/session\") and @app.post(\"/verify-token\")",
        "sharedServices\\Code\\security\\session_endpoint.py:7-28 — generate_session_token(\"SharedServices\")",
        "sharedServices\\Code\\security\\verify_token_endpoint.py:9-65 — verify_session_token(inbound, \"FindCare\")",
        "sharedServices\\Code\\Dockerfile:12 — comment \"/session, /verify-token, session_token\"",
        "evaluateCare\\Code\\app.py:111-139 — same pattern as SharedServices",
        "evaluateCare\\Code\\security\\session_endpoint.py:7-28 — generate_session_token(\"EvaluateCare\")",
        "evaluateCare\\Code\\security\\verify_token_endpoint.py:9-65 — verify endpoint",
        "evaluateCare\\Code\\security\\debug_verify_live_endpoint.py:7-38 — /debug/verify-live runs verify_session_token",
        "evaluateCare\\Code\\externalInterface\\evaluate_providers_endpoint.py:8-102 — session_token: SessionToken = Field(..., description=\"FindCare-signed session token for mutual authentication\")",
        "evaluateCare\\Code\\Dockerfile:12 — comment \"/session, /verify-token, /debug/*, session_token\"",
        "Code\\deploy\\localSmokeTestPyTest.py:140 — c.post(f\"{FINDCARE_URL}/session\").json()",
        "Code\\deploy\\localSmokeTestPyTest.py:194-204 — REQ-B-010 right-panel \"Server, serving security token:\" assertion",
        "Code\\deploy\\localSmokeTestPyTest.py:955-983 — REQ-019/020 enforce direct /verify-token (no proxy) + right-panel string check",
        "Code\\deploy\\localSmokeTestPyTest.py:1040,1057 — pytest.skip(\"S-002-REQ-B-002: security token panel suppressed in prod\")",
        "Code\\deploy\\devSmokeTestPyTest.py:140,194,204,773-801 — mirror of local smoke test",
        ".github\\workflows\\deploy-findcare-backend.yml:76-167 — FINDCARE_SIGNING_KEY_B64 → FINDCARE_SIGNING_KEY_PEM HF Space secret wiring",
    ]),
    ("Frontend (TypeScript/React) references", [
        "Code\\ConversationalUX\\FindCareChat\\frontend\\src\\components\\ChatWindow.tsx:33-38 — comment + getSessionToken(apiUrl) POSTs /session",
        "Code\\ConversationalUX\\FindCareChat\\frontend\\src\\components\\ChatWindow.tsx:261,270,276,300 — fetches token, sends as session_token: in payloads",
        "Code\\ConversationalUX\\FindCareChat\\frontend\\src\\components\\FindCareApp.tsx:58-60 — getSessionToken() POSTs ${API_URL}/session",
        "Code\\ConversationalUX\\FindCareChat\\frontend\\src\\components\\FindCareApp.tsx:363-386 — uses token, propagates session_token: data.session_token to parent",
    ]),
    ("Wrapper (Website/index.html) references", [
        "Website\\index.html:880-933 — wrapper-level cold-button flow: fetches /session, POSTs to cold service /verify-token, asserts result.session_token shape",
        "Website\\index.html:1032,1107 — comments: \"refreshTokenPanels also rotates the LEFT panel (FC /session) per REQ-022\"",
        "Website\\index.html:1177,1730 — renders \"Server, serving security token:\" string",
    ]),
    ("Library (FrontEndApplicationLib) references", [
        "FrontEndApplicationLib\\src\\chathealthy_frontend_lib\\session_token.py:5-127 — full module: SessionToken, SessionTokenVerification, generate_session_token, verify_session_token, CERTS_DIR",
        "FrontEndApplicationLib\\src\\chathealthy_frontend_lib\\__init__.py:12-27 — re-exports",
        "FrontEndApplicationLib\\src\\chathealthy_frontend_lib\\runtime_governance.py:13-107 — _extract_session_token for governance logging",
        "FrontEndApplicationLib\\pyproject.toml:4 — package description: \"token MANAGEMENT (verify, parse, advance nonce) lives here; token MANUFACTURE belongs 100% to SharedServices\"",
        "sharedServices\\APICatalogue\\openapi.json:44-323 — /session, /verify-token, SessionToken, SessionTokenVerification schemas",
        "evaluateCare\\APICatalogue\\openapi.json:44-480 — same as SharedServices",
    ]),
    ("Backlog REQ references (agile_backlog.json)", [
        "EPIC-002-F-001-S-012-REQ-B-007 (line 492) — display nonce + GUID + signed token as separate labeled fields",
        "EPIC-002-F-001-S-012-REQ-B-008 (line 512) — Fatal-Error splash on session-token verification timeout (non-prod)",
        "EPIC-002-F-001-S-012-REQ-B-009 (line 531) — session-verification panel matches NonProdSecuriytToken mock",
        "EPIC-002-F-001-S-012-REQ-B-010 (line 551) — verify-token response origin self-identifies the responding service",
        "EPIC-002-F-001-S-012-REQ-T-001 (line 571) — SharedServices mutual auth using FindCare-created session token",
        "EPIC-002-F-001-S-012-REQ-T-002 (line 591) — mutual-auth TLS 1.2+ for inter-service calls",
        "EPIC-002-F-001-S-012-REQ-T-004 (line 629) — startup primitive verification, exit 78/77/70",
        "EPIC-002-F-001-S-012-REQ-T-005 (line 647) — FINDCARE_CERT_PEM secret decoded to CERTS_DIR",
        "EPIC-002-F-001-S-012-REQ-T-006 (line 667) — origin field signed/verified per service",
        "EPIC-002-F-001-S-012-REQ-T-007 (line 687) — browser must call cold service /verify-token directly (no FindCare proxy); mentions BUG-VERIFY-TOKEN-PROXY-001",
        "EPIC-009-F-001-S-004-REQ-B-001 (line 1117) — \"GET /session must return a signed session token with origin=FindCare. Falls back to unsigned token if cert unavailable.\" STALE? — story is under EPIC-009 (Shared Services) which is now merged into EPIC-002, and the \"fallback to unsigned\" allowance contradicts EPIC-002-F-001-S-012-REQ-T-004's hard-exit rule.",
        "EPIC-009-F-001-S-008-REQ-B-001 area (line 4986) — \"When the hot button is clicked, the selected providers are sent to EvaluateCare with a signed session token.\"",
    ]),
    ("Bug references (bugs.json)", [
        "(none) — bugs.json contains BUG-003 through BUG-007; none mention session/token/nonce/GUID/security-panel. Note: code comments cite BUG-VERIFY-TOKEN-PROXY-001 (app.py / EPIC-002-F-001-S-012-REQ-T-007) and BUG-ARCH-GRAPH-EXEMPT-001 (app.py:594,628) and BUG-SEC-007 (capability audit) — none of these IDs exist in bugs.json. STALE?",
    ]),
    ("Design doc references (architecture/)", [
        "architecture\\ChatHealthyApplicationDesign\\ArchitectureDesignAndAuditDocs\\generate_pre_sql_capability_audit_V1.py:125-165 — CAP-R08/R13/R14 binding /session, /shared/verify-token, /evaluate/verify-token",
        "architecture\\ChatHealthyApplicationDesign\\ArchitectureDesignAndAuditDocs\\generate_pre_sql_capability_audit_V1.py:459-666 — CAP-B02/B03 startup primitive + cert bootstrap",
        "architecture\\ChatHealthyApplicationDesign\\ArchitectureDesignAndAuditDocs\\generate_pre_sql_capability_audit_V2.py:706,710 — references /session and proxy endpoints",
        "architecture\\ChatHealthyApplicationDesign\\ArchitectureDesignAndAuditDocs\\generate_design_V1.py:375 — names session_token.py as a \"shared cross-capability concern\"",
        "architecture\\FindCare\\ArchitectureDesignAndAuditDocs\\generate_codebase_reorganization_V4.py:337,380,549 — describes auth-token display + LEFT/RIGHT token panels",
        "architecture\\FindCare\\ArchitectureDesignAndAuditDocs\\generate_codebase_reorganization_V4.py:421,423 — V4 reorganization places auth-token CREATION/VERIFICATION under EPIC-002, display under EPIC-011; mentions session_token.py under EPIC-009. STALE? — EPIC-009 merged into EPIC-002.",
        "architecture\\FindCare\\ArchitectureDesignAndAuditDocs\\generate_codebase_reorganization_V7.py:73,499,1077,1098 — V7 also discusses auth-token display + session-token verification",
        "architecture\\FindCare\\ArchitectureDesignAndAuditDocs\\generate_langgraph_oo_doc_V6.py:486,557 — mentions Security capability subgraph + OAuth/session-token",
        "architecture\\DevOpsBuildDeployAndEnvironmentManagement\\ArchitectureDesignAndAuditDocs\\audit_chunks\\chunk_C_py.txt:9 — lists Code/Shared/session_token.py. STALE? — that path no longer exists; canonical is FrontEndApplicationLib\\src\\chathealthy_frontend_lib\\session_token.py.",
        "brain\\machine_artifacts\\content\\risk_acceptance.json:73 — accepted-risk note mentions adding nonce to session tokens",
    ]),
])

CONCEPT_2_GUID = ("2. GUID", [
    ("Backend (Python) references", [
        "Code\\deploy\\localSmokeTestPyTest.py:146 — \"\"\"Parse CH{nonce:34}{guid:32} into nonce and guid.\"\"\"",
        "Code\\deploy\\localSmokeTestPyTest.py:197-277 — REQ-T-003 GUID stability assertions; _verify_session_identity returns (nonce, guid)",
        "Code\\deploy\\localSmokeTestPyTest.py:759 — token_el = left.locator(\"#guiSessionCell, #guiSessionId\")",
        "Code\\deploy\\localSmokeTestPyTest.py:775,939,1047,1084,1119,1162 — GUID/nonce assertions across handoffs",
        "Code\\deploy\\devSmokeTestPyTest.py:146,197-270,617-628,759,890,923,956 — mirror",
        "Code\\ConversationalUX\\FindCareChat\\backend\\tests\\local_environment_user_ready.py:272 — page.locator(\"#guiSessionId\")",
        "Code\\ConversationalUX\\FindCareChat\\backend\\tests\\dev_environment_user_ready.py:249 — same",
    ]),
    ("Frontend (TypeScript/React) references", [
        "Code\\ConversationalUX\\FindCareChat\\frontend\\src\\components\\FindCareApp.tsx:349 — JSX id=\"guiSessionCell\" cell",
        "Code\\ConversationalUX\\FindCareChat\\frontend\\src\\components\\FindCareApp.tsx:364,386 — sendToParent('gui:session-display', {...})",
        "Code\\ConversationalUX\\FindCareChat\\frontend\\src\\components\\ChatWindow.tsx:264,303-306 — type: 'gui:session-display' postMessage; comment \"Show session GUID in control frame\"",
    ]),
    ("Wrapper (Website/index.html) references", [
        "Website\\index.html:935-941 — getElementById('guiSessionCell') || getElementById('guiSessionId'); injects guiSessionId div",
        "Website\\index.html:1080 — var sessionEl = document.getElementById('guiSessionId');",
        "Website\\index.html:1653,1678-1683 — gui:session-display handler resolves the GUID cell",
        "Website\\index.html:1141,1694 — extracts GUID substring from token string",
    ]),
    ("Library (FrontEndApplicationLib) references", [
        "FrontEndApplicationLib\\src\\chathealthy_frontend_lib\\session_token.py:61 — _session_guid: Optional[str] = None",
        "FrontEndApplicationLib\\src\\chathealthy_frontend_lib\\session_token.py:84-94 — token format CH{nonce:17}{created:17}{guid:32}, _session_guid = uuid.uuid4().hex",
    ]),
    ("Backlog REQ references (agile_backlog.json)", [
        "EPIC-002-F-001-S-012-REQ-B-007 (line 492) — GUID and nonce as separate labeled fields",
        "EPIC-002-F-001-S-012-REQ-T-003 (line 609) — \"Session token uses stable GUID and freshly regenerated nonce on each call\"",
        "EPIC-008-F-004-S-006-REQ-B-018 (line 9650) — Step 17: handoff yields fresh nonce, session GUID identical; element #guiSessionCell or #guiSessionId positioned below specialty list",
        "EPIC-008-F-004-S-006-REQ-B-020 (line 9690) — Step 19: EvaluateCare-side GUID equals original",
        "EPIC-008-F-004-S-006-REQ-B-025 (line 9789) — Step 24: GUID unchanged on return",
        "EPIC-008-F-004-S-006-REQ-B-029 / B-030 (lines 9889-9911) — Step 29/30: SharedServices preserves GUID",
    ]),
    ("Bug references (bugs.json)", [
        "(none)",
    ]),
    ("Design doc references (architecture/)", [
        "architecture\\FindCare\\ArchitectureDesignAndAuditDocs\\generate_codebase_reorganization_V4.py:549 — \"renders nonce + GUID + signature\"",
    ]),
])

CONCEPT_3_DISPLAY = ("3. Display Stuff (security panels)", [
    ("Backend (Python) references", [
        "Code\\deploy\\localSmokeTestPyTest.py:186,1040,1057 — refreshTokenPanels rotates LEFT and RIGHT in lockstep; pytest.skip refs to \"security token panel\"",
        "Code\\deploy\\devSmokeTestPyTest.py:186 — same comment",
        "architecture\\ChatHealthyApplicationDesign\\ArchitectureDesignAndAuditDocs\\generate_pre_sql_capability_audit_V1.py:165 — refers to \"security panel for the EvaluateCare button\"",
    ]),
    ("Frontend (TypeScript/React) references", [
        "Code\\ConversationalUX\\FindCareChat\\frontend\\src\\components\\FindCareApp.tsx:364,386 — sendToParent('gui:session-display', ...) (the wrapper's panel-update postMessage)",
        "Code\\ConversationalUX\\FindCareChat\\frontend\\src\\components\\ChatWindow.tsx:264,306 — same gui:session-display send",
    ]),
    ("Wrapper (Website/index.html) references", [
        "Website\\index.html:701 — comment about path that \"broke refreshTokenPanels\"",
        "Website\\index.html:862-874 — ensureRightPanelStructure builds #rightContent + #rightSecurity",
        "Website\\index.html:878-954 — window.refreshTokenPanels = function(owner) body; suppresses panel in prod per EPIC-008-F-004-S-004-REQ-B-002",
        "Website\\index.html:999-1108 — calls to refreshTokenPanels('evaluatecare'), ('findcare'), ('sharedservices')",
        "Website\\index.html:1138-1177 — window.buildTokenDisplayHtml (signed/Nonce/GUID rendering, REQ-B-007)",
        "Website\\index.html:1329,1428,1524 — additional refreshTokenPanels(...) calls",
        "Website\\index.html:1635-1689 — gui:session-display receiver paints the same buildTokenDisplayHtml into #guiSessionId",
    ]),
    ("Library (FrontEndApplicationLib) references", [
        "FrontEndApplicationLib\\src\\chathealthy_frontend_lib\\session_token.py — model only (no DOM); pyproject.toml:4 notes display is consumer concern",
        "(No widget code yet — EPIC-003-F-001 plans to extract rightSecurity panel; library currently does not contain it.)",
    ]),
    ("Backlog REQ references (agile_backlog.json)", [
        "EPIC-002-F-001-S-012-REQ-B-007 (line 492) — display contract for nonce/GUID/signed parts; right-panel/left-panel parity",
        "EPIC-002-F-001-S-012-REQ-B-009 (line 531) — both panels match NonProdSecuriytToken mock",
        "EPIC-008-F-004-S-004-REQ-B-002 (line 9150) — \"no security token panel\" in prod",
        "EPIC-008-F-004-S-005 (line 9172, story description) — \"Owns the environment banner, the security-token panel\"",
        "EPIC-008-F-004-S-005-REQ-B-005 (line 9258) — \"Security token panel\" matches the mock",
        "EPIC-008-F-004-S-006-REQ-B-019 (line 9670) — Step 18: \"Token components in distinct colors in both panels\"",
        "EPIC-008-F-004-S-006-REQ-B-026 (line 9810) — Step 25: #rightPanel self-identifies \"Server, serving security token: SharedServices\"",
        "EPIC-003 (line 1623, especially line 1637 F-001 description) — \"non-prod security panel (right-side rightSecurity panel)\" planned for EPIC-003 widget extraction",
    ]),
    ("Bug references (bugs.json)", [
        "(none)",
    ]),
    ("Design doc references (architecture/)", [
        "architecture\\FindCare\\ArchitectureDesignAndAuditDocs\\generate_codebase_reorganization_V4.py:337 — \"F-005 Auth-token display object — buildTokenDisplayHtml(st) and the LEFT/RIGHT token panels\"",
        "architecture\\FindCare\\ArchitectureDesignAndAuditDocs\\generate_codebase_reorganization_V4.py:380 — \"Website/index.html ... LEFT/RIGHT token panels\"",
        "architecture\\FindCare\\ArchitectureDesignAndAuditDocs\\generate_codebase_reorganization_V4.py:549 — \"buildTokenDisplayHtml(st) at lines 1011-1045 renders nonce + GUID + signature\". STALE? — actual current line is Website\\index.html:1138, not 1011-1045.",
        "architecture\\FindCare\\ArchitectureDesignAndAuditDocs\\generate_codebase_reorganization_V7.py:463-1077 — V7 plans auth_token_display/token_display.js; F-005 1-story plan",
        "architecture\\FindCare\\ArchitectureDesignAndAuditDocs\\generate_codebase_reorganization_V4.py:421 — \"the visual auth-token DISPLAY object is owned by EPIC-011\"; STALE? — EPIC-011 not present in current backlog grep results",
    ]),
])

CONCEPT_4_NONCE = ("4. NONCE", [
    ("Backend (Python) references", [
        "Code\\ConversationalUX\\FindCareChat\\backend\\app.py:603 — comment \"running with a token that has no nonce and no signature\"",
        "Code\\ConversationalUX\\FindCareChat\\backend\\tests\\local_environment_user_ready.py:200 — docstring \"Left panel shows session token starting with ch_ and nonce timestamps\"",
        "Code\\deploy\\localSmokeTestPyTest.py:146-277 — _parse_token extracts nonce; _verify_session_identity enforces freshness across handoffs (REQ-T-003)",
        "Code\\deploy\\localSmokeTestPyTest.py:458,745-756,775,789-1163 — full nonce-rotation assertions across all 6 handoffs",
        "Code\\deploy\\devSmokeTestPyTest.py:146-957 — mirror of above",
        "sharedServices\\Code\\security\\session_endpoint.py:28 — log \"Minted SharedServices session token: env=%s nonce_len=%d\"",
        "evaluateCare\\Code\\security\\session_endpoint.py:28 — log \"Minted EvaluateCare session token: env=%s nonce_len=%d\"",
    ]),
    ("Frontend (TypeScript/React) references", [
        "Code\\ConversationalUX\\FindCareChat\\frontend\\src\\components\\FindCareApp.tsx:384 — comment \"Update left panel with verified token (new nonce from EvaluateCare)\"",
    ]),
    ("Wrapper (Website/index.html) references", [
        "Website\\index.html:904 — comment \"smoke asserts cross-panel nonce equality\"",
        "Website\\index.html:1136,1140-1170 — extracts nonce substring; renders <b style=\"color:#d97706;\">Nonce:</b>",
        "Website\\index.html:1694,1724 — same in gui:session-display handler",
    ]),
    ("Library (FrontEndApplicationLib) references", [
        "FrontEndApplicationLib\\src\\chathealthy_frontend_lib\\session_token.py:29,84 — token format spec including {nonce:17} (timestamps)",
        "FrontEndApplicationLib\\pyproject.toml:4 — \"advance nonce\" lives in this package",
        "sharedServices\\APICatalogue\\openapi.json:187 — schema: CH{nonce:17}{created:17}{guid:32}",
        "evaluateCare\\APICatalogue\\openapi.json:337 — same schema",
    ]),
    ("Backlog REQ references (agile_backlog.json)", [
        "EPIC-002-F-001-S-012-REQ-B-007 (line 492) — nonce field in display (only owning service updates)",
        "EPIC-002-F-001-S-012-REQ-T-003 (line 609) — nonce regenerated fresh on every call (NIST 800-53 SC-23)",
        "EPIC-008-F-004-S-006-REQ-B-017 (line 9430-9432) — \"captures the originating session token for downstream nonce comparisons\"",
        "EPIC-008-F-004-S-006-REQ-B-018, B-019, B-020, B-025, B-026, B-029, B-030, B-031, B-032, B-033 (lines 9650-9972) — every cross-service handoff requires fresh nonce",
    ]),
    ("Bug references (bugs.json)", [
        "(none) — code comment in risk_acceptance.json:73 historical: \"Add nonce to session tokens\" (initial wiring authorization)",
    ]),
    ("Design doc references (architecture/)", [
        "architecture\\ChatHealthyApplicationDesign\\ArchitectureDesignAndAuditDocs\\generate_pre_sql_capability_audit_V1.py:125-159 — \"GET /session — generate signed session token (NONCE + cert-signed)\"; CAP-R13 mentions \"NONCE update\"",
        "architecture\\FindCare\\ArchitectureDesignAndAuditDocs\\generate_codebase_reorganization_V4.py:421,549 — nonce-based auth design; nonce rendering",
        "architecture\\FindCare\\ArchitectureDesignAndAuditDocs\\generate_langgraph_oo_doc_V6.py:557 — Security subgraph references",
        "brain\\machine_artifacts\\content\\risk_acceptance.json:73 — \"(2) Add nonce to session tokens\"",
    ]),
])

STALE_FLAGS = [
    "EPIC-009-F-001-S-004-REQ-B-001 (agile_backlog.json:1117) — sits under merged EPIC-009 and allows \"fallback to unsigned token\" which contradicts EPIC-002-F-001-S-012-REQ-T-004's hard-exit rule.",
    "BUG-VERIFY-TOKEN-PROXY-001, BUG-ARCH-GRAPH-EXEMPT-001, BUG-SEC-007 — cited in code comments (Code\\ConversationalUX\\FindCareChat\\backend\\app.py:594,628; capability audit) but NOT present in bugs.json.",
    "architecture\\DevOpsBuildDeployAndEnvironmentManagement\\ArchitectureDesignAndAuditDocs\\audit_chunks\\chunk_C_py.txt:9 — references Code/Shared/session_token.py, a path that no longer exists.",
    "architecture\\FindCare\\ArchitectureDesignAndAuditDocs\\generate_codebase_reorganization_V4.py:421,423,549 — places display under \"EPIC-011\" and session_token under EPIC-009; both attributions are out of date given the EPIC-009→EPIC-002 merge.",
    "generate_codebase_reorganization_V4.py:549 — line numbers it cites in Website\\index.html (1011-1045) are off; current buildTokenDisplayHtml is at line 1138.",
]


def render_concept(doc, title, sections):
    H(doc, title, level=1)
    for label, bullets in sections:
        H(doc, label, level=2)
        for b in bullets:
            BULLET(doc, b)


def main():
    doc = Document()

    H(doc, "Security / Session Token Inventory (V1)", level=0)
    P(doc, "Author: Skip Snow + Claude (research agent). 2026-05-08.")
    P(doc, "Substance produced by an inventory agent dispatched on 2026-05-08, "
           "mapping every reference to four related concepts (security/session "
           "token, GUID, display panels, NONCE) across code, wrapper, library, "
           "agile_backlog, bugs.json, and design docs. Intended as the seed "
           "reference for the EPIC-002 (Security And Shared Services) "
           "Authentication-feature work scheduled for 2026-05-09.")

    H(doc, "Summary", level=1)
    P(doc, "Across the codebase, the inventory found ~150 references total: "
           "~70 for session/security token, ~17 for GUID, ~30 for display "
           "panels, ~38 for nonce. The densest concentrations are in "
           "Code\\deploy\\localSmokeTestPyTest.py and devSmokeTestPyTest.py "
           "(every nonce/GUID assertion lives here, plus token panel "
           "scrutiny), Website\\index.html (the wrapper owns refreshTokenPanels, "
           "buildTokenDisplayHtml, rightSecurity/leftSecurity panels, and "
           "the gui:session-display postMessage receiver), and "
           "FrontEndApplicationLib\\src\\chathealthy_frontend_lib\\session_token.py "
           "(canonical generate_session_token / verify_session_token / "
           "SessionToken model). Three backend services (FindCare, "
           "SharedServices, EvaluateCare) each expose /session and "
           "/verify-token. bugs.json contains zero open bugs touching any of "
           "these four concepts. Backlog references concentrate in "
           "EPIC-002-F-001-S-012 (the security story) and "
           "EPIC-008-F-004-S-006 (the smoke-test step inventory), with an "
           "additional historical pointer at the now-merged "
           "EPIC-009-F-001-S-004.")

    for concept in (CONCEPT_1_TOKEN, CONCEPT_2_GUID,
                    CONCEPT_3_DISPLAY, CONCEPT_4_NONCE):
        render_concept(doc, *concept)

    H(doc, "STALE? flags collected", level=1)
    for i, line in enumerate(STALE_FLAGS, start=1):
        BULLET(doc, f"{i}. {line}")

    doc.save(str(OUT))
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
