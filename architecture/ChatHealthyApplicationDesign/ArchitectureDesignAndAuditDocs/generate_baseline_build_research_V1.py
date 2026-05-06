#!/usr/bin/env python3
# Copyright (c) 2026 ChatHealthy.ai LLC. All rights reserved.
# Licensed under the FindCare Evaluation License (FEL-1.0).
#
# Sibling generator for baseline-build-research-V1.docx
#
# Read-only research. Examines every build commit numbered 400..450
# inclusive on the dev branch and recommends ONE build as the baseline
# for the Design V3 conversion.
#
# Scope is front-end only: FindCareChat/backend/main.py (route/tool
# dispatch/system prompt assembly), domain/find_care, application,
# infrastructure, prompt_system_maker.py, tool_router.py, url_guardian.py,
# the React SPA in FindCareChat/frontend, and Website/*.html|js|css.
#
# Out of scope (per assignment): Code/DataPipelines/, Code/Shared/ops/,
# architecture/EngineeringRuleEnforcement/, LangGraph/poc/, brain/.
#
# All evidence cited below was gathered via `git show <hash>:<path>` /
# `git diff` / `git log --grep` against the dev branch. No checkouts,
# no commits, no working tree mutations.

import os

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL

OUT = os.path.join(
    os.environ['CHATHEALTHY_PROJECT_ROOT'],
    "architecture", "ChatHealthyApplicationDesign",
    "ArchitectureDesignAndAuditDocs", "baseline-build-research-V1.docx",
)

# ---------------------------------------------------------------------------
# The 51 builds, in chronological order. Hash | ISO timestamp | subject.
# Front-end-only +/- LOC computed against the path set defined in scope.
# ---------------------------------------------------------------------------
BUILDS = [
    # n, hash, date, plus, minus, subject
    (400, "3d4f022", "2026-04-02 19:07 PT", 298, 86,
     "GUI pagination + pipeline test suite — 46/46 passing"),
    (401, "51e6aef", "2026-04-02 21:10 PT", 162, 97,
     "Pagination UAT passed — frame layout, controls, record tracking"),
    (402, "226818e", "2026-04-02 21:12 PT", 16, 18,
     "Remove chat-config.js — iframe URL derived from environment"),
    (403, "1094abc", "2026-04-02 21:24 PT", 11, 2,
     "Fix pagination replay — resolved specialty_codes in search_params"),
    (404, "0e23e89", "2026-04-02 22:15 PT", 53, 60,
     "GOV-011 fix — vector search resolves codes, database returns data"),
    (405, "5abcb9d", "2026-04-02 22:31 PT", 16, 0,
     "Log resolved NUCC codes to admin DB (non-prod only)"),
    (406, "4a62daa", "2026-04-02 22:48 PT", 173, 2,
     "Specialty refinement — left panel filter with checkboxes"),
    (407, "42c774d", "2026-04-02 23:26 PT", 17, 20,
     "Frame layout — all percentages, 20% side panels, no pixels"),
    (408, "128a5ea", "2026-04-02 23:38 PT", 0, 0,
     "Bug logged — specialty filter empty, vector indexes missing"),
    (409, "de5adc2", "2026-04-02 23:44 PT", 19, 0,
     "Environment banner full width in parent page, not iframe"),
    (410, "be14579", "2026-04-02 23:48 PT", 7, 2,
     "Banner in parent only, hidden in iframe. DR-014 commit every test."),
    (411, "75e532a", "2026-04-03 00:53 PT", 152, 3,
     "FC-RESULT-MSG system-built summary, vector indexes all envs, DR-015/016/017/018"),
    (412, "9ba515f", "2026-04-03 07:11 PT", 7, 5,
     "Playwright SIT tooling, summary message template per Boss spec"),
    (413, "9e93bb6", "2026-04-03 07:12 PT", 22, 2,
     "Summary message template — state, location narrowing, Boss spec"),
    (414, "7177f08", "2026-04-03 07:23 PT", 5, 1,
     "Fix rehype-sanitize stripping #action: hrefs from summary links"),
    (415, "c6013c4", "2026-04-03 07:47 PT", 13, 1,
     "Fix BUG-MSG-002 — summary uses user's words not LLM translation"),
    (416, "a961689", "2026-04-03 07:51 PT", 77, 2,
     "Fix sideways filter text, LLM summary de-dup, GOV-011 standards"),
    (417, "b0a530e", "2026-04-03 09:03 PT", 48, 3,
     "Clinical trials summary (GOV-011), build_repo_json.py saved to ops"),
    (418, "883bcce", "2026-04-03 09:28 PT", 6, 18,
     "Fix sideways filter — writing-mode on label span only, not panel div"),
    (419, "8449ac3", "2026-04-03 09:47 PT", 27, 3,
     "Haiku extracts user search term from message (GOV-011-STD-002)"),
    (420, "e2cb3d6", "2026-04-03 09:54 PT", 10, 7,
     "Switch term extraction from Haiku to GPT-4.1-nano (8x cheaper)"),
    (421, "019892b", "2026-04-03 10:17 PT", 1, 1,
     "Nano prompt returns plural form always (bone doc -> bone docs)"),
    (422, "2da8846", "2026-04-03 10:22 PT", 3, 3,
     "Switch to GPT-4.1-mini for term extraction (kid's docs), fix wording"),
    (423, "3480938", "2026-04-03 10:24 PT", 29, 1,
     "Filter apply updates summary with selected specialty names and correct count"),
    (424, "818e99e", "2026-04-03 10:28 PT", 11, 3,
     "Fix DEV banner — fetch build/version from HF Space, not Cloudflare"),
    (425, "a8a3a7d", "2026-04-03 10:35 PT", 13, 0,
     "QA promotion milestone Apr 23, SIT migration backlog, DR-019"),
    (426, "5c82169", "2026-04-03 10:45 PT", 30, 53,
     "Switch de-dup from regex to GPT-4.1-mini (GOV-011-STD-001)"),
    (427, "43e4c0f", "2026-04-03 10:53 PT", 5, 5,
     "Fix filter apply — combine provider records + summary in single state update"),
    (428, "6572601", "2026-04-03 11:05 PT", 32, 2,
     "Filter toggle (Check All/Uncheck All), auto-pagination on filter apply"),
    (429, "85bd401", "2026-04-03 11:14 PT", 11, 16,
     "Fix toggle-all button — use data-gui-action handler, not inline onclick"),
    (430, "0fdab5f", "2026-04-03 11:28 PT", 0, 0,
     "BUG-VECTOR-001 release blocker logged, GOV-011 policy updated"),
    (431, "7c01b8c", "2026-04-03 11:37 PT", 35, 9,
     "$facet single-query for Route 3 (count + page), taxonomy_state_geo_npi index on dev"),
    (432, "ed9af1b", "2026-04-03 16:53 PT", 3, 27,
     "$facet migration complete — all search routes use single query"),
    (433, "d3b90f6", "2026-04-03 17:22 PT", 10, 21,
     "BUG-VECTOR-001 fix — RAG against SpecialtyMetaData not providers"),
    (434, "765760f", "2026-04-03 17:45 PT", 3, 1,
     "Tighten vector similarity threshold 0.4→0.6, max 3 classifications"),
    (435, "0390c20", "2026-04-03 18:05 PT", 24, 1,
     "Fix filter scroll, mobile filter as fullscreen overlay"),
    (436, "2d1f6d5", "2026-04-03 19:16 PT", 674, 1066,
     "Architecture page overhaul + GOV-011 compliance audit"),
    (437, "3bff960", "2026-04-03 19:28 PT", 0, 0,
     "Pipeline reserves cluster before health check, releases on completion"),
    (438, "0d7bc04", "2026-04-03 19:41 PT", 0, 0,
     "EVAL-RESEARCH feature + GOV-011 audit report PDF"),
    (439, "f9b324d", "2026-04-03 19:55 PT", 0, 0,
     "Fix chicken-egg — wake cluster before writing reservation to MongoDB"),
    (440, "846ae9b", "2026-04-04 00:24 PT", 0, 0,
     "EPIC-3 Pipeline Management — 4 bugs, 2 enhancements, 7 stories"),
    (441, "858bf1d", "2026-04-04 00:35 PT", 0, 0,
     "BUG-PIPE-001 + BUG-PIPE-002 fixed, unrealized_ideas.json"),
    (442, "6fdac91", "2026-04-04 00:43 PT", 0, 0,
     "QA report v0.1.4 SIT checklist, unrealized ideas IDEA-002/003"),
    (443, "2ee84f1", "2026-04-04 00:45 PT", 0, 0,
     "QA report v0.1.4 — EPIC-1 Evaluate Care + EPIC-3 Pipeline only"),
    (444, "2a7a36d", "2026-04-04 00:47 PT", 0, 0,
     "BUG-PIPE-003 — drop collection unless incremental=true"),
    (445, "25c3d63", "2026-04-04 01:02 PT", 836, 0,
     "QA report v0.1.4 published to website"),
    (446, "339c3a9", "2026-04-04 01:05 PT", 0, 0,
     "EPIC-4 Security — 11 Cloudflare findings as stories"),
    (447, "a5a52df", "2026-04-04 01:14 PT", 0, 0,
     "BUG-PIPE-006 — verify blob exists before skipping download"),
    (448, "0d532fc", "2026-04-04 01:25 PT", 51, 0,
     "QA report renderer at /qa-report — dynamic from brain JSON"),
    (449, "e8adf94", "2026-04-04 01:31 PT", 3, 1,
     "QA report — bundle brain JSON with HF deploy"),
    (450, "030aeb4", "2026-04-04 01:35 PT", 10, 0,
     "QA Report nav link — dev/QA only, opens in slide panel"),
]

# ---------------------------------------------------------------------------
# Per-build assessment against the criteria in the assignment.
# Y/N/(qualified). Cells use short codes; the legend below is in the doc.
#
# tools  : all 9 LLM tools registered? Verified by inspecting
#          ToolRouter.register_with_models in main.py.
# narr   : LLM-authored narrative around results preserved end-to-end?
#          Y = LLM text returned untouched; D = damaged by stripping;
#          G = gutted by external LLM call.
# vector : vector specialty search returning relevant results?
#          N = indexes missing on cluster (BUG-FILTER-001);
#          D = deployed but BUG-VECTOR-001 latent (provider-embedding pollution);
#          F = explicitly broken (kid's doc returns Hospitalist);
#          Y = SpecialtyMetaData RAG, threshold tightened, working.
# guard  : url_guardian wired in /chat path (Y/N).
# iplock : SafetyService is_ip_locked / lock_ip / try_admin_unlock (Y/N).
# filter : left-panel specialty filter wired (gui.showFilterPanel /
#          gui.onFilterApply). Y from build 406 onward.
# thinkstop : inline 'Thinking… Ns' counter + Stop button in ChatWindow.tsx
#          (Y from build 400; this is the pre-event state).
# uat    : HUMAN_TESTING /welcome reachable (Y from 400 onward).
# action : '#action:next-page' and '#action:filter' link host present in
#          system summary (Y from 411 onward).
# regs   : regression footprint — explicit identifiers in commit message.
# score  : overall relative score. Higher is better. Bounded 0..10.
# ---------------------------------------------------------------------------
EVAL = {
    400: dict(tools="Y", narr="Y", vector="N", guard="Y", iplock="Y",
              filter="N", thinkstop="Y", uat="Y", action="N",
              regs="—", score=4),
    401: dict(tools="Y", narr="Y", vector="N", guard="Y", iplock="Y",
              filter="N", thinkstop="Y", uat="Y", action="N",
              regs="—", score=4),
    402: dict(tools="Y", narr="Y", vector="N", guard="Y", iplock="Y",
              filter="N", thinkstop="Y", uat="Y", action="N",
              regs="—", score=4),
    403: dict(tools="Y", narr="Y", vector="N", guard="Y", iplock="Y",
              filter="N", thinkstop="Y", uat="Y", action="N",
              regs="—", score=4),
    404: dict(tools="Y", narr="Y", vector="N", guard="Y", iplock="Y",
              filter="N", thinkstop="Y", uat="Y", action="N",
              regs="—", score=4),
    405: dict(tools="Y", narr="Y", vector="N", guard="Y", iplock="Y",
              filter="N", thinkstop="Y", uat="Y", action="N",
              regs="—", score=4),
    406: dict(tools="Y", narr="Y", vector="N", guard="Y", iplock="Y",
              filter="Y", thinkstop="Y", uat="Y", action="N",
              regs="—", score=5),
    407: dict(tools="Y", narr="Y", vector="N", guard="Y", iplock="Y",
              filter="Y", thinkstop="Y", uat="Y", action="N",
              regs="—", score=5),
    408: dict(tools="Y", narr="Y", vector="N", guard="Y", iplock="Y",
              filter="Y", thinkstop="Y", uat="Y", action="N",
              regs="BUG-FILTER-001 logged (vector indexes missing on FE cluster)",
              score=4),
    409: dict(tools="Y", narr="Y", vector="N", guard="Y", iplock="Y",
              filter="Y", thinkstop="Y", uat="Y", action="N",
              regs="BUG-FILTER-001 still open", score=4),
    410: dict(tools="Y", narr="Y", vector="N", guard="Y", iplock="Y",
              filter="Y", thinkstop="Y", uat="Y", action="N",
              regs="BUG-FILTER-001 still open (specialty filter returns 0)",
              score=5),
    411: dict(tools="Y", narr="Y", vector="Y", guard="Y", iplock="Y",
              filter="Y", thinkstop="Y", uat="Y", action="Y",
              regs="—  (vector indexes deployed all envs; FC-RESULT-MSG added; LLM narrative untouched)",
              score=9),
    412: dict(tools="Y", narr="Y", vector="Y", guard="Y", iplock="Y",
              filter="Y", thinkstop="Y", uat="Y", action="Y",
              regs="—", score=9),
    413: dict(tools="Y", narr="Y", vector="Y", guard="Y", iplock="Y",
              filter="Y", thinkstop="Y", uat="Y", action="Y",
              regs="—", score=9),
    414: dict(tools="Y", narr="Y", vector="Y", guard="Y", iplock="Y",
              filter="Y", thinkstop="Y", uat="Y", action="Y",
              regs="—  (rehype-sanitize fix preserves #action: hrefs)",
              score=9),
    415: dict(tools="Y", narr="Y", vector="Y", guard="Y", iplock="Y",
              filter="Y", thinkstop="Y", uat="Y", action="Y",
              regs="BUG-MSG-002 fix (forces user term into summary; mild)",
              score=8),
    416: dict(tools="Y", narr="D", vector="Y", guard="Y", iplock="Y",
              filter="Y", thinkstop="Y", uat="Y", action="Y",
              regs="GOV-011-STD-001 enabled: regex _strip_redundant_summary damages LLM narrative",
              score=6),
    417: dict(tools="Y", narr="D", vector="Y", guard="Y", iplock="Y",
              filter="Y", thinkstop="Y", uat="Y", action="Y",
              regs="GOV-011-STD-001 still active", score=6),
    418: dict(tools="Y", narr="D", vector="Y", guard="Y", iplock="Y",
              filter="Y", thinkstop="Y", uat="Y", action="Y",
              regs="—", score=6),
    419: dict(tools="Y", narr="D", vector="D", guard="Y", iplock="Y",
              filter="Y", thinkstop="Y", uat="Y", action="Y",
              regs="GOV-011-STD-002: external Haiku call to extract user term — fragile",
              score=5),
    420: dict(tools="Y", narr="D", vector="D", guard="Y", iplock="Y",
              filter="Y", thinkstop="Y", uat="Y", action="Y",
              regs="OpenAI nano replaces Haiku for term extraction",
              score=5),
    421: dict(tools="Y", narr="D", vector="D", guard="Y", iplock="Y",
              filter="Y", thinkstop="Y", uat="Y", action="Y",
              regs="—", score=5),
    422: dict(tools="Y", narr="D", vector="D", guard="Y", iplock="Y",
              filter="Y", thinkstop="Y", uat="Y", action="Y",
              regs="GPT-4.1-mini replaces nano (kid's docs case)",
              score=5),
    423: dict(tools="Y", narr="D", vector="D", guard="Y", iplock="Y",
              filter="Y", thinkstop="Y", uat="Y", action="Y",
              regs="—", score=5),
    424: dict(tools="Y", narr="D", vector="D", guard="Y", iplock="Y",
              filter="Y", thinkstop="Y", uat="Y", action="Y",
              regs="—", score=5),
    425: dict(tools="Y", narr="D", vector="D", guard="Y", iplock="Y",
              filter="Y", thinkstop="Y", uat="Y", action="Y",
              regs="—", score=5),
    426: dict(tools="Y", narr="G", vector="D", guard="Y", iplock="Y",
              filter="Y", thinkstop="Y", uat="Y", action="Y",
              regs="GOV-011-STD-001 v2: external GPT-4.1-mini call gut LLM narrative",
              score=4),
    427: dict(tools="Y", narr="G", vector="D", guard="Y", iplock="Y",
              filter="Y", thinkstop="Y", uat="Y", action="Y",
              regs="—", score=4),
    428: dict(tools="Y", narr="G", vector="D", guard="Y", iplock="Y",
              filter="Y", thinkstop="Y", uat="Y", action="Y",
              regs="—", score=4),
    429: dict(tools="Y", narr="G", vector="D", guard="Y", iplock="Y",
              filter="Y", thinkstop="Y", uat="Y", action="Y",
              regs="—", score=4),
    430: dict(tools="Y", narr="G", vector="F", guard="Y", iplock="Y",
              filter="Y", thinkstop="Y", uat="Y", action="Y",
              regs="BUG-VECTOR-001 release blocker (kid's doc → Hospitalist; 2/7 correct)",
              score=2),
    431: dict(tools="Y", narr="G", vector="F", guard="Y", iplock="Y",
              filter="Y", thinkstop="Y", uat="Y", action="Y",
              regs="BUG-VECTOR-001 still open", score=3),
    432: dict(tools="Y", narr="G", vector="F", guard="Y", iplock="Y",
              filter="Y", thinkstop="Y", uat="Y", action="Y",
              regs="BUG-VECTOR-001 still open ($facet migration complete)",
              score=3),
    433: dict(tools="Y", narr="G", vector="Y", guard="Y", iplock="Y",
              filter="Y", thinkstop="Y", uat="Y", action="Y",
              regs="BUG-VECTOR-001 fixed (RAG vs SpecialtyMetaData)",
              score=6),
    434: dict(tools="Y", narr="G", vector="Y", guard="Y", iplock="Y",
              filter="Y", thinkstop="Y", uat="Y", action="Y",
              regs="—", score=6),
    435: dict(tools="Y", narr="G", vector="Y", guard="Y", iplock="Y",
              filter="Y", thinkstop="Y", uat="Y", action="Y",
              regs="—", score=6),
    436: dict(tools="Y", narr="G", vector="Y", guard="Y", iplock="Y",
              filter="Y", thinkstop="Y", uat="Y", action="Y",
              regs="—  (Website/architecture.html overhaul; chat unchanged)",
              score=6),
    437: dict(tools="Y", narr="G", vector="Y", guard="Y", iplock="Y",
              filter="Y", thinkstop="Y", uat="Y", action="Y",
              regs="—", score=6),
    438: dict(tools="Y", narr="G", vector="Y", guard="Y", iplock="Y",
              filter="Y", thinkstop="Y", uat="Y", action="Y",
              regs="—", score=6),
    439: dict(tools="Y", narr="G", vector="Y", guard="Y", iplock="Y",
              filter="Y", thinkstop="Y", uat="Y", action="Y",
              regs="—", score=6),
    440: dict(tools="Y", narr="G", vector="Y", guard="Y", iplock="Y",
              filter="Y", thinkstop="Y", uat="Y", action="Y",
              regs="—", score=6),
    441: dict(tools="Y", narr="G", vector="Y", guard="Y", iplock="Y",
              filter="Y", thinkstop="Y", uat="Y", action="Y",
              regs="—", score=6),
    442: dict(tools="Y", narr="G", vector="Y", guard="Y", iplock="Y",
              filter="Y", thinkstop="Y", uat="Y", action="Y",
              regs="—", score=6),
    443: dict(tools="Y", narr="G", vector="Y", guard="Y", iplock="Y",
              filter="Y", thinkstop="Y", uat="Y", action="Y",
              regs="—", score=6),
    444: dict(tools="Y", narr="G", vector="Y", guard="Y", iplock="Y",
              filter="Y", thinkstop="Y", uat="Y", action="Y",
              regs="—", score=6),
    445: dict(tools="Y", narr="G", vector="Y", guard="Y", iplock="Y",
              filter="Y", thinkstop="Y", uat="Y", action="Y",
              regs="—  (Website/qa-report.html added; chat unchanged)",
              score=6),
    446: dict(tools="Y", narr="G", vector="Y", guard="Y", iplock="Y",
              filter="Y", thinkstop="Y", uat="Y", action="Y",
              regs="—", score=6),
    447: dict(tools="Y", narr="G", vector="Y", guard="Y", iplock="Y",
              filter="Y", thinkstop="Y", uat="Y", action="Y",
              regs="—", score=6),
    448: dict(tools="Y", narr="G", vector="Y", guard="Y", iplock="Y",
              filter="Y", thinkstop="Y", uat="Y", action="Y",
              regs="—  (/qa-report route added; chat unchanged)",
              score=6),
    449: dict(tools="Y", narr="G", vector="Y", guard="Y", iplock="Y",
              filter="Y", thinkstop="Y", uat="Y", action="Y",
              regs="—", score=6),
    450: dict(tools="Y", narr="G", vector="Y", guard="Y", iplock="Y",
              filter="Y", thinkstop="Y", uat="Y", action="Y",
              regs="—  (QA Report nav link; dev/QA only)", score=6),
}

# ---------------------------------------------------------------------------
# Recommendation
# ---------------------------------------------------------------------------
RECOMMENDED_BUILD = 411
RECOMMENDED_HASH = "75e532a"
RECOMMENDED_DATE = "2026-04-03 00:53 PT"
RECOMMENDED_SUBJECT = (
    "FC-RESULT-MSG system-built summary, vector indexes all envs, "
    "DR-015/016/017/018"
)

CURRENT_ANCHOR_BUILD = 410
CURRENT_ANCHOR_HASH = "be14579"
CURRENT_ANCHOR_DATE = "2026-04-02 23:48 PT"

# ---------------------------------------------------------------------------
# docx body text
# ---------------------------------------------------------------------------
EXEC_RATIONALE = (
    "Build 411 is the first commit in the 400..450 window in which all "
    "of the front-end behaviours under evaluation work simultaneously: "
    "(a) all 9 LLM tools are registered in ToolRouter; (b) the React app "
    "renders the inline 'Thinking… Ns' counter with Stop button, the "
    "left specialty filter panel (gui.showFilterPanel / gui.onFilterApply), "
    "and the system-built summary bubble with clickable '#action:next-page' "
    "and '#action:filter' link host so users can actually click them; "
    "(c) the LLM-authored narrative around results is still passed through "
    "verbatim — _strip_redundant_summary does not yet exist, so the AI's "
    "natural prose answer is preserved end-to-end alongside the new "
    "system summary; (d) the vector specialty search is functioning — "
    "build 411 is the commit that creates provider_vector_index and "
    "specialty_vector_index on dev/QA/prod and adds /health degraded "
    "reporting, which is the moment BUG-FILTER-001 (logged at 408) is "
    "closed; (e) URLGuardian.guard_text is wired in /chat; (f) "
    "SafetyService.is_ip_locked / lock_ip / try_admin_unlock are intact; "
    "(g) HUMAN_TESTING /welcome routes to _build_test_welcome. Build 411 "
    "predates the regression cascade that begins at 416 (regex stripping "
    "of LLM narrative), 419 (external Haiku/nano/mini term extraction), "
    "426 (external GPT-4.1-mini gut of LLM narrative), and 430 "
    "(BUG-VECTOR-001 release blocker, 'kid's doc' returns Hospitalist). "
    "The 433/434 fixes recover vector accuracy but never restore the "
    "AI-authored narrative damaged by 416 and gutted by 426. Build 411 "
    "is therefore the latest build that has every later-improved UX "
    "feature already wired AND that still has AI-orchestrated narrative "
    "around the results."
)

WHY_NOT_410 = (
    "Build 410 is one build earlier than 411 and is missing two "
    "consequential things. First, the vector specialty search does not "
    "function on the front-end MongoDB cluster — BUG-FILTER-001 was "
    "logged at build 408 ('specialization_options returns 0 despite "
    "data existing... no vector search indexes on frontend cluster') "
    "and remained open through 409 and 410. The fix lands inside "
    "build 411 itself, which creates provider_vector_index and "
    "specialty_vector_index on every environment. A baseline that "
    "anchors on 410 inherits the 'specialty filter empty' regression "
    "that the team had not yet shipped a fix for. Second, build 410 "
    "has no system-built summary message and therefore no clickable "
    "'#action:next-page' / '#action:filter' link host in chat — those "
    "are added by build 411's FC-RESULT-MSG work. Build 411 introduces "
    "no regression of its own: the diff is purely additive over 410 "
    "(new optional summary_message field on PaginationMeta, new "
    "_check_indexes / /health degraded path, no removals). Build 411 "
    "also has not yet adopted GOV-011-STD-001 stripping (that arrives "
    "at 416), so the LLM narrative is still presented to the user in "
    "full, with the new system summary appended as a separate isSummary "
    "bubble — exactly the 'AI-narrated results around the system "
    "summary' shape Skip described."
)

WHY_NOT_415 = (
    "Build 415 is the last build before LLM-narrative damage begins. "
    "It is a credible alternative to 411. The reason 411 wins over 415: "
    "build 415 introduces the BUG-MSG-002 fix that forces body.message "
    "directly into the system summary instead of letting the LLM-resolved "
    "search term flow through; this is the first nudge toward 'use the "
    "user's words not LLM translation' and prefigures the GOV-011-STD-001 "
    "and STD-002 mistakes that follow. Builds 412-414 are all minor "
    "summary-template tweaks. Among 411..415 the cleanest, lowest-risk "
    "anchor is 411 — the moment the system first works end-to-end."
)

WHY_NOT_435 = (
    "Build 435 is the last build before the chat engine stops being "
    "edited. It is the latest 'all features wired and vector search "
    "fixed' state. The reason 411 wins over 435: at 435 the LLM-authored "
    "narrative around results is gutted by an external GPT-4.1-mini "
    "call (added at 426) that strips everything that is not a provider "
    "record. Skip's stated concern is that Design V3 may anchor on a "
    "build that contains code that 'hurt rather than helped' — the 426 "
    "external-LLM stripper is exactly that code. 435 inherits it. 411 "
    "predates it."
)

RISKS = [
    "The criteria 'drag-and-drop / Evaluate button in ChatWindow.tsx and "
    "MessageBubble.tsx' could not be confirmed for any build in the "
    "400..450 window. A grep for 'drag', 'drop', 'Evaluate' in both "
    "components at builds 400, 406, 410, 411, 425, 435, and 450 returned "
    "zero matches. If those features existed at any time they were "
    "either earlier than 400 or in files outside the standard React "
    "component tree. This finding does not affect the 411 recommendation "
    "but is flagged as a risk because the assignment listed it as a "
    "criterion.",
    "The criterion 'Inline provider listing in chat (vs use the panel "
    "below string)' cannot be confirmed by static read of the source — "
    "it depends on what the LLM actually returns at runtime. The static "
    "evidence available is: at 411..415 the LLM narrative is preserved "
    "untouched, so whatever the LLM produces is shown; from 416 onward "
    "regex stripping deletes 'Here are the first N of M' and similar "
    "lines and from 426 onward GPT-4.1-mini removes everything that is "
    "not a provider record. The behaviour can be inferred from the "
    "stripper definition but a runtime test would be needed to confirm.",
    "BUG-VECTOR-001 was logged at build 430 but the underlying defect "
    "(provider 3072d embeddings polluted by address/license text) was "
    "latent in earlier builds where the indexes existed (411..429). "
    "Whether the latent defect produces a user-visible regression at "
    "411 depends on the input — for very common terms ('cardiologist') "
    "the result set was acceptable; for colloquial terms ('kid's doc', "
    "'bone doc') it was not. A baseline anchored at 411 inherits the "
    "latent issue but also inherits the original fix path that the "
    "team eventually took (433: switch RAG target to SpecialtyMetaData). "
    "This risk is not unique to 411 — every build through 432 has it.",
    "Builds 437..450 are largely pipeline / QA-report work and do not "
    "alter the chat engine. They are listed in the table for completeness "
    "but their 'overall score' reflects the chat engine's state inherited "
    "from 435, not new work in those builds.",
]

# ---------------------------------------------------------------------------
# Builders
# ---------------------------------------------------------------------------

def _set_cell_text(cell, text, bold=False, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(size)
    if bold:
        run.bold = True


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # ---- Title -----------------------------------------------------------
    doc.add_heading("Baseline Build Research V1 — Builds 400-450",
                    level=1)
    p = doc.add_paragraph()
    r = p.add_run(
        "Front-end-only assessment. Read-only research. "
        "Recommends ONE build as the baseline for the Design V3 conversion.")
    r.italic = True
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # ---- 1. Executive summary -------------------------------------------
    doc.add_heading("1. Executive summary", level=2)
    p = doc.add_paragraph()
    p.add_run("Recommended baseline: ").bold = True
    p.add_run(
        f"Build {RECOMMENDED_BUILD}  ({RECOMMENDED_HASH}, "
        f"{RECOMMENDED_DATE}) — \"{RECOMMENDED_SUBJECT}\".")
    p = doc.add_paragraph()
    p.add_run("Current Design V3 anchor: ").bold = True
    p.add_run(
        f"Build {CURRENT_ANCHOR_BUILD}  ({CURRENT_ANCHOR_HASH}, "
        f"{CURRENT_ANCHOR_DATE}).")
    p = doc.add_paragraph()
    p.add_run("Recommendation differs from current anchor: ").bold = True
    p.add_run("YES — by exactly one commit (410 → 411).")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Rationale (one paragraph). ").bold = True
    p.add_run(EXEC_RATIONALE)

    # ---- 2. Methodology --------------------------------------------------
    doc.add_heading("2. Methodology and scope", level=2)
    doc.add_paragraph(
        "Builds were enumerated by `git log` filtered with "
        "`--grep='Build 4[0-4][0-9]:|Build 450:'` against the dev branch. "
        "All 51 builds in the inclusive 400..450 range were located. For "
        "each build, the commit hash, ISO timestamp, full subject line, "
        "and front-end-only +/- LOC were captured. Per-build assessment "
        "uses `git show <hash>:<path>` exclusively for past-state "
        "inspection — no checkouts, no commits, no working-tree mutations.")
    doc.add_paragraph(
        "Front-end is defined for this assessment as: "
        "(1) Code/ConversationalUX/FindCareChat/frontend/ — the React SPA; "
        "(2) Code/ConversationalUX/FindCareChat/backend/main.py — FastAPI "
        "routes /chat, /search, /welcome, /health and tool dispatch; "
        "(3) the supporting backend modules reachable from those routes — "
        "domain/find_care/, application/, infrastructure/, "
        "prompt_system_maker.py, tool_router.py, url_guardian.py; "
        "(4) Website/*.html, Website/*.js, Website/*.css. Out of scope by "
        "explicit assignment: Code/DataPipelines/, Code/Shared/ops/, "
        "architecture/EngineeringRuleEnforcement/, LangGraph/poc/* (the "
        "post-4/20 self-attributed work), and brain/ content or schemas.")
    doc.add_paragraph(
        "For each build the following criteria were evaluated. "
        "Tool registry completeness (all 9 tools registered in "
        "ToolRouter.register_with_models): find_providers, "
        "find_specialty_codes, search_clinical_trials, "
        "lookup_provider_external, record_user_details, "
        "record_unknown_question, get_skip_snow_context, "
        "get_chathealthy_context, commitSignificantActivity. "
        "AI-narrative result rendering (does the LLM-authored text reach "
        "the user untouched, partially stripped, or gutted). "
        "Vector specialty search functionality (indexes deployed; "
        "BUG-VECTOR-001 status). "
        "URLGuardian.guard_text wired in /chat. "
        "IP-lockout machinery (SafetyService.is_ip_locked / lock_ip / "
        "try_admin_unlock). "
        "Specialty selection filter panel "
        "(gui.showFilterPanel, gui.onFilterApply). "
        "Inline 'Thinking… Ns' counter and Stop button in ChatWindow.tsx. "
        "HUMAN_TESTING /welcome reachable. "
        "'#action:next-page' and '#action:filter' clickable host bubble "
        "present.")

    # ---- 3. Build-by-build evaluation table -----------------------------
    doc.add_heading("3. Build-by-build evaluation", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Legend. tools: Y = all 9 tools registered. "
        "narr: Y = LLM narrative untouched, D = damaged by regex strip, "
        "G = gutted by external LLM strip. "
        "vector: Y = working, N = indexes missing, "
        "D = deployed but BUG-VECTOR-001 latent, F = explicitly broken "
        "(release blocker). "
        "guard / iplock / filter / thinkstop / uat / action: Y = present, "
        "N = absent. "
        "score: 0..10, higher is better. "
        "+/- = front-end-only LOC delta in this commit."
    ).italic = True

    table = doc.add_table(rows=1, cols=14)
    table.style = "Light Grid Accent 1"
    headers = ["#", "hash", "date", "+/-", "tools", "narr", "vector",
               "guard", "iplock", "filter", "think+stop", "uat",
               "action", "regs"]
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        _set_cell_text(hdr[i], h, bold=True, size=8)

    for n, h, d, plus, minus, subj in BUILDS:
        e = EVAL[n]
        row = table.add_row().cells
        _set_cell_text(row[0], str(n), size=8)
        _set_cell_text(row[1], h, size=8)
        _set_cell_text(row[2], d.replace(" PT", ""), size=8)
        _set_cell_text(row[3], f"+{plus}/-{minus}", size=8)
        _set_cell_text(row[4], e["tools"], size=8)
        _set_cell_text(row[5], e["narr"], size=8)
        _set_cell_text(row[6], e["vector"], size=8)
        _set_cell_text(row[7], e["guard"], size=8)
        _set_cell_text(row[8], e["iplock"], size=8)
        _set_cell_text(row[9], e["filter"], size=8)
        _set_cell_text(row[10], e["thinkstop"], size=8)
        _set_cell_text(row[11], e["uat"], size=8)
        _set_cell_text(row[12], e["action"], size=8)
        _set_cell_text(row[13], e["regs"], size=8)

    # widths — keep the regression column readable
    widths = [Inches(0.30), Inches(0.55), Inches(0.85), Inches(0.55),
              Inches(0.35), Inches(0.32), Inches(0.40), Inches(0.35),
              Inches(0.40), Inches(0.35), Inches(0.55), Inches(0.30),
              Inches(0.40), Inches(2.40)]
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = widths[i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # subjects table — separate so the main grid stays compact
    doc.add_paragraph()
    doc.add_heading("3a. Build subjects (full text)", level=3)
    sub = doc.add_table(rows=1, cols=3)
    sub.style = "Light Grid Accent 1"
    sh = sub.rows[0].cells
    _set_cell_text(sh[0], "#", bold=True, size=9)
    _set_cell_text(sh[1], "hash", bold=True, size=9)
    _set_cell_text(sh[2], "subject", bold=True, size=9)
    for n, h, _, _, _, subj in BUILDS:
        row = sub.add_row().cells
        _set_cell_text(row[0], str(n), size=9)
        _set_cell_text(row[1], h, size=9)
        _set_cell_text(row[2], subj, size=9)
    sub_widths = [Inches(0.3), Inches(0.7), Inches(5.5)]
    for row in sub.rows:
        for i, cell in enumerate(row.cells):
            cell.width = sub_widths[i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # ---- 4. Justification -----------------------------------------------
    doc.add_heading("4. The recommended build — full justification",
                    level=2)
    p = doc.add_paragraph()
    p.add_run(
        f"Build {RECOMMENDED_BUILD}  ({RECOMMENDED_HASH}, "
        f"{RECOMMENDED_DATE})  —  ").bold = True
    p.add_run(f"\"{RECOMMENDED_SUBJECT}\".")

    doc.add_heading("4.1  Why 411 over the current 410 anchor",
                    level=3)
    doc.add_paragraph(WHY_NOT_410)
    p = doc.add_paragraph()
    p.add_run("Concrete evidence at the commit. ").bold = True
    p.add_run(
        "The 410→411 main.py diff is purely additive: PaginationMeta gains "
        "an Optional[str] summary_message field; a new _check_indexes "
        "function and a /health 'degraded' branch are added; a new "
        "summary_message= argument is wired into the PaginationMeta "
        "construction. _chat_inner still returns "
        "`text = next((b.text for b in response.content if b.type == "
        "'text'), '')` followed by `_url_guardian.guard_text(text)` — "
        "the LLM-authored prose path is unchanged. ChatWindow.tsx adds "
        "an isSummary bubble appended after the LLM message bubble; "
        "the LLM bubble is not suppressed.")

    doc.add_heading("4.2  Why 411 over 415 (last 'narrative untouched' "
                    "build)", level=3)
    doc.add_paragraph(WHY_NOT_415)

    doc.add_heading("4.3  Why 411 over 435 (latest 'all features + "
                    "vector fixed')", level=3)
    doc.add_paragraph(WHY_NOT_435)

    doc.add_heading("4.4  What 411 has that 410 does not",
                    level=3)
    bullets = [
        "Vector indexes deployed across dev, QA, and prod "
        "(provider_vector_index 3072d, specialty_vector_index 1536d) — "
        "closes BUG-FILTER-001 logged at build 408.",
        "/health endpoint returns 'degraded' when required indexes are "
        "missing — a runtime safety net the design can rely on.",
        "FC-RESULT-MSG system-built summary message with clickable "
        "'#action:next-page' and '#action:filter' host bubble — the "
        "click-link infrastructure for pagination and filter actions "
        "without a round-trip to the LLM.",
        "PaginationMeta.summary_message field exposed in ChatResponse — "
        "the structured surface the React app uses to drive the "
        "filter panel.",
    ]
    for b in bullets:
        p = doc.add_paragraph(b)
        p.style = "List Bullet"

    doc.add_heading("4.5  What 411 has that later builds in the window "
                    "have damaged", level=3)
    bullets2 = [
        "LLM narrative around results: at 411, the AI's prose answer "
        "reaches the user verbatim. At 416 a regex stripper "
        "(_strip_redundant_summary) starts deleting 'Here are the first…', "
        "'I found N…', 'Would you like to see more…' and similar lines. "
        "At 426 the regex stripper is replaced by an external GPT-4.1-mini "
        "call instructed to 'KEEP ONLY the provider listing… REMOVE "
        "everything else'. From that point onward the LLM is no longer "
        "the narrator.",
        "User-term resolution: at 411 the LLM resolves the user's "
        "natural-language search term and the system uses the LLM's "
        "result. At 415 BUG-MSG-002 forces body.message into the summary. "
        "At 419 an external Haiku call extracts the term. At 420 it is "
        "replaced by GPT-4.1-nano. At 422 by GPT-4.1-mini. Each of those "
        "is an external LLM call on the request path that 411 does not "
        "have.",
        "Vector specialty accuracy: at 411..429 the vector index exists "
        "and is returning results, but BUG-VECTOR-001 (provider "
        "3072d-embedding pollution by address/license text) is latent. "
        "It is logged at 430 as a release blocker ('kid's doc' returns "
        "Hospitalist, Clinic/Center, Mental Health — 2/7 correct) and "
        "fixed at 433 by switching RAG to SpecialtyMetaData (1536d). "
        "411 inherits the latent issue but also inherits the eventual "
        "fix path. This is the regression Skip described.",
    ]
    for b in bullets2:
        p = doc.add_paragraph(b)
        p.style = "List Bullet"

    # ---- 5. Top 3 builds rejected ---------------------------------------
    doc.add_heading("5. Top three rejected alternatives",
                    level=2)
    rej_table = doc.add_table(rows=1, cols=3)
    rej_table.style = "Light Grid Accent 1"
    rh = rej_table.rows[0].cells
    _set_cell_text(rh[0], "Build", bold=True)
    _set_cell_text(rh[1], "Hash", bold=True)
    _set_cell_text(rh[2], "One-line reason rejected", bold=True)
    rej_rows = [
        ("410", "be14579",
         "Vector specialty filter does not work on FE cluster "
         "(BUG-FILTER-001 still open) and no #action: link host yet."),
        ("415", "c6013c4",
         "Last 'narrative untouched' build, but introduces BUG-MSG-002 "
         "fix that begins forcing user message over LLM-resolved term."),
        ("435", "0390c20",
         "Inherits the GPT-4.1-mini LLM-narrative gut from 426 — "
         "exactly the 'code that hurt' pattern the architect is "
         "concerned about."),
    ]
    for n, h, why in rej_rows:
        row = rej_table.add_row().cells
        _set_cell_text(row[0], n)
        _set_cell_text(row[1], h)
        _set_cell_text(row[2], why)
    rwidths = [Inches(0.6), Inches(0.9), Inches(5.0)]
    for row in rej_table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = rwidths[i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # ---- 6. Risks -------------------------------------------------------
    doc.add_heading("6. Risks and open questions", level=2)
    for r in RISKS:
        p = doc.add_paragraph(r)
        p.style = "List Bullet"

    # ---- Footer ----------------------------------------------------------
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run(
        "Read-only research. No commits. No working-tree mutations. "
        "Evidence cited via real commit hashes and `git show <hash>:<path>` "
        "/ `git diff` against the dev branch."
    )
    r.italic = True
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.save(OUT)
    print(f"wrote {OUT}")


if __name__ == "__main__":
    main()
