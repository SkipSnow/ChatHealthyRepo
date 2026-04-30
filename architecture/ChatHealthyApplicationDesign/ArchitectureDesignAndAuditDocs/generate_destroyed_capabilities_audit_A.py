#!/usr/bin/env python3
# Copyright (c) 2026 ChatHealthy.ai LLC. All rights reserved.
# Licensed under the FindCare Evaluation License (FEL-1.0).
#
# Sibling generator for destroyed-capabilities-audit-A.docx
#
# Read-only audit. Compares the latest pre-2026-04-03 commit's front-end
# surface (FastAPI handlers that served the React UI; the React SPA itself;
# Website/index.html and other Website/*.html pages) against the current
# working tree, then filters out anything covered by a V2 BR.
#
# Pre-event commit: be14579a9200ba23227181b716589be6267ef63b (2026-04-02 23:48 PT)
# Identified by: git log --before="2026-04-03 00:00:00" --pretty=format:"%H %ad" --date=iso -1
#
# V2 BR source: architecture/.../design-V2.docx (table 1, BR rows).
#
# Inclusion test:
#   1. Behavior was visible at the pre-event commit on a front-end surface.
#   2. Behavior is no longer present in the current working tree.
#   3. No V2 BR clearly describes the user-facing behavior.
#
# Exclusion examples (out of scope by hard constraint):
#   - LangGraph/poc/* artifacts (post-4/20 by self-attribution).
#   - TR-level concerns ($vectorSearch, embedding model, framework version).
#   - Back-end services not directly serving a user-observable behavior.
#   - Capabilities the V2 BRs explicitly recover (FULL or PARTIAL).

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL

OUT = (
    r"c:/chatHealthy/findCare/architecture/ChatHealthyApplicationDesign/"
    r"ArchitectureDesignAndAuditDocs/destroyed-capabilities-audit-A.docx"
)

PRE_EVENT_COMMIT = "be14579a9200ba23227181b716589be6267ef63b"
PRE_EVENT_DATE   = "2026-04-02 23:48 PT (Build 410)"

# ---------------------------------------------------------------------------
# The list — destroyed pre-event front-end capabilities NOT covered by V2 BR.
# ---------------------------------------------------------------------------
ROWS = [
    {
        "n": 1,
        "capability": (
            "Inline 'Thinking… Ns' indicator with inline Stop button rendered "
            "inside the chat conversation. Pre-event the React app showed an "
            "assistant-style bubble with a live counter ('Thinking… 7s') and a "
            "'Stop' button next to it that the user clicked to dismiss the "
            "thinking spinner without abandoning the request. Today the React "
            "app only postMessages the timer to the parent control frame; "
            "when the React app is loaded outside the index.html parent "
            "(direct hf.space URL, or any host that does not forward "
            "gui:timer messages to a control frame), the user sees no thinking "
            "indicator and has no Stop button at all."
        ),
        "pre_evidence": (
            "Code/ConversationalUX/FindCareChat/frontend/src/components/"
            "ChatWindow.tsx @be14579a, lines 416-446 (showThinking JSX block "
            "rendering 'Thinking… {thinkSeconds}s' and 'Stop' button)."
        ),
        "current_evidence": (
            "Code/ConversationalUX/FindCareChat/frontend/src/components/"
            "ChatWindow.tsx, line 667 ('{/* Timer + Stop moved to parent "
            "control frame via gui:timer postMessage */}'). The inline JSX "
            "block was removed; only postMessage at lines 417-421 remains."
        ),
    },
    {
        "n": 2,
        "capability": (
            "UAT-mode welcome message in chat. Pre-event the /welcome "
            "endpoint returned a UAT-specific greeting (build_uat_welcome) "
            "when the HUMAN_TESTING env flag was set, surfacing the active "
            "UAT scenario list to the human tester directly inside the chat "
            "window's first message. Today /welcome unconditionally returns "
            "the static WELCOME_MESSAGE; the HUMAN_TESTING flag is loaded "
            "but never consulted for the welcome path. UAT testers no "
            "longer see the per-build UAT brief in the chat."
        ),
        "pre_evidence": (
            "Code/ConversationalUX/FindCareChat/backend/main.py @be14579a, "
            "line ~232: '@app.get(\"/welcome\")\\ndef welcome():\\n    "
            "return {\"message\": _build_test_welcome() if _HUMAN_TESTING "
            "else WELCOME_MESSAGE}'."
        ),
        "current_evidence": (
            "Code/ConversationalUX/FindCareChat/backend/main.py, lines "
            "572-574: '@app.get(\"/welcome\")\\ndef welcome():\\n    return "
            "{\"message\": WELCOME_MESSAGE}'. The HUMAN_TESTING branch was "
            "deleted; _HUMAN_TESTING is dead-loaded at lines 64-65."
        ),
    },
]

# ---------------------------------------------------------------------------
# V2-BR coverage rationale (kept alongside ROWS for the human reader; not in
# the deliverable's body per the 'no methodology section' constraint, but
# encoded in the exec summary so the architect can see the test that was run).
# ---------------------------------------------------------------------------
# Row 1 — V2 BRs scanned: BR-01 (conversational find-care UX) describes the
#   ranked-providers experience but says nothing about a thinking-indicator
#   or a user-facing 'Stop' affordance during model thinking. BR-13
#   (performance budgets) is a back-end SLO, not a user control. No BR
#   restores the inline thinking/Stop UI.
# Row 2 — V2 BRs scanned: BR-19 / BR-34 (deferred-work ledger) build a
#   back-end registry, not a chat welcome surface. BR-05 (10 static pages)
#   is unrelated. No BR restores the UAT-mode welcome path.

# ---------------------------------------------------------------------------
# Build the docx
# ---------------------------------------------------------------------------

def main():
    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # Title
    title = doc.add_heading("Destroyed Front-End Capabilities Audit (A)", level=1)

    # Executive summary — exactly three lines.
    p = doc.add_paragraph()
    p.add_run(f"Pre-event commit: {PRE_EVENT_COMMIT}").bold = True
    p.add_run(f"  ({PRE_EVENT_DATE})")
    p = doc.add_paragraph()
    p.add_run(
        f"Destroyed front-end capabilities NOT covered by any V2 BR: "
    )
    p.add_run(f"{len(ROWS)}").bold = True
    p = doc.add_paragraph(
        "Scope: front-end surfaces only (FindCareChat/backend/main.py routes, "
        "url_guardian, FindCareChat/frontend React SPA, Website/*.html). "
        "LangGraph/poc/* explicitly excluded (post-4/20 by self-attribution). "
        "Read-only audit; no commits, no code changes."
    )

    # Body — single table.
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "#"
    hdr[1].text = "Destroyed capability (user-facing description)"
    hdr[2].text = f"Pre-event evidence (file:line at {PRE_EVENT_COMMIT[:10]})"
    hdr[3].text = "Current evidence of absence (file:line in current tree)"
    for cell in hdr:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

    for r in ROWS:
        row = table.add_row().cells
        row[0].text = str(r["n"])
        row[1].text = r["capability"]
        row[2].text = r["pre_evidence"]
        row[3].text = r["current_evidence"]

    # Column widths — stop the # column from being unreadably narrow,
    # mirror Skip's note from V2 R66 about column-A readability.
    widths = [Inches(0.4), Inches(3.4), Inches(1.9), Inches(1.9)]
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = widths[i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # Footer line: V2 coverage tested
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("V2 BR coverage tested against ").italic = True
    r = p.add_run(
        "architecture/ChatHealthyApplicationDesign/"
        "ArchitectureDesignAndAuditDocs/design-V2.docx"
    )
    r.italic = True
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    p.add_run(
        "  (table 1, all 35 BR rows). No BR clearly describes the "
        "user-facing behavior of either listed item."
    ).italic = True

    doc.save(OUT)
    print(f"wrote {OUT}")


if __name__ == "__main__":
    main()
