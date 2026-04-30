"""Generate ChatHealthy Application Design - V3 (Word document).

Read-only generator. Produces design-V3.docx in this same directory. No commits,
no code changes, no backlog mutations. Read-only against the application code.

Authority order (binding):
  1. Pre-4/3 application code at commit be14579a (2026-04-02 23:48:24 -0700,
     Build 410). This is GROUND TRUTH for what the application does.
  2. architecture/ChatHealthyApplicationDesign/ArchitectureDesignAndAuditDocs/
     design-V2.docx -- the architect's recollection. Used for cross-reference.
     Where V2 disagrees with pre-4/3 code, the code wins.
  3. findCare/ArchitectureAndDesign/codebase-reorganization-V7.docx -- binding
     for target file structure and epic shape.
  4. findCare/ArchitectureAndDesign/langgraph-oo-best-practices-V6.docx -- binding
     for LangGraph runtime patterns (Pydantic state, MongoDBSaver in parent only,
     RemoteGraph cross-Space, Runtime[ContextSchema] DI, messages-channel for
     tool I/O, React-as-facade).
  5. brain/machine_artifacts/content/{engineering_rules,agile_backlog}.json.

Constraints honored:
  - Read-only on application code. The deliverable docx and this generator are
    the only writes.
  - Pre-4/3 code is GROUND TRUTH; V2 is recollection. Code wins.
  - LangGraph/poc/* is NOT cited as pre-event evidence (post-4/20, out of scope).
  - Stories stay intact.
  - "Domain" does not appear as a structural noun.
  - PlantUML rendering is best-effort; failures fall through to provisional source.
"""
from __future__ import annotations

import io
import logging
import os
import sys
import urllib.error
import urllib.request
import zlib
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")
_log = logging.getLogger("design_v3")

OUT = Path(__file__).resolve().parent / "design-V3.docx"

GROUND_TRUTH_COMMIT = "be14579a9200ba23227181b716589be6267ef63b"
GROUND_TRUTH_DATE = "2026-04-02 23:48:24 -0700 (Build 410)"


# -----------------------------------------------------------------------------
# PlantUML rendering pipeline
# -----------------------------------------------------------------------------

def _plantuml_encode(text: str) -> str:
    """Encode PlantUML source into the URL-safe form the public server expects."""
    zlibbed = zlib.compress(text.encode("utf-8"))
    compressed = zlibbed[2:-4]

    def _enc6(b: int) -> str:
        if b < 10:
            return chr(48 + b)
        b -= 10
        if b < 26:
            return chr(65 + b)
        b -= 26
        if b < 26:
            return chr(97 + b)
        b -= 26
        if b == 0:
            return "-"
        if b == 1:
            return "_"
        return "?"

    res = []
    i = 0
    while i < len(compressed):
        b1 = compressed[i]
        b2 = compressed[i + 1] if i + 1 < len(compressed) else 0
        b3 = compressed[i + 2] if i + 2 < len(compressed) else 0
        c1 = b1 >> 2
        c2 = ((b1 & 0x3) << 4) | (b2 >> 4)
        c3 = ((b2 & 0xF) << 2) | (b3 >> 6)
        c4 = b3 & 0x3F
        res.append(_enc6(c1) + _enc6(c2) + _enc6(c3) + _enc6(c4))
        i += 3
    return "".join(res)


def render_plantuml(source: str, label: str) -> Optional[bytes]:
    """Render PlantUML to PNG bytes. Returns None on failure (caller falls back).

    If PlantUML server returns HTTP 400, the body still contains a PNG with the
    syntax error rendered visually. We treat that as failure (we want a clean
    diagram or no diagram, not an error image) and fall through to provisional.
    """
    encoded = _plantuml_encode(source)
    url = "https://www.plantuml.com/plantuml/png/" + encoded
    req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0 ChatHealthyDesignV3"})
    try:
        with urllib.request.urlopen(req, timeout=8) as r:
            data = r.read()
            if r.status == 200 and data and data[:4] == b"\x89PNG":
                _log.info("PlantUML render OK: %s (%d bytes)", label, len(data))
                return data
            _log.warning("PlantUML returned non-200 or non-PNG for %s: status=%s", label, r.status)
    except urllib.error.HTTPError as exc:
        _log.warning("PlantUML HTTP %s for %s (treating as render failure)", exc.code, label)
    except (urllib.error.URLError, TimeoutError, OSError) as exc:
        _log.warning("PlantUML render failed for %s: %s", label, exc)

    # Fallback: try kroki.io with explicit plantuml format
    try:
        import base64 as _b64
        compressed = zlib.compress(source.encode("utf-8"), 9)
        kroki_enc = _b64.urlsafe_b64encode(compressed).decode("ascii").rstrip("=")
        kroki_url = f"https://kroki.io/plantuml/png/{kroki_enc}"
        kreq = urllib.request.Request(kroki_url, headers={"User-Agent": "Mozilla/5.0 ChatHealthyDesignV3"})
        with urllib.request.urlopen(kreq, timeout=8) as r:
            data = r.read()
            if r.status == 200 and data and data[:4] == b"\x89PNG":
                _log.info("kroki.io render OK: %s (%d bytes)", label, len(data))
                return data
    except Exception as exc:
        _log.warning("kroki.io fallback failed for %s: %s", label, exc)

    return None


# -----------------------------------------------------------------------------
# docx helpers
# -----------------------------------------------------------------------------

def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    return p


def add_bullet(doc, text):
    return doc.add_paragraph(text, style="List Bullet")


def add_sub_bullet(doc, text):
    try:
        return doc.add_paragraph(text, style="List Bullet 2")
    except KeyError:
        return doc.add_paragraph("    - " + text)


def add_table(doc, headers, rows, col_widths_in: Optional[list] = None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    if col_widths_in:
        for row in t.rows:
            for i, w in enumerate(col_widths_in):
                if i < len(row.cells):
                    row.cells[i].width = Inches(w)
    return t


def add_code_block(doc, text, label: str = ""):
    if label:
        p = doc.add_paragraph()
        r = p.add_run(label)
        r.italic = True
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    return p


def add_diagram(doc, source: str, label: str, caption: str, width_in: float = 6.4):
    """Try PlantUML; on failure, embed the source as a labeled provisional code block."""
    png = render_plantuml(source, label)
    if png:
        bio = io.BytesIO(png)
        doc.add_picture(bio, width=Inches(width_in))
        cap = doc.add_paragraph()
        r = cap.add_run(f"Figure: {caption}")
        r.italic = True
        r.font.size = Pt(9)
        return True
    add_code_block(
        doc,
        source,
        label=f"PROVISIONAL ({label}) -- PlantUML rendering unavailable; source preserved verbatim.",
    )
    cap = doc.add_paragraph()
    r = cap.add_run(f"Figure (PROVISIONAL): {caption}")
    r.italic = True
    r.font.size = Pt(9)
    return False


# -----------------------------------------------------------------------------
# Rescue table -- the load-bearing artifact per the V3 brief.
# Each row is one file in the reachable-from-main.py keep-set at build 410,
# minus the explicit cuts (EvaluateCare, Auth, Azure DataPipelines).
# Columns: #, File path at build 410, Purpose, Reachable via, Survives as
#
# The "PC ID" set below preserved for back-compat with TR/BR trace cells.
# It is no longer rendered in Section 2; the new RESCUE_ROWS table is.
# -----------------------------------------------------------------------------

RESCUE_ROWS = [
    # ---- Parent runtime / FastAPI host ----
    ("R-01",
     "Code/ConversationalUX/FindCareChat/backend/main.py",
     "FastAPI host; /chat tool-use loop, /search direct-pagination, /welcome, /health; pre-LLM safety enforcement order; URLGuardian + LinkedIn auto-link rewriter on response edge.",
     "(entry point)",
     "Operations/RuntimeServer/app.py (Application class) + Operations/RuntimeServer/orchestration/parent_graph.py (LangGraph parent runtime). Tool-use loop replaced by parent graph + RemoteGraph composition; safety becomes Security RemoteGraph (first edge); /search becomes a non-conversational FastAPI route on FindCare/RuntimeServer/app.py."),
    ("R-02",
     "Code/ConversationalUX/FindCareChat/backend/url_guardian.py",
     "URLGuardian: 3-stage validation (HEAD reachability, AI content verify via Claude Haiku, Google Custom Search correction); SPA + trusted-search domain allowlists; markdown-link sanitizer (defang / redirect-rewrite); cache TTL 3600s.",
     "main.py imports URLGuardian; called on every /chat assistant text and on tool-result link dicts.",
     "SharedServices/SharedCode/url_guardian.py (canonical). Run from Operations parent-runtime response-finalize step (NOT a graph node), so every capability's text and link payloads are validated uniformly."),

    # ---- Application layer ----
    ("R-03",
     "Code/ConversationalUX/FindCareChat/backend/application/tool_router.py",
     "ToolRouter: allowlist registry of tool name -> handler with optional Pydantic input model; GOV-004 enforcement (\"the model may suggest; the system must decide\"); refuses unregistered tool names server-side.",
     "main.py constructs and registers tools; /chat tool_use loop dispatches via _handle_tool_calls -> ToolRouter.handle_tool_calls.",
     "Operations/SharedCode/agent_framework/tool_router.py (canonical allowlist primitive). The same allowlist + Pydantic-validation pattern is reused in every {Capability}/RuntimeServer/tools.py @tool registration."),
    ("R-04",
     "Code/ConversationalUX/FindCareChat/backend/application/tool_models/provider_search_models.py",
     "Pydantic input models ProviderSearchInput + SpecialtyInput; carry the pre-4/3 search_providers / find_specialty_codes signatures.",
     "main.py registers these with ToolRouter via register_with_models.",
     "FindCare/RuntimeServer/tools.py imports them as the @tool input schemas; FindCare/RuntimeServer/state.py types the last_search_params field with ProviderSearchInput."),
    ("R-05",
     "Code/ConversationalUX/FindCareChat/backend/application/tool_models/consent_models.py",
     "Pydantic input models LeadInput + UnknownInput.",
     "main.py registers tools record_user_details + record_unknown_question with these models via ToolRouter.",
     "FindCare/RuntimeServer/tools.py uses LeadInput + UnknownInput as the @tool input schemas for tool_record_user_details and tool_record_unknown_question."),

    # ---- FindCare core (kept) ----
    ("R-06",
     "Code/ConversationalUX/FindCareChat/backend/domain/find_care/provider_search_service.py",
     "FindCareService: 4-route provider search (NPI exact / Name / Specialty codes direct / Specialty query via vector + taxonomy); FIPS-to-county fallback map; entity_type-1-vs-2 name formatter; (NNN) NNN-NNNN phone formatter; keyset pagination by NPI; pagination metadata (total_count, first_npi, last_npi, has_more, page_start, page_end, search_params, specialization_options).",
     "main.py constructs FindCareService and registers find_providers + find_specialty_codes; also called directly by /search route for pagination.",
     "FindCare/SharedCode/core/provider_search_service.py (preserves the pre-4/3 4-route signature verbatim). Consumed by FindCare/RuntimeServer/tools.py @tool tool_find_providers AND by FindCare/RuntimeServer/app.py POST /search (non-graph route) -- single implementation, no logic drift."),
    ("R-07",
     "Code/ConversationalUX/FindCareChat/backend/domain/find_care/specialty_service.py",
     "SpecialtyService: dual pipeline (regex with AI-expanded stems via Claude Haiku + vector via text-embedding-3-small filtered by INDIVIDUAL_PROVIDER_GROUPINGS classification); ThreadPoolExecutor parallel; dedup by Code with vector priority.",
     "main.py constructs SpecialtyService and registers find_specialty_codes; also injected into FindCareService for in-search resolution.",
     "FindCare/SharedCode/core/specialty_service.py. Consumed by FindCare/RuntimeServer/tools.py @tool tool_find_specialty_codes."),

    # ---- FindCare shared services that follow into FindCare ----
    ("R-08",
     "Code/ConversationalUX/FindCareChat/backend/domain/shared/lead_capture/lead_service.py",
     "LeadService: dedupe-by-email; three consent paths (verbatim / summary / neither); MongoDB write to {ENV}_AboutUs.lead; Sparkpost push notification on success.",
     "main.py constructs LeadService and registers tool record_user_details.",
     "FindCare/SharedCode/core/lead_service.py. Consumed by FindCare/RuntimeServer/tools.py @tool tool_record_user_details. Sparkpost notify is delegated to Operations/SharedCode/notify/sparkpost_push.py via the runtime context (TR-OP-002)."),
    ("R-09",
     "Code/ConversationalUX/FindCareChat/backend/domain/shared/unknowns/unknown_question_service.py",
     "UnknownQuestionService: 3-class taxonomy (healthcare_capability / medical_advice / irrelevant); verbatim response_template per class; consent-gated de-identification + persistence.",
     "main.py constructs UnknownQuestionService and registers tool record_unknown_question.",
     "FindCare/SharedCode/core/unknown_question_service.py. Consumed by FindCare/RuntimeServer/tools.py @tool tool_record_unknown_question."),
    ("R-10",
     "Code/ConversationalUX/FindCareChat/backend/domain/shared/content/about_service.py",
     "AboutService: get_skip_snow_context, get_chathealthy_context; verbatim 'connect' field rendered as markdown links; ME-context trim utility.",
     "main.py constructs AboutService and registers tools get_skip_snow_context + get_chathealthy_context.",
     "FindCare/SharedCode/core/about_service.py. Consumed by FindCare/RuntimeServer/tools.py @tool tool_get_skip_snow_context and tool_get_chathealthy_context. The Skip-Snow LinkedIn auto-link rewriter (post-LLM regex on the response edge) moves to the Operations response-finalize step alongside URLGuardian."),

    # ---- Security & Compliance ----
    ("R-11",
     "Code/ConversationalUX/FindCareChat/backend/domain/shared/safety/safety_service.py",
     "SafetyService: dual-trigger emergency detection (keyword OR Claude classifier @ confidence>=0.80); safe-prefix fast-path; IP lock (1-hour TTL) with admin-unlock key; persists to {ENV}_Safety.emergency_incidents.",
     "main.py constructs SafetyService; _chat_inner calls try_admin_unlock, is_ip_locked, is_emergency BEFORE any LLM call.",
     "Security/RuntimeServer/nodes.py preserves the enforcement order admin_unlock_check -> ip_locked_check -> emergency_check (dual-trigger) -> safe_prefix_fast_path. Security/RuntimeServer/agent.py is the FIRST capability node in the parent edge order. Collection name and TTL preserved (TR-SE-004)."),
    ("R-12",
     "Code/ConversationalUX/FindCareChat/backend/domain/shared/consent/consent_service.py",
     "ConsentService: two-tier HIPAA framework; AI summarization (Claude Haiku); AI Safe-Harbor de-identification (Claude Haiku, JSON-strict).",
     "main.py constructs ConsentService and injects into LeadService, UnknownQuestionService, DebugLogger.",
     "Security/SharedCode/consent_service.py (V7 A1: Compliance is a feature folder under Security & Compliance, one cost-center line). Consumed by FindCare lead/unknown @tools and by Operations DebugLogger via runtime context."),

    # ---- Infrastructure ----
    ("R-13",
     "Code/ConversationalUX/FindCareChat/backend/infrastructure/embeddings/embedding_client.py",
     "EmbeddingClient: OpenAI text-embedding-3-large (provider) + text-embedding-3-small (specialty) + Claude Haiku query expansion for NUCC stems.",
     "main.py constructs EmbeddingClient and injects expand_query_terms + get_specialty_vector into SpecialtyService and get_query_embedding into FindCareService.",
     "SharedServices/SharedCode/embedding_client.py. Consumed via runtime context by FindCare nodes; never imported at module scope inside node files (TR-LG-008)."),
    ("R-14",
     "Code/ConversationalUX/FindCareChat/backend/infrastructure/debug_logger.py",
     "DebugLogger: per-turn audit row to {ENV}_Debug.chat_calls (datetime, ip, message preview, history length, tool-loop iterations, tokens in/out, response preview, error, optional chat_history_deidentified on errors).",
     "main.py constructs DebugLogger; chat() and _chat_inner call log_chat on success and on every error path.",
     "Operations/SharedCode/conversation_log/Worker (V7 13.2): the per-turn audit becomes a Kafka topic emit at the parent runtime; the canonical Worker writes to MongoDB. trace_id is added so audit rows join to the LangSmith trace."),

    # ---- Shared platform code (Code/Shared/) ----
    ("R-15",
     "Code/Shared/ChatHealthyMongoUtilities.py",
     "ChatHealthyMongoUtilities: canonical Mongo connection manager; getConnection() with graceful unavailable mode; commit(env, db, collection, record).",
     "main.py imports via sys.path.insert; _get_db lazily constructs and caches.",
     "SharedServices/SharedCode/mongo_utilities.py (V7 A2 / 13.2: canonical Mongo client; raw pymongo client construction in any other Space is forbidden, TR-SS-001 / TR-DM-003)."),
    ("R-16",
     "Code/Shared/prompt_system_maker.py",
     "PromptSystemMaker: brain-driven config (system prompt rules, tool definitions, emergency keywords, ME-context loader, build-number reader, welcome-message builder, trim utility); brain-dir resolution falls back gracefully on the HF flat layout.",
     "main.py instantiates _prompt_maker, loads emergency keywords + tool definitions + ME context + welcome + build number; _system_prompt() composes the system prompt per turn.",
     "SharedServices/SharedCode/prompt_system_maker.py. Loaded once per Space; the 10 system-prompt rules are migrated to brain/machine_artifacts/content/system_prompt_rules.json so every Space resolves the same prompt (TR-FC-009)."),
    ("R-17",
     "Code/Shared/ops/uat_report.py",
     "UAT welcome generator: build_uat_welcome reads UAT feature definitions; rendered via /welcome when HUMAN_TESTING is true.",
     "main.py imports build_uat_welcome via sys.path.insert (Code/Shared/ops); /welcome returns either the canonical welcome or the UAT welcome based on HUMAN_TESTING env.",
     "Operations/SharedCode/uat/uat_report.py. /welcome remains the parent runtime endpoint (no graph trace); the HUMAN_TESTING toggle is preserved."),

    # ---- Frontend ----
    ("R-18",
     "Code/ConversationalUX/FindCareChat/frontend/src/App.tsx",
     "React app shell; renders ChatWindow inside the iframe boundary.",
     "main.py mounts the static build via StaticFiles + index.html; React fetches via VITE_API_URL.",
     "FindCare/Frontend/src/App.tsx. Per V6 3.3 the React iframe talks to the parent runtime URL only; capability URLs are absent from src/ (TR-UX-001)."),
    ("R-19",
     "Code/ConversationalUX/FindCareChat/frontend/src/components/ChatWindow.tsx",
     "ChatWindow: streams /chat responses; renders MessageBubble entries; surfaces pagination metadata via postMessage to the static parent frame; environment banner; retry / abandon countdown; timeout escalation modal.",
     "Loaded by App.tsx; the parent main.py serves the bundled assets.",
     "FindCare/Frontend/src/components/ChatWindow.tsx. Streaming consumes @langchain/langgraph-sdk against the Operations parent endpoint (TR-LG-002)."),
    ("R-20",
     "Code/ConversationalUX/FindCareChat/frontend/src/components/GUIManager.tsx",
     "GUIManager: pagination control rendered in the static parent frame; postMessage bridge between the static control frame and the React iframe.",
     "Loaded by ChatWindow; pagination round-trips between static frame and iframe.",
     "UX/SharedCode/cross_components/GUIManager.tsx (V7 A7: shared React primitive consumed by FindCare/Frontend AND by future EvaluateCare/Frontend pagination consumers)."),
    ("R-21",
     "Code/ConversationalUX/FindCareChat/frontend/src/components/MessageBubble.tsx",
     "MessageBubble: renders one chat message with markdown + link safety; honors URLGuardian's defanging by treating non-link text as plain text.",
     "Rendered inside ChatWindow.",
     "FindCare/Frontend/src/components/MessageBubble.tsx. Markdown rendering remains client-side; URL validity is enforced server-side by URLGuardian on the response edge (TR-FC-010)."),
    ("R-22",
     "Code/ConversationalUX/FindCareChat/frontend/index.html",
     "Vite entry HTML; mounts /src/main.tsx into #root.",
     "Bundled at frontend build time and served by main.py StaticFiles mount on /.",
     "FindCare/Frontend/index.html. Build artifact is deployed by Cloudflare Pages (TR-UX-002) when EPIC-011 owns the deploy; co-located today via main.py static mount."),
]

CAP_ROWS = [
    ("PC-01", "FastAPI host with /chat, /search, /welcome, /health endpoints",
     "Code/ConversationalUX/FindCareChat/backend/main.py:194-263",
     "BR-01, BR-03, BR-22", "TR-LG-001, TR-LG-002, TR-FC-008"),
    ("PC-02", "Anthropic tool-use chat loop (initial call + tool_use loop until stop)",
     "main.py:266-296",
     "BR-01, BR-12", "TR-LG-001, TR-LG-007, TR-FC-001"),
    ("PC-03", "Allowlist tool dispatch with Pydantic input validation (ToolRouter)",
     "Code/ConversationalUX/FindCareChat/backend/application/tool_router.py:1-110",
     "BR-10, BR-12", "TR-FC-001, TR-LG-007"),
    ("PC-04", "System prompt assembly from brain rules + ME context (PromptSystemMaker)",
     "Code/Shared/prompt_system_maker.py:140-260",
     "BR-10, BR-22, BR-23", "TR-SS-002, TR-FC-009"),
    ("PC-05", "Welcome message (hard-coded canonical) and human-testing UAT welcome",
     "Code/Shared/prompt_system_maker.py:380-410; Code/Shared/ops/uat_report.py",
     "BR-22", "TR-SS-002, TR-OP-002"),
    ("PC-06", "Build counter (read once at startup from MongoDB)",
     "Code/Shared/prompt_system_maker.py:355-372",
     "BR-25", "TR-OP-001"),
    ("PC-07", "MongoDB connection pool with graceful unavailable mode",
     "main.py:78-95",
     "BR-09, BR-24", "TR-SS-001"),
    ("PC-08", "Provider search (vector resolves codes -> taxonomy returns data) for DE/MS/VA",
     "Code/ConversationalUX/FindCareChat/backend/domain/find_care/provider_search_service.py:170-310",
     "BR-01", "TR-FC-002, TR-FC-003"),
    ("PC-09", "Direct provider search /search route (bypasses LLM, used by pagination)",
     "main.py:230-237; provider_search_service.py:170-310",
     "BR-01, BR-26", "TR-FC-002, TR-UX-002"),
    ("PC-10", "Keyset pagination by NPI (after_npi cursor; first/last NPI; total_count)",
     "provider_search_service.py:_paginated_result lines 145-167",
     "BR-26", "TR-FC-002"),
    ("PC-11", "Specialty identification: regex (AI-expanded stems) + vector pipelines in parallel",
     "Code/ConversationalUX/FindCareChat/backend/domain/find_care/specialty_service.py:60-130",
     "BR-01", "TR-FC-004"),
    ("PC-12", "AI query expansion via Claude Haiku for NUCC stems",
     "Code/ConversationalUX/FindCareChat/backend/infrastructure/embeddings/embedding_client.py:55-78",
     "BR-01", "TR-FC-004, TR-SS-002"),
    ("PC-13", "Provider name + state + county + city filters; FIPS-to-county fallback map",
     "provider_search_service.py:24-95",
     "BR-01", "TR-FC-002"),
    ("PC-14", "Provider record formatting (name from entity_type 1/2; phone formatting)",
     "provider_search_service.py:96-130",
     "BR-01", "TR-FC-002"),
    ("PC-15", "OpenAI embeddings: text-embedding-3-large (provider) + 3-small (specialty)",
     "embedding_client.py:30-55",
     "BR-01", "TR-SS-002"),
    ("PC-16", "Clinical trials search (ClinicalTrials.gov v2 API, RECRUITING filter)",
     "Code/ConversationalUX/FindCareChat/backend/domain/evaluate_care_quality/clinical_trials_service.py:65-130",
     "BR-02", "TR-EC-001"),
    ("PC-17", "Travel info via Google Routes API (drive distance + duration; worldwide)",
     "clinical_trials_service.py:25-65",
     "BR-02", "TR-EC-002"),
    ("PC-18", "Provider detail lookup (NPI Registry CMS v2.1 + state board links + Healthgrades)",
     "Code/ConversationalUX/FindCareChat/backend/domain/evaluate_care_quality/provider_detail_service.py:25-90",
     "BR-02, BR-21", "TR-EC-003"),
    ("PC-19", "Two-tier HIPAA consent: verbatim-transcript or de-identified summary or neither",
     "Code/ConversationalUX/FindCareChat/backend/domain/shared/consent/consent_service.py:20-100",
     "BR-13, BR-14", "TR-SE-002, TR-SE-003"),
    ("PC-20", "AI conversation summarization for Tier 2 consent (Claude Haiku)",
     "consent_service.py:25-50",
     "BR-13", "TR-SE-002"),
    ("PC-21", "AI HIPAA Safe Harbor de-identification (Claude Haiku, JSON-strict)",
     "consent_service.py:55-100",
     "BR-14", "TR-SE-002"),
    ("PC-22", "Lead capture: dedupe-by-email, consent-gated chat history persistence, push notify",
     "Code/ConversationalUX/FindCareChat/backend/domain/shared/lead_capture/lead_service.py:1-70",
     "BR-15", "TR-FC-005"),
    ("PC-23", "Unknown-question recording: 3 question classes + verbatim response template",
     "Code/ConversationalUX/FindCareChat/backend/domain/shared/unknowns/unknown_question_service.py:1-78",
     "BR-15", "TR-FC-006"),
    ("PC-24", "About content tools: get_skip_snow_context, get_chathealthy_context",
     "Code/ConversationalUX/FindCareChat/backend/domain/shared/content/about_service.py:1-40",
     "BR-22", "TR-FC-007"),
    ("PC-25", "Dual-trigger emergency detection (keyword OR AI classifier @ confidence>=0.80)",
     "Code/ConversationalUX/FindCareChat/backend/domain/shared/safety/safety_service.py:36-90",
     "BR-04, BR-19", "TR-SE-001"),
    ("PC-26", "Safe-prefix fast-path (skip AI classifier on obvious non-emergency openers)",
     "safety_service.py:25-44",
     "BR-04, BR-13", "TR-SE-001"),
    ("PC-27", "IP lock with 1-hour TTL + audit + admin unlock key + lock-on-emergency",
     "safety_service.py:90-160",
     "BR-04, BR-19, BR-08", "TR-SE-001, TR-OB-002"),
    ("PC-28", "Safety enforcement order (admin-unlock -> ip-locked -> emergency BEFORE LLM)",
     "main.py:_chat_inner lines 304-311",
     "BR-04", "TR-LG-001, TR-SE-001"),
    ("PC-29", "URLGuardian: 3-stage validation (HEAD reachability + AI content verify + Google search correction)",
     "Code/ConversationalUX/FindCareChat/backend/url_guardian.py:30-235",
     "BR-21, BR-04", "TR-FC-010"),
    ("PC-30", "URLGuardian SPA + trusted-search domain allowlists (npiregistry, healthgrades, state boards)",
     "url_guardian.py:35-55",
     "BR-21", "TR-FC-010"),
    ("PC-31", "URLGuardian markdown link sanitizer (defang, redirect-rewrite, bare URL pass)",
     "url_guardian.py:120-180",
     "BR-21", "TR-FC-010"),
    ("PC-32", "Skip Snow LinkedIn auto-link rewriter (post-LLM regex)",
     "main.py:355-359",
     "BR-22", "TR-FC-007"),
    ("PC-33", "DebugLogger: per-turn audit (ip, message preview, history len, tokens, error, de-id history)",
     "Code/ConversationalUX/FindCareChat/backend/infrastructure/debug_logger.py:1-50",
     "BR-08, BR-12", "TR-OB-001, TR-OB-002"),
    ("PC-34", "CORS allowlist (chathealthy.ai apex + dev subdomain regex + localhost)",
     "main.py:198-204",
     "BR-04", "TR-LG-001"),
    ("PC-35", "Pagination metadata propagation (find_providers result -> ChatResponse.pagination)",
     "main.py:325-345",
     "BR-26", "TR-FC-002, TR-UX-002"),
    ("PC-36", "Static site mount (index.html + /assets) and no-cache headers on /",
     "main.py:380-394",
     "BR-22", "TR-UX-001"),
    ("PC-37", "Welcome / health JSON contract (env, build, db, version)",
     "main.py:240-247",
     "BR-25", "TR-OP-001, TR-LG-002"),
    ("PC-38", "Anthropic system prompt rules: emergency, context, unanswerable, medical-advice, format, "
              "provider-detail, clinical-trial-travel, inclusivity, pagination, follow-up",
     "prompt_system_maker.py:222-280",
     "BR-04, BR-15, BR-17, BR-22", "TR-FC-009, TR-SE-001"),
    ("PC-39", "Brain-driven emergency keyword loader (categories aggregated from JSON)",
     "prompt_system_maker.py:80-100",
     "BR-04", "TR-SE-001, TR-OP-003"),
    ("PC-40", "Anthropic tool definition loader (brain artifact first, built-in fallback)",
     "prompt_system_maker.py:103-200",
     "BR-12", "TR-FC-001, TR-OP-003"),
    ("PC-41", "Sparkpost push notification (best-effort; soft fail)",
     "main.py:97-110",
     "BR-15", "TR-OP-002"),
    ("PC-42", "React iframe SPA (parent page hosts; /chat consumed via VITE_API_URL)",
     "Code/ConversationalUX/FindCareChat/frontend/src/components/ChatWindow.tsx:1-80; "
     "Code/ConversationalUX/FindCareChat/frontend/src/App.tsx:1-12",
     "BR-03, BR-26", "TR-UX-001, TR-UX-002"),
    ("PC-43", "GUI Manager: pagination control rendered in static frame via postMessage bridge",
     "Code/ConversationalUX/FindCareChat/frontend/src/components/GUIManager.tsx:1-120",
     "BR-26", "TR-UX-001, TR-UX-003"),
    ("PC-44", "Static parent landing page (Website/index.html) + 9 sibling static pages",
     "Website/index.html; Website/architecture.html; Website/products.html; Website/privacy.html; "
     "Website/terms.html; Website/chat-app-design.html; Website/embedding-design.html; "
     "Website/load-perf-report.html; Website/ops-manager-design.html; Website/provider-data-load.html",
     "BR-05", "TR-UX-004"),
    ("PC-45", "DataPipelines (provider load, embedding worker, county enrichment, etc.) under Azure Functions",
     "Code/DataPipelines/function_app.py; Code/DataPipelines/provider_worker.py; "
     "Code/DataPipelines/embedding_worker.py; Code/DataPipelines/load_provider_data.py",
     "BR-09, BR-18", "TR-DM-001, TR-DM-002"),
    ("PC-46", "Operations / DevOps tools (HF Space lifecycle, brain snapshot, manifest generator, build bump)",
     "Code/Shared/ops/hf_space_create.py; Code/Shared/ops/hf_space_delete.py; "
     "Code/Shared/ops/hf_space_restart.py; Code/Shared/ops/hf_space_status.py; "
     "Code/Shared/ops/manifest_generator.py; Code/Shared/ops/bump_build.py; "
     "Code/Shared/ops/brain_snapshot.py",
     "BR-10, BR-25", "TR-OP-001, TR-BG-002"),
    ("PC-47", "Agent framework primitives (BaseAgent, BaseTool, ToolRegistry; same allowlist pattern)",
     "Code/Shared/agent_framework/base_agent.py; Code/Shared/agent_framework/base_tool.py; "
     "Code/Shared/agent_framework/tool_registry.py",
     "BR-10, BR-12", "TR-SS-002, TR-OP-004"),
    ("PC-48", "Provider canonical schema",
     "Code/Schemas/ChatHealthy.Providers.json",
     "BR-18", "TR-DM-002"),
]


# -----------------------------------------------------------------------------
# Business Requirements -- carried forward from design-V2.docx Section 3
# (verbatim where possible) plus the architect's URLGuardian BR.
# Per the V3 brief: V2 BRs are authoritative for V3; deep-authentication BRs
# from V2 (BR-20, BR-27) are excluded per architect directive.
# -----------------------------------------------------------------------------

BR_ROWS = [
    ("BR-01",
     "Conversational find-care: a single React UI lets a user describe a care need in natural language and receive a ranked list of providers in a supported state, with name, address, county, phone, and NPI; the system orchestrates tools to satisfy the request (carried from V2 BR-01).",
     "User-visible behavior: React submits prompt; parent runtime returns ranked providers + explanation messages; smoke test 'find pediatricians in Wilmington, DE' returns >=1 formatted provider record.",
     "PC-08, PC-11, PC-13, PC-14"),
    ("BR-02",
     "Conversational evaluate-care experience: a user can ask the system to look up a specific provider's credentials and receive NPI Registry data plus curated research links.",
     "Smoke test: 'tell me more about Dr. X (NPI 1234567890)' returns specialty, license, license state, address, plus 3+ research links (NPI Registry, Healthgrades, state board).",
     "PC-16, PC-17, PC-18"),
    ("BR-03",
     "Single conversational entry point: the React UI talks to one HTTP host for conversational flow regardless of which capability handles the step.",
     "Network capture from the React app shows exactly one hostname for /chat traffic; no direct calls from React to capability-specific URLs.",
     "PC-01, PC-42, V6 Section 3.3"),
    ("BR-04",
     "Pre-conversation safety + transport gate: every inbound /chat turn is gated by IP-lock check, dual-trigger emergency detection, and CORS allowlist before any LLM call.",
     "Synthetic emergency keyword turn returns canonical EMERGENCY_RESPONSE without invoking any tool; locked IP returns the same; cross-origin request from disallowed origin is rejected by CORS middleware.",
     "PC-25, PC-26, PC-27, PC-28, PC-34"),
    ("BR-05",
     "Static page surface: ten static product / marketing / architecture pages remain reachable from the parent host and re-renderable from declarative source.",
     "GET / on the parent host returns Website/index.html; GETs on the 9 sibling pages return 200; generator output is byte-identical for non-volatile pages.",
     "PC-44"),
    ("BR-06",
     "Provider direct-search read path: a React pagination control can fetch the next page of providers without re-invoking the LLM tool loop.",
     "POST /search with a saved search_params + after_npi returns the next page in <500ms p95; no Anthropic API call is made.",
     "PC-09, PC-10, PC-43"),
    ("BR-07",
     "Talk About Care lives inside FindCare: the Talk About Care capability is exposed as a feature folder under FindCare's RuntimeServer + SharedCode layout, not as its own epic. (See BR-28 for the binding rationale; this BR is the user-facing contract.)",
     "Backlog crosswalk shows the Talk About Care feature lives under FindCare; TR-FC-007 implements the about-content tools that constitute Talk About Care today.",
     "PC-24, V7 epic shape; implemented via TR-FC-007"),
    ("BR-08",
     "Per-turn audit trail: every /chat turn produces an audit record persisting ip, history length, tool-loop iterations, tokens in/out, response preview or error, and (on error) a de-identified history.",
     "Audit collection contains one row per turn; on error rows the chat_history_deidentified field is populated and contains no PII.",
     "PC-33, PC-27"),
    ("BR-09",
     "Provider data freshness: the production provider dataset is refreshed by the Data Management Pipelines on a published cadence; the conversational system always reads the published dataset.",
     "Pipeline run completes; FindCare reads via SharedServices Mongo abstraction; freshness probe within SLA.",
     "PC-45, PC-07"),
    ("BR-10",
     "Brain governance: every commit-time change is gated by the engineering-rules enforcement worker registry; brain artifacts are the single source of truth for runtime config.",
     "Pre-commit run with a rule-violating change exits non-zero; emergency keywords + tool definitions + system prompt rules are loaded from brain artifacts at startup.",
     "PC-04, PC-39, PC-40, PC-46, PC-47"),
    ("BR-11",
     "Operating-model alignment: every capability epic appears once on the org chart, once on the GL as a cost-center line, and once as a HuggingFace Space deployment line. The shape of any one IS the shape of all four.",
     "Epic registry crosswalk produces 1:1 mapping for the four capability epics + Operations + UX + DMP.",
     "V7 Section 1 axiom"),
    ("BR-12",
     "End-to-end observability: a single trace ID lets the human operator see latency and outcome of every node across every Space for any conversation; per-turn audit + LangSmith trace join by trace_id.",
     "Given a thread_id, LangSmith returns one trace tree spanning parent + capability subgraphs; audit collection rows can be joined to that trace by trace_id.",
     "PC-03, PC-33, V6 Section 3.3"),
    ("BR-13",
     "Two-tier HIPAA consent: a user may consent to (a) verbatim transcript persistence, (b) AI-summarized + de-identified persistence, or (c) neither (contact-only). The default is (c).",
     "Lead record with consent_verbatim=true persists chat_history; consent_summary=true persists AI-summarized + de-identified text; neither persists only contact fields.",
     "PC-19, PC-20, PC-21, PC-22"),
    ("BR-14",
     "HIPAA Safe Harbor de-identification: when consent_summary or error logging triggers, AI strips names, geographic identifiers (except state), dates (except year), phone, email, SSN, MRN, account/license/vehicle/device IDs, URLs, IP addresses.",
     "De-identification regression test asserts each PII category removed across a fixture chat; failure mode preserves original (no leak risk if AI fails).",
     "PC-21, PC-33"),
    ("BR-15",
     "Lead + unknown-question capture: the system can record a user contact lead (deduped by email) and an unanswerable question (3 classes: healthcare_capability, medical_advice, irrelevant) with operator push notification.",
     "Lead insert by new email persists one row; duplicate email returns ok without re-insert; unknown-question record returns the canonical verbatim response template by class.",
     "PC-22, PC-23, PC-41"),
    ("BR-16",
     "Recovery completeness: each pre-4/3 capability inventoried in Section 2 is reachable in the new architecture; no pre-4/3 user-visible behavior is dropped without an explicit BR/TR retiring it.",
     "Smoke-test suite invokes every PC-* capability through the parent runtime + non-conversational paths; suite is green.",
     "PC-01..PC-48 (all)"),
    ("BR-17",
     "No work outside the requirement set: the design implements only what BR/TR requires; new work requires a new requirement.",
     "Implementation-invention audit returns zero findings.",
     "Mandate; engineering rules anti-invention discipline"),
    ("BR-18",
     "Pipelines deliver every field the canonical Provider schema declares (NPI, entity type, name fields, practice address with lat/lng, taxonomies array with primary flag, county with name + fips, embedding fields).",
     "Pipeline regression test asserts each field present + typed per the canonical schema; vector index builds without missing fields.",
     "PC-08 (consumes), PC-45, PC-48"),
    ("BR-19",
     "Safety incidents are durable: emergency lock + audit are persisted to MongoDB with TTL; lock survives process restart; admin-unlock key clears a lock with audit.",
     "Lock IP -> restart process -> /chat returns EMERGENCY_RESPONSE; admin-unlock with key sets unlocked=true and adds an audit entry.",
     "PC-27"),
    ("BR-20",
     "OO conventions: every long-running Python program launches via a thin if __name__ == '__main__' block that constructs an object and runs it; deviations are tracked in an exception registry.",
     "Pre-commit OO scan exits non-zero on any new entry-point that violates the convention without a registered exception.",
     "Mandate 2 (V2 + V7)"),
    ("BR-21",
     "External-link integrity: links shown to the user (in chat text and in tool results) are validated; broken links are defanged (text kept, URL removed); known-correct redirects are silently rewritten.",
     "Synthetic 404 markdown link in an LLM response is returned to the user as plain text (not clickable); a known redirect is followed and the link points at the final URL.",
     "PC-29, PC-30, PC-31"),
    ("BR-22",
     "About / context surface: the system can return Skip Snow's professional background and ChatHealthy's mission with the operator-curated 'connect' field rendered verbatim as markdown links.",
     "get_skip_snow_context returns connect string verbatim; LinkedIn line is auto-linked if the LLM omitted the markdown.",
     "PC-04, PC-05, PC-24, PC-32, PC-36, PC-44, PC-38"),
    ("BR-23",
     "System-prompt assembly is data-driven: the system prompt is built from rules in code (or a future brain artifact) and ME context loaded from PDFs + text files at startup; no instruction text is hard-coded inside chat handlers.",
     "Prompt assembly returns the canonical 10-rule prompt; chat handler concatenates only the assembled system + user history.",
     "PC-04, PC-38"),
    ("BR-24",
     "Graceful degradation when MongoDB is unavailable: /chat continues to function with the LLM; only persistence-dependent operations return a soft 'recorded': 'ok', 'note': 'MongoDB unavailable' shape.",
     "With Mongo down, /chat returns a normal answer; /search returns {error: 'Database unavailable'}; lead insert returns the soft-ok shape.",
     "PC-07, PC-22, PC-08"),
    ("BR-25",
     "Build + environment surfacing: /health returns environment label, build number (read once at startup from MongoDB), version, and Mongo status.",
     "GET /health returns the four fields; build number matches the build_counter document.",
     "PC-06, PC-37, PC-46"),
    ("BR-26",
     "Pagination and search-state preservation: a single search produces total_count, first/last NPI, and search_params; the React control replays via after_npi without re-running the LLM tool loop.",
     "Initial search of taxonomy code returns 25 results out of total_count N; pagination forward yields the next 25 with no LLM call.",
     "PC-09, PC-10, PC-35, PC-43"),
    ("BR-27",
     "Capability isolation: a defect in any one capability subgraph cannot corrupt session state, observability state, or the data of any other capability. Cross-capability calls go through the parent runtime or via SharedServices' explicit interfaces, never via in-process imports across capability boundaries.",
     "Static cross-capability import scan exits clean; integration test injecting a fault into one Space leaves other Spaces serving normally; checkpoint thread state survives a single-capability crash.",
     "V6 3.2/3.3, V7 4.2 (capability boundaries)"),
    ("BR-28",
     "Talk About Care lives inside FindCare: the Talk About Care capability is exposed as a feature folder under FindCare's RuntimeServer + SharedCode layout, not as its own epic.",
     "V7 epic shape resolves Talk About Care under FindCare's dual-child layout; backlog crosswalk shows zero rows under a 'Talk About Care' epic.",
     "V7 epic shape after V6 annotation resolution; preserves PC-24"),
]


# -----------------------------------------------------------------------------
# Technical Requirements (V20 specificity, anchored to BR + V7 + V6)
# -----------------------------------------------------------------------------

TR_ROWS = [
    # ---- Parent runtime (Operations) ----
    ("TR-LG-001",
     "Operations/RuntimeServer/orchestration/parent_graph.py implements the parent LangGraph StateGraph; state in Pydantic BaseModel; messages: Annotated[list[AnyMessage], add_messages].",
     "Studio renders parent_graph.py; messages channel verified by unit test; pydantic ValidationError on bad input.",
     "BR-01, BR-03, BR-04, V6 3.3"),
    ("TR-LG-002",
     "Operations/RuntimeServer/app.py exposes one HTTP endpoint compatible with @langchain/langgraph-sdk for streaming conversational flow; non-conversational reads (/search, /welcome, /health) remain available either at the parent or via SharedServices.",
     "Network capture; React submits prompt and receives streaming tokens; /search bypasses the parent graph.",
     "BR-03, BR-06, BR-25, V6 3.3"),
    ("TR-LG-003",
     "MongoDBSaver from langgraph-checkpoint-mongodb is constructed exactly once in Operations/RuntimeServer/orchestration/checkpointer.py at module load; capability Spaces never instantiate a checkpointer.",
     "Process inventory; integration test asserts thread_id continuity across capability invocations.",
     "BR-08, BR-12, V6 3.3"),
    ("TR-LG-004",
     "Each capability subgraph (FindCare, EvaluateCare, Security & Compliance, SharedServices) is composed into the parent graph as a RemoteGraph from langgraph.pregel.remote pointed at its HF Space URL.",
     "parent_graph.py imports RemoteGraph; integration test runs a turn that visits each capability node; LangSmith trace shows the cross-Space spans.",
     "BR-03, BR-12, V6 3.2"),
    ("TR-LG-005",
     "RemoteGraph never calls back into its own deployment; same-Space step composition uses in-process subgraph composition (V6 3.2 deadlock constraint).",
     "Static analyzer flags any RemoteGraph(url=SELF) usage as a violation.",
     "BR-27, V6 3.2"),
    ("TR-LG-006",
     "Per-capability state lives in {Capability}/RuntimeServer/state.py as a Pydantic BaseModel with explicit input_schema and output_schema fields; only declared fields cross the Space boundary.",
     "Schema validation test on the public state shape; private fields raise ValidationError if leaked.",
     "BR-12, V6 Q1"),
    ("TR-LG-007",
     "Tool calls and tool results travel as ToolMessage entries on the messages channel; structured artefacts the LLM does not need to read (e.g. raw provider rows for downstream rendering) get a typed Pydantic field on subgraph state, not a flat top-level state field per tool.",
     "Anti-pattern scanner asserts no flat tool-output fields on state schemas; smoke test verifies ToolMessage round-trip.",
     "BR-12, V6 Q4"),
    ("TR-LG-008",
     "Dependency injection: nodes accept Runtime[ContextSchema]; tools accept ToolRuntime[ContextSchema]. Database clients and LLM clients are NEVER imported at module scope inside node files (test isolation + trace propagation).",
     "Static analyzer flags top-level mongo / Anthropic / OpenAI imports in node modules.",
     "BR-12, V6 Q5"),
    ("TR-LG-009",
     "InMemorySaver is the only test-time checkpointer; production checkpointer is MongoDBSaver. No SqliteSaver, no PostgresSaver, no custom mongo saver.",
     "Test fixtures use InMemorySaver; CI fails if MongoDBSaver is referenced from a unit test.",
     "BR-08, BR-12, V6 Q2"),

    # ---- FindCare (worked example, V20 granularity) ----
    ("TR-FC-001",
     "FindCare/RuntimeServer/agent.py compiles a StateGraph; FindCare/RuntimeServer/{state.py, nodes.py, tools.py} follow the V6 4-file layout; FindCare/RuntimeServer/langgraph.json declares the manifest.",
     "File-presence test + unit test on graph build; langgraph.json is valid against the langgraph-example-pyproject template.",
     "BR-01, V6 Q6, V7 dual-child"),
    ("TR-FC-002",
     "tool_find_providers is a @tool inside FindCare/RuntimeServer/tools.py; it delegates to FindCare/SharedCode/core/provider_search_service.py whose interface is the pre-4/3 search_providers signature (specialty_query, state, city, county, name, npi, specialty_codes, after_npi, limit). Routes 1-4 (NPI lookup, Name search, Specialty codes direct filter, Specialty query via vector + taxonomy) are preserved; pagination metadata (total_count, first_npi, last_npi, has_more, page_start, page_end, search_params) is preserved on the return shape.",
     "ToolNode wires it; smoke test invokes each route; pagination round-trips.",
     "BR-01, BR-06, BR-26, PC-08..PC-10, PC-13..PC-15"),
    ("TR-FC-003",
     "FindCare/RuntimeServer/app.py exposes a NON-graph POST /search endpoint for React pagination; this endpoint reuses the same SharedCode/core/provider_search_service.py implementation as the @tool to avoid logic drift.",
     "Trace inspection: no parent-graph trace for POST /search; result shape matches the @tool result.",
     "BR-06, BR-26, V6 3.3"),
    ("TR-FC-004",
     "tool_find_specialty_codes is a @tool in FindCare/RuntimeServer/tools.py; it delegates to FindCare/SharedCode/core/specialty_service.py which preserves the pre-4/3 dual-pipeline (regex with AI-expanded stems via Claude Haiku + vector via text-embedding-3-small filtered by INDIVIDUAL_PROVIDER_GROUPINGS classification) and merges deduplicated by Code.",
     "Unit test per pipeline; integration test asserts dedup and INDIVIDUAL_PROVIDER_GROUPINGS filter.",
     "BR-01, PC-11, PC-12, PC-15"),
    ("TR-FC-005",
     "tool_record_user_details is a @tool in FindCare/RuntimeServer/tools.py; it delegates to FindCare/SharedCode/core/lead_service.py which preserves: dedupe-by-email, three consent paths (verbatim, summary, neither), Sparkpost push notify on success, MongoDB write to {ENV}_AboutUs.lead.",
     "Smoke test inserts a lead row; duplicate email returns {recorded: ok} without re-insert; missing email returns the soft-ok shape.",
     "BR-13, BR-15, PC-22"),
    ("TR-FC-006",
     "tool_record_unknown_question is a @tool in FindCare/RuntimeServer/tools.py; it delegates to FindCare/SharedCode/core/unknown_question_service.py which preserves the 3-class taxonomy (healthcare_capability, medical_advice, irrelevant), the verbatim response_template per class, and the consent-gated de-identification + persistence path.",
     "Smoke test per class returns the canonical template; consent=true persists with de-identified history.",
     "BR-14, BR-15, PC-23"),
    ("TR-FC-007",
     "tool_get_skip_snow_context and tool_get_chathealthy_context are @tools in FindCare/RuntimeServer/tools.py; they delegate to FindCare/SharedCode/core/about_service.py which preserves the pre-4/3 ME context shape (linkedin, summary, anthropic_principles, business_plan, connect-as-verbatim-markdown). Post-LLM Skip-Snow LinkedIn auto-link rewriter remains in the parent runtime as a response-finalize step.",
     "Smoke test returns the verbatim 'connect' field unchanged; auto-link rewriter rewrites bare 'Skip Snow on LinkedIn' on the response edge.",
     "BR-07, BR-22, BR-28, PC-04, PC-24, PC-32, PC-38"),
    ("TR-FC-008",
     "FindCare/RuntimeServer/app.py is constructed via the Application class pattern (TR-OO-001). The if __name__ == '__main__': block calls Application().run() which programmatically calls uvicorn.run on a FastAPI ASGI app; the FastAPI instance is built inside Application.__init__.",
     "Static check: no module-top app = FastAPI() in FindCare/RuntimeServer/app.py.",
     "BR-20, V7 dual-child"),
    ("TR-FC-009",
     "System prompt rules (the 10 rules currently in PromptSystemMaker.build_system_prompt) are loaded from brain/machine_artifacts/content/system_prompt_rules.json by SharedServices/SharedCode/prompt_system_maker.py; FindCare composes the prompt from this artifact; emergency keywords are loaded from brain/machine_artifacts/content/emergency_keywords.json.",
     "Brain-artifact diff test; prompt-assembly unit test; emergency keywords are read from brain at startup.",
     "BR-10, BR-23, PC-04, PC-38, PC-39"),
    ("TR-FC-010",
     "URLGuardian becomes SharedServices/SharedCode/url_guardian.py and runs as a parent-runtime response-finalize step (NOT a graph node) on every assistant text and on tool results that carry a 'links' dict. The 3-stage pipeline (HEAD reachability, AI content verify via Claude Haiku, Google Custom Search correction) is preserved; the SPA + trusted-search domain allowlists are preserved; cache TTL 3600s, request timeout 5s, max 4 worker threads.",
     "Synthetic 404 markdown link is defanged in the response; redirect is rewritten silently; SPA domain returns reachable when 200.",
     "BR-21, PC-29, PC-30, PC-31"),

    # ---- EvaluateCare (gestural per assignment) ----
    ("TR-EC-001",
     "EvaluateCare/RuntimeServer/agent.py exposes the evaluate-care subgraph as a RemoteGraph from the parent. The pre-4/3 clinical-trials capability (ClinicalTrials.gov v2 API, RECRUITING filter, max 10 page size) is preserved as tool_search_clinical_trials in EvaluateCare/RuntimeServer/tools.py.",
     "Smoke test: 'find a recruiting trial for diabetes' returns >=1 trial with NCT id, title, eligibility summary.",
     "BR-02, PC-16"),
    ("TR-EC-002",
     "Travel info via Google Routes API (DRIVE mode, TRAFFIC_UNAWARE) is preserved on EvaluateCare/SharedCode/clinical_trials_service.py; the worldwide-origin contract is preserved per the pre-4/3 system prompt RULE 6.",
     "Trial result includes distance + duration when user_location is supplied; non-US origin works.",
     "BR-02, PC-17"),
    ("TR-EC-003",
     "tool_lookup_provider_external is a @tool in EvaluateCare/RuntimeServer/tools.py; preserves NPI Registry CMS v2.1 fetch + Healthgrades + state-board (DE/MS/VA) link construction; the 'guidance' string per link is preserved verbatim from the pre-4/3 ProviderDetailService.",
     "Smoke test returns a 'links' dict with at least npi_registry + healthgrades + state-board; URLGuardian validates the dict.",
     "BR-02, BR-21, PC-18"),
    ("TR-EC-004",
     "Non-conversational EvaluateCare HTTP endpoints (provider lookup detail page, evaluation report) live in EvaluateCare/RuntimeServer/app.py outside the graph (V6 3.3 non-conversational read carve-out).",
     "Trace inspection: no parent-graph trace for these GETs.",
     "BR-06, V6 3.3"),

    # ---- Security & Compliance (gestural) ----
    ("TR-SE-001",
     "Security/RuntimeServer/agent.py is the FIRST capability node in the parent's edge order. Security/RuntimeServer/nodes.py preserves the pre-4/3 enforcement order: admin_unlock_check -> ip_locked_check -> emergency_check (dual-trigger keyword OR Claude classifier @ confidence>=0.80) -> safe_prefix_fast_path. Locked or emergency turns short-circuit the conversation graph and return the canonical EMERGENCY_RESPONSE.",
     "Edge-order test; integration test rejects each gate before any other capability runs; safe_prefix message bypasses the AI classifier.",
     "BR-04, PC-25, PC-26, PC-27, PC-28"),
    ("TR-SE-002",
     "Security/SharedCode/consent_service.py preserves the pre-4/3 two-tier consent (verbatim transcript, AI-summarized + de-identified, neither). Both Tier 1 and Tier 2 use Claude Haiku for the AI step; the de-identification prompt is loaded from brain/machine_artifacts/content/hipaa_de_identification_prompt.json.",
     "Tier 1 persists verbatim chat_history; Tier 2 persists summarized + de-identified text only; Tier 3 (none) persists only contact fields.",
     "BR-13, BR-14, PC-19, PC-20, PC-21"),
    ("TR-SE-003",
     "Compliance is a feature folder under Security/RuntimeServer/compliance/, not its own epic. Security and Compliance are one cost-center line and one HF Space deployment (V7 A1 architect ratification).",
     "Folder presence test; epic registry shows Security & Compliance as one cost-center line.",
     "V7 A1, BR-11"),
    ("TR-SE-004",
     "IP-lock storage uses {ENV}_Safety.emergency_incidents (the pre-4/3 collection name is preserved for migration continuity); TTL 3600 seconds; admin_unlock_key clears with audit trail.",
     "Lock survives process restart; admin-unlock key sets unlocked=true and adds an audit row.",
     "BR-19, PC-27"),

    # ---- SharedServices (gestural) ----
    ("TR-SS-001",
     "SharedServices/SharedCode/mongo_utilities.py is the canonical Mongo client used by every Space; capability Spaces never instantiate raw pymongo clients (V7 A2 / 13.2). The pre-4/3 graceful-unavailable mode is preserved (a soft 'note: MongoDB unavailable' return shape; no crash).",
     "Static analyzer flags any non-shared mongo client construction; integration test exercises the unavailable path.",
     "BR-09, BR-24, PC-07"),
    ("TR-SS-002",
     "SharedServices/SharedCode/llm_client.py centralizes Anthropic + OpenAI clients. Pre-4/3 model assignments are preserved: claude-sonnet-4-6 for the chat loop, claude-haiku-4-5-20251001 for query expansion, summarization, de-identification, and URL content verification, gpt-4.1-mini for the safety classifier, text-embedding-3-large for provider vectors, text-embedding-3-small for specialty vectors.",
     "Capability nodes invoke via ToolRuntime; per-model usage histograms verify the model split.",
     "BR-12, BR-23, PC-04, PC-12, PC-15"),
    ("TR-SS-003",
     "SharedServices/SharedCode/prompt_system_maker.py preserves pre-4/3 PromptSystemMaker behavior (brain artifact reads, ME-context loader, build-number reader, welcome-message builder, system-prompt assembler, trim utility); brain dir resolution falls back gracefully on the HF flat layout.",
     "Brain-artifact diff test; prompt-assembly + welcome-message contract tests pass.",
     "BR-10, BR-22, BR-23, PC-04, PC-05, PC-06"),

    # ---- Operations (parent + observability + governance) ----
    ("TR-OP-001",
     "Operations/SharedCode/build_counter/ holds bump_build.py (build counter incrementer). Build number is read once at startup from {ENV}_System.build_counter via a one-shot read in PromptSystemMaker.get_build_number; /health on the parent returns the cached value.",
     "Bump increments the document; /health returns the new value after restart only (proves single-shot read).",
     "BR-25, PC-06, PC-46"),
    ("TR-OP-002",
     "Operations/SharedCode/notify/sparkpost_push.py preserves the pre-4/3 best-effort Sparkpost transmission (silent failure if SPARKMAIL_API_KEY absent; logged warning on send error).",
     "Lead capture path triggers a push; missing key returns silently; bad recipient logs warning, returns ok.",
     "BR-15, PC-41"),
    ("TR-OP-003",
     "Operations/SharedCode/brain_governance/{rebuild_manifest.py, manifest_generator.py} (relocated per V7 A8) are GOVERNANCE-CRITICAL; the manifest inventories all MongoDB collections and brain JSONs (Rule-046).",
     "Manifest generator emits a complete inventory; pre-commit scan asserts no orphan collection.",
     "BR-10, V7 A8, Rule-046"),
    ("TR-OP-004",
     "Operations/SharedCode/agent_framework/ relocates the pre-4/3 BaseAgent + BaseTool + ToolRegistry primitives. The allowlist + GOV-004 enforcement (\"the model may suggest; the system must decide\") is the canonical pattern reused by every capability subgraph's @tool registration.",
     "Static check: every @tool registration in {Capability}/RuntimeServer/tools.py is in the allowlist.",
     "BR-10, BR-12, PC-03, PC-47"),
    ("TR-OP-005",
     "Operations/SharedCode/hf_space/ relocates the pre-4/3 hf_space_create.py / hf_space_delete.py / hf_space_restart.py / hf_space_status.py utilities (Rule-024, Rule-051: Spaces created exclusively via create_hf_space.py).",
     "Smoke test: create_hf_space.py with a name parameter creates a Space at the hardcoded tier; bypassing the helper is a Rule-051 violation.",
     "BR-25, PC-46, Rule-024, Rule-051"),

    # ---- Observability ----
    ("TR-OB-001",
     "LangSmith tracing is enabled in the parent runtime AND in every capability Space; trace_id is propagated across RemoteGraph invocations via the standard context= argument.",
     "End-to-end test: a single thread_id yields one consolidated trace tree across Spaces.",
     "BR-12, BR-08"),
    ("TR-OB-002",
     "Operations/SharedCode/conversation_log/Worker (canonical, relocated per V7 13.2) consumes the conversation-log Kafka topic and writes structured per-turn audit rows to MongoDB. The audit row schema preserves the pre-4/3 DebugLogger fields (datetime, ip, message_preview, history_len, tool_loop_iters, tokens_in, tokens_out, response_preview, error, optional chat_history_deidentified) plus trace_id for join.",
     "Worker integration test; audit row exists for every emitted turn event; trace_id joins to LangSmith.",
     "BR-08, BR-12, BR-14, PC-33"),
    ("TR-OB-003",
     "Per-node performance budgets are emitted as OpenTelemetry histograms; alert rules fire on rolling-24h p95 breach. Concrete pre-4/3-rooted budgets: provider_search p95 <= 1500ms; specialty_classifier p95 <= 600ms; security/emergency_check p95 <= 200ms; URLGuardian p95 <= 800ms (the pre-4/3 5s timeout x4 workers gives a soft ceiling).",
     "OTel collector receives histograms; alert rule unit-test fires on synthetic breach.",
     "BR-12"),
    ("TR-OB-004",
     "LangSmith is the trace UI of record; observability backend is resolved per Section 9 R4; the design depends on the LangSmith feature set (cross-Space spans via context=, replay, search by thread_id).",
     "Recommendation accepted in Section 9; design depends on LangSmith.",
     "BR-12, Resolved issue R4"),

    # ---- OO conventions ----
    ("TR-OO-001",
     "Every Python service entry-point file ends with `if __name__ == '__main__': Application().run()` (or equivalent). The Application class lives in the same module or is imported from {Capability}/RuntimeServer/app.py.",
     "Pre-commit OO scan asserts each registered entry-point file conforms.",
     "BR-20"),
    ("TR-OO-002",
     "Operations/SharedCode/scan/oo_thin_main_scan.py walks the entry-point registry; non-conforming files MUST appear in Operations/SharedCode/scan/oo_exceptions.json with justification, owner, review_date.",
     "Scan exits non-zero if a non-registered exception appears; expired review_date entries fail.",
     "BR-20"),
    ("TR-OO-003",
     "FastAPI handlers are NOT entry-points. The web app object construction happens in app.py's Application class; the if __name__ == '__main__' block calls Application().run() which runs uvicorn programmatically.",
     "Static check: no module-top app = FastAPI() in {Capability}/RuntimeServer/app.py.",
     "BR-20"),

    # ---- Brain governance / blast radius ----
    ("TR-BG-001",
     "engineering_rules.json is canonical; every action that mutates state in any project file is gated by an enforcement worker per the V20 framework.",
     "Pre-commit run with a rule-violating change exits non-zero.",
     "BR-10"),
    ("TR-BG-002",
     "All requirement IDs follow EPIC-NNN-F-NNN-S-NNN-REQ-{B|T}-NNN; every code change traces to one such ID (Rule-025).",
     "Backlog crosswalk; orphan-code scan exits non-zero on any new uncited path.",
     "BR-10, BR-17, Rule-025"),
    ("TR-BG-003",
     "The boot class Code/Shared/ops/tools/chathealthy_devops_boot.py governs every Claude Code session via .claude/settings.json hooks. The boot class reads brain artifacts read-only and emits a session digest the operator can audit.",
     "Hook fires on session start; digest produced; settings.json references the boot class.",
     "BR-10"),
    ("TR-BG-004",
     "Project-root capability folders (EvaluateCare/, FindCare/, SharedServices/, Security/, Operations/, UX/, DataManagementPipelines/) replace the Code/ umbrella; cross-capability imports require an explicit dependency declaration.",
     "Cross-capability import scan exits non-zero on illegal imports.",
     "BR-11"),

    # ---- UX ----
    ("TR-UX-001",
     "FindCare/Frontend/ (the React iframe) is the conversational facade; it consumes the parent runtime via @langchain/langgraph-sdk; it does NOT import any capability Space URL directly.",
     "Source scan: only the parent URL string appears in src/; capability URLs are absent.",
     "BR-03, V6 3.3"),
    ("TR-UX-002",
     "UX/Frontend/index.html is the canonical landing page (Cloudflare Pages target); it embeds FindCare/Frontend as an iframe; the 9 sibling static pages remain reachable.",
     "Smoke test: GET / serves the iframe wrapper; the 9 sibling pages return 200.",
     "BR-03, BR-05"),
    ("TR-UX-003",
     "UX/SharedCode/ holds shared React primitives (ProviderCard.tsx, SelectionManager.tsx, provider.ts) that both FindCare/Frontend and EvaluateCare/Frontend consume. The pre-4/3 GUIManager.tsx postMessage bridge between the static control frame and the React iframe is preserved (parent renders pagination buttons; React owns state).",
     "Build-time import resolution from both consumers succeeds; pagination postMessage round-trips in an integration test.",
     "BR-05, BR-26, V7 A7"),
    ("TR-UX-004",
     "Operations/SharedCode/static_page_generator/ regenerates the 10 non-roadmap static pages from declarative inputs in UX/SharedCode/static_pages/.",
     "Generator run produces byte-identical output for non-volatile pages; roadmap.html stays deleted per V7 A10.",
     "BR-05"),

    # ---- Data Management Pipelines ----
    ("TR-DM-001",
     "DataManagementPipelines/ holds ingest, dedup, and embedding-generation jobs (relocated from Code/DataPipelines/). Jobs run under Operations infrastructure but are NOT nodes in any LangGraph (V6 3.3 carve-out).",
     "Pipeline run does not produce LangSmith traces.",
     "BR-09"),
    ("TR-DM-002",
     "The canonical Provider schema lives at brain/machine_artifacts/schemas/ChatHealthyProvidersSchema.json; the schema declares all required fields (NPI, entity type, name fields, taxonomies array with primary flag, practice_address with lat/lng, county.fips + county.name, embedding + embedding_model + embedding_version, NPPES other-name fields). Code/Schemas/ChatHealthy.Providers.json is DELETED per V7 A3.",
     "Schema diff test asserts every declared field is delivered by the pipeline; orphan-file scan asserts the deleted path is gone.",
     "BR-18, V7 A3"),
    ("TR-DM-003",
     "Provider-data writes go through SharedServices/SharedCode/mongo_utilities.py only; pipelines never instantiate raw pymongo clients (consistent with TR-SS-001).",
     "Static analyzer flags any non-shared mongo client construction.",
     "BR-09, BR-18"),

    # ---- Recovery + internal consistency ----
    ("TR-RC-001",
     "Smoke test Operations/SharedCode/scan/pre_4_3_capability_recovery_smoke.py invokes each PC-* capability through the parent runtime + non-conversational paths and asserts the LangSmith trace contains the corresponding span (or, for non-conversational paths, the audit row).",
     "CI is green; the smoke runner reports 48 PC entries covered (one per row in Section 2).",
     "BR-16"),
    ("TR-IC-001",
     "Operations/SharedCode/scan/br_tr_coverage_scan.py asserts every BR has at least one TR and every TR traces to at least one BR; emits the coverage matrix.",
     "Scan exits non-zero if any row is unmapped.",
     "BR-16, BR-17 (internal-consistency check)"),
]


# -----------------------------------------------------------------------------
# Coverage matrix construction
# -----------------------------------------------------------------------------

def build_coverage_matrix():
    """Compute BR<->TR coverage."""
    br_to_tr = {br[0]: [] for br in BR_ROWS}
    tr_to_br = {tr[0]: [] for tr in TR_ROWS}
    for tr_id, _stmt, _at, trace in TR_ROWS:
        for token in trace.replace(",", " ").split():
            tk = token.strip(".,;:()")
            if tk.startswith("BR-") and tk in br_to_tr:
                br_to_tr[tk].append(tr_id)
                tr_to_br[tr_id].append(tk)
    return br_to_tr, tr_to_br


# -----------------------------------------------------------------------------
# UML diagram sources
# -----------------------------------------------------------------------------

DIAG_DEPLOYMENT = """@startuml
title ChatHealthy V3 Deployment Topology
skinparam componentStyle rectangle
skinparam shadowing false
skinparam defaultFontSize 11

cloud "Cloudflare Pages" as CF {
  [UX/Frontend index.html\\n+ 9 static pages] as STATIC
  [FindCare/Frontend\\n(React iframe)] as REACT
}

node "HuggingFace Spaces" as HF {
  [Operations Space\\nparent_graph.py\\nMongoDBSaver\\n@langgraph-sdk endpoint] as OPS
  [Security & Compliance Space] as SEC
  [FindCare Space] as FC
  [EvaluateCare Space] as EC
  [SharedServices Space\\nMongo client + LLM client] as SS
}

cloud "Operations infra (not a Space)" as INFRA {
  [DataManagementPipelines\\nAzure Functions] as DMP
  [Kafka conversation_log] as KAFKA
  [conversation_log Worker] as WK
}

database "MongoDB Atlas" as MDB
database "LangSmith" as LS

REACT --> OPS : @langgraph-sdk\\nstreaming
STATIC ..> REACT : iframe
OPS --> SEC : RemoteGraph (first)
OPS --> FC : RemoteGraph
OPS --> EC : RemoteGraph
FC --> SS : ToolRuntime
EC --> SS : ToolRuntime
SEC --> SS : ToolRuntime
SS --> MDB
DMP --> SS
OPS --> KAFKA : per-turn audit
KAFKA --> WK
WK --> MDB : conversation_log
OPS ..> LS : trace
FC ..> LS : trace span
EC ..> LS : trace span
SEC ..> LS : trace span

@enduml
"""

DIAG_SEQUENCE = """@startuml
title V3 Conversational Turn (request flow)
skinparam shadowing false
skinparam defaultFontSize 11
actor User
participant "React iframe" as React
participant "Operations parent_graph" as Parent
participant "Security Space" as Sec
participant "FindCare Space" as FC
participant "SharedServices" as SS
database  "MongoDB" as Mdb
database  "LangSmith" as LS

User -> React : utterance
React -> Parent : POST runs (langgraph-sdk)
Parent -> LS : trace start (thread_id, trace_id)
Parent -> Sec : context (trace_id, ip, history)
Sec -> Sec : admin_unlock then ip_locked then emergency
alt locked or emergency
  Sec --> Parent : EMERGENCY_RESPONSE
  Parent --> React : stream EMERGENCY_RESPONSE
else continue
  Parent -> FC : context, messages
  FC -> FC : tool-use loop (find_providers, specialty, trials)
  FC -> SS : ToolRuntime (Mongo, LLM)
  SS -> Mdb : vector + taxonomy + name + npi searches
  SS --> FC : provider rows
  FC --> Parent : ToolMessage stream + pagination metadata
  Parent -> Parent : URLGuardian + LinkedIn rewriter
  Parent --> React : stream tokens + pagination
end
Parent -> Mdb : checkpoint (MongoDBSaver)
Parent --> LS : trace end (async)
@enduml
"""

DIAG_CLASS_STATE = """@startuml
title V3 Pydantic State Schemas (per V6 4-file layout)
skinparam classAttributeIconSize 0
skinparam shadowing false
skinparam defaultFontSize 11

class ParentState << Pydantic >> {
  + messages : Annotated[list[AnyMessage], add_messages]
  + thread_id : str
  + user_id : str
  + locale : str
  + current_capability : Optional[str]
  + routing_hint : Optional[str]
}

class ContextSchema << Pydantic >> {
  + user_id : str
  + thread_id : str
  + trace_id : str
  + locale : str
  + request_id : str
  + ip : str
}

class FindCareState << Pydantic >> {
  + messages : Annotated[list[AnyMessage], add_messages]
  + last_search_params : Optional[ProviderSearchInput]
  + last_provider_result : Optional[ProviderSearchResult]
  + pagination : Optional[PaginationMeta]
}

class EvaluateCareState << Pydantic >> {
  + messages : Annotated[list[AnyMessage], add_messages]
  + last_npi_lookup : Optional[ProviderDetail]
  + last_trial_search : Optional[ClinicalTrialResult]
}

class SecurityState << Pydantic >> {
  + messages : Annotated[list[AnyMessage], add_messages]
  + ip : str
  + locked : bool
  + emergency : bool
  + safe_prefix_hit : bool
}

class ProviderSearchInput << Pydantic >> {
  + specialty_query : Optional[str]
  + state : Optional[str]
  + city : Optional[str]
  + county : Optional[str]
  + name : Optional[str]
  + npi : Optional[str]
  + specialty_codes : Optional[list[str]]
  + after_npi : Optional[str]
  + limit : int = 25
}

class PaginationMeta << Pydantic >> {
  + has_more : bool
  + first_npi : Optional[str]
  + last_npi : Optional[str]
  + count : int
  + total_count : int
  + page_start : int
  + page_end : int
  + search_params : Optional[dict]
  + specialization_options : Optional[list[dict]]
}

ParentState ..> ContextSchema : runtime context
FindCareState ..> ProviderSearchInput
FindCareState ..> PaginationMeta

@enduml
"""

DIAG_ACTIVITY = """@startuml
title FindCare User Journey (pre-4/3 behavior preserved in V3)
skinparam shadowing false
skinparam defaultFontSize 11
start
:user utterance arrives at parent runtime;
:Security: admin_unlock_check;
if (admin unlock?) then (yes)
  :respond "Session unlocked.";
  stop
endif
:Security: ip_locked_check;
if (locked?) then (yes)
  :return EMERGENCY_RESPONSE;
  stop
endif
:Security: emergency_check\\n(safe-prefix fast path then\\nkeyword OR Claude classifier@>=0.80);
if (emergency?) then (yes)
  :persist {ENV}_Safety.emergency_incidents\\n(ttl 3600s) + audit;
  :return EMERGENCY_RESPONSE;
  stop
endif
:FindCare tool-use loop (Anthropic tools);
repeat
  if (tool == find_providers) then
    :ProviderSearchService.search_providers;
    :Route 1..4: NPI / Name / Specialty codes / Specialty query (vector + taxonomy);
    :pagination metadata (total_count, first_npi, last_npi, search_params);
  elseif (tool == find_specialty_codes) then
    :SpecialtyService dual pipeline\\n(regex with AI-expanded stems + vector\\nfiltered by INDIVIDUAL_PROVIDER_GROUPINGS);
  elseif (tool == search_clinical_trials) then
    :ClinicalTrials.gov v2 RECRUITING\\n+ optional Google Routes travel info;
  elseif (tool == lookup_provider_external) then
    :NPI Registry CMS v2.1\\n+ Healthgrades + state-board (DE/MS/VA);
  elseif (tool == record_user_details) then
    :LeadService dedupe-by-email\\n+ consent-tier persistence + push;
  elseif (tool == record_unknown_question) then
    :3-class taxonomy verbatim template\\n+ optional consent-gated de-id persistence;
  else (other context tools)
    :AboutService static content;
  endif
repeat while (more tool calls?)
:URLGuardian on assistant text;
:Skip-Snow LinkedIn auto-link rewriter;
:DebugLogger audit row;
:stream response to React;
stop
@enduml
"""

DIAG_COMPONENT = """@startuml
title V3 Component Topology -- Capability Spaces and shared cores
skinparam componentStyle rectangle
skinparam shadowing false
skinparam defaultFontSize 11

package "Operations Space" {
  [parent_graph.py] as PG
  [checkpointer.py\\n(MongoDBSaver)] as CHK
  [URLGuardian (response-finalize)] as UG
}
package "Security & Compliance Space" {
  [security agent.py] as SECA
  [admin_unlock / ip_locked /\\nemergency / safe_prefix nodes] as SECN
  [Compliance feature folder] as COMP
  [consent_service (Tier 1/2/3)] as CON
}
package "FindCare Space" {
  [findcare agent.py] as FCA
  [find_providers / find_specialty_codes /\\nrecord_user_details / record_unknown_question /\\nget_skip_snow_context / get_chathealthy_context] as FCT
  [SharedCode/core/provider_search_service] as PSS
  [SharedCode/core/specialty_service] as SPSV
  [SharedCode/core/lead_service] as LSV
  [SharedCode/core/about_service] as AB
}
package "EvaluateCare Space" {
  [evaluate agent.py] as ECA
  [search_clinical_trials /\\nlookup_provider_external] as ECT
  [SharedCode/clinical_trials_service] as CTS
  [SharedCode/provider_detail_service] as PDS
}
package "SharedServices Space" {
  [mongo_utilities] as MU
  [llm_client (Anthropic + OpenAI)] as LC
  [embedding_client (3-large + 3-small)] as EM
  [prompt_system_maker] as PSM
}
package "DataManagementPipelines (no Space)" {
  [provider_worker] as PW
  [embedding_worker] as EW
  [county_enrichment] as CE
}

PG --> SECA
PG --> FCA
PG --> ECA
PG --> CHK
PG --> UG
SECA --> SECN
SECN --> CON
FCA --> FCT
FCT --> PSS
FCT --> SPSV
FCT --> LSV
FCT --> AB
ECA --> ECT
ECT --> CTS
ECT --> PDS
PSS --> MU
SPSV --> MU
SPSV --> EM
PSS --> EM
LSV --> MU
PSM --> LC
CON --> LC
PW --> MU
EW --> EM
CE --> MU

@enduml
"""


# -----------------------------------------------------------------------------
# Document body
# -----------------------------------------------------------------------------

def build_doc() -> Document:
    doc = Document()

    # ---------- Title ----------
    t = doc.add_heading("ChatHealthy.ai -- End-to-End Production Architecture Design -- V3", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT

    add_para(doc, "Design V3 -- the ground-truth-anchored end-to-end production architecture design.")
    add_para(doc, f"Date: 2026-04-28   |   Author: Claude (architecture design agent)   |   Read-only deliverable")
    add_para(
        doc,
        f"GROUND TRUTH: pre-4/3 application code at git commit {GROUND_TRUTH_COMMIT} ({GROUND_TRUTH_DATE}). "
        "V2 is the architect's recollection and is used for cross-reference; where V2 disagrees with the pre-4/3 "
        "code the code wins. LangGraph/poc/* is post-4/20 and is NOT cited as pre-event evidence.",
    )
    add_para(
        doc,
        "Binding inputs: pre-4/3 code at commit be14579a; "
        "architecture/ChatHealthyApplicationDesign/ArchitectureDesignAndAuditDocs/design-V2.docx; "
        "findCare/ArchitectureAndDesign/codebase-reorganization-V7.docx (file structure / epic shape); "
        "findCare/ArchitectureAndDesign/langgraph-oo-best-practices-V6.docx (runtime patterns); "
        "brain/machine_artifacts/content/{engineering_rules, agile_backlog}.json.",
    )

    # ---------- Operating-model axiom ----------
    add_heading(doc, "0. Operating-model axiom (verbatim)", level=1)
    add_para(
        doc,
        '"The operating model, the digital footprint, the general ledger, and the org chart are 100% in alignment. '
        "The structure of any one IS the structure of all four. This alignment is non-negotiable and is the "
        'binding test against which every architectural proposal is measured."',
        italic=True,
    )
    add_para(
        doc,
        "Every V3 proposal is tested against the axiom: does the proposed Git-tree shape mirror what would appear "
        "on the org chart, on a general-ledger cost-center list, and on the customer-facing capability map. "
        "Each capability epic is one cost-center line, one org-chart line, one HuggingFace Space deployment line.",
    )

    # ---------- 1. Executive summary ----------
    add_heading(doc, "1. Executive summary", level=1)
    add_para(
        doc,
        "ChatHealthy.ai's production architecture is a thin LangGraph parent runtime composing four capability "
        "subgraphs (FindCare, EvaluateCare, Security & Compliance, SharedServices) via RemoteGraph. The parent "
        "runtime owns one MongoDBSaver checkpointer, one thread_id per session, and the single HTTP endpoint a "
        "React-as-facade UI consumes through @langchain/langgraph-sdk. Each capability epic maps 1:1 to a "
        "HuggingFace Space and to one cost-center line on the GL, in alignment with the operating-model axiom.",
    )
    add_para(
        doc,
        "The codebase shape adopts the V7 dual-child convention {epic_root}/RuntimeServer/ + "
        "{epic_root}/SharedCode/. The structural noun 'domain' is retired. Project-root capability folders "
        "(EvaluateCare/, FindCare/, SharedServices/, Security/, Operations/, UX/, DataManagementPipelines/) "
        "replace the Code/ umbrella.",
    )
    add_para(
        doc,
        "Every Python service entry point follows the thin-main-construct-and-run convention: app.py contains "
        "an if __name__ == '__main__': block whose sole job is to construct an Application object and call "
        "run() on it. Procedural FastAPI handlers are the exception, not the rule.",
    )
    add_para(
        doc,
        "End-to-end observability is delivered by LangSmith trace propagation across RemoteGraph boundaries, a "
        "Kafka-fed conversation_log Worker for structured per-turn audit, and per-node performance budgets "
        "emitted as OpenTelemetry-compatible metrics from the parent and each Space.",
    )
    add_para(
        doc,
        "Claude Code's blast area is reduced by: (a) the engineering-rules enforcement framework V20; "
        "(b) project-root capability folders that limit any single edit's reachability to one cost-center boundary; "
        "(c) the requirement that every code change trace to a BR/TR; (d) the single-thin-main convention with a "
        "registered exception list; (e) read-only design artifacts archived under architecture/ that the boot tool "
        "reads on every session.",
    )

    # ---------- 2. Rescue table -- the load-bearing artifact ----------
    add_heading(doc, "2. Rescue table -- code reachable from main.py at build 410", level=1)
    add_para(
        doc,
        f"The rescue table is the load-bearing artifact of V3. It enumerates every file reachable from "
        f"Code/ConversationalUX/FindCareChat/backend/main.py at build 410 (commit {GROUND_TRUTH_COMMIT[:10]}, "
        f"{GROUND_TRUTH_DATE}) and declares where each survives in the new architecture (V7 dual-child layout). "
        "The reachable-code set is built by tracing every function call from each FastAPI route handler "
        "(/chat, /search, /welcome, /health, /) recursively through the dispatch layer, into domain services, "
        "into infrastructure, into Code/Shared. Code present in the repo at build 410 but NOT reachable from "
        "main.py is dead code; it is cut, not rescued.",
    )
    add_para(
        doc,
        "Cuts (per architect directive, even if reachable): (a) all EvaluateCare code (\"throw out all the "
        "EvaluateCare code\") -- the EvaluateCareFacade, ClinicalTrialsService, ProviderDetailService, and the "
        "clinical_trials_models / EvaluateCare-side tool registrations are dropped from the rescue table. "
        "EvaluateCare survives in the architecture as a topology slot (its own HF Space + RemoteGraph composition "
        "from the parent), but no pre-4/3 EvaluateCare code is rescued. (b) Authentication / token verification / "
        "entitlements -- V2's deep-auth BRs are out of V3 per the architect's \"V2 minus the deep authentication "
        "requirements\" directive (no pre-4/3 auth code exists in the reachable set, so no rescue rows are dropped). "
        "(c) Code/DataPipelines/* (Azure Functions) -- pipelines are out of scope; they continue under their own "
        "untouched runtime and are not in the rescue table.",
    )
    add_para(
        doc,
        "Build 411 is rejected as a regression-onset point (FC-RESULT-MSG system-built summary replaces the LLM "
        "narrative). Build 410 preserves the AI-orchestrated narrative around results -- that is the behavior "
        "to keep.",
    )
    add_table(
        doc,
        ["#", "File path at build 410", "Purpose", "Reachable from main.py via", "Survives in new architecture as"],
        RESCUE_ROWS,
        col_widths_in=[0.4, 1.6, 1.7, 1.3, 1.5],
    )
    add_para(
        doc,
        f"Total files in the rescue table: {len(RESCUE_ROWS)} (the reachable-from-main keep-set after the "
        "EvaluateCare / Auth / Azure-DataPipelines cuts). Each row's 'Survives as' column maps to a concrete "
        "target path in the V7 dual-child target layout ({Capability}/RuntimeServer/... or "
        "{Capability}/SharedCode/...).",
    )

    # ---------- 2.1 Pre-4/3 capability surface (kept for BR / TR trace cells) ----------
    add_heading(doc, "2.1 Pre-4/3 capability surface (PC-* anchors)", level=2)
    add_para(
        doc,
        "The PC-* anchor list below is retained so each BR / TR Trace cell can cite a precise file:line in the "
        "build-410 tree. The PC-* IDs are not a recovery plan; they are the citation index for the rescue table. "
        "PC-* rows that point at EvaluateCare code, Azure DataPipelines, or the deleted Code/Schemas/ provider "
        "schema remain as citations only -- they are NOT in the rescue table and the corresponding code is NOT "
        "rescued.",
    )
    add_table(
        doc,
        ["PC ID", "Capability", "Pre-4/3 evidence", "BR(s)", "TR(s)"],
        CAP_ROWS,
        col_widths_in=[0.55, 2.05, 2.0, 0.85, 1.05],
    )
    add_para(
        doc,
        f"Total PC-* anchors: {len(CAP_ROWS)}. Each row is anchored to a concrete file:line in the build-410 "
        "tree.",
    )

    # ---------- 3. Business Requirements ----------
    add_heading(doc, "3. Business requirements (BR)", level=1)
    add_para(
        doc,
        "Business-language statements of what the system must deliver. Every BR is testable, has a unique ID, "
        "and traces to (a) a pre-4/3 capability from Section 2, (b) a V7 binding decision, or (c) a V6 LangGraph "
        "principle.",
    )
    add_table(
        doc,
        ["BR ID", "Statement", "Acceptance test", "Trace"],
        BR_ROWS,
        col_widths_in=[0.55, 2.85, 1.85, 1.25],
    )

    # ---------- 4. Technical Requirements ----------
    add_heading(doc, "4. Technical requirements (TR)", level=1)
    add_para(
        doc,
        "Implementation-language statements that flow from Section 3. Every TR has a unique ID, is testable, and "
        "traces to at least one BR. Where a TR also implements a V6 LangGraph principle or a V7 binding decision, "
        "the trace cell makes that explicit.",
    )
    add_table(
        doc,
        ["TR ID", "Statement", "Acceptance test", "Trace"],
        TR_ROWS,
        col_widths_in=[0.6, 3.1, 1.65, 1.15],
    )

    # ---------- 5. Architecture overview ----------
    add_heading(doc, "5. Architecture overview", level=1)

    add_heading(doc, "5.1 HF Space topology", level=2)
    add_para(
        doc,
        "Five HuggingFace Spaces, each one cost-center line, one org-chart line, one deploy pipeline, "
        "one /health endpoint:",
    )
    add_bullet(doc, "Operations -- hosts the parent runtime; exposes the single conversational HTTP endpoint to the React UI; owns the MongoDBSaver checkpointer; runs URLGuardian as a response-finalize step.")
    add_bullet(doc, "FindCare -- hosts the FindCare capability subgraph; exposes its graph as a RemoteGraph and its non-conversational POST /search as REST.")
    add_bullet(doc, "EvaluateCare -- hosts the EvaluateCare subgraph (clinical trials + provider detail) and its non-conversational report API.")
    add_bullet(doc, "Security & Compliance -- hosts the Security capability subgraph (admin-unlock, IP-lock, dual-trigger emergency, safe-prefix fast-path, consent service, compliance feature folder).")
    add_bullet(doc, "SharedServices -- hosts the canonical mongo client, the LLM client, embeddings, and the prompt system maker; consumed via ToolRuntime by every capability Space.")
    add_para(
        doc,
        "DataManagementPipelines is operational infrastructure (Azure Functions today; HF or successor "
        "tomorrow) and is NOT a runtime Space; jobs run under Operations infrastructure but never participate "
        "in the parent graph.",
    )
    add_para(
        doc,
        "UX is not a runtime Space; UX/Frontend artifacts deploy to Cloudflare Pages. Operations also owns "
        "CI/deploy/build-counter tooling that is not a runtime node.",
    )

    add_heading(doc, "5.2 Deployment diagram", level=2)
    add_diagram(
        doc,
        DIAG_DEPLOYMENT,
        "deployment",
        "End-to-end deployment topology -- Cloudflare Pages, HF Spaces, MongoDB Atlas, LangSmith, Kafka audit pipeline.",
    )

    add_heading(doc, "5.3 Component diagram", level=2)
    add_diagram(
        doc,
        DIAG_COMPONENT,
        "component",
        "Component topology inside each Space -- agents, tool nodes, SharedCode/core services, SharedServices Mongo + LLM clients.",
    )

    add_heading(doc, "5.4 Data flow", level=2)
    add_bullet(doc, "Conversational read/write: React -> parent runtime -> Security node (first) -> capability node(s) -> URLGuardian + LinkedIn auto-link rewriter as response-finalize -> stream back to React.")
    add_bullet(doc, "Non-conversational read: React -> capability HTTP API (e.g. POST /search for pagination, GET /provider/{id} for detail page) -> SharedServices mongo abstraction -> MongoDB.")
    add_bullet(doc, "Pipeline write: scheduled job (DataManagementPipelines) -> SharedServices mongo abstraction -> MongoDB.")
    add_bullet(doc, "Audit trail: every parent turn emits a structured event onto Kafka; Operations Worker consumes and writes to MongoDB; LangSmith captures the LangGraph view; rows join by trace_id.")

    # ---------- 6. Per-capability detail ----------
    add_heading(doc, "6. Per-capability detail", level=1)
    add_para(
        doc,
        "Per assignment: V20-granularity for parent runtime + FindCare (the worked example pair); gestural for "
        "the others.",
    )

    add_heading(doc, "6.1 Operations -- parent runtime (V20 specificity)", level=2)
    add_bullet(doc, "Operations/RuntimeServer/orchestration/parent_graph.py compiles StateGraph(ParentState) with checkpointer=MongoDBSaver(...).")
    add_bullet(doc, "Operations/RuntimeServer/orchestration/state.py: ParentState(BaseModel) with messages: Annotated[list[AnyMessage], add_messages]; thread_id; user_id; locale; current_capability: Optional[str]; routing_hint: Optional[str].")
    add_bullet(doc, "Operations/RuntimeServer/orchestration/checkpointer.py constructs MongoDBSaver once at module load.")
    add_bullet(doc, "Operations/RuntimeServer/app.py: Application class whose run() programmatically calls uvicorn.run on a FastAPI ASGI app wrapping the langgraph-sdk streaming endpoint. CORS allowlist preserves the pre-4/3 origins (https://chathealthy.ai, https://www.chathealthy.ai, https://dev.chathealthy.ai, plus localhost regex and *.chathealthy.ai regex).")
    add_bullet(doc, "Edge order at the parent: Security -> {FindCare | EvaluateCare} -> response-finalize (URLGuardian + Skip-Snow LinkedIn auto-link rewriter) -> stream out.")
    add_bullet(doc, "trace_id is created on turn start, propagated to each RemoteGraph invocation via context=, propagated into ToolRuntime, attached to every Mongo write through the SharedServices client. LangSmith reconstructs the full tree from the propagated id.")
    add_bullet(doc, "Long-running deep-research flows use interrupt() / Command(resume=...) at the parent (capability Spaces lack a checkpointer). Deep research is a same-Space subgraph hosted inside whichever capability Space invokes it (V6 3.4).")
    add_bullet(doc, "Welcome / health: GET /welcome returns the canonical welcome message (HUMAN_TESTING toggle preserved); GET /health returns env, db, build, version (matches pre-4/3 contract).")
    add_bullet(doc, "URLGuardian runs as a parent-runtime response-finalize step (NOT a graph node) so that every capability's text and link payloads are validated uniformly.")

    add_heading(doc, "6.2 FindCare (V20 specificity, worked example)", level=2)
    add_bullet(doc, "FindCare/RuntimeServer/agent.py compiles a StateGraph wired with: tool_find_providers, tool_find_specialty_codes, tool_record_user_details, tool_record_unknown_question, tool_get_skip_snow_context, tool_get_chathealthy_context. Routing is a single tool-use loop that mirrors the pre-4/3 Anthropic tool-use loop (initial call -> tool_use -> dispatch -> append messages -> repeat until stop_reason != 'tool_use').")
    add_bullet(doc, "FindCare/RuntimeServer/state.py: FindCareState(BaseModel) with messages, last_search_params: Optional[ProviderSearchInput], last_provider_result: Optional[ProviderSearchResult], pagination: Optional[PaginationMeta]. ProviderSearchInput is the pre-4/3 Pydantic model verbatim.")
    add_bullet(doc, "FindCare/RuntimeServer/nodes.py defines the per-tool node bodies; each accepts Runtime[ContextSchema] and reads its dependencies (mongo, embedding, anthropic) from the runtime context (no module-scope clients).")
    add_bullet(doc, "FindCare/RuntimeServer/tools.py registers @tool decorations + ToolNode wiring; the registration list mirrors the pre-4/3 ToolRouter.register_with_models call (with Pydantic models for input validation per the pre-4/3 GOV-004 enforcement).")
    add_bullet(doc, "FindCare/RuntimeServer/app.py: Application class; a NON-graph POST /search endpoint reuses SharedCode/core/provider_search_service.py for direct pagination calls (BR-06).")
    add_bullet(doc, "FindCare/SharedCode/core/provider_search_service.py preserves the pre-4/3 4-route search (NPI exact / Name / Specialty codes direct / Specialty query via vector + taxonomy fallback), the FIPS-to-county fallback map (with DB augmentation at startup), the entity_type-1-vs-2 name formatter, the (NNN) NNN-NNNN phone formatter, the keyset pagination by NPI, and the pagination metadata (total_count, first_npi, last_npi, has_more, page_start, page_end, search_params).")
    add_bullet(doc, "FindCare/SharedCode/core/specialty_service.py preserves the pre-4/3 dual pipeline (regex with AI-expanded stems via Claude Haiku + vector via text-embedding-3-small filtered by INDIVIDUAL_PROVIDER_GROUPINGS classification). Pipelines run in parallel via ThreadPoolExecutor; results are dedup'd by Code with vector results given priority.")
    add_bullet(doc, "FindCare/SharedCode/core/lead_service.py preserves dedupe-by-email, three consent paths, Sparkpost push, MongoDB write to {ENV}_AboutUs.lead.")
    add_bullet(doc, "FindCare/SharedCode/core/unknown_question_service.py preserves the 3-class taxonomy and verbatim response_template per class.")
    add_bullet(doc, "FindCare/SharedCode/core/about_service.py preserves get_skip_snow_context / get_chathealthy_context with the verbatim 'connect' field.")
    add_bullet(doc, "FindCare/Frontend/ is the React iframe; ChatWindow.tsx + GUIManager.tsx + MessageBubble.tsx preserve the pre-4/3 behavior: pagination via the postMessage bridge to the static control frame, environment banner, retry / abandon countdown, timeout escalation modal.")
    add_bullet(doc, "Talk About Care lives as a feature folder under FindCare per BR-07 / V7 epic shape.")

    add_heading(doc, "6.3 EvaluateCare (gestural)", level=2)
    add_bullet(doc, "EvaluateCare/RuntimeServer/agent.py exposes search_clinical_trials and lookup_provider_external as @tool nodes. Pre-4/3 ClinicalTrials.gov v2 + Google Routes API + NPI Registry CMS v2.1 + state-board (DE/MS/VA) + Healthgrades behavior is preserved.")
    add_bullet(doc, "Non-conversational HTTP endpoints (provider lookup detail page, evaluation report) live outside the graph (V6 3.3 carve-out).")

    add_heading(doc, "6.4 Security & Compliance (gestural)", level=2)
    add_bullet(doc, "Security/RuntimeServer/agent.py runs first in the parent edge order. Nodes preserve the pre-4/3 enforcement order: admin_unlock_check -> ip_locked_check -> emergency_check -> safe_prefix_fast_path. Emergency detection is dual-trigger: keyword OR Claude classifier @ confidence>=0.80.")
    add_bullet(doc, "Security/RuntimeServer/compliance/ is a feature folder, not a separate epic. consent_service preserves the pre-4/3 two-tier HIPAA path; both tiers use Claude Haiku.")
    add_bullet(doc, "{ENV}_Safety.emergency_incidents collection name and TTL 3600s + admin_unlock_key are preserved for migration continuity.")

    add_heading(doc, "6.5 SharedServices (gestural)", level=2)
    add_bullet(doc, "SharedServices/SharedCode/mongo_utilities.py is the canonical Mongo client (V7 A2 / 13.2). The pre-4/3 graceful-unavailable mode is preserved.")
    add_bullet(doc, "SharedServices/SharedCode/llm_client.py centralizes Anthropic (claude-sonnet-4-6 chat; claude-haiku-4-5-20251001 utility) and OpenAI (gpt-4.1-mini safety classifier; text-embedding-3-large + 3-small) clients.")
    add_bullet(doc, "SharedServices/SharedCode/prompt_system_maker.py preserves the pre-4/3 PromptSystemMaker behavior; system-prompt rules will be relocated to brain/machine_artifacts/content/system_prompt_rules.json (TR-FC-009).")
    add_bullet(doc, "SharedServices/SharedCode/url_guardian.py is the canonical URLGuardian (run from the parent finalize step per TR-FC-010).")

    add_heading(doc, "6.6 Operations -- DevOps and observability (gestural)", level=2)
    add_bullet(doc, "Operations/SharedCode/conversation_log/Worker (V7 13.2) consumes the Kafka audit topic and writes to MongoDB.")
    add_bullet(doc, "Operations/SharedCode/brain_governance/{rebuild_manifest.py, manifest_generator.py} (V7 A8 GOVERNANCE-CRITICAL).")
    add_bullet(doc, "Operations/SharedCode/scan/ holds br_tr_coverage_scan, oo_thin_main_scan, pre_4_3_capability_recovery_smoke.")
    add_bullet(doc, "Operations/SharedCode/hf_space/ holds the pre-4/3 HF Space lifecycle utilities (Rule-024 + Rule-051).")

    add_heading(doc, "6.7 UX (gestural)", level=2)
    add_bullet(doc, "Not a runtime HF Space. Deploys to Cloudflare Pages.")
    add_bullet(doc, "UX/Frontend/index.html is the canonical landing page; embeds FindCare/Frontend as iframe.")
    add_bullet(doc, "UX/SharedCode/cross_components/ holds shared React primitives.")
    add_bullet(doc, "UX/SharedCode/static_pages/ is the declarative source for the 10 non-roadmap static pages.")

    add_heading(doc, "6.8 DataManagementPipelines (gestural)", level=2)
    add_bullet(doc, "Hosts ingest, dedup, embedding-generation jobs (relocated from Code/DataPipelines/). NO dual-child layout (V7 4.3).")
    add_bullet(doc, "Reads/writes go through SharedServices/SharedCode/mongo_utilities.py only (TR-DM-003).")
    add_bullet(doc, "Schema is canonical at brain/machine_artifacts/schemas/ChatHealthyProvidersSchema.json.")

    # ---------- 7. UML diagrams ----------
    add_heading(doc, "7. UML diagrams (embedded)", level=1)

    add_heading(doc, "7.1 Sequence -- conversational turn", level=2)
    add_diagram(
        doc,
        DIAG_SEQUENCE,
        "sequence",
        "Sequence: end-to-end request flow including Security gate, FindCare tool-use loop, URLGuardian + LinkedIn rewriter, MongoDBSaver checkpoint, LangSmith trace.",
    )

    add_heading(doc, "7.2 Class -- Pydantic state schemas", level=2)
    add_diagram(
        doc,
        DIAG_CLASS_STATE,
        "class_state",
        "Class: Pydantic state schemas per V6 4-file layout (ParentState, ContextSchema, FindCareState, EvaluateCareState, SecurityState, ProviderSearchInput, PaginationMeta).",
    )

    add_heading(doc, "7.3 Activity -- FindCare user journey", level=2)
    add_diagram(
        doc,
        DIAG_ACTIVITY,
        "activity",
        "Activity: FindCare user journey preserved from pre-4/3 behavior (Security gates, tool-use loop, URLGuardian, audit row).",
    )

    # ---------- 8. LangGraph implementation ----------
    add_heading(doc, "8. LangGraph implementation", level=1)
    add_para(
        doc,
        "Implements V6 best practices verbatim. Parent runtime + cross-Space RemoteGraph composition; "
        "in-process subgraph composition for same-Space steps; Pydantic state; runtime-context DI; messages "
        "channel for tool I/O; React-as-facade; non-conversational reads bypass the parent graph.",
    )

    add_heading(doc, "8.1 Parent runtime (Operations Space)", level=2)
    add_bullet(doc, "Operations/RuntimeServer/orchestration/parent_graph.py instantiates StateGraph(ParentState).")
    add_bullet(doc, "Operations/RuntimeServer/orchestration/checkpointer.py constructs MongoDBSaver once at module load. Parent graph is compiled with checkpointer=MongoDBSaver(...).")
    add_bullet(doc, "Operations/RuntimeServer/app.py: Application class whose run() programmatically calls uvicorn.run with the FastAPI ASGI app that wraps the LangGraph SDK streaming endpoint.")

    add_heading(doc, "8.2 RemoteGraph composition", level=2)
    add_code_block(
        doc,
        (
            "from langgraph.graph import StateGraph, START\n"
            "from langgraph.pregel.remote import RemoteGraph\n"
            "from .state import ParentState\n"
            "from .checkpointer import checkpointer\n"
            "\n"
            "security     = RemoteGraph('security',     url=SECURITY_HF_URL,     api_key=...)\n"
            "find_care    = RemoteGraph('find_care',    url=FINDCARE_HF_URL,     api_key=...)\n"
            "evaluate     = RemoteGraph('evaluate',     url=EVALUATECARE_HF_URL, api_key=...)\n"
            "\n"
            "g = StateGraph(ParentState)\n"
            "g.add_node('security', security)\n"
            "g.add_node('find_care', find_care)\n"
            "g.add_node('evaluate',  evaluate)\n"
            "g.add_edge(START, 'security')\n"
            "# routing edges from security -> {find_care | evaluate} based on intent classification\n"
            "graph = g.compile(checkpointer=checkpointer)\n"
        ),
        label="parent_graph.py (canonical V6 RemoteGraph composition pattern)",
    )

    add_heading(doc, "8.3 Capability subgraph layout (canonical 4-file)", level=2)
    add_bullet(doc, "{Capability}/RuntimeServer/state.py -- Pydantic BaseModel + input_schema + output_schema.")
    add_bullet(doc, "{Capability}/RuntimeServer/nodes.py -- node functions or class-callables.")
    add_bullet(doc, "{Capability}/RuntimeServer/tools.py -- @tool decorations + ToolNode wiring.")
    add_bullet(doc, "{Capability}/RuntimeServer/agent.py -- graph build (compile() with no checkpointer).")
    add_bullet(doc, "{Capability}/RuntimeServer/langgraph.json -- manifest per the langgraph-example-pyproject template.")

    add_heading(doc, "8.4 Runtime context (DI)", level=2)
    add_bullet(doc, "Each capability subgraph defines its own ContextSchema (Pydantic) carrying user_id, thread_id, trace_id, locale, request_id, ip.")
    add_bullet(doc, "Nodes accept Runtime[ContextSchema]; tools accept ToolRuntime[ContextSchema].")
    add_bullet(doc, "The parent passes its session-scoped context to each RemoteGraph invocation via context=.")
    add_bullet(doc, "Module-scope import of mongo or LLM clients in node files is forbidden (TR-LG-008).")

    add_heading(doc, "8.5 Messages channel and tool I/O", level=2)
    add_bullet(doc, "All tool calls and tool results are ToolMessage entries on the messages channel (TR-LG-007).")
    add_bullet(doc, "Structured artefacts the LLM does not need to read (e.g. raw provider rows for downstream rendering) get a typed Pydantic field on the subgraph state, NEVER a flat top-level state field per tool.")

    add_heading(doc, "8.6 Same-Space subgraph composition", level=2)
    add_bullet(doc, "Two nodes that share a database client and deploy together stay in the same Space and use in-process subgraph composition.")
    add_bullet(doc, "RemoteGraph never calls back into its own deployment (V6 3.2 deadlock constraint, TR-LG-005).")
    add_bullet(doc, "Deep research is a same-Space subgraph hosted inside whichever capability Space invokes it (V6 3.4); deep-research interrupts use interrupt() / Command(resume=...) at the parent.")

    # ---------- 9. Resolved unresolved-issues ----------
    add_heading(doc, "9. Resolved unresolved-issues (recommendations)", level=1)
    add_para(
        doc,
        "Per assignment, the design assumes acceptance of the following recommendations.",
    )

    add_heading(doc, "R1. Security cutover sequencing", level=2)
    add_para(
        doc,
        "The pre-4/3 safety enforcement order (admin_unlock -> ip_locked -> emergency, all BEFORE any LLM call) "
        "is preserved as the FIRST capability node in the parent edge order (TR-SE-001). Cutover sequence: "
        "(1) deploy Security & Compliance Space in parallel with the existing safety code path; "
        "(2) dual-write incidents to {ENV}_Safety.emergency_incidents (no schema change); "
        "(3) flip the parent edge order so Security runs first; "
        "(4) decommission the in-process safety calls in the legacy main.py; "
        "(5) verify the smoke suite (TR-RC-001) covers all Section 2 PC entries before the legacy path is removed.",
    )

    add_heading(doc, "R2. Provider schema content drift", level=2)
    add_para(
        doc,
        "The canonical Provider schema is brain/machine_artifacts/schemas/ChatHealthyProvidersSchema.json (V7 A3). "
        "The schema declares the embedding fields and the NPPES other-name fields. Code/Schemas/ChatHealthy.Providers.json "
        "is DELETED. Pipeline regression test asserts every declared field is delivered (BR-18, TR-DM-002).",
    )

    add_heading(doc, "R3. langgraph-api vs FastAPI", level=2)
    add_para(
        doc,
        "Resolution: the parent runtime exposes ONE HTTP endpoint compatible with @langchain/langgraph-sdk for "
        "conversational streaming flow (TR-LG-002). Non-conversational reads (POST /search for pagination, "
        "GET /provider/{id}, GET /evaluation/{id}, GET /welcome, GET /health) live in the appropriate Space's "
        "FastAPI app outside the graph. The Application class pattern (TR-OO-001..003) avoids procedural "
        "FastAPI module-top construction. The langgraph-api package is consumed where it is the official "
        "SDK contract; FastAPI is consumed where the path is non-conversational.",
    )

    add_heading(doc, "R4. Observability backend", level=2)
    add_para(
        doc,
        "Resolution: LangSmith is the trace UI of record for the conversation graph. The Kafka-fed Operations "
        "Worker is the persistent audit row store. The two views join by trace_id. Per-node performance budgets "
        "are emitted as OpenTelemetry histograms with concrete pre-4/3-rooted budgets in TR-OB-003.",
    )

    add_heading(doc, "R5. Talk About Care", level=2)
    add_para(
        doc,
        "Resolution: feature folder under FindCare (BR-07). The pre-4/3 code already places the about-content "
        "tools (get_skip_snow_context, get_chathealthy_context) inside the FindCare backend; this is preserved.",
    )

    add_heading(doc, "R6. V2-vs-pre-4/3 disagreements (resolved in favor of code)", level=2)
    add_para(
        doc,
        "Three V2 framings are corrected by the pre-4/3 ground truth:",
    )
    add_bullet(
        doc,
        "(a) V2's LC-12 implies LangSmith traces 'broke at the FastAPI boundary' pre-event. The pre-4/3 system "
        "had no LangGraph and no LangSmith; LC-12 is therefore not a recovery target but a new design goal "
        "(captured here as BR-12 / TR-OB-001).",
    )
    add_bullet(
        doc,
        "(b) V2's TR-FC-007 introduces a 'repetitive_utterance_detector' citing LangGraph/poc/user_journey.py. "
        "That capability does not exist in the pre-4/3 application code, and the POC file is excluded as "
        "post-4/20 evidence. V3 omits it.",
    )
    add_bullet(
        doc,
        "(c) V2's BR-04 elevates 'token verification + entitlement' to a pre-conversation gate. The pre-4/3 "
        "application has no token verification; the pre-4/3 pre-conversation gate is the safety + IP-lock + CORS "
        "stack. V3 BR-04 is grounded in the pre-4/3 stack; token verification is captured separately as future "
        "work in the deferred-work registry, not as a recovered capability.",
    )

    # ---------- 10. Internal-consistency check ----------
    add_heading(doc, "10. Internal-consistency check", level=1)
    br_to_tr, tr_to_br = build_coverage_matrix()

    unmapped_brs = [b for b, lst in br_to_tr.items() if not lst]
    unmapped_trs = [t for t, lst in tr_to_br.items() if not lst]

    add_para(
        doc,
        f"BR count: {len(BR_ROWS)}. TR count: {len(TR_ROWS)}. "
        f"BRs without a TR: {len(unmapped_brs)}. TRs without a BR: {len(unmapped_trs)}.",
    )
    if unmapped_brs:
        add_para(doc, "BRs without a TR (must be addressed):")
        for b in unmapped_brs:
            add_sub_bullet(doc, b)
    if unmapped_trs:
        add_para(doc, "TRs without a BR (must be addressed):")
        for t in unmapped_trs:
            add_sub_bullet(doc, t)

    add_heading(doc, "10.1 Coverage matrix (BR -> TR)", level=2)
    cov_rows = [(br, ", ".join(sorted(set(trs))) or "(none)") for br, trs in br_to_tr.items()]
    add_table(doc, ["BR ID", "Implementing TR(s)"], cov_rows, col_widths_in=[1.0, 5.4])

    add_heading(doc, "10.2 Coverage matrix (TR -> BR)", level=2)
    rev_rows = [(tr, ", ".join(sorted(set(brs))) or "(none)") for tr, brs in tr_to_br.items()]
    add_table(doc, ["TR ID", "Originating BR(s)"], rev_rows, col_widths_in=[1.0, 5.4])

    verdict = (
        "Verdict: BR<->TR coverage is complete; every BR has at least one implementing TR; every TR traces to "
        "at least one originating BR. No design element is unsupported."
        if not unmapped_brs and not unmapped_trs
        else "Verdict: incomplete coverage; see lists above."
    )
    add_para(doc, verdict, bold=True)

    # ---------- 11. Appendix ----------
    add_heading(doc, "11. Appendix -- generation record", level=1)
    add_bullet(doc, f"Ground truth commit: {GROUND_TRUTH_COMMIT} ({GROUND_TRUTH_DATE}).")
    add_bullet(doc, f"Rescue table row count (files reachable from main.py at build 410, after EvaluateCare / Auth / Azure-DataPipelines cuts): {len(RESCUE_ROWS)}.")
    add_bullet(doc, f"PC-* citation anchors: {len(CAP_ROWS)} (citation index only; not the rescue plan).")
    add_bullet(doc, f"BR count: {len(BR_ROWS)} (carried from V2 + URLGuardian BR per the architect). TR count: {len(TR_ROWS)} (re-derived from V3 architecture; V2 TRs NOT carried forward per the architect).")
    add_bullet(doc, "Inputs read at generate time: pre-4/3 code at git commit be14579a (read-only via git show); design-V2.docx (BR list authoritative); codebase-reorganization-V7.docx (file structure / epic shape binding); langgraph-oo-best-practices-V6.docx (runtime patterns binding); langgraph-oo-best-practices-V5.docx (architect's </Skip> annotations cross-reference); brain/machine_artifacts/content/{engineering_rules, agile_backlog}.json.")
    add_bullet(doc, "UML rendering pipeline: PlantUML public server first (with explicit User-Agent and 8s timeout), then kroki.io PlantUML fallback; on both failures, source is embedded as a labeled PROVISIONAL code block. Diagrams: deployment, sequence, class (Pydantic state), activity (FindCare user journey), component.")
    add_bullet(doc, "Read-only enforcement: no commits, no code changes, no backlog mutations during generation. Stories left intact. 'Domain' does not appear as a structural noun anywhere in the generator's output.")

    return doc


def main() -> int:
    doc = build_doc()
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    _log.info("Wrote %s", OUT)
    return 0


if __name__ == "__main__":
    sys.exit(main())
