# Copyright (c) 2026 Skip Snow. All rights reserved.
# Licensed under the FindCare Evaluation License (FEL-1.0).
#
# generate_pre_sql_capability_audit_V2.py
#
# V2 of the pre-event capability audit. Tighter scope than V1:
#   - PRE-4/3/2026 ONLY (not pre-4/19; the architect ratified that anything
#     from 4/3/2026 onward is post-regression work).
#   - BR-LEVEL ONLY (no TR-level concerns).
#   - DOES NOT cite anything from LangGraph/poc/ as a pre-event capability
#     (the POC is dated 2026-04-20 in its own header by self-attribution).
#
# Pre-4/3 reference: latest commit before 2026-04-03 00:00:00 -0700.
#   git log --before="2026-04-03 00:00:00" -n 1
#   -> be14579a9200ba23227181b716589be6267ef63b | 2026-04-02 23:48:24 -0700
#   -> commit subject: "[Attended] Build 410: Banner in parent only,
#                       hidden in iframe. DR-014 commit every test."
#
# Read-only:
#   - Pre-4/3 source inspected via `git show be14579a:<path>` only.
#   - No git checkout, no commits, no working-tree mutations.
#   - The only writes this generator performs are this generator file
#     itself and the deliverable .docx alongside it.
#
# Deliverable:
#   pre-sql-capability-audit-V2.docx (sibling of this script)

from __future__ import annotations

import os
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---------------------------------------------------------------------------
# Pre-4/3 reference commit
# ---------------------------------------------------------------------------
PRE_43_COMMIT = "be14579a9200ba23227181b716589be6267ef63b"
PRE_43_BUILD = 410
PRE_43_DATE = "2026-04-02 23:48:24 -0700"
PRE_43_SUBJECT = (
    "[Attended] Build 410: Banner in parent only, hidden in iframe. "
    "DR-014 commit every test."
)

# ---------------------------------------------------------------------------
# Business capability inventory (BR-level only)
# Each item: (cap_id, capability_name, source, user_facing_description,
#             v2_coverage, recommended_br_text)
# Source paths are read at commit be14579a via `git show`.
# ---------------------------------------------------------------------------
CAPABILITIES = [
    # ---------------- conversational core ----------------
    (
        "BC-01",
        "Natural-language conversational find-care",
        "Code/ConversationalUX/FindCareChat/backend/main.py:283",
        "User types a care need in natural language and receives a ranked "
        "list of providers with explanations from the assistant.",
        "FULL: BR-01",
        "",
    ),
    (
        "BC-02",
        "Pre-conversation safety gate (admin-unlock, IP-lock, emergency)",
        "Code/ConversationalUX/FindCareChat/backend/main.py:303 + "
        "domain/shared/safety/safety_service.py",
        "On every user turn the system checks first whether the user is "
        "trying an admin unlock key, then whether their IP is currently "
        "locked, then whether the message itself is an emergency; on "
        "emergency the conversation is locked with a 911 message.",
        "PARTIAL: BR-04, BR-16",
        "BR-04 names a security gate but does not bind its order; BR-16 "
        "names emergency lockout but does not bind the per-turn ordering. "
        "Add: 'Every conversational turn MUST be evaluated in this "
        "deterministic order BEFORE any LLM call: (1) admin-unlock attempt "
        "with the configured admin key (default unlock$123); (2) IP-lock "
        "check; (3) emergency classifier (keyword OR AI). On emergency, "
        "the user IP is locked for the configured duration (default 24h) "
        "and the canonical 911 emergency message is returned.'",
    ),
    (
        "BC-03",
        "Dual-trigger emergency detection (AI classifier OR keyword match)",
        "Code/ConversationalUX/FindCareChat/backend/domain/shared/safety/"
        "safety_service.py:50",
        "The system flags a message as an emergency when EITHER an AI "
        "classifier says so OR an emergency keyword is present, with a "
        "small allowlist of safe phrasing prefixes that skip the "
        "classifier (e.g. 'tell me about', 'find me').",
        "PARTIAL: BR-16",
        "BR-16 names emergency lockout consequences but not the detection "
        "method. Add: 'Emergency detection MUST combine an AI classifier "
        "and a keyword match using OR logic; obvious non-emergency "
        "phrasings (informational openers like \"tell me about\", "
        "\"find me\", \"who is\") MUST be allowed to skip the AI "
        "classifier as a fast path; the keyword list is configurable from "
        "brain artifact emergency_keywords.json.'",
    ),
    (
        "BC-04",
        "Periodic follow-up consent prompt every 5 user turns",
        "Code/ConversationalUX/FindCareChat/backend/main.py:308",
        "Every 5th user message in a conversation, the system prompts the "
        "user to consider sharing contact info or saving the conversation.",
        "GAP",
        "Add: 'After every 5 user utterances in a single session, the "
        "system MUST insert a follow-up prompt offering the user the "
        "option to leave contact details or save the conversation, "
        "without coercion. The cadence (5) is configurable.'",
    ),
    (
        "BC-05",
        "Cross-product handoff: ask FindCare assistant about clinical trials",
        "Code/ConversationalUX/FindCareChat/backend/main.py:202 + "
        "application/facades/evaluate_care_facade.py",
        "Within the same conversation, a user can ask for clinical trials "
        "for a condition; the FindCare assistant calls the EvaluateCare "
        "clinical-trials capability and returns trials filtered to "
        "RECRUITING, with optional drive distance and time from the "
        "user's location.",
        "PARTIAL: BR-33",
        "BR-33 describes a separate clinical-trials surface (right-panel "
        "diagnosis hierarchy + trial picker) but not the in-conversation "
        "tool. Add: 'The conversational assistant MUST be able to answer "
        "in-line clinical-trials questions by invoking a tool that "
        "queries ClinicalTrials.gov for RECRUITING studies for the "
        "user's condition, optionally enriched with drive-time and "
        "distance from the user's stated or inferred location.'",
    ),
    (
        "BC-06",
        "Cross-product handoff: provider detail lookup with state-board links",
        "Code/ConversationalUX/FindCareChat/backend/main.py:203 + "
        "domain/evaluate_care_quality/provider_detail_service.py",
        "When the user asks about a specific provider by name (and "
        "optionally NPI/state), the assistant returns the canonical "
        "NPI Registry record plus links to the state licensing board "
        "(DE/MS/VA) and a generic Healthgrades search URL.",
        "GAP",
        "Add: 'The conversational assistant MUST be able to look up a "
        "named provider and return: (a) the NPI Registry canonical "
        "record (name, credential, specialty, license, license state, "
        "address, phone, enumeration date); (b) a link to the state "
        "licensing board for the supported states (DE, MS, VA); (c) a "
        "Healthgrades search link constructed from the provider name. "
        "Sources MUST be cited in the response so the user can verify.'",
    ),
    (
        "BC-07",
        "Lead capture: assistant offers to record contact info",
        "Code/ConversationalUX/FindCareChat/backend/main.py:204 + "
        "domain/shared/lead_capture/lead_service.py",
        "Within the conversation, the assistant can offer to record the "
        "user's email, name, and a free-text note so a ChatHealthy team "
        "member can follow up. The user is deduplicated by email.",
        "FULL: BR-18, BR-19",
        "",
    ),
    (
        "BC-08",
        "Two-tier consent cascade (verbatim -> de-identified -> contact only)",
        "Code/ConversationalUX/FindCareChat/backend/domain/shared/consent/"
        "consent_service.py",
        "Before saving a conversation the user is offered three explicit "
        "consent levels in order: (1) verbatim transcript including any "
        "PHI; (2) a de-identified summary; (3) contact-fields only. If "
        "the user refuses all three, no record is kept.",
        "FULL: BR-18",
        "",
    ),
    (
        "BC-09",
        "Unanswerable-question handling with three classes",
        "Code/ConversationalUX/FindCareChat/backend/main.py:205 + "
        "domain/shared/unknowns/unknown_question_service.py",
        "When the system cannot or will not answer, the user sees a "
        "verbatim refusal categorized as 'medical_advice' (clinical "
        "refusal), 'irrelevant' (out-of-scope), or "
        "'healthcare_capability' (legitimate but unsupported), each with "
        "an offer to record the question so the product can improve.",
        "PARTIAL: BR-17",
        "BR-17 names the three branches in business prose. Tighten to: "
        "'Unanswerable-question handling MUST classify the utterance "
        "into exactly one of three classes (medical_advice, irrelevant, "
        "healthcare_capability) and present the class-specific verbatim "
        "refusal text to the user. The verbatim text is configurable "
        "from the brain. Recording the question is gated by the consent "
        "cascade in BC-08.'",
    ),
    (
        "BC-10",
        "About-the-product context tools (Skip Snow, ChatHealthy)",
        "Code/ConversationalUX/FindCareChat/backend/main.py:206-207 + "
        "domain/shared/content/about_service.py",
        "When the user asks who Skip Snow is or what ChatHealthy is, the "
        "assistant returns a configured business-plan summary, the "
        "Anthropic principles ChatHealthy follows, and a contact block "
        "with verbatim LinkedIn / email links that MUST be rendered as "
        "live links.",
        "FULL: TR-FC-006 (BR-10 governance)",
        "",
    ),
    (
        "BC-11",
        "Unreachable-URL defanging in assistant responses",
        "Code/ConversationalUX/FindCareChat/backend/main.py:347 + "
        "url_guardian.py",
        "Before any assistant response is shown, every URL it contains "
        "is verified to be reachable; URLs that fail are removed (text "
        "kept) so the user never sees a broken link.",
        "GAP",
        "Add: 'Every assistant response that reaches the user MUST first "
        "be passed through a URL-validation step that probes each URL "
        "and removes (or follows the redirect for) URLs that fail; the "
        "user-visible text remains intact, only the link is stripped. "
        "Trusted browser-only domains (e.g. healthgrades.com search "
        "URLs) and SPA shells the system already knows about MAY be "
        "exempted.'",
    ),
    (
        "BC-12",
        "Provider-search tool callable by the assistant",
        "Code/ConversationalUX/FindCareChat/backend/main.py:200 + "
        "domain/find_care/provider_search_service.py",
        "The assistant can search for providers given a colloquial "
        "specialty query, state, city/county, name, or NPI; the result "
        "is a paginated list (25 per page) with first/last NPI cursors.",
        "FULL: BR-01, BR-23, BR-24",
        "",
    ),
    (
        "BC-13",
        "Specialty identification tool callable by the assistant",
        "Code/ConversationalUX/FindCareChat/backend/main.py:201 + "
        "domain/find_care/specialty_service.py",
        "The assistant can convert a colloquial term ('shrink', 'foot "
        "doctor') into a list of NUCC specialty codes restricted to "
        "individual human providers (excludes facilities and groups).",
        "PARTIAL: BR-23",
        "BR-23 describes the user-visible filter UI but does not bind "
        "the underlying capability. Add: 'The system MUST support "
        "translating a colloquial care term to one or more NUCC "
        "specialty codes; results MUST be restricted to individual human "
        "provider groupings (no facilities, no agencies); both regex "
        "and semantic search MUST be combined and merged so neither "
        "alone is dispositive.'",
    ),
    (
        "BC-14",
        "Geographic scope: only DE / MS / VA returned providers",
        "Code/ConversationalUX/FindCareChat/backend/domain/find_care/"
        "provider_search_service.py:19",
        "If the user asks for providers outside DE, MS, or VA, the "
        "system tells the user that only those states are currently "
        "supported and offers to record the user's interest for "
        "future coverage.",
        "PARTIAL: BR-32",
        "BR-32 says all remaining states must be loaded eventually. "
        "Add: 'Until the pipeline (BR-32 / BR-31) loads remaining "
        "states, the conversational system MUST reject states outside "
        "the supported set with a user-facing 'we will note your "
        "interest' message and MUST log the state-of-interest for "
        "follow-up.'",
    ),
    (
        "BC-15",
        "Pagination of provider results (25-at-a-time, NPI cursor)",
        "Code/ConversationalUX/FindCareChat/backend/main.py:255-263 + "
        "domain/find_care/provider_search_service.py:147",
        "The user sees the first 25 providers with a 'has more' "
        "indicator, an explicit page_start/page_end label, and a way "
        "to fetch the next page without re-running the LLM.",
        "PARTIAL: BR-24",
        "BR-24 says 'first 25 providers ... cursor to scroll'. Tighten: "
        "'Provider results MUST be paginated 25 per page with a stable "
        "cursor (NPI keyset); each page MUST tell the user the total "
        "count, the current page range (page_start, page_end), and "
        "whether more results exist.'",
    ),
    (
        "BC-16",
        "Direct provider search bypassing the LLM (POST /search)",
        "Code/ConversationalUX/FindCareChat/backend/main.py:264-268",
        "The user-facing UI can fetch additional pages of provider "
        "results without paying for an LLM round-trip, by posting the "
        "same parameters that the LLM tool received.",
        "GAP",
        "Add: 'The system MUST expose a non-conversational provider "
        "search surface that accepts the same parameters as the "
        "in-conversation provider-search tool and returns identical "
        "results; this surface is what the UI uses for pagination, "
        "filter-apply, and any deterministic re-fetch that does not "
        "require the LLM.'",
    ),
    (
        "BC-17",
        "Welcome message that is environment-aware (UAT report in test mode)",
        "Code/ConversationalUX/FindCareChat/backend/main.py:269-271",
        "On first load the user sees a welcome message; in human-testing "
        "environments the welcome includes a live UAT status report so "
        "Skip and testers can see what is currently passing or failing.",
        "GAP",
        "Add: 'The conversational surface MUST present a welcome "
        "message on first load. In human-testing (non-prod) modes the "
        "welcome MUST additionally surface a live UAT status summary "
        "drawn from the QA report so testers see what is currently "
        "verified.'",
    ),
    (
        "BC-18",
        "Health surface with build, version, db status",
        "Code/ConversationalUX/FindCareChat/backend/main.py:273-277",
        "An external observer (operator, monitoring, or the React "
        "wrapper page) can poll a single endpoint and see whether the "
        "service is up, whether the database is reachable, what build "
        "is deployed, and what environment is serving.",
        "PARTIAL: BR-12",
        "BR-12 names end-to-end observability for tracing. Add: 'The "
        "system MUST publish a health surface that returns at minimum "
        "{status, db_state, env, build_number, app_version}; this is "
        "the surface the React wrapper uses to populate the environment "
        "banner with build and version.'",
    ),
    (
        "BC-19",
        "Per-turn debug audit row (chat_calls collection)",
        "Code/ConversationalUX/FindCareChat/backend/main.py:343 + "
        "infrastructure/debug_logger.py",
        "Every conversational turn produces a structured row in a "
        "debug-audit collection capturing originating IP, message "
        "preview, history length, tokens in/out, response, and any "
        "error class.",
        "PARTIAL: BR-08",
        "BR-08 names a per-turn audit record. Add: 'Each conversational "
        "turn MUST emit one structured audit row with at minimum: "
        "client IP, message preview, history length, tool-loop "
        "iterations, tokens in, tokens out, response (or null on "
        "error), error message (or null on success). The audit row is "
        "always emitted, including on error.'",
    ),
    (
        "BC-20",
        "Error-class triage: rate_limit, db_unavailable, internal",
        "Code/ConversationalUX/FindCareChat/backend/main.py:289-301",
        "When the conversation breaks, the user gets a typed error "
        "(rate_limit, db_unavailable, or internal) so the UI can react "
        "differently (e.g., retry vs degrade vs apology) and the "
        "operator can correlate root cause.",
        "GAP",
        "Add: 'On a conversational error the system MUST classify the "
        "error into rate_limit (LLM throttled), db_unavailable "
        "(persistence layer unreachable), or internal (anything else); "
        "the typed class MUST be returned to the client so the UI can "
        "react appropriately and a debug-audit row MUST still be "
        "written.'",
    ),
    (
        "BC-21",
        "CORS-allowed origins for chathealthy.ai sub-domains and localhost",
        "Code/ConversationalUX/FindCareChat/backend/main.py:218-223",
        "The browser-facing UI works from chathealthy.ai, "
        "www.chathealthy.ai, dev.chathealthy.ai, any *.chathealthy.ai "
        "sub-domain, and from localhost during development; other "
        "origins are blocked.",
        "PARTIAL: BR-05",
        "BR-05 covers the static page surface. Add: 'The conversational "
        "backend MUST accept browser requests from production "
        "(chathealthy.ai, www.chathealthy.ai), QA / dev "
        "(*.chathealthy.ai sub-domains), and developer localhost; all "
        "other origins MUST be rejected.'",
    ),
    (
        "BC-22",
        "Iframe-embedded chat with environment banner in parent only",
        "Website/index.html:593, 607-614, 638-645",
        "The Website main page hosts the conversational UI as an "
        "iframe; in non-production environments a colored banner above "
        "the iframe (NOT inside it) shows the current environment, "
        "build, and version so users on dev/QA know they are not in "
        "prod.",
        "PARTIAL: BR-05",
        "BR-05 covers static pages. Add: 'In non-production "
        "environments, the user-facing UI MUST display an "
        "environment-and-build banner that is rendered by the parent "
        "page (NOT inside the iframe) so the banner cannot be defaced "
        "or hidden by the embedded application.'",
    ),
    (
        "BC-23",
        "Iframe-embedded chat retry-on-load-failure (3s, 10 attempts)",
        "Website/index.html:683-695",
        "If the chat iframe fails to load (cold-start, transient "
        "network), the parent page automatically retries every 3 "
        "seconds up to 10 times before showing failure to the user.",
        "GAP",
        "Add: 'If the embedded conversational surface fails to load, "
        "the parent page MUST retry transparently for a configurable "
        "number of attempts (default 10) at a configurable interval "
        "(default 3s) before showing the user a load-failure state.'",
    ),
    (
        "BC-24",
        "Filter panel in left side panel with Check All / Uncheck All",
        "Website/index.html:737-768",
        "When the user issues a specialty query, a left-panel list of "
        "matching specialties appears; each has a checkbox; the user "
        "can Check All, Uncheck All, or toggle individual specialties; "
        "applying the filter triggers a server re-search with only "
        "the selected specialty codes.",
        "PARTIAL: BR-23",
        "BR-23 covers the filter UI in long form. Confirm and tighten: "
        "'The specialty filter list MUST appear in the left side panel "
        "with explicit Check All / Uncheck All controls, default-all-"
        "checked, scroll-to-11 visible, applied only on explicit "
        "Filter button click, with the resulting search retaining any "
        "providers the user already moved to the lower selection box.'",
    ),
    (
        "BC-25",
        "Multi-page legal panel: Privacy, Terms, Architecture, Roadmap",
        "Website/index.html:589-593, 786-789",
        "From the footer the user can open a slide-in panel to read "
        "Privacy Policy, Terms of Use, Roadmap, or Architecture "
        "without leaving the chat page.",
        "PARTIAL: BR-05",
        "BR-05 says ten static pages remain. Tighten: 'The 10 static "
        "pages (Privacy, Terms, Roadmap, Architecture, Products, "
        "Embedding-Design, Chat-App-Design, Provider-Data-Load, "
        "Load-Perf-Report, Ops-Manager-Design) MUST be reachable from "
        "the main page without leaving the chat experience (slide-in "
        "or new tab); the index page footer MUST surface at least "
        "Privacy, Terms, Roadmap, and Architecture.'",
    ),
    (
        "BC-26",
        "Hamburger / mobile nav with 'coming soon' tooltip pattern",
        "Website/index.html:790-810",
        "On small screens the user gets a hamburger menu; nav items "
        "marked 'coming soon' show a tooltip rather than a dead link.",
        "GAP",
        "Add: 'Navigation items that point to features not yet "
        "released MUST display a 'coming soon' tooltip on click "
        "rather than navigating to a 404 or empty page.'",
    ),
    (
        "BC-27",
        "Tool-call boundary: only allowlisted tools can be invoked",
        "Code/ConversationalUX/FindCareChat/backend/application/"
        "tool_router.py + main.py:199-209",
        "When the assistant tries to invoke a tool, the system "
        "verifies the tool is registered in the allowlist; "
        "unregistered tool names are rejected at the dispatch boundary "
        "before any business code runs.",
        "PARTIAL: BR-10",
        "BR-10 names the engineering-rules enforcement worker registry "
        "for commit-time gating. The runtime equivalent is unstated. "
        "Add: 'Every tool call by the assistant MUST be dispatched "
        "through an allowlist registry; tool names not registered for "
        "the assistant in this session MUST be rejected at the "
        "dispatch boundary, never reaching application code.'",
    ),
    (
        "BC-28",
        "Pydantic-validated tool inputs at the dispatch boundary",
        "Code/ConversationalUX/FindCareChat/backend/application/"
        "tool_router.py + tool_models/*.py",
        "When a tool has a typed input schema, the assistant's "
        "arguments are validated against it before the business code "
        "runs; bad arguments produce a deterministic error rather than "
        "an unsafe call.",
        "GAP",
        "Add: 'Every tool with a typed input contract MUST have its "
        "inputs validated by that contract at the dispatch boundary; "
        "validation failures MUST return a structured error to the "
        "assistant rather than entering business code.'",
    ),
    (
        "BC-29",
        "Brain-driven configuration for prompts, welcome, keywords, tools",
        "Code/Shared/prompt_system_maker.py + "
        "Code/ConversationalUX/FindCareChat/backend/main.py:152-159",
        "Every piece of LLM-facing text the user sees (system prompt, "
        "welcome message, emergency keyword list, tool descriptions) "
        "is loaded from a brain-controlled configuration so a "
        "non-engineer can change wording without a code change.",
        "GAP",
        "Add: 'All user-facing LLM configuration (system prompt, "
        "welcome message, emergency keyword list, tool descriptions / "
        "schemas) MUST be loaded at runtime from the brain "
        "configuration repository (brain/machine_artifacts/content/), "
        "never hard-coded in application code; non-engineers can "
        "review and change wording without a code change.'",
    ),
    (
        "BC-30",
        "Linkify Skip Snow LinkedIn mention if the LLM forgot",
        "Code/ConversationalUX/FindCareChat/backend/main.py:349",
        "If the assistant mentions 'Skip Snow on LinkedIn' as plain "
        "text without a hyperlink, the system rewrites it as a live "
        "Markdown link to skipsnow's LinkedIn profile.",
        "GAP",
        "Add: 'When the assistant references the founder by name in a "
        "context that should be a contact link, the system MUST "
        "rewrite the bare text as a live link to the canonical "
        "external profile (LinkedIn) so the user can reach the "
        "founder without copy-paste guesswork.'",
    ),
    (
        "BC-31",
        "Static page surface served from the same backend (HF deploy)",
        "Code/ConversationalUX/FindCareChat/backend/main.py:381-390",
        "When deployed to HuggingFace Space, the same backend that "
        "serves the conversation also serves the React/static index "
        "and assets, so a single URL covers both the page and the "
        "API.",
        "FULL: TR-UX-002 (BR-05)",
        "",
    ),
    (
        "BC-32",
        "Provider-search by name (last/first/organization OR-regex)",
        "Code/ConversationalUX/FindCareChat/backend/domain/find_care/"
        "provider_search_service.py:325",
        "The user can ask for a provider by name; the system "
        "case-insensitively matches against last name, first name, OR "
        "organization name and returns matches.",
        "GAP",
        "Add: 'Given a provider name (with optional state), the "
        "system MUST OR-search across last name, first name, and "
        "organization name fields with case-insensitive matching, "
        "returning the same paginated provider list shape as a "
        "specialty search.'",
    ),
    (
        "BC-33",
        "Provider-search by NPI (exact lookup)",
        "Code/ConversationalUX/FindCareChat/backend/domain/find_care/"
        "provider_search_service.py:317",
        "The user (or assistant) can supply an NPI and get back the "
        "single matching provider record, bypassing semantic search.",
        "GAP",
        "Add: 'Given an NPI, the system MUST return the single "
        "matching provider record (or an empty result with a "
        "user-facing message) without invoking the semantic search "
        "pipeline.'",
    ),
    (
        "BC-34",
        "FIPS county map (state + county name -> provider filter)",
        "Code/ConversationalUX/FindCareChat/backend/domain/find_care/"
        "provider_search_service.py:22, 56",
        "When the user narrows by county, the system resolves the "
        "county name to the FIPS code(s) the provider records use, so "
        "'Sussex County, DE' returns the right providers regardless "
        "of how the database stores the county.",
        "GAP",
        "Add: 'County filtering MUST resolve user-typed county names "
        "(case- and punctuation-insensitive) against the canonical "
        "FIPS map so the user gets the providers they meant; the FIPS "
        "map is loaded at service startup and seeded from the brain.'",
    ),
    (
        "BC-35",
        "AI-assisted query expansion for specialty matching (Haiku)",
        "Code/ConversationalUX/FindCareChat/backend/domain/find_care/"
        "specialty_service.py:69 + infrastructure/embeddings/"
        "embedding_client.py",
        "When the user uses colloquial language ('foot doctor'), an "
        "LLM expands the query into stems and synonyms ('podiatry', "
        "'podiatric') so both regex and semantic matching cover the "
        "intent.",
        "PARTIAL: BR-23",
        "BR-23 says 'No inappropriate providers are supplied'. Add: "
        "'Specialty matching MUST AI-expand colloquial care terms "
        "into related canonical stems before regex and semantic "
        "search, so plain-language queries surface the right NUCC "
        "groupings.'",
    ),
    (
        "BC-36",
        "Static FIPS-to-county seed (HACK ASN-4AFBDA)",
        "Code/ConversationalUX/FindCareChat/backend/domain/find_care/"
        "provider_search_service.py:22",
        "If the database is unreachable at startup, a small built-in "
        "FIPS-to-county map keeps Delaware and Mississippi counties "
        "queryable so the conversation does not fail open.",
        "GAP",
        "Add: 'The system MUST degrade gracefully on database "
        "unreachability for county lookup: a built-in canonical "
        "FIPS-to-county seed MUST cover the supported states so a "
        "transient DB outage does not break county-narrowed queries; "
        "this is a graceful-degradation requirement, not a forever "
        "alternative to the DB-loaded map.'",
    ),
    (
        "BC-37",
        "EvaluateCare facade-to-facade boundary",
        "Code/ConversationalUX/FindCareChat/backend/application/facades/"
        "evaluate_care_facade.py",
        "When EvaluateCare needs FindCare data (e.g. provider "
        "location), it calls the FindCare facade, never reaching "
        "into FindCare's internal services. The capability boundary "
        "is observable from the call graph.",
        "PARTIAL: BR-11",
        "BR-11 says each capability epic has its own org-chart and "
        "deploy line. Tighten: 'When one capability needs data from "
        "another, the cross-capability call MUST go through the "
        "called capability's public facade, never directly into its "
        "internal services; this is the architectural discipline that "
        "makes BR-11's cost-center boundary observable at call time, "
        "not just at deploy time.'",
    ),
    (
        "BC-38",
        "Founder / company context loaded from /me directory",
        "Code/ConversationalUX/FindCareChat/backend/main.py:161-167 + "
        "domain/shared/content/about_service.py",
        "When the user asks about Skip or ChatHealthy, the assistant "
        "draws from a versioned content directory (LinkedIn summary, "
        "business plan, Anthropic principles) so the answers stay "
        "current with the company's actual messaging.",
        "PARTIAL: TR-FC-006 (BR-10)",
        "TR-FC-006 names the tools but the source is unstated. Add: "
        "'Founder / company context MUST be loaded from a versioned "
        "content directory at startup; large fields (LinkedIn "
        "summary, business plan) MUST be trimmable so an over-long "
        "field does not blow the prompt budget.'",
    ),
    (
        "BC-39",
        "Push-notification on significant lead activity (SparkPost email)",
        "Code/ConversationalUX/FindCareChat/backend/main.py:103-113",
        "When a user leaves contact info, ChatHealthy receives an "
        "email so a human can follow up quickly without polling the "
        "database.",
        "GAP",
        "Add: 'When a user signals interest with a lead-capture event, "
        "ChatHealthy ops MUST receive an immediate notification so a "
        "human can follow up; the channel is configurable; failure "
        "to notify MUST NOT block the user-facing flow.'",
    ),
    (
        "BC-40",
        "Significant-activity audit write to MongoDB (commitSignificantActivity)",
        "Code/ConversationalUX/FindCareChat/backend/main.py:115-125 + "
        "main.py:208 (registered as LLM tool)",
        "Significant user-or-system events (lead recorded, unknown "
        "question recorded, etc.) are written to a Mongo audit "
        "collection so the operator can reconstruct the user's "
        "journey after the fact.",
        "GAP / WITHDRAW-CANDIDATE",
        "Add OR withdraw: 'Significant-activity events (lead "
        "recorded, unknown question recorded, emergency lock) MUST be "
        "persisted to an audit collection. NOTE: at pre-4/3 this "
        "capability was also registered as an LLM-callable tool; the "
        "V2 architect should confirm whether the LLM should EVER be "
        "permitted to write arbitrary audit rows. Recommended posture: "
        "remove from the LLM tool registry and keep it as an internal "
        "service only.'",
    ),
    (
        "BC-41",
        "PostMessage protocol between iframe and parent (gui:render, gui:filter)",
        "Website/index.html:715-770",
        "The chat iframe and the parent page communicate via "
        "structured postMessage events: gui:render to draw a control, "
        "gui:filter to show the filter panel, gui:filter-clear to "
        "clear it, gui:event for user clicks. The contract is "
        "narrow and explicit.",
        "GAP",
        "Add: 'The conversational iframe and its parent page MUST "
        "communicate exclusively through a documented set of "
        "postMessage event types (render, clear, filter render, "
        "filter clear, event); no direct DOM access across the iframe "
        "boundary; events with unknown types MUST be ignored.'",
    ),
]

# ---------------------------------------------------------------------------
# What V1 got wrong (explicit list)
# ---------------------------------------------------------------------------
V1_INVALID_FINDINGS = [
    (
        "V1 cited commit f836d4d0 (build 998, 2026-04-19) as the "
        "'pre-event' baseline.",
        "INVALID. Anything from 4/3/2026 onward is post-regression by "
        "the architect's V2 ratification. The correct pre-event "
        "baseline is be14579a (build 410, 2026-04-02 23:48). V1's "
        "entire commit reference is post-event.",
    ),
    (
        "V1 CAP-N01 .. CAP-N10 (chat_graph nodes, classify_graph nodes, "
        "qa_report_graph) at "
        "Code/ConversationalUX/FindCareChat/backend/application/chat_graph.py "
        "and application/graphs/*.py.",
        "INVALID as pre-event capabilities. None of those files exist "
        "at be14579a; they were created after 4/3 as part of the "
        "post-regression LangGraph wrap. They cannot be cited as "
        "pre-event capabilities.",
    ),
    (
        "V1 CAP-R03 (POST /classify), CAP-R04 / CAP-R05 (/qa-report), "
        "CAP-R08 (GET /session), CAP-R09..R14 (/evaluate/*, /shared/*, "
        "/transfer/*) at main.py line numbers >390.",
        "INVALID as pre-event routes. Pre-4/3 main.py is 394 lines "
        "and exposes only /search, /welcome, /health, /chat, /. The "
        "additional routes (/classify, /qa-report, /session, "
        "/evaluate/*, /shared/*, /transfer/*) were added AFTER 4/3 "
        "during the regression-recovery work.",
    ),
    (
        "V1 CAP-S10 (specialty_classifier.py), CAP-S11 (specialty_ranker.py), "
        "CAP-S12 (homeopathic_resolver.py).",
        "INVALID as pre-event services. None of those modules exist "
        "at be14579a. They were introduced after 4/3 as part of the "
        "post-regression specialty pipeline overhaul.",
    ),
    (
        "V1 reported $vectorSearch (CAP-S06) as a BR-level GAP and "
        "wrote a recommended BR specifying the index name "
        "(provider_vector_index), embedding model, and pipeline shape.",
        "INVALID at BR LEVEL. $vectorSearch index names, embedding "
        "model identity, and aggregation-pipeline shape are TR-level "
        "concerns. The right BR is the user-visible behavior "
        "('semantic specialty matching surfaces the right NUCC "
        "groupings for plain-language queries'); the index, model, "
        "and pipeline belong in TRs. V2 captures the user-visible "
        "behavior in BC-13 / BC-35.",
    ),
    (
        "V1 CAP-B02 (startup primitive verification, exit codes "
        "78/77/70) and CAP-B03 (cert bootstrap from PEM env vars to "
        "/tmp/ch_certs) reported as BR-level GAPs.",
        "INVALID at BR LEVEL. Process exit codes, PEM file paths, "
        "and cert bootstrap directories are TR / DevOps detail, not "
        "business requirements. The user-visible behavior 'in non-"
        "prod the user sees that mTLS/security primitives are "
        "healthy via the banner' is the BR (covered under BC-22 / "
        "BC-18); the primitive verification, exit codes, and PEM "
        "paths are TRs.",
    ),
    (
        "V1 CAP-S20 recommended BR text named the exact "
        "ClinicalTrials.gov v2 URL, the exact Google Routes API "
        "endpoint, and the overallStatus parameter values.",
        "INVALID at BR LEVEL. Specific external API hostnames, URL "
        "paths, and parameter keys are TR-level; the BR is the "
        "user-visible 'recruiting trials with drive-time' behavior. "
        "V2 captures the BR in BC-05.",
    ),
    (
        "V1 CAP-S23 recommended BR specifying OpenAI "
        "text-embedding-3-large as the embedding model, the index "
        "names, and that the same model serves both provider and "
        "specialty vector search.",
        "INVALID at BR LEVEL. The embedding model identity is "
        "TR-level. The user-visible BR is captured in BC-13 / BC-35.",
    ),
    (
        "V1 CAP-S25 (URLGuardian) recommended BR specifying 3600s "
        "cache TTL and 5s probe timeout in the BR text.",
        "INVALID at BR LEVEL. Cache TTL and probe timeout values are "
        "TR-level tuning. The BR is 'unreachable URLs are removed "
        "before display so the user never sees a broken link'; "
        "captured in V2 as BC-11.",
    ),
    (
        "V1 CAP-N03 (search_term_extract) recommended BR named "
        "gpt-4.1-mini specifically.",
        "INVALID at BR LEVEL. Choosing a specific LLM model is a "
        "TR / cost-tuning concern. The BR is 'the system extracts a "
        "colloquial search term from the user's message in parallel "
        "with safety checks' — but this is itself a node in "
        "post-4/3 chat_graph.py which does not exist at pre-4/3, so "
        "the underlying capability does not even exist pre-event.",
    ),
    (
        "V1 included LangGraph/poc/user_journey.py in its trace cells "
        "and methodology (the design-V2 BR table also references it).",
        "INVALID for V2 SCOPE. The POC's own header (LangGraph/poc/"
        "user_journey.py) dates it 2026-04-20, post-4/3. Any pre-4/3 "
        "audit MUST NOT cite it as evidence of pre-event capability.",
    ),
]


# ---------------------------------------------------------------------------
# Document generation
# ---------------------------------------------------------------------------
def _set_cell(cell, text: str, bold: bool = False, size: int = 9):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold


def _add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x0B, 0x7A, 0x75)
    return h


def _add_para(doc, text, size=11, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    return p


def build_doc(output_path: str) -> None:
    doc = Document()

    # Title
    title = doc.add_heading("Pre-SQL Capability Audit V2", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x0B, 0x7A, 0x75)

    _add_para(
        doc,
        "ChatHealthy.ai - Capability inventory of the codebase IMMEDIATELY "
        "BEFORE 2026-04-03 (the architect's ratified pre-regression cutoff), "
        "compared against design-V2.docx Business Requirements ONLY. "
        "BR-LEVEL findings only; no TR-level concerns are reported as gaps.",
        size=10,
    )
    _add_para(
        doc,
        "Date: 2026-04-28   |   Author: Claude (capability-audit agent, V2)   |   "
        "Read-only deliverable - no commits, no code changes, no checkouts.",
        size=10,
    )

    # 1. Executive summary
    _add_heading(doc, "1. Executive summary", level=1)
    full = sum(1 for c in CAPABILITIES if c[4].startswith("FULL"))
    partial = sum(1 for c in CAPABILITIES if c[4].startswith("PARTIAL"))
    gap = sum(
        1 for c in CAPABILITIES
        if c[4] == "GAP" or c[4].startswith("GAP")
    )
    total = len(CAPABILITIES)
    _add_para(doc, f"Pre-4/3 reference commit: {PRE_43_COMMIT}")
    _add_para(doc, f"Pre-4/3 commit date: {PRE_43_DATE}")
    _add_para(doc, f"Pre-4/3 build number: {PRE_43_BUILD}")
    _add_para(doc, f"Pre-4/3 commit subject: {PRE_43_SUBJECT}")
    _add_para(doc, "")
    _add_para(doc, f"Business capabilities enumerated (BR-level): {total}", bold=True)
    _add_para(doc, f"  FULL coverage by V2 BRs:    {full}")
    _add_para(doc, f"  PARTIAL coverage by V2 BRs: {partial}")
    _add_para(doc, f"  GAP (no V2 BR):             {gap}")
    _add_para(doc, "")
    _add_para(
        doc,
        "Headline finding: V2's BR set covers the high-level conversational "
        "find-care experience but is thin on (a) per-turn deterministic "
        "behaviors a user can directly observe (safety gating order, "
        "follow-up cadence, error-class handling), (b) reliability "
        "behaviors the user benefits from (URL defanging, iframe retry, "
        "graceful degradation on DB outage), (c) consent-and-record "
        "details that affect the user's stored data (notification on "
        "interest, audit-row obligations), and (d) cross-product "
        "boundary behaviors (facade-to-facade rule, postMessage "
        "protocol). Pre-4/3 code already exhibits these behaviors; they "
        "merit explicit BRs so V2 design and implementation can be "
        "tested against them.",
        size=10,
    )

    # 2. Methodology + scope
    _add_heading(doc, "2. Methodology and scope", level=1)
    _add_para(
        doc,
        "Pre-4/3 source was inspected ONLY via `git show " + PRE_43_COMMIT[:8] +
        ":<path>` (read-only). No `git checkout`, no commits, no working-"
        "tree mutations occurred during this audit. The deliverable docx "
        "and this generator script are the only writes performed.",
    )
    _add_para(
        doc,
        "V2 BRs were extracted from "
        "architecture/ChatHealthyApplicationDesign/"
        "ArchitectureDesignAndAuditDocs/design-V2.docx via python-docx "
        "(BR table = Table 1 of the document). BR numbering follows the "
        "explicit BR-NN headings and tolerates Skip's mid-table "
        "renumberings (BR-21 missing in source; BR-32 used twice).",
    )
    _add_para(
        doc,
        "Scope inclusions: pre-4/3 user-facing behaviors, system "
        "protections the user benefits from (e.g. URL defanging), "
        "conversational behaviors (safety, classification, follow-up "
        "cadence), and cross-product handoffs (FindCare <-> EvaluateCare).",
    )
    _add_para(
        doc,
        "Scope EXCLUSIONS (explicit, by V2 directive): "
        "(a) anything from LangGraph/poc/ - the POC dates itself "
        "2026-04-20 in its own header, which is post-4/3 and is "
        "post-event by definition. The POC is NEVER cited as a "
        "pre-event capability in this V2 audit. "
        "(b) TR-level concerns (specific database choice, embedding "
        "model identity, aggregation-pipeline operator names like "
        "$vectorSearch, framework version numbers, deployment topology, "
        "process exit codes, PEM file paths, port numbers, cache TTLs, "
        "specific API hostnames). These are NOT eligible to be reported "
        "as BR-level gaps.",
    )
    _add_para(
        doc,
        "Coverage classification: FULL means a V2 BR clearly describes "
        "the capability and constrains its observable behavior. PARTIAL "
        "means a V2 BR names the capability but leaves the behavior "
        "under-specified. GAP means no V2 BR describes the capability.",
    )

    # 3. Business capability inventory
    _add_heading(doc, "3. Business capability inventory", level=1)
    _add_para(
        doc,
        "Each row: capability ID, capability name, source file:line at "
        "the pre-4/3 commit, one-sentence user-facing description, V2 "
        "BR coverage, recommended BR text if PARTIAL or GAP. All source "
        "paths read at commit " + PRE_43_COMMIT[:8] + " via `git show`.",
        size=10,
    )

    table = doc.add_table(rows=1, cols=6)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    _set_cell(hdr[0], "ID", bold=True, size=9)
    _set_cell(hdr[1], "Business capability", bold=True, size=9)
    _set_cell(hdr[2], "Source file:line at " + PRE_43_COMMIT[:8],
              bold=True, size=9)
    _set_cell(hdr[3], "User-facing description", bold=True, size=9)
    _set_cell(hdr[4], "V2 BR coverage", bold=True, size=9)
    _set_cell(hdr[5], "Recommended BR text (PARTIAL/GAP only)",
              bold=True, size=9)

    for cap_id, name, src, desc, cov, rec in CAPABILITIES:
        row = table.add_row().cells
        _set_cell(row[0], cap_id, size=8)
        _set_cell(row[1], name, size=8)
        _set_cell(row[2], src, size=8)
        _set_cell(row[3], desc, size=8)
        _set_cell(row[4], cov, size=8)
        _set_cell(row[5], rec, size=8)

    # Force a sane column-A width so the ID is readable (Skip's V1 note)
    try:
        widths_in = [0.55, 1.4, 1.5, 1.95, 0.85, 2.25]
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                cell.width = Inches(widths_in[idx])
    except Exception:
        pass

    # 4. Gap summary at a glance
    _add_heading(doc, "4. Gap summary - PARTIAL and GAP at a glance", level=1)
    _add_para(
        doc,
        "GAP items (no V2 BR description). These are real, observable "
        "user-facing behaviors in the pre-4/3 code that the V2 BR set "
        "does not name.",
        bold=True, size=11,
    )
    gap_rows = [
        c for c in CAPABILITIES
        if c[4] == "GAP" or c[4].startswith("GAP")
    ]
    gap_table = doc.add_table(rows=1, cols=3)
    gap_table.style = "Light Grid Accent 1"
    gh = gap_table.rows[0].cells
    _set_cell(gh[0], "ID", bold=True, size=9)
    _set_cell(gh[1], "Business capability", bold=True, size=9)
    _set_cell(gh[2], "Source file:line", bold=True, size=9)
    for cap_id, name, src, _, _, _ in gap_rows:
        row = gap_table.add_row().cells
        _set_cell(row[0], cap_id, size=8)
        _set_cell(row[1], name, size=8)
        _set_cell(row[2], src, size=8)

    _add_para(doc, "")
    _add_para(
        doc,
        "PARTIAL items (named in a BR but observable behavior under-specified).",
        bold=True, size=11,
    )
    partial_rows = [c for c in CAPABILITIES if c[4].startswith("PARTIAL")]
    p_table = doc.add_table(rows=1, cols=4)
    p_table.style = "Light Grid Accent 1"
    ph = p_table.rows[0].cells
    _set_cell(ph[0], "ID", bold=True, size=9)
    _set_cell(ph[1], "Business capability", bold=True, size=9)
    _set_cell(ph[2], "Source file:line", bold=True, size=9)
    _set_cell(ph[3], "Existing V2 BR(s)", bold=True, size=9)
    for cap_id, name, src, _, cov, _ in partial_rows:
        row = p_table.add_row().cells
        _set_cell(row[0], cap_id, size=8)
        _set_cell(row[1], name, size=8)
        _set_cell(row[2], src, size=8)
        _set_cell(row[3], cov, size=8)

    # 5. Recommendations
    _add_heading(doc, "5. Recommendations - proposed BR additions / refinements",
                 level=1)
    _add_para(
        doc,
        "These are written at the specificity of Skip's V2-added BRs "
        "(BR-16..BR-35): one user-visible behavior per BR, in business "
        "language, with the testable acceptance criteria implicit in the "
        "behavior. No TR-level detail (no model names, no index names, "
        "no exit codes, no specific URLs).",
        size=10,
    )
    rec_rows = [c for c in CAPABILITIES if c[5]]
    for cap_id, name, _src, _desc, cov, rec in rec_rows:
        _add_para(doc, f"{cap_id} - {name}  [{cov}]", bold=True, size=10)
        _add_para(doc, rec, size=10)
        _add_para(doc, "")

    # 6. What V1 got wrong
    _add_heading(doc, "6. What V1 got wrong - explicit invalidations", level=1)
    _add_para(
        doc,
        "V1 referenced commit f836d4d0 (build 998, dated 2026-04-19) "
        "as the 'pre-event' baseline and treated post-4/3 artifacts "
        "as evidence of pre-event capability. The architect has now "
        "ratified that anything from 4/3/2026 onward is "
        "post-regression. Below are V1 findings that V2 invalidates "
        "and the reason.",
        size=10,
    )

    inv_table = doc.add_table(rows=1, cols=2)
    inv_table.style = "Light Grid Accent 1"
    ih = inv_table.rows[0].cells
    _set_cell(ih[0], "V1 finding", bold=True, size=9)
    _set_cell(ih[1], "Why V2 invalidates it", bold=True, size=9)
    for v1, why in V1_INVALID_FINDINGS:
        row = inv_table.add_row().cells
        _set_cell(row[0], v1, size=8)
        _set_cell(row[1], why, size=8)

    # 7. Caveats
    _add_heading(doc, "7. Caveats and ambiguities", level=1)
    _add_para(
        doc,
        "BR table caveats: design-V2 Table 1 contains a renumbering "
        "anomaly (BR-21 is absent; BR-32 appears twice with different "
        "statements). This audit treats every populated BR row as a "
        "constraint regardless of its ID and notes that a V2 erratum "
        "is needed to settle the BR-32 collision.",
        size=10,
    )
    _add_para(
        doc,
        "Capability BC-40 (significant-activity audit) is registered "
        "at pre-4/3 as both an internal commit utility AND an "
        "LLM-callable tool. The right V2 posture is most likely 'keep "
        "as internal service, remove from LLM tool registry' so the "
        "model cannot write arbitrary Mongo collections; the architect "
        "should ratify.",
        size=10,
    )
    _add_para(
        doc,
        "Cross-product handoff capabilities (BC-05, BC-06, BC-37) at "
        "pre-4/3 lived inside the FindCare backend (EvaluateCareFacade "
        "is a class inside FindCare's process). The capability boundary "
        "was logical, not a separate process. V2 architecture moves "
        "EvaluateCare to its own HF Space; the pre-4/3 user-facing "
        "BEHAVIOR (asking about clinical trials and provider detail "
        "in-conversation) is the BR; the inter-Space topology is TR.",
        size=10,
    )

    # 8. Final reply text duplicate (so the docx contains the same headline numbers)
    _add_heading(doc, "8. Headline numbers", level=1)
    _add_para(
        doc,
        f"Pre-4/3 commit: {PRE_43_COMMIT}, build {PRE_43_BUILD}, "
        f"dated {PRE_43_DATE}.",
    )
    _add_para(
        doc,
        f"Total business capabilities (BR-level): {total}. "
        f"FULL: {full}. PARTIAL: {partial}. GAP: {gap}.",
    )

    # Save
    doc.save(output_path)


if __name__ == "__main__":
    here = os.path.dirname(os.path.abspath(__file__))
    out = os.path.join(here, "pre-sql-capability-audit-V2.docx")
    build_doc(out)
    print(f"Wrote: {out}")
