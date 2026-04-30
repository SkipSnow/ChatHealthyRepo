# Copyright (c) 2026 ChatHealthy.ai LLC. All rights reserved.
#
# generate_pre_sql_capability_audit_V1.py
# Builds pre-sql-capability-audit-V1.docx
# Read-only audit. No commits, no code changes.

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "pre-sql-capability-audit-V1.docx",
)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_para(doc, text, bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if size:
        r.font.size = Pt(size)
    return p


def shade_cell(cell, fill_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def set_col_widths(table, widths_inches):
    """Set per-column widths so column A is not skinny (per Skip's note)."""
    for i, w in enumerate(widths_inches):
        for cell in table.columns[i].cells:
            cell.width = Inches(w)


def add_table(doc, headers, rows, col_widths=None, header_fill="1F4E79"):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(10)
        shade_cell(hdr[i], header_fill)
    for row in rows:
        rc = t.add_row().cells
        for i, val in enumerate(row):
            rc[i].text = str(val) if val is not None else ""
            for p in rc[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    if col_widths:
        set_col_widths(t, col_widths)
    return t


# ---------------------------------------------------------------------------
# Capability inventory
# ---------------------------------------------------------------------------
# Each entry: (id, capability, source, type, coverage, recommended_br)
# coverage values: "FULL: BR-N" | "PARTIAL: BR-N" | "GAP"

CAPABILITIES = [
    # --- Routes (FastAPI handlers exposing capabilities) ---
    ("CAP-R01", "POST /chat — conversational entry point invoking chat_graph",
     "Code/ConversationalUX/FindCareChat/backend/main.py:693",
     "Route",
     "PARTIAL: BR-01, BR-03",
     "BR-01/BR-03 name a single conversational entry but do not specify the per-turn graph contract (gate -> safety+search-term fan-out -> tool loop -> post-processing). Add: 'The /chat turn handler MUST execute exactly this node order: gate, parallel(safety_check, search_term_extract), merge, chat_loop, post-processing fan-out (url_guard + optional pagination/trials), assemble_response. Any additional node MUST be registered in a per-turn-graph manifest.'"),

    ("CAP-R02", "POST /search — direct provider search bypassing LLM (pagination/replay)",
     "Code/ConversationalUX/FindCareChat/backend/main.py:376",
     "Route",
     "GAP",
     "Add BR: 'The system MUST expose a non-conversational provider search endpoint (POST /search) that accepts SearchRequest (specialty_query, state, city, county, name, npi, specialty_codes, after_npi, limit) and returns ranked providers without invoking the LLM. Used by the React UI for pagination replay and filter changes (BR-23 mandates client-side filter refresh).' V2 BR-06 names provider detail GET but not /search POST."),

    ("CAP-R03", "POST /classify — utterance -> structured search params (safety gate -> vector specialty -> location parse -> homeo-generalists cache)",
     "Code/ConversationalUX/FindCareChat/backend/main.py:388 + application/graphs/classify_graph.py",
     "Route",
     "PARTIAL: BR-22",
     "BR-22 names 'utterance classifier' but does not constrain inputs/outputs. Add: 'POST /classify MUST take {message, ip}, run safety_gate first (IP lock + emergency), then return {specialties[code,name,can_prescribe,homeopathic,rank], homeopathic_generalists[], state, city, county, model}. State extraction limited to the supported states (DE/MS/VA/IN). The endpoint MUST be available without authentication for the initial filter-panel paint.'"),

    ("CAP-R04", "GET /qa-report — render editable QA-report HTML from MongoDB",
     "Code/ConversationalUX/FindCareChat/backend/main.py:418 + application/graphs/qa_report_graph.py",
     "Route",
     "GAP",
     "Add BR: 'A QA-report surface MUST exist at GET /qa-report that renders test-case status from the {ENV}_System.bugs collection, editable via dropdown (PASS/FAIL/DEFERRED/...). Disabled in prod. Used by Skip to record test outcomes against the canonical feature inventory.' Or explicitly mark as out-of-scope."),

    ("CAP-R05", "POST /qa-report — persist QA-report edits to MongoDB",
     "Code/ConversationalUX/FindCareChat/backend/main.py:427 + application/graphs/qa_report_submit_graph.py",
     "Route",
     "GAP",
     "Add BR companion to CAP-R04: 'POST /qa-report accepts the form post from the rendered report, updates each test_case status, recomputes summary counts, persists to {ENV}_System.bugs, and 303-redirects back to GET /qa-report. Disabled in prod.'"),

    ("CAP-R06", "GET /welcome — static welcome message (UAT-aware)",
     "Code/ConversationalUX/FindCareChat/backend/main.py:439",
     "Route",
     "GAP",
     "Add BR: 'The system MUST expose GET /welcome returning the welcome message text. In HUMAN_TESTING mode the welcome MUST include the live UAT report summary built from MongoDB.' Today there is no V2 BR for the initial paint message."),

    ("CAP-R07", "GET /health — degraded/ok status + index check + build/version/framework from Mongo",
     "Code/ConversationalUX/FindCareChat/backend/main.py:467",
     "Route",
     "PARTIAL: BR-09",
     "BR-09 covers freshness; does NOT specify health surface. Add: 'GET /health MUST return {status, db, env, build, version, framework, missing_indexes?}. build/version/framework MUST be sourced from {ENV}_System.version._id=current. Status MUST degrade if any required vector search index is missing (providers/provider_vector_index, SpecialtyMetaData/specialty_vector_index).'"),

    ("CAP-R08", "GET /session — generate signed session token (NONCE + cert-signed)",
     "Code/ConversationalUX/FindCareChat/backend/main.py:498",
     "Route",
     "PARTIAL: BR-20, BR-26",
     "BR-20 specifies the token shape (NONCE, role, encrypted user id) and BR-26 the GUI; the issuance endpoint contract is not bound. Add: 'GET /session MUST return a session_token signed with FindCare's x509 cert; token MUST be 68 chars, prefix CH, contain creation/last-use timestamps. Failure MUST return HTTP 500 with no fallback (security primitive — see BUG-SEC-007).'"),

    ("CAP-R09", "GET /evaluate/splash — proxy to EvaluateCare splash with mTLS + control transfer",
     "Code/ConversationalUX/FindCareChat/backend/main.py:530",
     "Route",
     "PARTIAL: BR-25",
     "BR-25 says 'rest of the page is now controlled by EvaluateCare server' but does not specify the proxy contract. Add: 'GET /evaluate/splash MUST proxy via mTLS (FindCare client cert + CA verify) to {EVALCARE_URL}/splash and return the rendered splash HTML, logging CONTROL TRANSFER. Failure MUST return a fallback service-unavailable HTML, never a stack trace.'"),

    ("CAP-R10", "POST /transfer/to-findcare — return page ownership FROM EvaluateCare",
     "Code/ConversationalUX/FindCareChat/backend/main.py:554",
     "Route",
     "GAP",
     "Add BR: 'When the user interacts with the FindCare filter while EvaluateCare owns the page, EvaluateCare MUST POST /transfer/to-findcare; FindCare MUST notify EvaluateCare of the transfer (mTLS round-trip), reclaim ownership, and return {owner: findcare, splash: WELCOME_MESSAGE}. Both services MUST log the transfer.'"),

    ("CAP-R11", "POST /evaluate/providers — proxy evaluation request to EvaluateCare with mTLS",
     "Code/ConversationalUX/FindCareChat/backend/main.py:579",
     "Route",
     "PARTIAL: BR-25",
     "BR-25 names the Evaluate Care button but does not bind the proxy. Add: 'POST /evaluate/providers MUST proxy {providers, session_token, question_summary} to {EVALCARE_URL}/evaluate/providers over mTLS (15s timeout) and return EvaluateCare's JSON unchanged. mTLS uses FindCare client cert + CA verify in non-HF environments; HF environment relies on platform TLS.'"),

    ("CAP-R12", "GET /shared/splash — proxy to SharedServices splash with mTLS + control transfer",
     "Code/ConversationalUX/FindCareChat/backend/main.py:619",
     "Route",
     "PARTIAL: BR-27",
     "BR-27 says 'Shared Services has no UI to present here.' Specify the contract: 'GET /shared/splash MUST proxy via mTLS to {SHARED_SERVICES_URL}/splash and return its splash HTML; on failure fall back to a service-unavailable card and never to FindCare content.'"),

    ("CAP-R13", "POST /shared/verify-token — proxy token verify to SharedServices over mTLS",
     "Code/ConversationalUX/FindCareChat/backend/main.py:643",
     "Route",
     "PARTIAL: BR-26",
     "BR-26 names the security-token GUI; the proxy contract is unspecified. Add: 'POST /shared/verify-token MUST proxy {session_token} to {SHARED_SERVICES_URL}/verify-token over mTLS and return its verification verdict (status, NONCE update, role, expiry, encrypted user id). FindCare MUST NOT verify the token in-process; only SharedServices does.'"),

    ("CAP-R14", "POST /evaluate/verify-token — proxy token verify to EvaluateCare over mTLS",
     "Code/ConversationalUX/FindCareChat/backend/main.py:668",
     "Route",
     "PARTIAL: BR-26",
     "Mirror of CAP-R13 for EvaluateCare. Add: 'POST /evaluate/verify-token MUST proxy {session_token} to {EVALCARE_URL}/verify-token over mTLS so the DEVOPS-BANNER-B006 banner can populate the security panel for the EvaluateCare button.'"),

    ("CAP-R15", "GET / — serve React static index (no-cache headers)",
     "Code/ConversationalUX/FindCareChat/backend/main.py:865",
     "Route",
     "FULL: TR-UX-002",
     "Covered by TR-UX-002 (Website/index.html embeds FindCare/Frontend as iframe). Confirm in V2 that the FindCare backend ALSO serves the SPA's index when SPACE_ID is set (HF deployment), with Cache-Control: no-cache."),

    # --- ChatGraph nodes (the in-process orchestration of /chat) ---
    ("CAP-N01", "chat_graph: gate node (LLM-needed routing hook)",
     "Code/ConversationalUX/FindCareChat/backend/application/chat_graph.py:114 (gate_node)",
     "Node",
     "GAP",
     "Add BR: 'The /chat turn MUST start with a gate node that decides whether the turn needs an LLM (today: always true; tomorrow: deterministic routes for pagination/filter). The gate is the documented hook for future LLM-bypass paths.'"),

    ("CAP-N02", "chat_graph: safety_check node (admin_unlock + ip_lock + is_emergency, fan-out)",
     "Code/ConversationalUX/FindCareChat/backend/application/chat_graph.py:131 (safety_check)",
     "Node",
     "PARTIAL: BR-16",
     "BR-16 names the user-id/IP lockout and admin unlock with default key 'unlock$123'. Add: 'safety_check MUST run admin_unlock check FIRST, then is_ip_locked, then is_emergency, on every /chat turn in parallel with search_term_extract. A locked IP returns the emergency response without invoking the LLM.'"),

    ("CAP-N03", "chat_graph: search_term_extract node (GPT-4.1-mini extracts colloquial term)",
     "Code/ConversationalUX/FindCareChat/backend/application/chat_graph.py:147 + main.py:717",
     "Node",
     "GAP",
     "Add BR: 'On every /chat turn the system MUST extract the colloquial provider/specialty search term from the user message via gpt-4.1-mini, in parallel with safety_check. The term is used to label the pagination summary and clinical-trial summary (e.g. \"bone doc\" not \"orthopedic surgeon\").' Today no V2 BR mentions colloquial-term preservation."),

    ("CAP-N04", "chat_graph: chat_loop node (LLM tool loop, multi-iter, tracks last provider/trials results)",
     "Code/ConversationalUX/FindCareChat/backend/application/chat_graph.py:179 (chat_loop)",
     "Node",
     "PARTIAL: BR-01",
     "BR-01 names a single ranked-providers turn; the iterative tool-loop with last_provider_result/last_trials_result tracking is unconstrained. Add: 'The /chat tool loop MUST iterate while stop_reason in (tool_calls, tool_use); MUST capture the most recent find_providers result as last_provider_result and the most recent search_clinical_trials result as last_trials_result for post-processing.'"),

    ("CAP-N05", "chat_graph: url_guard node (URLGuardian on every response)",
     "Code/ConversationalUX/FindCareChat/backend/application/chat_graph.py:275 (url_guard)",
     "Node",
     "GAP",
     "Add BR: 'Every /chat response text MUST be passed through URLGuardian (3600s cache, 5s probe timeout) which removes/replaces unreachable URLs. The Skip Snow LinkedIn marker MUST be linkified if the LLM emitted it as plain text.'"),

    ("CAP-N06", "chat_graph: build_pagination + strip_redundant_summary (system-built summary; LLM text de-duped)",
     "Code/ConversationalUX/FindCareChat/backend/application/chat_graph.py:283 + main.py:742",
     "Node",
     "PARTIAL: BR-23, BR-24",
     "BR-23/BR-24 describe the filter and provider list but not the system-built summary. Add: 'When find_providers returns total_count>0, the system MUST build a deterministic summary_message (specialty, geographic scope, count, page range, suggested specialization options). gpt-4.1-mini MUST then strip duplicate summary/pagination/filter-suggestion language from the LLM text so only the provider listing remains. This is GOV-011-STD-001/-002 (no LLM authorship of summary text).'"),

    ("CAP-N07", "chat_graph: build_trials node (ClinicalTrials.gov summary + travel)",
     "Code/ConversationalUX/FindCareChat/backend/application/chat_graph.py:327 (build_trials)",
     "Node",
     "PARTIAL: BR-33",
     "BR-33 mentions clinical trials but does not specify the summary contract. Add: 'When search_clinical_trials returns >0 trials, the system MUST emit a TrialsMeta {trial_count, condition, location, summary_message} and append a [View all on ClinicalTrials.gov](...) deep link with the user's colloquial term URL-encoded.'"),

    ("CAP-N08", "chat_graph: assemble_response + DebugLogger.log_chat (per-turn audit row)",
     "Code/ConversationalUX/FindCareChat/backend/application/chat_graph.py:367 + infrastructure/debug_logger.py",
     "Node",
     "PARTIAL: BR-08",
     "BR-08 names per-turn audit; the DebugLogger contract is narrower (chat_calls collection, dev only). Add: 'Every /chat turn MUST emit one row to {ENV}_Debug.chat_calls with {ip, message_preview[:200], history_len, tool_loop_iters, tokens_in, tokens_out, response_preview[:200], error}. On error, the de-identified chat_history MUST be attached. This is the per-turn debug surface separate from the LangSmith trace required by BR-12.'"),

    ("CAP-N09", "classify_graph: safety_gate node (IP lock + emergency BEFORE classification)",
     "Code/ConversationalUX/FindCareChat/backend/application/graphs/classify_graph.py:41 (safety_gate)",
     "Node",
     "PARTIAL: BR-04, BR-16",
     "BR-04 names a pre-conversation security gate, but classify is not strictly conversational. Add: 'The /classify endpoint MUST gate on safety BEFORE the specialty vector search runs; an IP-locked or emergency message MUST short-circuit to the emergency response.' V2 BR-04 covers /chat; classify needs explicit coverage."),

    ("CAP-N10", "classify_graph: classify_node (vector specialty + location parse + homeo-generalists cache)",
     "Code/ConversationalUX/FindCareChat/backend/application/graphs/classify_graph.py:58 (classify_node)",
     "Node",
     "PARTIAL: BR-22, BR-23",
     "BR-23 mentions the homeopathic checkbox; the classify-time pre-load of homeopathic_generalists for client-side cache is unspecified. Add: 'classify_node MUST also return homeopathic_generalists (Section=Individual && homeopathic_general=true) so the React panel can render the homeopathic-generalist option without a second round-trip.'"),

    # --- Tools (LLM-callable via ToolRouter) ---
    ("CAP-T01", "Tool: find_providers (FindCareService.search_providers)",
     "Code/ConversationalUX/FindCareChat/backend/main.py:206 (registration)",
     "Tool",
     "FULL: TR-FC-002",
     "Recovered by TR-FC-002 (and BR-01)."),

    ("CAP-T02", "Tool: find_specialty_codes (FindCareService.identify_specialty)",
     "Code/ConversationalUX/FindCareChat/backend/main.py:207",
     "Tool",
     "PARTIAL: TR-FC-004",
     "TR-FC-004 names specialty_classifier as a node, not a tool. The pre-SQL system also exposes find_specialty_codes as an LLM-callable tool. Add: 'find_specialty_codes MUST remain LLM-callable so the conversational LLM can resolve fuzzy specialty terms inside chat_loop without a separate /classify round-trip.'"),

    ("CAP-T03", "Tool: search_clinical_trials (EvaluateCareFacade.search_clinical_trials)",
     "Code/ConversationalUX/FindCareChat/backend/main.py:208",
     "Tool",
     "FULL: TR-FC-003",
     "Recovered by TR-FC-003."),

    ("CAP-T04", "Tool: lookup_provider_external (EvaluateCareFacade.get_provider_details — NPI Registry + state board + Healthgrades links)",
     "Code/ConversationalUX/FindCareChat/backend/main.py:209",
     "Tool",
     "GAP",
     "Add BR: 'The system MUST expose a tool that, given {provider_name, npi, state}, returns the live NPI Registry record (name, status, credential, specialty, license, license_state, address, phone, enumeration_date) plus a list of research links (Healthgrades, NPI Registry, state medical board for DE/MS/VA), each with copy-paste guidance for the user.' Not the same as TR-EC-001's evaluate_view; this is the external-research helper."),

    ("CAP-T05", "Tool: record_user_details (LeadService.record_user_details)",
     "Code/ConversationalUX/FindCareChat/backend/main.py:210",
     "Tool",
     "FULL: TR-FC-005",
     "Recovered by TR-FC-005, in line with BR-18/BR-19's consent cascade."),

    ("CAP-T06", "Tool: record_unknown_question (UnknownQuestionService.record)",
     "Code/ConversationalUX/FindCareChat/backend/main.py:211",
     "Tool",
     "FULL: TR-FC-005",
     "Recovered by TR-FC-005 / BR-17 (clinical refusal + irrelevant + healthcare_capability templates)."),

    ("CAP-T07", "Tool: get_skip_snow_context (AboutService.get_skip_snow_context)",
     "Code/ConversationalUX/FindCareChat/backend/main.py:212",
     "Tool",
     "FULL: TR-FC-006",
     "Recovered by TR-FC-006."),

    ("CAP-T08", "Tool: get_chathealthy_context (AboutService.get_chathealthy_context)",
     "Code/ConversationalUX/FindCareChat/backend/main.py:213",
     "Tool",
     "FULL: TR-FC-006",
     "Recovered by TR-FC-006."),

    ("CAP-T09", "Tool: commitSignificantActivity (generic Mongo write tool)",
     "Code/ConversationalUX/FindCareChat/backend/main.py:214 + main.py:118",
     "Tool",
     "GAP",
     "Add BR or DELETE: 'commitSignificantActivity is a generic Mongo write tool exposed to the LLM. Either: (a) bind it explicitly to LeadService/UnknownQuestionService and remove the LLM-callable shape (recommended for principle-of-least-tool); or (b) document the allowed databases/collections and validate inputs.' V2 has no BR for this — it is implementation-invention by V2's BR-30 standard."),

    ("CAP-T10", "ToolRouter: Pydantic-validated allowlist dispatch (rejects unregistered tools, validates inputs)",
     "Code/ConversationalUX/FindCareChat/backend/application/tool_router.py",
     "Service",
     "GAP",
     "Add BR: 'Every LLM tool call MUST pass through an allowlist registry; unregistered tool names MUST be blocked at the dispatch boundary; registered tools with Pydantic models MUST validate inputs before the handler runs (GOV-004: \"the model may suggest, the system must decide\"). The registry MUST log every dispatch with tool name and validation outcome.'"),

    ("CAP-T11", "ToolRouter: chat-history injection for record_user_details / record_unknown_question",
     "Code/ConversationalUX/FindCareChat/backend/application/tool_router.py:91",
     "Behavior",
     "GAP",
     "Add BR: 'When the LLM calls record_user_details or record_unknown_question, the dispatcher MUST inject the full (un-truncated) chat history into the call arguments before invoking the handler. The LLM MUST NOT be relied on to forward the conversation; the system owns the conversation record.'"),

    # --- Domain services (business behaviors) ---
    ("CAP-S01", "FindCareService: state allowlist (DE/MS/VA) with 'noted interest' fallback",
     "Code/ConversationalUX/FindCareChat/backend/domain/find_care/provider_search_service.py:280",
     "Behavior",
     "PARTIAL: BR-32",
     "BR-32 says 'all remaining states must be loaded into the front-end DB' (pipeline). Add: 'Until BR-32 is satisfied, FindCareService MUST reject states outside the SUPPORTED_STATES allowlist (DE/MS/VA) with {supported:false, state, message:\"FindCare is currently available in Delaware (DE), Mississippi (MS), and Virginia (VA) only. We've noted interest in {state}.\"}. The interest log MUST persist for sales/marketing.'"),

    ("CAP-S02", "FindCareService: NPI exact lookup route",
     "Code/ConversationalUX/FindCareChat/backend/domain/find_care/provider_search_service.py:317",
     "Behavior",
     "GAP",
     "Add BR: 'Given an NPI, the system MUST return the single matching provider record (or empty providers + message). Bypasses vector search entirely.'"),

    ("CAP-S03", "FindCareService: name search route (last/first/organization name regex, OR)",
     "Code/ConversationalUX/FindCareChat/backend/domain/find_care/provider_search_service.py:325",
     "Behavior",
     "GAP",
     "Add BR: 'Given a provider name (and optional state), the system MUST OR-search provider_last_name_legal_name, provider_first_name, and provider_organization_name_legal_business_name with case-insensitive regex; results paginate by NPI ascending.'"),

    ("CAP-S04", "FindCareService: specialty_codes direct filter route (city/county filters + specialization_options return)",
     "Code/ConversationalUX/FindCareChat/backend/domain/find_care/provider_search_service.py:336",
     "Behavior",
     "PARTIAL: BR-23",
     "BR-23 specifies the filter UI but not the server contract. Add: 'When called with specialty_codes, the system MUST filter providers by taxonomies.code in codes, plus optional state/city/county; MUST also return specialization_options with each option's Display Name, can_prescribe, homeopathic, code so the client can label the filter panel without a second call.'"),

    ("CAP-S05", "FindCareService: facet_query keyset pagination (after_npi cursor + total_count facet)",
     "Code/ConversationalUX/FindCareChat/backend/domain/find_care/provider_search_service.py:109",
     "Behavior",
     "PARTIAL: BR-24",
     "BR-24 says 'first 25 providers ... cursor to scroll'. Add: 'Pagination MUST use NPI-keyset (sort by NPI ascending, after_npi cursor); each response MUST include {has_more, first_npi, last_npi, count, total_count, page_start, page_end} so the React UI can render \"Showing X-Y of Z\".' The $facet aggregation that returns count + total in one Mongo round-trip is not specified anywhere in V2."),

    ("CAP-S06", "FindCareService: vector search fallback (provider embeddings, $vectorSearch with city/state/county filter)",
     "Code/ConversationalUX/FindCareChat/backend/domain/find_care/provider_search_service.py:152",
     "Behavior",
     "GAP",
     "Add BR: 'When no specialty_codes are supplied but a free-text specialty_query is, the system MUST execute a $vectorSearch on the providers collection (provider_vector_index, embedding path, text-embedding-3-large query vector) with optional state/city/county filter, returning the top-k by cosine similarity. This is the inferential-AI provider pick the regression event removed.'"),

    ("CAP-S07", "FindCareService: FIPS county map (state+county -> provider filter)",
     "Code/ConversationalUX/FindCareChat/backend/domain/find_care/provider_search_service.py:56",
     "Behavior",
     "GAP",
     "Add BR: 'County filter MUST resolve via the FIPS county map (state + canonical county name -> the provider documents whose practice_address.county matches). The map is loaded once at service start.'"),

    ("CAP-S08", "FindCareService.identify_specialty: legacy entry retained for tool path",
     "Code/ConversationalUX/FindCareChat/backend/domain/find_care/provider_search_service.py:487",
     "Behavior",
     "PARTIAL: TR-FC-004",
     "TR-FC-004 puts specialty_classifier as a graph node. Pre-SQL also has FindCareService.identify_specialty as a tool entry. Decide & document: keep both, or collapse, with explicit BR."),

    ("CAP-S09", "SpecialtyService.find_specialty_codes: vector search of NUCC SpecialtyMetaData (Individual section, score>0.4 threshold, fallback top-5)",
     "Code/ConversationalUX/FindCareChat/backend/domain/find_care/specialty_service.py:37",
     "Behavior",
     "PARTIAL: BR-23, TR-FC-004",
     "BR-23 says 'No inappropriate providers are supplied' and 'within the set the semantic query surfaces'. Add: 'find_specialty_codes MUST use $vectorSearch on SpecialtyMetaData with index=specialty_vector_index, numCandidates=883, limit=698, filter Section=Individual; MUST drop \"Deactivated\" rows; MUST keep score>0.4 matches (fallback to top 5 if none cross threshold); MUST add 1..n rank by descending similarity. Query vector built from last 5 user prompts + current query.'"),

    ("CAP-S10", "specialty_classifier: GPT-4.1-mini classifies NUCC specialties (can_prescribe, homeopathic) with MongoDB cache (specialty_classification collection)",
     "Code/ConversationalUX/FindCareChat/backend/domain/find_care/specialty_classifier.py",
     "Service",
     "GAP",
     "Add BR: 'NUCC specialty classification (can_prescribe, homeopathic) MUST be produced by gpt-4.1-mini and cached in {ENV}_PublicHealthData.specialty_classification with {code, name, can_prescribe, homeopathic, source, classified_at}. Cache is human-correctable. RISK-002 (AI classification) is accepted.' Today V2 has no BR for this batch classification."),

    ("CAP-S11", "specialty_ranker: GPT-4.1-mini reranks specialty options by query relevance (one call per /classify)",
     "Code/ConversationalUX/FindCareChat/backend/domain/find_care/specialty_ranker.py",
     "Service",
     "GAP",
     "Add BR: 'After vector specialty search, the system MUST rerank results by gpt-4.1-mini's per-query relevance assessment so \"find me a kids doc\" surfaces Pediatrics first and \"chest pain\" surfaces Cardiology first. One LLM call per /classify; same set, different order.'"),

    ("CAP-S12", "homeopathic_resolver: when homeopathic checkbox toggled, GPT-4.1-mini scores each homeopathic specialty against the query (strictly_compliant / loosely_compliant / out_of_scope)",
     "Code/ConversationalUX/FindCareChat/backend/domain/find_care/homeopathic_resolver.py",
     "Service",
     "PARTIAL: BR-23",
     "BR-23 says 'When homeopathic check box is checked then all homeopathic providers that could treat the condition are presented.' Specify: 'The homeopathic-checkbox expansion MUST consult homeopathic_resolver, which sends the user query + each homeopathic-flagged or alternative-keyword (naturo/chiro/acupunctur/holistic/integrative/herbal/...) specialty to Anthropic Claude and adds only strictly_compliant + loosely_compliant codes to the search filter. Already-present codes MUST be excluded.'"),

    ("CAP-S13", "SafetyService.is_emergency: AI primary (gpt-4.1-mini, three-gate prompt), keyword fallback only on AI failure, safe-prefix fast path",
     "Code/ConversationalUX/FindCareChat/backend/domain/shared/safety/safety_service.py:49",
     "Behavior",
     "PARTIAL: BR-16",
     "BR-16 names lockout + admin unlock; the detector contract is unspecified. Add: 'Emergency detection MUST be AI-primary (gpt-4.1-mini, response_format json_object) with a three-gate prompt requiring (1) specific body location, (2) acute onset/severity, (3) life-threat. Keyword match is fallback ONLY when the AI call fails. Safe prefixes (\"can you tell me more about...\", \"find me...\", \"who is...\") fast-path past the AI call. Confidence threshold 0.80.'"),

    ("CAP-S14", "SafetyService.is_ip_locked / lock_ip / try_admin_unlock (Mongo-backed, 1h default expiry, audit history)",
     "Code/ConversationalUX/FindCareChat/backend/domain/shared/safety/safety_service.py:102",
     "Behavior",
     "PARTIAL: BR-16",
     "BR-16 says default 24h with key 'unlock$123'. Pre-SQL code defaults to 1h (3600s) and reads ADMIN_UNLOCK_KEY from env. Reconcile in V2: confirm default 24h vs 1h; add: 'Locks persist in {ENV}_Safety.emergency_incidents with {ip, locked_at, expires_at, trigger_message[:500], chat_history}. Admin unlock sets {unlocked:true, unlocked_at, unlocked_by:admin}.'"),

    ("CAP-S15", "ConsentService.summarize_conversation: 2-3 sentence Claude Haiku summary for tier-2 consent",
     "Code/ConversationalUX/FindCareChat/backend/domain/shared/consent/consent_service.py:29",
     "Behavior",
     "PARTIAL: BR-18",
     "BR-18 names the consent cascade (verbatim -> de-identified -> contact only). Add: 'Tier-2 consent (de-identified summary) MUST produce a 2-3 sentence summary via Claude Haiku focused on what the user wanted; this summary is then de-identified.'"),

    ("CAP-S16", "ConsentService.de_identify: Claude Haiku HIPAA Safe Harbor de-identification, in-place mutation",
     "Code/ConversationalUX/FindCareChat/backend/domain/shared/consent/consent_service.py:52",
     "Behavior",
     "PARTIAL: BR-18",
     "BR-18 mentions de-identified transcript. Add: 'De-identification MUST follow HIPAA Safe Harbor (remove names, geographic identifiers except state, dates except year, phone, email, SSN, MRN, account/license/vehicle/device/IP/URL identifiers) via Claude Haiku, returning a JSON array of strings (one per message) that replaces the original chat_history content.'"),

    ("CAP-S17", "LeadService.record_user_details: dedupe-by-email + tier-aware storage (verbatim/summary/contact only) + email push",
     "Code/ConversationalUX/FindCareChat/backend/domain/shared/lead_capture/lead_service.py:32",
     "Behavior",
     "PARTIAL: BR-18, BR-19",
     "Add: 'record_user_details MUST dedupe by email (return ok if email already in {ENV}_AboutUs.lead); MUST store verbatim chat_history when consent_verbatim=true; otherwise store summarized+de-identified notes when consent_summary=true; otherwise store only contact fields. Every lead MUST trigger a SparkPost notification to NOTIFICATION_TO_EMAIL.'"),

    ("CAP-S18", "UnknownQuestionService.record: 3 question_class templates (healthcare_capability, medical_advice, irrelevant) + verbatim presentation",
     "Code/ConversationalUX/FindCareChat/backend/domain/shared/unknowns/unknown_question_service.py:46",
     "Behavior",
     "PARTIAL: BR-17",
     "BR-17 names the three branches (clinical refusal, nonsense, legitimate-non-clinical). Add: 'Templates are keyed healthcare_capability | medical_advice | irrelevant; each template MUST be presented to the user VERBATIM; consent flow re-uses the BR-18 cascade. Templates ARE configurable from the DB (BR-17 sentence: \"This value must be configurable from the Db\"); pre-SQL code has them inline with a TODO — bind a TR.'"),

    ("CAP-S19", "AboutService.get_skip_snow_context / get_chathealthy_context (LinkedIn URL, business plan, Anthropic principles, trim_fn)",
     "Code/ConversationalUX/FindCareChat/backend/domain/shared/content/about_service.py",
     "Service",
     "FULL: TR-FC-006",
     "Recovered by TR-FC-006."),

    ("CAP-S20", "ClinicalTrialsService: ClinicalTrials.gov v2 search (recruiting only, max 10 per page) + Google Routes API drive distance/time",
     "Code/ConversationalUX/FindCareChat/backend/domain/evaluate_care_quality/clinical_trials_service.py",
     "Service",
     "PARTIAL: BR-33",
     "BR-33 specifies hierarchy of services + diagnoses + distance via Google API in days/hours. Add: 'ClinicalTrialsService MUST query https://clinicaltrials.gov/api/v2/studies with overallStatus=RECRUITING, pageSize<=10, format=json; MUST enrich with Google Routes API distance (miles) and duration (Hh Mm) when user_location is supplied; MUST gracefully no-op travel when GOOGLE_MAPS_API_KEY is absent.'"),

    ("CAP-S21", "ProviderDetailService: NPI Registry lookup + state board + Healthgrades research links",
     "Code/ConversationalUX/FindCareChat/backend/domain/evaluate_care_quality/provider_detail_service.py",
     "Service",
     "GAP",
     "Add BR: 'Provider detail MUST hit NPI Registry (npiregistry.cms.hhs.gov/api?number={npi}&version=2.1) for canonical name/credential/specialty/license/address/phone/enumeration; MUST construct research-link cards for Healthgrades (search by name+state) and the state medical board for DE/MS/VA; each link MUST carry copy-paste guidance describing what the user will see on that site.' This is BR-06's read-path but with an external-source detail layer that V2 has no BR for."),

    ("CAP-S22", "EvaluateCareFacade: facade-to-facade entry (calls FindCareFacade for cross-domain provider location data)",
     "Code/ConversationalUX/FindCareChat/backend/application/facades/evaluate_care_facade.py",
     "Service",
     "GAP",
     "Add BR: 'EvaluateCare cross-domain calls (e.g., needing a provider's location) MUST go through FindCareFacade, NEVER directly to FindCareService. The facade-to-facade boundary is the V7 capability-folder rule (TR-BG-004) at the runtime call layer.'"),

    ("CAP-S23", "EmbeddingClient.get_query_embedding / get_specialty_vector / expand_query_terms (text-embedding-3-large + Claude Haiku expansion)",
     "Code/ConversationalUX/FindCareChat/backend/infrastructure/embeddings/embedding_client.py",
     "Service",
     "PARTIAL: TR-SS-002",
     "TR-SS-002 says SharedServices Tool Router exposes embeddings. Specify: 'EmbeddingClient MUST use OpenAI text-embedding-3-large for BOTH provider vector search AND specialty vector search (cross-collection RAG requires identical model). expand_query_terms uses Claude Haiku to produce keyword stems.'"),

    ("CAP-S24", "DebugLogger.log_chat: per-turn metadata + de-identified history on errors (Mongo {ENV}_Debug.chat_calls)",
     "Code/ConversationalUX/FindCareChat/backend/infrastructure/debug_logger.py",
     "Service",
     "PARTIAL: BR-08",
     "Already covered above as CAP-N08; included here as a separate service entry."),

    ("CAP-S25", "URLGuardian: URL liveness check with 3600s cache + 5s probe timeout (deletes/replaces dead URLs in LLM text)",
     "Code/ConversationalUX/FindCareChat/backend/main.py:23 (import) + url_guardian package",
     "Service",
     "GAP",
     "Add BR: 'Every LLM-emitted response MUST be passed through URLGuardian which: (a) probes URLs the LLM emitted with a 5s HEAD request, cached 3600s; (b) removes dead URLs and replaces them with the user-friendly equivalent (e.g., raw text for Skip Snow on LinkedIn -> [Skip Snow on LinkedIn](https://...)). The model is not allowed to publish broken or hallucinated URLs.'"),

    # --- Other Behaviors ---
    ("CAP-B01", "PromptSystemMaker: brain-driven system prompt assembly + welcome message + emergency keywords + tool definitions (single source of truth)",
     "Code/Shared/prompt_system_maker.py (loaded by main.py:55)",
     "Behavior",
     "GAP",
     "Add BR: 'All LLM configuration (system prompt, welcome message, emergency keywords, tool schemas) MUST be assembled by PromptSystemMaker reading from brain/machine_artifacts/content/. main.py MUST NOT contain prompt or keyword string literals. follow_up_check toggles every 5th user message.'"),

    ("CAP-B02", "Startup primitive verification (SEC-HTTPS-001-REQ-014): generate_session_token must produce 68-char CH-prefixed token or process exits 78/77/70",
     "Code/ConversationalUX/FindCareChat/backend/main.py:280",
     "Behavior",
     "GAP",
     "Add BR: 'Backend startup MUST verify the session-token primitive succeeds (well-formed 68-char CH-prefixed token, no underscore in the random portion). On failure: exit 78 (config), 77 (perm), or 70 (software). Service MUST NOT serve traffic if the primitive cannot initialize. Guards against BUG-SEC-007 fallback regression.'"),

    ("CAP-B03", "Cert bootstrap from HF Space env vars (FINDCARE_SIGNING_KEY_PEM, FINDCARE_CERT_PEM, CA_CERT_PEM -> /tmp/ch_certs)",
     "Code/ConversationalUX/FindCareChat/backend/main.py:237",
     "Behavior",
     "GAP",
     "Add BR: 'On HF Space startup the backend MUST decode FINDCARE_SIGNING_KEY_PEM, FINDCARE_CERT_PEM, CA_CERT_PEM (base64 PEM) into /tmp/ch_certs with mode 0600 and set CERTS_DIR. If env vars are absent, it falls through to the bind-mounted ../Shared/ops/certs path for local dev. Malformed base64 MUST raise so the startup-verification guard turns it into exit 78.'"),

    ("CAP-B04", "CORS: allow chathealthy.ai + localhost + *.chathealthy.ai sub-domains; credentials disabled",
     "Code/ConversationalUX/FindCareChat/backend/main.py:325",
     "Behavior",
     "FULL: BR-05",
     "Implicitly covered by BR-05 (static page surface) + TR-UX-002. Confirm in V2 the exact origin allowlist."),

    ("CAP-B05", "Request-log middleware (every request -> INFO line with method, path, status, elapsed_ms, x-forwarded-for IP)",
     "Code/ConversationalUX/FindCareChat/backend/main.py:315",
     "Behavior",
     "PARTIAL: BR-12",
     "BR-12 names end-to-end observability via LangSmith. Add: 'Every HTTP request MUST emit a structured access log line with method, path, status, elapsed_ms, and the x-forwarded-for chain (or client.host fallback). This is the Layer-7 audit complementing the LangGraph trace.'"),

    ("CAP-B06", "Error-class triage in /chat (rate_limit on 429/rate_limit, db_unavailable on connection issues, internal otherwise; debug log written for every error)",
     "Code/ConversationalUX/FindCareChat/backend/main.py:699",
     "Behavior",
     "GAP",
     "Add BR: 'On /chat error the system MUST classify into rate_limit (HTTP 429 / 'rate_limit' in str), db_unavailable ('unavailable'/'connection' in str), or internal; MUST emit a DebugLogger row even on error; MUST return ChatResponse(error=msg, error_type=type) so the React UI can render the right banner.'"),

    ("CAP-B07", "ME (skipsnow + chathealthy) context loading from local ME_DIR or HF flat layout, trimmed via PromptSystemMaker.trim",
     "Code/ConversationalUX/FindCareChat/backend/main.py:165",
     "Behavior",
     "PARTIAL: TR-FC-006",
     "TR-FC-006 names the tools but not the source. Add: 'ME context MUST be loaded at startup from $ME_DIR or the ChatHealthyWhoAmIChat/me/ fallback; large fields (linkedin, business_plan) MUST be trimmed to 3000 chars before serving as tool output.'"),

    ("CAP-B08", "Index health-check (providers/provider_vector_index, SpecialtyMetaData/specialty_vector_index) drives /health degraded state",
     "Code/ConversationalUX/FindCareChat/backend/main.py:444",
     "Behavior",
     "GAP",
     "Add BR: 'The required vector search indexes (providers/provider_vector_index, SpecialtyMetaData/specialty_vector_index) MUST be enumerated at the application level; missing indexes degrade /health and MUST log an ERROR. The list is the canonical declaration of \"what indexes the live app depends on\".'"),

    ("CAP-B09", "QA-report Mongo-or-file bootstrap pattern (Mongo source of truth, file fallback seeds Mongo on first read)",
     "Code/ConversationalUX/FindCareChat/backend/main.py:396",
     "Behavior",
     "GAP",
     "Add BR companion to CAP-R04: 'On first /qa-report load, if {ENV}_System.bugs has no qa_report_v014_sit doc, the system MUST seed it from brain/machine_artifacts/content/qa_report_v014_sit.json (with brain_data/ fallback for HF flat layout) and replace_one upsert. After bootstrap Mongo is the canonical source.'"),

    ("CAP-B10", "Strict consent-flow for record_user_details: REQUIRES email; otherwise short-circuits with note",
     "Code/ConversationalUX/FindCareChat/backend/domain/shared/lead_capture/lead_service.py:36",
     "Behavior",
     "PARTIAL: BR-18",
     "BR-18 says 'attempts to collect them. If the user says the final no no record is kept'. Add: 'record_user_details MUST short-circuit with {recorded:ok, note:\"Email required but not provided\"} if email is empty/whitespace; this prevents partial-record creation and matches BR-18's no-email-no-record rule.'"),
]


# ---------------------------------------------------------------------------
# Build doc
# ---------------------------------------------------------------------------
def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # Title
    add_heading(doc, "Pre-SQL Capability Audit V1", level=0)
    add_para(doc, "ChatHealthy.ai - Capability inventory of the codebase IMMEDIATELY BEFORE the 2026-04-19 SQL/graph-routing regression event, compared against design-V2.docx Business Requirements.", italic=True)
    add_para(doc, "Date: 2026-04-28   |   Author: Claude (capability-audit agent)   |   Read-only deliverable - no commits, no code changes, no checkouts.", italic=True)

    # ------------------------------------------------------------------
    # 1. Executive summary
    # ------------------------------------------------------------------
    add_heading(doc, "1. Executive summary", level=1)

    full_count = sum(1 for c in CAPABILITIES if c[4].startswith("FULL"))
    partial_count = sum(1 for c in CAPABILITIES if c[4].startswith("PARTIAL"))
    gap_count = sum(1 for c in CAPABILITIES if c[4].startswith("GAP"))
    total = len(CAPABILITIES)

    add_para(doc, f"Capabilities enumerated in pre-SQL code: {total}")
    add_para(doc, f"  FULL coverage by V2 BR/TR:    {full_count}")
    add_para(doc, f"  PARTIAL coverage by V2 BR/TR: {partial_count}")
    add_para(doc, f"  GAP (no V2 BR/TR):            {gap_count}")
    add_para(doc, "")
    add_para(doc, "Headline finding: V2's BR set covers the conversational shell well (LC-01..LC-12 all have a TR) but is thin on the ROUTE-LEVEL contracts (POST /search, POST /classify shape, /transfer, /shared/*, /evaluate/*), the AI-CLASSIFIER LAYER between the user and the provider DB (specialty_classifier batch, specialty_ranker, homeopathic_resolver), and the IMPLEMENTATION-LEVEL behaviors that protected the system from regression (URLGuardian, PromptSystemMaker, startup primitive verification, Pydantic-validated tool dispatch). These are precisely the layers the SQL-insertion event removed.", italic=False)

    # ------------------------------------------------------------------
    # 2. Pre-SQL commit identification
    # ------------------------------------------------------------------
    add_heading(doc, "2. Pre-SQL commit identification", level=1)
    add_para(doc, "Commit hash: f836d4d0cdd9493b2f07017f434cd0b9fba1925e", bold=True)
    add_para(doc, "Date: 2026-04-19 21:22:56 -0700", bold=True)
    add_para(doc, "Build number: 998", bold=True)
    add_para(doc, "Author: SkipSnow <skip.snow@gmail.com>")
    add_para(doc, "Subject: [Idiot Mode] data: build 998 - post-commit hook seeds new version.json shape (Refs: DEVOPS-DEPLOY-001-REQ-016)")
    add_para(doc, "Evidence:")
    add_para(doc, "  - version.json at this commit shows latest builds.build[].build_number = 998 (verified via 'git show f836d4d0:brain/machine_artifacts/content/version.json | tail -5').")
    add_para(doc, "  - All 26 BUG-LANGCHAIN-CONTAINER-* entries in bugs.json carry discovery_date 2026-04-19 (the regression-event date). Build 999 (the very next commit, hash f836d4d0 -> 02f404f4) carries cause EPIC-008-F-002-S-006-REQ-T-001 and files-touched include uat_handoff_2026_04_19_langgraph.docx, anchoring the regression at build 999.")
    add_para(doc, "  - design-V2 Section 2.1 explicitly states: 'BEFORE the regression event - build 998 (cause EPIC-008-F-004-S-001-REQ-T-002, on or just before 2026-04-19).' That is the commit this audit reads.")

    # ------------------------------------------------------------------
    # 3. Capability inventory table
    # ------------------------------------------------------------------
    add_heading(doc, "3. Capability inventory", level=1)
    add_para(doc, "Each row: capability ID, capability name, source file:line, type (Tool / Service / Node / Route / Behavior), V2 BR/TR coverage, recommended new BR text if PARTIAL or GAP. All source paths read at commit f836d4d0 via 'git show'.", italic=True)

    headers = ["ID", "Capability", "Source file:line", "Type", "V2 coverage", "Recommended BR text (PARTIAL/GAP only)"]
    rows = [(c[0], c[1], c[2], c[3], c[4], c[5]) for c in CAPABILITIES]
    # Wide column A for IDs (per Skip's "column A not skinny" annotation in design-V2)
    add_table(doc, headers, rows,
              col_widths=[0.9, 2.0, 2.1, 0.7, 1.1, 3.6])

    # ------------------------------------------------------------------
    # 4. Gap summary
    # ------------------------------------------------------------------
    add_heading(doc, "4. Gap summary (PARTIAL + GAP at a glance)", level=1)

    add_heading(doc, "4.1 GAP items (no V2 BR/TR coverage at all)", level=2)
    gap_rows = [(c[0], c[1], c[2]) for c in CAPABILITIES if c[4].startswith("GAP")]
    add_table(doc,
              ["ID", "Capability", "Source file:line"],
              gap_rows,
              col_widths=[0.9, 4.5, 2.6])

    add_heading(doc, "4.2 PARTIAL items (named in a BR/TR but contract under-specified)", level=2)
    partial_rows = [(c[0], c[1], c[4]) for c in CAPABILITIES if c[4].startswith("PARTIAL")]
    add_table(doc,
              ["ID", "Capability", "V2 BR/TR named"],
              partial_rows,
              col_widths=[0.9, 4.5, 2.6])

    # ------------------------------------------------------------------
    # 5. Recommendations (consolidated new BRs)
    # ------------------------------------------------------------------
    add_heading(doc, "5. Recommendations - proposed new BR text", level=1)
    add_para(doc, "These are written at the specificity of Skip's V2-added BRs (BR-16..BR-33) - one user-visible behavior per BR, with positive and negative test calls.", italic=True)

    rec_rows = [(c[0], c[5]) for c in CAPABILITIES if not c[4].startswith("FULL")]
    add_table(doc,
              ["Cap ID", "Proposed BR text"],
              rec_rows,
              col_widths=[0.9, 9.6])

    # ------------------------------------------------------------------
    # 6. Anomalies - regression-recovery items
    # ------------------------------------------------------------------
    add_heading(doc, "6. Anomalies - capabilities that existed pre-SQL but appear absent post-SQL", level=1)
    add_para(doc, "These are the items the architect specifically asked to recover. Each is a service or behavior that was wired into main.py at build 998 but whose connection to the live HTTP surface was removed when the LangChain containers became empty.", italic=True)

    anomalies = [
        ("A-01",
         "FindCareService vector search ($vectorSearch on providers collection, text-embedding-3-large query)",
         "Code/ConversationalUX/FindCareChat/backend/domain/find_care/provider_search_service.py:152",
         "This is the inferential AI provider pick. design-V2 Section 1 explicitly identifies its removal: 'Claude code eliminated ~9/12 functions with one refactor to eliminate a non-deterministic call to AI and substitute a deterministic call to the data base.' At build 998 the vector route is wired into FindCareService.search_providers and reachable via tool_find_providers; post-SQL the FastAPI handler became an empty container."),

        ("A-02",
         "specialty_ranker (per-query GPT-4.1-mini reranking of vector specialty results)",
         "Code/ConversationalUX/FindCareChat/backend/domain/find_care/specialty_ranker.py",
         "Reranks the same vector-search result set so the most clinically relevant specialty appears first per query. No V2 BR mentions reranking; Skip's BR-23 only says results 'are ranked by the most likely to be picked'."),

        ("A-03",
         "homeopathic_resolver (Anthropic Claude evaluating each homeopathic specialty against the user's query, three-class output)",
         "Code/ConversationalUX/FindCareChat/backend/domain/find_care/homeopathic_resolver.py",
         "When the user toggles the Homeopathic checkbox, this service expands the candidate set with strictly_compliant + loosely_compliant homeopathic specialties picked by an LLM. BR-23 implies expansion ('all homeopathic providers that could treat the condition') but does not bind the resolver."),

        ("A-04",
         "specialty_classifier (gpt-4.1-mini batch classification of NUCC codes -> can_prescribe + homeopathic, cached in Mongo)",
         "Code/ConversationalUX/FindCareChat/backend/domain/find_care/specialty_classifier.py",
         "Generates the can_prescribe + homeopathic flags BR-23 depends on. Without it, the prescriber filter and homeopathic filter cannot be implemented from data alone - they're derived attributes."),

        ("A-05",
         "URLGuardian (live URL validation + dead-link removal + Skip-Snow-LinkedIn linkification)",
         "Code/ConversationalUX/FindCareChat/backend/main.py:23 import + per-turn invocation in chat_graph.url_guard",
         "Prevented the LLM from publishing broken or hallucinated URLs to the user. No V2 BR mentions URL liveness. Easy regression target."),

        ("A-06",
         "Pydantic-validated allowlist tool dispatch (ToolRouter)",
         "Code/ConversationalUX/FindCareChat/backend/application/tool_router.py",
         "GOV-004's server-side enforcement layer ('the model may suggest, the system must decide'). V2 has TR-LG-007 (ToolMessage on the messages channel) but does not specify the allowlist + Pydantic validation contract. Without it, an LLM hallucinating a tool name or malformed args silently runs whatever Python function happens to be named that way."),

        ("A-07",
         "Per-turn DebugLogger row in {ENV}_Debug.chat_calls (with de-identified history on errors)",
         "Code/ConversationalUX/FindCareChat/backend/infrastructure/debug_logger.py",
         "Local Mongo-resident debug audit complementing LangSmith. BR-08/BR-12 cover the high-level audit but the chat_calls row is what Skip uses today to triage prompt failures."),

        ("A-08",
         "QA-report system (GET + POST /qa-report rendering and persisting test-case status; Mongo-or-file bootstrap)",
         "main.py:418, main.py:427, application/graphs/qa_report_graph.py, application/graphs/qa_report_submit_graph.py",
         "The human-facing QA dashboard: where Skip checks off test outcomes per feature. Not in V2 at all."),

        ("A-09",
         "Startup security primitive verification (session_token must produce well-formed token or process exits 78)",
         "main.py:280 (_startup_security_verification)",
         "The guard that catches BUG-SEC-007 fallback regressions. Critical security-defense-in-depth: the service refuses to serve traffic if it cannot mint a real signed token. Not in V2."),

        ("A-10",
         "GOV-011-STD-001 / -002 LLM-output stripping (build_pagination + strip_redundant_summary + _extract_user_search_term)",
         "main.py:717, main.py:742, chat_graph.build_pagination",
         "The system, NOT the LLM, owns the summary text. The LLM's pagination/filter prose is mechanically stripped via gpt-4.1-mini after the system has built the deterministic summary_message. This is what kept the LLM from inventing provider counts or pagination claims. V2 has no BR for this anti-invention pattern."),

        ("A-11",
         "Two-pane control transfer (FindCare <-> EvaluateCare via /evaluate/splash + /transfer/to-findcare; FindCare <-> SharedServices via /shared/splash) with mTLS round-trips logged on both sides",
         "main.py:530, main.py:554, main.py:619",
         "BR-25 says the page is 'controlled by EvaluateCare server' but does not specify the transfer protocol, the round-trip logging, or the failure UX. At build 998 the protocol is concrete and works; post-SQL the empty containers broke it."),

        ("A-12",
         "Cert bootstrap from HF Space env vars (FINDCARE_SIGNING_KEY_PEM / FINDCARE_CERT_PEM / CA_CERT_PEM -> /tmp/ch_certs)",
         "main.py:237 (_bootstrap_certs_from_env)",
         "The mechanism that lets HF Spaces (which lack bind-mounted cert dirs) participate in mTLS. Not in V2; without it, HF deployments cannot establish FindCare<->EvaluateCare<->SharedServices mTLS."),
    ]
    add_table(doc,
              ["ID", "Capability lost", "Source", "Why this is a regression-recovery item"],
              [(a[0], a[1], a[2], a[3]) for a in anomalies],
              col_widths=[0.7, 2.6, 2.5, 4.6])

    # ------------------------------------------------------------------
    # 7. Methodology + caveats
    # ------------------------------------------------------------------
    add_heading(doc, "7. Methodology and caveats", level=1)
    add_para(doc, "Pre-SQL state was inspected with 'git show f836d4d0:<path>' (read-only - no checkout, no working-tree mutation). All file:line citations are at that commit. Claude did NOT run the application; capabilities were read off the source.", italic=True)
    add_para(doc, "V2 BRs were extracted from architecture/ChatHealthyApplicationDesign/ArchitectureDesignAndAuditDocs/design-V2.docx via python-docx (XML extraction of all w:t runs in document order). BR numbering follows the explicit BR-NN headings; Skip's bracketed in-line annotations (<Skip: ...>) were preserved.", italic=True)
    add_para(doc, "Coverage classification: FULL means an explicit V2 BR or TR names the capability AND constrains its behavior. PARTIAL means a BR/TR names it but leaves the contract under-specified (e.g., does not bind the input/output shape, does not bind error behavior, does not bind which collection it reads). GAP means no BR or TR mentions it. By BR-30 ('No capability ... CAN exist if there is not a specific business requirement it traces to'), every PARTIAL and GAP entry is a candidate for either a new BR or for explicit deletion.", italic=True)
    add_para(doc, "Ambiguity notes: (1) commitSignificantActivity (CAP-T09) is registered as an LLM-callable tool but writes arbitrary Mongo collections; the right answer may be deletion not retention. (2) FindCareService.identify_specialty (CAP-S08) duplicates SpecialtyService.find_specialty_codes (CAP-S09); the architect should pick one. (3) Some 'graph-exempt' route comments (e.g., /health, /session, /evaluate/splash) explicitly disclaim graph routing - V2 should preserve that exemption list rather than force every handler through a graph.", italic=True)

    doc.save(OUT)
    print(f"Wrote {OUT}")
    print(f"Capabilities: {total}  FULL={full_count}  PARTIAL={partial_count}  GAP={gap_count}")


if __name__ == "__main__":
    build()
