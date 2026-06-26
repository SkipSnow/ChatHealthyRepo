#!/usr/bin/env python3
"""Export deployment deliverables for operators:

1. BUILD_DEPLOY_MINIMUM_FACTS_PROPOSAL.docx (from .md in this directory)
2. deployment_facts_inventory_<date>.xlsx + .csv — one row per fact IN deployment_architecture.json ONLY

The workbook is a pure projection of the persisted manifest. No proposed drafts,
no placeholder rows, no reachability audit merge, no objects outside the JSON.

Usage (repo root):
  python architecture/DevOpsBuildDeployAndEnvironmentManagement/export_deployment_deliverables.py
"""
from __future__ import annotations

import csv
import json
import re
from datetime import date
from pathlib import Path

ROOT = Path(__file__).resolve().parents[3]
DEVOPS = Path(__file__).resolve().parent.parent
MANIFEST = ROOT / "brain" / "machine_artifacts" / "content" / "deployment_architecture.json"
ENV_FILE = ROOT / "Code" / ".env"
MD_PATH = DEVOPS / "BUILD_DEPLOY_MINIMUM_FACTS_PROPOSAL.md"
TODAY = date.today().isoformat()

COLUMNS = [
    "fact_kind",
    "name",
    "address",
    "json_path",
    "purpose",
    "targets",
    "value",
    "retrieval_store",
    "retrieval_action",
]

RETRIEVAL_ACTIONS = {
    "local_env": "SecretsResolver.resolve(name, env) -> read key from Code/.env via _read_env_file",
    "hf_space_secret": "HF Space API set-secret at deploy (store stub in SecretsResolver today)",
    "cloudflare_env": "Cloudflare Pages env var API at deploy (store stub today)",
    "azure_function_app_setting": "az functionapp config appsettings set at deploy",
    "azure_automation_variable": "Set Azure Automation Variable on deploy account",
    "azure_automation_webhook": "SKIP resolve — upstream runbook deploy UPSERTs webhook URL onto consumer",
    "gha_secret": "GitHub Actions secret reference at workflow deploy",
    "env_name": "Deploy sets runtime value to env binding name (dev|qa|prod|local)",
    "local_cert_file": "Deploy reads repo-relative cert file, base64-encodes for runtime env",
    "peer_url": "Deploy resolves peer target node_address / HF URL via _hf_peer_url_for_build",
    "rename_from": "Deploy resolves sibling secret/variable entry then pushes under this name",
    "literal": "Deploy uses literal suffix after 'literal:' qualifier",
    "referenced_file": "Build/deploy copies bytes from git tree at source_location (Extractor)",
    "managed_file": "Crosswalk byte-check; Extractor materializes embedded_content to source_location",
    "environment_binding": "LocalDeploy/build reads node_address for env_binding from RecordLoader",
    "wholesale_copy": "hf_helpers._copy_tree from huggingface_source_set at build stage",
    "entitlement": "Deploy provisions role/membership per entitlements[] inventory",
}


def retrieval_action_for(value: str, fact_kind: str) -> tuple[str, str]:
    if fact_kind == "file":
        if value.startswith("managed"):
            return "", RETRIEVAL_ACTIONS["managed_file"]
        return "", RETRIEVAL_ACTIONS["referenced_file"]
    if fact_kind == "environment":
        return "", RETRIEVAL_ACTIONS["environment_binding"]
    if fact_kind == "wholesale_src":
        return "", RETRIEVAL_ACTIONS["wholesale_copy"]
    if fact_kind == "target_record":
        return "", "RecordLoader.load_collection(deployment_architecture.json)"
    if fact_kind == "entitlement":
        return "", RETRIEVAL_ACTIONS["entitlement"]

    if value == "local_env":
        return "local_env", RETRIEVAL_ACTIONS["local_env"]
    if value in RETRIEVAL_ACTIONS:
        return value, RETRIEVAL_ACTIONS[value]
    if value.startswith("local_cert_file:"):
        return "local_cert_file", RETRIEVAL_ACTIONS["local_cert_file"]
    if value.startswith("peer_url:"):
        return "peer_url", RETRIEVAL_ACTIONS["peer_url"]
    if value.startswith("rename_from:"):
        return "rename_from", RETRIEVAL_ACTIONS["rename_from"]
    if value.startswith("literal:"):
        return "literal", RETRIEVAL_ACTIONS["literal"]
    if value == "env_name":
        return "env_name", RETRIEVAL_ACTIONS["env_name"]
    return value, f"Deploy resolves qualifier {value!r} per _deploy_chain._resolve_qualifier"


def file_value_field(f: dict) -> str:
    """JSON fields stored on a files[] row (never secret bytes)."""
    parts = [
        f.get("disposition", ""),
        f.get("handler_type", ""),
    ]
    if f.get("content_hash"):
        parts.append(f"content_hash={f['content_hash']}")
    if f.get("feature_id"):
        parts.insert(0, f["feature_id"])
    return "|".join(p for p in parts if p)


def collect_fact_rows(manifest: dict) -> list[dict]:
    rows: list[dict] = []
    records = manifest.get("DeploymentTargetRecord", [])

    for idx, rec in enumerate(records):
        tid = rec["target_id"]
        tkind = rec.get("target_kind", "")
        base_path = f"DeploymentTargetRecord[{idx}]"

        rows.append({
            "fact_kind": "target_record",
            "name": tid,
            "address": "brain/machine_artifacts/content/deployment_architecture.json",
            "json_path": f"{base_path}.target_id",
            "purpose": f"Target identity; kind={tkind}",
            "targets": tid,
            "value": tid,
            "retrieval_store": "",
            "retrieval_action": "RecordLoader.load_collection(deployment_architecture.json)",
        })

        if rec.get("promote_chain_bound") is not None:
            rows.append({
                "fact_kind": "target_metadata",
                "name": "promote_chain_bound",
                "address": "brain/machine_artifacts/content/deployment_architecture.json",
                "json_path": f"{base_path}.promote_chain_bound",
                "purpose": "Promote-chain branch guard participation",
                "targets": tid,
                "value": str(rec["promote_chain_bound"]).lower(),
                "retrieval_store": "",
                "retrieval_action": "deploy_chathealthy branch-vs-env guard",
            })

        for ei, env in enumerate(rec.get("environments") or []):
            eb = env.get("env_binding", "")
            na = env.get("node_address", "")
            rows.append({
                "fact_kind": "environment",
                "name": eb,
                "address": "brain/machine_artifacts/content/deployment_architecture.json",
                "json_path": f"{base_path}.environments[{ei}].node_address",
                "purpose": f"Runtime binding for {tid} in env {eb}",
                "targets": tid,
                "value": na,
                "retrieval_store": "",
                "retrieval_action": RETRIEVAL_ACTIONS["environment_binding"],
            })

        for fi, f in enumerate(rec.get("files") or []):
            loc = f.get("source_location", "")
            disp = f.get("disposition", "referenced")
            val = file_value_field(f)
            store, action = retrieval_action_for(disp, "file")
            rows.append({
                "fact_kind": "file",
                "name": Path(loc).name if loc else "",
                "address": loc,
                "json_path": f"{base_path}.files[{fi}]",
                "purpose": f"Source file composition for {tid}",
                "targets": tid,
                "value": val,
                "retrieval_store": store,
                "retrieval_action": action,
            })

        for name, store_id in sorted((rec.get("secrets") or {}).items()):
            store, action = retrieval_action_for(store_id, "secret")
            addr = (
                str(ENV_FILE.relative_to(ROOT)).replace("\\", "/")
                if store_id == "local_env"
                else "brain/machine_artifacts/content/deployment_architecture.json"
            )
            rows.append({
                "fact_kind": "secret",
                "name": name,
                "address": addr,
                "json_path": f"{base_path}.secrets.{name}",
                "purpose": f"Secret key for {tid} runtime (value never in JSON)",
                "targets": tid,
                "value": store_id,
                "retrieval_store": store,
                "retrieval_action": action,
            })

        for name, qual in sorted((rec.get("variables") or {}).items()):
            store, action = retrieval_action_for(qual, "variable")
            rows.append({
                "fact_kind": "variable",
                "name": name,
                "address": "brain/machine_artifacts/content/deployment_architecture.json",
                "json_path": f"{base_path}.variables.{name}",
                "purpose": f"Deploy-computed variable for {tid}",
                "targets": tid,
                "value": qual,
                "retrieval_store": store,
                "retrieval_action": action,
            })

        for wi, item in enumerate(rec.get("huggingface_source_set") or []):
            src = item.get("src", "")
            rows.append({
                "fact_kind": "wholesale_src",
                "name": Path(src).name if src else "",
                "address": src,
                "json_path": f"{base_path}.huggingface_source_set[{wi}]",
                "purpose": f"Wholesale build-context copy root for {tid}",
                "targets": tid,
                "value": json.dumps(item, separators=(",", ":")),
                "retrieval_store": "",
                "retrieval_action": RETRIEVAL_ACTIONS["wholesale_copy"],
            })

        for ei, ent in enumerate(rec.get("entitlements") or []):
            kind = ent.get("kind", "")
            rows.append({
                "fact_kind": "entitlement",
                "name": kind,
                "address": "brain/machine_artifacts/content/deployment_architecture.json",
                "json_path": f"{base_path}.entitlements[{ei}]",
                "purpose": f"Access grant for {tid}",
                "targets": tid,
                "value": json.dumps(ent, separators=(",", ":")),
                "retrieval_store": "",
                "retrieval_action": RETRIEVAL_ACTIONS["entitlement"],
            })

    firm = manifest.get("firm") or {}
    if firm.get("git_identity"):
        rows.append({
            "fact_kind": "firm_metadata",
            "name": "git_identity",
            "address": "brain/machine_artifacts/content/deployment_architecture.json",
            "json_path": "firm.git_identity",
            "purpose": "Bot identity for promote/HF git pushes",
            "targets": "firm",
            "value": json.dumps(firm["git_identity"], separators=(",", ":")),
            "retrieval_store": "",
            "retrieval_action": "git config user.name/email at deploy",
        })

    return rows


def write_spreadsheet(rows: list[dict], base: Path) -> None:
    csv_path = base.with_suffix(".csv")
    xlsx_path = base.with_suffix(".xlsx")
    with csv_path.open("w", newline="", encoding="utf-8") as fh:
        w = csv.DictWriter(fh, fieldnames=COLUMNS)
        w.writeheader()
        w.writerows(rows)

    import openpyxl
    from openpyxl.utils import get_column_letter

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "deployment_facts"
    ws.append(COLUMNS)
    for row in rows:
        ws.append([row[c] for c in COLUMNS])
    widths = [14, 28, 50, 45, 40, 36, 48, 18, 55]
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w
    readme = wb.create_sheet("column_guide")
    readme.append(["Column", "Meaning"])
    for line in [
        ("fact_kind", "file | secret | variable | environment | target_record | wholesale_src | entitlement | firm_metadata"),
        ("name", "File basename or secret/variable key name or env_binding"),
        ("address", "Repo path for files; Code/.env for local_env secret VALUES at runtime; manifest path for JSON keys"),
        ("json_path", "Location inside deployment_architecture.json"),
        ("purpose", "Role of this fact for the target"),
        ("targets", "target_id (or firm) owning this fact"),
        ("value", "Exact string(s) stored in JSON — never secret VALUES"),
        ("retrieval_store", "Store id for SecretsResolver when applicable"),
        ("retrieval_action", "What deploy/build code does to obtain runtime bytes or values"),
    ]:
        readme.append(list(line))
    readme.append(["", ""])
    readme.append(["Purity rule", "Every row is persisted in deployment_architecture.json today. No proposed, empty, or audit-only rows."])
    wb.save(xlsx_path)
    print(f"Facts CSV:  {csv_path}")
    print(f"Facts XLSX: {xlsx_path} ({len(rows)} rows)")


def md_to_docx(md_path: Path, docx_path: Path) -> None:
    from docx import Document
    from docx.shared import Pt, RGBColor

    TEAL = RGBColor(0x0B, 0x7A, 0x75)
    text = md_path.read_text(encoding="utf-8")
    doc2 = Document()
    style2 = doc2.styles["Normal"]
    style2.font.name = "Calibri"
    style2.font.size = Pt(11)
    i = 0
    lines = text.splitlines()
    while i < len(lines):
        line = lines[i].rstrip()
        if line.startswith("| ") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|-") :
            headers = [c.strip() for c in line.strip("|").split("|")]
            i += 2
            body_rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                body_rows.append([c.strip() for c in lines[i].strip("|").split("|")])
                i += 1
            tbl = doc2.add_table(rows=1 + len(body_rows), cols=len(headers))
            tbl.style = "Table Grid"
            for ci, h in enumerate(headers):
                tbl.rows[0].cells[ci].text = h
            for ri, row in enumerate(body_rows):
                for ci, cell in enumerate(row):
                    if ci < len(headers):
                        tbl.rows[ri + 1].cells[ci].text = cell
            doc2.add_paragraph("")
            continue
        if line.startswith("# "):
            h = doc2.add_heading(line[2:].strip(), level=1)
            for run in h.runs:
                run.font.color.rgb = TEAL
        elif line.startswith("## "):
            doc2.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc2.add_heading(line[4:].strip(), level=3)
        elif line.startswith("```"):
            i += 1
            block = []
            while i < len(lines) and not lines[i].startswith("```"):
                block.append(lines[i])
                i += 1
            if block:
                doc2.add_paragraph("\n".join(block), style="Intense Quote")
        elif not line.strip():
            pass
        elif re.match(r"^\d+\.\s", line):
            doc2.add_paragraph(line, style="List Number")
        elif line.startswith("- "):
            doc2.add_paragraph(line[2:], style="List Bullet")
        else:
            doc2.add_paragraph(line)
        i += 1

    doc2.save(docx_path)
    print(f"Proposal DOCX: {docx_path}")


def main() -> None:
    manifest = json.loads(MANIFEST.read_text(encoding="utf-8"))
    rows = collect_fact_rows(manifest)
    base = DEVOPS / f"deployment_facts_inventory_{TODAY}"
    write_spreadsheet(rows, base)
    md_to_docx(MD_PATH, DEVOPS / "BUILD_DEPLOY_MINIMUM_FACTS_PROPOSAL.docx")
    print("Rows by fact_kind:")
    from collections import Counter
    for k, v in sorted(Counter(r["fact_kind"] for r in rows).items()):
        print(f"  {k}: {v}")


if __name__ == "__main__":
    main()
