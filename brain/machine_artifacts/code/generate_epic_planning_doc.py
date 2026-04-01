# Copyright (c) 2026 Skip Snow. All rights reserved.
# Licensed under the FindCare Evaluation License (FEL-1.0).
#
# generate_epic_planning_doc.py — Generates epic planning prompt Word doc.
# Function: epic planning doc and JSON production
# Usage: python generate_epic_planning_doc.py

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

# Title page
doc.add_paragraph('')
doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('ChatHealthy.ai')
r.font.size = Pt(28)
r.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Epic Planning Prompt Document')
r.font.size = Pt(18)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('v0.1.4 \u2014 Evaluate Care')
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Function: Epic Planning Doc and JSON Production')
r.font.size = Pt(11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('April 1, 2026 | Build 327')
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Copyright 2026 Skip Snow. All rights reserved.')
r.font.size = Pt(9)
r.italic = True

doc.add_page_break()

# ── Section 1: System Prompt ──────────────────────────────────
doc.add_heading('Section 1: System Prompt (Cached)', level=1)
doc.add_paragraph('Record ID: gpt_assignment_system_prompt')
doc.add_paragraph('Model: gpt-5.3-chat-latest')
doc.add_paragraph('Role: Enterprise Architect')
doc.add_paragraph('Cacheable: Yes (GOV-009)')

doc.add_heading('Identity', level=2)
doc.add_paragraph(
    'You are the Enterprise Architect for ChatHealthy.ai. You design, validate, and approve. '
    'You are precise, thorough, and never truncate required output.'
)

doc.add_heading('Governance', level=2)
for g in [
    'GOV-004: The model may suggest. The system must decide.',
    'Boss controls scope. You assess risk but may NOT defer mandatory items.',
    'Tradeoff suggestions go in tradeoff_suggestions only \u2014 never as a reason to defer.',
]:
    doc.add_paragraph(g, style='List Bullet')

doc.add_heading('Quality', level=2)
doc.add_paragraph('Thin outputs are rejected. Complete professional output required.')

doc.add_heading('Output Schema', level=2)
for s in [
    'assignment_id: string',
    'timestamp: ISO 8601 UTC',
    'architecture_status: pass | fail',
    'behavior_status: pass | fail',
    'risk: Low | Moderate | High | Critical',
    'gate_recommendation: auto | proceed_with_warning | escalate | block_escalate | block_boss_required',
    'issues: array (empty if none)',
    'notes: string',
    'scope_assessment: one entry per MANDATORY SCOPE item',
    'requirements: min 5, derive from code and context',
    'feature_set: ships + deferred',
    'acceptance_criteria: min 3 per shipping feature',
    'uat_scenarios: min 1 per acceptance criterion',
    'usage: tokens_in, tokens_out',
]:
    doc.add_paragraph(s, style='List Bullet')

doc.add_page_break()

# ── Section 2: User Prompt ────────────────────────────────────
doc.add_heading('Section 2: User Prompt', level=1)
doc.add_paragraph('Record ID: epic_planning_user_prompt_v014')
doc.add_paragraph('System Prompt Reference: gpt_assignment_system_prompt')

doc.add_heading('Planning for April 2026, Epic 2 Deliver credible quality module', level=2)
doc.add_paragraph('Date: 2026-04-01 | From: Boss | To: GPT (Enterprise Architect) | Iteration: 1 of 1000')

# Context
doc.add_heading('Context', level=3)
for c in [
    'current_version: v0.1.3',
    'current_build: 327',
    'next_version: v0.1.4',
    'sprint_cadence: weekly, Tuesday start, 7 days/week, PST',
]:
    doc.add_paragraph(c, style='List Bullet')

p = doc.add_paragraph()
r = p.add_run('Team')
r.bold = True
for t in [
    'Boss: CEO/QA, final gate authority',
    'Claude: Engineer, claude-opus-4-6',
    'GPT: Enterprise Architect, gpt-5.3-chat-latest',
]:
    doc.add_paragraph(t, style='List Bullet')

p = doc.add_paragraph()
r = p.add_run('Live Today')
r.bold = True
for l in [
    'FindCare Chat: React + FastAPI on HuggingFace Docker Space',
    'Provider search: DE, MS (MongoDB)',
    'Safety filter: dual-trigger (keyword + GPT-4.1-mini classifier)',
    'Clinical trials: ClinicalTrials.gov API',
    'About Us / Contact: HIPAA two-tier consent flow',
    'The Brain: GPTReader (27/27 SIT), manifest, governance, audit',
    'DataPipelines: Azure Functions, CrewAI, provider load, county enrichment',
    'Website: Cloudflare Pages (static)',
    'Maps: NONE \u2014 provider results are text only',
    'Distance: NONE \u2014 no distance calculation from user location to provider',
    'Find Care: STABLE \u2014 alpha release candidate, sufficient as-is',
]:
    doc.add_paragraph(l, style='List Bullet')

p = doc.add_paragraph()
r = p.add_run('Shipped in v0.1.3')
r.bold = True
for s in [
    'Brain infrastructure (GPTReader, manifest, governance sync)',
    'Safety filter hardening (false positive fix, GPT-4.1-mini, full audit trail)',
    'QA framework (human testing mode, build tracking)',
    'Corporate governance GOV-001 through GOV-009',
    'Four business components (Find Care, Evaluate Care, Discuss Care, The Brain)',
    'Production release with QA report and business artifacts gate',
]:
    doc.add_paragraph(s, style='List Bullet')

# Planning Model
doc.add_heading('Planning Model', level=3)
for l in [
    'Epic: Strategic capability stream with goal, description, success criteria, definition of done',
    'Feature: Every deliverable is a feature. Each belongs to frontend or backend. Every frontend feature lists backend dependencies.',
    'Story: Sprint-sized unit of work. Same-app dependencies by story ID. Cross-app dependencies by feature ID.',
    'Requirement: Boolean testable statement (true/false). Must trace to epic goal.',
]:
    doc.add_paragraph(l, style='List Bullet')

p = doc.add_paragraph()
r = p.add_run('Feature completeness rule: ')
r.bold = True
p.add_run(
    'A sprint must release entire features. All stories in a feature must land in the same sprint. '
    'A feature may NOT be split across sprints unless Boss explicitly authorizes an exception. '
    'Every sprint must deliver at least one complete feature. No empty sprints, no infrastructure-only sprints.'
)

p = doc.add_paragraph()
r = p.add_run('Requirement labels: ')
r.bold = True
p.add_run('Y (Passed), DEF (Deferred to future release), FAIL (Failed). Every requirement must carry exactly one label. No blanks.')

p = doc.add_paragraph()
r = p.add_run('Scope rule: ')
r.bold = True
p.add_run(
    'The scope of this plan cannot exceed what is defined or implied in this assignment. '
    'No feature, story, or requirement may be added beyond what Boss has authorized or implied. '
    'Scope may be deferred or dropped, never expanded. GPT owns the risk matrix \u2014 including the risk that '
    'the authorized scope is insufficient to deliver the epic. Claude and GPT collaborate on identifying risks, '
    'but GPT has final authority on what enters the risk matrix and the discretion to override Claude.'
)

doc.add_page_break()

# ── EPIC-1 ────────────────────────────────────────────────────
doc.add_heading('EPIC-1: Evaluate Care v0.1.4', level=2)
doc.add_paragraph('Capacity: 80%')

p = doc.add_paragraph()
r = p.add_run('Goal: ')
r.bold = True
p.add_run(
    'Release a credible version 1 of the Evaluate Care business component. '
    'Evaluate Care contains two separate and distinct capabilities: Evaluate Care for Providers '
    'and Evaluate Care for Clinical Trials. These are independent capabilities but may share measures '
    'where applicable. Each capability must deliver a quality score derived from a minimum of 3 credible '
    'measures, with a substantive, explainable answer showing what measures produced that score and how '
    'each measure contributed. If we can deliver 3 credible measures for both providers and clinical trials, '
    'we have an alpha component. Stretch goal: allow the user to adjust the algorithm to determine quality '
    'by adjusting the importance of the different measures that are used to determine quality.'
)

p = doc.add_paragraph()
r = p.add_run("Boss\u2019s non-binding suggestion: ")
r.bold = True
p.add_run(
    'Prescription behavior mapped to diseases. This is the secret AI-enhanced feature \u2014 '
    'the ability to analyze how a provider prescribes relative to the conditions they treat. '
    'This is a differentiator and should be treated as a premium quality measure.'
)

p = doc.add_paragraph()
r = p.add_run("Claude\u2019s non-binding suggestions for measures: ")
r.bold = True
doc.add_paragraph(
    'Provider: board certifications, patient ratings/reviews, malpractice or disciplinary history, '
    'years in practice, specialty match accuracy, acceptance of new patients',
    style='List Bullet'
)
doc.add_paragraph(
    'Clinical trials: trial phase (I/II/III/IV), enrollment status, relevance to user condition, '
    'sponsor credibility, proximity to user, recency of results',
    style='List Bullet'
)
doc.add_paragraph(
    'GPT and Claude are free to propose different measures, replace these, or add to them.',
    style='List Bullet'
)

p = doc.add_paragraph()
r = p.add_run('Success Criteria')
r.bold = True
for c in [
    "The system can tell the user what a provider\u2019s quality is and why we say so",
    "The system can tell the user what a clinical trial\u2019s quality/relevance is and why we say so",
    "Evaluate Care for Providers and Evaluate Care for Clinical Trials are separately functional",
    "Every quality assessment includes a quality score derived from a minimum of 3 distinct measures",
    "Every quality score is coupled with the measures that produced it",
    "Quality assessments are grounded in data (not fabricated by the LLM)",
    "Quality signals are explainable",
    "Every boolean requirement carries exactly one label (Y, DEF, or FAIL) \u2014 no requirement left blank",
    "No regressions from v0.1.3 Find Care baseline",
    "STRETCH: The user can adjust the weight of individual quality measures and the score recalculates",
]:
    doc.add_paragraph(c, style='List Bullet')

doc.add_paragraph('Definition of Done: Boss signs off as sufficient subsequent to UAT.')

doc.add_heading('EPIC-1 Feature Template', level=3)
for f in [
    'EVAL-SCORE-PROV: Provider Quality Scoring Engine (backend, Providers)',
    'EVAL-SCORE-CT: Clinical Trial Quality Scoring Engine (backend, Clinical Trials)',
    'EVAL-EXPLAIN: Score Explainability (backend, Shared)',
    'EVAL-EXPLAIN-UX: Score Explainability Display (frontend, Shared) \u2014 depends on EVAL-SCORE-PROV, EVAL-SCORE-CT, EVAL-EXPLAIN',
    'EVAL-P-MEASURE-1/2/3: Provider measures (GPT defines)',
    "EVAL-P-RX: Prescription Behavior Analysis (Boss\u2019s non-binding suggestion)",
    'EVAL-CT-MEASURE-1/2/3: Clinical trial measures (GPT defines)',
    'EVAL-WEIGHTS: User-Adjustable Quality Weights (STRETCH)',
]:
    doc.add_paragraph(f, style='List Bullet')

doc.add_page_break()

# ── EPIC-2 ────────────────────────────────────────────────────
doc.add_heading('EPIC-2: Maintenance', level=2)
doc.add_paragraph('Capacity: 20%')

p = doc.add_paragraph()
r = p.add_run('Goal: ')
r.bold = True
p.add_run(
    'Find Care enhancements (maps, distance, UX), architecture hardening, Brain enhancements, '
    'and operational hardening from the Perplexity external audit.'
)

p = doc.add_paragraph()
r = p.add_run('Success Criteria')
r.bold = True
for c in [
    'Provider search results render on an interactive Google Map with clickable markers',
    "Distance from user\u2019s location displayed wherever a provider address is shown",
    'Every function that presents an address to the user includes a map and distance',
    'Architecture refactor complete \u2014 monolith main.py replaced by domain services',
    'All HIGH-priority Perplexity audit items addressed or have documented deferral rationale',
    'Brand audit complete \u2014 ChatHealthy.ai used consistently everywhere',
    'Every boolean requirement carries exactly one label (Y, DEF, or FAIL)',
    'No regressions',
]:
    doc.add_paragraph(c, style='List Bullet')

doc.add_paragraph('Definition of Done: Boss signs off as sufficient subsequent to UAT.')

doc.add_heading('EPIC-2 Features', level=3)
for f in [
    'MAINT-MAPS: Maps & Mapping (frontend, HIGH)',
    'MAINT-DIST: Distance Propagation (backend, HIGH)',
    'MAINT-DIST-UX: Distance Display (frontend, HIGH) \u2014 depends on MAINT-DIST',
    'MAINT-ARCH: ARCH-001 Architecture Refactor (backend, HIGH)',
    'MAINT-F09: Emergency Keyword Expansion (backend, MEDIUM)',
    'MAINT-F10: Provider Search Cap (backend, LOW)',
    'MAINT-CONSENT: Consent Framework Bug Fixes (backend, HIGH)',
    'MAINT-TIMEOUT: UX Timeout Handling (frontend, MEDIUM)',
    'MAINT-BRAIN: Brain Enhancements (backend, GPT to recommend)',
    'BL-0.1.4-001: Server-Side Scope Enforcement (backend, HIGH)',
    'BL-0.1.4-002: Bearer Token Lifecycle (backend, HIGH)',
    'BL-0.1.4-003: Brain Lock TTL and Deadlock Recovery (backend, HIGH)',
    'BL-0.1.4-004: Independent Append-Only Audit Logging (backend, HIGH)',
    'BL-0.1.4-005: Local Disk Resilience (backend, MEDIUM)',
    'BL-0.1.4-006: Manifest Snapshot Freshness (backend, MEDIUM)',
    'BL-0.1.4-007: Health Check Endpoint (backend, MEDIUM)',
    'BL-0.1.4-008: Rate Limiting (backend, MEDIUM)',
    'BL-0.1.4-009: Mongo Query Size Estimation (backend, LOW)',
    'BL-0.1.4-010: Multi-Cluster Connection String Governance (backend, MEDIUM)',
    'BRAND-001: Brand Name Audit (both, HIGH)',
]:
    doc.add_paragraph(f, style='List Bullet')

doc.add_paragraph('Not all items will fit in 20% \u2014 prioritize by risk, tell Boss what defers to v0.1.5.')

doc.add_page_break()

# ── Sprints ───────────────────────────────────────────────────
doc.add_heading('Sprints', level=2)
for s in [
    'Sprint 1: Apr 1-5 \u2014 Development',
    'Sprint 2: Apr 7-12 \u2014 Development',
    'Sprint 3: Apr 14-19 \u2014 Development',
    'Sprint 4: Apr 21-26 \u2014 Development',
    'Sprint 5: Apr 26 - May 1 \u2014 UAT and release. No new development. Bug fixes only. Boss sign-off (GOV-007).',
]:
    doc.add_paragraph(s, style='List Bullet')

# ── Assignment DoD ────────────────────────────────────────────
doc.add_heading('Assignment Definition of Done', level=2)

p = doc.add_paragraph()
r = p.add_run('Schedule clarity: ')
r.bold = True
p.add_run('Every story assigned to a sprint. Pipeline execution time accounted for. No overload.')

p = doc.add_paragraph()
r = p.add_run('Traceability: ')
r.bold = True
p.add_run('Features have stories. Stories have requirements. Requirements trace to epic goal. Same-app deps by story ID. Cross-app deps by feature ID.')

p = doc.add_paragraph()
r = p.add_run('Completeness: ')
r.bold = True
p.add_run('All sections populated. Every frontend feature lists backend dependencies. Every sprint delivers a complete feature. Risk and gate sections present.')

# ── Deliverables ──────────────────────────────────────────────
doc.add_heading('Deliverables', level=2)

doc.add_paragraph(
    'This assignment produces two artifacts that must be consistent with each other in all content:'
)
doc.add_paragraph('1. Business document (Word format) \u2014 human-readable epic plan', style='List Bullet')
doc.add_paragraph('2. JSON \u2014 machine-readable epic plan matching the Assurance Output schema', style='List Bullet')
doc.add_paragraph(
    'GPT produces both. Claude validates compliance between the two. '
    'Neither is complete until both GPT and Claude agree they are good and consistent.'
)

doc.add_paragraph('')
p = doc.add_paragraph()
r = p.add_run('Iteration cache: ')
r.bold = True
p.add_run(
    'brain/machine_artifacts/.iteration_cache/ (committed to repo \u2014 GPT accesses via GPTReader). '
    'GPT shall review the last iteration as a starting place. '
    'If GPT recommends non-trivial changes, GPT must flag that to Claude.'
)

doc.add_paragraph('')
for d in [
    '1. Epic Definitions \u2014 validate or revise goals, criteria, definition of done',
    '2. Features Per Epic \u2014 GPT and Claude free to recommend different feature set',
    '3. Stories Per Feature \u2014 title, description, parent, sprint, dependencies, size',
    '4. Boolean Requirements Per Story \u2014 traces to epic goal, labels Y/DEF/FAIL',
    '5. Sprint Capability Map \u2014 per sprint: capability, stories, goal, capacity',
    '6. Risk Matrix \u2014 per feature (both epics) + per measure (EPIC-1)',
    '7. Gate Recommendation \u2014 overall risk, recommendation, issues',
]:
    doc.add_paragraph(d, style='List Bullet')

doc.add_paragraph('')
p = doc.add_paragraph()
r = p.add_run('Iteration discipline: ')
r.bold = True
p.add_run(
    'Work slowly and methodically. Keep track of where you are in the iteration path \u2014 '
    'the iteration counter is in the assignment header. You have up to 1000 iterations to produce '
    'an agreed plan. Do not rush. Each iteration should improve on the last. Claude and GPT iterate '
    'until both agree the plan is complete, consistent, and meets the Assignment Definition of Done. '
    'Claude increments the iteration counter on each submission.'
)

doc.add_paragraph('')
p = doc.add_paragraph()
r = p.add_run('Output: ')
r.bold = True
p.add_run('Return JSON matching the Assurance Output schema. Include usage block.')

# Save
out = 'brain/BusinessArtifacts/epic_planning_prompt_v014.docx'
doc.save(out)
print(f"Saved: {out}")
