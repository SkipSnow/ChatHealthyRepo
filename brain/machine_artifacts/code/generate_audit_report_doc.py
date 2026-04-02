# Copyright (c) 2026 Skip Snow. All rights reserved.
# Licensed under the FindCare Evaluation License (FEL-1.0).
#
# generate_audit_report_doc.py - Convert audit HTML to Word doc

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# Copyright
p = doc.add_paragraph("Copyright 2026 Skip Snow")
p.runs[0].font.size = Pt(9)

# Title
doc.add_heading("External Audit Report: Epic Plan Evaluate Care v0.1.4", level=1)

# Meta table
t = doc.add_table(rows=7, cols=2)
t.style = "Table Grid"
meta = [
    ("Audit Date", "April 1, 2026"),
    ("Audit Authority", "External audit requested by Skip Snow to assess the Evaluate Care epic plan against the ChatHealthy.ai Operational Business Architecture and to assess end-to-end traceability viability from Epic to Feature to Story to Test Case."),
    ("Audit Type", "Independent external audit of business-architecture compliance and traceability-model viability."),
    ("Prepared By", "Perplexity, powered by GPT-5.4."),
    ("Model Date", "April 1, 2026."),
    ("Source Artifacts", "epic_plan_evaluate_care_v014.docx; ChatHealthyOperationalBusinessArchitecture.docx."),
    ("Audited Subject", "Evaluate Care epic plan v0.1.4, including feature/stories/requirements/testability model and its alignment to the stated business architecture for Evaluate Care."),
]
for i, (k, v) in enumerate(meta):
    t.rows[i].cells[0].text = k
    t.rows[i].cells[1].text = v
    t.rows[i].cells[0].paragraphs[0].runs[0].bold = True

# Audit Scope
doc.add_heading("Audit Scope", level=2)
doc.add_paragraph(
    "This audit addresses two questions only: first, whether the Evaluate Care epic plan is compliant "
    "with the ChatHealthy.ai Operational Business Architecture; second, whether the traceability model "
    "running from Epic to Feature to Story to Test Case will work as specified."
)
doc.add_paragraph(
    "The audit is limited to the two source documents provided and does not validate code, runtime "
    "behavior, repository state, or undocumented assumptions outside those artifacts."
)

# Exec Summary
doc.add_heading("Executive Summary", level=2)
doc.add_paragraph(
    "The Evaluate Care epic plan is directionally aligned with the business architecture in that it "
    "clearly targets provider evaluation, clinical-trial evaluation, explainability, and user-adjustable "
    "weighting, all of which are consistent with the architecture\u2019s stated Evaluate Care capabilities "
    "and positioning."
)
doc.add_paragraph(
    "However, the plan is not fully compliant as written because it materially narrows or distorts some "
    "architecture commitments, especially around comparative provider evaluation, clinical-trials discovery "
    "and matching, premium-tier positioning, and the architecture\u2019s patent-positioned claim that the "
    "consumer tunes the algorithm."
)
doc.add_paragraph(
    "The traceability model is partially workable but not fully reliable as specified because the document "
    "demonstrates strong Feature-to-Story-to-Requirement decomposition while leaving material ambiguity "
    "around true Epic-to-Feature coverage, uniqueness of requirement IDs, and the final handoff from "
    "requirements to executable test cases."
)

# Source Summary
doc.add_heading("Source Summary", level=2)
doc.add_paragraph(
    "The business architecture defines Evaluate Care as a component that transforms fragmented public "
    "and clinical data into clear, consumer-comprehensible quality signals using proprietary AI synthesis, "
    "with capabilities that include clinical trials discovery and matching, provider credential analysis, "
    "quality signal aggregation and interpretation, and comparative provider evaluation."
)
doc.add_paragraph(
    "The epic plan defines Evaluate Care v0.1.4 as a 22-feature plan with 96 stories and 1459 boolean "
    "requirements, centered on provider scoring, clinical-trial scoring, explainability, score display, "
    "user-adjustable weights, measure extraction, caching, confidence indicators, and failsafes."
)

# Strengths
doc.add_heading("Strengths / What Works Well", level=2)
doc.add_paragraph(
    "The epic plan strongly reflects the architecture\u2019s core proposition that Evaluate Care should "
    "generate explainable quality signals rather than opaque outputs. The plan repeatedly emphasizes "
    "deterministic scoring, structured explainability, provenance, measure breakdowns, confidence "
    "indicators, and minimum-measure failsafes, all of which support consumer-comprehensible evaluation."
)
doc.add_paragraph(
    "The plan also shows unusually strong internal decomposition below the feature level. Features break "
    "into stories, stories break into boolean requirements, and many requirements are written in testable "
    "\u201cthe system must\u201d form, which is stronger than typical planning artifacts and gives a "
    "plausible base for later verification."
)

# Findings
doc.add_heading("Findings, Risks, and Weaknesses", level=2)

doc.add_heading("Architecture compliance", level=3)

findings_arch = [
    ("F1", "The epic plan does not demonstrate full compliance with the architecture\u2019s \u201cclinical trials discovery and matching\u201d capability because the plan focuses heavily on scoring and ranking returned trials, but the architecture explicitly frames Evaluate Care as including discovery and matching, not only post-retrieval scoring.",
     "A scoring engine can rank trials after retrieval, but that is not the same as proving that Evaluate Care itself delivers the architecture\u2019s stated discovery-and-matching capability, especially if matching logic remains primarily in upstream search tooling rather than in the Evaluate Care component."),
    ("F2", "The epic plan does not clearly fulfill the architecture\u2019s \u201ccomparative provider evaluation\u201d capability because the visible plan language centers on producing a provider score and explanation per entity, while the architecture explicitly calls out comparative evaluation as a named capability.",
     "Single-entity scoring is not equivalent to comparative evaluation unless the plan explicitly includes comparison workflows, comparison outputs, or side-by-side evaluation contracts across multiple providers."),
    ("F3", "The plan only partially aligns with the architecture\u2019s statement that the consumer tunes the algorithm. The architecture presents this as part of the Evaluate Care concept and even links it to the patent-positioned visualization, while the epic plan treats user-adjustable weights as a stretch feature rather than a core compliance item.",
     "Where the architecture describes consumer tuning as part of the core product conception, demoting it to stretch status creates a compliance gap unless the architecture itself is revised or that capability is explicitly deferred in a way the architecture recognizes."),
    ("F4", "The plan narrows the architecture\u2019s provider-analysis language toward deterministic measure scoring without clearly proving full \u201cprovider credential analysis\u201d breadth. The architecture frames Evaluate Care as synthesizing fragmented public and clinical data into quality signals, but the plan evidence shown in the artifact is heavily implementation-centric and does not clearly demonstrate coverage of the full business-level provider-credential-analysis promise.",
     "Architecture compliance is about capability coverage, not just technical scoring mechanics, and the current plan reads more like a scoring-engine plan than a full provider-analysis capability plan."),
    ("F5", "The plan does not show clear compliance with the architecture\u2019s pricing position for Evaluate Care. The architecture states that Evaluate Care is free initially with premium paid tiers for advanced insights and deeper analysis, but the epic plan extract does not show explicit packaging or traceability between base scoring features and premium-insight features.",
     "Where pricing tiering is part of the business architecture, the epic plan should at minimum identify which features support free baseline capability and which support premium insight depth, otherwise business-architecture traceability is incomplete."),
]

for fid, finding, why in findings_arch:
    p = doc.add_paragraph()
    r = p.add_run(f"{fid}. ")
    r.bold = True
    p.add_run(finding)
    p2 = doc.add_paragraph()
    r2 = p2.add_run("Why this matters: ")
    r2.italic = True
    p2.add_run(why)

doc.add_heading("Traceability and testability", level=3)

findings_trace = [
    ("F6", "The traceability chain is not fully reliable because requirement identifiers are visibly reused across stories and features in ways that can break unambiguous downstream linking. For example, multiple stories use repeated local IDs such as EVAL-REQ-001, EVAL-REQ-002, and similar patterns, rather than a globally unique requirement key across the artifact.",
     "If requirement IDs are not globally unique, test cases, defects, coverage reports, and release gates can mis-reference or overwrite each other unless every reference also carries feature and story context in a disciplined way."),
    ("F7", "The document demonstrates Epic-to-Feature and Feature-to-Story structure, but the final \u201cStory to Test Case\u201d link is not adequately specified in the visible artifact. The plan contains many requirements and test-oriented statements, yet the extract does not show a formal test-case object model, stable test-case IDs, or a true traceability matrix proving one-to-many mapping from requirement to test case.",
     "Boolean requirements are testable in concept, but the audit question is whether the matrix will work as specified, and without a formalized test-case layer the chain remains incomplete at execution time."),
    ("F8", "The plan mixes implementation assumptions and evidence gaps into the same planning structure, which weakens traceability reliability. For example, the explainability feature includes explicit GAP and ASSUMPTION language, showing that some features are being planned against missing underlying artifacts rather than against fully grounded source objects.",
     "Traceability matrices work best when every downstream item traces to stable upstream evidence; once assumptions and gaps are mixed into the same artifact without a strict classification model, downstream testability becomes more interpretive than deterministic."),
    ("F9", "The plan\u2019s requirement volume creates a manageability risk for the traceability matrix. With 1459 requirements across 96 stories and 22 features, the matrix may be technically possible but operationally fragile unless there is automation, global uniqueness, and disciplined cross-referencing between requirement IDs and test-case IDs.",
     "Scale by itself is not a defect, but at this volume any weakness in ID hygiene, matrix generation, or test-case normalization becomes a material execution risk."),
    ("F10", "The traceability chain is weakened by inconsistent abstraction levels across requirements. Some requirements are concrete and directly testable, while others are service-threshold, interpretability, or downstream-LLM-consumption statements that are harder to verify without separately defined test methods and acceptance thresholds.",
     "A matrix only \u201cworks\u201d if downstream tests can objectively validate the upstream statement; vague or mixed-abstraction requirements reduce that reliability even when phrased as \u201cmust.\u201d"),
    ("F11", "The plan does not yet prove complete architecture-to-plan traceability for all Evaluate Care business capabilities. While many plan features clearly support quality scoring and explainability, the mapping back to all four stated architecture capabilities is not explicitly rendered as an architecture compliance matrix.",
     "Without a direct architecture-to-feature matrix, compliance is inferred rather than demonstrated, which is weaker than a professional planning artifact should allow for a governed release."),
]

for fid, finding, why in findings_trace:
    p = doc.add_paragraph()
    r = p.add_run(f"{fid}. ")
    r.bold = True
    p.add_run(finding)
    p2 = doc.add_paragraph()
    r2 = p2.add_run("Why this matters: ")
    r2.italic = True
    p2.add_run(why)

# Audit Judgment
doc.add_heading("Audit Judgment", level=2)
t = doc.add_table(rows=3, cols=3)
t.style = "Table Grid"
t.rows[0].cells[0].text = "Question"
t.rows[0].cells[1].text = "Judgment"
t.rows[0].cells[2].text = "Basis"
for cell in t.rows[0].cells:
    cell.paragraphs[0].runs[0].bold = True

t.rows[1].cells[0].text = "1. Is the epic plan compliant with the business architecture?"
t.rows[1].cells[1].text = "Partially compliant, not fully compliant."
t.rows[1].cells[2].text = "The plan strongly supports explainable quality scoring and shared Evaluate Care themes, but it does not clearly prove full coverage of clinical-trials discovery and matching, comparative provider evaluation, consumer algorithm tuning as a core capability, premium-tier packaging, and full business-capability coverage."

t.rows[2].cells[0].text = "2. Will the traceability matrix from Epic to Feature to Story to Test Case work as specified?"
t.rows[2].cells[1].text = "Partially workable, not reliable enough as written."
t.rows[2].cells[2].text = "The plan shows strong decomposition through features, stories, and boolean requirements, but the visible artifact does not fully establish globally unique IDs, a formal test-case layer, or a complete end-to-end matrix proving stable execution traceability."

# Recommendations
doc.add_heading("Recommended Remediation", level=2)
recs = [
    ("R1 (addresses F1, F11)", "Add an explicit architecture compliance matrix mapping each Evaluate Care business-architecture capability to one or more epic features, stories, and tests."),
    ("R2 (addresses F2)", "Add one or more explicit comparative-provider-evaluation features or stories with defined outputs, such as side-by-side comparison views or ranked comparative explanation payloads."),
    ("R3 (addresses F3)", "Decide whether consumer-adjustable weighting is a required architecture capability or a deferrable enhancement, and make the architecture and epic plan say the same thing."),
    ("R4 (addresses F5)", "Add feature-level or story-level traceability to pricing tiers, distinguishing free baseline evaluation from premium advanced insight functionality."),
    ("R5 (addresses F6, F9)", "Enforce globally unique requirement IDs and globally unique future test-case IDs across the full artifact."),
    ("R6 (addresses F7, F10)", "Add a formal test-case model with test-case IDs, requirement-to-test mappings, and explicit verification methods for every requirement type."),
    ("R7 (addresses F8)", "Separate grounded evidence, assumptions, and gaps into distinct fields so traceability does not blur implemented source evidence with planning conjecture."),
]
for rid, rec in recs:
    p = doc.add_paragraph()
    r = p.add_run(f"{rid}. ")
    r.bold = True
    p.add_run(rec)

# Final Opinion
doc.add_heading("Final Opinion", level=2)
doc.add_paragraph(
    "The Evaluate Care epic plan is strong as an engineering-planning artifact, but it is not yet fully "
    "sufficient as a governed business-compliance artifact against the Operational Business Architecture."
)
doc.add_paragraph(
    "Its decomposition depth is a genuine strength, yet the plan should not be treated as "
    "architecture-compliant or fully traceability-safe until the architecture capability mapping, "
    "requirement-ID discipline, and true requirement-to-test-case traceability model are made explicit "
    "and unambiguous."
)

doc.save("brain/BusinessArtifacts/externalAudits/evaluate_care_epic_plan_external_audit.docx")
print("Saved: brain/BusinessArtifacts/externalAudits/evaluate_care_epic_plan_external_audit.docx")
