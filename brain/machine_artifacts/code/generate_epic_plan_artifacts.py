# Copyright (c) 2026 Skip Snow. All rights reserved.
# Licensed under the FindCare Evaluation License (FEL-1.0).
#
# generate_epic_plan_artifacts.py - Produces Word doc and JSON from plan_tree.json
# Hierarchy: Epic (flush left) -> Feature (0.25in) -> Story (0.5in) -> Requirement (0.75in)
# Tables have repeating headers. Every level is labeled.

import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

_REPO = Path(__file__).parent.parent.parent.parent
_TREE = _REPO / "brain" / "machine_artifacts" / ".iteration_cache" / "plan_tree.json"
_DOC_OUT = _REPO / "brain" / "BusinessArtifacts" / "epic_plan_evaluate_care_v014.docx"
_JSON_OUT = _REPO / "brain" / "machine_artifacts" / ".iteration_cache" / "epic_plan_v014_final.json"

with open(_TREE, encoding="utf-8") as f:
    tree = json.load(f)

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)

total_f = sum(len(e.get("features", [])) for e in tree["epics"])
total_s = sum(len(f.get("stories", [])) for e in tree["epics"] for f in e.get("features", []))
total_r = sum(len(s.get("requirements", [])) for e in tree["epics"] for f in e.get("features", []) for s in f.get("stories", []))


def set_repeat_header(table):
    """Make the first row of a table repeat on every page."""
    try:
        row = table.rows[0]
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        tblHeader = OxmlElement("w:tblHeader")
        trPr.append(tblHeader)
    except Exception:
        pass


def add_indented_heading(text, level, indent_inches=0):
    """Add a heading with left indent."""
    p = doc.add_heading(text, level=level)
    if indent_inches > 0:
        p.paragraph_format.left_indent = Inches(indent_inches)
    return p


def add_indented_para(text, indent_inches=0, bold=False, size=None, color=None):
    """Add a paragraph with left indent."""
    p = doc.add_paragraph()
    if indent_inches > 0:
        p.paragraph_format.left_indent = Inches(indent_inches)
    r = p.add_run(text)
    if bold:
        r.bold = True
    if size:
        r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor(*color)
    return p


# ── Title ──
doc.add_paragraph("")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("ChatHealthy.ai")
r.font.size = Pt(28)
r.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Epic Plan: Evaluate Care v0.1.4")
r.font.size = Pt(18)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("April 2, 2026 | Build 362").font.size = Pt(11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(f"{total_f} Features | {total_s} Stories | {total_r} Requirements")
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(11, 122, 117)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Copyright 2026 Skip Snow. All rights reserved.").font.size = Pt(9)

doc.add_page_break()

# ── Table of Contents ──
doc.add_heading("Table of Contents", level=1)

for epic in tree["epics"]:
    features = epic.get("features", [])
    if not features:
        continue
    p = doc.add_paragraph()
    r = p.add_run(f"Epic: {epic['epic_id']}: {epic['name']}")
    r.bold = True
    r.font.size = Pt(11)

    for feat in features:
        fid = feat.get("feature_id", "?")
        fname = feat.get("name", "?")
        stories = feat.get("stories", [])
        feat_reqs = sum(len(s.get("requirements", [])) for s in stories)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        run = p.add_run(f"Feature: {fid}: {fname} ({len(stories)} stories, {feat_reqs} reqs)")
        run.font.color.rgb = RGBColor(11, 122, 117)
        run.font.size = Pt(9)

doc.add_page_break()

# ── How to Use This Document ──
doc.add_heading("How to Use This Document", level=1)
doc.add_paragraph(
    "This document follows a strict hierarchy. Each level is labeled and indented:"
)
doc.add_paragraph("Epic (flush left) -- the strategic capability stream", style="List Bullet")
doc.add_paragraph("Feature (indented 0.25in) -- a deliverable within an epic", style="List Bullet")
doc.add_paragraph("Story (indented 0.5in) -- a sprint-sized unit of work within a feature", style="List Bullet")
doc.add_paragraph("Requirement (indented 0.75in, in tables) -- a boolean testable statement", style="List Bullet")

doc.add_paragraph(
    "Every header is prefixed with its level: 'Epic:', 'Feature:', 'Story:', 'Requirement:'. "
    "Tables have headers that repeat at the top of every page. "
    "To find a specific feature, use the Table of Contents. "
    "To review test cases for a story, find the story heading and read the requirement table below it."
)
doc.add_paragraph(
    "Requirement labels: Y (Passed), DEF (Deferred), FAIL (Failed). "
    "Labels are blank until UAT. Every requirement must carry exactly one label at release."
)

doc.add_page_break()

# ── Exec Summary ──
doc.add_heading("Executive Summary", level=1)
doc.add_paragraph(
    f"This document defines the epic plan for Evaluate Care v0.1.4. "
    f"The plan contains {total_f} features, {total_s} stories, and {total_r} boolean "
    f"testable requirements. Every feature was proposed by GPT (Enterprise Architect, gpt-5.3) "
    f"and accepted by Claude (Accountable Dev Manager, claude-opus-4-6). "
    f"Every requirement traces to the epic goal."
)
doc.add_paragraph(f"Gate: {tree.get('gate_recommendation', 'N/A')}")

# ── Feature Summary Table ──
doc.add_heading("Feature Summary", level=1)

for epic in tree["epics"]:
    features = epic.get("features", [])
    if not features:
        continue
    doc.add_heading(f"Epic: {epic['epic_id']}: {epic['name']}", level=2)
    t = doc.add_table(rows=1, cols=4)
    t.style = "Table Grid"
    for i, h in enumerate(["Feature ID", "Name", "Layer", "Description"]):
        t.rows[0].cells[i].text = h
        t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    set_repeat_header(t)
    for feat in features:
        row = t.add_row()
        row.cells[0].text = feat.get("feature_id", "?")
        row.cells[1].text = feat.get("name", "?")
        row.cells[2].text = feat.get("layer", "?")
        desc = feat.get("description", "")
        row.cells[3].text = desc[:120] + ("..." if len(desc) > 120 else "")

doc.add_page_break()

# ── Software Design ──
doc.add_heading("Software Design", level=1)

doc.add_heading("Data Flow", level=2)
flow = [
    "1. User asks about provider quality or clinical trial quality",
    "2. Claude Sonnet routes to evaluate_provider_quality() or evaluate_trial_quality()",
    "3. EvaluateCareFacade collects measures from provider/trial services",
    "4. Each measure normalized via MeasureNormalizationFramework (0-1)",
    "5. ScoringEngine computes weighted composite score (deterministic)",
    "6. ScoreExplainabilityService generates measure breakdown + provenance",
    "7. ConfidenceIndicator scores data completeness",
    "8. MeasureFailsafe validates minimum 3 measures; degrades if insufficient",
    "9. QualityScoreCache stores result",
    "10. EVAL-EXPLAIN-UX renders score card in chat",
]
for step in flow:
    doc.add_paragraph(step)

doc.add_heading("AI Calls", level=2)
t = doc.add_table(rows=1, cols=3)
t.style = "Table Grid"
for i, h in enumerate(["Model", "Purpose", "Invocation"]):
    t.rows[0].cells[i].text = h
    t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
set_repeat_header(t)
for model, purpose, inv in [
    ("Claude Sonnet 4.6", "Conversation + tool calling", "Per message"),
    ("Claude Haiku 4.5", "Query expansion, summarization, de-ID", "Per tool call"),
    ("GPT-4.1-mini", "Safety classification", "Per message"),
    ("text-embedding-3-large", "Provider vector search", "Per search"),
    ("text-embedding-3-small", "Specialty vector search", "Per specialty"),
    ("Google Routes API", "Travel distance for trials", "Per trial location"),
]:
    row = t.add_row()
    row.cells[0].text = model
    row.cells[1].text = purpose
    row.cells[2].text = inv

doc.add_heading("Object Model", level=2)
objects = [
    ("EvaluateCareFacade", "application/facades/",
     "evaluate_provider_quality(), evaluate_trial_quality(), get_score_explanation()",
     "Orchestrates scoring, measures, explainability"),
    ("ProviderScoringEngine", "domain/evaluate_care_quality/",
     "compute_score(), normalize(), apply_weights()",
     "Deterministic composite from provider measures"),
    ("ClinicalTrialScoringEngine", "domain/evaluate_care_quality/",
     "compute_score(), normalize(), apply_weights()",
     "Deterministic composite from trial measures"),
    ("ScoreExplainabilityService", "domain/evaluate_care_quality/",
     "explain(), format_breakdown(), trace_to_data()",
     "Measure values, contributions, provenance"),
    ("MeasureNormalizationFramework", "domain/evaluate_care_quality/",
     "normalize(), scale_to_range(), handle_missing()",
     "Consistent 0-1 normalization"),
    ("QualityScoreCache", "infrastructure/",
     "get(), put(), invalidate(), ttl_check()",
     "Read-through/write-through cache"),
    ("MeasureFailsafe", "domain/evaluate_care_quality/",
     "validate_minimum(), degrade_gracefully(), flag_insufficient()",
     "Minimum measure count enforcement"),
    ("ConfidenceIndicator", "domain/evaluate_care_quality/",
     "compute_confidence(), explain_confidence()",
     "Confidence based on data completeness"),
    ("DataProvenanceTracker", "infrastructure/",
     "attach(), trace(), audit_trail()",
     "Source, freshness, lineage per measure"),
]
t = doc.add_table(rows=1, cols=4)
t.style = "Table Grid"
for i, h in enumerate(["Class", "Location", "Methods", "Description"]):
    t.rows[0].cells[i].text = h
    t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
set_repeat_header(t)
for cls, loc, methods, desc in objects:
    row = t.add_row()
    row.cells[0].text = cls
    row.cells[1].text = loc
    row.cells[2].text = methods
    row.cells[3].text = desc

doc.add_page_break()

# ── Features / Stories / Requirements (hierarchical, indented) ──
for epic in tree["epics"]:
    features = epic.get("features", [])
    if not features:
        continue

    # Epic level — flush left
    doc.add_heading(f"Epic: {epic['epic_id']}: {epic['name']}", level=1)

    for feat in features:
        stories = feat.get("stories", [])
        feat_reqs = sum(len(s.get("requirements", [])) for s in stories)

        # Feature level — 0.25in indent
        add_indented_heading(
            f"Feature: {feat['feature_id']}: {feat['name']}",
            level=2, indent_inches=0.25
        )

        # Feature metadata table
        t = doc.add_table(rows=6, cols=2)
        t.style = "Table Grid"
        for i, (k, v) in enumerate([
            ("Layer", feat.get("layer", "?")),
            ("Priority", feat.get("priority", "?")),
            ("Evidence", str(feat.get("evidence", ""))[:200]),
            ("Accepted by", feat.get("accepted_by", "?")),
            ("Stories", str(len(stories))),
            ("Requirements", str(feat_reqs)),
        ]):
            t.rows[i].cells[0].text = k
            t.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            t.rows[i].cells[1].text = v

        for story in stories:
            reqs = story.get("requirements", [])

            # Story level — 0.5in indent
            add_indented_heading(
                f"Story: {story['story_id']}: {story['title']}",
                level=3, indent_inches=0.5
            )
            add_indented_para(story.get("description", ""), indent_inches=0.5)
            add_indented_para(
                f"Size: {story.get('size', '?')} | Sprint: {story.get('sprint', '?')} | Reqs: {len(reqs)}",
                indent_inches=0.5, size=9
            )
            if story.get("evidence"):
                add_indented_para(
                    f"Evidence: {story['evidence'][:150]}",
                    indent_inches=0.5, size=9
                )
            if story.get("dependencies"):
                add_indented_para(
                    f"Dependencies: {', '.join(story['dependencies'])}",
                    indent_inches=0.5, size=9
                )

            # Requirements table — 0.75in indent (via table position)
            if reqs:
                add_indented_para("Requirements:", indent_inches=0.75, bold=True, size=9)
                rt = doc.add_table(rows=1, cols=4)
                rt.style = "Table Grid"
                for i, h in enumerate(["Requirement ID", "Requirement", "Priority", "Status"]):
                    rt.rows[0].cells[i].text = h
                    rt.rows[0].cells[i].paragraphs[0].runs[0].bold = True
                set_repeat_header(rt)
                for req in reqs:
                    row = rt.add_row()
                    row.cells[0].text = req.get("req_id", "?")
                    row.cells[1].text = req.get("requirement", "?")[:150]
                    row.cells[2].text = req.get("priority", "?")
                    row.cells[3].text = req.get("status", "?")

        doc.add_page_break()

# ── Sprint Map ──
doc.add_heading("Sprint Capability Map", level=1)
for sprint in tree.get("sprint_map", []):
    if not isinstance(sprint, dict):
        continue
    doc.add_heading(f"Sprint {sprint.get('sprint', '?')}: {sprint.get('sprint_goal', '')}", level=2)
    doc.add_paragraph(f"Dates: {sprint.get('dates', '?')}")
    doc.add_paragraph(f"Capacity: {sprint.get('capacity_split', '?')}")
    shipped = sprint.get("features_shipped", [])
    if shipped:
        doc.add_paragraph(f"Features ({len(shipped)}): {', '.join(shipped)}")
    notes = sprint.get("notes", "")
    if notes:
        doc.add_paragraph(f"Notes: {notes}")

# ── Risk Matrix ──
doc.add_heading("Risk Matrix", level=1)
rm = tree.get("risk_matrix", {})
if isinstance(rm, dict):
    pf = rm.get("per_feature", [])
    if pf:
        doc.add_heading("Per Feature", level=2)
        t = doc.add_table(rows=1, cols=4)
        t.style = "Table Grid"
        for i, h in enumerate(["Feature", "Likelihood", "Impact", "Mitigation"]):
            t.rows[0].cells[i].text = h
            t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_repeat_header(t)
        for e in pf:
            if isinstance(e, dict):
                row = t.add_row()
                row.cells[0].text = e.get("feature_id", "?")
                row.cells[1].text = e.get("likelihood", "?")
                row.cells[2].text = e.get("impact", "?")
                row.cells[3].text = e.get("mitigation", "")[:100]

    pm = rm.get("per_measure", [])
    if pm:
        doc.add_heading("Per Measure", level=2)
        t = doc.add_table(rows=1, cols=5)
        t.style = "Table Grid"
        for i, h in enumerate(["Measure", "Availability", "Quality", "Credibility", "Mitigation"]):
            t.rows[0].cells[i].text = h
            t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_repeat_header(t)
        for e in pm:
            if isinstance(e, dict):
                row = t.add_row()
                row.cells[0].text = e.get("feature_id", "?")
                row.cells[1].text = e.get("data_source_availability", "?")
                row.cells[2].text = e.get("data_quality", "?")
                row.cells[3].text = e.get("credibility", "?")
                row.cells[4].text = e.get("mitigation", "")[:80]

# ── Rejected Candidates ──
rejected = tree.get("rejected_measure_candidates", [])
if rejected:
    doc.add_heading("Rejected Measure Candidates", level=1)
    doc.add_paragraph(
        "The following measures were evaluated and rejected by Boss. "
        "Preserved for institutional knowledge."
    )
    t = doc.add_table(rows=1, cols=4)
    t.style = "Table Grid"
    for i, h in enumerate(["ID", "Name", "Disposition", "Reason"]):
        t.rows[0].cells[i].text = h
        t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    set_repeat_header(t)
    for r in rejected:
        row = t.add_row()
        row.cells[0].text = r.get("id", "?")
        row.cells[1].text = r.get("name", "?")
        row.cells[2].text = r.get("disposition", "?")
        row.cells[3].text = r.get("reason_dropped", "")[:120]

# ── Gate ──
doc.add_heading("Gate Recommendation", level=1)
doc.add_paragraph(f"Gate: {tree.get('gate_recommendation', 'N/A')}")

# Save
doc.save(str(_DOC_OUT))
print(f"Word: {_DOC_OUT}")

# JSON
with open(_JSON_OUT, "w", encoding="utf-8") as f:
    json.dump(tree, f, indent=2, ensure_ascii=False)
print(f"JSON: {_JSON_OUT}")

print(f"\nTotals: {total_f} features, {total_s} stories, {total_r} requirements")
