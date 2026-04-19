"""Generate PromptFactoryDesign_V5.docx from V4 markdown with UML diagrams."""
import io
import os
import plantuml
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

REPO_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", ".."))
AUDIT_DIR = os.path.join(REPO_ROOT, "brain", "BusinessArtifacts", "Audits")


def render_uml(puml_text):
    """Render PlantUML text to PNG bytes via public server."""
    p = plantuml.PlantUML(url="https://www.plantuml.com/plantuml/img/")
    result = p.processes(puml_text)
    if result and len(result) > 100:
        return result
    return None


def add_table(doc, headers, rows, style="Table Grid"):
    """Add a formatted Word table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers), style=style)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph("")  # spacer


def add_uml_diagram(doc, puml_text, caption, width=Inches(6)):
    """Render and embed a UML diagram."""
    print(f"  Rendering: {caption}...")
    try:
        img_bytes = render_uml(puml_text)
        if img_bytes:
            stream = io.BytesIO(img_bytes)
            doc.add_picture(stream, width=width)
            last_para = doc.paragraphs[-1]
            last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap = doc.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.runs[0].italic = True if cap.runs else None
            print(f"    OK ({len(img_bytes)} bytes)")
            return True
        else:
            doc.add_paragraph(f"[Diagram rendering failed: {caption}]")
            print(f"    FAILED - empty response")
            return False
    except Exception as e:
        doc.add_paragraph(f"[Diagram error: {e}]")
        print(f"    ERROR: {e}")
        return False


def add_body(doc, text):
    """Add body text paragraph."""
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(10)
    return p


def add_bullet(doc, text):
    """Add bullet point."""
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(10)
    return p


def add_numbered(doc, text):
    """Add numbered list item."""
    p = doc.add_paragraph(text, style="List Number")
    for run in p.runs:
        run.font.size = Pt(10)
    return p


# ============================================================
# PUML DEFINITIONS
# ============================================================

PUML_PACKAGE_OVERVIEW = r"""@startuml
skinparam shadowing false
skinparam defaultFontSize 11
skinparam packageStyle rectangle
skinparam nodesep 40
skinparam ranksep 60

title Section 6.0 — Package Relationship Overview

package "PromptAssemblyFactory" as PAF #DDFFDD {
}

package "SharedServicesFactory" as SSF #FFDDDD {
}

package "FindCarePromptFactory" as FCPF #DDDDFF {
}

package "NarrativePromptFactory" as NPF #FFFFDD {
}

package "EvaluateCarePromptFactory" as ECPF #FFE4CC {
}

' Assembly depends on all
PAF --> SSF : <<uses>>\nsafety_check\npost_processing
PAF --> FCPF : <<routes to>>\nprovider search
PAF --> NPF : <<routes to>>\nnarrative questions
PAF --> ECPF : <<routes to>>\nclinical trials\nprovider evaluation

' SharedServices provides forces
SSF -[#red]-> PAF : <<force>> HALT\n(emergency/IP lock)
SSF -[#orange]-> PAF : <<force>> DOWNGRADE\n(off-topic >= 5)
SSF -[#blue]-> PAF : <<force>> OVERLAY\n(consent needed)

' Cross-epic awareness
FCPF ..> ECPF : selected_providers\n(handoff for evaluation)
SSF ..> FCPF : validates URLs\nin provider results
SSF ..> NPF : validates URLs\nin narrative text
SSF ..> ECPF : validates URLs\nin research links

@enduml"""

PUML_PKG_ASSEMBLY = r"""@startuml
skinparam shadowing false
skinparam defaultFontSize 9
skinparam classAttributeIconSize 0

title Section 6.1 — PromptAssemblyFactory (Internal)

class PromptAssemblyFactory {
  -_sub_factory_registry: dict[str, SubFactory]
  -_session_store: SessionStateManager
  -_slot_layout: list[str]
  +assemble(user_input: str, session: SessionState): Response
  +route(user_input: str, session: SessionState): str
  -_resolve_dependencies(parts: list[PromptPart]): list[PromptPart]
  -_check_forces(parts: list[PromptPart]): ForceResult
  -_sort_by_slot_and_priority(parts: list[PromptPart]): list[PromptPart]
}

class SessionStateManager {
  +last_state: Optional[str]
  +last_specialty: Optional[str]
  +last_city: Optional[str]
  +off_topic_count: int
  +consent_status: str
  +message_history: list[dict]
  +message_count: int
  +update(key: str, value: any): void
  +get(key: str): any
}

class SubFactoryRegistry {
  -_factories: dict[str, SubFactory]
  +register(epic_id: str, factory: SubFactory): void
  +get(epic_id: str): SubFactory
  +get_intent_declarations(): dict
}

class RoutingEngine {
  -_intent_map: dict[str, str]
  +classify_intent(user_input: str, declarations: dict): RoutingDecision
  +get_default_epic(): str
}

class RoutingDecision {
  +epic_id: str
  +intent: str
  +confidence: float
}

interface PromptPart <<interface>> {
  +slot: str
  +content: any
  +priority: int
  +forces: list[str]
  +depends_on: list[str]
}

class ForceResult {
  +halt: bool
  +downgrade: bool
  +overlay: bool
  +halt_response: Optional[str]
  +overlay_content: Optional[str]
}

PromptAssemblyFactory --> SessionStateManager
PromptAssemblyFactory --> SubFactoryRegistry
PromptAssemblyFactory --> RoutingEngine
PromptAssemblyFactory --> PromptPart : collects
RoutingEngine --> RoutingDecision : produces
PromptAssemblyFactory --> ForceResult : evaluates

@enduml"""

PUML_PKG_SHARED = r"""@startuml
skinparam shadowing false
skinparam defaultFontSize 9
skinparam classAttributeIconSize 0

title Section 6.2 — SharedServicesFactory (Internal)

class SharedServicesFactory {
  -_safety: SafetyService
  -_consent: ConsentService
  -_lead: LeadService
  -_url_guardian: URLGuardian
  -_unknown: UnknownQuestionService
  +get_pre_parts(user_msg: str, ip: str, session: dict): list[PromptPart]
  +get_post_parts(llm_response: str, session: dict, routing: str): list[PromptPart]
}

class SafetyService {
  -_get_db: callable
  -_keywords: list[str]
  -_openai_key: str
  +is_emergency(message: str): bool
  +is_ip_locked(ip: str): bool
  +lock_ip(ip: str, trigger: str, history: list): bool
  +try_admin_unlock(message: str, ip: str): bool
}
note right of SafetyService : Three-gate AI classifier:\nbody location + acute onset\n+ life-threat.\nAll three YES = emergency.

class ConsentService {
  -_anthropic_key: str
  +summarize_conversation(history: list): str
  +de_identify(history: list): void
}
note right of ConsentService : HIPAA two-tier consent.\nDe-identification via\nClaude Haiku.

class LeadService {
  -_get_db: callable
  -_consent: ConsentService
  +record_user_details(email: str, name: str, ...): dict
}

class URLGuardian {
  -_cache: dict
  +guard_text(text: str): str
  -_validate(url: str, context: str): tuple
  -_validate_batch(urls: list, context: str): dict
  -_ai_verify_content(url: str, context: str, content: str): bool
  -_find_correct_url(broken_url: str, context: str): str
}
note right of URLGuardian : Three stages:\n1. HEAD reachability\n2. AI content verify (Haiku)\n3. Google search correction

class UnknownQuestionService {
  -_consent: ConsentService
  +record(question: str, question_class: str, consent: str, ...): dict
}

SharedServicesFactory --> SafetyService
SharedServicesFactory --> ConsentService
SharedServicesFactory --> LeadService
SharedServicesFactory --> URLGuardian
SharedServicesFactory --> UnknownQuestionService
LeadService --> ConsentService

@enduml"""

PUML_PKG_FINDCARE = r"""@startuml
skinparam shadowing false
skinparam defaultFontSize 9
skinparam classAttributeIconSize 0

title Section 6.3 — FindCarePromptFactory (Internal)

class FindCarePromptFactory {
  -_find_care: FindCareService
  -_specialty: SpecialtyService
  -_tool_router: ToolRouter
  +get_parts(user_msg: str, session: dict): list[PromptPart]
  +declares_intents(): list[str]
}
note right of FindCarePromptFactory : Default epic for\nambiguous intents.\nIntents: provider_search,\nspecialty_lookup

class FindCareService {
  -_get_db: callable
  -_env: str
  -_get_embedding: callable
  -_specialty: SpecialtyService
  +search_providers(specialty_query: str, state: str, city: str, ...): dict
  +identify_specialty(query: str): dict
  -_facet_query(collection, filter: dict, after_npi: str, limit: int): tuple
  -_vector_search(embedding: list, state: str, city: str, county: str, limit: int): list
  -_build_summary_message(...): str
  -_format_provider(p: dict): dict
}
note right of FindCareService : 5 search strategies:\nNPI exact, name search,\nspecialty codes, specialty query,\ncounty fallback

class SpecialtyService {
  -_get_db: callable
  -_get_vector: callable
  +find_specialty_codes(query: str, chat_history: list): dict
}
note right of SpecialtyService : AI vector search against\nNUCC specialty embeddings.\ntext-embedding-3-large +\ncosine similarity.

class SpecialtyClassifier {
  +classify(specialty_code: str): dict
}
note bottom of SpecialtyClassifier : Classifies on can_prescribe\nand homeopathic dimensions\nvia GPT-4.1-mini

class SpecialtyRanker {
  +rank(specialties: list, query: str): list
}

class HomeopathicResolver {
  +evaluate(specialty: str, query: str): dict
}
note bottom of HomeopathicResolver : Evaluates alt-medicine\nspecialties via Claude Haiku

FindCarePromptFactory --> FindCareService
FindCarePromptFactory --> SpecialtyService
FindCareService --> SpecialtyService
SpecialtyService --> SpecialtyClassifier
SpecialtyService --> SpecialtyRanker
SpecialtyService --> HomeopathicResolver

@enduml"""

PUML_PKG_NARRATIVE = r"""@startuml
skinparam shadowing false
skinparam defaultFontSize 9
skinparam classAttributeIconSize 0

title Section 6.4 — NarrativePromptFactory (Internal)

class NarrativePromptFactory {
  -_about: AboutService
  -_llm: llm_client
  +get_parts(user_msg: str, session: dict): list[PromptPart]
  +declares_intents(): list[str]
}
note right of NarrativePromptFactory : Intents: narrative_question,\nabout_skip, about_chathealthy,\nmedical_info

class AboutService {
  -_me: dict
  +get_skip_snow_context(): dict
  +get_chathealthy_context(): dict
  +get_anthropic_context(): dict
}
note right of AboutService : Static content from\nme/ directory.\nSkip Snow bio,\nChatHealthy info,\nAnthropic principles.

class "llm_client" as LLM {
  +chat(model: str, messages: list, tools: list, system: str, ...): dict
}
note right of LLM : GPT-4.1 for narrative.\nCommodity LLM if\nDOWNGRADE active.

NarrativePromptFactory --> AboutService
NarrativePromptFactory --> LLM

@enduml"""

PUML_PKG_EVALUATE = r"""@startuml
skinparam shadowing false
skinparam defaultFontSize 9
skinparam classAttributeIconSize 0

title Section 6.5 — EvaluateCarePromptFactory (Internal)

class EvaluateCarePromptFactory {
  -_facade: EvaluateCareFacade
  +get_parts(user_msg: str, session: dict): list[PromptPart]
  +declares_intents(): list[str]
}
note right of EvaluateCarePromptFactory : Intents: clinical_trials,\nprovider_detail,\nquality_evaluation

class EvaluateCareFacade {
  -_clinical_trials: ClinicalTrialsService
  -_provider_detail: ProviderDetailService
  +search_clinical_trials(condition: str, location: str, ...): dict
  +get_provider_details(name: str, npi: str, state: str): dict
}

class ClinicalTrialsService {
  -_api_key: str
  +search(condition: str, location: str, user_location: str, max: int): dict
  -_get_travel_info(origin: str, destinations: list): dict
}
note right of ClinicalTrialsService : ClinicalTrials.gov API\n+ Google Routes API\nfor travel enrichment.

class ProviderDetailService {
  +lookup(provider_name: str, npi: str, state: str): dict
}
note right of ProviderDetailService : NPI Registry lookup +\nexternal research URLs:\nHealthgrades, NPI Registry,\nstate medical board.

EvaluateCarePromptFactory --> EvaluateCareFacade
EvaluateCareFacade --> ClinicalTrialsService
EvaluateCareFacade --> ProviderDetailService

@enduml"""


def load_puml_file(filename):
    """Load a .puml file from the audit directory."""
    path = os.path.join(AUDIT_DIR, filename)
    if os.path.exists(path):
        with open(path, encoding="utf-8") as f:
            return f.read()
    return None


def build_document():
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)

    # ================================================================
    # TITLE PAGE
    # ================================================================
    doc.add_paragraph("")
    doc.add_paragraph("")
    title = doc.add_heading("ChatHealthy FindCare", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_heading("Prompt Factory Architecture Design", level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta_items = [
        ("Version", "5.0"),
        ("Date", "April 15, 2026"),
        ("Authors", "Skip Snow (CTO/CEO), Claude Opus 4.6"),
        ("Pattern", "Assembly Factory with Neighbor-Aware Sub-Factory Composition"),
        ("Status", "Design \u2014 human review required"),
        ("Copyright", "ChatHealthy.ai LLC."),
    ]
    for label, value in meta_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{label}: ")
        run.bold = True
        run.font.size = Pt(11)
        run2 = p.add_run(value)
        run2.font.size = Pt(11)

    doc.add_page_break()

    # ================================================================
    # SECTION 1: Problem Statement
    # ================================================================
    doc.add_heading("1. Problem Statement", level=1)

    add_body(doc, "ChatHealthy FindCare has a fully functional backend with 10 domain services, a tool router, an LLM chat loop, safety filtering, lead capture, clinical trials search, URL validation, and provider search \u2014 all accessible through a /chat endpoint.")
    add_body(doc, "None of it is reachable by users.")
    add_body(doc, "The current production UI (FindCareApp.tsx) bypasses /chat entirely. It calls /classify (vector search for specialties) and /search (direct MongoDB query) \u2014 two endpoints that skip every backend capability except raw provider lookup. Six funded, implemented backend capabilities are invisible to all users.")
    add_body(doc, "A codebase audit (April 12, 2026) identified 13 findings. An independent GPT review validated all 13. The root cause: FindCareApp.tsx does not use the prompt-driven chat endpoint.")
    add_body(doc, "This design defines the architecture that fixes this by routing all user input through a PromptAssemblyFactory \u2014 ensuring every backend capability is reachable through one unified path.")

    # 1.1
    doc.add_heading("1.1 What Users Experience Today", level=2)
    add_table(doc,
        ["User Action", "What Happens", "What Should Happen"],
        [
            ['"Find me a bone doc in Virginia"', "Vector search, specialty codes, DB query, provider cards", "Same, plus AI triage, context carry-forward, specialty ranking"],
            ['"What about in Richmond?"', "New stateless /classify call \u2014 Virginia is lost", "System infers Virginia from prior turn, narrows to Richmond"],
            ['"What does a physiatrist do?"', "Provider cards (no explanation)", "Narrative LLM response with Markdown"],
            ['"Tell me about clinical trials for diabetes"', "Provider cards (wrong)", "Clinical trials search via ClinicalTrials.gov API"],
            ['"Who is Skip Snow?"', "Provider cards (wrong)", "About ChatHealthy narrative response"],
            ["User shares contact info", "Nothing", "Lead capture to MongoDB"],
            ["AI generates fake URL", "Reaches user", "URL Guardian validates before rendering"],
        ]
    )

    # 1.2
    doc.add_heading("1.2 Audit Findings Addressed by This Design", level=2)
    add_table(doc,
        ["Finding", "Priority", "Status"],
        [
            ["F1: FindCareApp /chat bypass", "CRITICAL", "Design-resolved"],
            ["F2: Context carry-forward regression", "CRITICAL", "Design-resolved"],
            ["F5: URL Guardian dead-coded", "HIGH", "Design-resolved"],
            ["F7: Specialty ranker limbo", "MEDIUM", "Design-resolved"],
            ["F8: LLM narrative response inaccessible", "MEDIUM", "Design-resolved"],
            ["F11: Dynamic context loading inefficiency", "LOW", "Design-resolved"],
        ]
    )

    # 1.3
    doc.add_heading("1.3 Audit Findings NOT Addressed by This Design", level=2)
    add_table(doc,
        ["Finding", "Priority", "Why Not"],
        [
            ["F3: UMLS CPT compliance", "HIGH", "Legal/compliance \u2014 handled by SharedServicesFactory"],
            ["F4: ChatWindow.tsx retirement", "HIGH", "UX governance decision"],
            ["F6: Regex specialty fallback", "MEDIUM", "Infrastructure resilience \u2014 specialty_service needs try/except"],
            ["F9: HuggingFace mTLS", "MEDIUM", "Network/security layer"],
            ["F10: Brain loop dead code", "LOW", "Internal tooling"],
            ["F12: Google Maps", "LOW", "Feature not built"],
            ["F13: EvaluateCare scoring", "LOW", "Feature not built"],
        ]
    )

    doc.add_page_break()

    # ================================================================
    # SECTION 2: Core Business Requirements
    # ================================================================
    doc.add_heading("2. Core Business Requirements", level=1)

    doc.add_heading("2.1 The Fundamental Requirement", level=2)
    add_body(doc, "One PromptAssemblyFactory produces dynamic prompts to do all the work. Every user input \u2014 whether it is a provider search, a conversational question, a clinical trials query, or an emergency \u2014 enters the system through a single assembly factory that combines prompt fragments from specialized sub-factories into a coherent LLM prompt carrying the user's history and context.")

    doc.add_heading("2.2 Business Requirements", level=2)
    add_table(doc,
        ["#", "Requirement"],
        [
            ["BR-01", "Every user input MUST be routed through the PromptAssemblyFactory. No UI component may bypass the factory by calling backend services directly."],
            ["BR-02", "The system MUST carry conversation context across turns. If a user says \"Virginia\" then \"Richmond\", the system MUST infer Virginia from the prior turn."],
            ["BR-03", "The system MUST be able to return both structured results (provider cards) and narrative responses (Markdown text) from the same entry point."],
            ["BR-04", "The consent workflow MUST occur if and only if: (a) consent has not already been obtained in this session, AND (b) the system needs to or wishes to persist user data."],
            ["BR-05", "All URLs in LLM responses MUST be validated before reaching the user."],
            ["BR-06", "Emergency/safety detection MUST occur on every request. Safety is managed by the SharedServicesFactory, not as standalone middleware."],
            ["BR-07", "Lead capture MUST trigger when the user provides contact information, regardless of conversation context."],
            ["BR-08", "Clinical trials search MUST be accessible through the same entry point as provider search."],
            ["BR-09", "The system MUST redact UMLS CPT codes from all user-facing output."],
            ["BR-10", "Each epic's prompt logic MUST be independently deployable and testable. A bug in one epic's sub-factory MUST NOT affect other epics."],
            ["BR-11", "If the user asks 5 consecutive off-topic questions unrelated to the product's mission, the system MUST downgrade to a commodity LLM and provide minimal responses until the user returns to on-topic queries. This is a Shared Services function."],
        ]
    )

    doc.add_page_break()

    # ================================================================
    # SECTION 3: Architecture
    # ================================================================
    doc.add_heading("3. Architecture", level=1)

    doc.add_heading("3.1 The Aircraft Assembly Pattern", level=2)
    add_body(doc, "The architecture uses the aircraft assembly pattern. There are five participants:")
    add_body(doc, "The PromptAssemblyFactory is the final assembly line. It knows the blueprint \u2014 where each part mounts, what forces each part creates, what order assembly proceeds, and how to wire the connections between sub-factories. It does NOT understand the content of any part. It is deliberately dumb.")
    add_body(doc, "Sub-factories (one per epic, plus SharedServicesFactory) each produce finished parts with a standard interface. Each sub-factory understands its own domain deeply. It knows nothing about how its output will be combined with other sub-factories' output by the assembly factory.")
    add_body(doc, "However \u2014 and this is critical \u2014 each sub-factory knows enough about its neighbors to design its contributions correctly. The wing knows the engine is there and what forces the engine applies. The engine knows it mounts on a wing and what structural loads the wing can bear. Neither knows the other's internals. Each sub-factory exposes an API that tells its neighbors what it needs and what it supplies. The neighbors use these APIs to design their own contributions, but never to understand or replicate each other's logic.")
    add_body(doc, "The PromptPart interface is the mounting spec \u2014 the standard connector between sub-factories and the assembly factory. Every sub-factory delivers a PromptPart. The assembly factory only interacts with parts through this interface.")
    add_body(doc, "The PromptAssemblyFactory reads the sub-factory APIs and wires the connections based on the blueprint. The sub-factories do not call each other directly.")

    doc.add_heading("3.2 Why Not a Monolithic Factory", level=2)
    add_body(doc, "A single factory class handling all nine epics would become a God Object \u2014 the exact maintenance and regression risk the audit identified. One class per epic means one failure domain, one test surface, and one deployment unit per capability cluster.")

    doc.add_heading("3.3 The PromptPart Interface \u2014 The Mounting Spec", level=2)
    add_body(doc, "Every sub-factory delivers PromptParts through this standard interface:")
    add_table(doc,
        ["Field", "Type", "Purpose"],
        [
            ["slot", "string", "Where this part mounts: system_context, user_intent, tools, constraints, or post_processing"],
            ["content", "any", "The prompt fragment \u2014 opaque to the assembly factory"],
            ["priority", "integer", "Assembly order within its slot. 0 = first, 100 = last"],
            ["forces", "list", "Structural effects: HALT, OVERLAY, DOWNGRADE, or NONE"],
            ["depends_on", "list", "What must be resolved before this part can be produced: session_state, routing_decision, llm_response"],
        ]
    )
    add_body(doc, "The assembly factory sorts parts by slot and priority, resolves dependency order, checks for force conflicts, and assembles. It never reads content.")

    doc.add_heading("3.4 What the PromptAssemblyFactory Knows (The Blueprint)", level=2)
    add_body(doc, "The assembly factory's intelligence is limited to exactly these six things:")
    add_numbered(doc, "Slot layout \u2014 the order of slots: system_context, then constraints, then user_intent, then tools, then LLM call, then post_processing.")
    add_numbered(doc, "Sub-factory registry \u2014 which sub-factories exist and their capability declarations (their APIs). Not their internals.")
    add_numbered(doc, "Routing rules \u2014 given user input plus session state, which epic sub-factory handles this request. This is the assembly factory's one non-trivial responsibility. Routing rules are explicit and declarative. Each epic sub-factory declares what intents it handles. The assembly factory matches user intent to declarations. Ambiguous intents default to FindCare (the primary epic).")
    add_numbered(doc, "Dependency resolution \u2014 if a PromptPart says it depends_on safety_status, the assembly factory ensures the safety check runs first. If it depends_on llm_response, the assembly factory ensures the LLM call happens before that part runs.")
    add_numbered(doc, "Force wiring \u2014 if any part says force=HALT, the assembly stops. If force=DOWNGRADE, the assembly swaps the LLM model. If force=OVERLAY, the assembly wraps the response with the overlay content.")
    add_numbered(doc, "Session state management \u2014 the assembly factory maintains session state (last_state, last_specialty, off_topic_count, consent_status, message_history) and passes it to sub-factories that require it.")

    doc.add_heading("3.5 What the PromptAssemblyFactory Does NOT Know", level=2)
    add_body(doc, "The assembly factory does not know:")
    items_not_known = [
        "What a specialty code is",
        "How clinical trials are searched",
        "What makes a URL valid or invalid",
        "What words are profane",
        "What constitutes an emergency",
        "How provider quality is scored",
        "What UMLS CPT codes look like",
        "What an off-topic question looks like",
        "How consent works",
        "How lead capture detects contact information",
    ]
    for item in items_not_known:
        add_bullet(doc, item)
    add_body(doc, "If we give it more than the six things above, it becomes a god object. If we give it less, it cannot assemble correctly.")

    doc.add_heading("3.6 The Routing Decision", level=2)
    add_body(doc, "The routing decision is the assembly factory's one non-trivial responsibility. It must classify user intent to pick the right epic sub-factory. This is the area most at risk of becoming god-like.")
    add_body(doc, "The routing rules should be:")
    add_bullet(doc, "Explicit and declarative, not a chain of if/else")
    add_bullet(doc, "Each epic sub-factory declares what intents it handles as part of its API")
    add_bullet(doc, "The assembly factory matches user intent to those declarations")
    add_bullet(doc, "Ambiguous intents default to FindCare (the primary epic)")
    add_body(doc, "The assembly factory does not understand WHY a particular intent maps to a particular sub-factory. It only knows the mapping.")

    doc.add_page_break()

    # ================================================================
    # SECTION 4: Sub-Factory Capability Declarations
    # ================================================================
    doc.add_heading("4. Sub-Factory Capability Declarations", level=1)
    add_body(doc, "This section defines each sub-factory completely: what it does, what real code it wraps, what capabilities it has, what it supplies to neighbors, what it requires from neighbors, what forces it creates, and what forces it responds to.")
    add_body(doc, "From a processing perspective, every attribute and function must be supervised by Skip and not auto-generated by Claude without supervision. This object will evolve during implementation.")

    # 4.1 SharedServicesFactory
    doc.add_heading("4.1 SharedServicesFactory", level=2)
    add_body(doc, "Domain: Cross-cutting concerns that apply to every request regardless of which epic handles it. Lives in the Shared Services epic. This is the one sub-factory that runs on EVERY request \u2014 both before and after the epic sub-factory.")
    add_body(doc, "Real code it wraps:")
    add_bullet(doc, "SafetyService (safety_service.py) \u2014 emergency detection via three-gate AI classifier (body location + acute onset + life-threat), IP locking, audit trail")
    add_bullet(doc, "ConsentService (consent_service.py) \u2014 HIPAA two-tier consent, de-identification via Claude Haiku")
    add_bullet(doc, "LeadService (lead_service.py) \u2014 contact capture to MongoDB")
    add_bullet(doc, "URLGuardian (url_guardian.py) \u2014 three-stage URL validation (HEAD check, AI content verify, Google search correction)")
    add_bullet(doc, "UnknownQuestionService (unknown_question_service.py) \u2014 off-topic classification and recording")
    add_bullet(doc, "CPT redaction (not yet implemented \u2014 compliance requirement)")

    doc.add_heading("Capabilities", level=3)
    add_table(doc,
        ["Capability", "Slot", "Priority", "Force", "Description"],
        [
            ["safety_check", "constraints", "0", "HALT", "Checks if message is an emergency (SafetyService.is_emergency \u2014 three-gate: body location + acute onset + life-threat). Also checks IP lock. If triggered, force=HALT and assembly stops immediately. Returns emergency response with 911 guidance."],
            ["off_topic_detection", "constraints", "10", "DOWNGRADE", "Tracks consecutive off-topic messages using session state. After 5 consecutive off-topic questions, force=DOWNGRADE \u2014 the assembly factory swaps GPT-4.1 for a commodity LLM. Resets when user returns to on-topic queries."],
            ["consent_check", "post_processing", "80", "OVERLAY", "Checks two conditions: (a) system needs to persist user data (lead capture, transcript storage), AND (b) consent has not yet been obtained in this session. If both true, force=OVERLAY \u2014 consent dialog wraps the normal response."],
            ["lead_capture", "post_processing", "70", "NONE", "Detects contact information in user input (email, phone, name). Triggers LeadService.record_user_details to MongoDB. Respects consent tier for de-identification."],
            ["unknown_recording", "post_processing", "75", "NONE", "If the routing decision classified the question as unanswerable, records the question via UnknownQuestionService for product improvement."],
            ["cpt_redaction", "post_processing", "85", "NONE", "Scans LLM response for UMLS CPT codes and redacts them. Compliance requirement \u2014 CPT codes cannot appear in consumer-facing output."],
            ["url_validation", "post_processing", "90", "NONE", "Validates all URLs in LLM response using URLGuardian. Three stages: HEAD reachability check, AI content verification via Claude Haiku, Google Custom Search for broken URL correction. Defangs links that cannot be validated."],
        ]
    )

    doc.add_heading("What SharedServicesFactory Supplies (API to Neighbors)", level=3)
    add_table(doc,
        ["Output", "Type", "Description"],
        [
            ["safety_status", "bool", "Is this request safe to process?"],
            ["ip_locked", "bool", "Is this IP address locked out?"],
            ["consent_status", "string", '"obtained" or "needed" or "not_needed"'],
            ["off_topic_count", "integer", "Consecutive off-topic message count"],
            ["validated_response", "string", "LLM response text with URLs validated and CPT codes redacted"],
        ]
    )

    doc.add_heading("What SharedServicesFactory Requires (API from Neighbors)", level=3)
    add_table(doc,
        ["Input", "Type", "Source", "Description"],
        [
            ["user_message", "string", "PromptAssemblyFactory", "Raw user input for safety check and off-topic detection"],
            ["user_ip", "string", "PromptAssemblyFactory", "IP address for IP locking"],
            ["session_state", "dict", "PromptAssemblyFactory", "consent_obtained, off_topic_count, message_count"],
            ["llm_response", "string", "PromptAssemblyFactory", "LLM output text \u2014 only for post_processing capabilities"],
            ["routing_decision", "string", "PromptAssemblyFactory", "Which epic was selected \u2014 for determining if question was on-mission"],
        ]
    )

    doc.add_heading("Forces Created by SharedServicesFactory", level=3)
    add_table(doc,
        ["Force", "When", "Effect on Assembly"],
        [
            ["HALT", "Emergency detected or IP locked", "Assembly stops immediately. Return emergency response. No epic sub-factory runs."],
            ["OVERLAY", "Consent needed and not yet obtained", "Normal response produced by epic sub-factory, then consent dialog overlaid on top. Epic runs normally."],
            ["DOWNGRADE", "5 or more consecutive off-topic messages", "Assembly swaps GPT-4.1 for commodity LLM. Epic sub-factory still runs but with degraded model quality."],
        ]
    )

    doc.add_heading("Forces SharedServicesFactory Responds To", level=3)
    add_body(doc, "SharedServicesFactory does not respond to forces from other sub-factories. It is always invoked. It is the first and last to run.")

    # 4.2 FindCarePromptFactory
    doc.add_heading("4.2 FindCarePromptFactory", level=2)
    add_body(doc, "Domain: Provider search, specialty resolution, geographic context carry-forward. The core FindCare epic. This is the primary epic \u2014 ambiguous intents default here.")
    add_body(doc, "Real code it wraps:")
    add_bullet(doc, "FindCareService (provider_search_service.py) \u2014 five search strategies: NPI exact lookup, provider name search, specialty codes direct filter, specialty query (vector-to-taxonomy), and county fallback")
    add_bullet(doc, "SpecialtyService (specialty_service.py) \u2014 AI vector search against NUCC specialty embeddings using text-embedding-3-large and cosine similarity")
    add_bullet(doc, "SpecialtyClassifier (specialty_classifier.py) \u2014 classifies specialties on can_prescribe and homeopathic dimensions via GPT-4.1-mini")
    add_bullet(doc, "SpecialtyRanker (specialty_ranker.py) \u2014 reorders specialty options by relevance to user query via GPT-4.1-mini")
    add_bullet(doc, "HomeopathicResolver (homeopathic_resolver.py) \u2014 evaluates alternative medicine specialties for compliance with user search query via Claude Haiku")

    doc.add_heading("Capabilities", level=3)
    add_table(doc,
        ["Capability", "Slot", "Priority", "Force", "Description"],
        [
            ["provider_search_prompt", "user_intent", "50", "NONE", "Builds prompt fragment for provider search including specialty codes, state, city, county, pagination context. Registers find_providers and find_specialty_codes tools with the tool router."],
            ["context_carry_forward", "system_context", "40", "NONE", "Injects prior turn's state, specialty, and city into the system context so conversational follow-ups work. \"What about in Richmond?\" resolves Virginia from the prior turn."],
            ["specialty_resolution", "tools", "50", "NONE", "Provides vector search for specialty codes. Includes homeopathic resolution and specialty ranking as part of the search flow."],
            ["filter_options", "post_processing", "60", "NONE", "Extracts specialty filter options from search results for the UI filter panel. Includes can_prescribe and homeopathic flags per specialty."],
        ]
    )

    doc.add_heading("What FindCarePromptFactory Supplies (API to Neighbors)", level=3)
    add_table(doc,
        ["Output", "Type", "Description"],
        [
            ["search_results", "dict", "Providers list, total_count, has_more flag, pagination cursor (after_npi)"],
            ["specialty_options", "list", "Specialty codes for UI filter panel with can_prescribe and homeopathic flags"],
            ["search_context", "dict", "last_state, last_city, last_specialty \u2014 for context carry-forward across turns"],
            ["response_type", "string", '"provider_cards" \u2014 tells the UI to render structured provider cards, not narrative text'],
        ]
    )

    doc.add_heading("What FindCarePromptFactory Requires (API from Neighbors)", level=3)
    add_table(doc,
        ["Input", "Type", "Source", "Description"],
        [
            ["safety_status", "bool", "SharedServicesFactory", "Do not build prompt if HALT"],
            ["session_state", "dict", "PromptAssemblyFactory", "last_state, last_city, last_specialty, message_history"],
            ["user_message", "string", "PromptAssemblyFactory", "Raw user input"],
            ["off_topic_count", "integer", "SharedServicesFactory", "If DOWNGRADE, skip expensive vector search"],
        ]
    )

    doc.add_heading("Forces Created by FindCarePromptFactory", level=3)
    add_body(doc, "FindCarePromptFactory creates no forces. It is a content producer.")

    doc.add_heading("Forces FindCarePromptFactory Responds To", level=3)
    add_table(doc,
        ["Force", "From", "My Response"],
        [
            ["HALT", "SharedServicesFactory", "Do not run. Do not build prompt. Do not call MongoDB or OpenAI."],
            ["DOWNGRADE", "SharedServicesFactory", "Skip vector search (expensive). Use simple text match fallback. Return fewer results."],
        ]
    )

    # 4.3 NarrativePromptFactory
    doc.add_heading("4.3 NarrativePromptFactory", level=2)
    add_body(doc, "Domain: Conversational questions that do not produce provider cards. Examples: \"What does a physiatrist do?\" \"Who is Skip Snow?\" \"What is ChatHealthy?\" These are questions that need a narrative Markdown response, not structured data.")
    add_body(doc, "Real code it wraps:")
    add_bullet(doc, "AboutService (about_service.py) \u2014 Skip Snow biography, ChatHealthy company info, Anthropic principles. Static content loaded from the me/ directory.")
    add_bullet(doc, "The general LLM chat path \u2014 system prompt plus user message produces narrative Markdown. No tool calls involved.")

    doc.add_heading("Capabilities", level=3)
    add_table(doc,
        ["Capability", "Slot", "Priority", "Force", "Description"],
        [
            ["narrative_prompt", "user_intent", "50", "NONE", "Builds prompt for open-ended conversational response. The LLM returns Markdown text, not structured provider data."],
            ["about_context", "system_context", "45", "NONE", "When intent is \"about\" (Skip Snow, ChatHealthy, Anthropic), injects the relevant About content into the system context."],
        ]
    )

    doc.add_heading("What NarrativePromptFactory Supplies (API to Neighbors)", level=3)
    add_table(doc,
        ["Output", "Type", "Description"],
        [
            ["narrative_text", "string", "Markdown response \u2014 not structured data"],
            ["response_type", "string", '"narrative" \u2014 tells the UI to render Markdown, not provider cards'],
            ["about_context", "dict", "Skip Snow bio, ChatHealthy company info \u2014 only when intent is \"about\""],
        ]
    )

    doc.add_heading("What NarrativePromptFactory Requires (API from Neighbors)", level=3)
    add_table(doc,
        ["Input", "Type", "Source", "Description"],
        [
            ["safety_status", "bool", "SharedServicesFactory", "Do not run if HALT"],
            ["user_message", "string", "PromptAssemblyFactory", "Raw user input"],
            ["routing_decision", "string", "PromptAssemblyFactory", "Confirmation that this is a narrative question, not a search"],
            ["session_state", "dict", "PromptAssemblyFactory", "message_history for conversational continuity"],
        ]
    )

    doc.add_heading("Forces Created by NarrativePromptFactory", level=3)
    add_body(doc, "NarrativePromptFactory creates no forces. It is a content producer.")

    doc.add_heading("Forces NarrativePromptFactory Responds To", level=3)
    add_table(doc,
        ["Force", "From", "My Response"],
        [
            ["HALT", "SharedServicesFactory", "Do not run."],
            ["DOWNGRADE", "SharedServicesFactory", "Use commodity LLM for response instead of GPT-4.1."],
        ]
    )

    # 4.4 EvaluateCarePromptFactory
    doc.add_heading("4.4 EvaluateCarePromptFactory", level=2)
    add_body(doc, "Domain: Provider quality evaluation, clinical trials search, provider detail lookup. The Evaluate Care epic. This sub-factory handles requests that go beyond finding a provider \u2014 evaluating provider quality, searching for clinical trials, and looking up detailed provider information.")
    add_body(doc, "Real code it wraps:")
    add_bullet(doc, "EvaluateCareFacade (evaluate_care_facade.py) \u2014 public interface wrapping clinical trials and provider detail services")
    add_bullet(doc, "ClinicalTrialsService (clinical_trials_service.py) \u2014 searches ClinicalTrials.gov for recruiting trials, optionally enriches with travel time via Google Routes API")
    add_bullet(doc, "ProviderDetailService (provider_detail_service.py) \u2014 NPI Registry lookup plus constructs external research URLs for Healthgrades, NPI Registry, and state medical board")

    doc.add_heading("Capabilities", level=3)
    add_table(doc,
        ["Capability", "Slot", "Priority", "Force", "Description"],
        [
            ["clinical_trials_prompt", "user_intent", "50", "NONE", "Builds prompt for clinical trials search. Registers search_clinical_trials tool with the tool router. Queries ClinicalTrials.gov API."],
            ["provider_detail_prompt", "user_intent", "50", "NONE", "Builds prompt for provider quality lookup. Registers lookup_provider_external tool. Queries NPI Registry API."],
            ["quality_context", "system_context", "45", "NONE", "Injects quality scoring context (measures, weights, explainability framework) into the system prompt when the intent is evaluate."],
        ]
    )

    doc.add_heading("What EvaluateCarePromptFactory Supplies (API to Neighbors)", level=3)
    add_table(doc,
        ["Output", "Type", "Description"],
        [
            ["trials_results", "dict", "Clinical trial list with NCT IDs, phases, locations, eligibility, travel info"],
            ["provider_details", "dict", "NPI details, research URLs (Healthgrades, NPI Registry, state board), quality scores"],
            ["response_type", "string", '"trials" or "provider_detail" or "quality_score" \u2014 tells the UI what to render'],
        ]
    )

    doc.add_heading("What EvaluateCarePromptFactory Requires (API from Neighbors)", level=3)
    add_table(doc,
        ["Input", "Type", "Source", "Description"],
        [
            ["safety_status", "bool", "SharedServicesFactory", "Do not run if HALT"],
            ["user_message", "string", "PromptAssemblyFactory", "Raw user input"],
            ["routing_decision", "string", "PromptAssemblyFactory", "Confirmation that this is an evaluate care request"],
            ["session_state", "dict", "PromptAssemblyFactory", "selected_providers from FindCare handoff \u2014 which providers the user selected for evaluation"],
        ]
    )

    doc.add_heading("Forces Created by EvaluateCarePromptFactory", level=3)
    add_body(doc, "EvaluateCarePromptFactory creates no forces. It is a content producer.")

    doc.add_heading("Forces EvaluateCarePromptFactory Responds To", level=3)
    add_table(doc,
        ["Force", "From", "My Response"],
        [
            ["HALT", "SharedServicesFactory", "Do not run."],
            ["DOWNGRADE", "SharedServicesFactory", "Skip quality scoring (expensive). Return basic provider details only."],
        ]
    )

    # 4.5 PromptAssemblyFactory
    doc.add_heading("4.5 PromptAssemblyFactory \u2014 The Blueprint", level=2)
    add_body(doc, "Domain: Assembly only. Knows the blueprint. Does not produce content. Deliberately dumb.")

    doc.add_heading("What the PromptAssemblyFactory Knows", level=3)
    add_numbered(doc, "Slot layout \u2014 the order of assembly slots: system_context first, then constraints, then user_intent, then tools, then the LLM call happens, then post_processing runs on the LLM response.")
    add_numbered(doc, "Sub-factory registry \u2014 which sub-factories exist, and their capability declarations (their supply/require/force APIs). Not their internals.")
    add_numbered(doc, "Routing rules \u2014 given user input plus session state, which epic sub-factory handles this request. Each epic sub-factory declares what intents it handles. The assembly factory matches user intent to those declarations. Ambiguous intents default to FindCare.")
    add_numbered(doc, "Dependency resolution \u2014 SharedServicesFactory.safety_status must resolve before any epic sub-factory runs. Post-processing parts depend on the LLM response existing.")
    add_numbered(doc, "Force wiring \u2014 HALT from SharedServicesFactory means skip all epic sub-factories and return the emergency response. DOWNGRADE means swap the LLM model to a commodity model. OVERLAY means wrap the final response with a consent dialog.")
    add_numbered(doc, "Session state management \u2014 the assembly factory maintains session state (last_state, last_specialty, off_topic_count, consent_status, message_history) and passes it to sub-factories that require it.")

    doc.add_heading("What the PromptAssemblyFactory Does NOT Know", level=3)
    items = [
        "What a specialty code is",
        "How clinical trials are searched",
        "What makes a URL valid or invalid",
        "What words are profane or dangerous",
        "What constitutes a medical emergency",
        "How provider quality is scored",
        "What UMLS CPT codes look like",
        "What an off-topic question looks like",
        "How consent works or when it is needed",
        "How lead capture detects contact information",
    ]
    for item in items:
        add_bullet(doc, item)
    add_body(doc, "If we give the assembly factory knowledge beyond the six items above, it becomes a god object. If we give it less, it cannot assemble correctly. The six items are the minimum viable blueprint.")

    doc.add_heading("The Assembly Sequence", level=3)
    steps = [
        "Step 1: Receive user input plus session state from the UI.",
        "Step 2: Ask SharedServicesFactory for safety_check (priority 0, slot=constraints). If force equals HALT, return the emergency response immediately. Stop. No epic sub-factory runs.",
        "Step 3: Route. Determine which epic sub-factory handles this request. The routing decision is the assembly factory's one non-trivial responsibility. Each sub-factory declares what intents it handles. The assembly factory matches user intent to those declarations.",
        "Step 4: Ask SharedServicesFactory for off_topic_detection (priority 10, slot=constraints). If force equals DOWNGRADE, note the LLM swap for step 7.",
        "Step 5: Ask the selected epic sub-factory for its PromptParts (priority 50). The sub-factory reads its required inputs from the assembly (session_state, safety_status, routing_decision). The sub-factory returns PromptParts with the content it produced. The assembly factory does not inspect that content.",
        "Step 6: Collect all PromptParts from all sources. Sort by slot, then by priority within each slot.",
        "Step 7: Assemble the final prompt from the sorted parts. Send the assembled prompt to the LLM. If DOWNGRADE was noted in step 4, send to the commodity LLM instead of GPT-4.1.",
        "Step 8: Receive the LLM response.",
        "Step 9: Ask SharedServicesFactory for post_processing parts (priorities 70 through 90). These include lead capture, unknown recording, consent check, CPT redaction, and URL validation. Apply each post-processing part to the LLM response in priority order.",
        "Step 10: If the consent check produced force=OVERLAY, wrap the processed response with the consent dialog.",
        "Step 11: Update session state. Save last_state, last_specialty, off_topic_count, consent_status, and message_history for the next turn.",
        "Step 12: Return the final response to the UI. The response includes response_type (from the epic sub-factory) so the UI knows whether to render provider cards, narrative Markdown, clinical trials, or a quality report.",
    ]
    for step in steps:
        add_body(doc, step)

    doc.add_page_break()

    # ================================================================
    # SECTION 5: Neighbor Awareness Matrix
    # ================================================================
    doc.add_heading("5. Neighbor Awareness Matrix", level=1)
    add_body(doc, "This matrix shows what each sub-factory knows about its neighbors. The rule: each sub-factory knows its neighbors' interfaces (what they supply, what forces they create) but not their implementations (how they produce those outputs). The wing knows the engine's thrust vector but not the combustion chamber design.")

    add_table(doc,
        ["Sub-Factory", "Knows About SharedServices", "Knows About FindCare", "Knows About Narrative", "Knows About EvaluateCare"],
        [
            ["SharedServicesFactory", "(self)", "Knows FindCare produces URLs in its provider results that need validation", "Knows Narrative produces text that may contain URLs", "Knows EvaluateCare produces URLs (research site links) that need validation"],
            ["FindCarePromptFactory", "Knows safety_status and off_topic_count affect whether it runs and how", "(self)", "Does not know about Narrative", "Knows EvaluateCare may receive its selected_providers for quality evaluation"],
            ["NarrativePromptFactory", "Knows safety_status affects whether it runs", "Does not know about FindCare", "(self)", "Does not know about EvaluateCare"],
            ["EvaluateCarePromptFactory", "Knows safety_status affects whether it runs", "Knows FindCare provides selected_providers for quality evaluation", "Does not know about Narrative", "(self)"],
            ["PromptAssemblyFactory", "Knows SharedServices runs first (safety) and last (post-processing). Knows its three forces (HALT, OVERLAY, DOWNGRADE).", "Knows FindCare handles search intents. Knows it is the default for ambiguous intents.", "Knows Narrative handles question/about intents.", "Knows EvaluateCare handles quality/trials intents."],
        ]
    )

    doc.add_page_break()

    # ================================================================
    # SECTION 6: UML Package Diagrams
    # ================================================================
    doc.add_heading("6. UML Diagrams", level=1)

    doc.add_heading("6.0 Package Relationship Overview", level=2)
    add_body(doc, "This diagram shows the relationship between all five packages with no internal objects expressed. It shows dependency arrows and force flows (HALT, OVERLAY, DOWNGRADE) between packages.")
    add_uml_diagram(doc, PUML_PACKAGE_OVERVIEW, "Figure 6.0 \u2014 Package Relationship Overview")

    doc.add_page_break()

    doc.add_heading("6.1 Package: PromptAssemblyFactory", level=2)
    add_body(doc, "Internal class diagram showing the assembly flow, session state management, routing logic, and force propagation. The PromptAssemblyFactory contains: a SubFactoryRegistry for tracking registered sub-factories, a RoutingEngine for classifying user intent, a SessionStateManager for maintaining conversation state, and a ForceResult for evaluating structural effects from PromptParts.")
    add_uml_diagram(doc, PUML_PKG_ASSEMBLY, "Figure 6.1 \u2014 PromptAssemblyFactory Internal Structure")

    doc.add_heading("6.2 Package: SharedServicesFactory", level=2)
    add_body(doc, "Internal class diagram showing safety detection, consent workflow, URL validation, CPT redaction, off-topic detection, and lead capture. SharedServicesFactory wraps five existing domain services (SafetyService, ConsentService, LeadService, URLGuardian, UnknownQuestionService) and provides pre-processing parts (safety_check, off_topic_detection) and post-processing parts (lead_capture, unknown_recording, consent_check, cpt_redaction, url_validation).")
    add_uml_diagram(doc, PUML_PKG_SHARED, "Figure 6.2 \u2014 SharedServicesFactory Internal Structure")

    doc.add_page_break()

    doc.add_heading("6.3 Package: FindCarePromptFactory", level=2)
    add_body(doc, "Internal class diagram showing provider search, specialty vector search, context carry-forward, and filter options. FindCarePromptFactory wraps FindCareService (five search strategies), SpecialtyService (vector search against NUCC embeddings), SpecialtyClassifier, SpecialtyRanker, and HomeopathicResolver.")
    add_uml_diagram(doc, PUML_PKG_FINDCARE, "Figure 6.3 \u2014 FindCarePromptFactory Internal Structure")

    doc.add_heading("6.4 Package: NarrativePromptFactory", level=2)
    add_body(doc, "Internal class diagram showing conversational responses and about content. NarrativePromptFactory is the simplest epic sub-factory \u2014 it wraps AboutService for static content and delegates to the LLM for open-ended narrative responses.")
    add_uml_diagram(doc, PUML_PKG_NARRATIVE, "Figure 6.4 \u2014 NarrativePromptFactory Internal Structure")

    doc.add_heading("6.5 Package: EvaluateCarePromptFactory", level=2)
    add_body(doc, "Internal class diagram showing clinical trials search, provider detail lookup, and quality scoring. EvaluateCarePromptFactory wraps EvaluateCareFacade, which in turn delegates to ClinicalTrialsService (ClinicalTrials.gov API + Google Routes) and ProviderDetailService (NPI Registry + research URLs).")
    add_uml_diagram(doc, PUML_PKG_EVALUATE, "Figure 6.5 \u2014 EvaluateCarePromptFactory Internal Structure")

    doc.add_page_break()

    # ================================================================
    # SECTION 7: UML Sequence Diagrams
    # ================================================================
    doc.add_heading("7. UML Sequence Diagrams", level=1)

    # Load existing puml files and update titles to match V5 numbering
    seq_diagrams = [
        ("7.1 Provider Search \u2014 Happy Path", "puml_seq_provider_search.puml"),
        ("7.2 Conversational Question \u2014 Narrative Response", "puml_seq_narrative.puml"),
        ("7.3 Clinical Trials Search", "puml_seq_clinical_trials.puml"),
        ("7.4 Emergency/Safety Detection (SharedServicesFactory HALT force)", "puml_seq_emergency.puml"),
        ("7.5 Context Carry-Forward \u2014 Virginia then Richmond", "puml_seq_context_carry.puml"),
        ("7.6 Consent Trigger (SharedServicesFactory OVERLAY force, conditional)", "puml_seq_consent.puml"),
        ("7.7 URL Validation on Tool Response (SharedServicesFactory post_processing)", "puml_seq_url_validation.puml"),
    ]

    for title, filename in seq_diagrams:
        doc.add_heading(title, level=2)
        puml_text = load_puml_file(filename)
        if puml_text:
            # Update the title line to match V5 section numbering
            import re
            puml_text = re.sub(r"title \d+\.\d+", f"title {title[:3]}", puml_text, count=1)
            add_uml_diagram(doc, puml_text, f"Figure {title}")
        else:
            doc.add_paragraph(f"[Diagram not found: {filename}]")
        doc.add_paragraph("")  # spacer

    doc.add_page_break()

    # ================================================================
    # SECTION 8: Assumptions
    # ================================================================
    doc.add_heading("8. Assumptions", level=1)
    add_table(doc,
        ["#", "Assumption"],
        [
            ["A-01", "The existing /chat endpoint and tool loop in main.py is the correct backend architecture. The PromptAssemblyFactory wraps it, not replaces it."],
            ["A-02", "FindCareApp.tsx will be modified to call the PromptAssemblyFactory endpoint instead of /classify and /search directly."],
            ["A-03", "The PromptAssemblyFactory runs server-side in the FastAPI backend, not client-side."],
            ["A-04", "Session state (consent status, conversation history, off-topic counter) is maintained per browser session via the existing session token mechanism."],
            ["A-05", "The ToolRouter and all registered tools remain unchanged. Sub-factories dispatch to them through the existing tool loop."],
            ["A-06", "No new LLM models are introduced. The existing GPT-4.1 chat model and GPT-4.1-mini for extraction/classification remain. A commodity LLM is used for off-topic downgrade."],
        ]
    )

    # ================================================================
    # SECTION 9: Known Gaps
    # ================================================================
    doc.add_heading("9. Known Gaps \u2014 What This Design Does NOT Cover", level=1)
    add_table(doc,
        ["#", "Gap", "Impact", "Mitigation"],
        [
            ["G-01", "UMLS CPT code compliance (F3)", "Legal pre-launch blocker", "SharedServicesFactory post-processing. Separate legal workstream required."],
            ["G-02", "ChatWindow.tsx formal retirement (F4)", "Governance debt", "File a story explicitly listing each capability: retire, re-implement, or defer."],
            ["G-03", "Regex specialty fallback (F6)", "Silent failure if vector index missing", "Add try/except in SpecialtyService with substring fallback. Independent of factory."],
            ["G-04", "HuggingFace mTLS (F9)", "Security gap in deployed environment", "Implement SEC-CERTAUTH-REMOTE. Independent of factory."],
            ["G-05", "Schema drift between factory output and tool contracts", "Silent degradation", "Extend SchemaDriftDetector to cover PromptPart schemas."],
            ["G-06", "FindCareApp UI rendering of narrative responses", "Users see only provider cards", "Frontend must detect response_type and render narrative Markdown when appropriate."],
        ]
    )

    # ================================================================
    # SECTION 10: Backlog
    # ================================================================
    doc.add_heading("10. Backlog", level=1)
    add_body(doc, "Pending \u2014 features, stories, requirements with B/T separation per engineering rule v4-035. Will be created after human approves the architecture.")

    # ================================================================
    # SAVE
    # ================================================================
    out_path = os.path.join(AUDIT_DIR, "PromptFactoryDesign_V5.docx")
    # If file is locked (open in Word), try removing first or use temp name
    try:
        if os.path.exists(out_path):
            os.remove(out_path)
    except PermissionError:
        out_path = os.path.join(AUDIT_DIR, "PromptFactoryDesign_V5_new.docx")
        print(f"  Original file locked, saving as: {os.path.basename(out_path)}")
    doc.save(out_path)
    print(f"\nSaved: {out_path}")
    print(f"Size: {os.path.getsize(out_path):,} bytes")
    return out_path


if __name__ == "__main__":
    build_document()
