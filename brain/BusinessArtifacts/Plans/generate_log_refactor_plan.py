"""One-shot generator — appends a plan record to business_plan.json and
produces the CTO-review Word doc for the 2026-04-17 conversation log service
refactor. Disposable per human's convention for BusinessArtifacts docx scripts.
"""
from __future__ import annotations

import json
import uuid
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

REPO_ROOT = Path(r"c:\chatHealthy\findCare")
BIZ_PLAN = REPO_ROOT / "brain" / "machine_artifacts" / "content" / "business_plan.json"
OUT_DIR = REPO_ROOT / "brain" / "BusinessArtifacts" / "Plans"
OUT_DOCX = OUT_DIR / "ConversationLogServiceRefactor_Plan_2026-04-17.docx"

PLAN_ID = "PLAN-LOG-SERVICE-REFACTOR-2026-04-17"
PLAN_UUID = str(uuid.uuid4())


# ── Plan content (single source of truth — used for both JSON and DOCX) ──

QUESTIONS_FOR_BOSS = [
    "File location: Human mentioned both 'Plan.json' and 'Business_Plan.json'. Plan.json does not exist; this plan was written to business_plan.json as a new record. If a separate plans.json file is preferred, human's call on Day 2.",
    "T011 was Claude-invented per human's memory — it is NOT. T011 exists in agile_backlog.json, source is 'human', and it specifies 50 messages or 30 seconds. Keep as-is, or revise on Day 2?",
    "Mongo batching claim: Claude overclaimed 'no cost increase' for smaller batches without checking MongoDB documentation. Research owed on Day 2 before any T011 change is proposed.",
    "B013 classification fields: Claude's proposed minimum set is {topic, intent, sensitivity, task_type}. Human edits on Day 2.",
    "B001 'thinking indicators' refinement: Claude's proposed v1 scope is `<system-reminder>` tag-strip + hook additionalContext-strip. Human's term 'philosophizing' may require a broader semantic definition — human's wording required on Day 2.",
    "Authorize inspection + retirement of conversation_log_purge_service.py and both conversation_log_agent.py copies (tools/ and tools/skill/) as part of this refactor? Human said 'sweep' in session — this question confirms in writing.",
]

EXECUTIVE_SUMMARY = (
    "Retire the orphaned conversation_log_hook.py file. Port content rules (binary "
    "handling, thinking-indicator stripping) into the Kafka consumer as part of a single "
    "unified AI call that also returns dialogue classification — turning the refactor "
    "into a business-value delivery (queryable dialogue DB), not just cleanup. Refactor "
    "producer and consumer Python files to class-based Main → run() pattern per human's "
    "architectural preference. Add a T requirement for Kafka broker TLS (currently "
    "PLAINTEXT, out of alignment with v4-029 spirit). Codify 'delete discarded code' as "
    "new engineering rule v4-037. All work across 4 days targeting completion Monday "
    "2026-04-20, preserving FindCare alpha priority via background-agent parallelism on "
    "Day 3."
)

SCHEDULE = [
    ("1", "Fri 2026-04-17", "Plan (this document). CTO review of scope.", "done"),
    ("2", "Sat 2026-04-18", "Requirements refinement — draft 7 artifacts, CTO review, write to agile_backlog.json + engineering_rules.json. MongoDB batching research completed before T011 change.", "pending"),
    ("3", "Sun 2026-04-19", "Foreground: Human + Claude on FindCare alpha. Background: subagent implements log-service refactor from Saturday's artifacts. EOD: Claude reviews agent diffs, runs pytests + full SIT, commits if green.", "pending"),
    ("4", "Mon 2026-04-20", "UAT — human in UAT-tester hat. Sign-off or rejection.", "pending"),
]

PROBLEMS = [
    ("P1", "Truncation violates B001", "conversation_log_hook.py:68-69 truncates content at 500 chars. B001 at agile_backlog.json:56471 says 'No truncation, no summarization.'", "Dead code — killed by hook deletion."),
    ("P2", "Binary handling has no backlog requirement", "Hook strips binary/base64, but no B or T requirement captures this. Human signed off verbally; not memorialized.", "Add B011."),
    ("P3", "Thinking-indicator stripping is imprecise", "B001 says 'Thinking indicators excluded' without defining scope. Hook's `<system-reminder>` strip is one interpretation; new Kafka path has no equivalent.", "Refine B001 wording with human-supplied definition."),
    ("P4", "conversation_log.json on disk is vestigial", "258 KB, last written 2026-04-16 18:02, not updated during active sessions (old hook is orphan). Raw unredacted content by design of the old path.", "Scrub or delete, after confirming Mongo path covers all reads."),
    ("P5", "Consumer redaction fail-open (SECURITY)", "conversation_log_consumer.py:106-108 catches redaction exceptions and stores UNREDACTED content. Silent gap.", "Defer — new deferred requirement B012. Performance cost + no DLQ-managing agent yet. Out of scope for this refactor."),
    ("P6", "Orphan hook file remains", "conversation_log_hook.py not wired in .claude/settings.json; referenced only by devops_boot.py:427 governance pattern list, test_prompt_log.py (imports), agile_backlog.json:56450 realized_by pointer.", "Delete per RISK-011 + new v4-037."),
    ("P7", "Producer and consumer are procedural, not class-based", "human's architectural standard: Main → class construction → run(), config via constructor. conversation_log_purge_service.py follows it; producer and consumer do not.", "Refactor both to the class pattern."),
    ("P8", "Kafka broker is PLAINTEXT", "docker-compose.yml uses PLAINTEXT:// listeners; no TLS. No T requirement currently covers Kafka transport security.", "Add T012 (new)."),
]

SOLUTIONS = [
    ("S1", "File a bug record", "'Old hook truncates at 500 chars, violating B001.' Resolution closes with hook deletion via RISK-011."),
    ("S2", "Add requirement B011", "Binary / base64 content replaced with '[binary content omitted]' before entering Kafka. Captures human's prior sign-off."),
    ("S3", "Refine B001", "Enumerate 'thinking indicators' concretely. Human's wording required."),
    ("S4", "Port content rules into consumer", "Integrate binary detection AND thinking-indicator detection into the single GPT-4.1-mini call per v4-034 (no regex/string matching for semantic work)."),
    ("S5", "Add requirement B013", "Single AI call returns redacted content + classification fields ({topic, intent, sensitivity, task_type} proposed; human edits). Creates queryable dialogue DB — the business justification for the refactor."),
    ("S6", "Add requirement T012", "Kafka broker MUST use TLS 1.2+ (and mTLS if confluent-kafka-python client supports, which it does). Reuse existing local CA under Code/Shared/ops/certs/ to issue kafka.crt/key."),
    ("S7", "Add deferred requirement B012", "Consumer redaction MUST fail loud (dead-letter queue, not silent fallback). Deferred until a DLQ-managing agent exists. Reason captured in backlog."),
    ("S8", "Refactor producer + consumer to class pattern", "Main → ClassName() → .run(). Config via constructor. Target: 1 C# file + 2 Python files for the whole service."),
    ("S9", "Retire orphan Python files", "Delete conversation_log_hook.py, conversation_log_purge_service.py, conversation_log_agent.py (both copies) — all replaced by Kafka path. RISK-011 covers the hook; human's session authorization covers the sweep."),
    ("S10", "Update callers", "Remove pattern from chathealthy_devops_boot.py:427; update test_prompt_log.py (the real regression tests move to the consumer); update agile_backlog.json:56450 realized_by."),
    ("S11", "Scrub conversation_log.json on disk", "After confirming Mongo path covers all reads. Flag: boot currently reads this file; may need to switch to Mongo source (separate scope question for Day 2)."),
    ("S12", "Add engineering rule v4-037", "'Orphan / discarded code MUST be deleted as part of the same task where its deadness is identified. Deletion via git commit with message naming (a) what replaced it, (b) evidence of safety (no callers, successor proven), (c) any backlog/test references updated.' Complement to v4-031."),
    ("S13", "Add engineering rule v4-038", "'Claude MUST NOT ship code that fails to meet any B* or T* backlog requirement without a risk acceptance record authorized by human, named in the commit message.' Memorializes human's mandate from the Day 1 session."),
]

REQUIREMENTS_DELTA = [
    ("New", "BUG-LOG-00X (next in sequence)", "Old hook violates B001 via 500-char truncation.", "must-have"),
    ("New", "B011", "Binary / base64 content replaced with '[binary content omitted]' before entering Kafka.", "must-have"),
    ("Refine", "B001", "Add concrete enumeration of 'thinking indicators.' human-supplied wording.", "must-have"),
    ("New deferred", "B012", "Consumer redaction fail-loud. Deferred: awaiting DLQ-managing agent.", "must-have / deferred"),
    ("New", "B013", "Single AI call returns redacted content + dialogue classification.", "must-have"),
    ("New", "T012", "Kafka broker uses TLS 1.2+ / mTLS. Reuse local CA.", "must-have"),
    ("Revisit", "T011", "Batch sizing 50 messages / 30 seconds. Claude overclaimed 'no cost increase' for smaller batches; MongoDB research owed Day 2 before any change.", "must-have (existing)"),
    ("New engineering rule", "v4-037", "Delete discarded code in the same task where its deadness is identified.", "standard"),
    ("New engineering rule", "v4-038", "Requirement deviation requires human risk acceptance, named in commit message.", "standard"),
]

RISKS = [
    ("Thinking-indicator false positive (strips legit content)", "Low", "Conservative tag-matching for structural markers; AI call only for classification, not for strip-decisions."),
    ("Binary detection false positive (strips legit short base64, e.g. SVG)", "Low", "Length threshold; specific prefixes (data:, explicit base64,)."),
    ("Consumer fail-loud causes observable crashes during OpenAI outage", "Deferred out of scope", "B012 deferred until DLQ-managing agent exists."),
    ("Meeting a B or T requirement is missed during refactor", "Low (human's own framing)", "Pytest each B*/T* maps to a green test; any deviation triggers v4-038 risk acceptance before commit."),
    ("MongoDB batch sizing change made without research", "Owed", "Day-2 research pre-dates any T011 change. If research confirms smaller batches are costly at scale, T011 stays at 50/30s."),
    ("In-flight Kafka messages ingested pre-refactor don't have new rules", "Low", "Drain through current consumer; new rules apply to messages written after deploy."),
]

TARGET_ARCHITECTURE = {
    "csharp_files": ["Worker.cs (unchanged — T005 supervisor, kills children on stop)"],
    "python_files": [
        "conversation_log_producer.py — class ProducerSidecar; __main__ constructs and calls .run()",
        "conversation_log_consumer.py — class ConversationLogConsumer; __main__ constructs and calls .run()",
    ],
    "deleted_python_files": [
        "conversation_log_hook.py (orphan — pre-Kafka)",
        "conversation_log_purge_service.py (replaced by Kafka consumer)",
        "conversation_log_agent.py (both copies — inspect first, delete if replaced by consumer; otherwise defer)",
    ],
    "ai_call_shape": (
        "Single GPT-4.1-mini call per utterance with structured output: "
        "{ redacted_content, binary_replaced, thinking_indicators_stripped, classification: { topic, intent, sensitivity, task_type } }. "
        "Deterministic format-level tag-strip may precede the AI call for <system-reminder> (v4-034 permits — structural, not intent)."
    ),
}

SIGNOFF_CRITERIA = [
    "Every B*/T* requirement for BRAIN-CONVERSATION-LOG has a passing pytest including negative scenarios (DR-024).",
    "End-to-end trace: new prompt → Kafka → consumer → Mongo. Show one record with: verbatim content (B001), PST+UTC timestamps + ch_key (B002), redacted credentials and profanity (B004), binary replaced (B011), thinking-indicators excluded per refined B001, classification fields populated (B013).",
    "Force a redaction API failure — show message held in Kafka (current behavior; DLQ is B012 deferred).",
    "Kafka broker on TLS — show broker config and producer/consumer successfully talking mTLS.",
    "RISK-011 referenced in every commit on the relevant branch.",
    "Bug record closed with reference to fixing commit.",
    "Pre-push scanner green. No bypass (v4-036).",
    "v4-037 and v4-038 added to engineering_rules.json and enforced on subsequent commits.",
]

NOT_IN_SCOPE = [
    "Kafka UI (separate earlier proposal; out of this refactor).",
    "conversation_log_purge_service.py features beyond retirement (the file goes; its logic doesn't carry forward).",
    "Changes to Windows Service supervisor (Worker.cs) beyond those mandated by architecture target.",
    "Kafka broker config beyond TLS enablement (no partition count change, no retention change, etc.).",
    "Porting TR11 truncation — it was the bug; its absence is the fix.",
    "Any deploy to QA or production (v4-006).",
]

REFERENCES = [
    ("RISK-011", "brain/machine_artifacts/content/risk_acceptance.json", "Authorizes orphan-hook deletion + v4-037 addition. Low risk: git repo preserves deleted content."),
    ("BRAIN-CONVERSATION-LOG story", "agile_backlog.json lines ~56434-56977", "All B001-B010 and T001-T011 requirements for the conversation log."),
    ("v4-022 alpha freeze", "engineering_rules.json", "Acknowledged tension; human's explicit authorization via RISK-011 supersedes per RISK-004/005/006 precedent."),
    ("v4-031 no dead code", "engineering_rules.json", "v4-037 complements this with file-level deletion discipline."),
    ("v4-034 AI-native intent classification", "engineering_rules.json", "Governs the no-regex-for-semantics rule. Redaction and classification go through the AI; structural tag-strip is allowed."),
    ("Feedback memories", "C:/Users/skips/.claude/projects/c--chatHealthy-findCare/memory/", "feedback_no_lying, feedback_no_regressions, feedback_meaningful_pytests, feedback_oo_discipline, feedback_prefer_oss, feedback_push_back — all applied in this plan."),
]


# ── 1. Append record to business_plan.json ──────────────────────────────────

def append_plan_record():
    with BIZ_PLAN.open(encoding="utf-8") as f:
        data = json.load(f)

    record = {
        "record_type": "plan",
        "plan_id": PLAN_ID,
        "title": "Conversation Log Service Refactor — 4-Day Plan",
        "date_authored": "2026-04-17",
        "authored_by": "Claude (Opus 4.7) under human CTO review",
        "status": "day_1_complete",
        "approval": "awaiting_day_2_cto_review",
        "risk_acceptances_referenced": ["RISK-011"],
        "artifact_docx": str(OUT_DOCX.relative_to(REPO_ROOT)).replace("\\", "/"),
        "questions_for_boss": QUESTIONS_FOR_BOSS,
        "executive_summary": EXECUTIVE_SUMMARY,
        "schedule": [
            {"day": d, "date": dt, "work": w, "status": s}
            for d, dt, w, s in SCHEDULE
        ],
        "problems": [
            {"id": pid, "title": t, "evidence": e, "disposition": d}
            for pid, t, e, d in PROBLEMS
        ],
        "solutions": [
            {"id": sid, "title": t, "detail": d}
            for sid, t, d in SOLUTIONS
        ],
        "requirements_delta": [
            {"kind": k, "id": rid, "detail": d, "priority": p}
            for k, rid, d, p in REQUIREMENTS_DELTA
        ],
        "risks": [
            {"risk": r, "severity": s, "mitigation": m}
            for r, s, m in RISKS
        ],
        "target_architecture": TARGET_ARCHITECTURE,
        "signoff_criteria": SIGNOFF_CRITERIA,
        "not_in_scope": NOT_IN_SCOPE,
        "references": [
            {"id": rid, "location": loc, "note": n}
            for rid, loc, n in REFERENCES
        ],
        "mongo_batching_research_owed": (
            "Day 2 — verify MongoDB documentation on insert_many batch sizing at ChatHealthy's "
            "traffic profile before proposing any change to T011 (50 messages / 30 seconds). "
            "human's pushback on Claude's unverified 'no cost increase' claim is correct."
        ),
        "ch_matrix_id": PLAN_UUID,
    }

    data.setdefault("records", []).append(record)
    data["record_count"] = len(data["records"])

    with BIZ_PLAN.open("w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)

    print(f"[business_plan.json] appended {PLAN_ID} (total records: {data['record_count']})")


# ── 2. Generate Word doc ────────────────────────────────────────────────────

def _set_font(run, size=11, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def generate_docx():
    doc = Document()

    # Title
    title = doc.add_heading("Conversation Log Service Refactor — 4-Day Plan", level=0)

    meta = doc.add_paragraph()
    meta.add_run("Plan ID: ").bold = True
    meta.add_run(PLAN_ID + "    ")
    meta.add_run("Date: ").bold = True
    meta.add_run("2026-04-17 (Day 1)    ")
    meta.add_run("Author: ").bold = True
    meta.add_run("Claude (Opus 4.7) under human CTO review")

    # Questions at top (per human's instruction)
    h = doc.add_heading("Open questions for human — please read first", level=1)
    for r in h.runs:
        _set_font(r, size=14, bold=True, color=RGBColor(0xB0, 0x00, 0x00))
    doc.add_paragraph(
        "Per session directive: questions raised here rather than returned to Claude in chat. "
        "Your answers drive Day 2 artifact drafting."
    )
    for i, q in enumerate(QUESTIONS_FOR_BOSS, 1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(q)

    # Executive summary
    doc.add_heading("Executive summary", level=1)
    doc.add_paragraph(EXECUTIVE_SUMMARY)

    # Schedule
    doc.add_heading("4-Day schedule", level=1)
    t = doc.add_table(rows=1, cols=4)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(["Day", "Date", "Work", "Status"]):
        hdr[i].text = h
    for day, date, work, status in SCHEDULE:
        row = t.add_row().cells
        row[0].text = day
        row[1].text = date
        row[2].text = work
        row[3].text = status

    # Problems
    doc.add_heading("Problems to solve", level=1)
    t = doc.add_table(rows=1, cols=4)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(["#", "Problem", "Evidence", "Disposition"]):
        hdr[i].text = h
    for pid, title, evidence, disp in PROBLEMS:
        row = t.add_row().cells
        row[0].text = pid
        row[1].text = title
        row[2].text = evidence
        row[3].text = disp

    # Solutions
    doc.add_heading("Solutions", level=1)
    for sid, title, detail in SOLUTIONS:
        p = doc.add_paragraph()
        p.add_run(f"{sid} — {title}. ").bold = True
        p.add_run(detail)

    # Target architecture
    doc.add_heading("Target architecture", level=1)
    doc.add_paragraph(
        "human's standard: 1 C# file + 1 Python producer + 1 Python consumer. "
        "Main instantiates a class and calls .run(); configuration flows via the constructor."
    )
    p = doc.add_paragraph()
    p.add_run("C# files: ").bold = True
    p.add_run(", ".join(TARGET_ARCHITECTURE["csharp_files"]))
    p = doc.add_paragraph()
    p.add_run("Python files (target state): ").bold = True
    for line in TARGET_ARCHITECTURE["python_files"]:
        doc.add_paragraph(line, style="List Bullet")
    p = doc.add_paragraph()
    p.add_run("Python files to delete: ").bold = True
    for line in TARGET_ARCHITECTURE["deleted_python_files"]:
        doc.add_paragraph(line, style="List Bullet")
    p = doc.add_paragraph()
    p.add_run("Unified AI call shape: ").bold = True
    p.add_run(TARGET_ARCHITECTURE["ai_call_shape"])

    # Requirements delta
    doc.add_heading("Requirements delta (for Day 2 drafting)", level=1)
    t = doc.add_table(rows=1, cols=4)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(["Kind", "ID", "Detail", "Priority"]):
        hdr[i].text = h
    for k, rid, detail, pri in REQUIREMENTS_DELTA:
        row = t.add_row().cells
        row[0].text = k
        row[1].text = rid
        row[2].text = detail
        row[3].text = pri

    # Risks
    doc.add_heading("Risks", level=1)
    t = doc.add_table(rows=1, cols=3)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(["Risk", "Severity", "Mitigation"]):
        hdr[i].text = h
    for risk, sev, mit in RISKS:
        row = t.add_row().cells
        row[0].text = risk
        row[1].text = sev
        row[2].text = mit

    # MongoDB research note
    doc.add_heading("MongoDB batching research — owed Day 2", level=1)
    doc.add_paragraph(
        "Claude's claim 'no cost increase' for smaller Mongo batches was not grounded in "
        "MongoDB documentation. Human pushed back correctly. Before any change to T011 (currently "
        "50 messages / 30 seconds, source: Human), Claude will research MongoDB's stated best practice "
        "for insert_many batch sizing at ChatHealthy's traffic profile (low-volume, latency-sensitive). "
        "T011 stays as-is unless research supports a change AND human approves."
    )

    # Sign-off criteria
    doc.add_heading("CTO sign-off criteria (Day 4 UAT)", level=1)
    for crit in SIGNOFF_CRITERIA:
        doc.add_paragraph(crit, style="List Bullet")

    # Not in scope
    doc.add_heading("Not in scope", level=1)
    for item in NOT_IN_SCOPE:
        doc.add_paragraph(item, style="List Bullet")

    # References
    doc.add_heading("References", level=1)
    for rid, loc, note in REFERENCES:
        p = doc.add_paragraph()
        p.add_run(f"{rid} — ").bold = True
        p.add_run(f"{loc}. ").italic = True
        p.add_run(note)

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer_run = footer.add_run(
        f"Generated {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} "
        f"by {Path(__file__).name}. Plan record also persisted to "
        "brain/machine_artifacts/content/business_plan.json."
    )
    footer_run.italic = True
    _set_font(footer_run, size=9)

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_DOCX))
    print(f"[docx] wrote {OUT_DOCX}")


if __name__ == "__main__":
    append_plan_record()
    generate_docx()
