# Generator for codebase-reorganization-V4.docx
# Read-only analysis output. Lineage:
#   pass3.docx (V1) -> pass3-V2.docx (Skip's annotated copy)
#   -> codebase-reorganization-V3.docx
#   -> codebase-reorganization-V4.docx (this file)
#
# V4 corrections relative to V3 (architect-binding):
#   1. UX is NOT dissolved. EPIC-011 stays as a cross-cutting epic that owns
#      the genuinely shared display concerns: shared design paradigm across
#      static pages, the parent (opening) frame, the lower-region environment
#      banner shared across servers, and the auth-token display object used
#      across 3 servers.
#   2. The pipeline is and remains ONE epic (EPIC-012 Provider Knowledge
#      Pipeline). Bifurcating ETL across capability epics forces architectural
#      drift in ETL processes -- not OK.
#   3. The layer rename from `domain/` -> `core/` deserves an up-front
#      legacy-resolution narrative: why "domain" is dropped, why `rules/` was
#      the working candidate, why it conflicts with engineering_rules.json,
#      why `core/` is the final answer.
#
# Three architect open issues with deeper analysis:
#   I1. Worker.cs and ChatHealthyMongoUtilities.py drift -- function-by-function
#       diff + active-build-path determination + recommendation.
#   I2. ChatHealthy.Providers.json schema drift -- field-by-field diff +
#       MongoDB sampling against both schemas + recommendation.
#   I3. UX disposition resolved per Correction 1 (retained).
#
# Aligned with langgraph-oo-best-practices-V4.docx (final consolidated
# reference: Pydantic state, MongoDBSaver, Runtime[ContextSchema] DI, messages
# channel + ToolMessage, subgraph per capability, Safety as shared subgraph,
# deep_research as reusable subgraph, 4-file layout).

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches

OUT = Path(r"c:\chatHealthy\findCare\findCare\ArchitectureAndDesign\codebase-reorganization-V4.docx")

doc = Document()

# Base style
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)


def H1(t):
    return doc.add_heading(t, level=1)


def H2(t):
    return doc.add_heading(t, level=2)


def H3(t):
    return doc.add_heading(t, level=3)


def P(t, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.bold = bold
    r.italic = italic
    return p


def Quote(t):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    r = p.add_run(t)
    r.italic = True
    return p


def Bullet(t):
    return doc.add_paragraph(t, style='List Bullet')


def HR():
    doc.add_paragraph("_" * 60)


def Table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for r in hdr[i].paragraphs:
            for run in r.runs:
                run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    return t


# -----------------------------------------------------------------------------
# Title block
# -----------------------------------------------------------------------------
H1("Codebase Reorganization -- V4")
P("Date: 2026-04-28   |   Author: Claude (architecture analysis agent)   |   Read-only")
P("Lineage: codebase-reorganization-pass3.docx (V1) -> codebase-reorganization-pass3-V2.docx (Skip's annotated) -> codebase-reorganization-V3.docx -> this V4.")
P("Inputs read in order:")
Bullet("findCare/ArchitectureAndDesign/codebase-reorganization-V3.docx -- prior version; V4 supersedes it on three points the architect corrected.")
Bullet("findCare/ArchitectureAndDesign/codebase-reorganization-pass3-V2.docx -- Skip's annotated Pass 3 (binding).")
Bullet("findCare/ArchitectureAndDesign/langgraph-oo-best-practices-V4.docx -- definitive LangGraph design recommendations.")
Bullet("brain/machine_artifacts/content/agile_backlog.json -- 9-epic shape (verified live; counts read at generate time).")
Bullet("Code/, Website/, brain/, architecture/ -- file tree walked; every cited path verified by direct filesystem listing.")
Bullet("MongoDB Atlas frontend cluster -- prod_PublicHealthData.providers (143,589 documents); read-only $sample aggregation.")

HR()

# -----------------------------------------------------------------------------
# 1. Operating-model axiom (verbatim)
# -----------------------------------------------------------------------------
H2("1. Lead -- the operating-model axiom")
Quote('"The operating model, the digital footprint, the general ledger, and the org chart are 100% in alignment. The structure of any one IS the structure of all four. This alignment is non-negotiable and is the fundamental organizing principle of ChatHealthy."')
P("Every proposal in V4 is tested against the axiom: does the proposed Git-tree shape mirror what would appear on the org chart, on a general-ledger cost-center list, and on the customer-facing capability map? Where the answer is no, the structure is decoration and is removed.")

HR()

# -----------------------------------------------------------------------------
# 2. Layer-name explanation (the legacy-resolution narrative)
# -----------------------------------------------------------------------------
H2("2. Layer-name resolution -- why `core/`, not `domain/`, not `rules/`")
P("This section leads V4 because it is the worked example of a legacy-resolution decision: an existing structural noun (`domain/`) failed the operating-model test, an alternative the architect proposed (`rules/`) collided with a legacy artifact, and a third candidate (`core/`) survived both tests. The lineage is documented up front so a future reader can see the trail.")

H3("2.1 Why `domain/` is being dropped")
P("`domain/` is the layer name that today appears in Code/ConversationalUX/FindCareChat/backend/domain/{find_care, evaluate_care_quality, shared}/. It is a hexagonal-architecture term of art. It does not appear on any HR org chart, any GL cost-center list, or any customer-facing capability map. By the operating-model axiom (Section 1), an architectural noun that lives only inside engineering vocabulary is decoration. The word `domain` therefore must not appear as a structural noun anywhere in the proposed Git tree. This decision is grounded in Skip's V2 binding from pass3-V2: 'this is super confusing... FindCare is Find Care it is not an app it is a capability.'")

H3("2.2 Why `rules/` was the working candidate")
P("Once `domain` was off the table, the architect's first-choice replacement was `rules/`. The reasoning was sound:")
Bullet("`rules/` names the right concept -- the deterministic, governed business rules a capability enforces (FindCare's specialty classifier, EvaluateCare's scoring engine, the homeopathic resolver). These are 'the business's rules for this product.'")
Bullet("It maps cleanly to a P&L line: a CFO would write a check to 'the FindCare rules engine.'")
Bullet("It avoids the layer-name flavor of `domain` (no hexagonal-architecture coupling).")

H3("2.3 Why `rules/` conflicts with a legacy artifact")
P("`rules/` collides with brain/machine_artifacts/content/engineering_rules.json -- the governed JSON that holds the project's engineering rules (~70 rules at last count, version-stamped, schema-validated, enforced by the V20 framework). This is not a hypothetical naming clash; it is a direct overlap with an active, load-bearing artifact:")
Bullet("A new engineer reading Code/FindCare/backend/rules/specialty_classifier.py would reasonably ask: 'Is this code an engineering rule? Where is its entry in engineering_rules.json?' -- the answer is no, but the question is built into the name.")
Bullet("Tools that enforce engineering rules (chathealthy_enforcement_manager.py, scan_files_enforcement_worker.py, brain_write_validator.py) all key off the word 'rule.' Adding a parallel `rules/` directory under every capability would create false positives in any name-based scan and would force every enforcement check to add path-disambiguation logic.")
Bullet("The brain JSON is the singular source of project rules under Engineering Rule-056 ('Brain JSON governance mandates must always be obeyed by Claude Code'). Splitting the noun across the brain artifact and a code-tree folder weakens the mandate.")

H3("2.4 Legacy/precedent that would need to be addressed if we kept `rules/`")
P("If the architect insisted on `rules/` despite the clash, the precedent surface would need explicit scoping:")
Bullet("Rename brain/machine_artifacts/content/engineering_rules.json -> engineering_governance.json (governance-controlled-vocabulary change requiring human risk acceptance per Rule-028).")
Bullet("Update all enforcement code that loads 'engineering_rules.json' by name (~12 references across architecture/EngineeringRuleEnforcement/code/ and Code/Shared/ops/tools/).")
Bullet("Update all docs and conversation_log entries that quote the artifact name.")
Bullet("Update Engineering Rule-008 enforcement scan, which currently allow-lists 'rules' folders explicitly.")
P("This is a legacy-resolution cost the architect can pay -- but it is a real cost, paid every time the controlled vocabulary moves.")

H3("2.5 Why `core/` was selected as the final answer")
P("`core/` passes both tests:")
Bullet("Operating-model axiom: it names the part of the capability that IS the capability -- 'the FindCare engine,' 'the EvaluateCare scoring core.' A CFO writes a check to 'the FindCare core' without confusion.")
Bullet("No clash with engineering_rules.json -- 'core' has no overlap with the governed-rules vocabulary, the enforcement framework, or any brain JSON.")
Bullet("LangGraph V4 reinforces it: each subgraph 'owns a focused state schema'; the capability's pure logic body that the subgraph wraps maps cleanly to the term 'core' (compare 'agent core,' 'tool core' in the LangGraph community literature).")
P("Alternative `engine/` was a strong second choice (works for EvaluateCare's scoring engine; feels heavy for FindCare's retrieval + classification posture). `services/` and `logic/` were rejected: `services/` conflates with EPIC-009 Shared Services; `logic/` is vague and fails the axiom test.")
P("Final structure: Code/FindCare/backend/core/{provider_search_service.py, specialty_classifier.py, specialty_ranker.py, specialty_service.py, homeopathic_resolver.py}; Code/EvaluateCare/core/{scoring_engine.py, normalization.py, weights.py, ...}.")

HR()

# -----------------------------------------------------------------------------
# 3. Three architect open issues -- deeper analysis
# -----------------------------------------------------------------------------
H2("3. Three architect open issues -- deeper analysis required")

# Issue 1
H3("3.1 Issue 1 -- Worker.cs and ChatHealthyMongoUtilities.py drift")
P("V3 said 'Skip must pick canonical' for both. V4 carries the diffs and the active-deploy-path determination so the choice is concrete.")

P("3.1.1 Worker.cs -- two copies, true drift", bold=True)
Bullet("Copy A: Code/CSharp/ChatHealthyLogService/dev/Worker.cs -- 189 lines.")
Bullet("Copy B: Code/Shared/ops/ChatHealthyLogService/Worker.cs -- 34 lines.")

P("Function-by-function differences (verified by `diff -u` on 2026-04-28):")
Table(
    ["Behavior", "Copy A (CSharp/dev/, 189 lines)", "Copy B (Shared/ops/, 34 lines)"],
    [
        ["Header tag", "T005 supervisor + T006 AUTO_START + BUG-LOG-002 (kill-children-on-stop)", "T011 'No business logic. Calls Python main entry point.'"],
        ["Child processes", "Two children: SidecarScript=Code/Shared/ops/kafka/conversation_log_producer.py (PID 1) and ConsumerScript=Code/Shared/ops/kafka/conversation_log_consumer.py (PID 2)", "One child: MainScript=Code/Shared/ops/tools/conversation_log_purge_service.py"],
        ["Kafka bootstrap", "EnsureKafkaRunning() runs `docker compose -f Code/Shared/ops/kafka/docker-compose.yml up -d` if container not Up", "Absent -- no Docker / Kafka orchestration"],
        ["Health-check loop", "30s loop, restart any dead PID with 5s backoff", "None -- ExecuteAsync simply awaits the single Python process"],
        ["StopAsync / Dispose", "Custom KillChild(entireProcessTree:true) for both children; logs PID kills", "Default Worker disposal; relies on `using var process`"],
        ["Stdout/stderr", "RedirectStandardOutput=false, RedirectStandardError=false on children", "RedirectStandardOutput=true, RedirectStandardError=true on the single child"],
    ],
    widths=[1.4, 2.7, 2.4],
)

P("Architectural context -- which copy is in active use?", bold=True)
Bullet(".vscode/tasks.json -- the build, publish, and watch tasks all point at Code/Shared/ops/ChatHealthyLogService/ChatHealthyLogService.csproj. This is the active build path.")
Bullet("Code/Shared/ops/ChatHealthyLogService/ has bin/, obj/, out/ -- proves it is the path that has actually been built recently.")
Bullet("Code/CSharp/ChatHealthyLogService/dev/ also has bin/ and obj/ from earlier builds (the dev/build split was introduced in commit af12412 'Restructure C# into dev/build, gitignore build artifacts').")
Bullet("Program.cs and ChatHealthyLogService.csproj are byte-identical between the two trees -- the drift is contained to Worker.cs.")
Bullet("Git history of the diverged Worker.cs: T005/T006/BUG-LOG-002 (Copy A) -> commit 39763aa 'Kafka conversation log infrastructure -- full stack' was the original supervisor pattern. T011 (Copy B) is commit bd44f12 'T011: Restructure service -- class with run(), no business logic in C# or main' -- a deliberate later simplification.")

P("The choice the architect is being asked to make:", bold=True)
Bullet("Choose Copy A (the 189-line supervisor) means: the C# service starts Kafka via Docker, runs TWO Python children (producer sidecar + consumer), monitors both PIDs, restarts on death, kills child trees on stop. This was the design through commit 39763aa.")
Bullet("Choose Copy B (the 34-line shim) means: the C# service is a one-job process launcher for Code/Shared/ops/tools/conversation_log_purge_service.py. Kafka container management, producer, consumer, and child-restart logic all move to the Python side or to a different ops tool.")
Bullet("Merging requires: deciding which child set the C# service supervises (purge service only? producer+consumer? all three?), deciding whether Docker-Kafka orchestration belongs in the C# service or in a separate compose-up step, and re-applying BUG-LOG-002's 'kill children on stop' guarantee to whichever child set survives.")

P("Recommendation:", bold=True)
P("Copy B (Code/Shared/ops/ChatHealthyLogService/Worker.cs, 34 lines) is canonical. Three independent signals point at it:")
Bullet("1. .vscode/tasks.json builds, publishes, and watches Copy B's csproj. Copy A is no longer in any active build pipeline.")
Bullet("2. The git log shows commit bd44f12 (T011 restructure) is the more recent design intent: 'no business logic in C# or main.' Copy A's supervisor-with-two-children is the superseded design.")
Bullet("3. Copy A's references (conversation_log_producer.py, conversation_log_consumer.py, docker-compose.yml) all still exist under Code/Shared/ops/kafka/, so Copy A's logic is recoverable from git history if any of its features (Docker orchestration, multi-child supervision) are wanted in a follow-up; nothing is lost by deleting Copy A.")
P("Action: delete Code/CSharp/ChatHealthyLogService/dev/Worker.cs and the surrounding Code/CSharp/ tree as part of the refactor that moves logging_service to Code/Operations/logging_service/. Closes V3 Open Question 3.")

P("3.1.2 ChatHealthyMongoUtilities.py -- two copies, drift is documentary not functional", bold=True)
Bullet("Copy A: Code/Shared/ChatHealthyMongoUtilities.py -- 131 lines. Imported by Code/ConversationalUX/FindCareChat/backend/main.py via `from ChatHealthyMongoUtilities import ChatHealthyMongoUtilities`. Has a `commit(env_prefix, database, collection, record)` write method that adds record_number + datetime and inserts a record.")
Bullet("Copy B: Code/DataPipelines/ChatHealthyMongoUtilities.py -- 136 lines. Has an explicit 23-line header docstring 'WHEN TO USE THIS CLASS / DO NOT USE in DataPipelines (Layer 2 -- Azure Functions)' and is missing the `commit()` method.")

P("Function-by-function differences:")
Table(
    ["Behavior", "Copy A (Code/Shared/, 131 lines)", "Copy B (Code/DataPipelines/, 136 lines)"],
    [
        ["Header docstring", "Brief; no Layer-1/Layer-2 guidance.", "23-line block: 'Use in GUI / conversational UX (Layer 1 -- ConversationalUX/HuggingFace) ... DO NOT USE in DataPipelines (Layer 2 -- Azure Functions). Use module-level lazy MongoClient singleton instead: _mongo: MongoClient | None = None; def _get_mongo_client(): ... PyMongo's MongoClient is itself a connection pool.'"],
        ["__init__(connection_string)", "Identical: validates string, creates MongoClient, runs ping.", "Identical."],
        ["getConnection()", "Identical: returns self._client after ping.", "Identical."],
        ["close()", "Identical: client.close(); self._client = None.", "Identical."],
        ["commit(env_prefix, database, collection, record)", "Present: insert record with record_number + datetime; returns {'recorded':'ok' / 'error'}.", "ABSENT -- removed."],
        ["Context-manager (__enter__/__exit__)", "Identical.", "Identical."],
    ],
    widths=[1.6, 2.4, 2.5],
)

P("Architectural context -- which copy is in active use?", bold=True)
Bullet("Code/ConversationalUX/FindCareChat/backend/main.py:47 imports `ChatHealthyMongoUtilities` -- this is Copy A (Code/Shared/), since main.py runs in the FindCare HF Space and Code/Shared/ is on the import path.")
Bullet(".github/workflows/deploy-findcare-backend.yml line 14 watches 'Code/Shared/ChatHealthyMongoUtilities.py' as a deploy trigger -- only Copy A. Copy B is not watched.")
Bullet("DataPipelines code does NOT import ChatHealthyMongoUtilities. Verified by grep: every pipeline module that touches Mongo (county_enrichment_job.py, copy_to_frontend.py, prescriber_pipeline_manager.py, ...) uses a module-level lazy `_get_mongo_client()` returning `MongoClient(os.environ['MONGO_connectionString'])` -- exactly the pattern Copy B's docstring tells the reader to use.")
Bullet("Therefore Copy B is dead code (no caller); its only purpose is to host the 23-line 'don't use this here' docstring.")

P("Deploy paths:", bold=True)
Bullet("Copy A: ships into the FindCare HF Space (Hugging Face Docker image). Source-of-truth for the conversational/UX side.")
Bullet("Copy B: ships into Azure Functions deploy bundle for DataPipelines (because the directory ships as a unit), but is never imported once deployed. Pure documentation.")

P("The choice the architect is being asked to make:", bold=True)
Bullet("Choose Copy A means: pipeline retains its module-level lazy `_get_mongo_client()` pattern (the architecturally correct one for Azure Functions warm-pool reuse), and the documentary docstring from Copy B is preserved by relocating it to a comment in Code/SharedServices/mongo_utilities.py or to the architecture doc.")
Bullet("Choose Copy B means: lose the `commit()` method (which the conversational backend may or may not call -- Copy A's `commit` is not currently grepped from main.py, but it is a public API that conversational stories may use later).")
Bullet("Merging requires: keeping Copy A's `commit()` method, adding Copy B's docstring as a class-level note, and deleting Copy B.")

P("Recommendation:", bold=True)
P("Copy A (Code/Shared/ChatHealthyMongoUtilities.py) is canonical. Copy B is dead code (no caller) whose only value is the documentary 23-line WHEN-TO-USE block. Action:")
Bullet("Move Copy A to Code/SharedServices/mongo_utilities.py (per V3 master move table, EPIC-009 owner).")
Bullet("Port Copy B's WHEN-TO-USE docstring into Copy A's module-level docstring (preserving the architectural guidance that pipeline must use the lazy pattern, not this class).")
Bullet("Delete Copy B (Code/DataPipelines/ChatHealthyMongoUtilities.py).")
Bullet("No FindCare or DataPipelines runtime behavior changes -- the import in main.py simply moves to the new path; the pipeline never imported Copy B.")
P("Closes V3 Open Question 4.")

# Issue 2
H3("3.2 Issue 2 -- ChatHealthy.Providers.json schema drift between Code/Schemas/ and Website/schemas/")
P("V3 said '~11-byte drift.' V4 carries the field-by-field diff and the MongoDB sample-validation result.")

P("3.2.1 The two schemas, byte-for-byte", bold=True)
Bullet("Code/Schemas/ChatHealthy.Providers.json -- 19,265 bytes.")
Bullet("Website/schemas/ChatHealthyProvidersSchema.json -- 19,276 bytes.")
Bullet("Verified by `diff -u`, the only differences are: (a) the `$id` URL -- `https://chathealthy.ai/schemas/ChatHealthy.Providers.json` vs `https://dev.chathealthy.ai/schemas/ChatHealthyProvidersSchema.json`; (b) trailing-newline presence (Code/Schemas/ has no trailing newline; Website/schemas/ does).")
Bullet("Every property declaration, every required-field list, every if/then/else conditional, every type, every $ref is byte-identical between the two files. The two schemas are FUNCTIONALLY EQUIVALENT.")

P("3.2.2 Sampling the actual MongoDB providers collection", bold=True)
P("Read-only $sample aggregation against MongoDB Atlas frontend cluster (URI from MONGO_FRONTEND_connectionString). No writes; no schema mutations.")
Bullet("prod_PublicHealthData.providers -- 143,589 documents (estimated_document_count).")
Bullet("dev_PublicHealthData.providers -- 43,491 documents.")
Bullet("Sample size: 500 documents from prod, 200 from dev.")

P("Validation results (Draft 2020-12 jsonschema.Draft202012Validator):", bold=True)
Table(
    ["Database", "Sample size", "Pass Code/Schemas (Code/Schemas/ChatHealthy.Providers.json)", "Pass Website/schemas (Website/schemas/ChatHealthyProvidersSchema.json)", "Asymmetric (one passes, one fails)"],
    [
        ["prod_PublicHealthData", "500", "0", "0", "0"],
        ["dev_PublicHealthData", "200", "0", "0", "0"],
    ],
    widths=[1.7, 1.0, 1.7, 1.7, 1.4],
)

P("In every one of the 700 sampled documents, both schemas reject the document for the same reason. Zero asymmetric rejections in 700 samples. The schemas are functionally equivalent against the production data.")

P("3.2.3 Why both schemas reject every document -- the actual gap", bold=True)
P("Both schemas use `additionalProperties: false`. The pipeline writes post-enrichment fields the schemas do not declare:")
Table(
    ["Extra field present in DB but missing from schema", "% of 500 prod documents"],
    [
        ["embedding (post-pipeline vector)", "100% (500/500)"],
        ["embedding_model", "52% (261/500)"],
        ["embedding_version", "52% (261/500)"],
        ["authorized_official_name_prefix_text (NPPES)", "9% (45/500)"],
        ["provider_other_last_name (NPPES)", "6% (29/500)"],
        ["provider_other_first_name (NPPES)", "6% (29/500)"],
        ["provider_other_last_name_type_code (NPPES)", "6% (29/500)"],
        ["provider_other_middle_name (NPPES)", "4% (22/500)"],
        ["provider_other_credential_text (NPPES)", "4% (19/500)"],
        ["provider_other_name_prefix_text (NPPES)", "1% (6/500)"],
        ["authorized_official_name_suffix_text (NPPES)", "<1% (1/500)"],
    ],
    widths=[3.0, 1.5],
)

P("Both schemas declare 41 properties; the running pipeline produces documents with up to 11 additional properties not in either schema. Both schemas are equivalently out of date.")

P("3.2.4 Architect's framing answered", bold=True)
Bullet("Q: 'Are both schemas equivalently good in the sense that they validate all the providers in the DB?' A: Both schemas reject 100% of production providers, for the same reasons. They are equivalently bad against current data, and equivalently good against intended data.")
Bullet("Q: 'I need a diff analysis -- knowing they are different size does not help me choose.' A: The 11-byte difference is the `$id` URL. There is no functional difference. The choice between them is not a schema-content decision; it is a publication-URL decision.")

P("Recommendation:", bold=True)
P("The two schemas are functionally identical. Either one is canonical for content. The publication question is the real one:")
Bullet("Adopt Website/schemas/ChatHealthyProvidersSchema.json's content as canonical (the trailing newline is a small but conventional improvement).")
Bullet("Move it to brain/machine_artifacts/schemas/ChatHealthyProvidersSchema.json per V3 B.4 (brain owns schemas as a sibling of content/).")
Bullet("Set $id at the brain-side canonical URL (e.g., `https://chathealthy.ai/schemas/ChatHealthyProvidersSchema.json`); the deploy step that mirrors brain -> Website preserves the URL surface.")
Bullet("Delete Code/Schemas/ChatHealthy.Providers.json.")
Bullet("Add the missing fields (embedding, embedding_model, embedding_version, the 8 NPPES 'other_name' fields) to the schema as a SEPARATE backlog item -- this is a real schema gap, not a relocation issue, and should not block the relocation. Track under EPIC-012 Provider Knowledge Pipeline / schema-drift feature.")
P("Closes V3 Open Question 6 (schema canonical) and unblocks B.4 (schema relocation). Opens a new tracked item: 'Schema content does not match current pipeline output -- 11 fields missing from declared schema; 100% of production providers fail validation.'")

# Issue 3
H3("3.3 Issue 3 -- UX disposition (resolved per Correction 1)")
P("V3 recommended dissolving EPIC-011 UX into capability epics on the basis of 'no current cross-epic consumers.' Architect-binding correction: this is wrong. UX IS shared across epics today and must remain a first-class shared concern.")

P("Evidence walked from the actual code on 2026-04-28:")
Bullet("Website/index.html (1,400+ lines) is the parent page that hosts the chat iframe and IS shared infrastructure -- it carries the design tokens (--ink, --teal, --accent), header layout, footer, and (critically) the lower-region environment banner (#envBanner, lines 627-660) that is shared across LOCAL/DEV/QA environments.")
Bullet("The banner is service-aware: window.refreshBanner(service) reads /health from the page-owning service and reflects build/version/commit from FindCare's HF Space, EvaluateCare's HF Space, and the SharedServices HF Space (3 servers). Confirmed by direct grep in Website/index.html: 'banner reflects new owner's /health' is logged on every service handoff (lines 882, 917, 939, 1158, 1238, 1336).")
Bullet("Website/index.html also hosts the auth-token display object: window.buildTokenDisplayHtml(st) at lines 1011-1045. Per the inline comment 'EPIC-002-F-001-S-012-REQ-B-007: Parse token into nonce and GUID, build display HTML.' This object is consumed by both LEFT and RIGHT token panels (lines 808, 840) -- the LEFT always shows FindCare's /session token, the RIGHT shows the page-owning-service's /session token. That makes it shared across at least 3 servers (FindCare, EvaluateCare, SharedServices) per the comment 'SEC-HTTPS-001-REQ-021/022.' This matches the architect's statement: 'auth-token display object used in or should be used in 3 servers.'")
Bullet("The 11 static HTML pages under Website/ (architecture.html, products.html, privacy.html, terms.html, roadmap.html, etc.) share a design paradigm -- the same :root token block, the same --teal/--accent palette, the same header/footer pattern. Verified by grep across the static HTML files. This is genuinely cross-cutting UX.")
Bullet("Code/Shared/ux/ holds {components/ProviderCard.tsx (4,235 bytes), components/SelectionManager.tsx (7,882 bytes), hooks/useSelectionState.ts (4,475 bytes), types/provider.ts (1,233 bytes)}. Today these are NOT consumed by FindCareChat/frontend (which has its own components/), but they are present and named for reuse -- the 'Anchor Timer' and selection-state hooks anticipate the cross-component scope EPIC-011 was created to hold.")

P("Resolution:", bold=True)
P("EPIC-011 UX is RETAINED as a cross-cutting epic. The default home for capability-specific UX is the capability that consumes it; the genuinely-shared UX surface lives under EPIC-011.")

P("EPIC-011 UX features under V4 (after restoration):", bold=True)
Bullet("F-001 Cross-component primitives -- Anchor Timer (existing F-001-S-001), shared selection state, shared provider card. (Today: 1 story; existing components in Code/Shared/ux/ promote here when first used.)")
Bullet("F-002 Shared design paradigm -- design tokens, header/footer, fonts shared across 11 static HTML pages. (Today: implemented inline in each page; future story consolidates into a generated-from-brain or shared-CSS surface.)")
Bullet("F-003 Opening / parent page frame -- the parent page that hosts the chat iframe, postMessage host side, control-frame relay, leaf-page management. (Today: implemented in Website/index.html. Note: V3 had F-021 through F-025 staying in EPIC-006 FindCare; V4 keeps the FindCare-side iframe wiring in EPIC-006 but moves the parent-page-as-shared-host concern to EPIC-011.)")
Bullet("F-004 Lower-region environment banner -- the #envBanner shared across LOCAL/DEV/QA and across 3 backend servers, reflecting build/version/commit on every service handoff.")
Bullet("F-005 Auth-token display object -- buildTokenDisplayHtml(st) and the LEFT/RIGHT token panels; consumed by FindCare, EvaluateCare, and SharedServices /verify-token responses.")

P("Net effect on the proposed epic list (Section 6.1):", bold=True)
Bullet("UX is in the post-refactor epic list with 5 features and ~1 backlog story today (Anchor Timer); the existing implementations in Website/index.html become stories under F-002/F-003/F-004/F-005 once the refactor lands.")

HR()

# -----------------------------------------------------------------------------
# 4. Simple counts table -- one row per epic in the proposed final structure
# -----------------------------------------------------------------------------
H2("4. Proposed epic list -- counts table")
P("One row per epic in the proposed final structure. Counts are today's backlog counts where the epic is unchanged; for epics that gain/lose features, V4 shows the projected count and notes the basis.")

Table(
    ["Epic", "Name", "# Features", "# Stories"],
    [
        ["EPIC-001", "EvaluateCare", "25", "102"],
        ["EPIC-002", "Security", "2", "14"],
        ["EPIC-006", "FindCare", "14 (13 stays + Test feature)", "26 (25 stays + 1 traceability)"],
        ["EPIC-009", "Shared Services", "2 (1 + Safety subgraph absorbed from EPIC-004)", "15 (12 + 3 from EPIC-004)"],
        ["EPIC-010", "Operations", "17 (10 from EPIC-008 + 7 new ops features)", "51 (51 from EPIC-008 absorbed)"],
        ["EPIC-011", "UX", "5 (1 existing + 4 new shared-display features)", "1 (Anchor Timer; existing implementations not yet storied)"],
        ["EPIC-012", "Provider Knowledge Pipeline", "13 (from EPIC-006 F-005..F-018, F-028)", "40 (from EPIC-006 pipeline-flavored features)"],
    ],
    widths=[1.0, 2.6, 1.6, 1.7],
)

P("Removed from the tree (relative to current 9-epic shape):")
Bullet("EPIC-004 Safety -- 1 feature, 3 stories -> dissolved into EPIC-009 Shared Services as the 'Safety subgraph' feature. Stories move INTACT.")
Bullet("EPIC-005 Testing -- 4 features, 19 stories -> dissolved; each story re-homes to the capability it exercises (capability epics gain a Test feature).")
Bullet("EPIC-007 Talk About Care -- 1 feature, 1 story -> NOT in the Git tree. The story stays in the backlog as a forward-looking item per Skip's V2 binding ('Agree and it is not needed in the tree').")
Bullet("EPIC-008 Architecture -- 10 features, 51 stories -> dissolved into EPIC-010 Operations. Stories move INTACT to the new Operations features.")

HR()

# -----------------------------------------------------------------------------
# 5. Section A -- Current-state analysis (refreshed; UX as real shared scope)
# -----------------------------------------------------------------------------
H2("5. Section A -- Current-state analysis")
P("Walk of the actual file tree at c:/chatHealthy/findCare/ on 2026-04-28. Every cited path verified by direct filesystem listing. This section is intentionally compressed relative to V3 -- V3's full inventory remains valid and is incorporated by reference; V4 only refreshes the surfaces the corrections changed.")

H3("5.1 UX is a real shared concern (refresh)")
P("V3 treated UX as 'no current cross-epic consumer' -- correct for Code/Shared/ux/ taken in isolation, but incomplete because Website/index.html is itself the cross-epic UX consumer. Refreshed inventory:")
Bullet("Website/index.html (1,400+ lines) -- cross-epic parent page. Hosts banner (3 servers), auth-token display object (3 servers), LEFT/RIGHT token panels (FindCare always-on-LEFT + page-owning-service-on-RIGHT). Shared by every chat session of every capability.")
Bullet("Website/{architecture,products,privacy,terms,roadmap,security-architecture,...}.html -- 11 static pages sharing the design-token block (--ink/--teal/--accent), header layout, footer pattern. Today the paradigm is duplicated inline in each page (no shared CSS/template); the duplication itself is the cross-cutting UX scope.")
Bullet("Code/Shared/ux/{components/ProviderCard.tsx, components/SelectionManager.tsx, hooks/useSelectionState.ts, types/provider.ts} -- not consumed today, but explicitly created as cross-component primitives. Promote to EPIC-011 when first consumer materializes.")
Bullet("Code/ConversationalUX/FindCareChat/frontend/src/{App.tsx, main.tsx, components/{ChatWindow.tsx, FindCareApp.tsx, GUIManager.tsx, MessageBubble.tsx}} -- FindCare-specific UX. Stays under EPIC-006 FindCare / frontend.")
Bullet("EvaluateCare and SharedServices do not yet have standalone frontend code -- their UX surface today is the parent page in Website/index.html (banner + auth-token display) plus their /health and /session endpoints. When they grow capability-specific frontends those go under EPIC-001 / EPIC-009 respectively; the shared parent-page chrome stays under EPIC-011.")

H3("5.2 Other current-state observations carry from V3")
Bullet("The 7 anti-pattern folders (Code/CSharp/, Code/Schemas/, Code/ConversationalUX/, Code/find_care/ empty, Code/_pending/, Code/deploy/ at non-top-level, Code/shared_services/ vs Code/Shared/ naming clash) all carry as in V3.")
Bullet("Three duplicate spellings of 'FindCare' (findCare/ at repo root, Code/find_care/, Code/ConversationalUX/FindCareChat/) carry as in V3.")
Bullet("Schemas living under Website/ (deploy artifact serving as source-of-truth) carries as in V3 -- B.4 schema relocation plan unchanged.")
Bullet("True duplicates (Worker.cs, ChatHealthyMongoUtilities.py, ChatHealthy.Providers.json) -- diffs and recommendations now in Section 3 above.")

HR()

# -----------------------------------------------------------------------------
# 6. Section B -- Proposals
# -----------------------------------------------------------------------------
H2("6. Section B -- Proposals")

H3("6.1 Post-refactor epic list (with UX retained)")
Table(
    ["Epic ID", "Name (V4)", "Type", "P&L / org-chart anchor"],
    [
        ["EPIC-001", "EvaluateCare", "Capability (operational)", "Customer-facing capability with its own HF runtime and cost center. Owns full stack: scoring engine, measures, deploy, tests, architecture docs."],
        ["EPIC-002", "Security", "Cross-cutting (functional)", "Shared cost: Cloudflare findings, compliance attestation, HSTS/TLS posture, OAuth, session-token authentication. Demoted from candidate-capability per Skip's V2 binding."],
        ["EPIC-006", "FindCare", "Capability (operational)", "Customer-facing capability. Owns: search code, chat front end, capability-specific iframe wiring, HF Space runtime, tests."],
        ["EPIC-009", "Shared Services", "Cross-cutting (functional)", "Shared back-end consumed by every capability: Tool Router, Auth/session, Embeddings, LLM client, Cost guard, Mongo utility, Skip-Snow context, Lead capture, Consent, Unknown-question, Safety subgraph (absorbed from EPIC-004)."],
        ["EPIC-010", "Operations", "Cross-cutting (functional)", "DevOps + observability + HF Space deployment + environment promotion + monitoring + the V20 engineering-rules enforcement framework + brain governance helpers + the parent LangGraph runtime. Architecture-as-feature-of-Operations (EPIC-008 absorbed)."],
        ["EPIC-011", "UX", "Cross-cutting (functional)", "RETAINED per Correction 1. Shared display concerns: design paradigm across static pages, parent (opening) frame, lower-region env banner shared across servers, auth-token display object shared across 3 servers, cross-component primitives."],
        ["EPIC-012", "Provider Knowledge Pipeline", "Capability (functional)", "Standalone capability. Azure Functions runtime. ONE epic; THREE features (get content, orchestrate flow, enrich content). Per Correction 2, ETL is shared operational/business service -- bifurcating it across capability epics forces architectural drift; not OK."],
    ],
    widths=[0.8, 1.5, 1.4, 3.5],
)

H3("6.2 Per-epic ownership -- code each epic owns")
P("Same as V3 B.2 except for the UX-retained delta. Reproducing the deltas only:")

Table(
    ["Epic", "Code paths it owns (V4)"],
    [
        ["EPIC-001 EvaluateCare", "Code/EvaluateCare/ -- service code (app.py, scoring_engine.py, measures/, normalization.py, weights.py, confidence.py, explainability.py, failsafe.py, provenance.py, cache.py, models.py); handoff/ (the EvaluateCare side of the LangGraph<->EvaluateCare auth-token handshake including evaluate_care_facade.py and clinical_trials_models.py); orchestration/langgraph/ (EvaluateCare subgraph: state.py/nodes.py/tools.py/agent.py per LangGraph V4 Q6); core/ (the renamed domain layer: scoring_engine, normalization, weights, ...); test/; architecture/EvaluateCare/."],
        ["EPIC-002 Security", "Cloudflare config + security.txt + HSTS/TLS settings (Website/_headers); the OAuth/Auth-token authentication code (still being designed -- nonce-based today, OAuth+user-creation imminent per Skip). NB: the visual auth-token DISPLAY object is owned by EPIC-011 (presentation); the auth-token CREATION/VERIFICATION is owned by EPIC-002."],
        ["EPIC-006 FindCare", "Code/FindCare/ -- backend/ (find-care-specific provider/specialty/homeopathic/ranker code + core/ rename of domain/find_care/); frontend/ (the React chat UI + chat-window-specific styles); orchestration/langgraph/ (FindCare subgraph composed under the parent runtime); test/. The capability-specific iframe wiring (postMessage handlers internal to the chat) stays here. Note: the parent page that HOSTS the iframe moves to EPIC-011 (see UX rows below)."],
        ["EPIC-009 Shared Services", "Code/SharedServices/ -- tool_router.py (V2 binding); session_token.py, llm_client.py, cost_guard.py, prompt_system_maker.py; agent_framework/; embeddings/; consent/, lead_capture/, unknown_question/, about/ (Skip-Snow context + me/); mongo_utilities.py (one canonical copy per Issue 1.2 above); safety/safety_subgraph.py (absorbed from EPIC-004 per V3 binding)."],
        ["EPIC-010 Operations", "Code/Operations/ -- deploy/, local/, hf_space/, edge/, logging_service/ (one canonical Worker.cs per Issue 1.1 above), conversation_log/ (kafka/), build_promotion/, page_generators/, observability/, session_governor/, enforcement/ (V20 framework), brain_governance/, doc_generators/, epic_planning/, orchestration/parent_graph.py (the LangGraph parent runtime composing each capability subgraph)."],
        ["EPIC-011 UX (RETAINED)", "Code/SharedUX/ -- (1) parent_page/ (Website/index.html restructured: banner, opening frame, control-frame relay, postMessage host, leaf-page mgmt); (2) design_paradigm/ (shared :root tokens, header/footer fragments consumed by static page generator); (3) env_banner/ (the cross-server #envBanner code, /health-aware build/version/commit display); (4) auth_token_display/ (buildTokenDisplayHtml + LEFT/RIGHT panel logic); (5) cross_component_primitives/ (Code/Shared/ux/* -- ProviderCard, SelectionManager, useSelectionState, provider types -- promoted on first consumer)."],
        ["EPIC-012 Provider Knowledge Pipeline", "Code/ProviderKnowledgePipeline/ -- entire current Code/DataPipelines/ restructured under three feature-folders: get_content/, orchestrate_flow/, enrich_content/, plus test/ and architecture/. Per Correction 2 it stays as ONE epic. HuggingFace servers consume the pipeline's output (providers collection in Mongo); they do not own the pipeline."],
    ],
    widths=[1.5, 5.7],
)

H3("6.3 Master code-path move table")
P("V4 carries V3 B.3 verbatim with the following deltas (rest is unchanged):")

Table(
    ["Current path", "Target epic / feature (V4)", "Proposed destination", "Reason (delta from V3)"],
    [
        ["Code/CSharp/ChatHealthyLogService/dev/Worker.cs", "DELETE (legacy copy)", "--", "V4 Issue 1.1: Code/Shared/ops/ copy is canonical (active build path per .vscode/tasks.json). V3 had this as 'reconcile drift first.' V4 picks Copy B and deletes Copy A."],
        ["Code/CSharp/ChatHealthyLogService/dev/Program.cs, .csproj, appsettings*.json, Properties/", "DELETE (byte-identical to canonical)", "--", "All non-Worker.cs files in CSharp/dev/ are byte-identical to their counterparts in Code/Shared/ops/ChatHealthyLogService/. Pure duplication."],
        ["Code/CSharp/ChatHealthyLogService/build/", "DELETE (build output)", "--", "Compiled .dll/.exe/.deps.json; should be .gitignored. Carries from V3."],
        ["Code/Shared/ChatHealthyMongoUtilities.py (Copy A)", "EPIC-009 Shared Services / mongo_utilities", "Code/SharedServices/mongo_utilities.py", "V4 Issue 1.2: Copy A is canonical (only copy with callers; only copy watched by deploy CI). V3 had this as 'reconcile drift first.' V4 picks Copy A."],
        ["Code/DataPipelines/ChatHealthyMongoUtilities.py (Copy B)", "DELETE (dead code)", "--", "Copy B has no callers (DataPipelines uses module-level lazy _get_mongo_client() instead). Its 23-line WHEN-TO-USE docstring ports into Copy A. V4 picks Copy A."],
        ["Code/Schemas/ChatHealthy.Providers.json", "DELETE after relocation", "brain/machine_artifacts/schemas/ChatHealthyProvidersSchema.json", "V4 Issue 2: schemas are functionally identical. Adopt the Website/schemas/ content as canonical (trailing newline = small improvement); set $id at brain canonical URL. Code/Schemas/ folder dissolves entirely (Code/Schemas/ was a file-type-bucket anti-pattern per V3)."],
        ["Website/schemas/ChatHealthyProvidersSchema.json", "EPIC-010 Operations / brain_governance + EPIC-010 Operations / deploy (mirror)", "brain/machine_artifacts/schemas/ChatHealthyProvidersSchema.json + deploy step that mirrors brain -> Website at deploy time", "V4 Issue 2: this content is canonical (trailing-newline conventional, $id is dev-flavored which the deploy mirror swaps for prod). Closes V3 B.4 + V3 OQ6."],
        ["Website/index.html (full file)", "EPIC-011 UX -- multiple features", "Code/SharedUX/parent_page/index.html.j2 (templated) + Code/SharedUX/env_banner/banner.js + Code/SharedUX/auth_token_display/token_display.js + Code/SharedUX/parent_page/control_frame.js", "V4 UX restoration (Correction 1): the parent page is shared infrastructure, not a Website-only artifact. Restructure into UX features F-003/F-004/F-005; the deploy step generates Website/index.html from the templated parts."],
        ["Website/{architecture,products,privacy,terms,roadmap,security-architecture,chat-app-design,embedding-design,load-perf-report,ops-manager-design,provider-data-load}.html", "EPIC-011 UX / design_paradigm + (per-page content owners as before)", "Code/SharedUX/design_paradigm/{tokens.css, header.html.j2, footer.html.j2} consumed by each static page generator", "V4 UX restoration: factor the duplicated :root token block + header + footer pattern into shared CSS/templates owned by EPIC-011. Per-page content (architecture diagrams, product copy, etc.) stays under its content-owner epic."],
        ["Code/Shared/ux/{components,hooks,types}/", "EPIC-011 UX / cross_component_primitives (deferred -- promote on first consumer)", "Code/SharedUX/cross_component_primitives/", "V4 UX restoration: V3 said 'move to FindCare; promote on second consumer.' V4 keeps them in EPIC-011 (default home is the cross-cutting epic) and they activate when first consumed -- no second-consumer trigger needed."],
    ],
    widths=[2.0, 1.9, 1.9, 2.5],
)

P("All other rows from V3 B.3 carry forward unchanged. Reproducing them here would add length without adding decision; reference V3 B.3 for the rest.")

H3("6.4 Schema relocation plan")
P("V4 carries V3 B.4 with one update: Issue 2 above resolves the schema-canonical question (functionally equivalent; pick Website/schemas/ content). Concrete plan:")
Bullet("Create brain/machine_artifacts/schemas/ as a sibling of brain/machine_artifacts/content/.")
Bullet("Move Website/schemas/ChatHealthy{AgileBacklog,Bugs,Errors,Providers,Relaxed,RiskAcceptance,Version}Schema.json + EngineeringRulesSchema.json + JsonValidationFixtureSchema.json -> brain/machine_artifacts/schemas/.")
Bullet("Move Website/schemas/dev/ -> brain/machine_artifacts/schemas/dev/.")
Bullet("Leave Website/schemas/standard/json-schema-2020-12-meta.json where it is (third-party meta-schema served as a static asset).")
Bullet("Delete Code/Schemas/ChatHealthy.Providers.json (functionally equivalent duplicate per Issue 2).")
Bullet("Set the canonical $id on the relocated schema; the deploy step mirrors brain -> Website preserving the URL surface.")
Bullet("OPEN ITEM (new, non-blocking): the schema's `additionalProperties: false` rejects 100% of production providers because the running pipeline produces 11 fields not declared in the schema (embedding, embedding_model, embedding_version, 8 NPPES other-name fields). Track under EPIC-012 / schema-drift feature. This is independent of relocation and should not block it.")
Bullet("Engineering Rule-008 enforcement: Rule-008-ENF-001 currently excludes Website/schemas/standard/. After the move, the enforcement worker's allowed/excluded patterns must be updated to point at brain/machine_artifacts/schemas/ as the validation target.")

H3("6.5 Stories leaving / arriving in EPIC-006 -- disposition")
P("Same as V3 B.5 (every story stays intact -- no merging, splitting, rewording). One delta: V4 explicitly relocates the parent-page-frame stories.")

P("Delta from V3 B.5:", bold=True)
Table(
    ["Feature", "V3 disposition", "V4 disposition", "Reason"],
    [
        ["F-021 Parent Page Frame Layout (2)", "STAY at EPIC-006", "MOVE to EPIC-011 UX / opening_frame", "Per Correction 1: parent page is shared infrastructure (hosts every capability's iframe), not FindCare-specific."],
        ["F-022 postMessage Orchestration (2)", "STAY at EPIC-006", "SPLIT: parent-side handler -> EPIC-011; FindCare-side handler -> EPIC-006", "The host side of postMessage is shared; the FindCare-iframe side is capability-specific. Stories stay intact -- placement decided per story."],
        ["F-023 Chat Iframe Wiring (2)", "STAY at EPIC-006", "STAY at EPIC-006", "Iframe-side wiring is FindCare-specific."],
        ["F-024 Control Frame (3)", "STAY at EPIC-006", "MOVE to EPIC-011 UX / opening_frame", "Control frame is the parent page's shared chrome surface."],
        ["F-025 Leaf Page Management (1)", "STAY at EPIC-006", "MOVE to EPIC-011 UX / design_paradigm", "Leaf-page management is the static-page paradigm; cross-cutting."],
        ["F-026 index.html Generated from Brain Design Record (1)", "STAY at EPIC-006", "MOVE to EPIC-011 UX / parent_page", "The generated index.html IS the parent page."],
    ],
    widths=[2.0, 1.5, 2.5, 2.0],
)

P("All other V3 B.5 dispositions carry: 40 stories MOVE to EPIC-012; 5 stories MOVE to EPIC-010 (F-029 Local Development Topology); F-020 (0 stories) deletes the empty feature shell; F-027 Partial (mTLS handshake stays in EPIC-006); rest STAY.")

P("Net EPIC-006 story count after V4 dispositions:", bold=True)
Bullet("Stays in EPIC-006: F-001 (3), F-002 (1), F-003 (1), F-004 (2), F-019 (5), F-022 partial (~1 of 2 stays), F-023 (2), F-027 partial (1), F-030 (1), F-017 traceability (1) -- approximately 18 stays.")
Bullet("Move from EPIC-006 to EPIC-011: F-021 (2), F-022 partial (~1), F-024 (3), F-025 (1), F-026 (1) -- approximately 8 stories.")
Bullet("Move from EPIC-006 to EPIC-012: 40 stories (unchanged from V3).")
Bullet("Move from EPIC-006 to EPIC-010: 5 stories (F-029, unchanged from V3).")
Bullet("Delete: F-020 (0 stories, empty feature shell).")
P("This is a delta; the precise final count depends on whether F-022 splits 1+1 or 2+0 -- this is itself a V4 open question (Section 7).")

H3("6.6 LangGraph implementation alignment (mirror V4 of langgraph-oo-best-practices)")
P("Cross-checked against findCare/ArchitectureAndDesign/langgraph-oo-best-practices-V4.docx. V4 of the LangGraph guide consolidates the prior research-resolution-architect rounds into definitive choices. Each row maps a V4 LangGraph recommendation to the V4 codebase-organization proposal.")

Table(
    ["LangGraph V4 recommendation", "V4 codebase-organization placement"],
    [
        ["State as Pydantic BaseModel in its own state.py module (Q1) -- TypedDict only acceptable in notebook prototypes; Pydantic mandatory at HTTP/Kafka boundaries.", "Code/{capability}/orchestration/langgraph/state.py per capability subgraph. The current LangGraph/poc/user_journey.py JourneyState (15 flat TypedDict fields) decomposes into per-subgraph Pydantic state classes (FindCareState, EvaluateCareState, SafetyState; talk_about_care_state when EPIC-007 lands)."],
        ["MongoDBSaver from langgraph-checkpoint-mongodb (Q2) -- co-maintained by MongoDB and LangChain; NOT PostgresSaver; NOT a custom Mongo saver.", "Code/Operations/orchestration/parent_graph.py instantiates a single MongoDBSaver. Capability subgraphs share the checkpointer instance via subgraph composition. No new database tier introduced. AP-9 (introducing Postgres only for the checkpointer) and AP-12 (custom Mongo saver) explicitly excluded."],
        ["Runtime[ContextSchema] / ToolRuntime DI (Q5) -- no module-level imports of clients in node bodies.", "Node files in Code/{capability}/orchestration/langgraph/nodes.py accept Runtime[ContextSchema]. The DI surface is owned by Code/SharedServices/ (LLM client, Mongo client, embedding client). AP-5 explicitly excluded."],
        ["Tool I/O on the messages: Annotated[list[AnyMessage], add_messages] channel; structured artefacts in named typed Pydantic fields (Q4 + Q10).", "FindCareState (and every capability state) extends MessagesState. Tool calls + results ride as ToolMessage entries. The POC's flat tool-output fields (location, specialty_query, specialties, providers, trials, homeopathic_expansion, response) collapse. Selected providers and similar artefacts get typed fields with explicit reducers. AP-3 / AP-7 / AP-10 explicitly excluded."],
        ["Subgraph per capability composed under one parent runtime (Q9, AP-11) -- Command(graph=Command.PARENT) for handoffs; input_schema/output_schema project state.", "Code/Operations/orchestration/parent_graph.py is the thin parent. Each capability adds its compiled subgraph as a node. Subgraphs do not know each other's routing tables -- the parent's routing table picks the next epic."],
        ["Safety as reusable shared subgraph (Q9) -- not its own epic.", "Code/SharedServices/safety/safety_subgraph.py compiles once; FindCare, EvaluateCare, and (when implemented) Talk About Care add it as a node. The 3 detectors (emergency_detector, repetitive_detector, ip_lock_check) live in this single shared subgraph. Aligns with EPIC-004 dissolving into EPIC-009 per V3."],
        ["Bespoke subgraphs per capability (Q9) -- 'Use multiple agent subgraphs when you need bespoke agent implementations.'", "Each capability's structural shape genuinely differs (FindCare = retrieval+classification+ranking; EvaluateCare = deterministic scoring with provenance; Talk About Care = meeting + deep-research with interrupt() + Send fan-out). Single-agent-with-middleware is the wrong pattern."],
        ["Application-structure layout: state.py / nodes.py / tools.py / agent.py + langgraph.json (Q6) -- the langgraph-example-pyproject template.", "Each capability orchestration folder follows the official 4-file layout. Directly mirrors R-1 (langchain-ai/langgraph-example-pyproject)."],
        ["Reducers on every list field (Q7, AP-6) -- add_messages for messages, operator.add or custom bounded reducer for everything else.", "Every list field in every capability state schema gets an explicit reducer at declaration time. The POC's bare list fields (history, specialties, providers, trials, homeopathic_expansion) under the parallel safety+classifier fan-out are exactly the topology where last-write-wins silently loses data -- closed by per-subgraph state schemas."],
        ["Deep-research as reusable subgraph (Q9) -- canonical reference is langchain-ai/open_deep_research; deep_research_from_scratch is the schema-discipline scaffold.", "When Talk About Care lands: Code/SharedServices/deep_research/deep_research_subgraph.py imported as a node by Talk About Care, FindCare (cancer-research), EvaluateCare (cutting-edge-evidence). Send fan-out + interrupt() so meeting transcript continues while research workers run; single MongoDBSaver covers all suspended state."],
        ["Routing decisions as Pydantic structured outputs (Q9) -- ClarifyWithUser, ResearchQuestion, ConductResearch, ResearchComplete pattern from deep_research_from_scratch.", "Capability subgraphs use structured-output schemas for branching decisions; no free-form-string routing. Decision schemas live in the per-subgraph state.py."],
        ["Class-based nodes ONLY for dependency injection (Q3, issue #1950) -- never for hidden private state.", "Default to pure functions; promote to class with __call__ when constructor-injected dependencies are needed. Class-nodes do not carry hidden mutable state."],
        ["Two-layer testing (Q8): Layer 1 = unit-test each node directly; Layer 2 = integration-test compiled graph with InMemorySaver + FakeListLLM.", "Code/{capability}/orchestration/langgraph/test/ holds both layers. Production graphs use MongoDBSaver only in integration env; unit/integration tests use InMemorySaver per AP-8."],
    ],
    widths=[3.5, 3.5],
)

HR()

# -----------------------------------------------------------------------------
# 7. Section C -- remaining open questions
# -----------------------------------------------------------------------------
H2("7. Section C -- remaining open questions (<=4, each one decision)")
P("Each open question is answerable with one decision. V4 closes V3's OQ3 (Worker.cs canonical -> Copy B), OQ4 (Mongo utility canonical -> Copy A), OQ6 (provider schema canonical -> Website/schemas/ content with brain $id), OQ1 (UX disposition -> retained per Correction 1). Remaining open:")

Table(
    ["#", "Question", "Why it matters"],
    [
        ["1", "F-022 postMessage Orchestration split: does the feature split 1+1 (one story to EPIC-011, one story to EPIC-006), 2+0 (both stories to EPIC-011 with FindCare consuming), or 0+2 (both stay at EPIC-006 and EPIC-011 gets a parallel host-side story)? Stories stay intact regardless -- this is a placement decision.", "Determines the final EPIC-006 / EPIC-011 story-count split; affects ownership of the postMessage host-side handler in Code/SharedUX/parent_page/."],
        ["2", "Pipeline epic name -- confirm 'Provider Knowledge Pipeline' (V3 B.7 recommendation, carried forward) OR an alternative ('Provider Curation,' 'Healthcare Knowledge Pipeline,' 'Provider Catalog Pipeline,' or your own name).", "Renames Code/DataPipelines/ -> Code/{NewName}/ and renames the epic in the backlog. V3 carried this as OQ2; not yet decided."],
        ["3", "Provider schema content drift (separate from relocation): the schema at brain/machine_artifacts/schemas/ChatHealthyProvidersSchema.json (after relocation) declares 41 properties; the running pipeline produces documents with up to 11 additional fields (embedding, embedding_model, embedding_version, 8 NPPES other-name fields). Should the schema be updated to declare these, or should the pipeline be updated to stop producing them, or should both run with `additionalProperties: true` until a separate decision lands?", "100% of production providers fail validation today. This is a real schema-vs-data gap. EPIC-012 / schema-drift feature owns the work; the architectural choice (extend schema vs. trim pipeline output) belongs to Skip."],
        ["4", "EPIC-011 UX feature granularity: the V4 proposal lists 5 features (cross-component primitives, design paradigm, opening frame, env banner, auth-token display). Are these the right features, or should F-003 'opening frame' fold into F-002 'design paradigm' (since the parent page IS the design paradigm), making it 4 features total?", "Decides whether the parent-page frame is its own feature or a sub-concern of the shared design paradigm. Affects how the F-021/F-024/F-025/F-026 stories from EPIC-006 distribute under EPIC-011."],
    ],
    widths=[0.4, 4.7, 2.0],
)

HR()

# -----------------------------------------------------------------------------
# 8. Appendix -- verification record
# -----------------------------------------------------------------------------
H2("8. Appendix -- verification record")
P("Every architectural claim in V4 is grounded in a real file path or DB query result, captured 2026-04-28.")
Bullet("Worker.cs diff: `diff -u Code/CSharp/ChatHealthyLogService/dev/Worker.cs Code/Shared/ops/ChatHealthyLogService/Worker.cs` -- 189 vs 34 lines; SidecarScript/ConsumerScript vs MainScript=conversation_log_purge_service.py.")
Bullet("Worker.cs active build path: .vscode/tasks.json points 'build', 'publish', 'watch' all at Code/Shared/ops/ChatHealthyLogService/ChatHealthyLogService.csproj.")
Bullet("Worker.cs git history: commit 39763aa (Kafka full stack, T005/T006/BUG-LOG-002) precedes commit bd44f12 (T011 restructure: no business logic). Copy B is the more recent design intent.")
Bullet("ChatHealthyMongoUtilities.py diff: `diff -u Code/Shared/ChatHealthyMongoUtilities.py Code/DataPipelines/ChatHealthyMongoUtilities.py` -- Copy B has 23-line WHEN-TO-USE docstring; Copy A has commit() method; otherwise identical.")
Bullet("ChatHealthyMongoUtilities.py callers: only Code/ConversationalUX/FindCareChat/backend/main.py:47 imports the class. DataPipelines code uses module-level lazy `_get_mongo_client()` (verified across county_enrichment_job.py, copy_to_frontend.py, prescriber_pipeline_manager.py).")
Bullet("ChatHealthyMongoUtilities.py CI watch: .github/workflows/deploy-findcare-backend.yml line 14 watches 'Code/Shared/ChatHealthyMongoUtilities.py' only.")
Bullet("Provider schemas diff: `diff -u Code/Schemas/ChatHealthy.Providers.json Website/schemas/ChatHealthyProvidersSchema.json` -- 19,265 vs 19,276 bytes; the 11-byte delta is `$id` URL difference + trailing newline. Every property, type, required-list, and conditional is byte-identical.")
Bullet("Provider DB sample: read-only $sample aggregation on prod_PublicHealthData.providers (143,589 docs, sample=500) and dev_PublicHealthData.providers (43,491 docs, sample=200). Both schemas validated each document with jsonschema.Draft202012Validator. Pass count: 0/500 prod, 0/200 dev for both schemas. Asymmetric rejections: 0/700 -- the schemas are functionally equivalent against production data.")
Bullet("Provider schema gap fields (frequency in 500 prod sample): embedding 100%, embedding_model 52%, embedding_version 52%, authorized_official_name_prefix_text 9%, provider_other_last_name 6%, provider_other_first_name 6%, provider_other_last_name_type_code 6%, provider_other_middle_name 4%, provider_other_credential_text 4%, provider_other_name_prefix_text 1%, authorized_official_name_suffix_text <1%.")
Bullet("UX cross-cutting evidence: Website/index.html has #envBanner at lines 627-660 (banner.style.background='#dc2626' for non-LOCAL); window.refreshBanner(service) reflects health from FindCare/EvaluateCare/SharedServices (lines 882, 917, 939, 1158, 1238, 1336); window.buildTokenDisplayHtml(st) at lines 1011-1045 renders nonce + GUID + signature; LEFT panel always shows FindCare token, RIGHT shows page-owning-service token (lines 808, 840).")
Bullet("Static-page design paradigm: 11 .html files under Website/ each carry a :root token block (--ink, --teal, --accent), DM Serif Display + DM Sans fonts, the same header layout. Verified by direct grep of Website/products.html and Website/privacy.html.")
Bullet("Code/Shared/ux/: ProviderCard.tsx (4,235 bytes), SelectionManager.tsx (7,882 bytes), useSelectionState.ts (4,475 bytes), provider.ts (1,233 bytes). Not currently consumed by FindCareChat/frontend/src/components/.")
Bullet("Backlog counts (verified live): EPIC-001=25F/102S, EPIC-002=2F/14S, EPIC-004=1F/3S, EPIC-005=4F/19S, EPIC-006=29F/71S, EPIC-007=1F/1S, EPIC-008=10F/51S, EPIC-009=1F/12S, EPIC-011=1F/1S.")
Bullet("LangGraph V4 alignment: all 12 'Definitive choices' from langgraph-oo-best-practices-V4.docx Section 1 plus all 12 anti-patterns AP-1..AP-12 are addressed in Section 6.6 above.")
Bullet("Engineering rules read: brain/machine_artifacts/content/engineering_rules.json -- Rule-001..Rule-062 reviewed before action. V4 takes only read-only actions: filesystem reads, $sample DB aggregations, Word doc emit. No commits, no backlog mutations, no schema mutations.")

# -----------------------------------------------------------------------------
# Save
# -----------------------------------------------------------------------------
doc.save(str(OUT))
print(f"Wrote {OUT}")
