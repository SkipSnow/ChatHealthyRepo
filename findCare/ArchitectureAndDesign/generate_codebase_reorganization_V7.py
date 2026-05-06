"""Generate codebase-reorganization-V7.docx.

Read-only generator. The V7 brief: this is the FINAL reorganization document
before the design-synthesis assignment, so every architectural choice must be
resolved with no "options to consider", every Skip annotation in V6 must be
quoted verbatim and answered, and a developer reading V7 alone (without V1-V6
in context) must understand the proposed structure end-to-end.

V7 imports the per-feature code map and per-epic orphan list from the V5
generator (those are the binding code-grounding tables) and overlays the
annotation resolutions and the dual-child sharedServices structure that V6
introduced.

No commits, no code changes, no backlog mutations. Pure docx emission.
"""

import json
import os
import sys
from collections import defaultdict
from docx import Document
from docx.shared import Pt

REPO = os.environ['CHATHEALTHY_PROJECT_ROOT']
BACKLOG = os.path.join(REPO, "brain/machine_artifacts/content/agile_backlog.json")
OUT = os.path.join(
    REPO, "findCare/ArchitectureAndDesign/codebase-reorganization-V7.docx"
)

# Import the binding code-grounding maps from the V5 generator.
sys.path.insert(0, os.path.join(REPO, "findCare/ArchitectureAndDesign"))
import generate_codebase_reorganization_V5 as V5  # noqa: E402

FEATURE_CODE = V5.FEATURE_CODE
FEATURE_FINAL_EPIC = V5.FEATURE_FINAL_EPIC
STORY_FINAL_EPIC_OVERRIDES = V5.STORY_FINAL_EPIC_OVERRIDES
ORPHANS = V5.ORPHANS

# V7 epic display names (carry forward from V5/V6; EPIC-012 named per binding).
V7_EPICS = [
    ("EPIC-001", "EvaluateCare"),
    ("EPIC-002", "Security & Compliance"),  # renamed per V6 Skip annotation #1
    ("EPIC-006", "FindCare"),
    ("EPIC-009", "Shared Services"),
    ("EPIC-010", "Operations"),
    ("EPIC-011", "UX"),
    ("EPIC-012", "Data Management Pipelines"),
]

# ----------------------------------------------------------------------------
# Skip annotations from V6 (extracted from the V6 docx body and tables).
# Each entry: (annotation_id, location_in_v6, verbatim_text, v7_resolution,
# v7_section_pointer)
# ----------------------------------------------------------------------------
SKIP_ANNOTATIONS_V6 = [
    {
        "id": "A1",
        "location": "V6 Section 2.3 (Security hosting recommendation), P35",
        "verbatim": (
            "</ skip Accepted. The CSO is a cost center in any digital "
            "healthcare company and it would be staffed by humans. The CSO "
            "is a C-level department, even though usually buried in the IT "
            "org and siloed from the compliance org. So from an organizing "
            "perspective I would say that compliance and security must be "
            "bundled together. Explicitly dose this make sense? >"
        ),
        "resolution": (
            "ACCEPTED. EPIC-002 is renamed Security & Compliance and owns "
            "one HF Space. The org-chart line is the CSO; the cost-center "
            "line is Security & Compliance; the deploy line is one Space. "
            "Three columns aligned, axiom honored. Compliance attestation, "
            "policy generation, audit-log surface, and security posture "
            "(OAuth, session-token verification, entitlements, HSTS/TLS) "
            "all live under EPIC-002. Compliance is NOT a separate epic; it "
            "is a feature folder under EPIC-002 (Code/Security/backend/"
            "compliance/). The Cloudflare findings page, the security-"
            "architecture page, and any future SOC2/HIPAA artifact "
            "generators all anchor here."
        ),
        "section_pointer": "Section 3 (Operating-model + epic shape) and "
        "Section 11 row 11 (master move table)",
    },
    {
        "id": "A2",
        "location": "V6 Section 4.2 (ChatHealthyMongoUtilities canonical), "
        "P43-P47 (multi-paragraph)",
        "verbatim": (
            "</ Skip While I agree with the premise that each "
            "HuggingFaceSpace must have access to Mongo my only quibble is "
            "having a root level directory named code. Perhaps shared "
            "services has two child: 1: {project root}/sharedServices/"
            "RuntimeServer/feachure_1   child: 2: {project root}/"
            "sharedServices/SharedCode/feachure_3   In this model the epic "
            "has one set of features but each feature manifests itself as a "
            "library that will be embedded everywhere, or an API that will "
            "present a facade and be managed by the workflow. >"
        ),
        "resolution": (
            "ACCEPTED. V7 adopts the dual-child layout for every cross-"
            "cutting epic that ships both an embedded library and a "
            "runtime server. The pattern is: {epic_root}/RuntimeServer/ "
            "(the HF Space process: app.py, Dockerfile, requirements.txt, "
            "endpoints) and {epic_root}/SharedCode/ (the library that "
            "FindCare and EvaluateCare embed at build time: pure Python "
            "modules, no FastAPI, no transport). A feature manifests as "
            "EITHER a library (lives only under SharedCode/) OR a facade "
            "(lives only under RuntimeServer/) OR both (a thin facade in "
            "RuntimeServer/ delegates to the library in SharedCode/). The "
            "root directory is no longer named 'Code'; the runtime tree "
            "moves to {project root}/SharedServices/, {project root}/"
            "Security/, etc. (capability-named roots, NOT a 'Code/' "
            "umbrella). Concrete shape under V7: SharedServices/"
            "RuntimeServer/{tool_router, llm_client, embeddings, consent, "
            "lead_capture, safety, who_am_i}/ + SharedServices/SharedCode/"
            "{mongo_utilities, agent_framework, cost_guard, prompt_system_"
            "maker, brain_loop}/ + SharedServices/Dockerfile + "
            "SharedServices/requirements.txt at the RuntimeServer/ child. "
            "ChatHealthyMongoUtilities lands at SharedServices/SharedCode/"
            "mongo_utilities.py (library, embedded by every Space)."
        ),
        "section_pointer": "Section 4 (Layer + folder convention) and "
        "Section 11 row 4 (master move table)",
    },
    {
        "id": "A3",
        "location": "V6 Section 4.3 (Provider schema canonical), P52",
        "verbatim": (
            "</Skip We should delete the one that is not on the website "
            "now, as the website is the current source of record >"
        ),
        "resolution": (
            "ACCEPTED. Code/Schemas/ChatHealthy.Providers.json is "
            "DELETED. The Website/schemas/ChatHealthyProvidersSchema.json "
            "content becomes canonical; it is relocated to brain/machine_"
            "artifacts/schemas/ChatHealthyProvidersSchema.json with the "
            "deploy step mirroring brain -> Website. The 11-field gap "
            "(production providers carry embedding + 8 NPPES other-name "
            "fields not declared by the schema) is tracked under EPIC-012 "
            "schema-drift feature; it does not block the relocation."
        ),
        "section_pointer": "Section 10 (Schema relocation plan) and "
        "Section 11 rows 6-7 (master move table)",
    },
    {
        "id": "A4",
        "location": "V6 Section 6 preamble (per-epic story-to-code "
        "tables intro), P59-P64 (multi-paragraph)",
        "verbatim": (
            "</ Skip Yes about no code found we should probably look to "
            "get rid of that noise.   My response prima facia to the "
            "tables below is:   For evaluate care we have a ton of code "
            "but we never exersized it and the design was made by Claude "
            "with no supervision so given our history it might not be of "
            "any value. I don't recommend we throw it out, but I'm not "
            "sure if we should port it either. If it relies on being in "
            "the same container as FindCare then at best we can refactor "
            "it. If it is all a jumble of procedural code with no reuse "
            "and a lack of abastactions then the best we can do is "
            "refactor it.   In other places it looks like we realize code "
            "for stories all over the place and I'm not sure that is a "
            "good thing or a bad thing but based on what I know of Claude "
            "Code's coding style chances are most of it needs to be mined "
            "and refactored. I don't really like that we have so much "
            "code to do so little work.   It is good that we have this "
            "map, because when we migrate to the new LangGraph "
            "architecture having this source to refactor from will be "
            "very helpful. >"
        ),
        "resolution": (
            "ACCEPTED with two concrete commitments: (1) NO-CODE-FOUND "
            "stories are surfaced as a discrete list (Section 8) and "
            "EACH is given an explicit V7 disposition: KEEP "
            "(future-work, anchor a future feature folder), or REMOVE "
            "(no business value; recommend story retirement). V7 does "
            "not mutate the backlog; it lists the dispositions for the "
            "design-synthesis agent to act on. (2) EvaluateCare code is "
            "treated as a refactor source, NOT a port. The V7 file-tree "
            "for EPIC-001 is described as 'rebuilt under EvaluateCare/"
            "RuntimeServer/ + EvaluateCare/SharedCode/ from the existing "
            "Code/evaluate_care/ source as study material'. The existing "
            "tree is preserved in the move table as the migration source "
            "(not deleted at move time); the design-synthesis agent "
            "specifies the migration shape per story."
        ),
        "section_pointer": "Section 6 (per-epic story-to-code tables, "
        "preamble), Section 8 (NO-CODE-FOUND disposition list)",
    },
    {
        "id": "A5",
        "location": "V6 Section 7 preamble (per-epic orphan-code tables "
        "intro), P81",
        "verbatim": (
            "</Skip All evaluate care code that is orphaned must be "
            "deleted>"
        ),
        "resolution": (
            "ACCEPTED. The four EPIC-001 orphan rows in V6 Section 7.1 "
            "(Code/evaluate_care/__init__.py, Code/evaluate_care/graphs/, "
            "Code/evaluate_care/measures/__init__.py, Code/evaluate_care/"
            "tests/) are tagged DELETE in the V7 orphan table. The "
            "design-synthesis agent does not need to find a home for "
            "them; they leave the tree."
        ),
        "section_pointer": "Section 7.1 (EPIC-001 orphan table) and "
        "Section 11 (master move table addendum)",
    },
    {
        "id": "A6",
        "location": "V6 Section 7.3 (EPIC-006 FindCare orphan-code "
        "table), P87",
        "verbatim": (
            "</ Skip: From reading this it looks like all of this code "
            "can be deleted and it is not referred to by other code, am "
            "I correct?>"
        ),
        "resolution": (
            "PARTIAL ACCEPT, with file-by-file verification. Of the 9 "
            "EPIC-006 orphan rows in V6: (a) Code/find_care/ (empty "
            "directory) -> DELETE. (b) The five empty-package __init__.py "
            "rows (backend/adapter, backend/domain/shared/url, backend/"
            "infrastructure/config, backend/infrastructure/external_apis, "
            "backend/infrastructure/mongo) -> DELETE the directories at "
            "migration time; the V7 FindCare/RuntimeServer/ tree does "
            "not need empty placeholders. (c) GUIManager.tsx, "
            "FindCareApp.tsx, MessageBubble.tsx -> KEEP and trace under "
            "the existing F-004 / F-019 / F-023 stories at refactor "
            "time; these are the React surface the iframe shows. They "
            "appear orphan only because no story names them by file "
            "name. Per V7 React-as-facade direction (LangGraph V6 "
            "Section 3.3) these components stay; they do not call the "
            "capability Spaces directly."
        ),
        "section_pointer": "Section 7.3 (EPIC-006 orphan table)",
    },
    {
        "id": "A7",
        "location": "V6 Section 7.6 (EPIC-011 UX orphan-code table), "
        "P93",
        "verbatim": (
            "</ Skip my bet is all of these or most of these are "
            "important  we were working very quickly and built a lot of "
            "valuable stuff n debugging sessions>"
        ),
        "resolution": (
            "ACCEPTED. All 14 EPIC-011 UX orphan rows are tagged KEEP "
            "in the V7 orphan table. The 3 React primitive files "
            "(ProviderCard.tsx, SelectionManager.tsx, provider.ts) "
            "move to UX/SharedCode/cross_component_primitives/; the "
            "11 static HTML pages remain in Website/ until "
            "EPIC-011-F-002 (design paradigm) lands templated parts "
            "(see annotation A10 for the roadmap.html exception). "
            "EPIC-011 grows new stories under F-002, F-004, F-005 to "
            "anchor the genuinely shared design surface. V7 commits "
            "to keeping these files."
        ),
        "section_pointer": "Section 7.6 (EPIC-011 orphan table) and "
        "Section 9 (UX requirements gap)",
    },
    {
        "id": "A8",
        "location": "V6 Section 7.5 (EPIC-010 Operations orphan-code "
        "table), Table 12 R12 cell on rebuild_manifest.py",
        "verbatim": (
            "</ Skip: looks like I have to look at these harder because "
            "this one is important. It creates our manifest of files "
            "which is an important governance object for the brain >"
        ),
        "resolution": (
            "ACCEPTED. Code/Shared/ops/rebuild_manifest.py and Code/"
            "Shared/ops/manifest_generator.py are explicitly tagged "
            "GOVERNANCE-CRITICAL in the V7 orphan table. Both move to "
            "Operations/SharedCode/brain_governance/ (NOT into "
            "RuntimeServer/, because they are batch-time tools the "
            "brain runs, not runtime endpoints). V7 recommends a new "
            "story under EPIC-010-F-007 (Brain governance / manifest "
            "generation) to anchor them; the design-synthesis agent "
            "is the right place to add that story to the backlog."
        ),
        "section_pointer": "Section 7.5 (EPIC-010 orphan table, "
        "rebuild_manifest.py row) and Section 11 (master move table)",
    },
    {
        "id": "A9",
        "location": "V6 Section 7.6 (EPIC-011 UX orphan-code table), "
        "Table 13 R4 cell on Website/architecture.html",
        "verbatim": (
            "</Skip: This is got to go but not until we have something "
            "to replace it with All of these web site pages are "
            "important and need stories. These the static web site "
            "pages  are some of what I was saying when I said we had "
            "static web pages that from my perspective are shared UX "
            "components when we do our next production push and soon I "
            "hope we have to refresh all of this content (the static "
            "web pages) >"
        ),
        "resolution": (
            "ACCEPTED. Website/architecture.html and the 9 sibling "
            "static pages (products, privacy, terms, security-"
            "architecture, chat-app-design, embedding-design, "
            "load-perf-report, ops-manager-design, provider-data-load) "
            "are tagged KEEP-PENDING-REPLACEMENT in V7. They stay in "
            "Website/ until EPIC-011-F-002 (Shared design paradigm) "
            "delivers the templated parts (UX/SharedCode/design_"
            "paradigm/{tokens.css, header.html.j2, footer.html.j2}) "
            "AND the static-page generator that consumes them. V7 "
            "commits to a new EPIC-011 feature 'Static page generator' "
            "(F-006) with 1 story per static page (10 stories) so each "
            "page has a story owner. Roadmap.html is the exception "
            "(see A10)."
        ),
        "section_pointer": "Section 7.6 (EPIC-011 orphan table), "
        "Section 9 (UX requirements gap), Section 11 (master move "
        "table row 9)",
    },
    {
        "id": "A10",
        "location": "V6 Section 7.6 (EPIC-011 UX orphan-code table), "
        "Table 13 R8 cell on Website/roadmap.html",
        "verbatim": (
            "</Skip we need to delete this ASAP>"
        ),
        "resolution": (
            "ACCEPTED. Website/roadmap.html is tagged DELETE in the V7 "
            "orphan table. It is removed from EPIC-011's territory at "
            "migration time. Code/Shared/ops/generate_roadmap_page.py "
            "(the generator) is also tagged DELETE in the Operations "
            "orphan table; with no roadmap page there is no generator."
        ),
        "section_pointer": "Section 7.5 (EPIC-010 orphan table, "
        "generate_roadmap_page.py row), Section 7.6 (EPIC-011 orphan "
        "table, roadmap.html row), Section 11 (master move table "
        "addendum row D2)",
    },
]


# ----------------------------------------------------------------------------
# NO-CODE-FOUND disposition list (per A4 commitment).
# ----------------------------------------------------------------------------
NO_CODE_DISPOSITIONS = [
    # (feature_id, feature_name, # stories, disposition, rationale)
    (
        "EPIC-001-F-011",
        "User-Adjustable Quality Weights - Mixing Board",
        5,
        "KEEP",
        "Mixing-board UX is a real EvaluateCare differentiator; design-"
        "synthesis assigns a feature folder under EvaluateCare/SharedCode/"
        "weights/ + UX/SharedCode/mixing_board/ for the slider UI",
    ),
    (
        "EPIC-001-F-023",
        "(EvaluateCare future feature)",
        1,
        "KEEP",
        "Anchor for a future EvaluateCare measure; surface to design-"
        "synthesis as a future-work story",
    ),
    (
        "EPIC-001-F-025",
        "(EvaluateCare future feature)",
        1,
        "KEEP",
        "Anchor for a future EvaluateCare measure; surface to design-"
        "synthesis as a future-work story",
    ),
    (
        "EPIC-006-F-020",
        "(FindCare future feature)",
        1,
        "RECOMMEND RETIREMENT",
        "No code, no clear capability anchor; surface to design-synthesis "
        "for backlog cleanup if no business sponsor",
    ),
    (
        "EPIC-007-F-001",
        "Talk About Care",
        1,
        "KEEP - PARK",
        "Per V5 binding, parked under EPIC-009 (deep_research subgraph "
        "host per LangGraph V6); revisit when Talk About Care reaches "
        "active development",
    ),
    (
        "EPIC-008-F-001",
        "(Architecture - absorbed into EPIC-010)",
        1,
        "RECOMMEND RETIREMENT",
        "Architecture epic was dissolved; story has no implementing code "
        "and no clear capability anchor",
    ),
    (
        "EPIC-008-F-006",
        "(Architecture - absorbed into EPIC-010)",
        1,
        "RECOMMEND RETIREMENT",
        "Same as F-001 above",
    ),
    (
        "EPIC-008-F-009",
        "(Architecture - absorbed into EPIC-010)",
        1,
        "RECOMMEND RETIREMENT",
        "Same as F-001 above",
    ),
]


# ----------------------------------------------------------------------------
# Master move table for V7 (carries forward V6 with annotation overlays).
# Each row: (current_path, target_epic_feature, proposed_destination, reason)
# ----------------------------------------------------------------------------
MASTER_MOVES_V7 = [
    (
        "Code/CSharp/ChatHealthyLogService/dev/Worker.cs",
        "DELETE (legacy)",
        "--",
        "Carries V6: Code/Shared/ops/ copy is canonical (active build).",
    ),
    (
        "Code/CSharp/ChatHealthyLogService/dev/* (non-Worker.cs files)",
        "DELETE (byte-identical duplicates)",
        "--",
        "Identical to Code/Shared/ops/ChatHealthyLogService/ counterparts.",
    ),
    (
        "Code/CSharp/ChatHealthyLogService/build/",
        "DELETE (build output)",
        "--",
        "Compiled .dll/.exe/.deps.json; should be .gitignored.",
    ),
    (
        "Code/Shared/ChatHealthyMongoUtilities.py",
        "EPIC-009 SharedServices / SharedCode",
        "SharedServices/SharedCode/mongo_utilities.py",
        "ACCEPTED A2: dual-child layout; library lives in SharedCode/, "
        "embedded by every Space.",
    ),
    (
        "Code/DataPipelines/ChatHealthyMongoUtilities.py",
        "DELETE (dead code)",
        "--",
        "Carries V6: no callers; WHEN-TO-USE docstring ports into "
        "SharedServices/SharedCode/mongo_utilities.py.",
    ),
    (
        "Code/Schemas/ChatHealthy.Providers.json",
        "DELETE (per A3)",
        "--",
        "ACCEPTED A3: Website/schemas/ copy is current source of "
        "record; the Code/Schemas/ duplicate is removed.",
    ),
    (
        "Website/schemas/* (except standard/json-schema-2020-12-meta.json)",
        "EPIC-010 Operations / brain_governance",
        "brain/machine_artifacts/schemas/*",
        "Brain owns schemas; deploy step mirrors brain -> Website.",
    ),
    (
        "Website/index.html (full file)",
        "EPIC-011 UX",
        "UX/SharedCode/parent_page/index.html.j2 + UX/SharedCode/"
        "env_banner/banner.js + UX/SharedCode/auth_token_display/"
        "token_display.js + UX/SharedCode/parent_page/control_frame.js",
        "Restructure into UX features F-002/F-003/F-004/F-005; deploy "
        "step generates Website/index.html from templated parts.",
    ),
    (
        "Website/{architecture,products,privacy,terms,security-"
        "architecture,chat-app-design,embedding-design,load-perf-"
        "report,ops-manager-design,provider-data-load}.html",
        "EPIC-011 UX / design_paradigm + new F-006 static page generator",
        "UX/SharedCode/design_paradigm/{tokens.css, header.html.j2, "
        "footer.html.j2} + UX/SharedCode/static_page_generator/",
        "ACCEPTED A9: keep pending replacement; new EPIC-011-F-006 "
        "feature anchors each static page with a story.",
    ),
    (
        "Website/roadmap.html",
        "DELETE (per A10)",
        "--",
        "ACCEPTED A10: 'we need to delete this ASAP'.",
    ),
    (
        "Code/Shared/ops/generate_roadmap_page.py",
        "DELETE (per A10)",
        "--",
        "ACCEPTED A10 cascade: with no roadmap page there is no "
        "generator.",
    ),
    (
        "Code/Shared/ux/{components,hooks,types}/",
        "EPIC-011 UX / SharedCode / cross_component_primitives",
        "UX/SharedCode/cross_component_primitives/",
        "ACCEPTED A7 (KEEP): the React primitives are valuable from "
        "debugging sessions; activate when first consumed.",
    ),
    (
        "Code/shared_services/* (auth + session-token + entitlement code)",
        "EPIC-002 Security & Compliance / RuntimeServer",
        "Security/RuntimeServer/{auth, entitlements, audit, compliance}/ "
        "+ Security/Dockerfile + Security/requirements.txt + "
        "Security/SharedCode/{token_models, policy_models}/",
        "ACCEPTED A1: own HF Space; Compliance is a feature folder "
        "under Security & Compliance, not a separate epic.",
    ),
    (
        "Code/shared_services/* (everything else)",
        "EPIC-009 SharedServices",
        "SharedServices/RuntimeServer/{tool_router, llm_client, "
        "embeddings, consent, lead_capture, safety, who_am_i, "
        "deep_research}/ + SharedServices/SharedCode/{mongo_utilities, "
        "agent_framework, cost_guard, prompt_system_maker, brain_loop}/ "
        "+ SharedServices/Dockerfile + SharedServices/requirements.txt",
        "ACCEPTED A2: dual-child layout. Each feature is library, "
        "facade, or both.",
    ),
    (
        "Code/Shared/ops/rebuild_manifest.py + manifest_generator.py",
        "EPIC-010 Operations / SharedCode / brain_governance",
        "Operations/SharedCode/brain_governance/{rebuild_manifest.py, "
        "manifest_generator.py}",
        "ACCEPTED A8: governance-critical; new story EPIC-010-F-007 "
        "anchors them.",
    ),
    (
        "Code/Shared/ops/* (everything else under ops/)",
        "EPIC-010 Operations",
        "Operations/RuntimeServer/{deploy, observability, hf_space_"
        "manager, build_promotion}/ + Operations/SharedCode/"
        "{engineering_rules, conversation_log, kafka, framework_"
        "version}/",
        "Carries V6 with dual-child overlay. ChatHealthyLogService "
        "lands at Operations/SharedCode/conversation_log/.",
    ),
    (
        "Code/DataPipelines/*",
        "EPIC-012 Data Management Pipelines",
        "DataManagementPipelines/{get_content, orchestrate_flow, "
        "enrich_content}/ (NOT dual-child; pipelines are scheduled "
        "jobs, not runtime servers)",
        "Carries V5/V6. Per LangGraph V6 Section 3.3, pipelines stay "
        "out of LangGraph entirely; no RuntimeServer/SharedCode split "
        "needed.",
    ),
    (
        "LangGraph/poc/user_journey.py",
        "EPIC-010 Operations / orchestration",
        "Operations/RuntimeServer/orchestration/parent_graph.py",
        "Per LangGraph V6: the parent runtime is thin; capability "
        "subgraphs added via RemoteGraph as nodes. The parent runtime "
        "is the React-facing endpoint.",
    ),
    (
        "Code/evaluate_care/* (entire tree)",
        "EPIC-001 EvaluateCare (refactor source per A4)",
        "EvaluateCare/RuntimeServer/{scoring_engine, explainability, "
        "measures, app}/ + EvaluateCare/SharedCode/{measures_lib, "
        "models, weights, normalization, provenance, confidence, "
        "failsafe, cache}/",
        "ACCEPTED A4: tree is a refactor source, not a port. "
        "EvaluateCare gains its own HF Space (already does in "
        "Code/evaluate_care/Dockerfile + app.py). Refactor at "
        "design-synthesis time per story.",
    ),
    (
        "Code/evaluate_care/__init__.py + graphs/ + measures/__init__.py "
        "+ tests/",
        "DELETE (per A5)",
        "--",
        "ACCEPTED A5: 'all evaluate care code that is orphaned must "
        "be deleted'.",
    ),
    (
        "Code/find_care/ (empty dir) + 5 empty __init__.py packages "
        "under FindCareChat/backend/{adapter, domain/shared/url, "
        "infrastructure/{config, external_apis, mongo}}",
        "DELETE (per A6)",
        "--",
        "ACCEPTED A6: confirmed not referenced; remove empty "
        "placeholders at migration time.",
    ),
    (
        "Code/ConversationalUX/FindCareChat/* (everything else)",
        "EPIC-006 FindCare",
        "FindCare/RuntimeServer/{provider_search, specialty_"
        "classifier, evaluate_care_facade, tool_router_client}/ + "
        "FindCare/SharedCode/{models, services, repositories}/ + "
        "FindCare/Frontend/{src, index.html, package.json}/",
        "FindCare gets dual-child for backend; the React frontend is "
        "an additional sibling FindCare/Frontend/ child (the iframe).",
    ),
]


# ----------------------------------------------------------------------------
# Helpers (small, copy of V5 helpers).
# ----------------------------------------------------------------------------
def load_backlog():
    with open(BACKLOG, "r", encoding="utf-8") as f:
        return json.load(f)


def all_stories():
    data = load_backlog()
    for e in data["epics"]["epic"]:
        eid = e["epic_id"]
        ename = e["name"]
        feats = (e.get("features") or {}).get("feature", [])
        for f in feats:
            fid = f["feature_id"]
            fname = f["name"]
            stories = (f.get("stories") or {}).get("story", [])
            if not stories:
                continue
            for s in stories:
                sid = s.get("story_id") or s.get("id") or ""
                sd = (s.get("description") or s.get("summary") or "").replace(
                    "\n", " "
                ).strip()
                yield eid, ename, fid, fname, sid, sd


def final_epic_for_story(orig_epic, fid, sid):
    if sid in STORY_FINAL_EPIC_OVERRIDES:
        return STORY_FINAL_EPIC_OVERRIDES[sid]
    if fid in FEATURE_FINAL_EPIC:
        return FEATURE_FINAL_EPIC[fid]
    return orig_epic


def code_for_feature(fid):
    val = FEATURE_CODE.get(fid)
    if val is None:
        return "NO CODE FOUND"
    return val


def render_code_paths(val):
    if val == "NO CODE FOUND":
        return "NO CODE FOUND"
    if isinstance(val, list):
        return "\n".join(val)
    return str(val)


# Doc primitives.
def add_para(doc, text, style=None, bold=False):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    r = p.add_run(text)
    if bold:
        r.bold = True
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def add_h1(doc, text):
    doc.add_heading(text, level=1)


def add_h2(doc, text):
    doc.add_heading(text, level=2)


def add_h3(doc, text):
    doc.add_heading(text, level=3)


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            t.rows[r_idx].cells[c_idx].text = "" if val is None else str(val)
    doc.add_paragraph()  # spacer
    return t


# ----------------------------------------------------------------------------
# Main emission
# ----------------------------------------------------------------------------
def main():
    doc = Document()

    # ---------- Title -------------------------------------------------------
    doc.add_heading("Codebase Reorganization -- V7", level=0)
    add_para(
        doc,
        "Date: 2026-04-28   |   Author: Claude (architecture analysis "
        "agent)   |   Read-only   |   FINAL reorganization doc before the "
        "design assignment",
    )
    add_para(
        doc,
        "Lineage: V1 (initial) -> V2 (Skip-annotated) -> V3 (resolution) "
        "-> V4 (consolidation) -> V5 (per-epic story-to-code maps + "
        "Security own-Space recommendation) -> V6 (architect annotations "
        "on V5) -> this V7 (annotation resolution; standalone; "
        "definitive).",
    )
    add_para(doc, "Inputs read in order:")
    add_bullet(
        doc,
        "findCare/ArchitectureAndDesign/codebase-reorganization-V6.docx "
        "-- architect-annotated V5; every </Skip ...> annotation is "
        "binding and is quoted verbatim in Section 2.",
    )
    add_bullet(
        doc,
        "findCare/ArchitectureAndDesign/codebase-reorganization-V5.docx "
        "-- baseline (per-epic story-to-code tables, orphan tables, UX "
        "5-features-1-story analysis, EPIC-012 named Data Management "
        "Pipelines, Security recommended for own HF Space).",
    )
    add_bullet(
        doc,
        "findCare/ArchitectureAndDesign/langgraph-oo-best-practices-V6."
        "docx -- definitive LangGraph implementation philosophy "
        "(RemoteGraph, MongoDBSaver in parent runtime only, subgraph-"
        "per-Space, React-as-facade).",
    )
    add_bullet(
        doc,
        "findCare/ArchitectureAndDesign/codebase-inventory-passV2.docx "
        "-- original Skip annotations on the inventory.",
    )
    add_bullet(
        doc,
        "brain/machine_artifacts/content/agile_backlog.json -- 9-epic "
        "backlog (verified live by direct read).",
    )
    add_bullet(
        doc,
        "Code/, Website/, LangGraph/, architecture/, brain/ -- file tree "
        "walked; every cited path verified by direct filesystem listing.",
    )

    # ---------- Section 1: operating-model axiom ---------------------------
    add_h1(doc, "1. Operating-model axiom (verbatim)")
    add_para(
        doc,
        '"The operating model, the digital footprint, the general '
        "ledger, and the org chart are 100% in alignment. The structure "
        "of any one IS the structure of all four. This alignment is "
        "non-negotiable and is the fundamental organizing principle of "
        'ChatHealthy."',
    )
    add_para(
        doc,
        "Every V7 proposal is tested against the axiom: does the "
        "proposed Git-tree shape mirror what would appear on the org "
        "chart, on a general-ledger cost-center list, and on the "
        "customer-facing capability map? Where the answer is no, the "
        "structure is decoration and is removed.",
    )

    # ---------- Section 2: V6 Skip annotations resolved verbatim ----------
    add_h1(doc, "2. Architect's V6 annotations -- verbatim, addressed")
    add_para(
        doc,
        "Ten </Skip ...> annotations were found in V6 (7 in body "
        "paragraphs, 3 in table cells). Each is quoted verbatim, "
        "addressed with a concrete V7 resolution, and pointed at the "
        "V7 section where the resolution lands. The annotations drove "
        "the renaming of EPIC-002 to Security & Compliance, the dual-"
        "child SharedCode/RuntimeServer convention adopted by every "
        "cross-cutting epic, the immediate deletion of Website/"
        "roadmap.html, and the new EPIC-011-F-006 static-page-"
        "generator feature.",
    )
    for a in SKIP_ANNOTATIONS_V6:
        add_h3(doc, f"2.{a['id']} {a['location']}")
        add_para(doc, "Verbatim annotation:", bold=True)
        add_para(doc, a["verbatim"])
        add_para(doc, "V7 resolution:", bold=True)
        add_para(doc, a["resolution"])
        add_para(doc, f"Resolution lands at: {a['section_pointer']}.")

    # ---------- Section 3: epic shape after resolution --------------------
    add_h1(doc, "3. Epic shape after V6 annotation resolution")
    add_para(
        doc,
        "Seven epics, each with one cost-center line, one org-chart "
        "line, and one HF Space deployment line (capability epics) or "
        "one cross-cutting deploy bundle (cross-cutting epics). The "
        "axiom is satisfied for every row.",
    )
    add_table(
        doc,
        ["Epic ID", "Name", "Type", "Owns HF Space?", "Org-chart line"],
        [
            (
                "EPIC-001",
                "EvaluateCare",
                "Capability",
                "Yes (existing: evaluate_care)",
                "EvaluateCare product line",
            ),
            (
                "EPIC-002",
                "Security & Compliance",
                "Cross-cutting",
                "Yes (NEW per A1: own Space)",
                "CSO + Compliance Officer (bundled per A1)",
            ),
            (
                "EPIC-006",
                "FindCare",
                "Capability",
                "Yes (existing: FindCareChat backend)",
                "FindCare product line",
            ),
            (
                "EPIC-009",
                "SharedServices",
                "Cross-cutting",
                "Yes (existing: shared_services)",
                "Platform / Shared Services",
            ),
            (
                "EPIC-010",
                "Operations",
                "Cross-cutting",
                "Partial (parent LangGraph runtime + ops tools; NO "
                "separate capability Space)",
                "DevOps + SRE",
            ),
            (
                "EPIC-011",
                "UX",
                "Cross-cutting",
                "No (assets ship inside other Spaces or as static "
                "Website/)",
                "Design + Front-end",
            ),
            (
                "EPIC-012",
                "Data Management Pipelines",
                "Cross-cutting",
                "No (Azure Functions + scheduled jobs; OUT of "
                "LangGraph per LangGraph V6 Section 3.3)",
                "Data Engineering",
            ),
        ],
    )

    # ---------- Section 4: layer + folder convention ---------------------
    add_h1(doc, "4. Layer + folder convention")
    add_h2(doc, "4.1 Layer-name resolution (carry forward)")
    add_para(
        doc,
        "Earlier passes considered renaming the implementation-layer "
        "noun from 'domain' to 'rules'. 'Rules' was rejected because it "
        "collides with the engineering-rules governance object "
        "(brain/machine_artifacts/content/engineering_rules.json) and "
        "because the implementation layer holds business policy, not "
        "compliance rules. The chosen layer-name is core/ (implementation "
        "layer) within each capability. The structural noun 'domain' "
        "does not appear anywhere in the proposed V7 tree. (Legacy "
        "history: this is the resolution recorded in V4 Section 3 and "
        "V5 Section 3; V7 confirms.)",
    )

    add_h2(doc, "4.2 Dual-child convention for cross-cutting epics (per A2)")
    add_para(
        doc,
        "Every cross-cutting epic that ships both a runtime server (HF "
        "Space) and an embedded library follows the dual-child layout:",
    )
    add_bullet(
        doc,
        "{epic_root}/RuntimeServer/ -- the HF Space process. Holds "
        "app.py, Dockerfile, requirements.txt, and feature folders "
        "that expose endpoints (FastAPI handlers, LangGraph nodes "
        "that run in this Space).",
    )
    add_bullet(
        doc,
        "{epic_root}/SharedCode/ -- pure-library code that other "
        "Spaces embed at build time. No FastAPI, no transport, no "
        "Dockerfile. Pure Python (or pure TypeScript, for UX).",
    )
    add_para(
        doc,
        "A feature manifests as one of: (a) a library only (lives in "
        "SharedCode/), (b) a facade only (lives in RuntimeServer/), "
        "(c) both -- a thin facade in RuntimeServer/ delegates to the "
        "library in SharedCode/.",
    )
    add_para(
        doc,
        "The root directory is no longer named 'Code'. Capability and "
        "cross-cutting epic roots sit at the project root: "
        "EvaluateCare/, FindCare/, SharedServices/, Security/, "
        "Operations/, UX/, DataManagementPipelines/. The Website/ tree "
        "remains as the deploy artifact for static pages and the "
        "schema mirror.",
    )

    add_h2(doc, "4.3 Which epics use dual-child")
    add_table(
        doc,
        ["Epic", "Dual-child?", "Reason"],
        [
            (
                "EPIC-001 EvaluateCare",
                "Yes",
                "RuntimeServer = HF Space; SharedCode = measures library "
                "embedded by FindCare for in-line scoring",
            ),
            (
                "EPIC-002 Security & Compliance",
                "Yes",
                "RuntimeServer = auth/entitlement/audit endpoints; "
                "SharedCode = token models, policy models",
            ),
            (
                "EPIC-006 FindCare",
                "Yes + Frontend sibling",
                "RuntimeServer = FastAPI backend; SharedCode = models "
                "& repositories; Frontend = React iframe",
            ),
            (
                "EPIC-009 SharedServices",
                "Yes",
                "RuntimeServer = tool_router/llm/embeddings/safety; "
                "SharedCode = mongo_utilities, agent_framework, "
                "cost_guard (embedded by every Space)",
            ),
            (
                "EPIC-010 Operations",
                "Yes",
                "RuntimeServer = parent LangGraph runtime + deploy tools "
                "+ observability; SharedCode = engineering_rules, "
                "conversation_log, kafka, brain_governance",
            ),
            (
                "EPIC-011 UX",
                "Yes (Frontend instead of RuntimeServer)",
                "Frontend = static page generator + design paradigm; "
                "SharedCode = cross-component primitives (TS), parent_"
                "page templates, env_banner, auth_token_display",
            ),
            (
                "EPIC-012 Data Management Pipelines",
                "No",
                "Pipelines are Azure Functions + scheduled jobs, not a "
                "long-running server. Three feature folders only "
                "(get_content, orchestrate_flow, enrich_content).",
            ),
        ],
    )

    # ---------- Section 5: counts table -----------------------------------
    add_h1(doc, "5. Counts table -- V7 epic list")
    add_para(
        doc,
        "Counts after V6 annotation resolution. EPIC-002 renamed per "
        "A1. EPIC-011 gains a new feature F-006 (Static page generator) "
        "per A9 with 10 stories (one per static page) added to the "
        "design-synthesis backlog (V7 surfaces the requirement; V7 does "
        "not mutate the backlog). The EPIC-001 orphan dirs are deleted "
        "per A5; this does not change story counts.",
    )
    counts = compute_counts()
    add_table(
        doc,
        ["Epic", "Name", "# Features (existing)", "# Stories (existing)",
         "New stories surfaced for design-synthesis"],
        [
            (eid, name, counts[eid]["features"], counts[eid]["stories"],
             new_for(eid))
            for eid, name in V7_EPICS
        ],
    )
    add_para(
        doc,
        f"Total existing stories across the 7 V7 epics: "
        f"{sum(counts[e]['stories'] for e, _ in V7_EPICS)}. New stories "
        f"surfaced (UX requirements gap A9 + UX gap from V5 F-002/F-004/"
        f"F-005 + EPIC-010-F-007 brain governance per A8): "
        f"{count_new_total()}.",
    )

    # ---------- Section 6: per-epic story-to-code tables ------------------
    add_h1(doc, "6. Per-epic story-to-code tables")
    add_para(
        doc,
        "One table per V7 epic. Each row is one story. Code path(s) cite "
        "real files under Code/, Website/, LangGraph/, brain/, or "
        "architecture/. 'NO CODE FOUND' rows are listed; their "
        "disposition is in Section 8 per annotation A4.",
    )
    rows_by_final_epic = group_stories()
    for eid, ename in V7_EPICS:
        add_h2(doc, f"6.{V7_EPICS.index((eid, ename)) + 1} {eid} {ename}")
        rows = rows_by_final_epic.get(eid, [])
        if not rows:
            add_para(doc, "No stories.")
            continue
        add_table(
            doc,
            ["Feature", "Story ID", "Story description", "Code path(s)"],
            [
                (
                    f"{r['fid']} {r['fname']}",
                    r["sid"],
                    r["sd"],
                    render_code_paths(r["code"]),
                )
                for r in rows
            ],
        )

    # ---------- Section 7: per-epic orphan-code tables --------------------
    add_h1(doc, "7. Per-epic orphan-code tables")
    add_para(
        doc,
        "Files that live in each V7 epic's territory but are NOT named "
        "by any story in the backlog. The V7 'Disposition' column is "
        "the explicit per-row commitment driven by V6 annotations A5, "
        "A6, A7, A8, A9, A10. The design-synthesis agent reads the "
        "Disposition column to know what to do with each file.",
    )

    # Build dispositions per orphan path. Defaults: KEEP for capability
    # epics, KEEP for Operations except where annotation says otherwise.
    DISPO = build_dispositions()

    for eid, ename in V7_EPICS:
        add_h2(doc, f"7.{V7_EPICS.index((eid, ename)) + 1} {eid} {ename} -- orphan code")
        items = ORPHANS.get(eid, [])
        if not items:
            add_para(doc, "No orphans surfaced.")
            continue
        add_table(
            doc,
            ["File path", "Why it's in this epic's territory",
             "Why it's orphan", "V7 disposition"],
            [(p, why_terr, why_orph, DISPO.get((eid, p), "KEEP (default)"))
             for p, why_terr, why_orph in items],
        )

    # ---------- Section 8: NO-CODE-FOUND disposition list ----------------
    add_h1(doc, "8. NO-CODE-FOUND disposition list (per A4)")
    add_para(
        doc,
        "Per Skip A4: 'about no code found we should probably look to "
        "get rid of that noise'. V7 surfaces every NO-CODE-FOUND "
        "feature with an explicit disposition. V7 does NOT mutate the "
        "backlog; it lists the dispositions for the design-synthesis "
        "agent to act on.",
    )
    add_table(
        doc,
        ["Feature", "Feature name", "# stories", "V7 disposition",
         "Rationale"],
        NO_CODE_DISPOSITIONS,
    )

    # ---------- Section 9: UX requirements gap ----------------------------
    add_h1(doc, "9. UX requirements gap and EPIC-011 new features")
    add_para(
        doc,
        "V5 surfaced 6 new stories under EPIC-011 (F-002 design "
        "paradigm: 3 stories; F-004 environment banner: 2 stories; "
        "F-005 auth-token display: 1 story). V7 adds, per A9, a new "
        "F-006 'Static page generator' feature with one story per "
        "non-roadmap static page (10 stories: architecture, products, "
        "privacy, terms, security-architecture, chat-app-design, "
        "embedding-design, load-perf-report, ops-manager-design, "
        "provider-data-load).",
    )
    add_para(
        doc,
        "Total new EPIC-011 stories surfaced for design-synthesis: 16 "
        "(6 from V5 + 10 from A9). V7 does not file these stories; the "
        "design-synthesis agent does.",
    )
    add_table(
        doc,
        ["Feature", "Existing stories", "New stories surfaced", "Total"],
        [
            ("F-001 Cross-component primitives (Anchor Timer)", 1, 0, 1),
            ("F-002 Shared design paradigm", 0, 3, 3),
            ("F-003 Opening / parent page frame", 7, 0, 7),
            ("F-004 Lower-region environment banner", 0, 2, 2),
            ("F-005 Auth-token display object", 0, 1, 1),
            ("F-006 Static page generator (NEW per A9)", 0, 10, 10),
        ],
    )

    # ---------- Section 10: schema relocation plan ------------------------
    add_h1(doc, "10. Schema relocation plan")
    add_bullet(
        doc,
        "Create brain/machine_artifacts/schemas/ as a sibling of "
        "brain/machine_artifacts/content/.",
    )
    add_bullet(
        doc,
        "Move Website/schemas/ChatHealthy{AgileBacklog,Bugs,Errors,"
        "Providers,Relaxed,RiskAcceptance,Version}Schema.json + "
        "EngineeringRulesSchema.json + JsonValidationFixtureSchema."
        "json -> brain/machine_artifacts/schemas/.",
    )
    add_bullet(
        doc,
        "Move Website/schemas/dev/ -> brain/machine_artifacts/schemas/"
        "dev/.",
    )
    add_bullet(
        doc,
        "Leave Website/schemas/standard/json-schema-2020-12-meta.json "
        "where it is (third-party meta-schema; static asset).",
    )
    add_bullet(
        doc,
        "DELETE Code/Schemas/ChatHealthy.Providers.json (per A3: "
        "Website is current source of record).",
    )
    add_bullet(
        doc,
        "Set canonical $id on the relocated schema; the deploy step "
        "mirrors brain -> Website preserving the URL surface.",
    )
    add_bullet(
        doc,
        "Provider schema 11-field gap (embedding, embedding_model, "
        "embedding_version, 8 NPPES other-name fields) tracked under "
        "EPIC-012 schema-drift feature; non-blocking for relocation.",
    )
    add_bullet(
        doc,
        "Engineering Rule-008-ENF-001 currently excludes Website/"
        "schemas/standard/. After the move, update the enforcement "
        "worker's allowed/excluded patterns to point at brain/machine_"
        "artifacts/schemas/ as the validation target.",
    )

    # ---------- Section 11: master move table -----------------------------
    add_h1(doc, "11. Master code-path move table")
    add_para(
        doc,
        "Refreshed for V7 to apply the dual-child layout (per A2), the "
        "Security rename (per A1), the Code/Schemas deletion (per A3), "
        "the EvaluateCare-as-refactor-source framing (per A4), the "
        "EPIC-001 orphan deletion (per A5), the EPIC-006 empty-package "
        "deletions (per A6), the brain_governance home for manifest "
        "tools (per A8), the static-page generator feature (per A9), "
        "and the roadmap.html + generator deletion (per A10).",
    )
    add_table(
        doc,
        ["Current path", "Target epic / feature (V7)",
         "Proposed destination", "Reason"],
        MASTER_MOVES_V7,
    )

    # ---------- Section 12: LangGraph implementation alignment ------------
    add_h1(doc, "12. LangGraph implementation alignment")
    add_para(
        doc,
        "Per langgraph-oo-best-practices-V6, the runtime topology has "
        "two orthogonal axes: a deployment axis (HF Spaces, one "
        "container per capability epic) and a runtime axis (graph "
        "nodes within a single conversation). V7's file-tree shape "
        "aligns with the V6 LangGraph philosophy as follows:",
    )
    add_h2(doc, "12.1 Parent runtime ownership (LangGraph V6 Section 3.3)")
    add_bullet(
        doc,
        "Operations/RuntimeServer/orchestration/parent_graph.py is the "
        "thin parent runtime. It owns: the messages: Annotated[list, "
        "add_messages] channel; the MongoDBSaver checkpointer (one "
        "instance, one thread_id per session); the routing edges "
        "between capability nodes; the HTTP endpoint that React talks "
        "to.",
    )
    add_bullet(
        doc,
        "Capability Spaces (FindCare, EvaluateCare, Security, "
        "SharedServices) do NOT own session checkpoints. They receive "
        "a focused projection of parent state and return a focused "
        "result.",
    )
    add_h2(doc, "12.2 RemoteGraph composition (LangGraph V6 Section 3.2)")
    add_bullet(
        doc,
        "Each capability epic exposes its LangGraph as a RemoteGraph "
        "from its HF Space URL. The parent runtime composes them: "
        "find_care = RemoteGraph('find_care', url=FINDCARE_HF_URL); "
        "evaluate_care = RemoteGraph('evaluate_care', url=EVALCARE_HF_"
        "URL); security = RemoteGraph('security', url=SECURITY_HF_URL).",
    )
    add_bullet(
        doc,
        "Within a Space, conventional in-process subgraph composition "
        "applies. RemoteGraph is reserved for cross-Space.",
    )
    add_bullet(
        doc,
        "RemoteGraph never calls back into its own deployment "
        "(deadlock constraint per the V6 binding); same-Space steps "
        "use in-process composition.",
    )
    add_h2(doc, "12.3 React-as-facade (LangGraph V6 Section 3.3)")
    add_bullet(
        doc,
        "FindCare/Frontend (the iframe React app) and the parent "
        "Website/index.html call exactly one URL for conversational "
        "flow: the parent runtime's HTTP endpoint hosted by "
        "Operations/RuntimeServer.",
    )
    add_bullet(
        doc,
        "Non-conversational reads (load a provider detail, render an "
        "evaluation report) MAY call the capability Space's HTTP API "
        "directly. Those endpoints live under the capability's "
        "RuntimeServer/ tree; they are NOT LangGraph endpoints.",
    )
    add_h2(doc, "12.4 What stays out of LangGraph entirely")
    add_bullet(
        doc,
        "EPIC-012 Data Management Pipelines: scheduled jobs, ingest, "
        "dedup, embedding generation -- not conversational steps. They "
        "run under Operations infrastructure but are NOT nodes in the "
        "parent graph. No dual-child layout for EPIC-012.",
    )
    add_bullet(
        doc,
        "EPIC-010 Operations sub-features that are CI/deploy/build-"
        "counter/smoke-test tooling: under Operations/, but NOT in "
        "parent_graph.py.",
    )
    add_h2(doc, "12.5 Concrete file paths for the parent runtime")
    add_bullet(
        doc,
        "Operations/RuntimeServer/orchestration/parent_graph.py -- "
        "thin parent graph; uses langgraph.graph.StateGraph; "
        "MongoDBSaver instantiated once at module load; capability "
        "RemoteGraphs added as nodes.",
    )
    add_bullet(
        doc,
        "Operations/RuntimeServer/orchestration/state.py -- parent "
        "state Pydantic BaseModel; messages: Annotated[list, "
        "add_messages]; per-capability state projection helpers.",
    )
    add_bullet(
        doc,
        "Operations/RuntimeServer/app.py -- FastAPI front door; "
        "exposes the LangGraph SDK streaming HTTP endpoint that React "
        "consumes via @langchain/langgraph-sdk.",
    )

    # ---------- Section 13: closing summary -------------------------------
    add_h1(doc, "13. Closing summary -- what V7 commits us to")
    add_para(
        doc,
        "V7 is the input to the design-synthesis assignment. The "
        "design-synthesis agent does not need to revisit any of these "
        "decisions; it consumes them and produces the file-structure + "
        "epic-shape design.",
    )
    add_h2(doc, "13.1 Architectural commitments")
    add_bullet(
        doc,
        "Seven epics: EPIC-001 EvaluateCare, EPIC-002 Security & "
        "Compliance, EPIC-006 FindCare, EPIC-009 SharedServices, "
        "EPIC-010 Operations, EPIC-011 UX, EPIC-012 Data Management "
        "Pipelines. Each epic is one cost-center line, one org-chart "
        "line, and (for capability + Security + SharedServices + "
        "Operations) one HF Space deploy line.",
    )
    add_bullet(
        doc,
        "Dual-child {epic_root}/RuntimeServer/ + {epic_root}/SharedCode/ "
        "convention for every cross-cutting epic that ships both an "
        "embedded library and a runtime server. Pipelines (EPIC-012) "
        "do NOT use dual-child.",
    )
    add_bullet(
        doc,
        "Project-root capability folders (EvaluateCare/, FindCare/, "
        "SharedServices/, Security/, Operations/, UX/, "
        "DataManagementPipelines/). NO 'Code/' umbrella.",
    )
    add_bullet(
        doc,
        "EPIC-002 owns its own HF Space; Compliance is a feature folder "
        "under Security & Compliance, not a separate epic.",
    )
    add_bullet(
        doc,
        "Layer-name 'core/' (not 'rules/', not 'domain/'). The "
        "structural noun 'domain' does not appear in the V7 tree.",
    )
    add_bullet(
        doc,
        "LangGraph implementation: thin parent runtime in Operations/"
        "RuntimeServer/orchestration/, capability subgraphs composed "
        "via RemoteGraph from each Space, MongoDBSaver in parent only, "
        "React-as-facade via @langchain/langgraph-sdk to one URL.",
    )
    add_bullet(
        doc,
        "Brain owns schemas. Website/schemas/ is the deploy mirror. "
        "Code/Schemas/ is deleted.",
    )
    add_h2(doc, "13.2 File-system commitments")
    add_bullet(
        doc,
        "Worker.cs canonical: Operations/SharedCode/conversation_log/"
        "Worker.cs (moved from Code/Shared/ops/ChatHealthyLogService/). "
        "Code/CSharp/ChatHealthyLogService/ is DELETED.",
    )
    add_bullet(
        doc,
        "ChatHealthyMongoUtilities canonical: SharedServices/SharedCode/"
        "mongo_utilities.py. Code/DataPipelines/ChatHealthyMongoUtilities."
        "py is DELETED.",
    )
    add_bullet(
        doc,
        "Provider schema canonical: brain/machine_artifacts/schemas/"
        "ChatHealthyProvidersSchema.json. Code/Schemas/ChatHealthy."
        "Providers.json DELETED. The 11-field gap is tracked under "
        "EPIC-012, not blocking.",
    )
    add_bullet(
        doc,
        "Website/roadmap.html + Code/Shared/ops/generate_roadmap_page."
        "py DELETED per A10.",
    )
    add_bullet(
        doc,
        "EPIC-001 EvaluateCare: existing Code/evaluate_care/ tree is "
        "the refactor source, not a port. Empty orphan dirs (__init__"
        ".py, graphs/, measures/__init__.py, tests/) DELETED per A5.",
    )
    add_h2(doc, "13.3 Backlog deltas surfaced for design-synthesis")
    add_bullet(
        doc,
        "EPIC-011 F-002, F-004, F-005, F-006: 16 new stories total to "
        "anchor the shared design paradigm, env banner, auth-token "
        "display, and static page generator.",
    )
    add_bullet(
        doc,
        "EPIC-010 F-007 (Brain governance / manifest generation): 1 "
        "new story to anchor rebuild_manifest.py and manifest_generator"
        ".py per A8.",
    )
    add_bullet(
        doc,
        "Recommended retirements (no code, no business sponsor): "
        "EPIC-006-F-020, EPIC-008-F-001, EPIC-008-F-006, EPIC-008-F-009 "
        "-- 4 stories surfaced for retirement.",
    )

    add_h2(doc, "13.4 Items the design-synthesis agent decides")
    add_bullet(
        doc,
        "The exact internal feature granularity inside each "
        "{epic}/RuntimeServer/ and {epic}/SharedCode/ tree (V7 names "
        "feature-folder candidates; design-synthesis chooses the "
        "final list).",
    )
    add_bullet(
        doc,
        "The cutover sequencing for the Security HF Space carve-out "
        "(parallel-run vs. atomic-flip). V6 open question 1.",
    )
    add_bullet(
        doc,
        "The concrete shape of EvaluateCare's refactor (which "
        "scoring measures are libraries vs. facades).",
    )

    # ---------- Section 14: appendix --------------------------------------
    add_h1(doc, "14. Appendix -- generation record")
    add_bullet(
        doc,
        "Backlog read at generate time: brain/machine_artifacts/"
        "content/agile_backlog.json -- 9 epics in the source backlog "
        "(EPIC-004, EPIC-005, EPIC-007, EPIC-008 still present in the "
        "JSON; V7 maps their stories into the 7 V7 epics per the "
        "binding decisions). V7 does NOT mutate the backlog.",
    )
    add_bullet(
        doc,
        "V6 Skip annotations: 10 found, all addressed in Section 2.",
    )
    add_bullet(
        doc,
        "Read-only enforcement: no commits, no code changes, no "
        "backlog mutations. The generator script is a read-only tool "
        "that emits a Word document.",
    )

    doc.save(OUT)
    print(f"Wrote {OUT}")


# ----------------------------------------------------------------------------
# Helpers used in main()
# ----------------------------------------------------------------------------
def compute_counts():
    """Return {epic_id: {features: int, stories: int}} for the 7 V7 epics."""
    by_epic_features = defaultdict(set)
    by_epic_stories = defaultdict(int)
    for orig_epic, _ename, fid, _fname, sid, _sd in all_stories():
        feid = final_epic_for_story(orig_epic, fid, sid)
        by_epic_features[feid].add(fid)
        by_epic_stories[feid] += 1
    return {
        eid: {"features": len(by_epic_features.get(eid, set())),
              "stories": by_epic_stories.get(eid, 0)}
        for eid, _ in V7_EPICS
    }


def new_for(eid):
    if eid == "EPIC-011":
        return "16 (F-002 x3, F-004 x2, F-005 x1, F-006 x10)"
    if eid == "EPIC-010":
        return "1 (F-007 brain governance)"
    return "0"


def count_new_total():
    return 16 + 1


def group_stories():
    rows_by_final_epic = defaultdict(list)
    for orig_epic, _ename, fid, fname, sid, sd in all_stories():
        feid = final_epic_for_story(orig_epic, fid, sid)
        rows_by_final_epic[feid].append({
            "fid": fid, "fname": fname, "sid": sid, "sd": sd,
            "code": code_for_feature(fid),
        })
    return rows_by_final_epic


def build_dispositions():
    """Return {(epic_id, file_path): disposition_string} for every orphan."""
    d = {}
    # EPIC-001: A5 says all orphans DELETE.
    for path, _, _ in ORPHANS["EPIC-001"]:
        d[("EPIC-001", path)] = "DELETE (per A5)"
    # EPIC-002: KEEP (security infra, traced by inference).
    for path, _, _ in ORPHANS["EPIC-002"]:
        d[("EPIC-002", path)] = "KEEP (anchor under EPIC-002 features)"
    # EPIC-006: A6 -- empty dirs DELETE, React components KEEP.
    for path, _, _ in ORPHANS["EPIC-006"]:
        if path == "Code/find_care/" or path.endswith("__init__.py"):
            d[("EPIC-006", path)] = "DELETE (per A6: empty placeholder)"
        else:
            d[("EPIC-006", path)] = (
                "KEEP (per A6: trace under existing F-004/F-019/F-023 "
                "at refactor)"
            )
    # EPIC-009: KEEP defaults; Dockerfile/requirements move under
    # SharedServices/RuntimeServer/.
    for path, _, _ in ORPHANS["EPIC-009"]:
        if path.endswith("__pycache__/"):
            d[("EPIC-009", path)] = "DELETE (build artifact; gitignore)"
        else:
            d[("EPIC-009", path)] = (
                "KEEP (move into SharedServices/{RuntimeServer or "
                "SharedCode}/)"
            )
    # EPIC-010: A8 (rebuild_manifest, manifest_generator) -> KEEP +
    # GOVERNANCE-CRITICAL; A10 generate_roadmap_page DELETE.
    for path, _, _ in ORPHANS["EPIC-010"]:
        if "rebuild_manifest" in path or "manifest_generator" in path:
            d[("EPIC-010", path)] = (
                "KEEP -- GOVERNANCE-CRITICAL (per A8: anchor under new "
                "EPIC-010-F-007 brain governance feature; move to "
                "Operations/SharedCode/brain_governance/)"
            )
        elif "generate_roadmap_page" in path:
            d[("EPIC-010", path)] = "DELETE (per A10 cascade)"
        elif path.endswith(".log"):
            d[("EPIC-010", path)] = "DELETE (run-time log; gitignore)"
        elif "certs;C/" in path:
            d[("EPIC-010", path)] = "DELETE (malformed empty dir)"
        elif path == "Code/CSharp/ChatHealthyLogService/dev/":
            d[("EPIC-010", path)] = (
                "DELETE (per master move row 1: Code/Shared/ops/ copy "
                "is canonical)"
            )
        elif path == "Code/Shared/cost_guard.py":
            d[("EPIC-010", path)] = (
                "RECLASSIFY -- not Operations orphan; SharedServices/"
                "SharedCode/cost_guard.py owns it"
            )
        else:
            d[("EPIC-010", path)] = (
                "KEEP (move into Operations/{RuntimeServer or "
                "SharedCode}/)"
            )
    # EPIC-011: A7 -- all KEEP; A9 -- static pages KEEP-PENDING-
    # REPLACEMENT under new F-006; A10 -- roadmap.html DELETE.
    for path, _, _ in ORPHANS["EPIC-011"]:
        if path.endswith("roadmap.html"):
            d[("EPIC-011", path)] = "DELETE (per A10)"
        elif path.startswith("Website/") and path.endswith(".html"):
            d[("EPIC-011", path)] = (
                "KEEP-PENDING-REPLACEMENT (per A9: anchor under new "
                "EPIC-011-F-006 static page generator)"
            )
        else:
            d[("EPIC-011", path)] = (
                "KEEP (per A7: move to UX/SharedCode/cross_component_"
                "primitives/)"
            )
    # EPIC-012: KEEP defaults; deletions named in V5 already.
    for path, _, _ in ORPHANS["EPIC-012"]:
        if "ChatHealthyMongoUtilities" in path:
            d[("EPIC-012", path)] = "DELETE (per master move row 5)"
        elif path == "Code/Schemas/ChatHealthy.Providers.json":
            d[("EPIC-012", path)] = "DELETE (per A3)"
        elif path == "Website/schemas/":
            d[("EPIC-012", path)] = (
                "RELOCATE to brain/machine_artifacts/schemas/ (per "
                "Section 10)"
            )
        elif path.endswith(".log"):
            d[("EPIC-012", path)] = "DELETE (run-time log; gitignore)"
        elif path.endswith(".http"):
            d[("EPIC-012", path)] = "KEEP (manual REST scratch)"
        elif "favicon.ico" in path or "legal/licence.txt" in path:
            d[("EPIC-012", path)] = "DELETE (stray asset)"
        else:
            d[("EPIC-012", path)] = (
                "KEEP (move into DataManagementPipelines/{get_content "
                "| orchestrate_flow | enrich_content}/)"
            )
    return d


if __name__ == "__main__":
    main()
