"""Generate langgraph-oo-best-practices-V1.docx.

Research-backed best practices for LangChain/LangGraph implementations using
OO principles. Every recommendation cites a specific URL fetched during
the research pass on 2026-04-28.

Tag convention in citations:
    (official)  — LangChain/LangGraph maintainer-authored
    (community) — third-party blog, tutorial, or independent repo
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor


OUT = Path(__file__).with_name("langgraph-oo-best-practices-V1.docx")


# ════════════════════════════════════════════════════════════════
# Citations dictionary — every URL referenced in the document.
# ════════════════════════════════════════════════════════════════

CITES = {
    # Official LangChain / LangGraph docs and blog
    "docs_graph_api":        ("official", "Use the graph API — LangChain Docs",
                              "https://docs.langchain.com/oss/python/langgraph/use-graph-api"),
    "docs_persistence":      ("official", "Persistence — LangChain Docs",
                              "https://docs.langchain.com/oss/javascript/langgraph/persistence"),
    "docs_app_structure":    ("official", "Application structure — LangChain Docs",
                              "https://docs.langchain.com/oss/python/langgraph/application-structure"),
    "docs_tools":            ("official", "Tools — LangChain Docs",
                              "https://docs.langchain.com/oss/python/langchain/tools"),
    "docs_context":          ("official", "Context overview — LangChain Docs",
                              "https://docs.langchain.com/oss/python/concepts/context"),
    "docs_test":             ("official", "Test — LangChain Docs",
                              "https://docs.langchain.com/oss/python/langgraph/test"),
    "docs_thinking":         ("official", "Thinking in LangGraph — LangChain Docs",
                              "https://docs.langchain.com/oss/python/langgraph/thinking-in-langgraph"),
    "blog_building":         ("official", "Building LangGraph: an agent runtime from first principles",
                              "https://www.langchain.com/blog/building-langgraph"),
    "blog_v1":               ("official", "LangChain and LangGraph reach v1.0",
                              "https://www.langchain.com/blog/langchain-langgraph-1dot0"),
    "blog_klarna":           ("official", "How Klarna's AI assistant redefined customer support",
                              "https://www.langchain.com/blog/customers-klarna"),
    "blog_in_prod":          ("official", "Is LangGraph used in production?",
                              "https://www.langchain.com/blog/is-langgraph-used-in-production"),
    "blog_v02_ckpt":         ("official", "LangGraph v0.2: customization with new checkpointer libraries",
                              "https://blog.langchain.com/langgraph-v0-2/"),
    # Official GitHub repos
    "gh_langgraph":          ("official", "langchain-ai/langgraph",
                              "https://github.com/langchain-ai/langgraph"),
    "gh_example_pyproject":  ("official", "langchain-ai/langgraph-example-pyproject",
                              "https://github.com/langchain-ai/langgraph-example-pyproject"),
    "gh_open_deep_research": ("official", "langchain-ai/open_deep_research",
                              "https://github.com/langchain-ai/open_deep_research"),
    "gh_deep_research_scratch": ("official", "langchain-ai/deep_research_from_scratch",
                              "https://github.com/langchain-ai/deep_research_from_scratch"),
    "gh_deepagents":         ("official", "langchain-ai/deepagents",
                              "https://github.com/langchain-ai/deepagents"),
    "gh_swarm":              ("official", "langchain-ai/langgraph-swarm-py",
                              "https://github.com/langchain-ai/langgraph-swarm-py"),
    "gh_supervisor":         ("official", "langchain-ai/langgraph-supervisor-py",
                              "https://github.com/langchain-ai/langgraph-supervisor-py"),
    "gh_issue_1950":         ("official", "Issue #1950: Passing private state with Class Node",
                              "https://github.com/langchain-ai/langgraph/issues/1950"),
    "gh_awesome":            ("community", "von-development/awesome-LangGraph",
                              "https://github.com/von-development/awesome-LangGraph"),
    # Community / third-party
    "blog_swarnendu":        ("community", "Swarnendu De — LangGraph Best Practices",
                              "https://www.swarnendu.de/blog/langgraph-best-practices/"),
    "blog_yasin_modular":    ("community", "Yassin Hashem — Scaling AI Agents Beyond Notebooks: A Modular Architecture",
                              "https://medium.com/@yasin162001/scaling-ai-agents-beyond-notebooks-a-modular-architecture-for-langgraph-in-production-4711764de464"),
    "blog_shaza_typing":     ("community", "Shaza Ali — Type Safety in LangGraph: TypedDict vs Pydantic",
                              "https://shazaali.substack.com/p/type-safety-in-langgraph-when-to"),
    "blog_pankaj_pyd":       ("community", "Pankaj Chandravanshi — LangGraph: State with Pydantic BaseModel",
                              "https://medium.com/fundamentals-of-artificial-intelligence/langgraph-state-with-pydantic-basemodel-023a2158ab00"),
    "blog_easton_2026":      ("community", "BetterLink Blog — LangGraph State Management in Practice (2026)",
                              "https://eastondev.com/blog/en/posts/ai/20260424-langgraph-agent-architecture/"),
    "blog_markaicode":       ("community", "Markaicode — Production Multi-Agent System with LangGraph",
                              "https://markaicode.com/langgraph-production-agent/"),
    "blog_clean_state":      ("community", "Vishal Lad — Clean State Architecture in LangGraph (Mar 2026)",
                              "https://medium.com/@ladvishal1985/everything-ive-learned-about-clean-state-architecture-in-langgraph-6d1352b0c00c"),
    "blog_unit_testing":     ("community", "Anirudh Sharma — Unit Testing LangGraph: Nodes and Flow Paths",
                              "https://medium.com/@anirudhsharmakr76/unit-testing-langgraph-testing-nodes-and-flow-paths-the-right-way-34c81b445cd6"),
    "blog_mock_llm":         ("community", "Matt Carvalho — How to Properly Mock LangChain LLM Execution",
                              "https://medium.com/@matgmc/how-to-properly-mock-langchain-llm-execution-in-unit-tests-python-76efe1b8707e"),
    "blog_tools_first":      ("community", "SitePoint — Implementing the Tools-First Pattern in LangGraph",
                              "https://www.sitepoint.com/implementing-the-tools-first-pattern-in-lang-graph/"),
    "gh_fastapi_template":   ("community", "wassim249/fastapi-langgraph-agent-production-ready-template",
                              "https://github.com/wassim249/fastapi-langgraph-agent-production-ready-template"),
}


def cite(key: str) -> str:
    tag, title, url = CITES[key]
    return f"[{tag}] {title} — {url}"


# ════════════════════════════════════════════════════════════════
# Document builder helpers
# ════════════════════════════════════════════════════════════════

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def add_citation(doc, key):
    """Indented bullet that is the citation line for a recommendation."""
    tag, title, url = CITES[key]
    p = doc.add_paragraph(style="List Bullet 2")
    run = p.add_run(f"[{tag}] {title} — ")
    run.font.size = Pt(10)
    run.italic = True
    link = p.add_run(url)
    link.font.size = Pt(10)
    link.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p


# ════════════════════════════════════════════════════════════════
# Build the document
# ════════════════════════════════════════════════════════════════

def build():
    doc = Document()

    # Title
    title = doc.add_heading("LangChain / LangGraph OO Best-Practices Guide — V1", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sr = sub.add_run("Research-backed reference for refactoring procedural LangGraph "
                     "scripts into OO production code")
    sr.italic = True; sr.font.size = Pt(11)
    meta = doc.add_paragraph()
    meta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    mr = meta.add_run("Compiled 2026-04-28  ·  Skip Snow / ChatHealthy  ·  "
                      "Sources tagged [official] = LangChain/LangGraph maintainers, "
                      "[community] = third-party")
    mr.font.size = Pt(9); mr.italic = True

    # ───────────────────────────────────────────────
    # 1. Executive Summary
    # ───────────────────────────────────────────────
    add_heading(doc, "1. Executive Summary", level=1)

    add_para(doc,
        "The dominant production pattern in late-2025 / 2026 LangGraph is: "
        "(a) define state as a typed schema in its own file, with TypedDict acceptable "
        "for prototypes but a Pydantic BaseModel preferred for production validation; "
        "(b) carry conversation as messages: list[AnyMessage] with the add_messages "
        "reducer rather than as bespoke fields; (c) decompose the graph into "
        "subgraphs each owning a focused state schema (Single-Responsibility applied "
        "to graphs); (d) implement nodes as small pure functions or as callable "
        "classes with __call__ when they need encapsulated dependencies, returning "
        "partial state updates, never mutating; (e) inject runtime dependencies "
        "(DB clients, LLM clients, config) via the LangGraph context / Runtime API "
        "rather than module-level imports; (f) persist with PostgresSaver in production "
        "(or AsyncPostgresSaver) and reserve InMemorySaver / SqliteSaver for tests "
        "and local development; (g) lay the project out as state.py / nodes.py / "
        "tools.py / agent.py — the layout the official langgraph-example template "
        "ships with. The ChatHealthy POC user_journey.py — 336 lines, single file, "
        "15 flat TypedDict fields, bare top-level functions, tool outputs as "
        "top-level state keys — directly violates anti-patterns that the community "
        "literature explicitly names: the God State Object, the Monolithic Graph, "
        "and Overloaded Messages.")

    add_para(doc, "Key supporting sources for this summary:", italic=True)
    for k in ("docs_graph_api", "docs_app_structure", "docs_persistence",
              "docs_context", "blog_building", "blog_clean_state",
              "blog_easton_2026", "blog_markaicode"):
        add_citation(doc, k)

    # ───────────────────────────────────────────────
    # 2. Per-question findings
    # ───────────────────────────────────────────────
    add_heading(doc, "2. Per-Question Findings", level=1)

    # ---- Q1
    add_heading(doc, "Q1. State schema home — TypedDict, Pydantic BaseModel, dataclass, or JSON Schema?", level=2)
    add_para(doc, "Recommendation: define state in a dedicated state.py module. "
                  "TypedDict is the documented baseline and is what the LangGraph "
                  "official tutorials use first; Pydantic BaseModel is the production "
                  "upgrade once external inputs, validation, or checkpoint serialisation "
                  "matter. Dataclass is a third option, mainly chosen when Pydantic's "
                  "recursive validation cost is unacceptable.")
    add_para(doc, "Rationale grounded in the sources: the official graph-API doc states "
                  "explicitly that 'State in LangGraph can be a TypedDict, Pydantic "
                  "model, or dataclass' and presents TypedDict as the primary example. "
                  "The community guidance (Shaza Ali; Pankaj Chandravanshi; BetterLink "
                  "2026) converges on a hybrid: TypedDict for internal graph state, "
                  "Pydantic at the boundaries (inputs, outputs, tool I/O). Markaicode's "
                  "production write-up is the strongest single statement: 'Use a "
                  "Pydantic BaseModel the moment you step into production. It provides "
                  "validation, serialisation for checkpointing, and IDE autocompletion.' "
                  "All sources agree state should NOT live as a 15-field flat blob in "
                  "the graph build script.")
    for k in ("docs_graph_api", "blog_shaza_typing", "blog_pankaj_pyd",
              "blog_easton_2026", "blog_markaicode", "docs_app_structure"):
        add_citation(doc, k)

    # ---- Q2
    add_heading(doc, "Q2. State persistence — which checkpointer in production?", level=2)
    add_para(doc, "Recommendation: PostgresSaver (or AsyncPostgresSaver) for production. "
                  "InMemorySaver for tests and ephemeral local dev. SqliteSaver only "
                  "for single-process experimentation. Use a separate Store "
                  "(PostgresStore / RedisStore) for cross-thread long-term memory.")
    add_para(doc, "Rationale grounded in the sources: the official persistence doc "
                  "labels InMemorySaver 'suitable for development and testing,' "
                  "SqliteSaver 'ideal for experimentation and local workflows,' and "
                  "PostgresSaver 'used in LangSmith. Ideal for using in production.' "
                  "The official LangGraph v0.2 announcement details the on-disk "
                  "optimisations of the Postgres checkpointer (per-channel versioned "
                  "writes, cursor reads). BetterLink's 2026 write-up goes further: "
                  "'Skip SqliteSaver entirely; use PostgresSaver with async "
                  "(AsyncPostgresSaver) for production to avoid write-lock bottlenecks.' "
                  "Klarna's published case study confirms LangGraph state is "
                  "'immutable and checkpointed after every step' in production.")
    for k in ("docs_persistence", "blog_v02_ckpt", "blog_easton_2026", "blog_klarna"):
        add_citation(doc, k)

    # ---- Q3
    add_heading(doc, "Q3. Node organization — top-level functions, classes, or class-based callables?", level=2)
    add_para(doc, "Recommendation: default to small pure functions returning partial "
                  "state updates. Promote a node to a class with __call__ when it needs "
                  "constructor-injected dependencies (LLM client, DB handle, config) "
                  "or when several closely-related nodes share state through "
                  "instance attributes. Avoid the procedural top-level-function-per-step "
                  "pattern beyond a prototype.")
    add_para(doc, "Rationale grounded in the sources: the official graph-API doc "
                  "defines a node as 'just a Python function that reads our graph's "
                  "state and makes updates to it' and explicitly warns 'Nodes should "
                  "return updates to the state directly, instead of mutating the state.' "
                  "The same doc presents the class-with-__call__ pattern as a valid "
                  "alternative for parameterised nodes. Issue #1950 documents a known "
                  "limitation when class-nodes use private state, so classes are best "
                  "used for dependency-injection rather than hidden mutable internals. "
                  "The 'Thinking in LangGraph' doc reinforces fine-grained nodes "
                  "because 'durable execution creates checkpoints at node boundaries.'")
    for k in ("docs_graph_api", "docs_thinking", "gh_issue_1950"):
        add_citation(doc, k)

    # ---- Q4
    add_heading(doc, "Q4. Tool design — @tool, BaseTool, and how tool I/O lives in state", level=2)
    add_para(doc, "Recommendation: use the @tool decorator for the common case "
                  "(typed args, docstring as description). Drop to a BaseTool "
                  "subclass when you need explicit args_schema, sync+async behaviour, "
                  "or tool-level state. Wire tools into the graph via the prebuilt "
                  "ToolNode. Model tool I/O as ToolMessage entries on the messages "
                  "channel — NOT as separate top-level state fields per tool.")
    add_para(doc, "Rationale grounded in the sources: the official tools doc states "
                  "'The simplest way to create a tool is with the @tool decorator. "
                  "By default, the function's docstring becomes the tool's description.' "
                  "It identifies config and runtime as reserved parameter names, "
                  "implying tools must use ToolRuntime for dependency injection. "
                  "The community 'tools-first' write-up (SitePoint) calls out the "
                  "anti-pattern of bolting tools on after the prompt: 'Most LangGraph "
                  "agents that reach production carry a structural flaw: their tools "
                  "were bolted on after the prompt was written, resulting in "
                  "hallucinated parameters and silent failures.' The 'Clean State "
                  "Architecture' post warns explicitly against overloading the messages "
                  "list with tool outputs and recommends storing structured artefacts "
                  "(retrieved_docs, providers, trials) in dedicated typed fields, but "
                  "the LLM-visible conversation transcript stays as messages.")
    for k in ("docs_tools", "blog_tools_first", "blog_clean_state", "docs_graph_api"):
        add_citation(doc, k)

    # ---- Q5
    add_heading(doc, "Q5. Dependency injection — how do clients/config reach nodes and tools?", level=2)
    add_para(doc, "Recommendation: pass dependencies via the LangGraph runtime "
                  "context (Runtime[ContextSchema] for nodes, ToolRuntime[ContextSchema] "
                  "for tools), supplied at graph.invoke(...) time. For class-based "
                  "nodes, accept the dependency in __init__ and bind the instance into "
                  "the graph. Do NOT import database clients or LLM clients at module "
                  "scope inside node files.")
    add_para(doc, "Rationale grounded in the sources: the official Context overview "
                  "states 'Runtime context is a form of dependency injection' and shows "
                  "the canonical pattern: def node(state, runtime: Runtime[ContextSchema]) "
                  "with graph.invoke({...}, context={'user_name': 'John Smith'}). "
                  "Swarnendu De's best-practices write-up echoes this: 'Move provider "
                  "choices, model names, and feature flags to config, and inject them "
                  "with configurable inputs at runtime.' The fastapi-langgraph "
                  "production template (2.2k stars) demonstrates the same idea at "
                  "scale by isolating clients in a services/ layer that nodes consume.")
    for k in ("docs_context", "blog_swarnendu", "gh_fastapi_template"):
        add_citation(doc, k)

    # ---- Q6
    add_heading(doc, "Q6. Module / file structure", level=2)
    add_para(doc, "Recommendation: at minimum the official 4-file layout — state.py "
                  "(schema), nodes.py (node functions / classes), tools.py (tool "
                  "definitions), agent.py (graph construction). Beyond a single agent, "
                  "decompose by capability: a base/ directory for shared utilities and "
                  "a per-agent folder containing models/, node/, workflow/. "
                  "Subgraphs become first-class units of decomposition.")
    add_para(doc, "Rationale grounded in the sources: the official Application "
                  "Structure doc gives the canonical layout (state.py / nodes.py / "
                  "tools.py / agent.py inside my_agent/utils/ + my_agent/agent.py) "
                  "and states 'all project code lies within' the package directory. "
                  "The langgraph-example-pyproject template ships exactly this layout. "
                  "The 'Scaling AI Agents Beyond Notebooks' Medium piece extends the "
                  "structure for multi-agent production: per-agent models/, node/, "
                  "workflow/ folders + a base/ for shared code, justified as "
                  "'Each component lives in its own place,' enabling reuse across "
                  "workflows. The fastapi-langgraph production template (2.2k stars) "
                  "uses an even fuller layout (api/ + core/langgraph/ + services/ + "
                  "schemas/ + models/) showing how this scales into a real service.")
    for k in ("docs_app_structure", "gh_example_pyproject", "blog_yasin_modular",
              "gh_fastapi_template"):
        add_citation(doc, k)

    # ---- Q7
    add_heading(doc, "Q7. State immutability and reducers", level=2)
    add_para(doc, "Recommendation: nodes return partial-update dicts; never mutate "
                  "incoming state. Annotate any list/dict field that more than one "
                  "node may write — use add_messages for the messages channel and "
                  "operator.add (or a custom reducer) for other accumulating lists. "
                  "Cap unbounded growth via custom reducers.")
    add_para(doc, "Rationale grounded in the sources: the official graph-API doc "
                  "is unambiguous — 'Nodes should return updates to the state directly, "
                  "instead of mutating the state' and 'If no reducer function is "
                  "explicitly specified then it is assumed that all updates to the "
                  "key should override it.' The 'Building LangGraph' design post "
                  "explains the underlying Pregel/BSP model: channels are versioned "
                  "and updates are merged deterministically, which depends on "
                  "immutable returns. The community 'Clean State Architecture' post "
                  "names 'missing reducers on list fields' and 'unbounded state "
                  "growth' as explicit anti-patterns ('Always give it a reducer, "
                  "even if only one node writes to it now').")
    for k in ("docs_graph_api", "blog_building", "blog_clean_state"):
        add_citation(doc, k)

    # ---- Q8
    add_heading(doc, "Q8. Testing — node isolation, graph integration, LLM mocking", level=2)
    add_para(doc, "Recommendation: two-layer split. Unit-test nodes by calling the "
                  "function (or instance) directly with a synthetic state dict and "
                  "asserting on the returned partial update — no graph compilation "
                  "needed. Integration-test the compiled graph with an in-memory "
                  "checkpointer, mocked LLM (FakeListLLM or a custom fake), and "
                  "stubbed tools; assert on final state and the chosen edges. "
                  "Use graph.nodes['name'] for targeted in-graph node access.")
    add_para(doc, "Rationale grounded in the sources: the official Test doc says "
                  "'Compiled LangGraph agents expose references to each individual "
                  "node as graph.nodes. You can take advantage of this to test "
                  "individual nodes' and recommends per-test compilation with a fresh "
                  "checkpointer. Anirudh Sharma's testing post advocates the two-level "
                  "split: 'For node-level unit tests, you don't need pytest — you can "
                  "call the node function directly, which is fast and deterministic.' "
                  "Matt Carvalho's mocking post documents the FakeListLLM technique "
                  "and the reasons mocking .invoke() directly is brittle.")
    for k in ("docs_test", "blog_unit_testing", "blog_mock_llm"):
        add_citation(doc, k)

    # ---- Q9
    add_heading(doc, "Q9. Which OO principles does the community emphasize?", level=2)
    add_para(doc, "Recommendation: apply Single-Responsibility at the subgraph level "
                  "(each subgraph owns a focused slice of state); apply Encapsulation "
                  "via class-based nodes whose dependencies are constructor-injected; "
                  "prefer Composition over Inheritance — assemble graphs and subgraphs "
                  "rather than subclassing node base classes; use Dependency Injection "
                  "via the runtime context, not module-level imports. Treat "
                  "graph-as-a-state-machine: pure transitions, immutable updates, "
                  "explicit interfaces between subgraphs.")
    add_para(doc, "Rationale grounded in the sources: the 'Clean State Architecture' "
                  "post explicitly invokes the SRP analogy — 'Single Responsibility "
                  "Principle — the same idea used for classes and services — applied "
                  "to agent graphs.' BetterLink's 2026 piece confirms the subgraph-as-"
                  "module pattern: 'Split complex Agent into multiple Subgraphs, each "
                  "Subgraph is independent state machine, can be tested, reused "
                  "individually.' The 'Building LangGraph' design post justifies the "
                  "low-abstraction stance ('little to no abstraction at all. Instead, "
                  "we focused on control and durability') — meaning the OO discipline "
                  "comes from your code, not from a framework base class.")
    for k in ("blog_clean_state", "blog_easton_2026", "blog_building", "docs_thinking"):
        add_citation(doc, k)

    # ---- Q10
    add_heading(doc, "Q10. Dialogue / message / history modeling", level=2)
    add_para(doc, "Recommendation: model conversation as messages: "
                  "Annotated[list[AnyMessage], add_messages] — either by extending "
                  "the prebuilt MessagesState or by declaring the field manually. "
                  "Tool calls and tool results live as ToolMessage entries on the "
                  "same channel. Keep non-conversational artefacts (retrieved docs, "
                  "structured tool payloads) in separate typed fields.")
    add_para(doc, "Rationale grounded in the sources: the official graph-API doc "
                  "introduces the pattern with 'LangGraph includes a built-in reducer "
                  "add_messages' and the prebuilt MessagesState 'so that we can have' "
                  "the simpler pattern. The 'Clean State Architecture' post names "
                  "'Overloading Messages Field' as an anti-pattern: 'The messages list "
                  "is easy to append to, and before long it can end up holding "
                  "retrieved documents, intermediate reasoning, tool outputs, and more' "
                  "— prescribing dedicated fields like retrieved_docs instead. "
                  "LangChain Academy's StateGraph + MessagesState walkthrough (community "
                  "DeepWiki) confirms this is the dominant convention in tutorials.")
    for k in ("docs_graph_api", "blog_clean_state"):
        add_citation(doc, k)

    # ───────────────────────────────────────────────
    # 3. Anti-patterns
    # ───────────────────────────────────────────────
    add_heading(doc, "3. Anti-Patterns to Avoid", level=1)

    add_para(doc, "The eight anti-patterns below are each named in the cited source. "
                  "Items marked POC are concretely present in "
                  "LangGraph/poc/user_journey.py.")

    # AP-1
    add_heading(doc, "AP-1. The God State Object  (POC)", level=3)
    add_para(doc, "Quote from source: 'It's tempting to create one big state class "
                  "that every node in your graph can read and write… State that "
                  "belongs to everyone belongs to no one.' POC reality: JourneyState "
                  "has 15 unrelated fields (input, session flags, safety flags, "
                  "classifier output, six tool-output fields, dispatch metadata) "
                  "all flat in one TypedDict.")
    add_citation(doc, "blog_clean_state")

    # AP-2
    add_heading(doc, "AP-2. Monolithic Graph (no subgraph decomposition)  (POC)", level=3)
    add_para(doc, "Quote from source: 'I had a single LangGraph graph with 22 nodes "
                  "and one huge state object… adding new features and debugging "
                  "failures required tracing through state mutations across unrelated "
                  "nodes — essentially building a monolith.' POC reality: 14 nodes "
                  "and the build_user_journey() function all live in one 336-line file.")
    add_citation(doc, "blog_clean_state")

    # AP-3
    add_heading(doc, "AP-3. Tool outputs as flat top-level state fields  (POC)", level=3)
    add_para(doc, "Source guidance: tool I/O should ride the messages channel as "
                  "ToolMessage entries; structured artefacts go into named typed "
                  "fields — but only the artefact, not one field per tool wired "
                  "directly into top-level state. POC reality: location, "
                  "specialty_query, specialties, homeopathic_expansion, providers, "
                  "trials are all flat top-level keys.")
    add_citation(doc, "docs_graph_api")
    add_citation(doc, "blog_clean_state")

    # AP-4
    add_heading(doc, "AP-4. Procedural top-level node functions with no encapsulation  (POC)", level=3)
    add_para(doc, "Source guidance: nodes are pure functions OR classes-with-__call__ "
                  "when they have dependencies. The official Application Structure "
                  "puts node code under a nodes.py module and the modular-architecture "
                  "Medium piece pushes further to per-capability folders. "
                  "POC reality: every node is a top-level def in user_journey.py "
                  "with no module boundary, no class wrappers, and stub bodies.")
    add_citation(doc, "docs_app_structure")
    add_citation(doc, "blog_yasin_modular")

    # AP-5
    add_heading(doc, "AP-5. Module-level imports of clients instead of dependency injection", level=3)
    add_para(doc, "Source guidance: 'Runtime context is a form of dependency "
                  "injection… optimize the LLM context by providing dependencies "
                  "(like database connections, user IDs, or API clients) to your "
                  "tools and nodes at runtime rather than hardcoding them.' "
                  "Anti-pattern in any LangGraph code is a top-of-file import of "
                  "a singleton DB client referenced inside node bodies.")
    add_citation(doc, "docs_context")

    # AP-6
    add_heading(doc, "AP-6. Missing reducers on list fields  (POC risk)", level=3)
    add_para(doc, "Quote from source: 'If you don't set a reducer, LangGraph uses "
                  "last-write-wins by default… Always give it a reducer, even if "
                  "only one node writes to it now.' POC reality: history, "
                  "specialties, providers, trials, and homeopathic_expansion are "
                  "all bare list with no annotation; the parallel safety+classifier "
                  "fan-out is exactly the topology where last-write-wins silently "
                  "loses data.")
    add_citation(doc, "blog_clean_state")
    add_citation(doc, "docs_graph_api")

    # AP-7
    add_heading(doc, "AP-7. Overloaded messages field", level=3)
    add_para(doc, "Quote: 'The messages list is easy to append to, and before long "
                  "it can end up holding retrieved documents, intermediate reasoning, "
                  "tool outputs, and more.' Fix: keep messages for the conversation; "
                  "store retrieval/tool artefacts in named typed fields.")
    add_citation(doc, "blog_clean_state")

    # AP-8
    add_heading(doc, "AP-8. SqliteSaver / InMemorySaver in production", level=3)
    add_para(doc, "Quote (official): 'InMemorySaver: suitable for development and "
                  "testing… SqliteSaver: ideal for experimentation and local "
                  "workflows… PostgresSaver: ideal for using in production.' Quote "
                  "(community 2026): 'Skip SqliteSaver entirely; use PostgresSaver "
                  "with async (AsyncPostgresSaver) for production to avoid write-lock "
                  "bottlenecks.'")
    add_citation(doc, "docs_persistence")
    add_citation(doc, "blog_easton_2026")

    # ───────────────────────────────────────────────
    # 4. Reference implementations
    # ───────────────────────────────────────────────
    add_heading(doc, "4. Reference Implementations", level=1)

    add_para(doc, "Five real GitHub repositories that exemplify the patterns in this "
                  "guide. Star counts and dates were captured during the research "
                  "pass on 2026-04-28.")

    # R1
    add_heading(doc, "R-1. langchain-ai/langgraph (the framework itself)", level=3)
    add_para(doc, "30.7k stars. Authoritative source for the graph runtime, "
                  "checkpointer interfaces, and example notebooks under examples/. "
                  "Read for: Pregel/BSP model and channel implementation.")
    add_citation(doc, "gh_langgraph")

    # R2
    add_heading(doc, "R-2. langchain-ai/langgraph-example-pyproject", level=3)
    add_para(doc, "52 stars but officially curated. The canonical layout: "
                  "my_agent/utils/{state.py, nodes.py, tools.py} + my_agent/agent.py "
                  "+ langgraph.json. Read for: the minimal correct file structure.")
    add_citation(doc, "gh_example_pyproject")
    add_citation(doc, "docs_app_structure")

    # R3
    add_heading(doc, "R-3. langchain-ai/open_deep_research", level=3)
    add_para(doc, "11.3k stars. Production-grade research agent. src/ layout, "
                  "configurable models for distinct LLM roles (summariser, "
                  "researcher, compressor, writer), legacy/ folder preserves an "
                  "older workflow + multi-agent supervisor implementation for "
                  "reference. Read for: subgraph decomposition and configurable "
                  "model injection.")
    add_citation(doc, "gh_open_deep_research")

    # R4
    add_heading(doc, "R-4. langchain-ai/deepagents", level=3)
    add_para(doc, "21.9k stars; latest release 2026-04-15. Agent harness with "
                  "planning, filesystem, sub-agent delegation. LangGraph-native; "
                  "compiled graphs are returned for downstream composition. "
                  "Read for: how a complex agent stays composable.")
    add_citation(doc, "gh_deepagents")

    # R5
    add_heading(doc, "R-5. wassim249/fastapi-langgraph-agent-production-ready-template", level=3)
    add_para(doc, "2.2k stars (community). Realistic FastAPI service skeleton: "
                  "app/api/v1/ + app/core/langgraph/ + app/services/ + app/schemas/ "
                  "+ app/models/. Demonstrates JWT, rate limiting, Langfuse "
                  "observability, mem0+pgvector long-term memory. Read for: how "
                  "the 4-file template scales into a deployable service.")
    add_citation(doc, "gh_fastapi_template")

    # R6 (bonus index)
    add_heading(doc, "R-6 (bonus). langchain-ai/deep_research_from_scratch", level=3)
    add_para(doc, "708 stars. Tutorial repo (5 notebooks). Notable because it "
                  "explicitly uses Pydantic schemas for routing decisions "
                  "(ClarifyWithUser, ResearchQuestion). Read for: an end-to-end "
                  "Pydantic-state example.")
    add_citation(doc, "gh_deep_research_scratch")

    add_para(doc, "Index of the wider ecosystem (community-curated):")
    add_citation(doc, "gh_awesome")

    # ───────────────────────────────────────────────
    # 5. Bibliography
    # ───────────────────────────────────────────────
    add_heading(doc, "5. Citations Bibliography", level=1)

    def section(label, keys):
        add_heading(doc, label, level=2)
        for k in keys:
            tag, title, url = CITES[k]
            p = doc.add_paragraph()
            r = p.add_run(f"[{tag}] {title}\n")
            r.font.size = Pt(10)
            r.bold = True
            r2 = p.add_run(url)
            r2.font.size = Pt(10)
            r2.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    section("Official LangChain / LangGraph documentation", [
        "docs_graph_api", "docs_persistence", "docs_app_structure",
        "docs_tools", "docs_context", "docs_test", "docs_thinking",
    ])
    section("Official LangChain blog", [
        "blog_building", "blog_v1", "blog_klarna", "blog_in_prod", "blog_v02_ckpt",
    ])
    section("Official GitHub repositories", [
        "gh_langgraph", "gh_example_pyproject", "gh_open_deep_research",
        "gh_deep_research_scratch", "gh_deepagents", "gh_swarm",
        "gh_supervisor", "gh_issue_1950",
    ])
    section("Community blog posts and tutorials", [
        "blog_swarnendu", "blog_yasin_modular", "blog_shaza_typing",
        "blog_pankaj_pyd", "blog_easton_2026", "blog_markaicode",
        "blog_clean_state", "blog_unit_testing", "blog_mock_llm",
        "blog_tools_first",
    ])
    section("Community repositories", [
        "gh_awesome", "gh_fastapi_template",
    ])

    # Recency caveat
    add_heading(doc, "Recency note", level=2)
    add_para(doc, "All cited sources were retrieved on 2026-04-28. None of the "
                  "URLs cited here are older than 2024; the official LangChain "
                  "blog post 'LangGraph v0.2 checkpointers' is from 2024 and is "
                  "the oldest source in this bibliography. Three community pieces "
                  "(BetterLink Apr-2026, Vishal Lad Mar-2026, Markaicode 2025) are "
                  "explicitly 2025–2026.")

    # No-claim-without-source caveat
    add_heading(doc, "Sources searched but no firm answer found", level=2)
    add_para(doc, "Klarna's published case study confirms LangGraph use and "
                  "checkpointing posture but does not disclose state-schema choice "
                  "or directory layout. The LangGraph 1.0 announcement asserts "
                  "stability but does not pick TypedDict vs Pydantic. No claim in "
                  "this document is sourced from those gaps.")
    for k in ("blog_klarna", "blog_v1"):
        add_citation(doc, k)

    # Save
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
