# Generator for codebase-reorganization-V3.docx
# Read-only analysis output. Lineage:
#   pass3.docx (V1) -> pass3-V2.docx (Skip's annotated copy) -> V3 (this file).
# Honors all V2 binding directives Skip set in conversation since Pass 3.
# Aligned with langgraph-oo-best-practices-V3.docx (Pydantic state,
# MongoDBSaver, runtime-context DI, messages channel, subgraphs per epic,
# Safety as shared subgraph).

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(r"c:\chatHealthy\findCare\findCare\ArchitectureAndDesign\codebase-reorganization-V3.docx")

doc = Document()

# Set base style
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)


def H1(text):
    p = doc.add_heading(text, level=1)
    return p


def H2(text):
    p = doc.add_heading(text, level=2)
    return p


def H3(text):
    p = doc.add_heading(text, level=3)
    return p


def P(text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    return p


def Quote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    r = p.add_run(text)
    r.italic = True
    return p


def Bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p


def HR():
    doc.add_paragraph("_" * 60)


def Table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for r in hdr[i].paragraphs:
            for run in r.runs:
                run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    return t


# ---------------------------------------------------------------- Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Codebase Reorganization — V3")
r.bold = True
r.font.size = Pt(20)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("Date: 2026-04-28   |   Author: Claude (architecture analysis agent)   |   Read-only").italic = True

P("Lineage: codebase-reorganization-pass3.docx (V1) → codebase-reorganization-pass3-V2.docx (Skip's annotated copy) → this V3.")
P("Inputs read in order:")
Bullet("findCare/ArchitectureAndDesign/codebase-reorganization-pass3-V2.docx — Skip's </Skip> annotations on every paragraph and table cell. Treated as binding.")
Bullet("findCare/ArchitectureAndDesign/langgraph-oo-best-practices-V3.docx — research-backed LangGraph patterns (Section 6: epic-as-subgraph, MongoDBSaver, deep_research as reference, Pydantic schemas for routing).")
Bullet("findCare/ArchitectureAndDesign/codebase-inventory-passV2.docx — Pass 1 inventory + Skip's first round of </Skip> annotations.")
Bullet("brain/machine_artifacts/content/agile_backlog.json — current 9-epic shape (verified live).")
Bullet("Code/ — file tree walked three times: inventory, target-mapping, verification.")

HR()

# ---------------------------------------------------------------- Lead
H1("Lead — the operating-model axiom")
Quote('"The operating model, the digital footprint, the general ledger, and the org chart are 100% in alignment. The structure of any one IS the structure of all four. This alignment is non-negotiable and is the fundamental organizing principle of ChatHealthy."')
P(
    "Every proposal in this document is tested against the axiom: does the proposed Git-tree shape mirror what would appear on the org chart, on a "
    "general-ledger cost-center list, and on the customer-facing capability map? Where the answer is no, the structure is decoration and is removed. "
    "The word \"domain\" is the worked example: it corresponds to no department, no payroll line, and no customer-facing capability — so it disappears "
    "from the Git tree as an epic, as a folder, and as a hexagonal-layer name. The same test is applied to every architectural noun that follows."
)

HR()

# ================================================================ Section A
H1("Section A — Current-state analysis")
P(
    "Walk of the actual file tree at c:/chatHealthy/findCare/ on 2026-04-28. Every cited path was verified by direct filesystem listing or git ls-files; "
    "every \"empty\" claim was confirmed against the working tree. This section is unchanged in shape from Pass 1 / Pass 3 and is intentionally redundant "
    "with Pass 3 V2 so V3 stands as a self-contained record."
)

H2("A.1 Top-level repo layout")
Table(
    ["Path", "Role today", "Anchored to a business artifact?", "Notes"],
    [
        ["Code/", "All implementation", "No (it is a file-type bucket: \"the code\")", "7 first-level subdirs of mixed character: capability epics, file-type buckets, scratch."],
        ["Website/", "Cloudflare Pages target — static HTML + schemas", "Partially", "Website/schemas/ holds the schemas-of-record — that is a deploy detail leaking into source-of-truth."],
        ["brain/", "Governed JSON content + business artifacts", "Yes — this IS the business artifact root", "machine_artifacts/content/ holds 27 governed JSONs; enforcement_code/validate_all_json.py."],
        ["architecture/", "Design + audit docs", "Partially", "Holds EngineeringRuleEnforcement/ (V1–V20 design history), DevOpsBuildDeployAndEnvironmentManagement/, FindCare/."],
        ["LangGraph/", "POC for user-journey orchestration", "Yes — runtime tooling for the conversational front end", "Single 336-line POC: LangGraph/poc/user_journey.py + langgraph.json."],
        ["findCare/", "Holds findCare/ArchitectureAndDesign/", "No — duplicates the architecture/ root", "Three different spellings of \"FindCare\" coexist (findCare/, Code/find_care/, Code/ConversationalUX/FindCareChat/)."],
        ["docs/", "ADR + manifests + machine-brain-claude-spec.md", "Partially — ADR is a real artifact", "Small doc surface."],
        ["Analysis/", "embedding-design-2026-03-24.md + a pptx", "No — scratch", "Two tracked files."],
        ["Legal/", "Licence.txt", "Yes — license is a real artifact", "But duplicated at Code/DataPipelines/legal/licence.txt — a third home exists."],
        ["claude_thoughts/, gpt_thoughts/", ".gitkeep only", "No", "Empty placeholders."],
        ["enrichment_content/", "specialty_classification_gpt41.json", "No", "Single file; pipeline data masquerading as a top-level concern."],
        ["test_output/, debug_*.png, *.log, *.bak", "Scratch outputs", "No", "Top-level pollution."],
        ["C:temp/", "Empty Windows-path-mishap directory", "No — bug", "Confirmed empty; not in git."],
        ["start_local.bat, migrate_dev_to_qa.py, pipeline.http, pipeline_status.http, README.md, ROADMAP.md, readmeRestore.txt", "Top-level scripts and docs", "Mixed", "Operational, one-shot, dev-time fixtures, project docs."],
    ],
    widths=[1.6, 1.6, 1.7, 2.0],
)

# ---- A.2
H2("A.2 Code/ — what's actually there")
P("Code/ has 8 first-level subdirectories (CSharp, ConversationalUX, DataPipelines, Schemas, Shared, _pending, deploy, evaluate_care) plus shared_services/ and find_care/ + logging_config.py + skip_pipelineTest.http. Each is classified by what it contains, not by its name.")

H3("A.2.1 Code/CSharp/ — file-type bucket (anti-pattern)")
P("Holds exactly one thing: ChatHealthyLogService/ (a Windows logging service). Subdivided into dev/ (source) and build/ (compiled .dll/.exe/.deps.json). 6 source files tracked.")
P("The directory exists purely because the service is C# and everything else is Python/TypeScript. That is segregation by file type — Skip's V2 binding rules this anti-pattern out: \"this is the wrong idea, we should not segregate the tree by file types.\"")
P("Verified duplicate: Worker.cs differs between Code/CSharp/ChatHealthyLogService/dev/ and Code/Shared/ops/ChatHealthyLogService/. True drift.")

H3("A.2.2 Code/ConversationalUX/ — layer-name parent (anti-pattern)")
P("Two chat products living under a layer-name parent. \"ConversationalUX\" is itself an IT-internal abstraction (it is not a P&L line; it is a description of the technology surface).")
P("FindCareChat/backend (167 tracked files):")
Bullet("application/ — tool_router.py, facades/{find_care_facade.py, evaluate_care_facade.py}, tool_models/{provider_search_models.py, clinical_trials_models.py, consent_models.py}")
Bullet("domain/find_care/ — provider_search_service.py, specialty_classifier.py, specialty_ranker.py, specialty_service.py, homeopathic_resolver.py")
Bullet("domain/evaluate_care_quality/ — clinical_trials_service.py, provider_detail_service.py")
Bullet("domain/shared/ — consent/, lead_capture/, safety/, unknowns/, content/about_service.py, url/ (only __init__.py)")
Bullet("infrastructure/ — debug_logger.py, embeddings/embedding_client.py; config/, external_apis/, mongo/ are all __init__.py-only")
Bullet("adapter/ — only __init__.py (empty)")
Bullet("static/, tests/ (29 modules), main.py, url_guardian.py, Dockerfile, requirements.txt")
P("FindCareChat/frontend (8 tracked source files): Vite + TypeScript + React. index.html, src/{App.tsx, main.tsx, components/{ChatWindow.tsx, FindCareApp.tsx, GUIManager.tsx, MessageBubble.tsx}, ux (single-line file)}, public/, package.json, vite.config.ts.")
P("ChatHealthyWhoAmIChat/ (7 tracked files): me/ holds 5 PDFs/text describing Skip Snow + ChatHealthy + Anthropic principles + the business plan + the summary file; tests/ has only __init__.py. Skip's binding directive: this is foundational shared content; the LangGraph references it as tool_get_skip_snow_context. Leave content as-is on first port.")

H3("A.2.3 Code/DataPipelines/ — flat 90-file directory")
P("All Azure Functions pipeline code lives here in one flat namespace. Notable shape:")
Bullet("Entry: function_app.py (Azure Functions binding); host.json")
Bullet("Get-content workers: data_fetcher_base.py, prescriber_data_fetcher.py, gpt_reader.py, icd10_loader.py, load_specialty_data.py, zip_county_crosswalk_loader.py")
Bullet("Orchestration / cluster lifecycle: atlas_cluster_manager.py, cluster_lifecycle_manager.py, idle_monitor.py, instance_warmer.py, otp_manager.py, pipeline_worker_base.py, prescriber_pipeline_manager.py, provider_load_manager.py, sync_gateway_agent.py, ops_manager/{ops_agent.py, alert_tool.py, audit_trail.py, cluster_tool.py, evidence_package.py, triage_tool.py}")
Bullet("Enrichment workers: county_economic_enrichment.py, county_enrichment_job.py, embedding_worker.py, provider_embedding.py, crosswalk_builder.py, prescriber_enrichment_job.py, prescriber_evaluate_care_pipeline.py")
Bullet("Quality: quality_gate.py, schema_drift_detector.py, discrepancy_reporter.py, validate_provider_load.py, qa_provider_load.py")
Bullet("Delivery: copy_to_frontend.py, promote_data_fn.py")
Bullet("Analytics: count_providers_by_state.py")
Bullet("Infrastructure: auth.py, blob_client.py, ChatHealthyMongoUtilities.py")
Bullet("Tests: tests/ has 25 pytest modules covering compliance, idempotency, schema drift, quality gates, traceability, e2e")
Bullet("Misc: pipeline_code_review.json, pipeline_files.json, pipeline_health.py, prescriber-pipeline-test.http, requirements-pipeline.txt")
Bullet("Stray: legal/licence.txt (third copy of the project license), favicon.ico (orphan), overnight_pipeline.log")
P("Skip's V2 binding: pipeline divides into three natural groups — get content, orchestrate flow, enrich content. Pipeline is task-divided, not content-divided. Pipeline does NOT belong inside a front-end-associated epic. Azure Functions runtime; LangGraph for the front end. Bifurcation acceptable.")

H3("A.2.4 Code/Schemas/ — file-type bucket (anti-pattern)")
P("Contains exactly one file: ChatHealthy.Providers.json (~19KB). The actual schemas-of-record live at Website/schemas/. Anti-pattern: directory named after the file extension, like CSharp/.")

H3("A.2.5 Code/Shared/ — 90 tracked files of \"infrastructure\"")
P("This is the big one. Skip's V2 directive: \"all of these files need to be distributed to other features, and this is not the proper place for any of these files, and for the most part FindCare is not the proper place for these files.\" Code/Shared/ as a structural home is dissolved.")
P("Contents by cluster:")
Bullet("agent_framework/ — base_agent.py, base_tool.py, tool_registry.py (LangGraph/agent primitives)")
Bullet("Brain governance helpers: brain_auth.py, brain_loop.py, brain_schema_validator.py, machine_brain.py")
Bullet("LLM + cost + session: llm_client.py, cost_guard.py, prompt_system_maker.py, session_token.py")
Bullet("Mongo utility: ChatHealthyMongoUtilities.py — differs from Code/DataPipelines/ChatHealthyMongoUtilities.py (verified by diff). True drift.")
Bullet("ops/ — Caddyfile, certs/, ChatHealthyLogService/ (second C# copy), FullRegressionAudit/, kafka/ (conversation_log producer/consumer/docker-compose), atlas_cluster_toggle.py, bump_build.py, conversation_log_hook.py, daily_conversation_commit.sh, epic_planning_runner.py, framework_version.py, generate_roadmap_page.py, generate_security_page.py, hf_space_*.py (4 scripts), local_admin_server.py, manifest_generator.py, promote_build.py, promote_data.py, rebuild_manifest.py, start_local.bat, sync_to_mongo.py, uat_report.py, unattended_monitor.py")
Bullet("ops/tools/ — bash_rule_guard.py, boot_probe.py, brain_write_validator.py, build_backlog_story_view.py, capture_brain_hashes.py, chathealthy_devops_boot.py (the boot class governing every session), conversation_log_agent.py, conversation_log_purge_service.py, create_hf_space.py, gen_ops_design_docx.py, kill_dev_servers.ps1, kill_zombies.py, lineage_orchestrator.py, log_instructions_loaded.py, pipeline_trigger.py, preCommitScan.py, regression_runner.py, riskWaiver_gate.py, scan_http.py, skill/ (SKILL.md + conversation_log_agent.py), tests/")
Bullet("ux/ — components/{ProviderCard.tsx, SelectionManager.tsx}, hooks/useSelectionState.ts, types/provider.ts (TS — not consumed; the FindCareChat frontend has its own components/)")
Bullet("tests/ — test_brain_schema_validator.py only")
Bullet("Stray: ops/certs;C/ — malformed Windows-path artifact, empty, untracked")

H3("A.2.6 Code/_pending/")
P("Gitignored scratch (with one tracked .gitignore + one tracked README.docx). Subfolders: backlog_consolidation/, compliance_crosswalk/{src/, tests/}, compliance_crosswalk_third_framework/, deferred_work_pattern/, engineering_rules_redesign/{src/}.")
P("Skip's V2: \"we do need a gitignored scratch directory for experiments etc. I believe there is a bunch of places where we do this but they should all be gathered together.\" Code/_pending/ is the right home — but stray scratch lives elsewhere too (top-level test_output/, claude_thoughts/, gpt_thoughts/, debug_*.png, *.bak files).")

H3("A.2.7 Code/deploy/ — 13 tracked files")
P("Top-level deploy + smoke harness:")
Bullet("deploy_app_dev.sh (canonical, recently exercised), deploy_dev.sh, deploy_localhost.sh, deploy_localhost_docker.sh")
Bullet("_localhost_server.py, _start_findcare.py, _start_evalcare.py, _start_shared.py, _start_website.py")
Bullet("apiSmokeTestPyTest.py, devSmokeTestPyTest.py, localSmokeTestPyTest.py")
Bullet("tests/")
P("Top-level deploy/ does NOT exist (verified). Skip's binding: smoke harness belongs 100% to FindCare as a Playwright story. Multiple per-environment deploy scripts exist for now (consolidation deferred per Skip's binding). The mechanism re-homes to EPIC-010 Operations.")

H3("A.2.8 Code/evaluate_care/ — 36 tracked files")
P("Self-contained EvaluateCare FastAPI service: Dockerfile, app.py, cache.py, confidence.py, explainability.py, failsafe.py, models.py, normalization.py, provenance.py, scoring_engine.py, weights.py, plus measures/ (13 measure modules) and graphs/ subdir (empty cache only) and tests/ (6 test modules).")
P("Skip's binding: in-scope right now is the LangGraph→EvaluateCare auth-token handoff demo. The 24+ measure features (EPIC-001 has 102 stories across 25 features) remain elaborated in the backlog for later.")

H3("A.2.9 Code/find_care/ — 1 tracked file (.gitkeep)")
P("Empty placeholder. Skip's binding: \"This is to be deprecated.\" The real FindCare code lives at Code/ConversationalUX/FindCareChat/backend/domain/find_care/.")

H3("A.2.10 Code/shared_services/ — 3 tracked files")
P("Minimal HF Space deployment: app.py, Dockerfile, requirements.txt. Today essentially a thin proxy / placeholder. The naming clash with Code/Shared/ is a real source of confusion (capital-S \"Shared\" vs lowercase-s \"shared_services\" as the discriminator).")

# ---- A.3
H2("A.3 Website/")
P("Cloudflare Pages target. Tracked files include:")
Bullet("Static HTML pages: index.html, architecture.html, chat-app-design.html, embedding-design.html, load-perf-report.html, ops-manager-design.html, privacy.html, products.html, provider-data-load.html, roadmap.html, security-architecture.html, terms.html")
Bullet("_headers (Cloudflare config), Images/, _redirects (in frontend dist)")
Bullet("schemas/ — 9 *Schema.json files (AgileBacklog, Bugs, Errors, Providers, Relaxed, RiskAcceptance, Version, EngineeringRules, JsonValidationFixture) + dev/EngineeringRulesSchema.json + standard/json-schema-2020-12-meta.json")
Bullet("tests/ — conftest.py, test_website.py")
P("Two structural problems: (1) the schemas-of-record live under Website/, which is a deploy artifact; (2) several .html files are design docs that duplicate content also written in architecture/ — the website is being used as a doc-publishing surface and a marketing/legal surface at the same time.")

# ---- A.4
H2("A.4 brain/")
P("Three children: BusinessArtifacts/ + BusinessArtifacts.lnk Windows shortcut; enforcement_code/ — validate_all_json.py; machine_artifacts/ — content/ (27 governed JSONs: agile_backlog, architecture, bugs, business_plan, controlled_vocabularies, conversation_log, daily_punch_list_with_results_and_accomplishments, design, emergency_keywords, engineering_rules, errors, external_audits, governance, governance_matrix, legal, project_manifest, prompts, regression_report, risk_acceptance, schema, security, SecurityAuditControls, token_usage, traceability_matrix, unrealized_ideas, version, work_log, conversation_log_fallback.jsonl); diagrams/ (puml + svg); ClaudCodeAndCursorFiles/; tonight_context.md.")
P("This is the only top-level directory whose name unambiguously names a business artifact. The 27 JSONs in content/ are the source of truth.")

# ---- A.5
H2("A.5 LangGraph/")
P("Single child: poc/{README.md, langgraph.json, requirements.txt, user_journey.py}. The 336-line user_journey.py POC is the orchestration spine for the conversational front end. Per Skip's V2: \"We do use an Azure workflow tool and not LangGraph for the pipeline and though that bifurcates the front end and back end architectures I do not object to that.\"")

# ---- A.6
H2("A.6 architecture/ and findCare/")
P("architecture/ holds three subfolders: DevOpsBuildDeployAndEnvironmentManagement/{ArchitectureAndDesign/, code/bell_ringer.py}; EngineeringRuleEnforcement/{ArchitectureDesignAndAuditDocs/ (V1–V20 design history), code/{chathealthy_enforcement_manager.py, enforcement_worker.py, scan_files_enforcement_worker.py, requirements.txt}, locks/, tests/, __init__.py, requirements.txt}; FindCare/{ArchitectureDesignAndAuditDocs/} (mostly empty stub).")
P("findCare/ at repo root holds findCare/ArchitectureAndDesign/ — the Pass-N design docs + 2 user-journey xlsx files + the langgraph-oo-best-practices V1/V2/V3 set. This name collides with Code/find_care/ and with the FindCareChat backend. Three different spellings, three different roles.")

# ---- A.7
H2("A.7 Cross-cutting observations")
P("True duplicates (drift confirmed):")
Bullet("ChatHealthyMongoUtilities.py — Code/DataPipelines/ vs Code/Shared/ — files differ.")
Bullet("Worker.cs — Code/CSharp/ChatHealthyLogService/dev/ vs Code/Shared/ops/ChatHealthyLogService/ — files differ. (Skip's binding: ChatHealthyLogService consolidates under EPIC-010 Operations.)")
Bullet("ChatHealthy.Providers.json (Code/Schemas/) vs ChatHealthyProvidersSchema.json (Website/schemas/) — sizes differ ~11 bytes.")
Bullet("Licence.txt — Legal/Licence.txt vs Code/DataPipelines/legal/licence.txt — three top-level homes for one artifact (one tracked).")
P("Empty / near-empty directories:")
Bullet("Code/find_care/ — only .gitkeep")
Bullet("Code/Schemas/ — one file")
Bullet("Code/ConversationalUX/FindCareChat/backend/adapter/ — only __init__.py")
Bullet("Code/ConversationalUX/FindCareChat/backend/infrastructure/{config,external_apis,mongo}/ — __init__.py only")
Bullet("Code/ConversationalUX/FindCareChat/backend/domain/shared/url/ — __init__.py only")
Bullet("Code/ConversationalUX/ChatHealthyWhoAmIChat/tests/ — __init__.py only")
Bullet("claude_thoughts/, gpt_thoughts/ — .gitkeep only")
Bullet("Code/evaluate_care/graphs/ — __pycache__ only")
Bullet("C:temp/ and Code/Shared/ops/certs;C/ — empty Windows-path-mishap directories, untracked")
P("Anti-pattern folders (segregation by file type or layer, not by capability):")
Bullet("Code/CSharp/ — named after the language")
Bullet("Code/Schemas/ — named after the file format")
Bullet("Code/ConversationalUX/ — named after the layer concept")

HR()

# ================================================================ Section B
H1("Section B — Proposals")

# ---- B.0
H2("B.0 Test for every proposal")
P(
    "Each proposal in B.1–B.10 was tested against the operating-model axiom: would this folder, this epic, or this layer name appear on an HR org chart "
    "or a GL cost-center list? If not, it is decoration and is removed or renamed. The word \"domain\" fails this test (no department, no payroll line, "
    "no customer-facing capability), so it does not appear as a structural noun anywhere in the proposed Git tree. Section B.8 carries the same logic "
    "into the hexagonal layer that today is called domain/."
)

# ---- B.1
H2("B.1 Proposed epic list (post-refactor)")
P("Honors all binding decisions Skip recorded since Pass 3 V2: capability epics own full stack; Safety dissolves into Shared Services; Architecture dissolves into Operations; Testing dissolves into per-capability test features; EPIC-007 is not in the tree; EPIC-011 disposition is recommended in B.9; pipeline epic is renamed in B.7.")
Table(
    ["Epic ID", "Name (V3 proposal)", "Type", "P&L / org-chart anchor"],
    [
        ["EPIC-001", "EvaluateCare", "Capability (operational)", "Customer-facing capability with its own HuggingFace runtime and cost center. Owns full stack: scoring engine, measures, deploy, tests, architecture docs."],
        ["EPIC-002", "Security", "Cross-cutting (functional)", "Shared cost: Cloudflare findings, compliance attestation, HSTS/TLS posture. Demoted from candidate-capability — Skip's V2 binding read security as a shared concern."],
        ["EPIC-006", "FindCare", "Capability (operational)", "Customer-facing capability. Owns full stack: search code, chat front end, frame layout, HF Space runtime, tests, deploy artifacts (FindCare slice), architecture docs."],
        ["EPIC-009", "Shared Services", "Cross-cutting (functional)", "The shared back-end consumed by every capability: Tool Router, Auth/session, Embeddings, LLM client, Cost guard, Mongo utility, Skip-Snow context, Lead capture, Consent, Unknown-question, and — V3 — the Safety subgraph (emergency_detector, repetitive_detector, ip_lock_check)."],
        ["EPIC-010", "Operations", "Cross-cutting (functional)", "DevOps + observability + HF Space deployment + environment promotion + monitoring + the engineering-rules enforcement framework (V20). \"Architecture\" as a discipline lives here as a feature."],
        ["EPIC-011", "UX (recommendation in B.9)", "Cross-cutting (functional)", "Disposition recommended below: thin shared-façade epic — see B.9 for justification."],
        ["EPIC-012", "Provider Knowledge Pipeline (rename — see B.7)", "Capability (functional)", "Standalone capability. Azure Functions runtime. ONE epic; THREE features (get content, orchestrate flow, enrich content) per Skip's V2 binding. Not a FindCare child."],
    ],
    widths=[0.9, 2.1, 1.5, 2.5],
)
P("Removed from the tree (relative to current backlog):")
Bullet("EPIC-004 Safety — dissolved. The 3 stories under EPIC-004-F-001 Emergency Lockout System move INTACT to EPIC-009 Shared Services as a feature \"Safety subgraph.\" Per Skip's V2: \"this is a feature not an epic.\" LangGraph V3 Section 6.1 confirms safety detectors are properly modeled as a thin shared subgraph invoked by every capability subgraph — not as a top-level epic.")
Bullet("EPIC-005 Testing — dissolved. The 19 stories across 4 features re-home to the capability they exercise; \"Test\" is added as a feature on every capability epic.")
Bullet("EPIC-007 Talk About Care — NOT in the tree. The 1 story (EPIC-007-F-001-S-001 \"animated avatars + cameras\") has zero implementing code (verified by grep). The capability is real on the roadmap; it is not yet in code; per Skip's binding it stays in the backlog as a forward-looking item but the epic shell is not added to the Git tree until implementation begins.")
Bullet("EPIC-008 Architecture — dissolved as an epic; becomes EPIC-010-F-001 \"Engineering rules enforcement framework\" (the load-bearing V20 feature) plus EPIC-010-F-002 \"Brain governance helpers.\"")

# ---- B.2
H2("B.2 Per-epic ownership — what code each epic owns")
P("Each capability epic owns its full stack. Tests are a feature of the capability, with smoke as a story (not a feature). HF Space deployment artifacts live in EPIC-010 Operations under per-capability deployment stories; the capability owns the handshake/handoff code that lives inside its service.")
Table(
    ["Epic", "Code paths it owns (post-refactor)"],
    [
        ["EPIC-001 EvaluateCare",
         "Code/EvaluateCare/ — service code (app.py, scoring_engine.py, measures/, normalization.py, weights.py, confidence.py, explainability.py, failsafe.py, provenance.py, cache.py, models.py); handoff/ (the EvaluateCare side of the LangGraph→EvaluateCare auth-token handshake, including evaluate_care_facade.py and clinical_trials_models.py); orchestration/langgraph/ (EvaluateCare subgraph); test/ (existing tests/ modules + smoke story); architecture/EvaluateCare/ (existing architecture/FindCare/ pattern, mirrored)."],
        ["EPIC-002 Security",
         "Cloudflare config + security.txt + HSTS/TLS settings (Website/_headers); the OAuth/Auth-token feature (still being designed — Skip's V2 note: \"Authorization, which is now encapsulated by using the auth token NONCE structure but we will quickly introduce OAuth and the creation of users\"). No large code body today."],
        ["EPIC-006 FindCare",
         "Code/FindCare/ — backend/ (find-care-specific provider/specialty/homeopathic/ranker code), frontend/ (the React chat UI), parent_frame/ (when implemented — the parent page, postMessage, iframe, control frame, leaf pages), index_html_generator/, orchestration/langgraph/ (FindCare subgraph composed under the parent runtime), test/ (smoke story + Playwright SIT + backend tests + frontend tests), architecture/FindCare/ (existing architecture/FindCare/ folder consolidated). The FindCare-side mTLS handshake of the HF Space."],
        ["EPIC-009 Shared Services",
         "Code/SharedServices/ — tool_router.py (V2 binding: moved from FindCareChat backend; can route to any capability); session_token.py, llm_client.py, cost_guard.py, prompt_system_maker.py; agent_framework/ (base_agent, base_tool, tool_registry); embeddings/embedding_client.py; consent/, lead_capture/, unknown_question/, about/ (Skip-Snow context — me/ content + about_service.py); mongo_utilities.py (one canonical copy); safety/ (the safety subgraph: emergency_detector + repetitive_detector + ip_lock_check + emergency_keywords governance — invoked by every capability subgraph)."],
        ["EPIC-010 Operations",
         "Code/Operations/ — deploy/ (per-environment deploy scripts; consolidation deferred per V2 binding); local/ (start_local.bat, _start_*.py, _localhost_server.py, kill_dev_servers.ps1, kill_zombies.py); hf_space/ (create_hf_space.py, hf_space_create/delete/restart/status, pipeline_trigger.py); edge/ (Caddyfile, certs/); logging_service/ (one canonical ChatHealthyLogService); conversation_log/ (kafka/ producer+consumer, conversation_log_hook.py, conversation_log_agent.py, conversation_log_purge_service.py, daily_conversation_commit.sh); build_promotion/ (manifest_generator, rebuild_manifest, promote_build, promote_data, bump_build, framework_version, sync_to_mongo); page_generators/ (generate_roadmap_page, generate_security_page, uat_report); observability/ (debug_logger, unattended_monitor, atlas_cluster_toggle); session_governor/ (chathealthy_devops_boot.py, boot_probe.py, log_instructions_loaded.py); enforcement/ (the engineering-rules enforcement framework V20 — chathealthy_enforcement_manager.py, enforcement_worker.py, scan_files_enforcement_worker.py, bash_rule_guard.py, brain_write_validator.py, capture_brain_hashes.py, preCommitScan.py, riskWaiver_gate.py, scan_http.py, lineage_orchestrator.py — plus URL Guardian per V2 binding); brain_governance/ (brain_auth.py, brain_loop.py, brain_schema_validator.py, machine_brain.py, validate_all_json.py); doc_generators/ (build_backlog_story_view.py, gen_ops_design_docx.py, FullRegressionAudit/ as testing's regression-audit tooling); epic_planning/ (epic_planning_runner.py)."],
        ["EPIC-011 UX",
         "(Disposition recommendation in B.9.) If retained as thin shared-façade epic: Code/SharedUX/ — cross-component primitives (Anchor Timer when implemented; the shared cursor UX; the dev banner; the index.html generator if it stays cross-capability). The current Code/Shared/ux/{ProviderCard.tsx, SelectionManager.tsx, useSelectionState.ts, provider.ts} are NOT consumed today; they re-home to FindCare initially and promote here only when a second consumer exists."],
        ["EPIC-012 Provider Knowledge Pipeline",
         "Code/ProviderKnowledgePipeline/ — entire current Code/DataPipelines/ restructured under three feature-folders: get_content/ (data_fetcher_base, prescriber_data_fetcher, gpt_reader, icd10_loader, load_specialty_data, zip_county_crosswalk_loader); orchestrate_flow/ (atlas_cluster_manager, cluster_lifecycle_manager, idle_monitor, instance_warmer, otp_manager, pipeline_worker_base, prescriber_pipeline_manager, provider_load_manager, sync_gateway_agent, ops_manager/, function_app.py, host.json, pipeline_health.py, quality_gate.py, schema_drift_detector.py, discrepancy_reporter.py, validate_provider_load.py, qa_provider_load.py); enrich_content/ (county_economic_enrichment, county_enrichment_job, embedding_worker, provider_embedding, crosswalk_builder, prescriber_enrichment_job, prescriber_evaluate_care_pipeline, copy_to_frontend, promote_data_fn, count_providers_by_state); test/; architecture/ProviderKnowledgePipeline/."],
    ],
    widths=[1.5, 5.5],
)

# ---- B.3
H2("B.3 Master code-path move table")
P("Every \"current path\" was verified to exist on disk; every \"proposed destination\" aligns with B.2 ownership. Order is current top-level → child.")

# Build the move table
moves = [
    ("Code/CSharp/ChatHealthyLogService/dev/", "EPIC-010 Operations / logging_service", "Code/Operations/logging_service/", "Anti-pattern (file-type bucket). Merge with Code/Shared/ops/ChatHealthyLogService/ as canonical (Worker.cs drift must be resolved first — Open Question 3)."),
    ("Code/CSharp/ChatHealthyLogService/build/", "DELETE (build output)", "—", "Compiled .dll/.exe; should be .gitignored."),
    ("Code/ConversationalUX/FindCareChat/backend/main.py", "EPIC-006 FindCare / backend", "Code/FindCare/backend/main.py", "Capability-anchored root."),
    ("Code/ConversationalUX/FindCareChat/backend/url_guardian.py", "EPIC-010 Operations / enforcement", "Code/Operations/enforcement/url_guardian.py", "V2 binding: URL Guardian is now embedded into the engineering rules enforcement feature, which lives as a feature of Operations."),
    ("Code/ConversationalUX/FindCareChat/backend/application/tool_router.py", "EPIC-009 Shared Services", "Code/SharedServices/tool_router.py", "V2 binding: Tool Router lives in Shared Services (it can route to any capability)."),
    ("Code/ConversationalUX/FindCareChat/backend/application/facades/find_care_facade.py", "EPIC-006 FindCare / backend", "Code/FindCare/backend/find_care_facade.py", "Find-care-specific facade."),
    ("Code/ConversationalUX/FindCareChat/backend/application/facades/evaluate_care_facade.py", "EPIC-001 EvaluateCare / handoff", "Code/EvaluateCare/handoff/evaluate_care_facade.py", "EvaluateCare handoff is the EvaluateCare epic's responsibility."),
    ("Code/ConversationalUX/FindCareChat/backend/application/tool_models/provider_search_models.py", "EPIC-006 FindCare / backend", "Code/FindCare/backend/models/provider_search_models.py", "FindCare contract."),
    ("Code/ConversationalUX/FindCareChat/backend/application/tool_models/clinical_trials_models.py", "EPIC-001 EvaluateCare / handoff", "Code/EvaluateCare/handoff/clinical_trials_models.py", "EvaluateCare contract."),
    ("Code/ConversationalUX/FindCareChat/backend/application/tool_models/consent_models.py", "EPIC-009 Shared Services / consent", "Code/SharedServices/consent/consent_models.py", "Owned by EPIC-009."),
    ("Code/ConversationalUX/FindCareChat/backend/domain/find_care/", "EPIC-006 FindCare / backend", "Code/FindCare/backend/rules/ (rename — see B.8)", "\"Domain\" word disappears. Layer renamed per B.8."),
    ("Code/ConversationalUX/FindCareChat/backend/domain/evaluate_care_quality/", "EPIC-001 EvaluateCare / handoff", "Code/EvaluateCare/handoff/findcare_side/", "Cross-epic call from FindCare to EvaluateCare."),
    ("Code/ConversationalUX/FindCareChat/backend/domain/shared/consent/", "EPIC-009 Shared Services / consent", "Code/SharedServices/consent/", "Shared service."),
    ("Code/ConversationalUX/FindCareChat/backend/domain/shared/lead_capture/", "EPIC-009 Shared Services / lead_capture", "Code/SharedServices/lead_capture/", "Shared service."),
    ("Code/ConversationalUX/FindCareChat/backend/domain/shared/safety/", "EPIC-009 Shared Services / safety_subgraph", "Code/SharedServices/safety/safety_service.py", "V3 binding: Safety is a feature of Shared Services, not its own epic. The safety nodes become a shared subgraph that capability subgraphs invoke (LangGraph V3 §6.1)."),
    ("Code/ConversationalUX/FindCareChat/backend/domain/shared/unknowns/", "EPIC-009 Shared Services / unknown_question", "Code/SharedServices/unknown_question/", "Shared service."),
    ("Code/ConversationalUX/FindCareChat/backend/domain/shared/content/about_service.py", "EPIC-009 Shared Services / about", "Code/SharedServices/about/about_service.py", "Skip-Snow context tool."),
    ("Code/ConversationalUX/FindCareChat/backend/domain/shared/url/", "DELETE (empty)", "—", "Only __init__.py."),
    ("Code/ConversationalUX/FindCareChat/backend/infrastructure/debug_logger.py", "EPIC-010 Operations / observability", "Code/Operations/observability/debug_logger.py", "Operations infra."),
    ("Code/ConversationalUX/FindCareChat/backend/infrastructure/embeddings/embedding_client.py", "EPIC-009 Shared Services / embeddings", "Code/SharedServices/embeddings/embedding_client.py", "Shared embedding infrastructure."),
    ("Code/ConversationalUX/FindCareChat/backend/infrastructure/{config,external_apis,mongo}/", "DELETE (empty packages)", "—", "Only __init__.py in each."),
    ("Code/ConversationalUX/FindCareChat/backend/adapter/", "DELETE (empty)", "—", "Only __init__.py."),
    ("Code/ConversationalUX/FindCareChat/backend/static/", "EPIC-006 FindCare / backend", "Code/FindCare/backend/static/", "Backend-served assets."),
    ("Code/ConversationalUX/FindCareChat/backend/tests/ (29 modules)", "EPIC-006 FindCare / test", "Code/FindCare/test/ (smoke story = local_environment_user_ready.py + dev_environment_user_ready.py + test_sit_browser.py; remaining 26 backend tests are non-smoke stories under the same Test feature)", "V2 binding: tests are a feature of the capability; smoke is a story under the Test feature."),
    ("Code/ConversationalUX/FindCareChat/frontend/", "EPIC-006 FindCare / frontend", "Code/FindCare/frontend/", "FindCare's chat UI."),
    ("Code/ConversationalUX/FindCareChat/backend/Dockerfile", "EPIC-006 FindCare / backend", "Code/FindCare/backend/Dockerfile", "FindCare HF Space build (the EPIC-010 deploy mechanism consumes it)."),
    ("Code/ConversationalUX/ChatHealthyWhoAmIChat/me/", "EPIC-009 Shared Services / about", "Code/SharedServices/about/me/", "Foundational shared content; tool_get_skip_snow_context references it. Per V2: leave content as-is on first port."),
    ("Code/ConversationalUX/ChatHealthyWhoAmIChat/tests/", "EPIC-009 Shared Services / about", "Code/SharedServices/about/tests/", "Co-located tests."),
    ("Code/DataPipelines/ (entire flat 90-file directory)", "EPIC-012 Provider Knowledge Pipeline (rename — see B.7)", "Code/ProviderKnowledgePipeline/{get_content/, orchestrate_flow/, enrich_content/, test/}", "V2 binding: one pipeline epic; three task-divided features per Skip's natural divisions."),
    ("Code/DataPipelines/ChatHealthyMongoUtilities.py", "EPIC-009 Shared Services (canonical) / EPIC-012 (consumer)", "Code/SharedServices/mongo_utilities.py — reconcile drift first (Open Question 4)", "True drift today; EPIC-012 imports from EPIC-009."),
    ("Code/DataPipelines/legal/licence.txt", "DELETE", "—", "Third copy of license; canonical is Legal/Licence.txt."),
    ("Code/DataPipelines/favicon.ico", "DELETE", "—", "Orphan."),
    ("Code/DataPipelines/overnight_pipeline.log", "DELETE / scratch", "Code/_pending/scratch/", "Output log; should not be in repo."),
    ("Code/Schemas/ChatHealthy.Providers.json", "DELETE after reconciliation; canonical to brain/", "brain/machine_artifacts/schemas/ChatHealthyProvidersSchema.json", "Anti-pattern: file-type bucket. Reconcile ~11-byte drift first."),
    ("Code/Shared/agent_framework/", "EPIC-009 Shared Services / agent_framework", "Code/SharedServices/agent_framework/", "Tool/agent primitives the Tool Router uses."),
    ("Code/Shared/brain_auth.py, brain_loop.py, brain_schema_validator.py, machine_brain.py", "EPIC-010 Operations / brain_governance", "Code/Operations/brain_governance/", "Brain governance helpers; Architecture-as-feature-of-Operations per V2 binding."),
    ("Code/Shared/cost_guard.py, llm_client.py, prompt_system_maker.py, session_token.py", "EPIC-009 Shared Services", "Code/SharedServices/{llm_client.py, cost_guard.py, prompt_system_maker.py, session_token.py}", "Shared LLM and session infra."),
    ("Code/Shared/ChatHealthyMongoUtilities.py", "EPIC-009 Shared Services (canonical)", "Code/SharedServices/mongo_utilities.py", "Reconcile with Code/DataPipelines/ copy first (Open Question 4)."),
    ("Code/Shared/ops/Caddyfile, ops/certs/", "EPIC-010 Operations / edge", "Code/Operations/edge/{Caddyfile, certs/}", "Edge / TLS infra."),
    ("Code/Shared/ops/ChatHealthyLogService/", "EPIC-010 Operations / logging_service", "Code/Operations/logging_service/", "Reconcile Worker.cs drift with Code/CSharp/.../dev/ first (Open Question 3)."),
    ("Code/Shared/ops/FullRegressionAudit/", "EPIC-010 Operations / doc_generators (regression-audit tooling) — invoked from per-capability Test features", "Code/Operations/doc_generators/regression_audit/", "V3 binding: Testing dissolves; the regression-audit tool is operations infrastructure invoked by capability test stories."),
    ("Code/Shared/ops/kafka/", "EPIC-010 Operations / conversation_log", "Code/Operations/conversation_log/kafka/", "Conversation-log infrastructure. (Logging-as-rule is viable per V2 binding because the producer call is a Kafka publish, not a Mongo write.)"),
    ("Code/Shared/ops/atlas_cluster_toggle.py, conversation_log_hook.py, daily_conversation_commit.sh", "EPIC-010 Operations / observability and conversation_log", "Code/Operations/observability/, Code/Operations/conversation_log/", "Operations scripts."),
    ("Code/Shared/ops/hf_space_create.py, hf_space_delete.py, hf_space_restart.py, hf_space_status.py", "EPIC-010 Operations / hf_space", "Code/Operations/hf_space/", "HF Space operations (per-capability deployment stories live here)."),
    ("Code/Shared/ops/local_admin_server.py, sync_to_mongo.py, unattended_monitor.py", "EPIC-010 Operations / observability and local", "Code/Operations/observability/, Code/Operations/local/", "Operations runtime."),
    ("Code/Shared/ops/manifest_generator.py, rebuild_manifest.py, promote_build.py, promote_data.py, bump_build.py, framework_version.py", "EPIC-010 Operations / build_promotion", "Code/Operations/build_promotion/", "Build/promote/version pipeline."),
    ("Code/Shared/ops/generate_roadmap_page.py, generate_security_page.py, uat_report.py", "EPIC-010 Operations / page_generators", "Code/Operations/page_generators/", "Static page generators (operations cron)."),
    ("Code/Shared/ops/epic_planning_runner.py", "EPIC-010 Operations / epic_planning", "Code/Operations/epic_planning/", "Architecture-as-feature-of-Operations: epic-planning tooling lives here."),
    ("Code/Shared/ops/start_local.bat", "EPIC-010 Operations / local", "Code/Operations/local/start_local.bat", "Operations script."),
    ("Code/Shared/ops/tools/chathealthy_devops_boot.py + boot_probe.py + log_instructions_loaded.py", "EPIC-010 Operations / session_governor", "Code/Operations/session_governor/", "The boot class governs every Claude session; load-bearing for every commit."),
    ("Code/Shared/ops/tools/bash_rule_guard.py, brain_write_validator.py, preCommitScan.py, riskWaiver_gate.py, scan_http.py, capture_brain_hashes.py", "EPIC-010 Operations / enforcement", "Code/Operations/enforcement/", "Engineering-rule enforcement helpers (parallel to architecture/EngineeringRuleEnforcement/code/)."),
    ("Code/Shared/ops/tools/build_backlog_story_view.py, gen_ops_design_docx.py", "EPIC-010 Operations / doc_generators", "Code/Operations/doc_generators/", "Design-doc generators."),
    ("Code/Shared/ops/tools/conversation_log_agent.py + conversation_log_purge_service.py + lineage_orchestrator.py", "EPIC-010 Operations / conversation_log and enforcement", "Code/Operations/conversation_log/, Code/Operations/enforcement/", "Conversation-log + lineage operations."),
    ("Code/Shared/ops/tools/create_hf_space.py + pipeline_trigger.py", "EPIC-010 Operations / hf_space", "Code/Operations/hf_space/", "HF Space operations."),
    ("Code/Shared/ops/tools/regression_runner.py", "EPIC-010 Operations / doc_generators (cross-capability regression runner)", "Code/Operations/doc_generators/regression_audit/", "Cross-capability regression runner; per-capability tests still live in their epic's test/ folder."),
    ("Code/Shared/ops/tools/kill_dev_servers.ps1 + kill_zombies.py", "EPIC-010 Operations / local", "Code/Operations/local/", "Local dev operations."),
    ("Code/Shared/ops/tools/skill/", "EPIC-010 Operations / session_governor", "Code/Operations/session_governor/skill/", "SKILL.md is design/discipline metadata used by the boot class."),
    ("Code/Shared/ux/components/{ProviderCard.tsx, SelectionManager.tsx} + hooks/useSelectionState.ts + types/provider.ts", "EPIC-006 FindCare / frontend (initial); promote to EPIC-011 if/when a second consumer exists", "Code/FindCare/frontend/shared_ux/", "Today not consumed (FindCareChat/frontend has its own components/). Move with the only candidate consumer; promote to a shared-façade epic only on second use."),
    ("Code/Shared/tests/test_brain_schema_validator.py", "EPIC-010 Operations / brain_governance / tests", "Code/Operations/brain_governance/tests/", "Co-locate with brain_schema_validator.py."),
    ("Code/Shared/ops/certs;C/", "DELETE", "—", "Malformed Windows-path artifact, empty."),
    ("Code/Shared/ChatHealthyLogService note: see also CSharp/ row above", "EPIC-010 Operations / logging_service", "Code/Operations/logging_service/", "Single canonical source for the C# logging service."),
    ("Code/_pending/", "Stays as gitignored scratch", "Code/_pending/", "Right home. Top-level scratch (test_output/, claude_thoughts/, gpt_thoughts/, debug_*.png, *.bak, .log, enrichment_content/, Analysis/) consolidates here."),
    ("Code/deploy/deploy_*.sh", "EPIC-010 Operations / deploy", "Code/Operations/deploy/", "Per-environment deploy. V2 binding: defer the consolidation refactor; document the situation."),
    ("Code/deploy/_start_*.py + _localhost_server.py", "EPIC-010 Operations / local", "Code/Operations/local/", "Local dev process starters."),
    ("Code/deploy/apiSmokeTestPyTest.py + devSmokeTestPyTest.py + localSmokeTestPyTest.py + tests/", "EPIC-006 FindCare / test (smoke story)", "Code/FindCare/test/smoke/", "V2 binding: smoke is a story under FindCare's Test feature, not its own feature."),
    ("Code/evaluate_care/", "EPIC-001 EvaluateCare", "Code/EvaluateCare/", "Rename to capability case. Existing tests/, measures/, graphs/ move with it."),
    ("Code/find_care/ (.gitkeep)", "DELETE", "—", "Empty placeholder; V2 binding: deprecated."),
    ("Code/shared_services/ (HF Space stub)", "EPIC-009 Shared Services (rename home)", "Code/SharedServices/{app.py, Dockerfile, requirements.txt}", "Capital-S to deconflict with Code/Shared/."),
    ("Code/logging_config.py", "EPIC-009 Shared Services", "Code/SharedServices/logging_config.py", "App-wide Python logging config; consumed by every service."),
    ("Code/skip_pipelineTest.http", "EPIC-010 Operations / dev_fixtures", "Code/Operations/dev_fixtures/skip_pipelineTest.http", "Dev-time test fixture."),
    ("Website/schemas/Chathealthy*Schema.json + EngineeringRulesSchema.json + JsonValidationFixtureSchema.json", "EPIC-010 Operations / brain_governance (owner) + EPIC-010 Operations / deploy (mechanism)", "brain/machine_artifacts/schemas/ + a deploy step that mirrors brain → Website at deploy time", "V2 binding: \"the website is not the natural place for these schemas. But a deployment detail. Thus I would like to figure out how to solve this problem so the Brain JSONs and the schemas can live side by side or as parent and child.\" See B.4."),
    ("Website/schemas/standard/json-schema-2020-12-meta.json", "Stays at Website/schemas/standard/", "—", "Third-party meta-schema; legitimate static asset."),
    ("Website/schemas/dev/EngineeringRulesSchema.json", "Move with parent to brain/machine_artifacts/schemas/dev/", "—", "Dev override of canonical."),
    ("Website/{architecture, chat-app-design, embedding-design, ops-manager-design, provider-data-load, security-architecture, load-perf-report}.html", "EPIC-010 Operations / page_generators (generated; deploy mirrors to Website/)", "architecture/{respective-design-folder}/ → generator → Website/", "Design docs leaking into the marketing surface."),
    ("Website/{index, products, privacy, terms, roadmap}.html, _headers, Images/", "Stays at Website/", "Website/", "Genuine marketing/legal surface."),
    ("Website/tests/", "EPIC-006 FindCare / test (website integration)", "Code/FindCare/test/website/", "Website integration tests; the only consumer is FindCare today."),
    ("LangGraph/poc/user_journey.py + langgraph.json + requirements.txt + README.md", "Refactored: parent runtime under EPIC-010 / orchestration; per-capability subgraphs under each capability epic", "Code/Operations/orchestration/parent_graph.py + Code/FindCare/orchestration/langgraph/, Code/EvaluateCare/orchestration/langgraph/, Code/SharedServices/safety/safety_subgraph.py", "Per langgraph-oo-best-practices V3 §6.1: each capability gets its own LangGraph subgraph composed under one parent runtime via Command(graph=Command.PARENT). Safety as a shared subgraph invoked by capability subgraphs."),
    ("architecture/EngineeringRuleEnforcement/", "EPIC-010 Operations / enforcement", "stays at architecture/EngineeringRuleEnforcement/ (design-side); code/ may consolidate under Code/Operations/enforcement/", "Design-side artifacts stay; code-side consolidates with Code/Shared/ops/tools/ enforcement helpers."),
    ("architecture/DevOpsBuildDeployAndEnvironmentManagement/", "EPIC-010 Operations / architecture", "architecture/Operations/", "Renamed to match the new epic."),
    ("architecture/FindCare/ArchitectureDesignAndAuditDocs/", "EPIC-006 FindCare / architecture", "stays at architecture/FindCare/", "Design-discipline anchor for FindCare."),
    ("findCare/ArchitectureAndDesign/ (Pass 1, 2, 3, 3-V2, V3 docs + xlsx + langgraph-oo-best-practices V1/V2/V3)", "EPIC-006 FindCare / architecture (consolidate)", "Merge into architecture/FindCare/ArchitectureDesignAndAuditDocs/", "Eliminate the dual home (findCare/ vs architecture/FindCare/)."),
    ("docs/adr/, docs/manifests/, docs/machine-brain-claude-spec.md", "EPIC-010 Operations / architecture", "architecture/{adr/, manifests/, machine_brain/}", "Consolidate doc surfaces."),
    ("Analysis/, claude_thoughts/, gpt_thoughts/, enrichment_content/, test_output/", "Code/_pending/scratch/ (unified)", "Code/_pending/scratch/{analysis,claude_thoughts,gpt_thoughts,enrichment_content,test_output}/", "V2: one gitignored scratch home."),
    ("debug_*.png, evaluate_test.png, regression_test.png, unit_test_result.png, load_*.png, *.log, *.bak", "DELETE or move to Code/_pending/scratch/", "—", "Top-level pollution."),
    ("C:temp/", "DELETE", "—", "Windows-path mishap."),
    ("Legal/Licence.txt", "Stays", "Legal/Licence.txt", "Single canonical home for the project license."),
    ("brain_v0.1.3_design.json.bak (top-level)", "DELETE", "—", "Backup file; should not be in repo."),
    ("start_local.bat (top-level), pipeline.http, pipeline_status.http, migrate_dev_to_qa.py, migrate_progress.log, readmeRestore.txt", "EPIC-010 Operations / local + dev_fixtures or DELETE one-shots", "Code/Operations/local/start_local.bat; Code/Operations/dev_fixtures/ for the .http files; archive then delete one-shots", "Top-level scripts."),
    ("brain/", "Stays at brain/ + adds brain/machine_artifacts/schemas/", "—", "Owner of the JSON-of-record + their schemas (see B.4)."),
]

# Render
Table(
    ["Current path", "Target epic / feature", "Proposed destination", "Reason"],
    moves,
    widths=[1.85, 1.6, 1.85, 1.7],
)

# ---- B.4
H2("B.4 Schema relocation plan")
P("V2 directive verbatim: \"the website is not the natural place for these schemas. But a deployment detail. Thus I would like to figure out how to solve this problem so the Brain JSONs and the schemas can live side by side or as parent and child.\"")
P("Concrete proposal:")
Bullet("Create brain/machine_artifacts/schemas/ as a sibling of brain/machine_artifacts/content/.")
Bullet("Move Website/schemas/ChatHealthy{AgileBacklog,Bugs,Errors,Providers,Relaxed,RiskAcceptance,Version}Schema.json + EngineeringRulesSchema.json + JsonValidationFixtureSchema.json → brain/machine_artifacts/schemas/.")
Bullet("Move Website/schemas/dev/ → brain/machine_artifacts/schemas/dev/.")
Bullet("Leave Website/schemas/standard/json-schema-2020-12-meta.json where it is — third-party meta-schema served as a static asset.")
Bullet("Reconcile Code/Schemas/ChatHealthy.Providers.json against brain/machine_artifacts/schemas/ChatHealthyProvidersSchema.json (size differs ~11 bytes; pick canonical — Open Question 5).")
Bullet("Add a deploy step under EPIC-010 Operations that mirrors brain/machine_artifacts/schemas/ → Website/schemas/ at deploy time. Cloudflare Pages still serves the same URL surface; the source-of-truth has moved to the brain.")
Bullet("Update any code that loads schemas by URL (e.g., $schema references) to point at the deployed Website URL while validators continue to read from the brain-side path locally.")
P("Engineering Rule 8 enforcement: Rule-008-ENF-001 currently excludes Website/schemas/standard/ from JSON validation. After the move, the enforcement worker's allowed/excluded patterns must be updated to point at brain/machine_artifacts/schemas/ as the validation target.")

# ---- B.5
H2("B.5 Story disposition — what stays in EPIC-006 vs. what leaves vs. what is deleted")
P("Skip's binding: stories MUST stay intact (no merging, splitting, rewording). Below is the disposition by feature; story counts are exact (verified against agile_backlog.json).")
Table(
    ["Current EPIC-006 feature (story count)", "Disposition", "Destination epic / feature"],
    [
        ["F-001 Provider Search and Selection (3)", "STAY", "EPIC-006 FindCare — search core"],
        ["F-002 Specialty Filter (1)", "STAY", "EPIC-006 FindCare — search core"],
        ["F-003 Handoff to EvaluateCare (1)", "STAY", "EPIC-006 FindCare — handoff button (the EvaluateCare side stays in EPIC-001)"],
        ["F-004 Search Result Presentation — System-Built Summary (2)", "STAY", "EPIC-006 FindCare — result presentation"],
        ["F-005 Index Tuning — Frontend and Pipeline Clusters (3)", "MOVE", "EPIC-012 Provider Knowledge Pipeline — index tuning is a pipeline concern with a FindCare consumer"],
        ["F-006 Pipeline Lifecycle and Error Handling (16)", "MOVE", "EPIC-012 Provider Knowledge Pipeline — orchestrate flow"],
        ["F-007 Provider Data Model — In-Place Enrichment (1)", "MOVE", "EPIC-012 Provider Knowledge Pipeline — enrich content"],
        ["F-008 Incremental Provider Updates (2)", "MOVE", "EPIC-012 Provider Knowledge Pipeline — orchestrate flow"],
        ["F-009 Bulletproof ETL Process (3)", "MOVE", "EPIC-012 Provider Knowledge Pipeline — orchestrate flow"],
        ["F-010 Provider Data Quality (4)", "MOVE", "EPIC-012 Provider Knowledge Pipeline — enrich content"],
        ["F-011 Inter-Stage Quality Gates (1)", "MOVE", "EPIC-012 Provider Knowledge Pipeline — orchestrate flow"],
        ["F-013 Schema Drift Detection (1)", "MOVE", "EPIC-012 Provider Knowledge Pipeline — orchestrate flow"],
        ["F-014 Idempotent Pipeline Resume (1)", "MOVE", "EPIC-012 Provider Knowledge Pipeline — orchestrate flow"],
        ["F-015 v4-001D Compliance — DataFetcherBase (1)", "MOVE", "EPIC-012 Provider Knowledge Pipeline — get content"],
        ["F-016 CopyToFrontEnd Parity (2)", "MOVE", "EPIC-012 Provider Knowledge Pipeline — enrich content (delivery)"],
        ["F-017 Requirement-to-Test Traceability (1)", "MOVE", "EPIC-006 FindCare — Test feature (per Skip's binding: tests are a feature of the capability that owns the requirement)"],
        ["F-018 FindCarePipeline SpecialtyMetaData Orchestration (4)", "MOVE", "EPIC-012 Provider Knowledge Pipeline — enrich content"],
        ["F-019 Mobile Phone Support (5)", "STAY", "EPIC-006 FindCare — frontend"],
        ["F-020 Native Mobile Apps (0)", "DELETE the empty feature", "Empty feature; no traceability surface (Open Question 8 — Pass 3 V2 confirmed delete)"],
        ["F-021 Parent Page Frame Layout (2)", "STAY", "EPIC-006 FindCare — frontend frame"],
        ["F-022 postMessage Orchestration (2)", "STAY", "EPIC-006 FindCare — frontend frame"],
        ["F-023 Chat Iframe Wiring (2)", "STAY", "EPIC-006 FindCare — frontend frame"],
        ["F-024 Control Frame (3)", "STAY", "EPIC-006 FindCare — frontend frame"],
        ["F-025 Leaf Page Management (1)", "STAY", "EPIC-006 FindCare — frontend frame"],
        ["F-026 index.html Generated from Brain Design Record (1)", "STAY", "EPIC-006 FindCare — index.html generator"],
        ["F-027 FindCare Service (HuggingFace Space) (1)", "PARTIAL — STAY for the FindCare-side mTLS handshake; the cross-Space deploy mechanism moves to EPIC-010", "EPIC-006 FindCare for the story; EPIC-010 Operations owns the HF Space deployment mechanism"],
        ["F-028 Data Pipeline (Azure Functions) (1)", "MOVE", "EPIC-012 Provider Knowledge Pipeline — pipeline service deployment story"],
        ["F-029 Local Development Topology (5)", "MOVE", "EPIC-010 Operations — local dev topology is operations"],
        ["F-030 Provider UX Display (1)", "STAY", "EPIC-006 FindCare — result presentation"],
    ],
    widths=[2.5, 1.5, 3.0],
)
P("Resulting EPIC-006 story count: 25 stays (across F-001/-002/-003/-004/-019/-021/-022/-023/-024/-025/-026/-027/-030 + the relocated F-017 traceability story under the new Test feature).")
P("Stories moved out of EPIC-006: 40 to EPIC-012 + 5 to EPIC-010 + 1 stays at EPIC-006 (F-017 traceability re-homes inside EPIC-006 under the Test feature) = 45 stories that leave the F-005..F-018/F-028/F-029 ranges. (F-020 deletion drops 0 stories, but the empty feature shell is removed.)")
P("Stories arriving in EPIC-006: 0 (FindCare gains a new Test feature whose stories include the smoke harness, which is currently in Code/deploy/, plus the 26 backend test modules currently in Code/ConversationalUX/FindCareChat/backend/tests/ — these currently are not first-class stories, so they need to be added under EPIC-025 governance once the refactor lands; or alternatively moved as a single F-017 sibling per Skip's preference — Open Question 5).")

# ---- B.6
H2("B.6 Stories with no implementing code — deletion candidates")
P("Every claim below was verified by grep against Code/. Story IDs and patterns recorded.")
Table(
    ["Story", "Grep evidence", "Disposition"],
    [
        ["EPIC-006-F-020 Native Mobile Apps — Apple App Store and Google Play (0 stories)", "Pattern \"Native.?Mobile|App Store|Google Play\" in Code/ → 0 matches.", "DELETE the empty feature shell. Re-add when iOS/Android work begins."],
        ["EPIC-007-F-001-S-001 LLMs must have animated avatars, users must be allowed to use their cameras", "Pattern \"animated.avatar|user.camera|chatroom.ux\" in Code/ → 0 matches.", "Future story; KEEP in backlog (capability is on the roadmap), flag as no-implementation. Per binding: EPIC-007 itself is NOT in the Git tree."],
        ["EPIC-011-F-001-S-001 Anchor Timer", "Pattern \"Anchor.?Timer|cross.component.timer\" in Code/ → 0 matches.", "Future story; KEEP in backlog; flag as no-implementation. Disposition of EPIC-011 itself in B.9."],
        ["EPIC-006-F-027-S-001 + EPIC-006-F-028-S-001 + EPIC-001-F-024-S-001 mTLS for service-to-service communication", "Pattern \"mTLS|m_tls\" in Code/ → 14 matches across tests + deploy + main.py wiring; the service-to-service handshake is partially implemented but not closed out.", "KEEP all three; they are open work. None are deletion candidates."],
        ["EPIC-008-F-001 Architecture Governance (0 stories)", "0 stories in feature.", "Per V3 binding (Architecture dissolves into Operations), the empty feature shell is removed when EPIC-008 dissolves."],
        ["EPIC-008-F-006 The Brain — Agentic Infrastructure (0 stories)", "0 stories in feature.", "Same — removed on EPIC-008 dissolution."],
    ],
    widths=[2.5, 2.5, 2.0],
)
P("True deletion candidates: EPIC-006-F-020 (and the empty EPIC-008 features once that epic dissolves into EPIC-010). The other no-code stories are forward-looking and should remain in the backlog with a clear status.")

# ---- B.7
H2("B.7 Pipeline epic name — recommendation")
P("V2 binding (verbatim): \"This epic is for all data pipelines. We must name it appropriately it is a corporate capability and is not separated by domain.\"")
P("Recommendation: rename to EPIC-012 Provider Knowledge Pipeline.")
P("Reasoning:")
Bullet("\"Provider Knowledge\" names what the pipeline produces — the structured, enriched, embedded knowledge about providers (clinicians, facilities, trials, prescribers) that every front-end capability consumes. It is the corporate capability the pipeline realizes, not the technology that runs it.")
Bullet("\"Data Pipeline\" describes the technology surface (Azure Functions, ETL workers). Naming an epic by its technology violates Engineering Rule-031: \"Every epic is named for the business component it realizes, not for the technology used.\"")
Bullet("\"Provider Data Pipeline\" (the Pass 3 V2 working name) is slightly better but still anchored on the artifact \"Data\" rather than the value \"Knowledge.\" The pipeline does not produce raw data — it produces curated, embedded, scored, lineage-tracked knowledge about providers. The CFO sees this as a knowledge-asset cost center.")
Bullet("Alternative considered and rejected: \"Provider Curation.\" Curation describes one of the three task-divided features (the orchestrate-flow / enrich step), not the whole pipeline.")
Bullet("Alternative considered and rejected: \"Healthcare Knowledge Graph.\" Premature — the pipeline does not yet produce a graph; it produces relational provider records with embeddings. Reserve the name for when the graph emerges.")

# ---- B.8
H2("B.8 Layer rename — domain/ → ?")
P("V2 binding: the word \"domain\" disappears from the Git tree. Skip's working candidate is rules/. Recommendation below.")
P("Recommendation: rename domain/ → core/.")
P("Reasoning:")
Bullet("\"Core\" passes the operating-model axiom: it names the part of the capability that IS the capability — the business logic that any reasonable CFO/COO would describe as \"the FindCare engine\" or \"the EvaluateCare scoring core.\" It maps cleanly to a P&L line (this is what the cost center is paying for).")
Bullet("\"Rules\" (Skip's working candidate) is close — it names the right concept (the deterministic, governed business rules) but it conflates with brain/machine_artifacts/content/engineering_rules.json. A new engineer reading rules/find_care/specialty_classifier.py would reasonably ask \"is this code an engineering rule?\" Naming clash with an active artifact is a real cost.")
Bullet("LangGraph V3 §6.1 reinforces the choice: \"each subgraph owns a focused state schema\" and \"compiled subgraphs are first-class units of decomposition.\" The capability core is the body of code the subgraph wraps; \"core\" matches the language in the LangGraph community literature (e.g., \"agent core,\" \"tool core\").")
Bullet("Alternative considered: \"engine\" (e.g., find_care_engine/). Strong second choice. Slightly more anthropomorphic than \"core\" — works well for EvaluateCare (scoring engine) but feels heavy for FindCare (which is more retrieval + classification than a heavy-engine).")
Bullet("Alternative considered and rejected: \"services\" (e.g., find_care/services/). Conflates with EPIC-009 Shared Services and with the existing *_service.py file naming convention.")
Bullet("Alternative considered and rejected: \"logic\" (e.g., find_care/logic/). Vague; fails the axiom test (no CFO writes a check to \"the logic department\").")
P("Resulting structure example: Code/FindCare/backend/core/{provider_search_service.py, specialty_classifier.py, specialty_ranker.py, specialty_service.py, homeopathic_resolver.py}. Code/EvaluateCare/core/{scoring_engine.py, normalization.py, ...}.")

# ---- B.9
H2("B.9 EPIC-011 UX — disposition recommendation")
P("V2 binding (verbatim): \"this gets a bit fuzzy because we have tightly coupled the UX parts of the React app with the orchestration part of the React app. I believe we are moving the orchestration aspects to LangGraph, and if we do then the React app becomes a façade for the actual orchestration taking place in LangGraph.\"")
P("Recommendation: dissolve EPIC-011 UX into the capability epics, with one narrow exception. The default home for UX components is the capability that consumes them. A UX component is promoted into a thin shared-façade epic ONLY when a second consumer materializes.")
P("Reasoning:")
Bullet("Today there is exactly one consumer for every UX component in the repo. Code/Shared/ux/{ProviderCard.tsx, SelectionManager.tsx, useSelectionState.ts, provider.ts} is currently consumed by zero pages — FindCareChat/frontend has its own components/. There is no shared cost line.")
Bullet("The operating-model axiom requires every cross-cutting epic to correspond to a real shared cost. A cross-cutting UX epic that holds two unused components and one unimplemented Anchor Timer story is decoration.")
Bullet("Once orchestration moves to LangGraph (per V2 binding), the React app becomes a façade that renders capability subgraph state. The natural home for FindCare UX is then Code/FindCare/frontend/. The natural home for EvaluateCare UX (when built) is Code/EvaluateCare/frontend/. The natural home for Talk About Care UX is Code/TalkAboutCare/frontend/ (when that epic enters the tree).")
Bullet("The narrow exception: when a second consumer materializes (e.g., the Anchor Timer is consumed by both FindCare and Talk About Care, or the dev banner is consumed by every capability frontend), promote the component to a thin shared-façade epic Code/SharedUX/. The promotion rule must require two real consumers, not a hypothetical second consumer.")
Bullet("Stories that survive: EPIC-011-F-001-S-001 Anchor Timer (no code; flagged as no-implementation in B.6) stays in the backlog as a forward-looking item under whatever Code/SharedUX/ epic exists when the second consumer arrives. Until then, it is parked.")
P("Net effect: EPIC-011 UX is removed from the proposed epic list in B.1 with a note that it is reinstated as Code/SharedUX/ on the second-consumer trigger. (If Skip prefers to keep EPIC-011 as a thin placeholder so the Anchor Timer story has a home today, that is the alternative — see Open Question 1.)")

# ---- B.10
H2("B.10 LangGraph implementation alignment")
P("Cross-checked against findCare/ArchitectureAndDesign/langgraph-oo-best-practices-V3.docx. Each row maps a V3 best-practice recommendation to the V3 codebase-organization proposal.")
Table(
    ["LangGraph V3 best practice", "Where it lands in the V3 codebase organization"],
    [
        ["State as Pydantic models in their own files (V3 §Q1 + §6.1)",
         "Code/{capability}/orchestration/langgraph/state.py per capability subgraph. The current LangGraph/poc/user_journey.py JourneyState (15 flat TypedDict fields) decomposes into per-subgraph Pydantic state classes (FindCareState, EvaluateCareState, SafetyState, talk_about_care_state when EPIC-007 enters)."],
        ["MongoDBSaver for persistence, not PostgresSaver (V3 §Q2 + §6.2)",
         "The parent runtime in Code/Operations/orchestration/parent_graph.py instantiates a single MongoDBSaver (langgraph-checkpoint-mongodb, MongoDB+LangChain co-maintained). Capability subgraphs share the same checkpointer instance under namespaced keys (the {node_name}:{task_id} pattern). No new database introduced."],
        ["Runtime-context dependency injection — no module-level imports of clients in node bodies (V3 §Q5 + AP-5)",
         "Node files in Code/{capability}/orchestration/langgraph/nodes.py accept Runtime[ContextSchema] (or ToolRuntime[ContextSchema]) at invocation time. The DI surface is owned by Code/SharedServices/ (LLM client, Mongo client, embedding client). No node body imports a singleton client."],
        ["Tool I/O on the messages: Annotated[list[AnyMessage], add_messages] channel — not flat top-level state fields (V3 §Q4 + §Q10 + AP-3)",
         "FindCareState (and every capability state) extends MessagesState. Tool outputs (provider list, trial list, specialty list, homeopathic expansion) ride as ToolMessage entries on the messages channel; only structured artefacts that need a typed home land in named fields (e.g., selected_providers as a typed list with an explicit reducer). The current POC's flat tool-output fields (location, specialty_query, specialties, providers, trials, homeopathic_expansion, response) collapse into ToolMessage entries."],
        ["Subgraph per capability composed under one parent runtime (V3 §6.1)",
         "Code/Operations/orchestration/parent_graph.py compiles a thin parent that adds each capability subgraph as a node via add_node(find_care_subgraph). Handoffs between epics use Command(graph=Command.PARENT). input_schema / output_schema on each subgraph define what crosses the boundary; keys NOT in the schema do not flow."],
        ["Safety as a reusable shared subgraph invoked by capability graphs (V3 §6.1; rejects safety-as-its-own-epic)",
         "Code/SharedServices/safety/safety_subgraph.py compiles once and is added as a node by FindCare, EvaluateCare, and (when implemented) Talk About Care. The 3 safety detectors (emergency_detector, repetitive_detector, ip_lock_check) live inside this single shared subgraph. Aligns with EPIC-004 Safety dissolving into EPIC-009 Shared Services per V3 binding."],
        ["Bespoke subgraphs per capability (V3 §6.1: \"Use multiple agent subgraphs when you need bespoke agent implementations\")",
         "Justified: each capability's structural shape genuinely differs. FindCare is retrieval + classification + ranking. EvaluateCare is deterministic scoring with provenance. Talk About Care is meeting + deep-research with interrupt() + Send fan-out. A single agent with middleware is the wrong pattern."],
        ["Application-structure layout: state.py / nodes.py / tools.py / agent.py (V3 §Q6 + langgraph-example-pyproject)",
         "Each capability orchestration folder follows the official 4-file layout: Code/{capability}/orchestration/langgraph/{state.py, nodes.py, tools.py, agent.py} + langgraph.json at the orchestration root."],
        ["Reducers on every list field, especially under fan-out (V3 §Q7 + AP-6)",
         "Every list field in every capability state schema gets an explicit reducer (add_messages for messages; operator.add or a custom bounded reducer for everything else). The POC's bare list fields (history, specialties, providers, trials, homeopathic_expansion) under the parallel safety+classifier fan-out are exactly the topology where last-write-wins silently loses data — this risk is closed by the per-subgraph state schemas."],
        ["Deep-research as a reusable subgraph (V3 §6.3, NOT an anti-pattern)",
         "When Talk About Care lands: Code/SharedServices/deep_research/deep_research_subgraph.py imported as a node by Talk About Care, FindCare (cancer-research), and EvaluateCare (cutting-edge-evidence). Uses Send fan-out + interrupt() so the meeting transcript continues while research workers run concurrently. Single MongoDBSaver covers all suspended subgraph state."],
    ],
    widths=[3.5, 3.5],
)

HR()

# ================================================================ Section C
H1("Section C — Open questions for Skip")
P("Each is answerable with one decision.")
Table(
    ["#", "Question", "Why it matters"],
    [
        ["1", "EPIC-011 UX disposition — confirm the recommendation in B.9 (dissolve into capabilities; reinstate as Code/SharedUX/ on the second-consumer trigger) OR keep EPIC-011 as a thin placeholder so the Anchor Timer story has a home today.", "Affects whether EPIC-011 appears in the post-refactor epic list at all."],
        ["2", "Pipeline epic name — confirm \"Provider Knowledge Pipeline\" (B.7 recommendation) OR pick an alternative (\"Provider Curation,\" \"Healthcare Knowledge Pipeline,\" \"Provider Catalog Pipeline,\" or your own name).", "Renames Code/DataPipelines/ → Code/{NewName}/ and renames the epic in the backlog."],
        ["3", "Layer name — confirm \"core\" (B.8 recommendation) OR \"rules\" (working candidate) OR \"engine\" (alternative). The word \"domain\" is removed regardless.", "Renames every domain/ folder under capability backends."],
        ["4", "ChatHealthyLogService canonical: Worker.cs differs between Code/CSharp/ChatHealthyLogService/dev/ and Code/Shared/ops/ChatHealthyLogService/. Which is source-of-truth?", "Cannot consolidate to Code/Operations/logging_service/ until you pick one."],
        ["5", "ChatHealthyMongoUtilities.py canonical: differs between Code/Shared/ and Code/DataPipelines/. Which is source-of-truth?", "Cannot consolidate to Code/SharedServices/mongo_utilities.py until you pick one."],
        ["6", "Code/Schemas/ChatHealthy.Providers.json (~19,265 bytes) vs Website/schemas/ChatHealthyProvidersSchema.json (~19,276 bytes): which is canonical, and is the resolved file the one that goes to brain/machine_artifacts/schemas/?", "Schema relocation (B.4) cannot start until this is decided."],
    ],
    widths=[0.4, 3.6, 3.0],
)

HR()

# ================================================================ Appendix
H1("Appendix — verification record")
P("Every cited path was verified to exist (or, in the case of \"empty\" or \"missing\", verified absent) by direct filesystem listing and by git ls-files. Notable verifications repeated from Pass 3 V2:")
Bullet("Top-level deploy/ does NOT exist — only Code/deploy/.")
Bullet("C:temp/ exists at repo root, is empty, and is not in git.")
Bullet("Code/Shared/ops/certs;C/ exists, is empty, and is not in git.")
Bullet("Code/find_care/ contains exactly one tracked file (.gitkeep) — verified by git ls-files.")
Bullet("ChatHealthyMongoUtilities.py: diff between Code/Shared/ and Code/DataPipelines/ confirmed differs.")
Bullet("ChatHealthyLogService: diff between Code/CSharp/.../dev/ and Code/Shared/ops/.../ confirms drift.")
Bullet("Backlog: 9 epics. EPIC-001 has 25 features and 102 stories (verified). EPIC-006 has 29 features and 71 stories (verified). The pipeline-flavored EPIC-006 features alone (F-005, F-006, F-007, F-008, F-009, F-010, F-011, F-013, F-014, F-015, F-016, F-018, F-028) account for 40 stories — Skip's V2 \"if a pipeline epic has 39 stories, the shape is wrong\" constraint is honored by re-homing them to EPIC-012.")
Bullet("EPIC-007 Talk About Care: 1 feature, 1 story, 0 implementing matches in Code/. Per binding, NOT in tree.")
Bullet("EPIC-011 UX: 1 feature, 1 story (Anchor Timer), 0 implementing matches in Code/.")
Bullet("EPIC-008 Architecture: 10 features, 51 stories. F-001 and F-006 have 0 stories (empty feature shells). Per binding, dissolves into EPIC-010 Operations.")

# Save
doc.save(str(OUT))
print(f"Wrote {OUT}")
print(f"Size: {OUT.stat().st_size} bytes")
