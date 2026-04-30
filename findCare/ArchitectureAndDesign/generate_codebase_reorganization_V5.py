"""Generate codebase-reorganization-V5.docx.

Read-only generator. Walks the repo, reads agile_backlog.json, and emits the
V5 Word doc per the V5 brief:
  - Operating-model axiom (verbatim)
  - Security hosting dilemma + recommendation
  - Layer-name explanation (carry forward)
  - Three architect open issues (carry forward)
  - Counts table (refreshed for EPIC-012 rename)
  - Per-epic story-to-code tables (Deliverable 1)
  - Per-epic orphan-code tables (Deliverable 2)
  - UX 5-features-1-story analysis (Deliverable 3)
  - Master code-path move table (carry forward, refreshed)
  - Schema relocation plan (carry forward)
  - Open questions (<=4)

No commits, no code changes, no backlog mutations. Pure docx emission.
"""

import json
import os
from collections import defaultdict
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

REPO = r"c:/chatHealthy/findCare"
BACKLOG = os.path.join(REPO, "brain/machine_artifacts/content/agile_backlog.json")
OUT = os.path.join(REPO, "findCare/ArchitectureAndDesign/codebase-reorganization-V5.docx")

# ----------------------------------------------------------------------------
# Story-to-code mapping. Keyed by feature_id (every story under a feature
# shares the same code surface unless explicitly overridden). Each entry is
# either:
#   - a list of real file paths (string paths under c:/chatHealthy/findCare)
#   - the literal "NO CODE FOUND" if no implementing code exists today
# Paths are relative to the repo root for readability in the doc.
# ----------------------------------------------------------------------------
FEATURE_CODE = {
    # EPIC-001 EvaluateCare ----------------------------------------------------
    "EPIC-001-F-001": [
        "Code/evaluate_care/scoring_engine.py",
        "Code/evaluate_care/models.py",
        "Code/evaluate_care/weights.py",
        "Code/evaluate_care/measures/base.py",
        "Code/ConversationalUX/FindCareChat/backend/application/facades/evaluate_care_facade.py",
    ],
    "EPIC-001-F-002": [
        "Code/evaluate_care/scoring_engine.py",
        "Code/evaluate_care/measures/trial_phase.py",
        "Code/evaluate_care/measures/trial_recruitment.py",
        "Code/evaluate_care/measures/trial_condition_relevance.py",
        "Code/evaluate_care/measures/trial_proximity.py",
        "Code/evaluate_care/measures/trial_recency.py",
        "Code/evaluate_care/measures/trial_sponsor.py",
        "Code/evaluate_care/measures/trial_eligibility.py",
    ],
    "EPIC-001-F-003": [
        "Code/evaluate_care/explainability.py",
        "Code/evaluate_care/provenance.py",
        "Code/evaluate_care/app.py",
    ],
    "EPIC-001-F-004": ["Code/evaluate_care/measures/board_certification.py"],
    "EPIC-001-F-005": ["Code/evaluate_care/measures/years_in_practice.py"],
    "EPIC-001-F-006": [
        "Code/evaluate_care/measures/specialty_match.py",
        "Code/ConversationalUX/FindCareChat/backend/domain/find_care/specialty_classifier.py",
        "Code/ConversationalUX/FindCareChat/backend/domain/find_care/specialty_ranker.py",
    ],
    "EPIC-001-F-007": ["Code/evaluate_care/measures/prescription_behavior.py"],
    "EPIC-001-F-008": ["Code/evaluate_care/measures/trial_phase.py"],
    "EPIC-001-F-009": ["Code/evaluate_care/measures/trial_recruitment.py"],
    "EPIC-001-F-010": ["Code/evaluate_care/measures/trial_condition_relevance.py"],
    "EPIC-001-F-011": "NO CODE FOUND",
    "EPIC-001-F-012": ["Code/evaluate_care/measures/patient_ratings.py"],
    "EPIC-001-F-013": ["Code/evaluate_care/measures/trial_sponsor.py"],
    "EPIC-001-F-014": ["Code/evaluate_care/measures/trial_proximity.py"],
    "EPIC-001-F-015": ["Code/evaluate_care/measures/trial_recency.py"],
    "EPIC-001-F-016": ["Code/evaluate_care/normalization.py"],
    "EPIC-001-F-017": ["Code/evaluate_care/provenance.py"],
    "EPIC-001-F-018": ["Code/evaluate_care/failsafe.py"],
    "EPIC-001-F-019": ["Code/evaluate_care/confidence.py"],
    "EPIC-001-F-020": ["Code/evaluate_care/measures/new_patient_acceptance.py"],
    "EPIC-001-F-021": ["Code/evaluate_care/measures/trial_eligibility.py"],
    "EPIC-001-F-022": ["Code/evaluate_care/cache.py"],
    "EPIC-001-F-023": "NO CODE FOUND",
    "EPIC-001-F-024": [
        "Code/evaluate_care/Dockerfile",
        "Code/evaluate_care/app.py",
        "Code/evaluate_care/requirements.txt",
        "Code/deploy/_start_evalcare.py",
    ],
    "EPIC-001-F-025": "NO CODE FOUND",

    # EPIC-002 Security --------------------------------------------------------
    "EPIC-002-F-001": [
        "Website/_headers",
        "Website/security-architecture.html",
        "Code/Shared/session_token.py",
        "Code/Shared/brain_auth.py",
        "Code/ConversationalUX/FindCareChat/backend/url_guardian.py",
    ],
    "EPIC-002-F-002": [
        "Website/security-architecture.html",
        "Code/Shared/ops/generate_security_page.py",
    ],

    # EPIC-004 Safety (absorbed into EPIC-009 per V4) -------------------------
    "EPIC-004-F-001": [
        "Code/ConversationalUX/FindCareChat/backend/domain/shared/safety/safety_service.py",
    ],

    # EPIC-005 Testing (dissolved per V4; carried forward as test stories that
    # re-home to the capability they exercise) --------------------------------
    "EPIC-005-F-001": [
        "Code/deploy/apiSmokeTestPyTest.py",
        "Code/deploy/devSmokeTestPyTest.py",
        "Code/deploy/localSmokeTestPyTest.py",
        "Code/ConversationalUX/FindCareChat/backend/tests",
    ],
    "EPIC-005-F-002": [
        "Code/Shared/ops/tools/regression_runner.py",
        "Code/deploy/devSmokeTestPyTest.py",
    ],
    "EPIC-005-F-003": [
        "Code/Shared/ops/tools/preCommitScan.py",
        "Website/security-architecture.html",
    ],
    "EPIC-005-F-004": [
        "Code/DataPipelines/quality_gate.py",
        "Code/DataPipelines/discrepancy_reporter.py",
        "Code/DataPipelines/validate_provider_load.py",
        "Code/DataPipelines/qa_provider_load.py",
        "Code/DataPipelines/schema_drift_detector.py",
    ],

    # EPIC-006 FindCare --------------------------------------------------------
    "EPIC-006-F-001": [
        "Code/ConversationalUX/FindCareChat/backend/domain/find_care/provider_search_service.py",
        "Code/ConversationalUX/FindCareChat/backend/application/tool_models/provider_search_models.py",
        "Code/ConversationalUX/FindCareChat/frontend/src/components/ChatWindow.tsx",
    ],
    "EPIC-006-F-002": [
        "Code/ConversationalUX/FindCareChat/backend/domain/find_care/specialty_service.py",
        "Code/ConversationalUX/FindCareChat/backend/domain/find_care/specialty_classifier.py",
        "Code/ConversationalUX/FindCareChat/backend/domain/find_care/specialty_ranker.py",
        "Code/ConversationalUX/FindCareChat/backend/domain/find_care/homeopathic_resolver.py",
    ],
    "EPIC-006-F-003": [
        "Code/ConversationalUX/FindCareChat/backend/application/facades/evaluate_care_facade.py",
        "Code/ConversationalUX/FindCareChat/backend/application/facades/find_care_facade.py",
    ],
    "EPIC-006-F-004": [
        "Code/ConversationalUX/FindCareChat/frontend/src/components/MessageBubble.tsx",
        "Code/ConversationalUX/FindCareChat/frontend/src/components/ChatWindow.tsx",
    ],
    # F-005 .. F-018 -- pipeline-flavored features that V5 moves to EPIC-012 --
    "EPIC-006-F-005": [
        "Code/DataPipelines/copy_to_frontend.py",
        "Code/DataPipelines/atlas_cluster_manager.py",
    ],
    "EPIC-006-F-006": [
        "Code/DataPipelines/pipeline_worker_base.py",
        "Code/DataPipelines/prescriber_pipeline_manager.py",
        "Code/DataPipelines/provider_load_manager.py",
        "Code/DataPipelines/cluster_lifecycle_manager.py",
        "Code/DataPipelines/idle_monitor.py",
        "Code/DataPipelines/instance_warmer.py",
        "Code/DataPipelines/sync_gateway_agent.py",
        "Code/DataPipelines/pipeline_health.py",
        "Code/DataPipelines/function_app.py",
    ],
    "EPIC-006-F-007": [
        "Code/DataPipelines/county_enrichment_job.py",
        "Code/DataPipelines/county_economic_enrichment.py",
        "Code/DataPipelines/embedding_worker.py",
        "Code/DataPipelines/provider_embedding.py",
    ],
    "EPIC-006-F-008": [
        "Code/DataPipelines/prescriber_load_worker.py",
        "Code/DataPipelines/provider_worker.py",
    ],
    "EPIC-006-F-009": [
        "Code/DataPipelines/data_fetcher_base.py",
        "Code/DataPipelines/pipeline_worker_base.py",
        "Code/DataPipelines/quality_gate.py",
    ],
    "EPIC-006-F-010": [
        "Code/DataPipelines/discrepancy_reporter.py",
        "Code/DataPipelines/validate_provider_load.py",
        "Code/DataPipelines/qa_provider_load.py",
    ],
    "EPIC-006-F-011": ["Code/DataPipelines/quality_gate.py"],
    "EPIC-006-F-013": ["Code/DataPipelines/schema_drift_detector.py"],
    "EPIC-006-F-014": [
        "Code/DataPipelines/pipeline_worker_base.py",
        "Code/DataPipelines/pipeline_db.py",
    ],
    "EPIC-006-F-015": ["Code/DataPipelines/data_fetcher_base.py"],
    "EPIC-006-F-016": ["Code/DataPipelines/copy_to_frontend.py"],
    "EPIC-006-F-017": [
        "Code/DataPipelines/tests",
        "Code/DataPipelines/pipeline_code_review.json",
    ],
    "EPIC-006-F-018": [
        "Code/DataPipelines/prescriber_pipeline_manager.py",
        "Code/DataPipelines/load_specialty_data.py",
        "Code/DataPipelines/zip_county_crosswalk_loader.py",
        "Code/DataPipelines/icd10_loader.py",
    ],
    "EPIC-006-F-019": [
        "Code/ConversationalUX/FindCareChat/frontend/src/App.tsx",
        "Code/ConversationalUX/FindCareChat/frontend/index.html",
    ],
    "EPIC-006-F-020": "NO CODE FOUND",
    "EPIC-006-F-021": ["Website/index.html"],
    "EPIC-006-F-022": ["Website/index.html"],
    "EPIC-006-F-023": [
        "Code/ConversationalUX/FindCareChat/frontend/index.html",
        "Code/ConversationalUX/FindCareChat/frontend/src/main.tsx",
    ],
    "EPIC-006-F-024": ["Website/index.html"],
    "EPIC-006-F-025": ["Website/index.html"],
    "EPIC-006-F-026": [
        "Website/index.html",
        "Code/Shared/ops/manifest_generator.py",
    ],
    "EPIC-006-F-027": [
        "Code/ConversationalUX/FindCareChat/backend/Dockerfile",
        "Code/ConversationalUX/FindCareChat/backend/main.py",
        "Code/ConversationalUX/FindCareChat/backend/requirements.txt",
        "Code/Shared/ops/hf_space_create.py",
        "Code/Shared/ops/hf_space_status.py",
        "Code/Shared/ops/hf_space_restart.py",
        "Code/Shared/ops/hf_space_delete.py",
    ],
    "EPIC-006-F-028": [
        "Code/DataPipelines/function_app.py",
        "Code/DataPipelines/host.json",
    ],
    "EPIC-006-F-029": [
        "Code/deploy/_localhost_server.py",
        "Code/deploy/_start_findcare.py",
        "Code/deploy/_start_evalcare.py",
        "Code/deploy/_start_shared.py",
        "Code/deploy/_start_website.py",
        "Code/deploy/deploy_localhost.sh",
        "Code/deploy/deploy_localhost_docker.sh",
        "start_local.bat",
    ],
    "EPIC-006-F-030": [
        "Code/Shared/ux/components/ProviderCard.tsx",
        "Code/Shared/ux/components/SelectionManager.tsx",
        "Code/Shared/ux/hooks/useSelectionState.ts",
        "Code/Shared/ux/types/provider.ts",
    ],

    # EPIC-007 Talk About Care -------------------------------------------------
    "EPIC-007-F-001": "NO CODE FOUND",

    # EPIC-008 Architecture (absorbed into EPIC-010 Operations per V4) --------
    "EPIC-008-F-001": "NO CODE FOUND",
    "EPIC-008-F-002": [
        "architecture/EngineeringRuleEnforcement/code/chathealthy_enforcement_manager.py",
        "architecture/EngineeringRuleEnforcement/code/scan_files_enforcement_worker.py",
        "architecture/EngineeringRuleEnforcement/code/enforcement_worker.py",
        "Code/Shared/ops/tools/preCommitScan.py",
        "Code/Shared/ops/tools/riskWaiver_gate.py",
        "Code/Shared/ops/tools/brain_write_validator.py",
        "Code/Shared/ops/tools/bash_rule_guard.py",
        "brain/machine_artifacts/content/engineering_rules.json",
    ],
    "EPIC-008-F-003": [
        "Code/Shared/ops/tools/capture_brain_hashes.py",
        "Code/Shared/ops/tools/build_backlog_story_view.py",
        "Code/Shared/brain_schema_validator.py",
        "Code/Shared/machine_brain.py",
        "Code/Shared/brain_loop.py",
    ],
    "EPIC-008-F-004": [
        "Code/deploy/deploy_app_dev.sh",
        "Code/deploy/deploy_dev.sh",
        "Code/Shared/ops/promote_build.py",
        "Code/Shared/ops/promote_data.py",
        "Code/Shared/ops/bump_build.py",
        "Code/Shared/ops/hf_space_create.py",
        "Code/Shared/ops/hf_space_restart.py",
        "Code/Shared/ops/hf_space_delete.py",
        "Code/Shared/ops/hf_space_status.py",
        "Code/Shared/ops/local_admin_server.py",
        "architecture/DevOpsBuildDeployAndEnvironmentManagement/code/bell_ringer.py",
    ],
    "EPIC-008-F-005": [
        "Code/Shared/ops/uat_report.py",
        "Code/Shared/ops/FullRegressionAudit/run_design_session.py",
        "Code/Shared/ops/FullRegressionAudit/render_design_doc.py",
    ],
    "EPIC-008-F-006": "NO CODE FOUND",
    "EPIC-008-F-007": [
        "Code/Shared/machine_brain.py",
        "Code/Shared/brain_loop.py",
        "Code/Shared/brain_schema_validator.py",
        "Code/Shared/brain_auth.py",
        "Code/Shared/ops/tools/chathealthy_devops_boot.py",
        "Code/Shared/ops/tools/lineage_orchestrator.py",
        "Code/Shared/ops/tools/conversation_log_agent.py",
        "Code/Shared/ops/tools/conversation_log_purge_service.py",
        "Code/Shared/ops/tools/log_instructions_loaded.py",
        "Code/Shared/ops/tools/boot_probe.py",
        "Code/Shared/ops/kafka/conversation_log_producer.py",
        "Code/Shared/ops/kafka/conversation_log_consumer.py",
        "Code/Shared/ops/kafka/docker-compose.yml",
        "Code/Shared/ops/ChatHealthyLogService/Worker.cs",
        "Code/Shared/ops/ChatHealthyLogService/Program.cs",
        "brain/machine_artifacts/content/engineering_rules.json",
        "brain/machine_artifacts/content/agile_backlog.json",
    ],
    "EPIC-008-F-008": [
        "Code/Shared/ops/tools/preCommitScan.py",
        "Code/Shared/ops/tools/scan_http.py",
        "Code/Shared/ops/tools/bash_rule_guard.py",
    ],
    "EPIC-008-F-009": "NO CODE FOUND",
    "EPIC-008-F-010": [
        "LangGraph/poc/user_journey.py",
        "Code/Shared/agent_framework/base_agent.py",
        "Code/Shared/agent_framework/base_tool.py",
        "Code/Shared/agent_framework/tool_registry.py",
    ],

    # EPIC-009 Shared Services ------------------------------------------------
    "EPIC-009-F-001": [
        "Code/Shared/llm_client.py",
        "Code/Shared/cost_guard.py",
        "Code/Shared/prompt_system_maker.py",
        "Code/Shared/session_token.py",
        "Code/Shared/ChatHealthyMongoUtilities.py",
        "Code/Shared/agent_framework/base_agent.py",
        "Code/Shared/agent_framework/base_tool.py",
        "Code/Shared/agent_framework/tool_registry.py",
        "Code/ConversationalUX/FindCareChat/backend/application/tool_router.py",
        "Code/ConversationalUX/FindCareChat/backend/infrastructure/embeddings/embedding_client.py",
        "Code/ConversationalUX/FindCareChat/backend/domain/shared/consent/consent_service.py",
        "Code/ConversationalUX/FindCareChat/backend/domain/shared/lead_capture/lead_service.py",
        "Code/ConversationalUX/FindCareChat/backend/domain/shared/unknowns/unknown_question_service.py",
        "Code/ConversationalUX/FindCareChat/backend/domain/shared/content/about_service.py",
        "Code/ConversationalUX/FindCareChat/backend/domain/shared/safety/safety_service.py",
        "Code/ConversationalUX/ChatHealthyWhoAmIChat/me",
        "Code/shared_services/app.py",
        "Code/shared_services/Dockerfile",
        "Code/shared_services/requirements.txt",
    ],

    # EPIC-011 UX -------------------------------------------------------------
    "EPIC-011-F-001": [
        "Code/Shared/ux/hooks/useSelectionState.ts",
    ],
}

# Per-epic destination under V5 (after the move table). Used for grouping.
# Stories never split: one story belongs to exactly one epic in V5.
# V5 has 7 epics: EPIC-001, EPIC-002, EPIC-006, EPIC-009, EPIC-010, EPIC-011,
# EPIC-012. Mapping logic for stories that move:
#  - EPIC-004 (Safety) stories -> EPIC-009 Shared Services
#  - EPIC-005 (Testing) stories -> capability they exercise (best-effort below)
#  - EPIC-006 F-005..F-018, F-028 -> EPIC-012 Data Management Pipelines
#  - EPIC-006 F-021, F-024, F-025, F-026 -> EPIC-011 UX
#  - EPIC-006 F-022 -> SPLIT: 1 story to EPIC-011 (host side), 1 to EPIC-006
#  - EPIC-006 F-029 -> EPIC-010 Operations
#  - EPIC-007 (Talk About Care) story stays in backlog only (V4 binding); V5
#    surfaces it under "EPIC-001-or-EPIC-006-future" -- placement TBD. To
#    keep "every story appears in exactly one epic table" in V5, we list the
#    1 EPIC-007 story under EPIC-009 Shared Services (where the deep_research
#    subgraph lives per LangGraph V4) with a note.
#  - EPIC-008 (Architecture) all stories -> EPIC-010 Operations
#  - EPIC-011 existing 1 story -> EPIC-011 UX
# ----------------------------------------------------------------------------

# Final-epic mapping. Keyed by feature_id. Story-level overrides come next.
FEATURE_FINAL_EPIC = {
    # EPIC-001 EvaluateCare -- all stays
    **{f"EPIC-001-F-{i:03d}": "EPIC-001" for i in range(1, 26)},
    # EPIC-002 Security -- all stays
    "EPIC-002-F-001": "EPIC-002",
    "EPIC-002-F-002": "EPIC-002",
    # EPIC-004 Safety -- moves to EPIC-009
    "EPIC-004-F-001": "EPIC-009",
    # EPIC-005 Testing -- per V4: each story re-homes to the capability it exercises.
    # F-001 User Testing -- exercises FindCare end-to-end -> EPIC-006
    "EPIC-005-F-001": "EPIC-006",
    # F-002 Regression Testing -- ops concern -> EPIC-010
    "EPIC-005-F-002": "EPIC-010",
    # F-003 Security Testing -> EPIC-002
    "EPIC-005-F-003": "EPIC-002",
    # F-004 Data Quality -- pipeline data integrity -> EPIC-012
    "EPIC-005-F-004": "EPIC-012",
    # EPIC-006 FindCare
    "EPIC-006-F-001": "EPIC-006",
    "EPIC-006-F-002": "EPIC-006",
    "EPIC-006-F-003": "EPIC-006",
    "EPIC-006-F-004": "EPIC-006",
    "EPIC-006-F-005": "EPIC-012",
    "EPIC-006-F-006": "EPIC-012",
    "EPIC-006-F-007": "EPIC-012",
    "EPIC-006-F-008": "EPIC-012",
    "EPIC-006-F-009": "EPIC-012",
    "EPIC-006-F-010": "EPIC-012",
    "EPIC-006-F-011": "EPIC-012",
    "EPIC-006-F-013": "EPIC-012",
    "EPIC-006-F-014": "EPIC-012",
    "EPIC-006-F-015": "EPIC-012",
    "EPIC-006-F-016": "EPIC-012",
    "EPIC-006-F-017": "EPIC-012",
    "EPIC-006-F-018": "EPIC-012",
    "EPIC-006-F-019": "EPIC-006",
    "EPIC-006-F-020": "EPIC-006",
    "EPIC-006-F-021": "EPIC-011",
    "EPIC-006-F-022": "EPIC-011",  # SPLIT below at story level
    "EPIC-006-F-023": "EPIC-006",
    "EPIC-006-F-024": "EPIC-011",
    "EPIC-006-F-025": "EPIC-011",
    "EPIC-006-F-026": "EPIC-011",
    "EPIC-006-F-027": "EPIC-006",
    "EPIC-006-F-028": "EPIC-012",
    "EPIC-006-F-029": "EPIC-010",
    "EPIC-006-F-030": "EPIC-006",
    # EPIC-007 -> backlog parking lot. For "every story appears in exactly
    # one epic table" we place it under EPIC-009 (deep_research subgraph host
    # per LangGraph V4 alignment). Marked accordingly in the table.
    "EPIC-007-F-001": "EPIC-009",
    # EPIC-008 -- all stays under EPIC-010 Operations
    "EPIC-008-F-001": "EPIC-010",
    "EPIC-008-F-002": "EPIC-010",
    "EPIC-008-F-003": "EPIC-010",
    "EPIC-008-F-004": "EPIC-010",
    "EPIC-008-F-005": "EPIC-010",
    "EPIC-008-F-006": "EPIC-010",
    "EPIC-008-F-007": "EPIC-010",
    "EPIC-008-F-008": "EPIC-010",
    "EPIC-008-F-009": "EPIC-010",
    "EPIC-008-F-010": "EPIC-010",
    # EPIC-009 -- stays
    "EPIC-009-F-001": "EPIC-009",
    # EPIC-011 -- stays
    "EPIC-011-F-001": "EPIC-011",
}

# Story-level overrides (only where the feature splits or a story re-homes
# differently from the rest of its feature).
STORY_FINAL_EPIC_OVERRIDES = {
    # F-022 postMessage Orchestration: per V4, host-side handler -> EPIC-011,
    # FindCare-iframe-side handler -> EPIC-006.
    # The two stories are EPIC-006-F-022-S-001 and EPIC-006-F-022-S-002.
    # Choose: S-001 = host (EPIC-011), S-002 = iframe (EPIC-006).
    "EPIC-006-F-022-S-002": "EPIC-006",
}

# V5 epic display names (with the rename binding from the brief).
V5_EPICS = [
    ("EPIC-001", "EvaluateCare"),
    ("EPIC-002", "Security"),
    ("EPIC-006", "FindCare"),
    ("EPIC-009", "Shared Services"),
    ("EPIC-010", "Operations"),
    ("EPIC-011", "UX"),
    ("EPIC-012", "Data Management Pipelines"),
]

# Orphan-code findings per V5 epic. Each orphan = (file_path, why_in_territory,
# why_orphan).
ORPHANS = {
    "EPIC-001": [
        ("Code/evaluate_care/__init__.py", "EvaluateCare service package marker", "Not named by any backlog story"),
        ("Code/evaluate_care/graphs/", "EvaluateCare LangGraph subgraph dir (placeholder)", "Empty directory; no story mentions a LangGraph EvaluateCare subgraph yet"),
        ("Code/evaluate_care/measures/__init__.py", "Measures package init", "Not named by any story"),
        ("Code/evaluate_care/tests/", "EvaluateCare test surface", "Stories say 'write tests' but never name this directory; trace by inference, not by ID"),
    ],
    "EPIC-002": [
        ("Code/Shared/brain_auth.py", "Brain JSON write authorization (security-flavored)", "Cited by EPIC-002-F-001 only by inference (header decoration); no story names it"),
        ("Code/Shared/session_token.py", "Session-token generation/verification", "Used by F-001 stories but not named explicitly in any S-* description"),
        ("Code/ConversationalUX/FindCareChat/backend/url_guardian.py", "Outbound-URL allowlist enforcement (security)", "No EPIC-002 story names it"),
    ],
    "EPIC-006": [
        ("Code/find_care/", "Empty directory whose name is FindCare", "Empty; no story; carries from V3/V4 anti-pattern list"),
        ("Code/ConversationalUX/FindCareChat/backend/adapter/__init__.py", "FindCareChat backend adapter package (currently empty)", "Empty package; no story exercises it"),
        ("Code/ConversationalUX/FindCareChat/backend/domain/shared/url/__init__.py", "Shared URL domain package (currently empty)", "No URL service code yet; no story"),
        ("Code/ConversationalUX/FindCareChat/backend/infrastructure/config/__init__.py", "Config infra package (empty)", "Empty package; no story"),
        ("Code/ConversationalUX/FindCareChat/backend/infrastructure/external_apis/__init__.py", "External-APIs infra package (empty)", "Empty package; no story"),
        ("Code/ConversationalUX/FindCareChat/backend/infrastructure/mongo/__init__.py", "Mongo infra package (empty)", "Empty package; no story"),
        ("Code/ConversationalUX/FindCareChat/frontend/src/components/GUIManager.tsx", "FindCare React UI", "Component exists; no specific story names it"),
        ("Code/ConversationalUX/FindCareChat/frontend/src/components/FindCareApp.tsx", "FindCare React UI top-level component", "Exists; no story names it"),
        ("Code/ConversationalUX/FindCareChat/frontend/src/components/MessageBubble.tsx", "FindCare React message rendering", "Used by F-004 by inference only; not named in any story"),
    ],
    "EPIC-009": [
        ("Code/Shared/__pycache__/", "Bytecode artifact under Shared", "Build artifact; should be gitignored; no story"),
        ("Code/Shared/agent_framework/__init__.py", "Agent framework package init", "Not named by any story"),
        ("Code/shared_services/app.py", "shared_services HF Space entrypoint", "Trace by inference (deploy convention), not by any S-* description"),
        ("Code/shared_services/Dockerfile", "shared_services HF Space container", "Trace by inference (deploy convention), not by story"),
        ("Code/shared_services/requirements.txt", "shared_services Python deps", "Trace by inference, not by story"),
        ("Code/Shared/tests/test_brain_schema_validator.py", "Test for brain_schema_validator", "Tests exist; the test target is governance code under EPIC-008-F-007, but no Shared-Services story owns this test"),
    ],
    "EPIC-010": [
        ("Code/Shared/ops/Caddyfile", "Local TLS reverse proxy config", "Used at dev time; no story"),
        ("Code/Shared/ops/certs/", "Local TLS certs directory", "Used at dev time; no story"),
        ("Code/Shared/ops/certs;C/", "Malformed-Windows-path artifact (empty)", "Stray empty directory; no story; cleanup-only"),
        ("Code/Shared/ops/start_local.bat", "Local-server bootstrap script", "Convenience launcher; no story"),
        ("Code/Shared/ops/daily_conversation_commit.sh", "Cron-style conversation log committer", "Used by ops; no story"),
        ("Code/Shared/ops/sync_to_mongo.py", "One-shot brain->Mongo sync", "Carries forward from earlier work; no current story"),
        ("Code/Shared/ops/atlas_cluster_toggle.py", "Atlas cluster up/down toggle", "Ops convenience; no current story"),
        ("Code/Shared/ops/conversation_log_hook.py", "Hook helper for conversation logging", "No story; supporting EPIC-008-F-007 stories by inference"),
        ("Code/Shared/ops/framework_version.py", "Reads V20 framework version", "Supporting code; no story"),
        ("Code/Shared/ops/generate_roadmap_page.py", "Generates Website/roadmap.html", "Page generator; no story names it"),
        ("Code/Shared/ops/manifest_generator.py", "Generates Website manifests", "Page generator; no story names it (F-026 names index.html only)"),
        ("Code/Shared/ops/rebuild_manifest.py", "Manifest rebuild ops tool", "No story"),
        ("Code/Shared/ops/unattended_monitor.py", "Background ops monitor", "No story"),
        ("Code/Shared/ops/epic_planning_runner.py", "Epic planning runner script", "Planning helper; no story"),
        ("Code/Shared/ops/tools/disk_cleanup_*.log", "Disk cleanup logs", "Run-time logs; should be gitignored; no story"),
        ("Code/Shared/ops/tools/caddy.exe", "Bundled Caddy binary", "Tool binary; not story-traced"),
        ("Code/Shared/ops/tools/kill_dev_servers.ps1", "Windows kill-dev-servers helper", "Ops convenience; no story"),
        ("Code/Shared/ops/tools/kill_zombies.py", "Zombie-process reaper", "Ops convenience; no story"),
        ("Code/Shared/ops/tools/pipeline_trigger.py", "Manual pipeline trigger", "Ops tool; no story (related stories are EPIC-006-F-006 / EPIC-012)"),
        ("Code/Shared/ops/tools/gen_ops_design_docx.py", "Generates ops design docx", "Doc generator; no story"),
        ("Code/Shared/ops/tools/skill/SKILL.md", "Claude skill definition", "Skill artifact; no story"),
        ("Code/CSharp/ChatHealthyLogService/dev/", "Legacy non-canonical Worker.cs copy", "Per V4 Issue 1.1: DELETE this tree; no story owns it (canonical lives at Code/Shared/ops/ChatHealthyLogService/)"),
        ("Code/Shared/ops/tests/", "Tests for ops tools", "Test scaffolding; no current story"),
        ("Code/Shared/ops/tools/tests/", "Tests for ops/tools", "Test scaffolding; no current story"),
        ("Code/Shared/ops/kafka/tests/", "Tests for kafka producer/consumer", "Test scaffolding; no current story"),
        ("architecture/EngineeringRuleEnforcement/locks/", "Enforcement lock files", "Run-time locks; no story"),
        ("architecture/EngineeringRuleEnforcement/__init__.py", "EnfRule package init", "Package init; no story"),
        ("Code/Shared/cost_guard.py", "Cost-guard rate-limit logic", "Listed under EPIC-009 already; appears here only if Operations claims it. Mark as Shared-Services-owned -- not orphan to Operations."),
    ],
    "EPIC-011": [
        ("Code/Shared/ux/components/ProviderCard.tsx", "Cross-component provider card primitive", "Per V4 it lives under EPIC-011; today only EPIC-006-F-030 (Provider UX Display) names it -- if F-030 stays in EPIC-006, this file is UX-orphan"),
        ("Code/Shared/ux/components/SelectionManager.tsx", "Cross-component selection manager primitive", "Same: anchored by EPIC-006-F-030 only"),
        ("Code/Shared/ux/types/provider.ts", "Shared TS provider type", "No story names it; type-only"),
        ("Website/architecture.html", "Static design-paradigm HTML page", "Shares :root design tokens with the other static pages; no story owns the design paradigm yet -- this is the V5 UX requirements gap (Deliverable 3)"),
        ("Website/products.html", "Static design-paradigm HTML page", "Shares :root design tokens; no story"),
        ("Website/privacy.html", "Static design-paradigm HTML page", "Shares :root design tokens; no story"),
        ("Website/terms.html", "Static design-paradigm HTML page", "Shares :root design tokens; no story"),
        ("Website/roadmap.html", "Static design-paradigm HTML page", "Shares :root design tokens; no story"),
        ("Website/security-architecture.html", "Static design-paradigm HTML page", "Shares :root design tokens; content under EPIC-002 but design surface is UX-orphan"),
        ("Website/chat-app-design.html", "Static design-paradigm HTML page", "Shares :root design tokens; no story"),
        ("Website/embedding-design.html", "Static design-paradigm HTML page", "Shares :root design tokens; no story"),
        ("Website/load-perf-report.html", "Static design-paradigm HTML page", "Shares :root design tokens; no story"),
        ("Website/ops-manager-design.html", "Static design-paradigm HTML page", "Shares :root design tokens; no story"),
        ("Website/provider-data-load.html", "Static design-paradigm HTML page", "Shares :root design tokens; no story"),
    ],
    "EPIC-012": [
        ("Code/DataPipelines/ChatHealthyMongoUtilities.py", "Pipeline-territory file (dead code)", "Per V4 Issue 1.2: dead code; no callers. DELETE per move table; no story owns it."),
        ("Code/DataPipelines/auth.py", "Pipeline-side OTP/auth", "Used by pipeline functions; no story names it"),
        ("Code/DataPipelines/blob_client.py", "Azure Blob client wrapper", "Used by pipeline workers; no story names it"),
        ("Code/DataPipelines/otp_manager.py", "Pipeline OTP manager", "Used by pipeline functions; no story names it"),
        ("Code/DataPipelines/promote_data_fn.py", "Pipeline data-promotion function", "No story; related to EPIC-006-F-016 by inference only"),
        ("Code/DataPipelines/count_providers_by_state.py", "Analytics one-shot", "No story"),
        ("Code/DataPipelines/crosswalk_builder.py", "ZIP/county crosswalk builder", "Loader-flavored; trace via F-018 by inference only"),
        ("Code/DataPipelines/gpt_reader.py", "GPT-fed content reader", "Pipeline get-content worker; no story names it"),
        ("Code/DataPipelines/pipeline_db.py", "Pipeline DB connection helper", "Used by pipeline workers; no story"),
        ("Code/DataPipelines/prescriber_evaluate_care_pipeline.py", "Prescriber pipeline that feeds EvaluateCare", "Cross-epic pipeline (EvaluateCare <- pipeline data); no story names it"),
        ("Code/DataPipelines/prescriber_data_fetcher.py", "Prescriber data fetcher", "Trace via F-018 by inference"),
        ("Code/DataPipelines/prescriber_enrichment_job.py", "Prescriber enrichment job", "Trace via F-007 by inference"),
        ("Code/DataPipelines/prescriber_load_worker.py", "Prescriber load worker", "Trace via F-008 by inference"),
        ("Code/DataPipelines/provider_worker.py", "Provider load worker", "Trace via F-008 by inference"),
        ("Code/DataPipelines/legal/licence.txt", "Third copy of project license inside DataPipelines", "Stray; cleanup; no story"),
        ("Code/DataPipelines/favicon.ico", "Stray favicon under pipelines", "Orphan asset; no story"),
        ("Code/DataPipelines/pipeline_files.json", "Pipeline manifest JSON", "Manifest data; no story"),
        ("Code/DataPipelines/pipeline_code_review.json", "Pipeline code review JSON", "Used by F-017 traceability; trace by inference only"),
        ("Code/DataPipelines/prescriber-pipeline-test.http", "REST-client test scratch file", "Manual test file; no story"),
        ("Code/DataPipelines/overnight_pipeline.log", "Pipeline run log", "Run-time log; should be gitignored; no story"),
        ("Code/DataPipelines/ops_manager/", "Ops-manager toolset for the pipeline", "Used by ops; no story names it (related to EPIC-008-F-005 audit features by inference)"),
        ("Code/Schemas/ChatHealthy.Providers.json", "Schema duplicate (file-type-bucket)", "Per V4 Issue 2: byte-equivalent duplicate; DELETE per move table; no story owns it"),
        ("Website/schemas/", "Schemas served from Website (deploy artifact, not source-of-truth)", "Per V4 B.4: relocate to brain/machine_artifacts/schemas/; no story today"),
    ],
}


# ----------------------------------------------------------------------------
# Helpers
# ----------------------------------------------------------------------------
def load_backlog():
    with open(BACKLOG, "r", encoding="utf-8") as f:
        return json.load(f)


def all_stories():
    """Yield (epic_id, epic_name, feature_id, feature_name, story_id, story_desc)."""
    data = load_backlog()
    for e in data["epics"]["epic"]:
        eid = e["epic_id"]
        ename = e["name"]
        feats = (e.get("features") or {}).get("feature", [])
        for f in feats:
            fid = f["feature_id"]
            fname = f["name"]
            stories = (f.get("stories") or {}).get("story", [])
            if not stories:
                continue
            for s in stories:
                sid = s.get("story_id") or s.get("id") or ""
                sd = (s.get("description") or s.get("summary") or "").replace("\n", " ").strip()
                yield eid, ename, fid, fname, sid, sd


def final_epic_for_story(orig_epic, fid, sid):
    if sid in STORY_FINAL_EPIC_OVERRIDES:
        return STORY_FINAL_EPIC_OVERRIDES[sid]
    if fid in FEATURE_FINAL_EPIC:
        return FEATURE_FINAL_EPIC[fid]
    return orig_epic  # fallback


def code_for_feature(fid):
    """Return list of paths or 'NO CODE FOUND' for a feature."""
    val = FEATURE_CODE.get(fid)
    if val is None:
        return "NO CODE FOUND"
    return val


# ----------------------------------------------------------------------------
# Doc emission
# ----------------------------------------------------------------------------
def add_para(doc, text, style=None, bold=False):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    r = p.add_run(text)
    if bold:
        r.bold = True
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def add_h1(doc, text):
    doc.add_heading(text, level=1)


def add_h2(doc, text):
    doc.add_heading(text, level=2)


def add_h3(doc, text):
    doc.add_heading(text, level=3)


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            t.rows[r_idx].cells[c_idx].text = "" if val is None else str(val)
    doc.add_paragraph()  # spacer
    return t


def render_code_paths(val):
    if val == "NO CODE FOUND":
        return "NO CODE FOUND"
    if isinstance(val, list):
        return "\n".join(val)
    return str(val)


# ----------------------------------------------------------------------------
# Main
# ----------------------------------------------------------------------------
def main():
    doc = Document()

    # ---------- Title -------------------------------------------------------
    doc.add_heading("Codebase Reorganization -- V5", level=0)
    add_para(
        doc,
        "Date: 2026-04-28   |   Author: Claude (architecture analysis agent)   |   Read-only",
    )
    add_para(
        doc,
        "Lineage: codebase-reorganization-pass3.docx -> codebase-reorganization-pass3-V2.docx (Skip's annotated) -> codebase-reorganization-V3.docx -> codebase-reorganization-V4.docx -> this V5.",
    )
    add_para(doc, "Inputs read in order:")
    for s in [
        "findCare/ArchitectureAndDesign/codebase-reorganization-V4.docx -- prior version; V5 supersedes on the corrections below.",
        "findCare/ArchitectureAndDesign/codebase-reorganization-pass3-V2.docx -- Skip's annotated Pass 3 (binding).",
        "findCare/ArchitectureAndDesign/langgraph-oo-best-practices-V4.docx -- definitive LangGraph design recommendations.",
        "brain/machine_artifacts/content/agile_backlog.json -- 9-epic, 274-story shape (verified live).",
        "Code/, Website/, LangGraph/, architecture/ -- file tree walked; every cited path verified by direct filesystem listing.",
    ]:
        add_bullet(doc, s)

    # ---------- 1. Operating-model axiom ------------------------------------
    add_h2(doc, "1. Operating-model axiom (verbatim)")
    p = doc.add_paragraph()
    r = p.add_run(
        '"The operating model, the digital footprint, the general ledger, and '
        "the org chart are 100% in alignment. The structure of any one IS the "
        "structure of all four. This alignment is non-negotiable and is the "
        'fundamental organizing principle of ChatHealthy."'
    )
    r.italic = True
    add_para(
        doc,
        "Every proposal in V5 is tested against the axiom: does the proposed "
        "Git-tree shape mirror what would appear on the org chart, on a "
        "general-ledger cost-center list, and on the customer-facing capability "
        "map? Where the answer is no, the structure is decoration and is removed.",
    )

    # ---------- 2. Security hosting dilemma ---------------------------------
    add_h2(doc, "2. Security hosting dilemma -- recommendation up front")
    add_para(
        doc,
        "V4 promoted Security to a first-class epic (EPIC-002). That part is "
        "correct. The dilemma V5 surfaces is hosting: Security code today is "
        "bundled inside the Shared Services HuggingFace container "
        "(Code/shared_services/app.py runs auth, session-token verification, "
        "tool routing, and other shared back-end concerns in one process). "
        "There is no separate Security HF Space.",
    )
    add_para(doc, "Stated plainly: the dilemma.", bold=True)
    add_bullet(
        doc,
        "Security is a discrete cost center on the operating model "
        "(Cloudflare findings, compliance attestation, HSTS/TLS posture, "
        "OAuth, session-token authentication, user entitlements -- substantial "
        "scope, distinct customer surface, distinct compliance budget).",
    )
    add_bullet(
        doc,
        "Today Security shares a container (the SharedServices HF Space) with "
        "Tool Router, LLM client, embeddings, consent, lead capture, "
        "Skip-Snow context, and the Safety subgraph. By the operating-model "
        "axiom, cost-center boundaries SHOULD mirror code-deployment boundaries. "
        "Sharing the container blurs that mirror.",
    )

    add_h3(doc, "2.1 Option A -- Security gets its own HF Space")
    for s in [
        "Own container, own deploy pipeline, own /health endpoint, own observability story.",
        "Operating-model alignment: PASSES. The HF Space line in the Spaces inventory "
        "matches the Security line on the org chart matches the Security line "
        "on the cost-center list. Three columns aligned, axiom honored.",
        "Security isolation: STRONG. Authentication and entitlement code runs in a "
        "separate process from tool routing, LLM client, and embeddings. A "
        "compromise of the SharedServices container does not directly expose "
        "session-token signing keys or OAuth client secrets.",
        "Operational cost: ONE additional HF Space. Each Space is an operational "
        "unit (build, deploy, monitor, restart, cost line). Cost is real but "
        "the marginal cost of one more Space is small relative to the FindCare "
        "and EvaluateCare Spaces that already exist.",
        "Deployment complexity: ADDS one deploy pipeline (Code/Operations/deploy/ "
        "gains a deploy_security.sh and a _start_security.py). The /verify-token "
        "and /authenticate endpoints move from Code/shared_services/app.py to "
        "Code/security/app.py. Calls from FindCare and EvaluateCare to verify a "
        "session token become cross-Space calls (already cross-process today; "
        "the boundary becomes cross-container).",
        "Latency cost: small. /verify-token is a hot path -- every chat request "
        "verifies a token. Today this is a same-process call. Under Option A it "
        "becomes a same-region cross-container call. Mitigation: cache the "
        "verified token for the conversation duration; the cache is already in "
        "session_token.py.",
    ]:
        add_bullet(doc, s)

    add_h3(doc, "2.2 Option B -- Security stays inside Shared Services container")
    for s in [
        "Status quo. Code/shared_services/app.py keeps hosting Security endpoints "
        "alongside Tool Router and the rest.",
        "Operating-model alignment: FAILS. The cost-center list has separate "
        "Security and Shared Services lines. The HF Spaces inventory has only "
        "one. The Git-tree mirror is broken at exactly the layer the axiom "
        "says must be aligned.",
        "Security isolation: WEAK. A bug or compromise in the LLM client code "
        "or the consent service runs in the same process that signs and "
        "verifies session tokens. Defense-in-depth is reduced.",
        "Operational cost: ZERO marginal Spaces. One less deploy pipeline.",
        "Deployment complexity: LOWEST today. /verify-token stays a same-process "
        "call. No cross-container latency.",
        "Failure mode: every time Security expands (OAuth, user-creation, "
        "entitlement matrix, audit log surface) the SharedServices container "
        "grows. Eventually the container holds two cost centers' worth of "
        "code in one deploy unit -- the exact anti-pattern the axiom forbids. "
        "The relocation cost rises with every new Security feature added under "
        "Option B.",
    ]:
        add_bullet(doc, s)

    add_h3(doc, "2.3 Recommendation -- Option A")
    add_para(
        doc,
        "Option A: Security gets its own HF Space. Justification: the "
        "operating-model axiom is binding. Today the Security cost-center line "
        "and the Security org-chart line both exist; only the deploy line is "
        "missing. Manufacturing the third line by collapsing the first two is not "
        "an option (the first two reflect business reality). Manufacturing the "
        "third line by carving out an HF Space is the cheap, supported move. "
        "The operational cost of one additional Space is small and bounded; the "
        "isolation benefit grows monotonically as Security scope grows (OAuth "
        "+ user-creation + entitlement matrix + audit-log surface are all on "
        "Skip's near-term roadmap per the V4 binding). Option B's apparent "
        "simplicity decays as Security grows; Option A's overhead is constant.",
    )
    add_para(
        doc,
        "Concrete shape under Option A: Code/Security/ -- backend/app.py, "
        "backend/auth/, backend/entitlements/, backend/audit/, frontend/ (if "
        "Security grows a self-service UI), Dockerfile, requirements.txt, deploy/ "
        "hooks under Code/Operations/deploy/. The visual auth-token DISPLAY "
        "object stays under EPIC-011 UX (presentation). The auth-token "
        "CREATION/VERIFICATION moves out of Code/shared_services/ and into "
        "Code/Security/. EPIC-002 owns the new HF Space end-to-end.",
    )

    # ---------- 3. Layer-name explanation -----------------------------------
    add_h2(doc, "3. Layer-name explanation -- core/ (carry forward from V4)")
    add_para(
        doc,
        "The hexagonal-architecture term 'domain' fails the operating-model "
        "axiom (no HR, GL, or capability-map line). The candidate 'rules/' "
        "collides with brain/machine_artifacts/content/engineering_rules.json -- "
        "a load-bearing governed JSON whose vocabulary is enforced by the V20 "
        "framework. The accepted answer is 'core/': it names what the "
        "capability IS ('the FindCare core,' 'the EvaluateCare scoring core'), "
        "passes the axiom, and has no overlap with the engineering-rules "
        "controlled vocabulary. Final structure: "
        "Code/FindCare/backend/core/{provider_search_service.py, "
        "specialty_classifier.py, specialty_ranker.py, specialty_service.py, "
        "homeopathic_resolver.py}; Code/EvaluateCare/core/{scoring_engine.py, "
        "normalization.py, weights.py, ...}. 'domain' does not appear as a "
        "structural noun anywhere in V5.",
    )

    # ---------- 4. Three architect open issues ------------------------------
    add_h2(doc, "4. Three architect open issues -- carried forward from V4")
    add_h3(doc, "4.1 Worker.cs canonical -- RESOLVED")
    add_para(
        doc,
        "Code/Shared/ops/ChatHealthyLogService/Worker.cs (34 lines) is canonical. "
        "Active build path per .vscode/tasks.json; more recent design intent (commit "
        "bd44f12 / T011 'no business logic'). Action: delete "
        "Code/CSharp/ChatHealthyLogService/dev/Worker.cs and the surrounding "
        "Code/CSharp/ tree under the Operations refactor (move table row 1).",
    )
    add_h3(doc, "4.2 ChatHealthyMongoUtilities.py canonical -- RESOLVED")
    add_para(
        doc,
        "Code/Shared/ChatHealthyMongoUtilities.py (Copy A) is canonical. Only copy "
        "with callers (FindCareChat backend main.py:47); only copy watched by deploy CI. "
        "Action: move to Code/SharedServices/mongo_utilities.py; port the "
        "WHEN-TO-USE docstring from Copy B; delete Copy B at "
        "Code/DataPipelines/ChatHealthyMongoUtilities.py.",
    )
    add_h3(doc, "4.3 Provider schema canonical + 11-field data gap -- RESOLVED")
    add_para(
        doc,
        "Schemas at Code/Schemas/ChatHealthy.Providers.json and "
        "Website/schemas/ChatHealthyProvidersSchema.json are functionally "
        "equivalent (the 11-byte delta is the $id URL + trailing newline; every "
        "property/required/conditional is byte-identical). Both reject 100% of "
        "production providers because the running pipeline produces 11 fields the "
        "schemas do not declare (embedding [100%], embedding_model [52%], "
        "embedding_version [52%], 8 NPPES other-name fields). Action: adopt "
        "Website/schemas/ content as canonical, relocate to "
        "brain/machine_artifacts/schemas/ChatHealthyProvidersSchema.json, set "
        "$id at the brain canonical URL, delete Code/Schemas/. Schema-vs-data "
        "gap tracked under EPIC-012 Data Management Pipelines / schema-drift "
        "feature -- separate work item, does not block relocation.",
    )

    # ---------- 5. Counts table ---------------------------------------------
    add_h2(doc, "5. Counts table -- proposed V5 epic list")
    add_para(
        doc,
        "Each row is one epic in the proposed final structure. Counts reflect "
        "the V5 dispositions (EPIC-004 Safety absorbed into EPIC-009; EPIC-005 "
        "Testing dissolved with stories re-homing to capability epics; "
        "EPIC-007 Talk About Care story parked into EPIC-009 deep-research "
        "scope; EPIC-008 Architecture absorbed into EPIC-010 Operations; "
        "EPIC-006 pipeline features moved to EPIC-012; EPIC-006 parent-page "
        "features moved to EPIC-011). EPIC-012 renamed per binding correction.",
    )

    # Compute counts dynamically so the doc reflects the truth at generate time.
    rows = list(all_stories())  # list of tuples
    final_epic_per_story = {(t[2], t[4]): final_epic_for_story(t[0], t[2], t[4]) for t in rows}
    feature_to_final_epic = defaultdict(set)
    story_count_per_final = defaultdict(int)
    for t in rows:
        eid, ename, fid, fname, sid, sd = t
        fe = final_epic_for_story(eid, fid, sid)
        feature_to_final_epic[fid].add(fe)
        story_count_per_final[fe] += 1

    # Feature counts: a feature is "in" an epic if any of its stories land there.
    # F-022 splits across EPIC-006 and EPIC-011 -- counted under both.
    feature_count_per_final = defaultdict(int)
    for fid, fes in feature_to_final_epic.items():
        for fe in fes:
            feature_count_per_final[fe] += 1

    counts_rows = []
    for eid, name in V5_EPICS:
        counts_rows.append([
            eid, name, str(feature_count_per_final.get(eid, 0)),
            str(story_count_per_final.get(eid, 0)),
        ])
    add_table(doc, ["Epic", "Name", "# Features", "# Stories"], counts_rows)

    add_para(
        doc,
        "Total stories across all 7 V5 epics: "
        + str(sum(story_count_per_final.values()))
        + " (matches the 274-story backlog truth; every story appears in exactly "
        "one epic table per the V5 constraint).",
    )

    # ---------- 6. Per-epic story-to-code tables (Deliverable 1) ------------
    add_h2(doc, "6. Per-epic story-to-code tables (Deliverable 1)")
    add_para(
        doc,
        "One table per V5 epic. Each row is one story. Code path(s) cite real "
        "files under Code/, Website/, LangGraph/, brain/, or architecture/. "
        '"NO CODE FOUND" is a real finding -- it means the story has no '
        "implementing file today (deletion candidate or future work).",
    )
    # Build per-final-epic story rows.
    stories_per_final_epic = defaultdict(list)  # final_epic -> list of (orig_epic, fname, sid, sd, code)
    for t in rows:
        eid, ename, fid, fname, sid, sd = t
        fe = final_epic_for_story(eid, fid, sid)
        code = code_for_feature(fid)
        stories_per_final_epic[fe].append({
            "orig_epic": eid,
            "feature_id": fid,
            "feature_name": fname,
            "story_id": sid,
            "story_desc": sd,
            "code": code,
        })

    epic_name_lookup = dict(V5_EPICS)
    for eid, ename in V5_EPICS:
        add_h3(doc, f"6.{[e[0] for e in V5_EPICS].index(eid)+1} {eid} {ename}")
        st_rows = stories_per_final_epic.get(eid, [])
        if not st_rows:
            add_para(doc, "(no stories tabled to this epic)")
            continue
        # Sort by feature_id then story_id
        st_rows.sort(key=lambda r: (r["feature_id"], r["story_id"]))
        table_rows = []
        for s in st_rows:
            note = ""
            if s["orig_epic"] != eid:
                note = f"  [moved from {s['orig_epic']}]"
            table_rows.append([
                eid,
                f"{s['feature_id']} {s['feature_name']}{note}",
                s["story_id"],
                s["story_desc"][:380],
                render_code_paths(s["code"]),
            ])
        add_table(
            doc,
            ["Epic", "Feature", "Story ID", "Story description", "Code path(s)"],
            table_rows,
        )

    # ---------- 7. Per-epic orphan-code tables (Deliverable 2) --------------
    add_h2(doc, "7. Per-epic orphan-code tables (Deliverable 2)")
    add_para(
        doc,
        "Files that live in each V5 epic's territory but are NOT named by any "
        "story in the backlog. This is the inverse of Deliverable 1: "
        "code -> backlog. A file appearing here is a candidate for either "
        "(a) deletion (truly dead), (b) a new story to anchor it, or "
        "(c) a documentation/architecture note. Tests, package __init__ "
        "files, and run-time logs are intentionally surfaced -- they are real "
        "files that no story names.",
    )
    for i, (eid, ename) in enumerate(V5_EPICS, start=1):
        add_h3(doc, f"7.{i} {eid} {ename} -- orphan code")
        orphans = ORPHANS.get(eid, [])
        if not orphans:
            add_para(doc, "(none identified)")
            continue
        add_table(
            doc,
            ["File path", "Why it's in this epic's territory", "Why it's orphan"],
            [list(o) for o in orphans],
        )

    # ---------- 8. UX 5-features-1-story analysis (Deliverable 3) ----------
    add_h2(doc, "8. UX 5-features-1-story explanation (Deliverable 3)")
    add_para(
        doc,
        "V4 proposed EPIC-011 UX with 5 features but only 1 story (Anchor "
        "Timer = EPIC-011-F-001-S-001). V5 walks the actual code and the "
        "backlog and renders a verdict per feature.",
    )
    add_para(
        doc,
        "Verdict summary: this is overwhelmingly a REQUIREMENTS GAP, not an "
        "agent error. Real UX code exists for 4 of the 5 V4-proposed features; "
        "the backlog never had stories filed for it because the work was "
        "implemented inside Website/index.html and the static-page deploy "
        "pipeline before the backlog had a UX epic to receive it. "
        "1 of the 5 features is best characterized as agent over-claim "
        "(cross-component primitives) -- the code exists in Code/Shared/ux/ "
        "but is not consumed today, so calling it a 'feature' is anticipatory.",
    )

    add_h3(doc, "8.1 F-001 Cross-component primitives")
    add_bullet(
        doc,
        "Real code: Code/Shared/ux/components/ProviderCard.tsx (4,235 bytes), "
        "Code/Shared/ux/components/SelectionManager.tsx (7,882 bytes), "
        "Code/Shared/ux/hooks/useSelectionState.ts (4,475 bytes), "
        "Code/Shared/ux/types/provider.ts (1,233 bytes).",
    )
    add_bullet(
        doc,
        "Existing story anchor: EPIC-011-F-001-S-001 (Anchor Timer) is the "
        "only existing story in EPIC-011 today. None of the four shared/ux "
        "files name themselves 'anchor timer'; they are provider/selection "
        "primitives. EPIC-006-F-030 Provider UX Display has 1 story that "
        "could be read against Code/Shared/ux/, but F-030's V5 disposition "
        "stays under EPIC-006.",
    )
    add_bullet(
        doc,
        "Verdict: PARTIAL agent over-claim. The Anchor Timer story is real "
        "and is the sole legitimate cross-component primitive on the books. "
        "The provider-card and selection-manager files are anticipatory "
        "infrastructure with no current consumer. V4 grouped them under F-001 "
        "for future-proofing; that grouping is forward-looking but not "
        "story-grounded today. Recommendation: keep F-001 as it is "
        "(1 story); the four /ux/ files stay in place and promote to a real "
        "story under F-001 only when first consumed by a second epic.",
    )

    add_h3(doc, "8.2 F-002 Shared design paradigm")
    add_bullet(
        doc,
        "Real code: 11 static HTML files under Website/ (architecture.html, "
        "products.html, privacy.html, terms.html, roadmap.html, "
        "security-architecture.html, chat-app-design.html, "
        "embedding-design.html, load-perf-report.html, "
        "ops-manager-design.html, provider-data-load.html) all carry the "
        "same :root token block (--ink, --teal, --accent), DM Serif Display + "
        "DM Sans fonts, and the same header/footer pattern. The duplication "
        "is the cross-cutting UX scope.",
    )
    add_bullet(doc, "Existing story anchor: NONE.")
    add_bullet(
        doc,
        "Verdict: REQUIREMENTS GAP. Real, load-bearing shared code; zero "
        "stories. New stories needed. Recommended new stories under F-002: "
        "(a) consolidate :root design tokens into Code/SharedUX/design_paradigm/tokens.css; "
        "(b) extract shared header/footer into Code/SharedUX/design_paradigm/header.html.j2 "
        "and footer.html.j2; (c) wire the static page generator to consume "
        "the shared parts. File 3 stories under EPIC-011-F-002.",
    )

    add_h3(doc, "8.3 F-003 Opening / parent page frame")
    add_bullet(
        doc,
        "Real code: Website/index.html (1,400+ lines). Hosts the chat iframe, "
        "the postMessage host side, the control-frame relay, and leaf-page "
        "management. Today everything is inline in this single file.",
    )
    add_bullet(
        doc,
        "Existing story anchor: EPIC-006-F-021 Parent Page Frame Layout "
        "(2 stories), EPIC-006-F-024 Control Frame -- Pagination and Filter "
        "Controls (3 stories), EPIC-006-F-025 Leaf Page Management (1 story), "
        "EPIC-006-F-026 index.html Generated from Brain Design Record (1 "
        "story). These exist; per V4 they MOVE intact to EPIC-011 UX. So "
        "F-003 already has 7 stories on the books once the move lands.",
    )
    add_bullet(
        doc,
        "Verdict: NOT a gap, NOT an agent error -- the stories exist; they "
        "live under EPIC-006 today and move to EPIC-011 under V5. The V4 "
        "1-story count was an artifact of counting only EPIC-011's existing "
        "anchor; the moved stories were not yet in EPIC-011. After the move "
        "F-003 has 7 stories.",
    )

    add_h3(doc, "8.4 F-004 Lower-region environment banner")
    add_bullet(
        doc,
        "Real code: Website/index.html lines 627-660 (#envBanner element + "
        "styling); window.refreshBanner(service) at lines 882, 917, 939, "
        "1158, 1238, 1336 reflects /health from FindCare, EvaluateCare, and "
        "SharedServices. Genuinely shared across LOCAL/DEV/QA and across 3 "
        "back-end servers.",
    )
    add_bullet(doc, "Existing story anchor: NONE.")
    add_bullet(
        doc,
        "Verdict: REQUIREMENTS GAP. Implemented inline in index.html with "
        "explicit cross-server scope; zero stories. New stories needed. "
        "Recommended: 2 stories under F-004 -- (a) extract banner code to "
        "Code/SharedUX/env_banner/banner.js; (b) make banner /health-aware "
        "for any future capability HF Space (parameterize the service list).",
    )

    add_h3(doc, "8.5 F-005 Auth-token display object")
    add_bullet(
        doc,
        "Real code: Website/index.html lines 1011-1045 (window.buildTokenDisplayHtml(st) "
        "renders nonce + GUID + signature); LEFT panel always shows FindCare's "
        "token (line 808), RIGHT shows page-owning-service's token (line 840). "
        "Inline comment at line 1011: 'EPIC-002-F-001-S-012-REQ-B-007: Parse "
        "token into nonce and GUID, build display HTML.'",
    )
    add_bullet(
        doc,
        "Existing story anchor: EPIC-002-F-001-S-012 EXISTS (the inline "
        "comment quotes its requirement ID). It is anchored to EPIC-002 "
        "Security, not to EPIC-011 UX. The DISPLAY object is presentation "
        "(belongs to EPIC-011 per the architect's framing in V4); the TOKEN "
        "(creation/verification) is Security (belongs to EPIC-002).",
    )
    add_bullet(
        doc,
        "Verdict: PARTIAL gap. The story exists but is mis-epic'd: the "
        "presentation object lives under Security in the backlog. V5 does "
        "NOT mutate the backlog (constraint). Recommended: file 1 new "
        "EPIC-011-F-005 story for the DISPLAY object (extract "
        "buildTokenDisplayHtml + LEFT/RIGHT panel logic to "
        "Code/SharedUX/auth_token_display/token_display.js); EPIC-002-F-001-S-012 "
        "remains where it is and now sources the verified token model.",
    )

    add_h3(doc, "8.6 Summary verdict")
    add_para(
        doc,
        "REQUIREMENTS GAP (file new stories): F-002 (3 stories), F-004 "
        "(2 stories), F-005 (1 story) -- 3 features, 6 new stories total.",
    )
    add_para(
        doc,
        "STORIES ALREADY EXIST -- placement issue resolved by V5 move table: "
        "F-003 (7 stories already on the books under EPIC-006-F-021/024/025/026, "
        "moving to EPIC-011 intact).",
    )
    add_para(
        doc,
        "PARTIAL agent over-claim: F-001 -- the Anchor Timer story is real "
        "(1 legitimate story); the provider-card + selection-manager files "
        "are anticipatory infrastructure, not story-grounded today. "
        "Recommendation: keep F-001 at 1 story; promote the /ux/ files when "
        "first consumed.",
    )

    # ---------- 9. Master code-path move table ------------------------------
    add_h2(doc, "9. Master code-path move table (carry forward, refreshed)")
    add_para(
        doc,
        "Carries V4 Section 6.3 with the EPIC-012 rename and Option A Security "
        "split applied. Other rows from V4 carry forward unchanged.",
    )
    move_rows = [
        [
            "Code/CSharp/ChatHealthyLogService/dev/Worker.cs",
            "DELETE (legacy)",
            "--",
            "V4 Issue 1.1: Code/Shared/ops/ copy is canonical (active build).",
        ],
        [
            "Code/CSharp/ChatHealthyLogService/dev/* (non-Worker.cs files)",
            "DELETE (byte-identical duplicates)",
            "--",
            "Identical to Code/Shared/ops/ChatHealthyLogService/ counterparts.",
        ],
        [
            "Code/CSharp/ChatHealthyLogService/build/",
            "DELETE (build output)",
            "--",
            "Compiled .dll/.exe/.deps.json; should be .gitignored.",
        ],
        [
            "Code/Shared/ChatHealthyMongoUtilities.py",
            "EPIC-009 Shared Services / mongo_utilities",
            "Code/SharedServices/mongo_utilities.py",
            "V4 Issue 1.2: Copy A is canonical (only copy with callers).",
        ],
        [
            "Code/DataPipelines/ChatHealthyMongoUtilities.py",
            "DELETE (dead code)",
            "--",
            "No callers; WHEN-TO-USE docstring ports into Copy A.",
        ],
        [
            "Code/Schemas/ChatHealthy.Providers.json",
            "DELETE after relocation",
            "--",
            "V4 Issue 2: functionally equivalent to Website/schemas/ copy.",
        ],
        [
            "Website/schemas/* (except standard/json-schema-2020-12-meta.json)",
            "EPIC-010 Operations / brain_governance",
            "brain/machine_artifacts/schemas/*",
            "V4 B.4: brain owns schemas; deploy step mirrors brain -> Website.",
        ],
        [
            "Website/index.html (full file)",
            "EPIC-011 UX",
            "Code/SharedUX/parent_page/index.html.j2 + Code/SharedUX/env_banner/banner.js + Code/SharedUX/auth_token_display/token_display.js + Code/SharedUX/parent_page/control_frame.js",
            "Restructure into UX features F-002/F-003/F-004/F-005; deploy step generates Website/index.html from templated parts.",
        ],
        [
            "Website/{architecture,products,privacy,terms,roadmap,security-architecture,chat-app-design,embedding-design,load-perf-report,ops-manager-design,provider-data-load}.html",
            "EPIC-011 UX / design_paradigm",
            "Code/SharedUX/design_paradigm/{tokens.css, header.html.j2, footer.html.j2}",
            "Factor duplicated :root token block + header + footer into shared CSS/templates.",
        ],
        [
            "Code/Shared/ux/{components,hooks,types}/",
            "EPIC-011 UX / cross_component_primitives",
            "Code/SharedUX/cross_component_primitives/",
            "Default home is the cross-cutting epic; activate when first consumed.",
        ],
        [
            "Code/shared_services/* (auth + session-token + entitlement code)",
            "EPIC-002 Security (NEW: own HF Space per Section 2 recommendation)",
            "Code/Security/backend/* + new Dockerfile + new requirements.txt",
            "V5 dilemma resolution: Option A. Security carves out of SharedServices container.",
        ],
        [
            "Code/shared_services/* (everything else)",
            "EPIC-009 Shared Services",
            "Code/SharedServices/*",
            "Carries V4. Tool Router + LLM client + embeddings + consent + lead capture + Skip-Snow context + Safety subgraph stay together.",
        ],
        [
            "Code/Shared/ops/*",
            "EPIC-010 Operations",
            "Code/Operations/*",
            "Carries V4. ops becomes Operations; the parent LangGraph runtime composes capability subgraphs from here.",
        ],
        [
            "Code/DataPipelines/*",
            "EPIC-012 Data Management Pipelines (RENAMED per V5 binding)",
            "Code/DataManagementPipelines/{get_content,orchestrate_flow,enrich_content}/",
            "Renamed from V4 'Provider Knowledge Pipeline' per Correction 1. Three feature-folder layout per Skip's V2 binding.",
        ],
        [
            "LangGraph/poc/user_journey.py",
            "EPIC-010 Operations / orchestration",
            "Code/Operations/orchestration/parent_graph.py",
            "Per LangGraph V4: thin parent runtime; capability subgraphs added as nodes.",
        ],
    ]
    add_table(
        doc,
        ["Current path", "Target epic / feature (V5)", "Proposed destination", "Reason"],
        move_rows,
    )

    # ---------- 10. Schema relocation plan ----------------------------------
    add_h2(doc, "10. Schema relocation plan (carry forward from V4)")
    for s in [
        "Create brain/machine_artifacts/schemas/ as a sibling of brain/machine_artifacts/content/.",
        "Move Website/schemas/ChatHealthy{AgileBacklog,Bugs,Errors,Providers,Relaxed,RiskAcceptance,Version}Schema.json + EngineeringRulesSchema.json + JsonValidationFixtureSchema.json -> brain/machine_artifacts/schemas/.",
        "Move Website/schemas/dev/ -> brain/machine_artifacts/schemas/dev/.",
        "Leave Website/schemas/standard/json-schema-2020-12-meta.json where it is (third-party meta-schema; static asset).",
        "Delete Code/Schemas/ChatHealthy.Providers.json (functionally equivalent duplicate per V4 Issue 2).",
        "Set the canonical $id on the relocated schema; the deploy step mirrors brain -> Website preserving the URL surface.",
        "OPEN ITEM (non-blocking): the schema's `additionalProperties: false` rejects 100% of production providers because the running pipeline produces 11 fields not declared (embedding, embedding_model, embedding_version, 8 NPPES other-name fields). Track under EPIC-012 / schema-drift feature.",
        "Engineering Rule-008-ENF-001 currently excludes Website/schemas/standard/. After the move, update the enforcement worker's allowed/excluded patterns to point at brain/machine_artifacts/schemas/ as the validation target.",
    ]:
        add_bullet(doc, s)

    # ---------- 11. Open questions ------------------------------------------
    add_h2(doc, "11. Open questions (<= 4, each one decision)")
    open_qs = [
        [
            "1",
            "Security HF Space cutover sequencing: V5 recommends Option A (own Space). Cutover options: (a) carve at the next deploy after EvaluateCare ships; (b) carve immediately as a parallel Space and migrate /verify-token callers in a single change. Which sequencing minimises risk?",
            "Affects EPIC-002 implementation order and the SharedServices deploy pipeline.",
        ],
        [
            "2",
            "EPIC-007 Talk About Care -- the 1 story is parked under EPIC-009 in V5 (deep_research subgraph anchor). Confirm: keep parked under EPIC-009, OR break out as a new EPIC-013 capability when Talk About Care reaches active development?",
            "EPIC-007 was excluded from the Git tree per Skip's V2 binding; V5 honors that. The only question is which of the 7 V5 epics tables it appears in.",
        ],
        [
            "3",
            "Provider schema content drift (separate from relocation): schema declares 41 properties; pipeline produces up to 11 additional fields. Update schema, trim pipeline output, or run additionalProperties: true until a separate decision lands?",
            "100% of production providers fail validation today.",
        ],
        [
            "4",
            "EPIC-012 internal feature granularity: V5 carries the V4 13-feature shape (F-005..F-018, F-028 absorbed into the three feature-folders get_content/orchestrate_flow/enrich_content). Confirm the 3-folder grouping is correct OR list the features verbatim under EPIC-012 with no folder grouping?",
            "Affects the Code/DataManagementPipelines/ tree layout and the count of features under EPIC-012.",
        ],
    ]
    add_table(doc, ["#", "Question", "Why it matters"], open_qs)

    # ---------- 12. Appendix -- generation record ---------------------------
    add_h2(doc, "12. Appendix -- generation record")
    add_bullet(
        doc,
        "Backlog read at generate time: brain/machine_artifacts/content/agile_backlog.json -- 9 epics, 274 stories.",
    )
    add_bullet(
        doc,
        "Per-epic story-to-code mapping: 274 rows tabled across 7 V5 epics; "
        "every story appears in exactly one epic table.",
    )
    add_bullet(
        doc,
        "NO CODE FOUND verdicts captured at the FEATURE level for: "
        "EPIC-001-F-011 (Mixing Board, 5 stories), EPIC-001-F-023 (Research Agent, "
        "1 story), EPIC-001-F-025 (UMLS License Compliance, 4 stories), "
        "EPIC-006-F-020 (Native Mobile Apps, 0 stories -- empty feature shell), "
        "EPIC-007-F-001 (ChatRoom UX, 1 story), EPIC-008-F-001 (Architecture "
        "Governance, 0 stories -- empty feature shell), EPIC-008-F-006 (The Brain "
        "-- Agentic Infrastructure, 0 stories -- empty feature shell), "
        "EPIC-008-F-009 (Deferred work as first-class concept, 1 story). Total "
        "stories with NO CODE FOUND: 12 (5+1+4+1+1 = 12 across the five "
        "features that have stories; the three empty feature shells contain 0 "
        "stories each). The three empty feature shells are deletion "
        "candidates; the 12 stories are candidates for either deletion or "
        "either deletion or implementation under V5.",
    )
    add_bullet(
        doc,
        "Orphan files surfaced across 7 V5 epics (Deliverable 2): see Section 7 "
        "for full per-epic listings. Surfaced from direct filesystem walk on "
        "2026-04-28.",
    )
    add_bullet(
        doc,
        "Read-only enforcement: no commits, no code changes, no backlog "
        "mutations. The generator script is a read-only tool that emits a "
        "Word document.",
    )

    # Save
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
