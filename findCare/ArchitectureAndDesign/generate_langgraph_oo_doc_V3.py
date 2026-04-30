"""Generate langgraph-oo-best-practices-V3.docx.

V3 carries forward V2's content (which is V1 + four architect ``</Skip>`` inline
annotations) and ADDS a new top-level section, "Architect questions from V2
annotations," with one subsection per annotation. Each new subsection quotes
the architect's note verbatim, then answers it with research-backed citations
fetched on 2026-04-28 and a clear recommendation.

Tag convention in citations:
    [official]  - LangChain / LangGraph / MongoDB-maintainer-authored
    [community] - third-party blog, tutorial, or independent repo
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor


OUT = Path(__file__).with_name("langgraph-oo-best-practices-V3.docx")


# ============================================================
# Citations dictionary - V1 set + new sources for the architect
# question section.
# ============================================================

CITES = {
    # --- Carried forward from V1/V2 ---
    "docs_graph_api":        ("official", "Use the graph API - LangChain Docs",
                              "https://docs.langchain.com/oss/python/langgraph/use-graph-api"),
    "docs_persistence":      ("official", "Persistence - LangChain Docs",
                              "https://docs.langchain.com/oss/javascript/langgraph/persistence"),
    "docs_app_structure":    ("official", "Application structure - LangChain Docs",
                              "https://docs.langchain.com/oss/python/langgraph/application-structure"),
    "docs_tools":            ("official", "Tools - LangChain Docs",
                              "https://docs.langchain.com/oss/python/langchain/tools"),
    "docs_context":          ("official", "Context overview - LangChain Docs",
                              "https://docs.langchain.com/oss/python/concepts/context"),
    "docs_test":             ("official", "Test - LangChain Docs",
                              "https://docs.langchain.com/oss/python/langgraph/test"),
    "docs_thinking":         ("official", "Thinking in LangGraph - LangChain Docs",
                              "https://docs.langchain.com/oss/python/langgraph/thinking-in-langgraph"),
    "blog_building":         ("official", "Building LangGraph: an agent runtime from first principles",
                              "https://www.langchain.com/blog/building-langgraph"),
    "blog_v1":               ("official", "LangChain and LangGraph reach v1.0",
                              "https://www.langchain.com/blog/langchain-langgraph-1dot0"),
    "blog_klarna":           ("official", "How Klarna's AI assistant redefined customer support",
                              "https://www.langchain.com/blog/customers-klarna"),
    "blog_in_prod":          ("official", "Is LangGraph used in production?",
                              "https://www.langchain.com/blog/is-langgraph-used-in-production"),
    "blog_v02_ckpt":         ("official", "LangGraph v0.2: customization with new checkpointer libraries",
                              "https://blog.langchain.com/langgraph-v0-2/"),
    "gh_langgraph":          ("official", "langchain-ai/langgraph",
                              "https://github.com/langchain-ai/langgraph"),
    "gh_example_pyproject":  ("official", "langchain-ai/langgraph-example-pyproject",
                              "https://github.com/langchain-ai/langgraph-example-pyproject"),
    "gh_open_deep_research": ("official", "langchain-ai/open_deep_research",
                              "https://github.com/langchain-ai/open_deep_research"),
    "gh_deep_research_scratch": ("official", "langchain-ai/deep_research_from_scratch",
                              "https://github.com/langchain-ai/deep_research_from_scratch"),
    "gh_deepagents":         ("official", "langchain-ai/deepagents",
                              "https://github.com/langchain-ai/deepagents"),
    "gh_swarm":              ("official", "langchain-ai/langgraph-swarm-py",
                              "https://github.com/langchain-ai/langgraph-swarm-py"),
    "gh_supervisor":         ("official", "langchain-ai/langgraph-supervisor-py",
                              "https://github.com/langchain-ai/langgraph-supervisor-py"),
    "gh_issue_1950":         ("official", "Issue #1950: Passing private state with Class Node",
                              "https://github.com/langchain-ai/langgraph/issues/1950"),
    "gh_awesome":            ("community", "von-development/awesome-LangGraph",
                              "https://github.com/von-development/awesome-LangGraph"),
    "blog_swarnendu":        ("community", "Swarnendu De - LangGraph Best Practices",
                              "https://www.swarnendu.de/blog/langgraph-best-practices/"),
    "blog_yasin_modular":    ("community", "Yassin Hashem - Scaling AI Agents Beyond Notebooks: A Modular Architecture",
                              "https://medium.com/@yasin162001/scaling-ai-agents-beyond-notebooks-a-modular-architecture-for-langgraph-in-production-4711764de464"),
    "blog_shaza_typing":     ("community", "Shaza Ali - Type Safety in LangGraph: TypedDict vs Pydantic",
                              "https://shazaali.substack.com/p/type-safety-in-langgraph-when-to"),
    "blog_pankaj_pyd":       ("community", "Pankaj Chandravanshi - LangGraph: State with Pydantic BaseModel",
                              "https://medium.com/fundamentals-of-artificial-intelligence/langgraph-state-with-pydantic-basemodel-023a2158ab00"),
    "blog_easton_2026":      ("community", "BetterLink Blog - LangGraph State Management in Practice (2026)",
                              "https://eastondev.com/blog/en/posts/ai/20260424-langgraph-agent-architecture/"),
    "blog_markaicode":       ("community", "Markaicode - Production Multi-Agent System with LangGraph",
                              "https://markaicode.com/langgraph-production-agent/"),
    "blog_clean_state":      ("community", "Vishal Lad - Clean State Architecture in LangGraph (Mar 2026)",
                              "https://medium.com/@ladvishal1985/everything-ive-learned-about-clean-state-architecture-in-langgraph-6d1352b0c00c"),
    "blog_unit_testing":     ("community", "Anirudh Sharma - Unit Testing LangGraph: Nodes and Flow Paths",
                              "https://medium.com/@anirudhsharmakr76/unit-testing-langgraph-testing-nodes-and-flow-paths-the-right-way-34c81b445cd6"),
    "blog_mock_llm":         ("community", "Matt Carvalho - How to Properly Mock LangChain LLM Execution",
                              "https://medium.com/@matgmc/how-to-properly-mock-langchain-llm-execution-in-unit-tests-python-76efe1b8707e"),
    "blog_tools_first":      ("community", "SitePoint - Implementing the Tools-First Pattern in LangGraph",
                              "https://www.sitepoint.com/implementing-the-tools-first-pattern-in-lang-graph/"),
    "gh_fastapi_template":   ("community", "wassim249/fastapi-langgraph-agent-production-ready-template",
                              "https://github.com/wassim249/fastapi-langgraph-agent-production-ready-template"),

    # --- New for V3 - architect annotation answers ---
    "docs_handoffs":         ("official", "Multi-agent handoffs - LangChain Docs",
                              "https://docs.langchain.com/oss/python/langchain/multi-agent/handoffs"),
    "docs_interrupts":       ("official", "Interrupts - LangChain Docs",
                              "https://docs.langchain.com/oss/python/langgraph/interrupts"),
    "docs_add_memory":       ("official", "Memory (add-memory) - LangChain Docs",
                              "https://docs.langchain.com/oss/python/langgraph/add-memory"),
    "docs_custom_ckpt":      ("official", "How to use a custom checkpointer - LangChain Docs",
                              "https://docs.langchain.com/langsmith/custom-checkpointer"),
    "ref_checkpoints":       ("official", "checkpoints API reference - LangGraph",
                              "https://reference.langchain.com/python/langgraph/checkpoints"),
    "pypi_mongo_ckpt":       ("official", "langgraph-checkpoint-mongodb on PyPI (MongoDB + LangChain co-maintained)",
                              "https://pypi.org/project/langgraph-checkpoint-mongodb/"),
    "mongo_atlas_lg":        ("official", "Integrate MongoDB with LangGraph - MongoDB Atlas Docs",
                              "https://www.mongodb.com/docs/atlas/ai-integrations/langgraph/"),
    "mongo_atlas_lg_agents": ("official", "Build an AI Agent with LangGraph and MongoDB Atlas",
                              "https://www.mongodb.com/docs/atlas/ai-integrations/langgraph/build-agents/"),
    "mongo_lc_readthedocs":  ("official", "MongoDBSaver class - LangChain MongoDB Read the Docs",
                              "https://langchain-mongodb.readthedocs.io/en/latest/langgraph_checkpoint_mongodb/saver/langgraph.checkpoint.mongodb.saver.MongoDBSaver.html"),
    "mongo_blog":            ("official", "Checkpointers and Native Parent-Child Retrievers with LangChain and MongoDB",
                              "https://www.mongodb.com/company/blog/innovation/checkpointers-native-parent-child-retrievers-with-langchain-mongodb"),
    "deepwiki_subgraphs":    ("community", "Graph composition and nested graphs - DeepWiki (langchain-ai/langgraph)",
                              "https://deepwiki.com/langchain-ai/langgraph/3.6-graph-composition-and-nested-graphs"),
    "deepwiki_hitl":         ("community", "Human-in-the-loop and interrupts - DeepWiki (langchain-ai/langgraph)",
                              "https://deepwiki.com/langchain-ai/langgraph/3.7-human-in-the-loop-and-interrupts"),
    "blog_send_api":         ("community", "Sreeni - Leveraging LangGraph's Send API for Dynamic and Parallel Workflow Execution",
                              "https://dev.to/sreeni5018/leveraging-langgraphs-send-api-for-dynamic-and-parallel-workflow-execution-4pgd"),
    "blog_handoffs_tds":     ("community", "Towards Data Science - How Agent Handoffs Work in Multi-Agent Systems",
                              "https://towardsdatascience.com/how-agent-handoffs-work-in-multi-agent-systems/"),
    "blog_multiagent_2026":  ("community", "Mager - LangGraph: Build Stateful Multi-Agent Systems That Don't Crash (2026)",
                              "https://www.mager.co/blog/2026-03-12-langgraph-deep-dive/"),
    "blog_scaling_lg":       ("community", "AI Practitioner - Scaling LangGraph Agents: Parallelization, Subgraphs, and Map-Reduce Trade-Offs",
                              "https://aipractitioner.substack.com/p/scaling-langgraph-agents-parallelization"),
    "dev_mongo_ltm":         ("community", "DEV (MongoDB) - LangGraph With MongoDB: Building Conversational Long-Term Memory",
                              "https://dev.to/mongodb/langgraph-with-mongodb-building-conversational-long-term-memory-for-intelligent-ai-agents-2pcn"),
}


def cite_text(key: str) -> str:
    tag, title, url = CITES[key]
    return f"[{tag}] {title} - {url}"


# ============================================================
# Document builder helpers
# ============================================================

def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    return p


def add_quote(doc, text):
    """Indented italic block for the architect's verbatim note."""
    p = doc.add_paragraph(style="Intense Quote")
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(11)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p


def add_citation(doc, key):
    tag, title, url = CITES[key]
    p = doc.add_paragraph(style="List Bullet 2")
    run = p.add_run(f"[{tag}] {title} - ")
    run.font.size = Pt(10)
    run.italic = True
    link = p.add_run(url)
    link.font.size = Pt(10)
    link.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p


# ============================================================
# Build
# ============================================================

def build():
    doc = Document()

    # Title
    title = doc.add_heading("LangChain / LangGraph OO Best-Practices Guide - V3", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sr = sub.add_run("V2 content carried forward + new section answering the four "
                     "architect </Skip> annotations on V2")
    sr.italic = True
    sr.font.size = Pt(11)
    meta = doc.add_paragraph()
    meta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    mr = meta.add_run("Compiled 2026-04-28  -  Skip Snow / ChatHealthy  -  "
                      "Sources tagged [official] = LangChain/LangGraph or MongoDB "
                      "co-maintained, [community] = third-party")
    mr.font.size = Pt(9)
    mr.italic = True

    # ---------------------------------------------------
    # 1. Executive Summary
    # ---------------------------------------------------
    add_heading(doc, "1. Executive Summary", level=1)

    add_para(doc,
        "The dominant production pattern in late-2025 / 2026 LangGraph is: "
        "(a) define state as a typed schema in its own file, with TypedDict acceptable "
        "for prototypes but a Pydantic BaseModel preferred for production validation; "
        "(b) carry conversation as messages: list[AnyMessage] with the add_messages "
        "reducer rather than as bespoke fields; (c) decompose the graph into "
        "subgraphs each owning a focused state schema (Single-Responsibility applied "
        "to graphs); (d) implement nodes as small pure functions or as callable "
        "classes with __call__ when they need encapsulated dependencies, returning "
        "partial state updates, never mutating; (e) inject runtime dependencies "
        "(DB clients, LLM clients, config) via the LangGraph context / Runtime API "
        "rather than module-level imports; (f) persist with a production-grade "
        "checkpointer - PostgresSaver is the LangChain reference, but MongoDBSaver "
        "(co-maintained by MongoDB and LangChain) is an officially supported "
        "alternative and is the right choice for ChatHealthy because Mongo is "
        "already in the stack; (g) lay the project out as state.py / nodes.py / "
        "tools.py / agent.py - the layout the official langgraph-example template "
        "ships with. The ChatHealthy POC user_journey.py - 336 lines, single file, "
        "15 flat TypedDict fields, bare top-level functions, tool outputs as "
        "top-level state keys - directly violates anti-patterns that the community "
        "literature explicitly names: the God State Object, the Monolithic Graph, "
        "and Overloaded Messages.")

    add_para(doc,
        "V3 specifically: section 6 (new) answers the four </Skip> annotations the "
        "architect placed on V2. Headlines: (1) one-graph-per-front-end-epic with "
        "state-handoffs IS supported and is the more common multi-agent pattern in "
        "LangGraph - it is implemented via subgraph composition or compiled-graph-as-"
        "node, with state passed by selective projection through input_schema / "
        "output_schema and routed via Command(graph=Command.PARENT); (2) prefer the "
        "officially-supported MongoDBSaver over PostgresSaver because MongoDB is "
        "already in ChatHealthy's stack and the MongoDB-LangChain co-maintained "
        "checkpointer covers the entire BaseCheckpointSaver contract including "
        "async; (3) open_deep_research is a reference implementation, NOT an anti-"
        "pattern - the Talk About Care async-research use case is exactly what the "
        "subgraph + Send + interrupt() machinery is designed for; (4) yes - the "
        "deep_research_from_scratch Pydantic-schemas-for-routing-decisions pattern "
        "(ClarifyWithUser, ResearchQuestion, ConductResearch, ResearchComplete) is "
        "the right model for the cancer-research use case, with one note about "
        "scoping versus full supervisor.")

    add_para(doc, "Key supporting sources for this summary:", italic=True)
    for k in ("docs_graph_api", "docs_app_structure", "docs_persistence",
              "docs_context", "blog_building", "blog_clean_state",
              "blog_easton_2026", "blog_markaicode",
              "pypi_mongo_ckpt", "mongo_atlas_lg", "docs_handoffs",
              "docs_interrupts", "deepwiki_subgraphs", "gh_deep_research_scratch"):
        add_citation(doc, k)

    # ---------------------------------------------------
    # 2. Per-question findings (V1/V2 content carried forward)
    # ---------------------------------------------------
    add_heading(doc, "2. Per-Question Findings", level=1)

    # Q1
    add_heading(doc, "Q1. State schema home - TypedDict, Pydantic BaseModel, dataclass, or JSON Schema?", level=2)
    add_para(doc, "Recommendation: define state in a dedicated state.py module. "
                  "TypedDict is the documented baseline and is what the LangGraph "
                  "official tutorials use first; Pydantic BaseModel is the production "
                  "upgrade once external inputs, validation, or checkpoint serialisation "
                  "matter. Dataclass is a third option, mainly chosen when Pydantic's "
                  "recursive validation cost is unacceptable.")
    add_para(doc, "Rationale grounded in the sources: the official graph-API doc states "
                  "explicitly that 'State in LangGraph can be a TypedDict, Pydantic "
                  "model, or dataclass' and presents TypedDict as the primary example. "
                  "The community guidance (Shaza Ali; Pankaj Chandravanshi; BetterLink "
                  "2026) converges on a hybrid: TypedDict for internal graph state, "
                  "Pydantic at the boundaries (inputs, outputs, tool I/O). Markaicode's "
                  "production write-up is the strongest single statement: 'Use a "
                  "Pydantic BaseModel the moment you step into production.'")
    for k in ("docs_graph_api", "blog_shaza_typing", "blog_pankaj_pyd",
              "blog_easton_2026", "blog_markaicode", "docs_app_structure"):
        add_citation(doc, k)

    # Q2
    add_heading(doc, "Q2. State persistence - which checkpointer in production?", level=2)
    add_para(doc, "Recommendation (revised in V3 section 6.2): in the general "
                  "LangChain reference, PostgresSaver is named the production target "
                  "and InMemorySaver / SqliteSaver are dev-only. For ChatHealthy "
                  "specifically, MongoDBSaver (co-maintained by MongoDB and LangChain) "
                  "is the correct choice because MongoDB is already in our stack. "
                  "InMemorySaver remains the right choice for tests.")
    add_para(doc, "Rationale grounded in the sources: the official persistence doc "
                  "labels InMemorySaver 'suitable for development and testing,' "
                  "SqliteSaver 'ideal for experimentation and local workflows,' and "
                  "PostgresSaver 'ideal for using in production.' MongoDB's official "
                  "Atlas integration documents MongoDBSaver as a first-party "
                  "implementation that 'persists agent state in MongoDB' and "
                  "'enables human-in-the-loop, memory, time travel, and fault-"
                  "tolerance for your LangGraph agents.' See section 6.2 for the "
                  "Postgres-vs-Mongo trade-off applied to ChatHealthy.")
    for k in ("docs_persistence", "blog_v02_ckpt", "blog_easton_2026",
              "blog_klarna", "pypi_mongo_ckpt", "mongo_atlas_lg"):
        add_citation(doc, k)

    # Q3
    add_heading(doc, "Q3. Node organization - top-level functions, classes, or class-based callables?", level=2)
    add_para(doc, "Recommendation: default to small pure functions returning partial "
                  "state updates. Promote a node to a class with __call__ when it needs "
                  "constructor-injected dependencies (LLM client, DB handle, config) "
                  "or when several closely-related nodes share state through "
                  "instance attributes. Avoid the procedural top-level-function-per-step "
                  "pattern beyond a prototype.")
    add_para(doc, "Rationale grounded in the sources: the official graph-API doc "
                  "defines a node as 'just a Python function that reads our graph's "
                  "state and makes updates to it' and explicitly warns 'Nodes should "
                  "return updates to the state directly, instead of mutating the state.' "
                  "Issue #1950 documents a known limitation when class-nodes use "
                  "private state, so classes are best used for dependency-injection "
                  "rather than hidden mutable internals.")
    for k in ("docs_graph_api", "docs_thinking", "gh_issue_1950"):
        add_citation(doc, k)

    # Q4
    add_heading(doc, "Q4. Tool design - @tool, BaseTool, and how tool I/O lives in state", level=2)
    add_para(doc, "Recommendation: use the @tool decorator for the common case "
                  "(typed args, docstring as description). Drop to a BaseTool "
                  "subclass when you need explicit args_schema, sync+async behaviour, "
                  "or tool-level state. Wire tools into the graph via the prebuilt "
                  "ToolNode. Model tool I/O as ToolMessage entries on the messages "
                  "channel - NOT as separate top-level state fields per tool.")
    for k in ("docs_tools", "blog_tools_first", "blog_clean_state", "docs_graph_api"):
        add_citation(doc, k)

    # Q5
    add_heading(doc, "Q5. Dependency injection - how do clients/config reach nodes and tools?", level=2)
    add_para(doc, "Recommendation: pass dependencies via the LangGraph runtime "
                  "context (Runtime[ContextSchema] for nodes, ToolRuntime[ContextSchema] "
                  "for tools), supplied at graph.invoke(...) time. For class-based "
                  "nodes, accept the dependency in __init__ and bind the instance "
                  "into the graph. Do NOT import database clients or LLM clients at "
                  "module scope inside node files.")
    for k in ("docs_context", "blog_swarnendu", "gh_fastapi_template"):
        add_citation(doc, k)

    # Q6
    add_heading(doc, "Q6. Module / file structure", level=2)
    add_para(doc, "Recommendation: at minimum the official 4-file layout - state.py, "
                  "nodes.py, tools.py, agent.py. Beyond a single agent, decompose "
                  "by capability: a base/ directory for shared utilities and a "
                  "per-agent folder containing models/, node/, workflow/. Subgraphs "
                  "become first-class units of decomposition.")
    for k in ("docs_app_structure", "gh_example_pyproject", "blog_yasin_modular",
              "gh_fastapi_template"):
        add_citation(doc, k)

    # Q7
    add_heading(doc, "Q7. State immutability and reducers", level=2)
    add_para(doc, "Recommendation: nodes return partial-update dicts; never mutate "
                  "incoming state. Annotate any list/dict field that more than one "
                  "node may write - use add_messages for the messages channel and "
                  "operator.add (or a custom reducer) for other accumulating lists. "
                  "Cap unbounded growth via custom reducers.")
    for k in ("docs_graph_api", "blog_building", "blog_clean_state"):
        add_citation(doc, k)

    # Q8
    add_heading(doc, "Q8. Testing - node isolation, graph integration, LLM mocking", level=2)
    add_para(doc, "Recommendation: two-layer split. Unit-test nodes by calling the "
                  "function (or instance) directly with a synthetic state dict and "
                  "asserting on the returned partial update - no graph compilation "
                  "needed. Integration-test the compiled graph with an in-memory "
                  "checkpointer, mocked LLM (FakeListLLM or a custom fake), and "
                  "stubbed tools; assert on final state and the chosen edges.")
    for k in ("docs_test", "blog_unit_testing", "blog_mock_llm"):
        add_citation(doc, k)

    # Q9
    add_heading(doc, "Q9. Which OO principles does the community emphasize?", level=2)

    # Reproduce the architect's V2 annotation verbatim (as quoted block).
    add_para(doc, "Architect annotation on V2 (verbatim, preserved here for context; "
                  "the full answer is in section 6.1):", italic=True)
    add_quote(doc,
        "</Skip It is my concept that every front end epic will have its own graph "
        "and when a handoff from one graph to another takes place the necasary parts "
        "of state would flow from one graph to another please think through if this "
        "degree of granularity is sufficient. Here below is our current graph:>")

    add_para(doc, "Recommendation: apply Single-Responsibility at the subgraph level "
                  "(each subgraph owns a focused slice of state); apply Encapsulation "
                  "via class-based nodes whose dependencies are constructor-injected; "
                  "prefer Composition over Inheritance - assemble graphs and subgraphs "
                  "rather than subclassing node base classes; use Dependency Injection "
                  "via the runtime context, not module-level imports. Treat "
                  "graph-as-a-state-machine: pure transitions, immutable updates, "
                  "explicit interfaces between subgraphs.")
    for k in ("blog_clean_state", "blog_easton_2026", "blog_building", "docs_thinking"):
        add_citation(doc, k)

    # Q10
    add_heading(doc, "Q10. Dialogue / message / history modeling", level=2)
    add_para(doc, "Recommendation: model conversation as messages: "
                  "Annotated[list[AnyMessage], add_messages] - either by extending "
                  "the prebuilt MessagesState or by declaring the field manually. "
                  "Tool calls and tool results live as ToolMessage entries on the "
                  "same channel. Keep non-conversational artefacts (retrieved docs, "
                  "structured tool payloads) in separate typed fields.")
    for k in ("docs_graph_api", "blog_clean_state"):
        add_citation(doc, k)

    # ---------------------------------------------------
    # 3. Anti-patterns (V1/V2 carried forward)
    # ---------------------------------------------------
    add_heading(doc, "3. Anti-Patterns to Avoid", level=1)

    add_para(doc, "The eight anti-patterns below are each named in the cited source. "
                  "Items marked POC are concretely present in "
                  "LangGraph/poc/user_journey.py.")

    add_heading(doc, "AP-1. The God State Object  (POC)", level=3)
    add_para(doc, "Quote from source: 'It's tempting to create one big state class "
                  "that every node in your graph can read and write... State that "
                  "belongs to everyone belongs to no one.' POC reality: JourneyState "
                  "has 15 unrelated fields all flat in one TypedDict.")
    add_citation(doc, "blog_clean_state")

    add_heading(doc, "AP-2. Monolithic Graph (no subgraph decomposition)  (POC)", level=3)
    add_para(doc, "Quote from source: 'I had a single LangGraph graph with 22 nodes "
                  "and one huge state object... essentially building a monolith.' "
                  "POC reality: 14 nodes and the build_user_journey() function all "
                  "live in one 336-line file.")
    add_citation(doc, "blog_clean_state")

    add_heading(doc, "AP-3. Tool outputs as flat top-level state fields  (POC)", level=3)
    add_para(doc, "Source guidance: tool I/O should ride the messages channel as "
                  "ToolMessage entries; structured artefacts go into named typed "
                  "fields - but only the artefact, not one field per tool wired "
                  "directly into top-level state.")
    add_citation(doc, "docs_graph_api")
    add_citation(doc, "blog_clean_state")

    add_heading(doc, "AP-4. Procedural top-level node functions with no encapsulation  (POC)", level=3)
    add_para(doc, "Source guidance: nodes are pure functions OR classes-with-__call__ "
                  "when they have dependencies. POC reality: every node is a top-"
                  "level def in user_journey.py with no module boundary, no class "
                  "wrappers, and stub bodies.")
    add_citation(doc, "docs_app_structure")
    add_citation(doc, "blog_yasin_modular")

    add_heading(doc, "AP-5. Module-level imports of clients instead of dependency injection", level=3)
    add_para(doc, "Anti-pattern: top-of-file import of a singleton DB client "
                  "referenced inside node bodies. Use Runtime[ContextSchema] / "
                  "ToolRuntime[ContextSchema] instead.")
    add_citation(doc, "docs_context")

    add_heading(doc, "AP-6. Missing reducers on list fields  (POC risk)", level=3)
    add_para(doc, "Quote from source: 'If you don't set a reducer, LangGraph uses "
                  "last-write-wins by default... Always give it a reducer, even if "
                  "only one node writes to it now.' POC reality: history, "
                  "specialties, providers, trials, and homeopathic_expansion are "
                  "all bare list with no annotation; the parallel safety+classifier "
                  "fan-out is exactly the topology where last-write-wins silently "
                  "loses data.")
    add_citation(doc, "blog_clean_state")
    add_citation(doc, "docs_graph_api")

    add_heading(doc, "AP-7. Overloaded messages field", level=3)
    add_para(doc, "Quote: 'The messages list is easy to append to, and before long "
                  "it can end up holding retrieved documents, intermediate reasoning, "
                  "tool outputs, and more.' Fix: keep messages for the conversation; "
                  "store retrieval/tool artefacts in named typed fields.")
    add_citation(doc, "blog_clean_state")

    add_heading(doc, "AP-8. SqliteSaver / InMemorySaver in production", level=3)

    # Architect's V2 annotation reproduced verbatim
    add_para(doc, "Architect annotation on V2 (verbatim, preserved here; full answer "
                  "in section 6.2):", italic=True)
    add_quote(doc, "</Skip Is there a compelling reason to use Postgress over Mongo?>")

    add_para(doc, "Quote (official): 'InMemorySaver: suitable for development and "
                  "testing... SqliteSaver: ideal for experimentation and local "
                  "workflows... PostgresSaver: ideal for using in production.' "
                  "Quote (community 2026): 'Skip SqliteSaver entirely; use "
                  "PostgresSaver with async (AsyncPostgresSaver) for production to "
                  "avoid write-lock bottlenecks.' V3 update: this anti-pattern is "
                  "narrow - the issue is using a dev-only checkpointer in "
                  "production, NOT specifically requiring Postgres. MongoDBSaver, "
                  "co-maintained by MongoDB and LangChain, is also production-grade.")
    add_citation(doc, "docs_persistence")
    add_citation(doc, "blog_easton_2026")
    add_citation(doc, "pypi_mongo_ckpt")
    add_citation(doc, "mongo_atlas_lg")

    # ---------------------------------------------------
    # 4. Reference implementations (V1/V2 carried forward)
    # ---------------------------------------------------
    add_heading(doc, "4. Reference Implementations", level=1)

    add_para(doc, "Five real GitHub repositories that exemplify the patterns in this "
                  "guide. Star counts and dates were captured during the research "
                  "pass on 2026-04-28.")

    add_heading(doc, "R-1. langchain-ai/langgraph (the framework itself)", level=3)
    add_para(doc, "30.7k stars. Authoritative source for the graph runtime, "
                  "checkpointer interfaces, and example notebooks under examples/.")
    add_citation(doc, "gh_langgraph")

    add_heading(doc, "R-2. langchain-ai/langgraph-example-pyproject", level=3)
    add_para(doc, "52 stars but officially curated. The canonical layout: "
                  "my_agent/utils/{state.py, nodes.py, tools.py} + my_agent/agent.py "
                  "+ langgraph.json.")
    add_citation(doc, "gh_example_pyproject")
    add_citation(doc, "docs_app_structure")

    add_heading(doc, "R-3. langchain-ai/open_deep_research", level=3)

    # Architect's V2 annotation reproduced verbatim
    add_para(doc, "Architect annotation on V2 (verbatim, preserved here; full answer "
                  "in section 6.3):", italic=True)
    add_quote(doc,
        "</Skip we have need for this, espeshially in 'Talk about care'  if this is "
        "an anti pattern I'm  a bit confused. Use case: User has stage 4 cancer and "
        "wants to know about cutting edge research where no FDA approval of on "
        "label drugs exists. User is in 'Talk about care' and the bots, people and "
        "principle (the person who is hosting the 'Talk about care meeting) decide "
        "together the only way to understand what cutting edge drugs are out there "
        "is deep research. LangGraph challenge in relation to this pattern is that "
        "for me the natural thing to do is establish the deep research node in the "
        "Talk About care sub graph and have an agent go off and do the research "
        "while the meeting participants continue to talk. Help me understand how "
        "we will model this use case. There is probably cause for deep research to "
        "be required by find care and evaluate care as well.>")

    add_para(doc, "11.3k stars. Production-grade research agent. src/ layout, "
                  "configurable models for distinct LLM roles (summariser, "
                  "researcher, compressor, writer), legacy/ folder preserves an "
                  "older workflow + multi-agent supervisor implementation for "
                  "reference. NOTE: open_deep_research is not an anti-pattern; it "
                  "is a reference pattern - see section 6.3 for application to "
                  "Talk About Care.")
    add_citation(doc, "gh_open_deep_research")

    add_heading(doc, "R-4. langchain-ai/deepagents", level=3)
    add_para(doc, "21.9k stars; latest release 2026-04-15. Agent harness with "
                  "planning, filesystem, sub-agent delegation. LangGraph-native; "
                  "compiled graphs are returned for downstream composition.")
    add_citation(doc, "gh_deepagents")

    add_heading(doc, "R-5. wassim249/fastapi-langgraph-agent-production-ready-template", level=3)
    add_para(doc, "2.2k stars (community). Realistic FastAPI service skeleton: "
                  "app/api/v1/ + app/core/langgraph/ + app/services/ + "
                  "app/schemas/ + app/models/.")
    add_citation(doc, "gh_fastapi_template")

    add_heading(doc, "R-6 (bonus). langchain-ai/deep_research_from_scratch", level=3)

    # Architect's V2 annotation reproduced verbatim
    add_para(doc, "Architect annotation on V2 (verbatim, preserved here; full answer "
                  "in section 6.4):", italic=True)
    add_quote(doc,
        "</ Skip see note with use case  above should we follow the model "
        "suggested by this pattern>")

    add_para(doc, "708 stars. Tutorial repo (5 notebooks: scoping, research_agent, "
                  "research_agent_mcp, research_supervisor, full_agent). Notable "
                  "because it explicitly uses Pydantic schemas for routing decisions "
                  "(ClarifyWithUser, ResearchQuestion in scoping; ConductResearch, "
                  "ResearchComplete as supervisor delegation tools).")
    add_citation(doc, "gh_deep_research_scratch")

    add_para(doc, "Index of the wider ecosystem (community-curated):")
    add_citation(doc, "gh_awesome")

    # ---------------------------------------------------
    # 5. Bibliography (V1/V2 carried forward)
    # ---------------------------------------------------
    add_heading(doc, "5. Citations Bibliography", level=1)

    def section(label, keys):
        add_heading(doc, label, level=2)
        for k in keys:
            tag, title, url = CITES[k]
            p = doc.add_paragraph()
            r = p.add_run(f"[{tag}] {title}\n")
            r.font.size = Pt(10)
            r.bold = True
            r2 = p.add_run(url)
            r2.font.size = Pt(10)
            r2.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    section("Official LangChain / LangGraph documentation", [
        "docs_graph_api", "docs_persistence", "docs_app_structure",
        "docs_tools", "docs_context", "docs_test", "docs_thinking",
        "docs_handoffs", "docs_interrupts", "docs_add_memory",
        "docs_custom_ckpt", "ref_checkpoints",
    ])
    section("Official LangChain blog", [
        "blog_building", "blog_v1", "blog_klarna", "blog_in_prod", "blog_v02_ckpt",
    ])
    section("Official MongoDB / LangChain co-maintained sources (new in V3)", [
        "pypi_mongo_ckpt", "mongo_atlas_lg", "mongo_atlas_lg_agents",
        "mongo_lc_readthedocs", "mongo_blog",
    ])
    section("Official GitHub repositories", [
        "gh_langgraph", "gh_example_pyproject", "gh_open_deep_research",
        "gh_deep_research_scratch", "gh_deepagents", "gh_swarm",
        "gh_supervisor", "gh_issue_1950",
    ])
    section("Community blog posts and tutorials", [
        "blog_swarnendu", "blog_yasin_modular", "blog_shaza_typing",
        "blog_pankaj_pyd", "blog_easton_2026", "blog_markaicode",
        "blog_clean_state", "blog_unit_testing", "blog_mock_llm",
        "blog_tools_first",
        "deepwiki_subgraphs", "deepwiki_hitl",
        "blog_send_api", "blog_handoffs_tds",
        "blog_multiagent_2026", "blog_scaling_lg", "dev_mongo_ltm",
    ])
    section("Community repositories", [
        "gh_awesome", "gh_fastapi_template",
    ])

    add_heading(doc, "Recency note", level=2)
    add_para(doc, "All cited sources were retrieved on 2026-04-28. None of the "
                  "URLs cited here are older than 2024; the official LangChain "
                  "blog post 'LangGraph v0.2 checkpointers' is from 2024 and is "
                  "the oldest source in this bibliography.")

    add_heading(doc, "Sources searched but no firm answer found", level=2)
    add_para(doc, "Klarna's published case study confirms LangGraph use and "
                  "checkpointing posture but does not disclose state-schema choice "
                  "or directory layout. The LangGraph 1.0 announcement asserts "
                  "stability but does not pick TypedDict vs Pydantic. No claim in "
                  "this document is sourced from those gaps.")
    for k in ("blog_klarna", "blog_v1"):
        add_citation(doc, k)

    # =========================================================
    # 6. Architect questions from V2 annotations  (NEW)
    # =========================================================
    add_heading(doc, "6. Architect Questions from V2 Annotations", level=1)
    add_para(doc, "V3 was triggered by four </Skip> annotations the architect placed "
                  "inline in V2. Each annotation is preserved verbatim below, then "
                  "answered with research-backed citations and a clear recommendation.",
                  italic=True)

    # ===== 6.1 - Annotation 1: granularity of one-graph-per-front-end-epic =====
    add_heading(doc, "6.1. Annotation on V2 P70 (Q9, OO principles): "
                     "graph-per-front-end-epic granularity", level=2)
    add_para(doc, "Architect's note (verbatim):", italic=True)
    add_quote(doc,
        "</Skip It is my concept that every front end epic will have its own graph "
        "and when a handoff from one graph to another takes place the necasary parts "
        "of state would flow from one graph to another please think through if this "
        "degree of granularity is sufficient. Here below is our current graph:>")

    add_heading(doc, "Recommendation", level=3)
    add_para(doc,
        "Confirm with refinement. One graph per front-end epic is a sound primary "
        "decomposition - it matches the LangGraph community's 'subgraph as "
        "independent state machine' guidance and is one of the two patterns the "
        "official handoff doc explicitly supports. The refinement: do NOT model "
        "epics as separately invoked top-level graphs by default. Compose them as "
        "subgraphs (each compiled once, embedded as a node in a thin parent graph) "
        "so that interrupts, checkpoints, and state propagation are handled by the "
        "single LangGraph runtime via the parent's BaseCheckpointSaver. Drop to "
        "separately invoked top-level graphs only when an epic must run on a "
        "different process / deployment / SLA.")

    add_heading(doc, "Why this is the right granularity", level=3)
    add_para(doc,
        "1) The official LangGraph handoff doc presents two patterns: 'single "
        "agent with middleware' and 'multiple agent subgraphs,' and notes the "
        "guidance 'Use single agent with middleware for most handoff use cases - "
        "it's simpler. Only use multiple agent subgraphs when you need bespoke "
        "agent implementations.' ChatHealthy's epics ARE bespoke agent "
        "implementations (Find Care, Talk About Care, Evaluate Care each have "
        "different tools, different state shapes, different LLM personae), so the "
        "multi-subgraph pattern is the appropriate one.")
    add_para(doc,
        "2) The DeepWiki survey of LangGraph composition documents that a "
        "CompiledStateGraph 'implements the PregelProtocol and can be passed to "
        "StateGraph.add_node()' - i.e., a compiled subgraph IS a node. State "
        "passes by 'selective projection': downward 'the parent passes only "
        "state keys that exist in the child's input_schema' and upward 'the "
        "subgraph returns values from its output_schema, which are written back "
        "using the parent's channel reducers.' This is exactly the architect's "
        "'necessary parts of state would flow from one graph to another' "
        "specified.")
    add_para(doc,
        "3) The 2026 BetterLink piece reinforces the same architecture: 'Split "
        "complex Agent into multiple Subgraphs, each Subgraph is independent "
        "state machine, can be tested, reused individually.'")

    add_heading(doc, "Implementation primitives for the handoff", level=3)
    add_bullet(doc, "input_schema / output_schema on each subgraph define the "
                    "wire interface between epics. Keys NOT in the schema do not "
                    "cross the boundary.")
    add_bullet(doc, "Command(graph=Command.PARENT) is the canonical primitive a "
                    "subgraph uses to direct the parent to the next epic. The "
                    "DeepWiki composition page documents: 'Subgraphs can direct "
                    "operations to parent graphs via Command(graph=Command.PARENT). "
                    "Internally, this raises a ParentCommand exception that the "
                    "parent execution loop catches and processes against its own "
                    "routing table - enabling sophisticated hierarchical control "
                    "without explicit parent awareness of the subgraph's internal "
                    "decisions.'")
    add_bullet(doc, "Send is the right primitive when one epic must fan out to "
                    "many parallel sub-tasks (e.g., research-many-trials at "
                    "once); see section 6.3.")
    add_bullet(doc, "Checkpointing: subgraph writes go through the same "
                    "BaseCheckpointSaver as the parent under a namespaced key "
                    "({node_name}:{task_id}), so a single checkpointer instance "
                    "covers all epics. This matters for the Mongo decision in "
                    "6.2.")

    add_heading(doc, "When the epic-as-graph pattern is wrong", level=3)
    add_bullet(doc, "Two epics that share most of their state and most of their "
                    "tools should be ONE subgraph with conditional routing, not "
                    "two subgraphs. The community guidance on multi-agent vs "
                    "single-agent is explicit: 'over 60% of Agent production "
                    "incidents relate to state management,' and gratuitous graph "
                    "boundaries multiply the surface area where keys can be "
                    "dropped or mis-projected.")
    add_bullet(doc, "An epic that is just a deterministic pipeline (no LLM "
                    "decisions) should be a tool, not a graph - the LangChain "
                    "team explicitly state in 'Building LangGraph' that they "
                    "chose 'little to no abstraction... we focused on control "
                    "and durability.' Spinning up a graph per pure-function "
                    "epic is over-engineering.")

    add_heading(doc, "Apply to the current POC", level=3)
    add_para(doc,
        "user_journey.py today is a single 336-line monolith. The architect's "
        "model maps cleanly: Find Care subgraph (specialty + provider + trial "
        "search), Talk About Care subgraph (the meeting + deep-research subgraph - "
        "see 6.3), Evaluate Care subgraph (the second terminal). The fan-out "
        "safety detectors (emergency_detector, ip_lock_check, repetitive_detector, "
        "classifier) belong in a thin per-utterance pre-processing subgraph - "
        "shared across all three epics rather than re-implemented in each.")

    add_heading(doc, "Sources", level=3)
    for k in ("docs_handoffs", "deepwiki_subgraphs", "blog_easton_2026",
              "blog_handoffs_tds", "blog_multiagent_2026", "blog_scaling_lg",
              "blog_building", "gh_swarm", "gh_supervisor"):
        add_citation(doc, k)

    # ===== 6.2 - Annotation 2: Postgres vs Mongo =====
    add_heading(doc, "6.2. Annotation on V2 P111 (AP-8): Postgres vs Mongo for the "
                     "checkpointer", level=2)
    add_para(doc, "Architect's note (verbatim):", italic=True)
    add_quote(doc, "</Skip Is there a compelling reason to use Postgress over Mongo?>")

    add_heading(doc, "Recommendation", level=3)
    add_para(doc,
        "Use the officially-supported MongoDBSaver. There is NO compelling "
        "reason to introduce Postgres just to host the LangGraph checkpointer "
        "when ChatHealthy already runs MongoDB at scale. The "
        "langgraph-checkpoint-mongodb package is co-maintained by MongoDB and "
        "LangChain, exposes the full BaseCheckpointSaver contract (including "
        "async via aput / aget / alist), and is documented by MongoDB Atlas as "
        "the canonical way to persist LangGraph state. Reserve a possible Postgres "
        "introduction for an explicit future need (multi-DB transactional join "
        "with another Postgres-only system) and document that as a separate "
        "decision.")

    add_heading(doc, "Evidence the Mongo checkpointer is production-grade and "
                     "officially supported", level=3)
    add_bullet(doc, "PyPI 'langgraph-checkpoint-mongodb' lists the maintainers as "
                    "MongoDB (10gen) and LangChain - making this an officially-"
                    "supported integration rather than a community project. "
                    "Latest release 0.3.1 on 2026-01-22, requires Python 3.10+, "
                    "Trusted Publishing verified.")
    add_bullet(doc, "MongoDB Atlas's official AI-integration docs document "
                    "MongoDBSaver as a first-party feature: it 'persists agent "
                    "state in MongoDB' and 'enables human-in-the-loop, memory, "
                    "time travel, and fault-tolerance for your LangGraph agents.' "
                    "No 'Preview' or 'Beta' labels.")
    add_bullet(doc, "MongoDB Atlas additionally provides MongoDBStore for "
                    "long-term memory (cross-thread, semantic-search "
                    "compatible). If we use Atlas Vector Search anyway for RAG, "
                    "checkpointer + store + vector store are 'consolidated... in "
                    "a single database, simplifying your architecture and "
                    "reducing operational complexity' - the official MongoDB "
                    "value proposition.")

    add_heading(doc, "Does MongoDBSaver actually meet the checkpointer "
                     "interface?", level=3)
    add_para(doc, "Yes. The LangGraph checkpoints reference defines the contract:")
    add_bullet(doc, "get_tuple(config) - retrieve a checkpoint by thread / id")
    add_bullet(doc, "list(config, *, filter, before, limit) - list checkpoints")
    add_bullet(doc, "put(config, checkpoint, metadata, new_versions) - persist a "
                    "checkpoint")
    add_bullet(doc, "put_writes(config, writes, task_id, task_path) - store "
                    "pending channel writes")
    add_bullet(doc, "delete_thread(thread_id) - remove all checkpoints for a "
                    "thread")
    add_bullet(doc, "get_next_version() - generate channel version IDs (default "
                    "integer sequence)")
    add_bullet(doc, "Async equivalents (aput / aput_writes / aget_tuple / alist / "
                    "adelete_thread) for async graph execution")
    add_para(doc, "The MongoDB Read-the-Docs MongoDBSaver page and the PyPI "
                  "package both confirm both sync and async methods are "
                  "implemented. The LangChain forum discussion of custom-"
                  "checkpointer construction lists exactly the same method set.")

    add_heading(doc, "Trade-offs to be aware of", level=3)
    add_bullet(doc, "Postgres has slightly more mature on-disk versioning per "
                    "the official LangGraph v0.2 announcement (per-channel "
                    "versioned writes, cursor reads). MongoDB's implementation "
                    "satisfies the same contract but the official LangChain "
                    "blog post does not call out Mongo-specific optimisations.")
    add_bullet(doc, "If a future need is multi-row ACID join across the "
                    "checkpoint and another relational table, Postgres is "
                    "stronger. ChatHealthy currently has no such requirement.")
    add_bullet(doc, "Don't build a custom MongoSaver - the official one already "
                    "exists. Building a parallel implementation would violate "
                    "the engineering rules' don't-reinvent-the-wheel posture and "
                    "would also have to track LangGraph's BaseCheckpointSaver "
                    "interface drift over time.")

    add_heading(doc, "Sources", level=3)
    for k in ("pypi_mongo_ckpt", "mongo_atlas_lg", "mongo_atlas_lg_agents",
              "mongo_lc_readthedocs", "mongo_blog", "dev_mongo_ltm",
              "ref_checkpoints", "docs_custom_ckpt", "docs_persistence",
              "blog_v02_ckpt"):
        add_citation(doc, k)

    # ===== 6.3 - Annotation 3: Talk About Care async deep research =====
    add_heading(doc, "6.3. Annotation on V2 P125 (R-3, open_deep_research): "
                     "async deep-research while a meeting continues", level=2)
    add_para(doc, "Architect's note (verbatim):", italic=True)
    add_quote(doc,
        "</Skip we have need for this, espeshially in 'Talk about care'  if this is "
        "an anti pattern I'm  a bit confused. Use case: User has stage 4 cancer and "
        "wants to know about cutting edge research where no FDA approval of on "
        "label drugs exists. User is in 'Talk about care' and the bots, people and "
        "principle (the person who is hosting the 'Talk about care meeting) decide "
        "together the only way to understand what cutting edge drugs are out there "
        "is deep research. LangGraph challenge in relation to this pattern is that "
        "for me the natural thing to do is establish the deep research node in the "
        "Talk About care sub graph and have an agent go off and do the research "
        "while the meeting participants continue to talk. Help me understand how "
        "we will model this use case. There is probably cause for deep research to "
        "be required by find care and evaluate care as well.>")

    add_heading(doc, "First, the confusion", level=3)
    add_para(doc,
        "open_deep_research is NOT an anti-pattern. V2 listed it under "
        "'Reference Implementations' (R-3) precisely BECAUSE it is the canonical "
        "LangGraph pattern for long-running research. The anti-patterns section "
        "in V2 (AP-1 .. AP-8) does not include deep research; the only nearby "
        "anti-pattern is 'Monolithic Graph' (AP-2), which open_deep_research "
        "explicitly avoids by decomposing into supervisor + worker subgraphs.")

    add_heading(doc, "Recommendation", level=3)
    add_para(doc,
        "Yes - put the deep-research subgraph inside the Talk About Care "
        "subgraph (and reuse it from Find Care and Evaluate Care). Run it "
        "concurrently with the meeting using LangGraph's Send API for "
        "fan-out, an async checkpointer (MongoDBSaver async API per 6.2) so "
        "the runtime can suspend and resume across multiple worker tasks, and "
        "an interrupt() in a 'wait for research' node so the parent meeting "
        "graph yields control to the host's UI while the supervisor completes. "
        "When research completes, resume with Command(resume=research_payload), "
        "which appears as a state update to the meeting subgraph and is "
        "rendered as a system message into the meeting transcript.")

    add_heading(doc, "How LangGraph models concurrent long-running work", level=3)
    add_bullet(doc, "Send API. The official graph-API doc and the community "
                    "Send-API write-ups document that 'in LangGraph, map-reduce "
                    "is implemented using the Send API, which enables dynamic "
                    "task creation at runtime where the number and configuration "
                    "of parallel tasks are determined by the graph's state.' For "
                    "the cancer-research use case, the supervisor emits one "
                    "Send per cutting-edge-trial-or-paper to research in "
                    "parallel.")
    add_bullet(doc, "Subgraph as a node, with its own checkpoint namespace. "
                    "DeepWiki on graph composition: 'Each subgraph invocation "
                    "receives its own namespace using the format "
                    "{node_name}:{task_id}... All checkpoints share the same "
                    "thread_id but differ by namespace, enabling independent "
                    "state history.' The deep-research subgraph thus suspends "
                    "and resumes independently of the meeting transcript.")
    add_bullet(doc, "interrupt() and Command(resume=...). The official "
                    "Interrupts doc: 'To use interrupt, you need... A "
                    "checkpointer to persist the graph state.' When the meeting "
                    "node calls interrupt() awaiting the research result, the "
                    "graph 'saves the current graph state' and 'waits "
                    "indefinitely until you resume execution.' Resume with "
                    "Command(resume=value), where 'the resume value becomes the "
                    "return value of the interrupt() call.'")
    add_bullet(doc, "Async-throughout. The official add-memory / persistence "
                    "docs note that for async graph execution the checkpointer "
                    "must implement aput / aput_writes / aget_tuple / alist - "
                    "MongoDBSaver does (per 6.2). Async graph execution is what "
                    "lets the supervisor await many Send-emitted research "
                    "workers concurrently without blocking the meeting "
                    "subgraph.")

    add_heading(doc, "Three options for the meeting-versus-research handshake", level=3)
    add_bullet(doc, "Option A - interrupt-and-wait (simplest, recommended). The "
                    "Talk About Care subgraph contains a 'await_research' node "
                    "that calls interrupt() with payload {'kind': "
                    "'research_in_flight', 'subject': '...'}. The host UI "
                    "renders a non-blocking banner ('research in progress') and "
                    "the meeting transcript continues via UI-side streaming "
                    "until the research subgraph completes. The application "
                    "layer then resumes the graph with the research payload, "
                    "which the meeting node turns into a system message. This "
                    "is the official human-in-the-loop pattern documented in "
                    "the Interrupts doc.")
    add_bullet(doc, "Option B - background thread + polling. The deep-research "
                    "subgraph runs on a separate thread_id (its own checkpoint "
                    "namespace) launched by the meeting node, and the meeting "
                    "polls a Mongo collection for completion. Works, but loses "
                    "the elegance of LangGraph's resume - the application layer "
                    "becomes responsible for marshaling. Use only if the "
                    "research must outlive the meeting session.")
    add_bullet(doc, "Option C - Send fan-out fully inside the meeting. The "
                    "supervisor in the meeting subgraph emits Send tasks and "
                    "awaits them inline; the meeting transcript pauses on a "
                    "spinner. Discouraged for the stage-4 cancer use case "
                    "because it freezes the conversation while research runs, "
                    "violating the architect's requirement that 'meeting "
                    "participants continue to talk.'")

    add_heading(doc, "Recommended structure for Talk About Care + deep research", level=3)
    add_bullet(doc, "Talk About Care subgraph state schema includes "
                    "research_requests: list[ResearchQuestion] (Pydantic), "
                    "research_results: dict[str, ResearchReport], and "
                    "pending_research: set[str].")
    add_bullet(doc, "deep_research subgraph (its own input_schema = "
                    "{ResearchQuestion}, output_schema = {ResearchReport}) is "
                    "imported as a node both by Talk About Care and by Find "
                    "Care and Evaluate Care - one implementation, three "
                    "consumers, satisfying the architect's note about reuse.")
    add_bullet(doc, "Inside deep_research the supervisor emits Send-fan-out "
                    "to per-source workers (clinicaltrials.gov, PubMed, "
                    "preprint servers) - each running concurrently under "
                    "asyncio.gather inside the supervisor as in "
                    "open_deep_research / deep_research_from_scratch.")
    add_bullet(doc, "The meeting subgraph resumes via Command(resume=...) and "
                    "renders the consolidated ResearchReport as a "
                    "ToolMessage / system utterance into the transcript.")

    add_heading(doc, "Sources", level=3)
    for k in ("docs_interrupts", "deepwiki_hitl", "deepwiki_subgraphs",
              "docs_handoffs", "blog_send_api", "blog_scaling_lg",
              "gh_open_deep_research", "gh_deepagents",
              "docs_add_memory", "blog_multiagent_2026"):
        add_citation(doc, k)

    # ===== 6.4 - Annotation 4: deep_research_from_scratch pattern =====
    add_heading(doc, "6.4. Annotation on V2 P137 (R-6, deep_research_from_scratch): "
                     "should we follow this model?", level=2)
    add_para(doc, "Architect's note (verbatim):", italic=True)
    add_quote(doc,
        "</ Skip see note with use case  above should we follow the model "
        "suggested by this pattern>")

    add_heading(doc, "Recommendation", level=3)
    add_para(doc,
        "Yes - follow the deep_research_from_scratch pattern. It is the "
        "right scaffold for the stage-4 cancer use case in 6.3 and for any "
        "future deep-research need across Find Care and Evaluate Care. "
        "Specifically: adopt the 5-stage decomposition (Scope -> Research "
        "Agent -> MCP-aware tool layer -> Supervisor -> Full Agent), and "
        "adopt its Pydantic-schema-for-routing-decisions discipline "
        "(ClarifyWithUser, ResearchQuestion, ConductResearch, "
        "ResearchComplete) as the wire types between meeting and research.")

    add_heading(doc, "What deep_research_from_scratch actually contains", level=3)
    add_bullet(doc, "Notebook 1 'scoping': defines ClarifyWithUser (Pydantic "
                    "structured output: 'Determines if additional context is "
                    "needed from the user') and ResearchQuestion ('Transforms "
                    "conversations into detailed research inquiries'). State "
                    "management, structured outputs, conditional routing.")
    add_bullet(doc, "Notebook 2 'research_agent': iterative agent with custom "
                    "search tools - this is the per-source worker.")
    add_bullet(doc, "Notebook 3 'research_agent_mcp': MCP integration so "
                    "external tool servers can be plugged in - directly "
                    "applicable if ChatHealthy wants to call out to "
                    "clinicaltrials.gov / PubMed via standardised MCP servers.")
    add_bullet(doc, "Notebook 4 'research_supervisor': defines ConductResearch "
                    "and ResearchComplete as supervisor delegation tools, "
                    "workers run in parallel via asyncio.gather, context "
                    "isolation per topic.")
    add_bullet(doc, "Notebook 5 'full_agent': the end-to-end composition - "
                    "subgraph composition + the full pipeline.")

    add_heading(doc, "Why this is the right model for ChatHealthy", level=3)
    add_bullet(doc, "Pydantic schemas as routing decisions is the exact pattern "
                    "the V1/V2 doc already recommends in Q1 (Pydantic at "
                    "boundaries) - deep_research_from_scratch is the "
                    "concrete example of how this scales beyond a single "
                    "boundary. Each schema is testable in isolation.")
    add_bullet(doc, "ClarifyWithUser maps directly onto the architect's "
                    "described meeting dynamic where 'the bots, people and "
                    "principle decide together' - that decision IS the "
                    "ClarifyWithUser branch. Bots emit a ClarifyWithUser to "
                    "the host; host approves; ResearchQuestion is then "
                    "constructed; ConductResearch dispatches to the workers.")
    add_bullet(doc, "The supervisor / worker split keeps the deep-research "
                    "subgraph reusable across Find Care, Talk About Care, and "
                    "Evaluate Care without any of those epics knowing about "
                    "specific search providers - exactly the input_schema / "
                    "output_schema discipline from 6.1.")
    add_bullet(doc, "MCP support (Notebook 3) gives a standard wire format "
                    "for adding new sources later - a stage-4 cancer user "
                    "may need preprint servers and trial registries we "
                    "haven't yet integrated.")

    add_heading(doc, "Caveat - not every notebook applies as-is", level=3)
    add_bullet(doc, "deep_research_from_scratch is a tutorial repo (708 "
                    "stars). Production hardening - retries, rate limits, "
                    "observability, JWT auth - is NOT in scope for those "
                    "notebooks. The fastapi-langgraph production template "
                    "(R-5) and open_deep_research (R-3) are the references "
                    "for that hardening; deep_research_from_scratch is the "
                    "reference for the schema discipline.")
    add_bullet(doc, "If we adopt MongoDBStore as our long-term memory layer "
                    "(per 6.2), the Notebook 5 'full_agent' wiring needs to "
                    "be adapted from a Postgres / in-memory store to "
                    "MongoDBStore - tractable, since MongoDBStore "
                    "implements the same Store interface.")

    add_heading(doc, "Sources", level=3)
    for k in ("gh_deep_research_scratch", "gh_open_deep_research",
              "blog_send_api", "blog_scaling_lg",
              "deepwiki_subgraphs", "deepwiki_hitl",
              "docs_interrupts", "docs_handoffs",
              "mongo_atlas_lg_agents"):
        add_citation(doc, k)

    # ---------------------------------------------------
    # Save
    # ---------------------------------------------------
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
